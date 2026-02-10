# -*- coding: utf-8 -*-
"""
โปรแกรมรวมไฟล์ Word รายงานออกแบบโครงสร้างชั้นทาง
Pavement Design Report Merger
Version 3.0 (Refactored)

โดย: ภาควิชาครุศาสตร์โยธา มจพ.

การปรับปรุงจาก v2.0:
- [ข้อ 1] ลด code ซ้ำซ้อน: ใช้ SECTION_CONFIG + render_upload_section()
- [ข้อ 3] ลดความซับซ้อนของ merge logic: ไม่สร้าง header_doc แยก ไม่ต้อง save/reload BytesIO ซ้อน
- [ข้อ 3] ยังคงใช้ docxcompose.Composer เพื่อรักษา formatting ของไฟล์ต้นฉบับ (รูปภาพ ตาราง styles)
- [ข้อ 4] เพิ่ม validate_docx_file() ตรวจสอบไฟล์ก่อน merge
- [ข้อ 5] เพิ่ม progress bar จริง แทน spinner
"""

import streamlit as st
import os
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docxcompose.composer import Composer
import io

# ═══════════════════════════════════════════════════════════════
# Configuration: โครงสร้างรายงาน (แก้ไขที่เดียวจบ)
# ═══════════════════════════════════════════════════════════════
SECTION_CONFIG = [
    {
        "group_icon": "📊",
        "group_title": "1. การคำนวณ Truck Factor",
        "items": [
            {
                "key": "truck_factor",
                "title": "การคำนวณ Truck Factor",
                "label": "**การคำนวณ Truck Factor** (ถ้ามี)",
                "uploader_label": "เลือกไฟล์ Truck Factor",
                "help": "ไฟล์รายงานการคำนวณ Truck Factor",
                "report_title": "การคำนวณ Truck Factor",
            }
        ],
    },
    {
        "group_icon": "📈",
        "group_title": "2. การคำนวณ ESALs (Equivalent Single Axle Loads)",
        "columns": 2,
        "items": [
            {
                "key": "esals_ac",
                "title": "2.1 ESALs สำหรับผิวทางลาดยาง",
                "label": "**2.1 ESALs สำหรับผิวทางลาดยาง** (Flexible Pavement)",
                "uploader_label": "เลือกไฟล์ ESALs ผิวทางลาดยาง",
                "help": "ไฟล์รายงานการคำนวณ ESALs สำหรับผิวทางลาดยาง (AC)",
                "report_title": "การคำนวณ ESALs สำหรับผิวทางลาดยาง (Flexible Pavement)",
            },
            {
                "key": "esals_concrete",
                "title": "2.2 ESALs สำหรับผิวทางคอนกรีต",
                "label": "**2.2 ESALs สำหรับผิวทางคอนกรีต** (Rigid Pavement)",
                "uploader_label": "เลือกไฟล์ ESALs ผิวทางคอนกรีต",
                "help": "ไฟล์รายงานการคำนวณ ESALs สำหรับผิวทางคอนกรีต",
                "report_title": "การคำนวณ ESALs สำหรับผิวทางคอนกรีต (Rigid Pavement)",
            },
        ],
    },
    {
        "group_icon": "🔬",
        "group_title": "3. การวิเคราะห์ค่า CBR ที่เปอร์เซ็นต์ไทล์",
        "items": [
            {
                "key": "cbr_analysis",
                "title": "การวิเคราะห์ค่า CBR ที่เปอร์เซ็นต์ไทล์",
                "label": "**การวิเคราะห์ค่า CBR ที่เปอร์เซ็นต์ไทล์**",
                "uploader_label": "เลือกไฟล์วิเคราะห์ CBR",
                "help": "ไฟล์รายงานการวิเคราะห์ค่า CBR ที่เปอร์เซ็นต์ไทล์ (Percentile Analysis)",
                "report_title": "การวิเคราะห์ค่า CBR ที่เปอร์เซ็นต์ไทล์",
            }
        ],
    },
    {
        "group_icon": "🛤️",
        "group_title": "4. การออกแบบผิวทางลาดยาง (Flexible Pavement)",
        "items": [
            {
                "key": "ac_design",
                "title": "การออกแบบผิวทางลาดยาง",
                "label": "**การออกแบบผิวทางลาดยาง (AC)**",
                "uploader_label": "เลือกไฟล์ออกแบบ AC",
                "help": "ไฟล์รายงานการออกแบบผิวทางแอสฟัลต์ตามวิธี AASHTO 1993",
                "report_title": "การออกแบบผิวทางลาดยาง (Flexible Pavement)",
            }
        ],
    },
    {
        "group_icon": "🏗️",
        "group_title": "5. การออกแบบผิวทางคอนกรีต (Rigid Pavement)",
        "columns": 2,
        "items": [
            {
                "key": "jpcp_jrcp_design",
                "title": "5.1 การออกแบบ JPCP/JRCP",
                "label": "**5.1 การออกแบบ JPCP/JRCP**",
                "caption": "Jointed Plain/Reinforced Concrete Pavement",
                "uploader_label": "เลือกไฟล์ออกแบบ JPCP/JRCP",
                "help": "ไฟล์รายงานการออกแบบผิวทาง JPCP หรือ JRCP",
                "report_title": "การออกแบบผิวทางคอนกรีต JPCP/JRCP",
            },
            {
                "key": "crcp_design",
                "title": "5.2 การออกแบบ CRCP",
                "label": "**5.2 การออกแบบ CRCP**",
                "caption": "Continuously Reinforced Concrete Pavement",
                "uploader_label": "เลือกไฟล์ออกแบบ CRCP",
                "help": "ไฟล์รายงานการออกแบบผิวทาง CRCP",
                "report_title": "การออกแบบผิวทางคอนกรีต CRCP",
            },
        ],
    },
    {
        "group_icon": "📐",
        "group_title": "6. การคำนวณ Corrected Modulus of Subgrade Reaction (k-value)",
        "columns": 2,
        "items": [
            {
                "key": "k_value_jpcp_jrcp",
                "title": "6.1 k-value สำหรับ JPCP/JRCP",
                "label": "**6.1 k-value สำหรับ JPCP/JRCP**",
                "uploader_label": "เลือกไฟล์ k-value JPCP/JRCP",
                "help": "ไฟล์รายการคำนวณ Corrected k-value สำหรับ JPCP/JRCP",
                "report_title": "การคำนวณ Corrected Modulus of Subgrade Reaction (k-value) สำหรับ JPCP/JRCP",
            },
            {
                "key": "k_value_crcp",
                "title": "6.2 k-value สำหรับ CRCP",
                "label": "**6.2 k-value สำหรับ CRCP**",
                "uploader_label": "เลือกไฟล์ k-value CRCP",
                "help": "ไฟล์รายการคำนวณ Corrected k-value สำหรับ CRCP",
                "report_title": "การคำนวณ Corrected Modulus of Subgrade Reaction (k-value) สำหรับ CRCP",
            },
        ],
    },
    {
        "group_icon": "💰",
        "group_title": "7. การประมาณราคาค่าก่อสร้าง",
        "items": [
            {
                "key": "cost_estimate",
                "title": "การประมาณราคาค่าก่อสร้าง",
                "label": "**การประมาณราคาค่าก่อสร้าง** (ถ้ามี)",
                "uploader_label": "เลือกไฟล์ประมาณราคา",
                "help": "ไฟล์รายงานการประมาณราคาค่าก่อสร้าง",
                "report_title": "การประมาณราคาค่าก่อสร้าง",
            }
        ],
    },
]


# ═══════════════════════════════════════════════════════════════
# ตั้งค่าหน้าเว็บ
# ═══════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="โปรแกรมรวมรายงานออกแบบโครงสร้างชั้นทาง",
    page_icon="🛣️",
    layout="wide"
)

# CSS สำหรับตกแต่งหน้าเว็บ
st.markdown("""
<style>
    .main-header {
        font-size: 28px;
        font-weight: bold;
        color: #1E3A5F;
        text-align: center;
        padding: 20px;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border-radius: 10px;
        margin-bottom: 20px;
    }
    .sub-header {
        font-size: 18px;
        color: #4A5568;
        text-align: center;
        margin-bottom: 30px;
    }
    .file-section {
        background-color: #F7FAFC;
        padding: 15px;
        border-radius: 10px;
        margin-bottom: 10px;
        border-left: 4px solid #667eea;
    }
    .section-header {
        background-color: #C6F6D5;
        padding: 10px 15px;
        border-radius: 8px;
        margin: 15px 0 10px 0;
        font-weight: bold;
        color: #276749;
        border-left: 4px solid #38A169;
    }
    .success-box {
        background-color: #C6F6D5;
        padding: 15px;
        border-radius: 10px;
        border-left: 4px solid #38A169;
    }
    .warning-box {
        background-color: #FEFCBF;
        padding: 15px;
        border-radius: 10px;
        border-left: 4px solid #D69E2E;
    }
    .stButton>button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        font-weight: bold;
        padding: 10px 30px;
        border-radius: 25px;
        border: none;
        font-size: 16px;
    }
    .stButton>button:hover {
        background: linear-gradient(135deg, #764ba2 0%, #667eea 100%);
    }
</style>
""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════
# Utility Functions
# ═══════════════════════════════════════════════════════════════

def set_thai_font(run, font_name="TH Sarabun New", font_size=15):
    """ตั้งค่าฟอนต์ไทยและขนาด"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)


def set_page_margins(section):
    """ตั้งค่าหน้ากระดาษ A4 แนวตั้ง กั้นหน้า-หลัง 2.5 cm"""
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.orientation = WD_ORIENT.PORTRAIT
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)


def validate_docx_file(file):
    """
    [ข้อ 4] ตรวจสอบว่าไฟล์เป็น .docx ที่ valid หรือไม่
    Returns: (is_valid: bool, error_message: str)
    """
    try:
        file_bytes = file.read()
        file.seek(0)
        doc = Document(io.BytesIO(file_bytes))
        # ตรวจสอบว่ามีเนื้อหาอย่างน้อย 1 paragraph
        if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
            return False, "ไฟล์ว่างเปล่า ไม่มีเนื้อหา"
        return True, ""
    except Exception as e:
        return False, f"ไฟล์เสียหายหรือไม่ใช่ไฟล์ .docx ที่ถูกต้อง ({str(e)})"


def get_all_items():
    """ดึงรายการ item ทั้งหมดจาก SECTION_CONFIG ตามลำดับ"""
    items = []
    for group in SECTION_CONFIG:
        for item in group["items"]:
            items.append(item)
    return items


# ═══════════════════════════════════════════════════════════════
# [ข้อ 3] Refactored Merge Logic
# ═══════════════════════════════════════════════════════════════

def create_cover_and_toc(uploaded_files, project_name, report_date):
    """
    สร้างเอกสาร master ที่มีหน้าปก + สารบัญ
    [ข้อ 3] แยกเป็นฟังก์ชันชัดเจน ไม่ซ้อน BytesIO
    """
    doc = Document()
    section = doc.sections[0]
    set_page_margins(section)

    # ─── หน้าปก ───
    spacer = doc.add_paragraph()
    spacer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacer.add_run("\n\n\n\n\n")

    main_title = doc.add_paragraph()
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = main_title.add_run("รายงานการออกแบบโครงสร้างชั้นทาง")
    set_thai_font(run, font_size=24)
    run.font.bold = True

    if project_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"\n{project_name}")
        set_thai_font(run, font_size=20)
        run.font.bold = True

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(f"\n\n\n\n{report_date}")
    set_thai_font(run, font_size=16)

    doc.add_page_break()

    # ─── สารบัญ ───
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run("สารบัญ")
    set_thai_font(run, font_size=18)
    run.font.bold = True

    doc.add_paragraph()

    # สร้างรายการสารบัญอัตโนมัติจากไฟล์ที่อัปโหลด
    all_items = get_all_items()
    section_num = 1
    for item in all_items:
        if uploaded_files.get(item["key"]) is not None:
            toc_para = doc.add_paragraph()
            run = toc_para.add_run(f"{section_num}. {item['report_title']}")
            set_thai_font(run, font_size=15)
            section_num += 1

    doc.add_page_break()

    return doc


def create_section_header_doc(section_num, title):
    """
    [ข้อ 3] สร้างเอกสารหัวข้อแบบ minimal
    ยังคงต้องสร้างเป็น Document แยกเพราะ Composer.append() ต้องการ Document object
    แต่ลดขั้นตอนซ้ำซ้อนลง (ไม่ต้อง save → BytesIO → reload)
    """
    header_doc = Document()
    p = header_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"{section_num}. {title}")
    set_thai_font(run, font_size=18)
    run.font.bold = True
    header_doc.add_paragraph()  # ระยะห่างก่อนเนื้อหา
    return header_doc


def merge_documents(uploaded_files, project_name, report_date, progress_callback=None):
    """
    [ข้อ 3] Refactored merge logic
    
    การเปลี่ยนแปลงหลักจาก v2.0:
    1. ไม่ต้อง save merged_doc → BytesIO → reload เป็น master_doc อีกต่อไป
       → สร้าง cover+toc document แล้วส่งเข้า Composer โดยตรง
    2. ไม่ต้อง save header_doc → BytesIO → reload
       → ส่ง header_doc เข้า Composer.append() โดยตรง (docxcompose รองรับ)
    3. ฟังก์ชัน copy_table() ที่ไม่ได้ใช้งานถูกลบออก
       (Composer จัดการ copy tables + images + formatting ให้อัตโนมัติ)
    """
    # สร้างเอกสาร master (ปก + สารบัญ)
    master_doc = create_cover_and_toc(uploaded_files, project_name, report_date)
    composer = Composer(master_doc)

    # นับจำนวนไฟล์ที่ต้อง merge สำหรับ progress bar
    all_items = get_all_items()
    active_items = [(item, uploaded_files[item["key"]]) 
                    for item in all_items 
                    if uploaded_files.get(item["key"]) is not None]
    total = len(active_items)

    # รวมเนื้อหาจากแต่ละไฟล์
    for idx, (item, file) in enumerate(active_items):
        section_num = idx + 1

        # เพิ่มหัวข้อ section
        header_doc = create_section_header_doc(section_num, item["report_title"])
        composer.append(header_doc)

        # เพิ่มเนื้อหาจากไฟล์ต้นฉบับ
        file_bytes = file.read()
        file.seek(0)
        source_doc = Document(io.BytesIO(file_bytes))
        composer.append(source_doc)

        # [ข้อ 5] อัปเดต progress bar
        if progress_callback:
            progress_callback((idx + 1) / total, f"กำลังรวม: {item['report_title']}")

    return composer.doc


# ═══════════════════════════════════════════════════════════════
# [ข้อ 1] UI Rendering Functions (ลด code ซ้ำซ้อน)
# ═══════════════════════════════════════════════════════════════

def render_single_uploader(item):
    """แสดง file uploader สำหรับ 1 รายการ"""
    st.markdown('<div class="file-section">', unsafe_allow_html=True)
    st.markdown(item["label"])
    if "caption" in item:
        st.caption(item["caption"])
    uploaded = st.file_uploader(
        item["uploader_label"],
        type=['docx'],
        key=item["key"],
        help=item["help"],
    )
    st.markdown('</div>', unsafe_allow_html=True)
    return uploaded


def render_upload_sections():
    """
    [ข้อ 1] แสดง upload sections ทั้งหมดจาก SECTION_CONFIG
    ลด code ซ้ำซ้อนจาก ~150 บรรทัด เหลือ loop เดียว
    """
    uploaded_files = {}

    for group in SECTION_CONFIG:
        # แสดงหัวข้อ group
        st.markdown(
            f'<div class="section-header">{group["group_icon"]} {group["group_title"]}</div>',
            unsafe_allow_html=True
        )

        n_cols = group.get("columns", 1)
        items = group["items"]

        if n_cols > 1 and len(items) > 1:
            # แสดงแบบหลายคอลัมน์
            cols = st.columns(n_cols)
            for i, item in enumerate(items):
                with cols[i % n_cols]:
                    uploaded_files[item["key"]] = render_single_uploader(item)
        else:
            # แสดงแบบคอลัมน์เดียว
            for item in items:
                uploaded_files[item["key"]] = render_single_uploader(item)

    return uploaded_files


def render_file_status(uploaded_files):
    """แสดงสถานะไฟล์ที่อัปโหลด"""
    st.markdown("### 📊 สถานะไฟล์ที่อัปโหลด")

    all_items = get_all_items()
    file_count = sum(1 for item in all_items if uploaded_files.get(item["key"]) is not None)

    # แสดงในรูปแบบ 3 คอลัมน์
    cols = st.columns(3)
    for i, item in enumerate(all_items):
        with cols[i % 3]:
            if uploaded_files.get(item["key"]) is not None:
                st.success(f"{item['title']}: ✅ อัปโหลดแล้ว")
            else:
                st.warning(f"{item['title']}: ⬜ ยังไม่อัปโหลด")

    st.markdown(f"### 📈 อัปโหลดแล้ว: **{file_count}** จาก **{len(all_items)}** ไฟล์")
    return file_count


# ═══════════════════════════════════════════════════════════════
# Main Application
# ═══════════════════════════════════════════════════════════════

def main():
    # หัวข้อหลัก
    st.markdown('<div class="main-header">🛣️ โปรแกรมรวมรายงานออกแบบโครงสร้างชั้นทาง</div>', unsafe_allow_html=True)
    st.markdown('<div class="sub-header">Pavement Structure Design Report Merger v3.0</div>', unsafe_allow_html=True)

    # ข้อมูลโครงการ
    st.markdown("### 📋 ข้อมูลโครงการ")
    col1, col2 = st.columns(2)
    with col1:
        project_name = st.text_input("ชื่อโครงการ", placeholder="กรอกชื่อโครงการ")
    with col2:
        report_date = st.date_input("วันที่รายงาน", datetime.now())
        report_date_str = report_date.strftime("%d/%m/%Y")

    st.markdown("---")

    # อัปโหลดไฟล์
    st.markdown("### 📁 อัปโหลดไฟล์รายงาน")
    st.info("💡 อัปโหลดไฟล์ Word (.docx) สำหรับแต่ละส่วนของรายงาน ไฟล์ที่มีเครื่องหมาย (ถ้ามี) สามารถเว้นว่างได้")

    uploaded_files = render_upload_sections()

    st.markdown("---")

    # แสดงสถานะ
    file_count = render_file_status(uploaded_files)

    st.markdown("---")

    # ปุ่มรวมไฟล์
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        merge_button = st.button("🔄 รวมไฟล์และสร้างรายงาน", use_container_width=True)

    if merge_button:
        if file_count == 0:
            st.error("❌ กรุณาอัปโหลดไฟล์อย่างน้อย 1 ไฟล์")
        else:
            # [ข้อ 4] ตรวจสอบไฟล์ก่อน merge
            validation_errors = []
            all_items = get_all_items()
            for item in all_items:
                file = uploaded_files.get(item["key"])
                if file is not None:
                    is_valid, error_msg = validate_docx_file(file)
                    if not is_valid:
                        validation_errors.append(f"❌ **{item['title']}**: {error_msg}")

            if validation_errors:
                st.error("พบไฟล์ที่มีปัญหา กรุณาตรวจสอบและอัปโหลดใหม่:")
                for err in validation_errors:
                    st.markdown(err)
            else:
                # [ข้อ 5] Progress bar จริง
                progress_bar = st.progress(0, text="เริ่มต้นรวมไฟล์...")

                def update_progress(fraction, text):
                    progress_bar.progress(fraction, text=text)

                try:
                    merged_doc = merge_documents(
                        uploaded_files,
                        project_name,
                        report_date_str,
                        progress_callback=update_progress
                    )

                    progress_bar.progress(1.0, text="✅ รวมไฟล์เรียบร้อยแล้ว!")

                    with tempfile.TemporaryDirectory() as temp_dir:
                        base_filename = "รายงานออกแบบโครงสร้างชั้นทาง"
                        if project_name:
                            base_filename = f"รายงานออกแบบ_{project_name.replace(' ', '_')}"

                        docx_path = os.path.join(temp_dir, f"{base_filename}.docx")
                        merged_doc.save(docx_path)

                        st.markdown('<div class="success-box">', unsafe_allow_html=True)
                        st.success(f"✅ รวมไฟล์เรียบร้อยแล้ว! ({file_count} ไฟล์)")
                        st.markdown('</div>', unsafe_allow_html=True)

                        st.markdown("### 📥 ดาวน์โหลดรายงาน")

                        with open(docx_path, 'rb') as f:
                            docx_data = f.read()
                        st.download_button(
                            label="📄 ดาวน์โหลดไฟล์ Word (.docx)",
                            data=docx_data,
                            file_name=f"{base_filename}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )

                except Exception as e:
                    st.error(f"❌ เกิดข้อผิดพลาด: {str(e)}")
                    st.exception(e)

    # Footer
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #718096; font-size: 14px;">
        <p>พัฒนาโดย รศ.ดร.อิทธิพล มีผล // ภาควิชาครุศาสตร์โยธา คณะครุศาสตร์อุตสาหกรรม </p>
        <p>มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ</p>
        <p>© 2025 - Pavement Design Report Merger v3.0</p>
    </div>
    """, unsafe_allow_html=True)


if __name__ == "__main__":
    main()
