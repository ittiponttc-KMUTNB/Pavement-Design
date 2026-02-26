"""
================================================================================
AASHTO 1993 Flexible Pavement Design - Streamlit Application (Version 6)
================================================================================
‡πÅ‡∏≠‡∏õ‡∏û‡∏•‡∏¥‡πÄ‡∏Ñ‡∏ä‡∏±‡∏ô‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö Flexible Pavement ‡∏ï‡∏≤‡∏°‡∏ß‡∏¥‡∏ò‡∏µ AASHTO 1993
‡∏õ‡∏£‡∏±‡∏ö‡∏õ‡∏£‡∏∏‡∏á‡∏ï‡∏≤‡∏°‡∏°‡∏≤‡∏ï‡∏£‡∏ê‡∏≤‡∏ô‡∏Å‡∏£‡∏°‡∏ó‡∏≤‡∏á‡∏´‡∏•‡∏ß‡∏á (DOH Thailand)

[V6 Improvements ‚Äî UX/UI Overhaul + ‡∏ä‡∏∑‡πà‡∏≠‡∏ß‡∏±‡∏™‡∏î‡∏∏‡πÉ‡∏´‡∏°‡πà]
- Quick Summary Banner (PASS/FAIL) ‡∏ñ‡∏≤‡∏ß‡∏£‡∏≠‡∏¢‡∏π‡πà‡πÄ‡∏´‡∏ô‡∏∑‡∏≠ Tabs
- ‡∏õ‡∏∏‡πà‡∏° "‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì" ‡∏ä‡∏±‡∏î‡πÄ‡∏à‡∏ô ‡∏≠‡∏¢‡∏π‡πà‡πÉ‡∏ï‡πâ Tab 2 (‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á)
- AC Sublayer UI ‡πÄ‡∏£‡∏µ‡∏¢‡∏ö‡∏á‡πà‡∏≤‡∏¢‡∏Ç‡∏∂‡πâ‡∏ô (‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡πÄ‡∏î‡∏µ‡∏¢‡∏ß radio + number_input)
- Sidebar ‡∏•‡∏î‡∏Ç‡∏ô‡∏≤‡∏î ‚Äî ‡∏¢‡πâ‡∏≤‡∏¢‡∏ê‡∏≤‡∏ô‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏ß‡∏±‡∏™‡∏î‡∏∏‡πÑ‡∏õ‡∏≠‡∏¢‡∏π‡πà‡πÉ‡∏ô Tab 2
- Export buttons ‡∏à‡∏±‡∏î‡πÉ‡∏´‡∏°‡πà 2 ‡πÅ‡∏ñ‡∏ß + use_container_width
- Refactor _short_name ‡πÄ‡∏õ‡πá‡∏ô function ‡πÄ‡∏î‡∏µ‡∏¢‡∏ß
- Sensitivity Analysis ‡πÉ‡∏ô Tab 3 ‡∏û‡∏£‡πâ‡∏≠‡∏° label ‡∏ä‡∏±‡∏î‡πÄ‡∏à‡∏ô
- ‡∏ä‡∏∑‡πà‡∏≠‡∏ß‡∏±‡∏™‡∏î‡∏∏‡πÉ‡∏´‡∏°‡πà‡∏ï‡∏≤‡∏°‡∏°‡∏≤‡∏ï‡∏£‡∏ê‡∏≤‡∏ô ‡∏ó‡∏•.:
    CTB ‚Üí ‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏´‡∏¥‡∏ô‡∏Ñ‡∏•‡∏∏‡∏Å‡∏õ‡∏£‡∏±‡∏ö‡∏õ‡∏£‡∏∏‡∏á‡∏Ñ‡∏∏‡∏ì‡∏†‡∏≤‡∏û‡∏î‡πâ‡∏ß‡∏¢‡∏õ‡∏π‡∏ô‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå (Cement Treated Base)
    Wearing ‚Üí ‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡πÅ‡∏≠‡∏™‡∏ü‡∏±‡∏•‡∏ï‡πå‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï (AC. Wearing Course)
    Binder ‚Üí ‡∏£‡∏≠‡∏á‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡πÅ‡∏≠‡∏™‡∏ü‡∏±‡∏•‡∏ï‡πå‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï (AC. Binder Course)
    Base Course ‚Üí ‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡πÅ‡∏≠‡∏™‡∏ü‡∏±‡∏•‡∏ï‡πå‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï (AC. Base Course)

Author: ‡∏£‡∏®.‡∏î‡∏£.‡∏≠‡∏¥‡∏ó‡∏ò‡∏¥‡∏û‡∏• ‡∏°‡∏µ‡∏ú‡∏• // ‡∏†‡∏≤‡∏Ñ‡∏ß‡∏¥‡∏ä‡∏≤‡∏Ñ‡∏£‡∏∏‡∏®‡∏≤‡∏™‡∏ï‡∏£‡πå‡πÇ‡∏¢‡∏ò‡∏≤ // ‡∏°‡∏à‡∏û.
Version: 6.0
================================================================================
"""

import streamlit as st
import numpy as np
import json
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import matplotlib.font_manager as fm
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ================================================================================
# CUSTOM ROOT-FINDING (‡πÅ‡∏ó‡∏ô scipy.optimize.brentq)
# ================================================================================

def brentq(f, a, b, xtol=1e-12, maxiter=200):
    fa, fb = f(a), f(b)
    if fa * fb > 0:
        raise ValueError(f"f(a) and f(b) must have different signs")
    if abs(fa) < xtol:
        return a
    if abs(fb) < xtol:
        return b
    c, fc = a, fa
    d = e = b - a
    for _ in range(maxiter):
        if fb * fc > 0:
            c, fc = a, fa
            d = e = b - a
        if abs(fc) < abs(fb):
            a, b, c = b, c, b
            fa, fb, fc = fb, fc, fb
        tol1 = 2.0 * 2.2e-16 * abs(b) + 0.5 * xtol
        m = 0.5 * (c - b)
        if abs(m) <= tol1 or fb == 0.0:
            return b
        if abs(e) >= tol1 and abs(fa) > abs(fb):
            s = fb / fa
            if a == c:
                p = 2.0 * m * s
                q = 1.0 - s
            else:
                q = fa / fc
                r = fb / fc
                p = s * (2.0 * m * q * (q - r) - (b - a) * (r - 1.0))
                q = (q - 1.0) * (r - 1.0) * (s - 1.0)
            if p > 0:
                q = -q
            else:
                p = -p
            if 2.0 * p < min(3.0 * m * q - abs(tol1 * q), abs(e * q)):
                e = d
                d = p / q
            else:
                d = m
                e = m
        else:
            d = m
            e = m
        a, fa = b, fb
        if abs(d) > tol1:
            b += d
        else:
            b += tol1 if m > 0 else -tol1
        fb = f(b)
    return b


# ================================================================================
# PAGE CONFIGURATION
# ================================================================================

st.set_page_config(
    page_title="Flexible Pavement Design (AASHTO 1993) v6",
    page_icon="üõ£Ô∏è",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ================================================================================
# MATERIAL DATABASE - ‡∏ï‡∏≤‡∏°‡∏°‡∏≤‡∏ï‡∏£‡∏ê‡∏≤‡∏ô‡∏Å‡∏£‡∏°‡∏ó‡∏≤‡∏á‡∏´‡∏•‡∏ß‡∏á (DOH Thailand)
# ================================================================================

MATERIALS = {
    # ============ ‡∏ä‡∏±‡πâ‡∏ô‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á (Surface) ============
    "‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡∏•‡∏≤‡∏î‡∏¢‡∏≤‡∏á AC": {
        "layer_coeff": 0.40, "drainage_coeff": 1.0,
        "mr_psi": 362500, "mr_mpa": 2500,
        "layer_type": "surface", "color": "#1C1C1C",
        "short_name": "AC", "english_name": "Asphalt Concrete"
    },
    "‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡∏•‡∏≤‡∏î‡∏¢‡∏≤‡∏á PMA": {
        "layer_coeff": 0.40, "drainage_coeff": 1.0,
        "mr_psi": 536500, "mr_mpa": 3700,
        "layer_type": "surface", "color": "#2C2C2C",
        "short_name": "PMA", "english_name": "Polymer Modified Asphalt"
    },

    # ============ ‡∏ä‡∏±‡πâ‡∏ô‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á (Base) ============
    "‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏´‡∏¥‡∏ô‡∏Ñ‡∏•‡∏∏‡∏Å‡∏õ‡∏£‡∏±‡∏ö‡∏õ‡∏£‡∏∏‡∏á‡∏Ñ‡∏∏‡∏ì‡∏†‡∏≤‡∏û‡∏î‡πâ‡∏ß‡∏¢‡∏õ‡∏π‡∏ô‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå (Cement Treated Base)": {
        "layer_coeff": 0.18, "drainage_coeff": 1.0,
        "mr_psi": 174000, "mr_mpa": 1200,
        "layer_type": "base", "color": "#78909C",
        "short_name": "CTB", "english_name": "Cement Treated Base"
    },
    "‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏´‡∏¥‡∏ô‡∏Ñ‡∏•‡∏∏‡∏Å‡∏ú‡∏™‡∏°‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå UCS 24.5 ksc.": {
        "layer_coeff": 0.15, "drainage_coeff": 1.0,
        "mr_psi": 123250, "mr_mpa": 850,
        "layer_type": "base", "color": "#607D8B",
        "short_name": "MOD.CRB", "english_name": "Mod.Crushed Rock Base"
    },
    "‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏´‡∏¥‡∏ô‡∏Ñ‡∏•‡∏∏‡∏Å CBR 80%": {
        "layer_coeff": 0.13, "drainage_coeff": 1.0,
        "mr_psi": 50750, "mr_mpa": 350,
        "layer_type": "base", "color": "#795548",
        "short_name": "CAB", "english_name": "Crushed Rock Base"
    },
    "‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏î‡∏¥‡∏ô‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå UCS 17.5 ksc.": {
        "layer_coeff": 0.13, "drainage_coeff": 1.0,
        "mr_psi": 50750, "mr_mpa": 350,
        "layer_type": "base", "color": "#8D6E63",
        "short_name": "SCB", "english_name": "Soil Cement Base"
    },
    "‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏´‡∏°‡∏∏‡∏ô‡πÄ‡∏ß‡∏µ‡∏¢‡∏ô (Recycling)": {
        "layer_coeff": 0.15, "drainage_coeff": 1.0,
        "mr_psi": 123250, "mr_mpa": 850,
        "layer_type": "base", "color": "#5D4037",
        "short_name": "RAP", "english_name": "Recycled Asphalt Pavement"
    },

    # ============ ‡∏ä‡∏±‡πâ‡∏ô‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á (Subbase) ============
    "‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏°‡∏ß‡∏•‡∏£‡∏ß‡∏° CBR 25%": {
        "layer_coeff": 0.10, "drainage_coeff": 1.0,
        "mr_psi": 21750, "mr_mpa": 150,
        "layer_type": "subbase", "color": "#FFB74D",
        "short_name": "GSB", "english_name": "Aggregate Subbase"
    },

    # ============ ‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Ñ‡∏±‡∏î‡πÄ‡∏•‡∏∑‡∏≠‡∏Å (Selected Material) ============
    "‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Ñ‡∏±‡∏î‡πÄ‡∏•‡∏∑‡∏≠‡∏Å ‡∏Å": {
        "layer_coeff": 0.08, "drainage_coeff": 1.0,
        "mr_psi": 14504, "mr_mpa": 100,
        "layer_type": "selected", "color": "#FFF176",
        "short_name": "SM-A", "english_name": "Selected Material"
    },

    # ============ ‡πÑ‡∏°‡πà‡πÉ‡∏ä‡πâ‡∏ß‡∏±‡∏™‡∏î‡∏∏ (Skip layer) ============
    "‡πÑ‡∏°‡πà‡πÉ‡∏ä‡πâ‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Ñ‡∏±‡∏î‡πÄ‡∏•‡∏∑‡∏≠‡∏Å (‡πÉ‡∏ä‡πâ‡∏î‡∏¥‡∏ô‡∏ó‡∏≤‡∏á‡∏ó‡∏£‡∏û)": {
        "layer_coeff": 0.00, "drainage_coeff": 1.0,
        "mr_psi": 0, "mr_mpa": 0,
        "layer_type": "none", "color": "#D7CCC8",
        "short_name": "NONE", "english_name": "None"
    }
}

# ================================================================================
# PRESET STRUCTURES - ‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏°‡∏≤‡∏ï‡∏£‡∏ê‡∏≤‡∏ô ‡∏ó‡∏•.
# ================================================================================

PRESET_STRUCTURES = {
    "--- ‡πÄ‡∏•‡∏∑‡∏≠‡∏Å‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏°‡∏≤‡∏ï‡∏£‡∏ê‡∏≤‡∏ô ---": None,
    "AC + CTB + GSB + SM (‡∏°‡∏≤‡∏ï‡∏£‡∏ê‡∏≤‡∏ô‡∏´‡∏•‡∏±‡∏Å)": {
        "description": "‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á AC / ‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á CTB / ‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á GSB / ‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Ñ‡∏±‡∏î‡πÄ‡∏•‡∏∑‡∏≠‡∏Å",
        "num_layers": 4,
        "layers": [
            {"material": "‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡∏•‡∏≤‡∏î‡∏¢‡∏≤‡∏á AC", "thickness_cm": 15.0},
            {"material": "‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏´‡∏¥‡∏ô‡∏Ñ‡∏•‡∏∏‡∏Å‡∏õ‡∏£‡∏±‡∏ö‡∏õ‡∏£‡∏∏‡∏á‡∏Ñ‡∏∏‡∏ì‡∏†‡∏≤‡∏û‡∏î‡πâ‡∏ß‡∏¢‡∏õ‡∏π‡∏ô‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå (Cement Treated Base)", "thickness_cm": 15.0},
            {"material": "‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏°‡∏ß‡∏•‡∏£‡∏ß‡∏° CBR 25%", "thickness_cm": 15.0},
            {"material": "‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Ñ‡∏±‡∏î‡πÄ‡∏•‡∏∑‡∏≠‡∏Å ‡∏Å", "thickness_cm": 30.0},
        ]
    },
    "AC + MOD.CRB + GSB + SM": {
        "description": "‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á AC / ‡∏´‡∏¥‡∏ô‡∏Ñ‡∏•‡∏∏‡∏Å‡∏ú‡∏™‡∏°‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå / ‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á GSB / ‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Ñ‡∏±‡∏î‡πÄ‡∏•‡∏∑‡∏≠‡∏Å",
        "num_layers": 4,
        "layers": [
            {"material": "‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡∏•‡∏≤‡∏î‡∏¢‡∏≤‡∏á AC", "thickness_cm": 15.0},
            {"material": "‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏´‡∏¥‡∏ô‡∏Ñ‡∏•‡∏∏‡∏Å‡∏ú‡∏™‡∏°‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå UCS 24.5 ksc.", "thickness_cm": 20.0},
            {"material": "‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏°‡∏ß‡∏•‡∏£‡∏ß‡∏° CBR 25%", "thickness_cm": 15.0},
            {"material": "‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Ñ‡∏±‡∏î‡πÄ‡∏•‡∏∑‡∏≠‡∏Å ‡∏Å", "thickness_cm": 30.0},
        ]
    },
    "AC + CAB + GSB + SM": {
        "description": "‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á AC / ‡∏´‡∏¥‡∏ô‡∏Ñ‡∏•‡∏∏‡∏Å CBR 80% / ‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á GSB / ‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Ñ‡∏±‡∏î‡πÄ‡∏•‡∏∑‡∏≠‡∏Å",
        "num_layers": 4,
        "layers": [
            {"material": "‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡∏•‡∏≤‡∏î‡∏¢‡∏≤‡∏á AC", "thickness_cm": 15.0},
            {"material": "‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏´‡∏¥‡∏ô‡∏Ñ‡∏•‡∏∏‡∏Å CBR 80%", "thickness_cm": 20.0},
            {"material": "‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏°‡∏ß‡∏•‡∏£‡∏ß‡∏° CBR 25%", "thickness_cm": 15.0},
            {"material": "‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Ñ‡∏±‡∏î‡πÄ‡∏•‡∏∑‡∏≠‡∏Å ‡∏Å", "thickness_cm": 30.0},
        ]
    },
    "AC + SCB + GSB + SM": {
        "description": "‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á AC / ‡∏î‡∏¥‡∏ô‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå / ‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á GSB / ‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Ñ‡∏±‡∏î‡πÄ‡∏•‡∏∑‡∏≠‡∏Å",
        "num_layers": 4,
        "layers": [
            {"material": "‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡∏•‡∏≤‡∏î‡∏¢‡∏≤‡∏á AC", "thickness_cm": 15.0},
            {"material": "‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏î‡∏¥‡∏ô‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå UCS 17.5 ksc.", "thickness_cm": 20.0},
            {"material": "‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏°‡∏ß‡∏•‡∏£‡∏ß‡∏° CBR 25%", "thickness_cm": 15.0},
            {"material": "‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Ñ‡∏±‡∏î‡πÄ‡∏•‡∏∑‡∏≠‡∏Å ‡∏Å", "thickness_cm": 30.0},
        ]
    },
    "AC + CTB + GSB (‡πÑ‡∏°‡πà‡πÉ‡∏ä‡πâ SM)": {
        "description": "‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á AC / ‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á CTB / ‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á GSB (‡πÑ‡∏°‡πà‡∏°‡∏µ‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Ñ‡∏±‡∏î‡πÄ‡∏•‡∏∑‡∏≠‡∏Å)",
        "num_layers": 3,
        "layers": [
            {"material": "‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡∏•‡∏≤‡∏î‡∏¢‡∏≤‡∏á AC", "thickness_cm": 15.0},
            {"material": "‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏´‡∏¥‡∏ô‡∏Ñ‡∏•‡∏∏‡∏Å‡∏õ‡∏£‡∏±‡∏ö‡∏õ‡∏£‡∏∏‡∏á‡∏Ñ‡∏∏‡∏ì‡∏†‡∏≤‡∏û‡∏î‡πâ‡∏ß‡∏¢‡∏õ‡∏π‡∏ô‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå (Cement Treated Base)", "thickness_cm": 20.0},
            {"material": "‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏°‡∏ß‡∏•‡∏£‡∏ß‡∏° CBR 25%", "thickness_cm": 20.0},
        ]
    },
    "PMA + CTB + GSB + SM": {
        "description": "‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á PMA / ‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á CTB / ‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á GSB / ‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Ñ‡∏±‡∏î‡πÄ‡∏•‡∏∑‡∏≠‡∏Å",
        "num_layers": 4,
        "layers": [
            {"material": "‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡∏•‡∏≤‡∏î‡∏¢‡∏≤‡∏á PMA", "thickness_cm": 15.0},
            {"material": "‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏´‡∏¥‡∏ô‡∏Ñ‡∏•‡∏∏‡∏Å‡∏õ‡∏£‡∏±‡∏ö‡∏õ‡∏£‡∏∏‡∏á‡∏Ñ‡∏∏‡∏ì‡∏†‡∏≤‡∏û‡∏î‡πâ‡∏ß‡∏¢‡∏õ‡∏π‡∏ô‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå (Cement Treated Base)", "thickness_cm": 15.0},
            {"material": "‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏°‡∏ß‡∏•‡∏£‡∏ß‡∏° CBR 25%", "thickness_cm": 15.0},
            {"material": "‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Ñ‡∏±‡∏î‡πÄ‡∏•‡∏∑‡∏≠‡∏Å ‡∏Å", "thickness_cm": 30.0},
        ]
    },
    "AC + RAP + GSB + SM": {
        "description": "‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á AC / ‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á Recycling / ‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á GSB / ‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Ñ‡∏±‡∏î‡πÄ‡∏•‡∏∑‡∏≠‡∏Å",
        "num_layers": 4,
        "layers": [
            {"material": "‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡∏•‡∏≤‡∏î‡∏¢‡∏≤‡∏á AC", "thickness_cm": 15.0},
            {"material": "‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏´‡∏°‡∏∏‡∏ô‡πÄ‡∏ß‡∏µ‡∏¢‡∏ô (Recycling)", "thickness_cm": 20.0},
            {"material": "‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏°‡∏ß‡∏•‡∏£‡∏ß‡∏° CBR 25%", "thickness_cm": 15.0},
            {"material": "‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Ñ‡∏±‡∏î‡πÄ‡∏•‡∏∑‡∏≠‡∏Å ‡∏Å", "thickness_cm": 30.0},
        ]
    },
}

# ================================================================================
# RELIABILITY TABLE
# ================================================================================

RELIABILITY_ZR = {
    50: -0.000, 60: -0.253, 70: -0.524, 75: -0.674, 80: -0.841,
    85: -1.037, 90: -1.282, 91: -1.340, 92: -1.405, 93: -1.476,
    94: -1.555, 95: -1.645, 96: -1.751, 97: -1.881, 98: -2.054,
    99: -2.327, 99.9: -3.090
}

# ================================================================================
# DRAINAGE COEFFICIENT TABLE
# ================================================================================

DRAINAGE_TABLE = {
    "Excellent": {"description": "‡∏£‡∏∞‡∏ö‡∏≤‡∏¢‡∏ô‡πâ‡∏≥‡∏î‡∏µ‡πÄ‡∏¢‡∏µ‡πà‡∏¢‡∏° (< 2 ‡∏ä‡∏°.)",
                  "values": {"<1%": 1.40, "1-5%": 1.35, "5-25%": 1.30, ">25%": 1.20}},
    "Good":      {"description": "‡∏£‡∏∞‡∏ö‡∏≤‡∏¢‡∏ô‡πâ‡∏≥‡∏î‡∏µ (1 ‡∏ß‡∏±‡∏ô)",
                  "values": {"<1%": 1.35, "1-5%": 1.25, "5-25%": 1.15, ">25%": 1.00}},
    "Fair":      {"description": "‡∏£‡∏∞‡∏ö‡∏≤‡∏¢‡∏ô‡πâ‡∏≥‡∏û‡∏≠‡πÉ‡∏ä‡πâ (1 ‡∏™‡∏±‡∏õ‡∏î‡∏≤‡∏´‡πå)",
                  "values": {"<1%": 1.25, "1-5%": 1.15, "5-25%": 1.05, ">25%": 0.80}},
    "Poor":      {"description": "‡∏£‡∏∞‡∏ö‡∏≤‡∏¢‡∏ô‡πâ‡∏≥‡πÑ‡∏°‡πà‡∏î‡∏µ (1 ‡πÄ‡∏î‡∏∑‡∏≠‡∏ô)",
                  "values": {"<1%": 1.15, "1-5%": 1.05, "5-25%": 0.80, ">25%": 0.60}},
    "Very Poor": {"description": "‡∏£‡∏∞‡∏ö‡∏≤‡∏¢‡∏ô‡πâ‡∏≥‡πÑ‡∏°‡πà‡∏î‡∏µ‡∏°‡∏≤‡∏Å (‡πÑ‡∏°‡πà‡∏£‡∏∞‡∏ö‡∏≤‡∏¢)",
                  "values": {"<1%": 1.05, "1-5%": 0.80, "5-25%": 0.60, ">25%": 0.40}},
}

# DOH AC Sublayer Thickness Standards (mm)
DOH_THICKNESS_STANDARDS = {
    "‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡πÅ‡∏≠‡∏™‡∏ü‡∏±‡∏•‡∏ï‡πå‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï (AC. Wearing Course)":    [40, 45, 50, 55, 60, 65, 70],
    "‡∏£‡∏≠‡∏á‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡πÅ‡∏≠‡∏™‡∏ü‡∏±‡∏•‡∏ï‡πå‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï (AC. Binder Course)":  [40, 45, 50, 55, 60, 65, 70, 75, 80],
    "‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡πÅ‡∏≠‡∏™‡∏ü‡∏±‡∏•‡∏ï‡πå‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï (AC. Base Course)":      [0, 70, 75, 80, 85, 90, 95, 100]
}

# ================================================================================
# HELPER: ‡∏ä‡∏∑‡πà‡∏≠‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏¢‡πà‡∏≠ (Refactored ‚Äî ‡πÄ‡∏î‡∏µ‡∏¢‡∏ß‡∏Å‡∏±‡∏ô‡∏ó‡∏±‡πâ‡∏á‡∏£‡∏∞‡∏ö‡∏ö)
# ================================================================================

def short_material_name(mat_name: str) -> str:
    """‡πÅ‡∏õ‡∏•‡∏á‡∏ä‡∏∑‡πà‡∏≠‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏¢‡∏≤‡∏ß‡πÄ‡∏õ‡πá‡∏ô‡∏ä‡∏∑‡πà‡∏≠‡∏¢‡πà‡∏≠‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡πÅ‡∏™‡∏î‡∏á‡πÉ‡∏ô‡∏ï‡∏≤‡∏£‡∏≤‡∏á/‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô"""
    mapping = {
        "‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏´‡∏¥‡∏ô‡∏Ñ‡∏•‡∏∏‡∏Å‡∏õ‡∏£‡∏±‡∏ö‡∏õ‡∏£‡∏∏‡∏á‡∏Ñ‡∏∏‡∏ì‡∏†‡∏≤‡∏û‡∏î‡πâ‡∏ß‡∏¢‡∏õ‡∏π‡∏ô‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå (Cement Treated Base)": "‡∏´‡∏¥‡∏ô‡∏Ñ‡∏•‡∏∏‡∏Å‡∏õ‡∏£‡∏±‡∏ö‡∏õ‡∏£‡∏∏‡∏á‡∏Ñ‡∏∏‡∏ì‡∏†‡∏≤‡∏û‡∏î‡πâ‡∏ß‡∏¢‡∏õ‡∏π‡∏ô‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå (CTB)",
        "‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏´‡∏¥‡∏ô‡∏Ñ‡∏•‡∏∏‡∏Å‡∏ú‡∏™‡∏°‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå UCS 24.5 ksc.":   "‡∏´‡∏¥‡∏ô‡∏Ñ‡∏•‡∏∏‡∏Å‡∏ú‡∏™‡∏°‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå UCS ‚â• 24.5 ksc",
        "‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏´‡∏¥‡∏ô‡∏Ñ‡∏•‡∏∏‡∏Å CBR 80%":                    "‡∏´‡∏¥‡∏ô‡∏Ñ‡∏•‡∏∏‡∏Å CBR ‚â• 80%",
        "‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏î‡∏¥‡∏ô‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå UCS 17.5 ksc.":          "‡∏î‡∏¥‡∏ô‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå UCS ‚â• 17.5 ksc",
        "‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏´‡∏°‡∏∏‡∏ô‡πÄ‡∏ß‡∏µ‡∏¢‡∏ô (Recycling)":         "‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏´‡∏°‡∏∏‡∏ô‡πÄ‡∏ß‡∏µ‡∏¢‡∏ô (Recycling)",
        "‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏°‡∏ß‡∏•‡∏£‡∏ß‡∏° CBR 25%":             "‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏°‡∏ß‡∏•‡∏£‡∏ß‡∏° CBR ‚â• 25%",
        "‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡∏•‡∏≤‡∏î‡∏¢‡∏≤‡∏á AC":                           "‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡∏•‡∏≤‡∏î‡∏¢‡∏≤‡∏á AC",
        "‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡∏•‡∏≤‡∏î‡∏¢‡∏≤‡∏á PMA":                          "‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡∏•‡∏≤‡∏î‡∏¢‡∏≤‡∏á PMA",
        "‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Ñ‡∏±‡∏î‡πÄ‡∏•‡∏∑‡∏≠‡∏Å ‡∏Å":                           "‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Ñ‡∏±‡∏î‡πÄ‡∏•‡∏∑‡∏≠‡∏Å ‡∏Å",
    }
    return mapping.get(mat_name, mat_name)


# ================================================================================
# CORE CALCULATION FUNCTIONS
# ================================================================================

def aashto_1993_equation(SN, W18, Zr, So, delta_psi, Mr):
    log_W18 = np.log10(W18)
    term1 = Zr * So
    term2 = 9.36 * np.log10(SN + 1) - 0.20
    numerator = np.log10(delta_psi / (4.2 - 1.5))
    denominator = 0.4 + (1094 / ((SN + 1) ** 5.19))
    term3 = numerator / denominator
    term4 = 2.32 * np.log10(Mr) - 8.07
    right_side = term1 + term2 + term3 + term4
    return right_side - log_W18


def calculate_sn_for_layer(W18, Zr, So, delta_psi, Mr):
    def f(SN):
        return aashto_1993_equation(SN, W18, Zr, So, delta_psi, Mr)
    try:
        return round(brentq(f, 0.01, 25.0, xtol=1e-6, maxiter=100), 2)
    except ValueError:
        return None


def calculate_w18_supported(SN, Zr, So, delta_psi, Mr):
    term1 = Zr * So
    term2 = 9.36 * np.log10(SN + 1) - 0.20
    numerator = np.log10(delta_psi / (4.2 - 1.5))
    denominator = 0.4 + (1094 / ((SN + 1) ** 5.19))
    term3 = numerator / denominator
    term4 = 2.32 * np.log10(Mr) - 8.07
    log_W18 = term1 + term2 + term3 + term4
    return 10 ** log_W18


def calculate_layer_thicknesses(W18, Zr, So, delta_psi, subgrade_mr, layers, ac_sublayers=None):
    results = {
        'layers': [], 'sn_values': [], 'subgrade_mr': subgrade_mr,
        'total_sn_required': None, 'total_sn_provided': 0,
        'ac_sublayers': ac_sublayers, 'warnings': []
    }

    active_layers = [l for l in layers if l['material'] != "‡πÑ‡∏°‡πà‡πÉ‡∏ä‡πâ‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Ñ‡∏±‡∏î‡πÄ‡∏•‡∏∑‡∏≠‡∏Å (‡πÉ‡∏ä‡πâ‡∏î‡∏¥‡∏ô‡∏ó‡∏≤‡∏á‡∏ó‡∏£‡∏û)"]
    if not active_layers:
        results['warnings'].append("‚ö†Ô∏è ‡πÑ‡∏°‡πà‡∏°‡∏µ‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏ó‡∏µ‡πà active")
        return results

    num_layers = len(active_layers)
    sn_values = []

    for i in range(num_layers - 1):
        mr_current = MATERIALS[active_layers[i]['material']]['mr_psi']
        mr_next = MATERIALS[active_layers[i + 1]['material']]['mr_psi']
        if mr_current < mr_next:
            results['warnings'].append(
                f"‚ö†Ô∏è ‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏µ‡πà {i+1} ‡∏°‡∏µ‡∏Ñ‡πà‡∏≤ Mr = {mr_current:,} psi ‡∏ï‡πà‡∏≥‡∏Å‡∏ß‡πà‡∏≤‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏µ‡πà {i+2} ‡∏ó‡∏µ‡πà‡∏°‡∏µ Mr = {mr_next:,} psi "
                f"‚Äî ‡∏õ‡∏Å‡∏ï‡∏¥‡∏ä‡∏±‡πâ‡∏ô‡∏ö‡∏ô‡∏Ñ‡∏ß‡∏£‡∏°‡∏µ‡∏Ñ‡πà‡∏≤ Mr ‡∏™‡∏π‡∏á‡∏Å‡∏ß‡πà‡∏≤‡∏ä‡∏±‡πâ‡∏ô‡∏•‡πà‡∏≤‡∏á"
            )

    for i in range(num_layers):
        if i == num_layers - 1:
            mr_below = subgrade_mr
        else:
            mat_below = MATERIALS[active_layers[i + 1]['material']]
            mr_below = mat_below['mr_psi']
        sn_i = calculate_sn_for_layer(W18, Zr, So, delta_psi, mr_below)
        if sn_i is None:
            results['warnings'].append(
                f"‚ö†Ô∏è ‡πÑ‡∏°‡πà‡∏™‡∏≤‡∏°‡∏≤‡∏£‡∏ñ‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì SN ‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏µ‡πà {i+1} ‡πÑ‡∏î‡πâ ‚Äî ‡∏Ñ‡πà‡∏≤ W18 ‡∏≠‡∏≤‡∏à‡∏™‡∏π‡∏á‡πÄ‡∏Å‡∏¥‡∏ô‡πÑ‡∏õ"
            )
        sn_values.append({'layer_index': i + 1, 'mr_below': mr_below, 'sn_required': sn_i})

    results['sn_values'] = sn_values
    results['total_sn_required'] = calculate_sn_for_layer(W18, Zr, So, delta_psi, subgrade_mr)

    if results['total_sn_required'] is None:
        results['warnings'].append("‚ö†Ô∏è ‡πÑ‡∏°‡πà‡∏™‡∏≤‡∏°‡∏≤‡∏£‡∏ñ‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì SN_required ‡πÑ‡∏î‡πâ ‚Äî ‡∏•‡∏≠‡∏á‡∏õ‡∏£‡∏±‡∏ö W18, Reliability ‡∏´‡∏£‡∏∑‡∏≠ CBR")

    cumulative_sn = 0
    for i, layer in enumerate(active_layers):
        mat = MATERIALS[layer['material']]
        a_i = layer.get('layer_coeff', mat['layer_coeff'])
        m_i = layer.get('drainage_coeff', 1.0)
        sn_required_at_layer = sn_values[i]['sn_required'] if sn_values[i]['sn_required'] else 0

        if a_i > 0 and m_i > 0:
            remaining_sn = max(0, sn_required_at_layer - cumulative_sn)
            min_thickness_inch = remaining_sn / (a_i * m_i)
            min_thickness_cm = min_thickness_inch * 2.54
        else:
            min_thickness_inch = 0
            min_thickness_cm = 0

        design_thickness_cm = layer['thickness_cm']
        design_thickness_inch = design_thickness_cm / 2.54
        sn_contribution = a_i * design_thickness_inch * m_i
        cumulative_sn += sn_contribution
        is_ok = design_thickness_cm >= min_thickness_cm

        layer_ac_sublayers = ac_sublayers if i == 0 and ac_sublayers is not None else None

        results['layers'].append({
            'layer_no': i + 1,
            'material': layer['material'],
            'short_name': mat['short_name'],
            'english_name': mat.get('english_name', mat['short_name']),
            'mr_psi': mat['mr_psi'], 'mr_mpa': mat['mr_mpa'],
            'a_i': a_i, 'm_i': m_i,
            'sn_required_at_layer': sn_required_at_layer,
            'min_thickness_inch': round(min_thickness_inch, 2),
            'min_thickness_cm': round(min_thickness_cm, 1),
            'design_thickness_cm': design_thickness_cm,
            'design_thickness_inch': round(design_thickness_inch, 2),
            'sn_contribution': round(sn_contribution, 4),
            'cumulative_sn': round(cumulative_sn, 2),
            'is_ok': is_ok, 'color': mat['color'],
            'ac_sublayers': layer_ac_sublayers
        })

    results['total_sn_provided'] = round(cumulative_sn, 2)
    return results


def check_design(sn_required, sn_provided):
    if sn_required is None:
        return {'status': 'ERROR', 'passed': False,
                'message': '‡πÑ‡∏°‡πà‡∏™‡∏≤‡∏°‡∏≤‡∏£‡∏ñ‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì SN_required ‡πÑ‡∏î‡πâ', 'safety_margin': 0}
    safety_margin = sn_provided - sn_required
    passed = sn_provided >= sn_required
    return {
        'status': 'OK' if passed else 'NG',
        'passed': passed,
        'safety_margin': round(safety_margin, 2),
        'message': f"SN_provided ({sn_provided:.2f}) {'‚â•' if passed else '<'} SN_required ({sn_required:.2f})"
    }


# ================================================================================
# SENSITIVITY ANALYSIS
# ================================================================================

def plot_sensitivity_cbr(W18, Zr, So, delta_psi, current_cbr):
    cbr_range = np.linspace(2, 20, 50)
    sn_values = []
    for cbr in cbr_range:
        mr = 1500 * cbr
        sn = calculate_sn_for_layer(W18, Zr, So, delta_psi, mr)
        sn_values.append(sn if sn else np.nan)

    fig, ax = plt.subplots(figsize=(7, 4))
    ax.plot(cbr_range, sn_values, 'b-', linewidth=2.5, label='SN required')

    current_mr  = 1500 * current_cbr
    current_sn  = calculate_sn_for_layer(W18, Zr, So, delta_psi, current_mr)
    if current_sn:
        ax.plot(current_cbr, current_sn, 'ro', markersize=12,
                label=f'Current: CBR={current_cbr:.1f}%, SN={current_sn:.2f}')
        # annotate
        ax.annotate(
            f'CBR={current_cbr:.1f}%\nSN={current_sn:.2f}',
            xy=(current_cbr, current_sn),
            xytext=(current_cbr + 1.5, current_sn + 0.3),
            fontsize=9,
            arrowprops=dict(arrowstyle='->', color='red', lw=1.2),
            color='red'
        )

    ax.set_xlabel('CBR (%)', fontsize=12)
    ax.set_ylabel('SN Required', fontsize=12)
    ax.set_title(
        'Sensitivity: SN Required vs CBR\n'
        '(Effect of Subgrade CBR on Required SN)',
        fontsize=11, fontweight='bold'
    )
    ax.legend(fontsize=10)
    ax.grid(True, alpha=0.3)
    ax.set_xlim(left=2)
    try:
        plt.tight_layout()
    except Exception:
        pass
    return fig


def plot_sensitivity_w18(Zr, So, delta_psi, Mr, current_w18):
    w18_range = np.logspace(5, 8.5, 50)
    sn_values = []
    for w18 in w18_range:
        sn = calculate_sn_for_layer(w18, Zr, So, delta_psi, Mr)
        sn_values.append(sn if sn else np.nan)

    fig, ax = plt.subplots(figsize=(7, 4))
    ax.semilogx(w18_range, sn_values, 'g-', linewidth=2.5, label='SN required')

    current_sn = calculate_sn_for_layer(current_w18, Zr, So, delta_psi, Mr)
    if current_sn:
        ax.semilogx(current_w18, current_sn, 'ro', markersize=12,
                    label=f'Current: W18={current_w18/1e6:.2f}M, SN={current_sn:.2f}')
        ax.annotate(
            f'W18={current_w18/1e6:.2f}M\nSN={current_sn:.2f}',
            xy=(current_w18, current_sn),
            xytext=(current_w18 * 0.15, current_sn + 0.4),
            fontsize=9,
            arrowprops=dict(arrowstyle='->', color='red', lw=1.2),
            color='red'
        )

    ax.set_xlabel('W18 (ESALs)', fontsize=12)
    ax.set_ylabel('SN Required', fontsize=12)
    ax.set_title(
        'Sensitivity: SN Required vs W18\n'
        '(Effect of Cumulative Traffic Load on Required SN)',
        fontsize=11, fontweight='bold'
    )
    ax.legend(fontsize=10)
    ax.grid(True, alpha=0.3)
    try:
        plt.tight_layout()
    except Exception:
        pass
    return fig


# ================================================================================
# VISUALIZATION FUNCTIONS
# ================================================================================

def _get_thai_fonts():
    import os
    sys_candidates = [
        ('/usr/share/fonts/truetype/tlwg/Garuda.ttf', '/usr/share/fonts/truetype/tlwg/Garuda-Bold.ttf'),
        ('/usr/share/fonts/opentype/tlwg/Garuda.otf', '/usr/share/fonts/opentype/tlwg/Garuda-Bold.otf'),
        ('/usr/share/fonts/truetype/tlwg/Loma.ttf', '/usr/share/fonts/truetype/tlwg/Loma-Bold.ttf'),
        ('/usr/share/fonts/opentype/tlwg/Loma.otf', '/usr/share/fonts/opentype/tlwg/Loma-Bold.otf'),
        ('/usr/share/fonts/truetype/noto/NotoSansThai-Regular.ttf', '/usr/share/fonts/truetype/noto/NotoSansThai-Bold.ttf'),
    ]
    for reg, bold in sys_candidates:
        if os.path.exists(reg):
            fp_r = fm.FontProperties(fname=reg)
            fp_b = fm.FontProperties(fname=bold) if os.path.exists(bold) else fm.FontProperties(fname=reg)
            return fp_r, fp_b, True
    return (fm.FontProperties(family='DejaVu Sans'),
            fm.FontProperties(family='DejaVu Sans', weight='bold'), False)


@st.cache_resource
def get_cached_thai_fonts():
    return _get_thai_fonts()


def plot_pavement_section(layers_result, subgrade_mr=None, subgrade_cbr=None, lang='en'):
    plt.rcParams['font.family'] = 'DejaVu Sans'
    thai_font = thai_font_bold = None
    has_thai = False
    if lang == 'th':
        thai_font, thai_font_bold, has_thai = get_cached_thai_fonts()
        if not has_thai:
            lang = 'en'

    def _fp(bold=False):
        if has_thai:
            return {'fontproperties': thai_font_bold if bold else thai_font}
        return {}

    if not layers_result:
        fig, ax = plt.subplots(figsize=(12, 8))
        ax.text(0.5, 0.5, 'No layers defined', ha='center', va='center', fontsize=14)
        ax.axis('off')
        return fig

    valid_layers = [l for l in layers_result if l.get('design_thickness_cm', 0) > 0]
    if not valid_layers:
        fig, ax = plt.subplots(figsize=(12, 8))
        ax.text(0.5, 0.5, 'No valid layers', ha='center', va='center', fontsize=14)
        ax.axis('off')
        return fig

    # Expand AC sublayers ‚Äî ‡πÉ‡∏ä‡πâ‡∏ä‡∏∑‡πà‡∏≠‡πÉ‡∏´‡∏°‡πà
    expanded_layers = []
    for layer in valid_layers:
        ac_sub = layer.get('ac_sublayers', None)
        if ac_sub is not None and layer['layer_no'] == 1:
            sub_info = [
                ('wearing', '#1C1C1C',
                 '‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡πÅ‡∏≠‡∏™‡∏ü‡∏±‡∏•‡∏ï‡πå‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï (AC. Wearing Course)',
                 'AC. Wearing Course'),
                ('binder',  '#333333',
                 '‡∏£‡∏≠‡∏á‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡πÅ‡∏≠‡∏™‡∏ü‡∏±‡∏•‡∏ï‡πå‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï (AC. Binder Course)',
                 'AC. Binder Course'),
                ('base',    '#4A4A4A',
                 '‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡πÅ‡∏≠‡∏™‡∏ü‡∏±‡∏•‡∏ï‡πå‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï (AC. Base Course)',
                 'AC. Base Course'),
            ]
            for key, color, th_name, en_name in sub_info:
                if ac_sub[key] > 0:
                    expanded_layers.append({
                        'design_thickness_cm': ac_sub[key],
                        'material': th_name if lang == 'th' else en_name,
                        'english_name': en_name, 'short_name': key[:2].upper() + 'C',
                        'color': color, 'mr_mpa': layer['mr_mpa'], 'is_sublayer': True
                    })
        else:
            expanded_layers.append(layer)

    draw_layers = expanded_layers
    total_thickness = sum(l['design_thickness_cm'] for l in draw_layers)

    fig, ax = plt.subplots(figsize=(12, 9))
    width = 3
    x_center = 7
    x_start = x_center - width / 2
    min_display_height = 6
    display_heights = [max(l['design_thickness_cm'], min_display_height) for l in draw_layers]
    total_display = sum(display_heights)
    dark_colors = ['#1C1C1C', '#2C2C2C', '#333333', '#4A4A4A', '#78909C', '#607D8B',
                   '#795548', '#8D6E63', '#5D4037', '#6D4C41', '#455A64']

    y_current = total_display
    for i, layer in enumerate(draw_layers):
        thickness = layer['design_thickness_cm']
        display_h = display_heights[i]
        color = layer.get('color', '#CCCCCC')
        e_mpa = layer.get('mr_mpa', 0)
        is_sublayer = layer.get('is_sublayer', False)

        if lang == 'th':
            name = layer.get('material', layer.get('short_name', f'Layer {i+1}'))
        else:
            name = layer.get('english_name', layer.get('short_name', f'Layer {i+1}'))

        y_bottom = y_current - display_h
        ls = '--' if is_sublayer else '-'
        lw = 1 if is_sublayer else 2
        rect = mpatches.Rectangle((x_start, y_bottom), width, display_h,
                                  linewidth=lw, linestyle=ls, edgecolor='black', facecolor=color)
        ax.add_patch(rect)
        yc = y_bottom + display_h / 2
        tc = 'white' if color in dark_colors else 'black'
        fs_center = 14 if is_sublayer else 16
        ax.text(x_center, yc, f'{thickness:.0f} cm',
                ha='center', va='center', fontsize=fs_center, fontweight='bold', color=tc)
        fs_name = 12 if is_sublayer else 14
        ax.text(x_start - 0.5, yc, name,
                ha='right', va='center', fontsize=fs_name, fontweight='bold', color='black', **_fp(True))
        if e_mpa and e_mpa > 0 and not is_sublayer:
            ax.text(x_start + width + 0.5, yc, f'E = {e_mpa:,.0f} MPa',
                    ha='left', va='center', fontsize=12, color='#0066CC')
        y_current = y_bottom

    # Subgrade
    sg_h = 6
    sg_yb = -sg_h
    ax.add_patch(mpatches.Rectangle((x_start, sg_yb), width, sg_h,
        linewidth=2, edgecolor='black', facecolor='#D7CCC8', hatch='///'))
    text_box_h = 3.5
    text_box_w = width * 0.85
    ax.add_patch(mpatches.FancyBboxPatch(
        (x_center - text_box_w / 2, sg_yb + (sg_h - text_box_h) / 2),
        text_box_w, text_box_h, boxstyle="round,pad=0.2",
        facecolor='#EFEBE9', edgecolor='#8D6E63', linewidth=1.5, alpha=0.95))

    sg_label = '‡∏î‡∏¥‡∏ô‡πÄ‡∏î‡∏¥‡∏° (Subgrade)' if lang == 'th' else 'Subgrade'
    if subgrade_cbr:
        sg_label += f'\nCBR = {subgrade_cbr:.1f}%'
    ax.text(x_center, sg_yb + sg_h / 2, sg_label,
            ha='center', va='center', fontsize=12, fontweight='bold', color='#5D4037', **_fp(True))
    if subgrade_mr:
        ax.text(x_start + width + 0.5, sg_yb + sg_h / 2, f'Mr = {subgrade_mr:,} psi',
                ha='left', va='center', fontsize=12, color='#0066CC')

    # Total thickness arrow
    ax.annotate('', xy=(x_start + width + 3.5, total_display),
                xytext=(x_start + width + 3.5, 0),
                arrowprops=dict(arrowstyle='<->', color='red', lw=2))
    total_label = f'‡∏£‡∏ß‡∏°\n{total_thickness:.0f} cm' if lang == 'th' else f'Total\n{total_thickness:.0f} cm'
    ax.text(x_start + width + 4, total_display / 2, total_label,
            ha='left', va='center', fontsize=14, color='red', fontweight='bold', **_fp(True))

    margin = 10
    ax.set_xlim(0, 15)
    ax.set_ylim(-sg_h - 4, total_display + margin)
    ax.axis('off')

    title_text = '‡∏£‡∏π‡∏õ‡∏ï‡∏±‡∏î‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á' if lang == 'th' else 'Pavement Structure'
    ax.set_title(title_text, fontsize=20, fontweight='bold', pad=20, **_fp(True))

    box_text = (f'‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡∏£‡∏ß‡∏°‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á: {total_thickness:.0f} cm'
                if lang == 'th' else f'Total Pavement Thickness: {total_thickness:.0f} cm')
    ax.text(x_center, -sg_h - 2, box_text,
            ha='center', va='center', fontsize=15, fontweight='bold', **_fp(True),
            bbox=dict(boxstyle='round', facecolor='lightyellow', alpha=0.9, edgecolor='orange'))
    try:
        plt.tight_layout()
    except Exception:
        pass
    return fig


def get_figure_as_bytes(fig):
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    buf.seek(0)
    return buf


# ================================================================================
# WORD EXPORT FUNCTIONS
# ================================================================================

def set_thai_font(run, size_pt=15, bold=False):
    run.font.name = 'TH SarabunPSK'
    run.font.size = Pt(size_pt)
    run.bold = bold
    try:
        run._element.rPr.rFonts.set(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cs', 'TH SarabunPSK')
    except Exception:
        pass


def set_equation_font(run, size_pt=11, bold=False, italic=True):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic


def add_thai_paragraph(doc, text, size_pt=15, bold=False, alignment=None):
    para = doc.add_paragraph()
    if alignment:
        para.alignment = alignment
    run = para.add_run(text)
    set_thai_font(run, size_pt, bold)
    return para


def add_equation_paragraph(doc, text, size_pt=11, bold=False, italic=True):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_equation_font(run, size_pt, bold, italic)
    return para


def add_table_header_shading(cell, fill_hex='BDD7EE'):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), fill_hex)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(shading)


def create_word_report(project_title, inputs, calc_results, design_check, fig):
    """‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô Word ‡πÅ‡∏ö‡∏ö‡∏¢‡πà‡∏≠"""
    doc = Document()

    title = doc.add_heading('‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö Flexible Pavement', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_thai_font(run, size_pt=24, bold=True)

    heading1 = doc.add_heading(f'‡πÇ‡∏Ñ‡∏£‡∏á‡∏Å‡∏≤‡∏£: {project_title}', level=1)
    for run in heading1.runs:
        set_thai_font(run, size_pt=18, bold=True)

    add_thai_paragraph(doc, f'‡∏ß‡∏±‡∏ô‡∏ó‡∏µ‡πà‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö: {datetime.now().strftime("%d/%m/%Y %H:%M")}', size_pt=15)

    # Section 1
    h2 = doc.add_heading('1. ‡∏ß‡∏¥‡∏ò‡∏µ‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö', level=2)
    for run in h2.runs:
        set_thai_font(run, size_pt=16, bold=True)
    add_thai_paragraph(doc,
        '‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ñ‡∏ô‡∏ô‡πÉ‡∏ä‡πâ‡∏ß‡∏¥‡∏ò‡∏µ AASHTO 1993 Guide for Design of Pavement Structures '
        '‡∏ï‡∏≤‡∏°‡∏°‡∏≤‡∏ï‡∏£‡∏ê‡∏≤‡∏ô‡∏Å‡∏£‡∏°‡∏ó‡∏≤‡∏á‡∏´‡∏•‡∏ß‡∏á ‡πÇ‡∏î‡∏¢‡πÉ‡∏ä‡πâ‡∏™‡∏°‡∏Å‡∏≤‡∏£‡∏´‡∏•‡∏±‡∏Å‡∏î‡∏±‡∏á‡∏ô‡∏µ‡πâ:', size_pt=15)
    add_equation_paragraph(doc,
        'log‚ÇÅ‚ÇÄ(W‚ÇÅ‚Çà) = Z·µ£¬∑S‚Çí + 9.36¬∑log‚ÇÅ‚ÇÄ(SN+1) - 0.20 + '
        'log‚ÇÅ‚ÇÄ(ŒîPSI/2.7) / [0.4 + 1094/(SN+1)‚Åµ¬∑¬π‚Åπ] + 2.32¬∑log‚ÇÅ‚ÇÄ(M·µ£) - 8.07',
        size_pt=11, italic=True)

    # Section 2: Inputs
    h2_2 = doc.add_heading('2. ‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏ô‡∏≥‡πÄ‡∏Ç‡πâ‡∏≤ (Design Inputs)', level=2)
    for run in h2_2.runs:
        set_thai_font(run, size_pt=16, bold=True)

    input_table = doc.add_table(rows=1, cols=3)
    input_table.style = 'Table Grid'
    for i, h in enumerate(['‡∏û‡∏≤‡∏£‡∏≤‡∏°‡∏¥‡πÄ‡∏ï‡∏≠‡∏£‡πå', '‡∏Ñ‡πà‡∏≤', '‡∏´‡∏ô‡πà‡∏ß‡∏¢']):
        cell = input_table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_thai_font(r, size_pt=15, bold=True)
        add_table_header_shading(cell)

    for param, value, unit in [
        ('Design ESALs (W‚ÇÅ‚Çà)', f'{inputs["W18"]:,.0f}', '18-kip ESAL'),
        ('Reliability (R)', f'{inputs["reliability"]}', '%'),
        ('Z·µ£', f'{inputs["Zr"]:.3f}', '-'),
        ('S‚Çí', f'{inputs["So"]:.2f}', '-'),
        ('P‚ÇÄ', f'{inputs["P0"]:.1f}', '-'),
        ('P‚Çú', f'{inputs["Pt"]:.1f}', '-'),
        ('ŒîPSI', f'{inputs["delta_psi"]:.1f}', '-'),
        ('CBR ‡∏î‡∏¥‡∏ô‡πÄ‡∏î‡∏¥‡∏°', f'{inputs.get("CBR", "-")}', '%'),
        ('M·µ£ = 1500√óCBR', f'{inputs["Mr"]:,.0f}', 'psi'),
    ]:
        row = input_table.add_row()
        for j, val in enumerate([param, value, unit]):
            row.cells[j].text = ''
            p = row.cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j != 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            set_thai_font(r, size_pt=15)

    # Section 3: Material Properties
    h2_3 = doc.add_heading('3. ‡∏Ñ‡∏∏‡∏ì‡∏™‡∏°‡∏ö‡∏±‡∏ï‡∏¥‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á', level=2)
    for run in h2_3.runs:
        set_thai_font(run, size_pt=16, bold=True)

    mat_table = doc.add_table(rows=1, cols=6)
    mat_table.style = 'Table Grid'
    for i, h in enumerate(['‡∏ä‡∏±‡πâ‡∏ô', '‡∏ß‡∏±‡∏™‡∏î‡∏∏', 'a·µ¢', 'm·µ¢', 'M·µ£ (psi)', 'E (MPa)']):
        cell = mat_table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_thai_font(r, size_pt=15, bold=True)
        add_table_header_shading(cell)

    for layer in calc_results['layers']:
        row = mat_table.add_row()
        for j, val in enumerate([
            str(layer['layer_no']), short_material_name(layer['material']),
            f'{layer["a_i"]:.2f}', f'{layer["m_i"]:.2f}',
            f'{layer["mr_psi"]:,}', f'{layer["mr_mpa"]:,}'
        ]):
            row.cells[j].text = ''
            p = row.cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j != 1 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            set_thai_font(r, size_pt=15)

    # AC Sublayer breakdown
    ac_sub = calc_results.get('ac_sublayers', None)
    if ac_sub is not None:
        doc.add_paragraph()
        add_thai_paragraph(doc, '‡∏£‡∏≤‡∏¢‡∏•‡∏∞‡πÄ‡∏≠‡∏µ‡∏¢‡∏î‡∏ä‡∏±‡πâ‡∏ô‡∏¢‡πà‡∏≠‡∏¢‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á AC:', size_pt=15, bold=True)
        sub_table = doc.add_table(rows=1, cols=3)
        sub_table.style = 'Table Grid'
        for i, h in enumerate(['‡∏ä‡∏±‡πâ‡∏ô‡∏¢‡πà‡∏≠‡∏¢', '‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤ (cm)', '‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤ (mm)']):
            cell = sub_table.rows[0].cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(h)
            set_thai_font(r, size_pt=15, bold=True)
            add_table_header_shading(cell)
        for name, thick_cm in [
            ('‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡πÅ‡∏≠‡∏™‡∏ü‡∏±‡∏•‡∏ï‡πå‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï (AC. Wearing Course)', ac_sub['wearing']),
            ('‡∏£‡∏≠‡∏á‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡πÅ‡∏≠‡∏™‡∏ü‡∏±‡∏•‡∏ï‡πå‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï (AC. Binder Course)', ac_sub['binder']),
            ('‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡πÅ‡∏≠‡∏™‡∏ü‡∏±‡∏•‡∏ï‡πå‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï (AC. Base Course)', ac_sub['base']),
            ('‡∏£‡∏ß‡∏°', ac_sub['total']),
        ]:
            row = sub_table.add_row()
            for j, val in enumerate([name, f'{thick_cm:.1f}', f'{thick_cm*10:.0f}']):
                row.cells[j].text = ''
                p = row.cells[j].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j != 0 else WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(val)
                set_thai_font(r, size_pt=15)

    # Section 4: Step-by-step
    h2_4 = doc.add_heading('4. ‡∏Å‡∏≤‡∏£‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡πÅ‡∏ï‡πà‡∏•‡∏∞‡∏ä‡∏±‡πâ‡∏ô', level=2)
    for run in h2_4.runs:
        set_thai_font(run, size_pt=16, bold=True)

    for layer in calc_results['layers']:
        doc.add_paragraph()
        ln   = layer['layer_no']
        a_i  = layer['a_i']
        m_i  = layer['m_i']
        d_in = layer['design_thickness_inch']
        d_cm = layer['design_thickness_cm']
        sn_at     = layer['sn_required_at_layer']
        d_min_in  = layer['min_thickness_inch']
        d_min_cm  = layer['min_thickness_cm']
        sn_cont   = layer['sn_contribution']
        sn_cum    = layer['cumulative_sn']
        is_ok     = layer['is_ok']

        add_thai_paragraph(doc,
            f'‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏µ‡πà {ln}: {short_material_name(layer["material"])}',
            size_pt=15, bold=True)
        add_thai_paragraph(doc,
            f'  ‚Ä¢ Mr = {layer["mr_psi"]:,} psi = {layer["mr_mpa"]:,} MPa\n'
            f'  ‚Ä¢ a_{ln} = {a_i:.2f}   m_{ln} = {m_i:.2f}',
            size_pt=15)

        # --- SN ‡∏à‡∏≤‡∏Å‡∏™‡∏°‡∏Å‡∏≤‡∏£ AASHTO (Times New Roman 11pt) ---
        add_thai_paragraph(doc, '‡∏Å‡∏≤‡∏£‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì SN:', size_pt=15, bold=True)
        add_equation_paragraph(doc,
            f'‡∏à‡∏≤‡∏Å‡∏™‡∏°‡∏Å‡∏≤‡∏£ AASHTO 1993:   SN_{ln} = {sn_at:.2f}',
            size_pt=11, bold=True, italic=True)

        # --- ‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡∏Ç‡∏±‡πâ‡∏ô‡∏ï‡πà‡∏≥ + ‡πÅ‡∏ó‡∏ô‡∏Ñ‡πà‡∏≤‡∏ï‡∏±‡∏ß‡πÄ‡∏•‡∏Ç ---
        add_thai_paragraph(doc, '‡∏Å‡∏≤‡∏£‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡∏Ç‡∏±‡πâ‡∏ô‡∏ï‡πà‡∏≥:', size_pt=15, bold=True)
        if ln == 1:
            add_equation_paragraph(doc,
                f'D_1 >= SN_1 / (a_1 √ó m_1)',
                size_pt=11, italic=True)
            add_equation_paragraph(doc,
                f'D_1 >= {sn_at:.2f} / ({a_i:.2f} √ó {m_i:.2f})  =  {d_min_in:.2f} in  =  {d_min_cm:.1f} cm',
                size_pt=11, bold=True, italic=False)
        else:
            prev_sn = calc_results['layers'][ln-2]['cumulative_sn']
            add_equation_paragraph(doc,
                f'D_{ln} >= (SN_{ln} ‚àí SN_{ln-1}) / (a_{ln} √ó m_{ln})',
                size_pt=11, italic=True)
            add_equation_paragraph(doc,
                f'D_{ln} >= ({sn_at:.2f} ‚àí {prev_sn:.2f}) / ({a_i:.2f} √ó {m_i:.2f})'
                f'  =  {d_min_in:.2f} in  =  {d_min_cm:.1f} cm',
                size_pt=11, bold=True, italic=False)

        # --- ‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡∏ó‡∏µ‡πà‡πÄ‡∏•‡∏∑‡∏≠‡∏Å ---
        add_thai_paragraph(doc, '‡πÄ‡∏•‡∏∑‡∏≠‡∏Å‡πÉ‡∏ä‡πâ‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤:', size_pt=15, bold=True)
        add_equation_paragraph(doc,
            f'D_{ln}(design)  =  {d_cm:.0f} cm  ({d_in:.2f} in)',
            size_pt=11, bold=True, italic=False)

        # --- SN contribution + ‡πÅ‡∏ó‡∏ô‡∏Ñ‡πà‡∏≤‡∏ï‡∏±‡∏ß‡πÄ‡∏•‡∏Ç ---
        add_thai_paragraph(doc, 'SN contribution:', size_pt=15, bold=True)
        add_equation_paragraph(doc,
            f'ŒîSN_{ln} = a_{ln} √ó D_{ln} √ó m_{ln}'
            f'  =  {a_i:.2f} √ó {d_in:.2f} √ó {m_i:.2f}  =  {sn_cont:.3f}',
            size_pt=11, italic=True)
        add_equation_paragraph(doc,
            f'Œ£SN  =  {sn_cum:.2f}',
            size_pt=11, bold=True, italic=False)

        status = '‚úì OK ‚Äî ‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡πÄ‡∏û‡∏µ‡∏¢‡∏á‡∏û‡∏≠' if is_ok else f'‚úó NG ‚Äî ‡∏ï‡πâ‡∏≠‡∏á‡πÄ‡∏û‡∏¥‡πà‡∏°‡∏≠‡∏µ‡∏Å {d_min_cm - d_cm:.1f} cm'
        add_equation_paragraph(doc, f'‡∏™‡∏ñ‡∏≤‡∏ô‡∏∞: {status}', size_pt=11, bold=True, italic=False)

    # Section 5: SN Summary Table
    h2_5 = doc.add_heading('5. ‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏™‡∏£‡∏∏‡∏õ‡∏Å‡∏≤‡∏£‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì Structural Number', level=2)
    for run in h2_5.runs:
        set_thai_font(run, size_pt=16, bold=True)

    sn_table = doc.add_table(rows=1, cols=8)
    sn_table.style = 'Table Grid'
    for i, h in enumerate(['‡∏ä‡∏±‡πâ‡∏ô', '‡∏ß‡∏±‡∏™‡∏î‡∏∏', 'a·µ¢', 'm·µ¢', 'D·µ¢ (‡∏ô‡∏¥‡πâ‡∏ß)', 'D·µ¢ (‡∏ã‡∏°.)', 'ŒîSN·µ¢', 'Œ£SN']):
        cell = sn_table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_thai_font(r, size_pt=15, bold=True)
        add_table_header_shading(cell)
    for layer in calc_results['layers']:
        row = sn_table.add_row()
        for j, val in enumerate([
            str(layer['layer_no']), short_material_name(layer['material']),
            f'{layer["a_i"]:.2f}', f'{layer["m_i"]:.2f}',
            f'{layer["design_thickness_inch"]:.2f}', f'{layer["design_thickness_cm"]:.0f}',
            f'{layer["sn_contribution"]:.3f}', f'{layer["cumulative_sn"]:.2f}'
        ]):
            row.cells[j].text = ''
            p = row.cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j != 1 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            set_thai_font(r, size_pt=15)

    doc.add_paragraph()
    add_equation_paragraph(doc, '‡∏™‡∏π‡∏ï‡∏£: SN = Œ£(a·µ¢ √ó D·µ¢ √ó m·µ¢)', size_pt=11, italic=True)

    # Section 6: Design Check
    h2_6 = doc.add_heading('6. ‡∏ú‡∏•‡∏Å‡∏≤‡∏£‡∏ï‡∏£‡∏ß‡∏à‡∏™‡∏≠‡∏ö‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö', level=2)
    for run in h2_6.runs:
        set_thai_font(run, size_pt=16, bold=True)

    result_table = doc.add_table(rows=4, cols=2)
    result_table.style = 'Table Grid'
    for i, (param, value) in enumerate([
        ('SN Required (‡∏à‡∏≤‡∏Å‡∏™‡∏°‡∏Å‡∏≤‡∏£ AASHTO)', f'{calc_results["total_sn_required"]:.2f}'),
        ('SN Provided (‡∏à‡∏≤‡∏Å‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á)', f'{calc_results["total_sn_provided"]:.2f}'),
        ('Safety Margin', f'{design_check["safety_margin"]:.2f}'),
        ('‡∏ú‡∏•‡∏Å‡∏≤‡∏£‡∏ï‡∏£‡∏ß‡∏à‡∏™‡∏≠‡∏ö', '‡∏ú‡πà‡∏≤‡∏ô (OK)' if design_check['passed'] else '‡πÑ‡∏°‡πà‡∏ú‡πà‡∏≤‡∏ô (NG)'),
    ]):
        for j, val in enumerate([param, value]):
            result_table.rows[i].cells[j].text = ''
            p = result_table.rows[i].cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            set_thai_font(r, size_pt=15)

    doc.add_paragraph()
    w18_sup = calculate_w18_supported(
        calc_results['total_sn_provided'], inputs['Zr'], inputs['So'], inputs['delta_psi'], inputs['Mr'])
    add_thai_paragraph(doc, f'W‚ÇÅ‚Çà ‡∏ó‡∏µ‡πà‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏£‡∏≠‡∏á‡∏£‡∏±‡∏ö‡πÑ‡∏î‡πâ = {w18_sup/1e6:,.2f} ‡∏•‡πâ‡∏≤‡∏ô ESALs', size_pt=15, bold=True)

    if design_check['passed']:
        add_thai_paragraph(doc,
            f'‡∏™‡∏£‡∏∏‡∏õ: ‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö‡∏ú‡πà‡∏≤‡∏ô‡πÄ‡∏Å‡∏ì‡∏ë‡πå SN_provided ({calc_results["total_sn_provided"]:.2f}) '
            f'‚â• SN_required ({calc_results["total_sn_required"]:.2f})', size_pt=15, bold=True)
    else:
        add_thai_paragraph(doc, '‡∏™‡∏£‡∏∏‡∏õ: ‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö‡πÑ‡∏°‡πà‡∏ú‡πà‡∏≤‡∏ô‡πÄ‡∏Å‡∏ì‡∏ë‡πå ‡∏Å‡∏£‡∏∏‡∏ì‡∏≤‡∏õ‡∏£‡∏±‡∏ö‡πÄ‡∏û‡∏¥‡πà‡∏°‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á',
                           size_pt=15, bold=True)

    # Section 7: Figure
    h2_7 = doc.add_heading('7. ‡∏†‡∏≤‡∏û‡∏ï‡∏±‡∏î‡∏Ç‡∏ß‡∏≤‡∏á‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ñ‡∏ô‡∏ô', level=2)
    for run in h2_7.runs:
        set_thai_font(run, size_pt=16, bold=True)
    fig_bytes = get_figure_as_bytes(fig)
    doc.add_picture(fig_bytes, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Section 8: Summary table
    h2_8 = doc.add_heading('8. ‡∏™‡∏£‡∏∏‡∏õ‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏ó‡∏µ‡πà‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö', level=2)
    for run in h2_8.runs:
        set_thai_font(run, size_pt=16, bold=True)

    structure_rows = _build_structure_rows(calc_results, inputs.get('CBR', 3.0))
    sum_table = doc.add_table(rows=1 + len(structure_rows), cols=3)
    sum_table.style = 'Table Grid'
    sum_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['‡∏•‡∏≥‡∏î‡∏±‡∏ö', '‡∏ä‡∏ô‡∏¥‡∏î‡∏ß‡∏±‡∏™‡∏î‡∏∏', '‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤ (‡∏ã‡∏°.)']):
        cell = sum_table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_thai_font(r, size_pt=15, bold=True)
        add_table_header_shading(cell)
    for i, (num, mat_name, thickness) in enumerate(structure_rows):
        row = sum_table.rows[i + 1]
        for j, (val, align) in enumerate([
            (str(num), WD_ALIGN_PARAGRAPH.CENTER),
            (mat_name, WD_ALIGN_PARAGRAPH.LEFT),
            (thickness, WD_ALIGN_PARAGRAPH.CENTER)
        ]):
            row.cells[j].text = ''
            p = row.cells[j].paragraphs[0]
            p.alignment = align
            r = p.add_run(val)
            set_thai_font(r, size_pt=15)

    doc.add_paragraph()
    add_thai_paragraph(doc,
        '‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô‡∏ô‡∏µ‡πâ‡∏™‡∏£‡πâ‡∏≤‡∏á‡πÇ‡∏î‡∏¢‡πÅ‡∏≠‡∏õ‡∏û‡∏•‡∏¥‡πÄ‡∏Ñ‡∏ä‡∏±‡∏ô AASHTO 1993 Flexible Pavement Design v6.0\n'
        '‡∏û‡∏±‡∏í‡∏ô‡∏≤‡πÇ‡∏î‡∏¢ ‡∏£‡∏®.‡∏î‡∏£.‡∏≠‡∏¥‡∏ó‡∏ò‡∏¥‡∏û‡∏• ‡∏°‡∏µ‡∏ú‡∏• // ‡∏†‡∏≤‡∏Ñ‡∏ß‡∏¥‡∏ä‡∏≤‡∏Ñ‡∏£‡∏∏‡∏®‡∏≤‡∏™‡∏ï‡∏£‡πå‡πÇ‡∏¢‡∏ò‡∏≤ // ‡∏°‡∏à‡∏û.',
        size_pt=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes


def _build_structure_rows(calc_results, cbr_val):
    """‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏£‡∏≤‡∏¢‡∏Å‡∏≤‡∏£‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏™‡∏£‡∏∏‡∏õ (‡πÉ‡∏ä‡πâ‡∏£‡πà‡∏ß‡∏°‡∏Å‡∏±‡∏ô‡∏ó‡∏±‡πâ‡∏á 2 report)"""
    structure_rows = []
    row_num = 1
    ac_sub = calc_results.get('ac_sublayers', None)
    first_layer = calc_results['layers'][0] if calc_results.get('layers') else None

    if ac_sub is not None and first_layer:
        for key, label in [
            ('wearing', '‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡πÅ‡∏≠‡∏™‡∏ü‡∏±‡∏•‡∏ï‡πå‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï (AC. Wearing Course)'),
            ('binder',  '‡∏£‡∏≠‡∏á‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡πÅ‡∏≠‡∏™‡∏ü‡∏±‡∏•‡∏ï‡πå‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï (AC. Binder Course)'),
            ('base',    '‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡πÅ‡∏≠‡∏™‡∏ü‡∏±‡∏•‡∏ï‡πå‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï (AC. Base Course)'),
        ]:
            if ac_sub.get(key, 0) > 0:
                structure_rows.append((row_num, label, f"{ac_sub[key]:.0f}"))
                row_num += 1
        for layer in calc_results['layers'][1:]:
            structure_rows.append((row_num, short_material_name(layer['material']),
                                   f"{layer['design_thickness_cm']:.0f}"))
            row_num += 1
    else:
        for layer in calc_results.get('layers', []):
            structure_rows.append((row_num, short_material_name(layer['material']),
                                   f"{layer['design_thickness_cm']:.0f}"))
            row_num += 1

    structure_rows.append((row_num, '‡∏î‡∏¥‡∏ô‡∏Ñ‡∏±‡∏ô‡∏ó‡∏≤‡∏á', f'CBR ‚â• {cbr_val:.1f} %'))
    return structure_rows


def set_thai_distribute(para):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = para._element.get_or_add_pPr()
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'thaiDistribute')
    pPr.append(jc)


def create_word_report_intro(project_title, inputs, calc_results, design_check, fig, report_settings):
    """‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô Word ‡πÅ‡∏ö‡∏ö‡∏ó‡∏µ‡πà‡∏õ‡∏£‡∏∂‡∏Å‡∏©‡∏≤ (‡∏£‡∏ß‡∏°‡∏Å‡∏±‡∏ö‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô‡∏≠‡∏∑‡πà‡∏ô)"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'TH SarabunPSK'
    style.font.size = Pt(15)
    try:
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')
    except Exception:
        pass

    sec_no   = report_settings.get('section_number', '4.4')
    tbl_inp  = report_settings.get('table_number_inputs', '4-8')
    tbl_mat  = report_settings.get('table_number_materials', '4-9')
    fig_no   = report_settings.get('figure_number', '4-8')
    sec_title = report_settings.get('section_title', '‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡∏•‡∏≤‡∏î‡∏¢‡∏≤‡∏á (Flexible Pavement)')
    tbl_cap_inp  = report_settings.get('table_caption_inputs', '‡∏Ñ‡πà‡∏≤‡∏û‡∏≤‡∏£‡∏≤‡∏°‡∏¥‡πÄ‡∏ï‡∏≠‡∏£‡πå‡∏ó‡∏µ‡πà‡πÉ‡∏ä‡πâ‡πÉ‡∏ô‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡∏¢‡∏∑‡∏î‡∏´‡∏¢‡∏∏‡πà‡∏ô')
    tbl_cap_mat  = report_settings.get('table_caption_materials', '‡∏Ñ‡πà‡∏≤‡∏™‡∏±‡∏°‡∏õ‡∏£‡∏∞‡∏™‡∏¥‡∏ó‡∏ò‡∏¥‡πå‡πÅ‡∏•‡∏∞‡∏Ñ‡πà‡∏≤‡πÇ‡∏°‡∏î‡∏π‡∏•‡∏±‡∏™‡∏Ç‡∏≠‡∏á‡∏ß‡∏±‡∏™‡∏î‡∏∏‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á')
    fig_cap  = report_settings.get('figure_caption', '‡∏£‡∏π‡∏õ‡∏ï‡∏±‡∏î‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏ó‡∏µ‡πà‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö')

    RED   = RGBColor(0xCC, 0x00, 0x00)
    GREEN = RGBColor(0, 112, 0)
    BLUE  = RGBColor(0x00, 0x47, 0xAB)

    def _run(para, text, size=15, bold=False, italic=False, color=None, underline=False):
        r = para.add_run(text)
        r.font.name = 'TH SarabunPSK'
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        r.underline = underline
        if color:
            r.font.color.rgb = color
        try:
            r._element.rPr.rFonts.set(qn('w:cs'), 'TH SarabunPSK')
        except Exception:
            pass
        return r

    def _eq(para, text, size=11, bold=False, italic=True):
        """Times New Roman 11pt ‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡∏™‡∏°‡∏Å‡∏≤‡∏£"""
        r = para.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        try:
            r._element.rPr.rFonts.set(qn('w:cs'), 'Times New Roman')
        except Exception:
            pass
        return r

    def _eq_para(text, indent_cm=2.0, bold=False, italic=True, align=WD_ALIGN_PARAGRAPH.LEFT):
        """Paragraph ‡∏™‡∏°‡∏Å‡∏≤‡∏£ Times New Roman 11pt"""
        p = _para(indent_cm=indent_cm)
        p.alignment = align
        _eq(p, text, bold=bold, italic=italic)
        return p

    def _heading(text, level, size):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = 'TH SarabunPSK'
            run.font.size = Pt(size)
        return h

    def _para(indent_cm=0, space_before=0, space_after=4):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(indent_cm)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        return p

    def _tbl_cell(cell, text, align=WD_ALIGN_PARAGRAPH.CENTER, size=15, bold=False, fill=None):
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = align
        r = p.add_run(text)
        set_thai_font(r, size_pt=size, bold=bold)
        if fill:
            add_table_header_shading(cell, fill)

    def _fig_caption(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = 'TH SarabunPSK'
        r.font.size = Pt(14)
        r.italic = True

    # ===== ‡∏´‡∏±‡∏ß‡∏Ç‡πâ‡∏≠‡∏´‡∏•‡∏±‡∏Å =====
    h_main = _heading(f'{sec_no}  {sec_title}', level=2, size=16)

    # ===== ‡πÄ‡∏Å‡∏£‡∏¥‡πà‡∏ô‡∏ô‡∏≥ =====
    W18 = inputs.get('W18', 0)
    reliability = inputs.get('reliability', 90)
    CBR = inputs.get('CBR', 5.0)
    Mr = inputs.get('Mr', 7500)
    sn_req  = calc_results.get('total_sn_required', 0)
    sn_prov = calc_results.get('total_sn_provided', 0)
    total_thick = sum(l['design_thickness_cm'] for l in calc_results.get('layers', []))
    num_layers  = len(calc_results.get('layers', []))
    passed_txt  = '‡∏ú‡πà‡∏≤‡∏ô‡πÄ‡∏Å‡∏ì‡∏ë‡πå' if design_check.get('passed') else '‡πÑ‡∏°‡πà‡∏ú‡πà‡∏≤‡∏ô‡πÄ‡∏Å‡∏ì‡∏ë‡πå'

    p_intro = _para(indent_cm=0, space_before=6)
    p_intro.paragraph_format.first_line_indent = Cm(1.25)
    set_thai_distribute(p_intro)
    _run(p_intro, '‡∏ñ‡∏ô‡∏ô‡∏•‡∏≤‡∏î‡∏¢‡∏≤‡∏á‡∏ã‡∏∂‡πà‡∏á‡∏õ‡∏£‡∏∞‡∏Å‡∏≠‡∏ö‡∏î‡πâ‡∏ß‡∏¢‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏á‡∏≤‡∏ô‡∏ó‡∏≤‡∏á‡∏´‡∏•‡∏≤‡∏¢‡∏ä‡∏ô‡∏¥‡∏î ‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ñ‡∏ô‡∏ô‡πÅ‡∏ö‡∏ö‡∏¢‡∏∑‡∏î‡∏´‡∏¢‡∏∏‡πà‡∏ô '
         '(Flexible Pavement) ‡πÉ‡∏ä‡πâ‡∏ß‡∏¥‡∏ò‡∏µ AASHTO 1993 Guide for Design of Pavement Structures '
         '‡πÇ‡∏î‡∏¢‡∏û‡∏¥‡∏à‡∏≤‡∏£‡∏ì‡∏≤‡∏õ‡∏±‡∏à‡∏à‡∏±‡∏¢‡∏î‡πâ‡∏≤‡∏ô‡∏õ‡∏£‡∏¥‡∏°‡∏≤‡∏ì‡∏à‡∏£‡∏≤‡∏à‡∏£‡∏™‡∏∞‡∏™‡∏° ESALs ‡∏Ñ‡∏ß‡∏≤‡∏°‡∏ô‡πà‡∏≤‡πÄ‡∏ä‡∏∑‡πà‡∏≠‡∏ñ‡∏∑‡∏≠ ‡πÅ‡∏•‡∏∞‡∏Ñ‡∏∏‡∏ì‡∏™‡∏°‡∏ö‡∏±‡∏ï‡∏¥‡∏Ç‡∏≠‡∏á‡∏î‡∏¥‡∏ô‡∏£‡∏≠‡∏á‡∏£‡∏±‡∏ö '
         '‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡πÇ‡∏Ñ‡∏£‡∏á‡∏Å‡∏≤‡∏£‡∏ô‡∏µ‡πâ‡∏ó‡∏µ‡πà‡∏õ‡∏£‡∏∂‡∏Å‡∏©‡∏≤‡πÑ‡∏î‡πâ‡∏Å‡∏≥‡∏´‡∏ô‡∏î‡∏Ñ‡πà‡∏≤‡∏û‡∏≤‡∏£‡∏≤‡∏°‡∏¥‡πÄ‡∏ï‡∏≠‡∏£‡πå‡∏´‡∏•‡∏±‡∏Å‡πÉ‡∏ô‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö ‡πÑ‡∏î‡πâ‡πÅ‡∏Å‡πà '
         f'‡∏õ‡∏£‡∏¥‡∏°‡∏≤‡∏ì W\u2081\u2088 = ')
    _run(p_intro, f'{W18:,.0f}', bold=True, color=BLUE)
    _run(p_intro, f' 18-kip ESALs ‡∏ó‡∏µ‡πà‡∏£‡∏∞‡∏î‡∏±‡∏ö‡∏Ñ‡∏ß‡∏≤‡∏°‡∏ô‡πà‡∏≤‡πÄ‡∏ä‡∏∑‡πà‡∏≠‡∏ñ‡∏∑‡∏≠ (Reliability) = ')
    _run(p_intro, f'{reliability}', bold=True, color=BLUE)
    _run(p_intro, f' % ‡πÇ‡∏î‡∏¢‡∏°‡∏µ‡∏î‡∏¥‡∏ô‡πÄ‡∏î‡∏¥‡∏°‡∏Ñ‡πà‡∏≤ CBR = ')
    _run(p_intro, f'{CBR:.1f}', bold=True, color=BLUE)
    _run(p_intro, f' % (M\u1D63 = ')
    _run(p_intro, f'{Mr:,.0f}', bold=True, color=BLUE)
    _run(p_intro, f' psi) ‡∏ú‡∏•‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö‡πÑ‡∏î‡πâ‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á ')
    _run(p_intro, f'{num_layers}', bold=True, color=BLUE)
    _run(p_intro, f' ‡∏ä‡∏±‡πâ‡∏ô ‡∏ó‡∏µ‡πà SN\u200B_required = ')
    _run(p_intro, f'{sn_req:.2f}', bold=True, color=BLUE)
    _run(p_intro, ' ‡πÅ‡∏•‡∏∞ SN\u200B_provided = ')
    _run(p_intro, f'{sn_prov:.2f}', bold=True, color=BLUE)
    _run(p_intro, f' ‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡∏£‡∏ß‡∏° ')
    _run(p_intro, f'{total_thick:.0f}', bold=True, color=BLUE)
    _run(p_intro, f' ‡∏ã‡∏°. ‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö')
    _run(p_intro, passed_txt, bold=True, color=GREEN if design_check.get('passed') else RED)
    _run(p_intro, f' ‡∏î‡∏±‡∏á‡πÅ‡∏™‡∏î‡∏á‡∏ú‡∏•‡∏Å‡∏≤‡∏£‡∏ß‡∏¥‡πÄ‡∏Ñ‡∏£‡∏≤‡∏∞‡∏´‡πå‡πÉ‡∏ô‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏ó‡∏µ‡πà ')
    _run(p_intro, tbl_inp, bold=True)
    _run(p_intro, ' ‡πÅ‡∏•‡∏∞‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏ó‡∏µ‡πà ')
    _run(p_intro, tbl_mat, bold=True)
    _run(p_intro, ' ‡πÅ‡∏•‡∏∞‡∏£‡∏π‡∏õ‡∏ó‡∏µ‡πà ')
    _run(p_intro, fig_no, bold=True)

    # ===== {sec_no}.1 ‡∏ß‡∏¥‡∏ò‡∏µ‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö =====
    _heading(f'{sec_no}.1  ‡∏ß‡∏¥‡∏ò‡∏µ‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö', level=3, size=15)
    p_meth = _para(indent_cm=0, space_before=4)
    p_meth.paragraph_format.first_line_indent = Cm(1.25)
    _run(p_meth, '‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ñ‡∏ô‡∏ô‡πÉ‡∏ä‡πâ‡∏ß‡∏¥‡∏ò‡∏µ AASHTO 1993 Guide for Design of Pavement Structures '
         '‡∏ï‡∏≤‡∏°‡∏°‡∏≤‡∏ï‡∏£‡∏ê‡∏≤‡∏ô‡∏Å‡∏£‡∏°‡∏ó‡∏≤‡∏á‡∏´‡∏•‡∏ß‡∏á ‡πÇ‡∏î‡∏¢‡πÉ‡∏ä‡πâ‡∏™‡∏°‡∏Å‡∏≤‡∏£‡∏´‡∏•‡∏±‡∏Å‡∏î‡∏±‡∏á‡∏ô‡∏µ‡πâ')

    _eq_para(
        'log10(W18) = Zr¬∑So + 9.36¬∑log10(SN+1) - 0.20\n'
        '                   + log10(ŒîPSI/2.7) / [0.4 + 1094/(SN+1)^5.19]\n'
        '                   + 2.32¬∑log10(Mr) - 8.07',
        indent_cm=2.5, italic=True
    )

    # ===== {sec_no}.2 ‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏ô‡∏≥‡πÄ‡∏Ç‡πâ‡∏≤ + ‡∏ï‡∏≤‡∏£‡∏≤‡∏á =====
    _heading(f'{sec_no}.2  ‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏ô‡∏≥‡πÄ‡∏Ç‡πâ‡∏≤ (Design Inputs)', level=3, size=15)
    p_intro2 = _para(indent_cm=0, space_before=4)
    p_intro2.paragraph_format.first_line_indent = Cm(1.25)
    set_thai_distribute(p_intro2)
    _run(p_intro2,
         '‡πÉ‡∏ô‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ñ‡∏ô‡∏ô‡∏¢‡∏∑‡∏î‡∏´‡∏¢‡∏∏‡πà‡∏ô ‡∏Å‡∏≤‡∏£‡∏Å‡∏≥‡∏´‡∏ô‡∏î‡∏Ñ‡πà‡∏≤‡∏û‡∏≤‡∏£‡∏≤‡∏°‡∏¥‡πÄ‡∏ï‡∏≠‡∏£‡πå‡∏ô‡∏≥‡πÄ‡∏Ç‡πâ‡∏≤ (Design Inputs) ‡∏ñ‡∏∑‡∏≠‡πÄ‡∏õ‡πá‡∏ô'
         '‡∏Ç‡∏±‡πâ‡∏ô‡∏ï‡∏≠‡∏ô‡∏™‡∏≥‡∏Ñ‡∏±‡∏ç‡∏ó‡∏µ‡πà‡∏°‡∏µ‡∏ú‡∏•‡πÇ‡∏î‡∏¢‡∏ï‡∏£‡∏á‡∏ï‡πà‡∏≠‡∏Ñ‡∏ß‡∏≤‡∏°‡∏ñ‡∏π‡∏Å‡∏ï‡πâ‡∏≠‡∏á‡πÅ‡∏•‡∏∞‡∏Ñ‡∏ß‡∏≤‡∏°‡∏ô‡πà‡∏≤‡πÄ‡∏ä‡∏∑‡πà‡∏≠‡∏ñ‡∏∑‡∏≠‡∏Ç‡∏≠‡∏á‡πÅ‡∏ö‡∏ö‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ñ‡∏ô‡∏ô‡∏ó‡∏µ‡πà‡∏ï‡πâ‡∏≠‡∏á‡∏Å‡∏≤‡∏£ '
         '‡πÄ‡∏ô‡∏∑‡πà‡∏≠‡∏á‡∏à‡∏≤‡∏Å‡∏Ñ‡πà‡∏≤‡∏û‡∏≤‡∏£‡∏≤‡∏°‡∏¥‡πÄ‡∏ï‡∏≠‡∏£‡πå‡πÅ‡∏ï‡πà‡∏•‡∏∞‡∏ï‡∏±‡∏ß‡∏™‡∏∞‡∏ó‡πâ‡∏≠‡∏ô‡πÉ‡∏´‡πâ‡πÄ‡∏´‡πá‡∏ô‡∏™‡∏†‡∏≤‡∏û‡∏Å‡∏≤‡∏£‡πÉ‡∏ä‡πâ‡∏á‡∏≤‡∏ô‡∏à‡∏£‡∏¥‡∏á‡∏Ç‡∏≠‡∏á‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ñ‡∏ô‡∏ô '
         '‡πÑ‡∏î‡πâ‡πÅ‡∏Å‡πà ‡∏õ‡∏£‡∏¥‡∏°‡∏≤‡∏ì‡∏Å‡∏≤‡∏£‡∏à‡∏£‡∏≤‡∏à‡∏£‡∏ï‡∏•‡∏≠‡∏î‡∏≠‡∏≤‡∏¢‡∏∏‡∏Å‡∏≤‡∏£‡πÉ‡∏ä‡πâ‡∏á‡∏≤‡∏ô ‡∏£‡∏∞‡∏î‡∏±‡∏ö‡∏Ñ‡∏ß‡∏≤‡∏°‡∏ô‡πà‡∏≤‡πÄ‡∏ä‡∏∑‡πà‡∏≠‡∏ñ‡∏∑‡∏≠‡∏ó‡∏µ‡πà‡∏¢‡∏≠‡∏°‡∏£‡∏±‡∏ö‡πÑ‡∏î‡πâ '
         '‡∏£‡∏ß‡∏°‡∏ñ‡∏∂‡∏á‡∏Ñ‡∏∏‡∏ì‡∏™‡∏°‡∏ö‡∏±‡∏ï‡∏¥‡∏Ç‡∏≠‡∏á‡∏î‡∏¥‡∏ô‡∏£‡∏≠‡∏á‡∏£‡∏±‡∏ö‡πÉ‡∏ô‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏µ‡πà‡πÇ‡∏Ñ‡∏£‡∏á‡∏Å‡∏≤‡∏£ ‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡πÇ‡∏Ñ‡∏£‡∏á‡∏Å‡∏≤‡∏£‡∏ô‡∏µ‡πâ '
         '‡∏ó‡∏µ‡πà‡∏õ‡∏£‡∏∂‡∏Å‡∏©‡∏≤‡πÑ‡∏î‡πâ‡∏Å‡∏≥‡∏´‡∏ô‡∏î‡∏Ñ‡πà‡∏≤‡∏û‡∏≤‡∏£‡∏≤‡∏°‡∏¥‡πÄ‡∏ï‡∏≠‡∏£‡πå‡∏´‡∏•‡∏±‡∏Å‡∏ó‡∏µ‡πà‡πÉ‡∏ä‡πâ‡πÉ‡∏ô‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö‡∏ï‡∏≤‡∏°‡πÅ‡∏ô‡∏ß‡∏ó‡∏≤‡∏á‡∏Ç‡∏≠‡∏á AASHTO '
         '‡∏ã‡∏∂‡πà‡∏á‡∏õ‡∏£‡∏∞‡∏Å‡∏≠‡∏ö‡∏î‡πâ‡∏ß‡∏¢‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏î‡πâ‡∏≤‡∏ô‡∏Ñ‡∏ß‡∏≤‡∏°‡∏™‡∏≤‡∏°‡∏≤‡∏£‡∏ñ‡πÉ‡∏ô‡∏Å‡∏≤‡∏£‡∏£‡∏±‡∏ö‡∏ô‡πâ‡∏≥‡∏´‡∏ô‡∏±‡∏Å‡∏Ç‡∏≠‡∏á‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á '
         '‡∏õ‡∏£‡∏¥‡∏°‡∏≤‡∏ì‡∏à‡∏£‡∏≤‡∏à‡∏£‡∏ó‡∏µ‡πà‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ñ‡∏ô‡∏ô‡∏ï‡πâ‡∏≠‡∏á‡∏£‡∏≠‡∏á‡∏£‡∏±‡∏ö‡∏ï‡∏•‡∏≠‡∏î‡∏≠‡∏≤‡∏¢‡∏∏‡∏Å‡∏≤‡∏£‡πÉ‡∏ä‡πâ‡∏á‡∏≤‡∏ô '
         '‡∏£‡∏ß‡∏°‡∏ñ‡∏∂‡∏á‡∏Ñ‡∏∏‡∏ì‡∏™‡∏°‡∏ö‡∏±‡∏ï‡∏¥‡∏Ç‡∏≠‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏î‡∏¥‡∏ô‡∏ó‡∏µ‡πà‡∏ï‡πâ‡∏≠‡∏á‡∏ã‡πà‡∏≠‡∏°‡∏ö‡∏≥‡∏£‡∏∏‡∏á‡∏´‡∏£‡∏∑‡∏≠‡∏õ‡∏£‡∏±‡∏ö‡∏õ‡∏£‡∏∏‡∏á‡πÉ‡∏´‡∏°‡πà '
         f'‡∏£‡∏ß‡∏°‡∏ñ‡∏∂‡∏á‡∏Ñ‡∏∏‡∏ì‡∏™‡∏°‡∏ö‡∏±‡∏ï‡∏¥‡∏Ç‡∏≠‡∏á‡∏î‡∏¥‡∏ô‡∏ä‡∏±‡πâ‡∏ô‡∏£‡∏≠‡∏á‡∏£‡∏±‡∏ö ‡∏£‡∏≤‡∏¢‡∏•‡∏∞‡πÄ‡∏≠‡∏µ‡∏¢‡∏î‡∏Ç‡∏≠‡∏á‡∏Ñ‡πà‡∏≤‡∏û‡∏≤‡∏£‡∏≤‡∏°‡∏¥‡πÄ‡∏ï‡∏≠‡∏£‡πå‡∏ó‡∏±‡πâ‡∏á‡∏´‡∏°‡∏î‡πÅ‡∏™‡∏î‡∏á‡πÉ‡∏ô‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏ó‡∏µ‡πà {tbl_inp}')
    p_tbl1_cap = _para(indent_cm=0, space_before=4)
    p_tbl1_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_tbl1_cap, f'‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏ó‡∏µ‡πà {tbl_inp}  {tbl_cap_inp}', bold=True)

    inp_tbl = doc.add_table(rows=1, cols=3)
    inp_tbl.style = 'Table Grid'
    inp_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['‡∏û‡∏≤‡∏£‡∏≤‡∏°‡∏¥‡πÄ‡∏ï‡∏≠‡∏£‡πå', '‡∏Ñ‡πà‡∏≤', '‡∏´‡∏ô‡πà‡∏ß‡∏¢']):
        _tbl_cell(inp_tbl.rows[0].cells[j], h, bold=True, fill='D9E2F3')
    for param, value, unit in [
        ('Design ESALs (W\u2081\u2088)', f'{W18:,.0f}', '18-kip ESAL'),
        ('Reliability (R)', f'{reliability}', '%'),
        ('Z\u1D63', f'{inputs.get("Zr", 0):.3f}', '-'),
        ('S\u2080', f'{inputs.get("So", 0):.2f}', '-'),
        ('P\u2080', f'{inputs.get("P0", 0):.1f}', '-'),
        ('P\u209C', f'{inputs.get("Pt", 0):.1f}', '-'),
        ('\u0394PSI', f'{inputs.get("delta_psi", 0):.1f}', '-'),
        ('CBR ‡∏î‡∏¥‡∏ô‡πÄ‡∏î‡∏¥‡∏°', f'{CBR:.1f}', '%'),
        ('M\u1D63 = 1,500\u00D7CBR', f'{Mr:,.0f}', 'psi'),
    ]:
        row = inp_tbl.add_row()
        _tbl_cell(row.cells[0], param, align=WD_ALIGN_PARAGRAPH.LEFT)
        _tbl_cell(row.cells[1], value)
        _tbl_cell(row.cells[2], unit)

    # ===== {sec_no}.3 ‡∏Ñ‡∏∏‡∏ì‡∏™‡∏°‡∏ö‡∏±‡∏ï‡∏¥‡∏ß‡∏±‡∏™‡∏î‡∏∏ + ‡∏ï‡∏≤‡∏£‡∏≤‡∏á =====
    _heading(f'{sec_no}.3  ‡∏Ñ‡∏∏‡∏ì‡∏™‡∏°‡∏ö‡∏±‡∏ï‡∏¥‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á', level=3, size=15)
    p_intro3 = _para(indent_cm=0, space_before=4)
    p_intro3.paragraph_format.first_line_indent = Cm(1.25)
    set_thai_distribute(p_intro3)
    _run(p_intro3,
         '‡∏ß‡∏±‡∏™‡∏î‡∏∏‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡πÅ‡∏ï‡πà‡∏•‡∏∞‡∏ä‡∏ô‡∏¥‡∏î‡∏°‡∏µ‡∏Ñ‡πà‡∏≤‡∏™‡∏±‡∏°‡∏õ‡∏£‡∏∞‡∏™‡∏¥‡∏ó‡∏ò‡∏¥‡πå‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á (Layer Coefficient) '
         '‡πÅ‡∏•‡∏∞‡∏Ñ‡πà‡∏≤‡∏™‡∏±‡∏°‡∏õ‡∏£‡∏∞‡∏™‡∏¥‡∏ó‡∏ò‡∏¥‡πå‡∏Å‡∏≤‡∏£‡∏£‡∏∞‡∏ö‡∏≤‡∏¢‡∏ô‡πâ‡∏≥ (Drainage Coefficient) '
         '‡πÇ‡∏î‡∏¢‡∏ó‡∏µ‡πà‡∏õ‡∏£‡∏∂‡∏Å‡∏©‡∏≤‡πÄ‡∏•‡∏∑‡∏≠‡∏Å‡πÉ‡∏ä‡πâ‡∏ß‡∏±‡∏™‡∏î‡∏∏‡πÅ‡∏•‡∏∞‡πÅ‡∏™‡∏î‡∏á‡∏Ñ‡πà‡∏≤‡∏™‡∏±‡∏°‡∏õ‡∏£‡∏∞‡∏™‡∏¥‡∏ó‡∏ò‡∏¥‡πå‡∏£‡∏ß‡∏°‡∏ñ‡∏∂‡∏á‡∏Ñ‡πà‡∏≤‡πÇ‡∏°‡∏î‡∏π‡∏•‡∏±‡∏™‡∏Ç‡∏≠‡∏á‡∏ß‡∏±‡∏™‡∏î‡∏∏'
         f'‡∏ï‡πà‡∏≤‡∏á ‡πÜ ‡∏î‡∏±‡∏á‡πÅ‡∏™‡∏î‡∏á‡πÉ‡∏ô‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏ó‡∏µ‡πà {tbl_mat}')
    p_tbl2_cap = _para(indent_cm=0, space_before=4)
    p_tbl2_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_tbl2_cap, f'‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏ó‡∏µ‡πà {tbl_mat}  {tbl_cap_mat}', bold=True)

    mat_tbl = doc.add_table(rows=1, cols=6)
    mat_tbl.style = 'Table Grid'
    mat_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['‡∏ä‡∏±‡πâ‡∏ô', '‡∏ß‡∏±‡∏™‡∏î‡∏∏', 'a\u1D62', 'm\u1D62', 'M\u1D63 (psi)', 'E (MPa)']):
        _tbl_cell(mat_tbl.rows[0].cells[j], h, bold=True, fill='D9E2F3')
    for layer in calc_results.get('layers', []):
        row = mat_tbl.add_row()
        _tbl_cell(row.cells[0], str(layer['layer_no']))
        _tbl_cell(row.cells[1], short_material_name(layer['material']), align=WD_ALIGN_PARAGRAPH.LEFT)
        _tbl_cell(row.cells[2], f'{layer["a_i"]:.2f}')
        _tbl_cell(row.cells[3], f'{layer["m_i"]:.2f}')
        _tbl_cell(row.cells[4], f'{layer["mr_psi"]:,}')
        _tbl_cell(row.cells[5], f'{layer["mr_mpa"]:,}')

    # AC Sublayer table (‡∏ñ‡πâ‡∏≤‡∏°‡∏µ)
    ac_sub = calc_results.get('ac_sublayers', None)
    if ac_sub:
        p_sub = _para(indent_cm=0, space_before=6)
        _run(p_sub, '‡∏£‡∏≤‡∏¢‡∏•‡∏∞‡πÄ‡∏≠‡∏µ‡∏¢‡∏î‡∏ä‡∏±‡πâ‡∏ô‡∏¢‡πà‡∏≠‡∏¢‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡πÅ‡∏≠‡∏™‡∏ü‡∏±‡∏•‡∏ï‡πå‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï:', bold=True)
        sub_tbl = doc.add_table(rows=1, cols=3)
        sub_tbl.style = 'Table Grid'
        sub_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, h in enumerate(['‡∏ä‡∏±‡πâ‡∏ô‡∏¢‡πà‡∏≠‡∏¢', '‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤ (cm)', '‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤ (mm)']):
            _tbl_cell(sub_tbl.rows[0].cells[j], h, bold=True, fill='D9E2F3')
        for key, label in [
            ('wearing', '‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡πÅ‡∏≠‡∏™‡∏ü‡∏±‡∏•‡∏ï‡πå‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï (AC. Wearing Course)'),
            ('binder',  '‡∏£‡∏≠‡∏á‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡πÅ‡∏≠‡∏™‡∏ü‡∏±‡∏•‡∏ï‡πå‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï (AC. Binder Course)'),
            ('base',    '‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡πÅ‡∏≠‡∏™‡∏ü‡∏±‡∏•‡∏ï‡πå‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï (AC. Base Course)'),
            ('total',   '‡∏£‡∏ß‡∏°'),
        ]:
            val = ac_sub.get(key, 0)
            if key == 'total' or val > 0:
                row = sub_tbl.add_row()
                _tbl_cell(row.cells[0], label, align=WD_ALIGN_PARAGRAPH.LEFT)
                _tbl_cell(row.cells[1], f'{val:.1f}')
                _tbl_cell(row.cells[2], f'{val*10:.0f}' if key != 'total' else f'{val*10:.0f}')

    # ===== {sec_no}.4 ‡∏Ç‡∏±‡πâ‡∏ô‡∏ï‡∏≠‡∏ô‡∏Å‡∏≤‡∏£‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì =====
    _heading(f'{sec_no}.4  ‡∏Ç‡∏±‡πâ‡∏ô‡∏ï‡∏≠‡∏ô‡∏Å‡∏≤‡∏£‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á', level=3, size=15)
    p_intro4 = _para(indent_cm=0, space_before=4)
    p_intro4.paragraph_format.first_line_indent = Cm(1.25)
    set_thai_distribute(p_intro4)
    _run(p_intro4,
         '‡∏Å‡∏≤‡∏£‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡∏Ç‡∏±‡πâ‡∏ô‡∏ï‡πà‡∏≥‡∏Ç‡∏≠‡∏á‡πÅ‡∏ï‡πà‡∏•‡∏∞‡∏ä‡∏±‡πâ‡∏ô ‡πÉ‡∏ä‡πâ‡∏´‡∏•‡∏±‡∏Å‡∏Å‡∏≤‡∏£‡∏ß‡πà‡∏≤ Structural Number (SN) ‡∏ó‡∏µ‡πà‡∏à‡∏∏‡∏î‡πÉ‡∏î ‡πÜ '
         '‡∏ï‡πâ‡∏≠‡∏á‡∏°‡∏≤‡∏Å‡∏Å‡∏ß‡πà‡∏≤‡∏´‡∏£‡∏∑‡∏≠‡πÄ‡∏ó‡πà‡∏≤‡∏Å‡∏±‡∏ö SN ‡∏ó‡∏µ‡πà‡∏ï‡πâ‡∏≠‡∏á‡∏Å‡∏≤‡∏£ ‡πÇ‡∏î‡∏¢‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡∏à‡∏≤‡∏Å‡∏Ñ‡πà‡∏≤ M\u1D63 ‡∏Ç‡∏≠‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ñ‡∏±‡∏î‡πÑ‡∏õ')

    for layer in calc_results.get('layers', []):
        sn_at    = layer['sn_required_at_layer']
        layer_no = layer['layer_no']
        a_i      = layer['a_i']
        m_i      = layer['m_i']
        d_in     = layer['design_thickness_inch']
        d_cm     = layer['design_thickness_cm']
        d_min_in = layer['min_thickness_inch']
        d_min_cm = layer['min_thickness_cm']
        sn_cont  = layer['sn_contribution']
        sn_cum   = layer['cumulative_sn']
        is_ok    = layer['is_ok']

        doc.add_paragraph()
        hdr_p = _para(indent_cm=1.0, space_before=6)
        _run(hdr_p, f'‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏µ‡πà {layer_no}: {short_material_name(layer["material"])}',
             bold=True, underline=True)

        p_mat = _para(indent_cm=1.5)
        _run(p_mat, '‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏ß‡∏±‡∏™‡∏î‡∏∏:', bold=True)
        p_mat2 = _para(indent_cm=2.0)
        _run(p_mat2,
             f'\u2022 Mr = {layer["mr_psi"]:,} psi  =  {layer["mr_mpa"]:,} MPa\n'
             f'\u2022 Layer Coefficient (a{layer_no}) = {a_i:.2f}\n'
             f'\u2022 Drainage Coefficient (m{layer_no}) = {m_i:.2f}')

        # --- ‡∏Å‡∏≤‡∏£‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì SN ---
        p_sn = _para(indent_cm=1.5)
        _run(p_sn, '‡∏Å‡∏≤‡∏£‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì SN:', bold=True)
        # SN_N notation (‡πÑ‡∏°‡πà‡πÉ‡∏ä‡πâ subscript 0)
        _eq_para(f'‡∏à‡∏≤‡∏Å‡∏™‡∏°‡∏Å‡∏≤‡∏£ AASHTO 1993:   SN_{layer_no} = {sn_at:.2f}',
                 indent_cm=2.0, bold=True, italic=True)

        # --- ‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡∏Ç‡∏±‡πâ‡∏ô‡∏ï‡πà‡∏≥ ‡∏û‡∏£‡πâ‡∏≠‡∏°‡πÅ‡∏ó‡∏ô‡∏Ñ‡πà‡∏≤‡∏ï‡∏±‡∏ß‡πÄ‡∏•‡∏Ç ---
        p_th = _para(indent_cm=1.5)
        _run(p_th, '‡∏Å‡∏≤‡∏£‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡∏Ç‡∏±‡πâ‡∏ô‡∏ï‡πà‡∏≥:', bold=True)
        if layer_no == 1:
            # ‡∏™‡∏°‡∏Å‡∏≤‡∏£‡∏ó‡∏±‡πà‡∏ß‡πÑ‡∏õ
            _eq_para(
                f'D_1 >= SN_1 / (a_1 √ó m_1)',
                indent_cm=2.5, italic=True
            )
            # ‡πÅ‡∏ó‡∏ô‡∏Ñ‡πà‡∏≤‡∏ï‡∏±‡∏ß‡πÄ‡∏•‡∏Ç
            _eq_para(
                f'D_1 >= {sn_at:.2f} / ({a_i:.2f} √ó {m_i:.2f})  =  {d_min_in:.2f} in  =  {d_min_cm:.1f} cm',
                indent_cm=2.5, bold=True, italic=False
            )
        else:
            prev_sn = calc_results['layers'][layer_no - 2]['cumulative_sn']
            # ‡∏™‡∏°‡∏Å‡∏≤‡∏£‡∏ó‡∏±‡πà‡∏ß‡πÑ‡∏õ
            _eq_para(
                f'D_{layer_no} >= (SN_{layer_no} ‚àí SN_{layer_no-1}) / (a_{layer_no} √ó m_{layer_no})',
                indent_cm=2.5, italic=True
            )
            # ‡πÅ‡∏ó‡∏ô‡∏Ñ‡πà‡∏≤‡∏ï‡∏±‡∏ß‡πÄ‡∏•‡∏Ç
            _eq_para(
                f'D_{layer_no} >= ({sn_at:.2f} ‚àí {prev_sn:.2f}) / ({a_i:.2f} √ó {m_i:.2f})'
                f'  =  {d_min_in:.2f} in  =  {d_min_cm:.1f} cm',
                indent_cm=2.5, bold=True, italic=False
            )

        # --- ‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡∏ó‡∏µ‡πà‡πÄ‡∏•‡∏∑‡∏≠‡∏Å‡πÉ‡∏ä‡πâ ---
        p_d = _para(indent_cm=1.5)
        _run(p_d, '‡πÄ‡∏•‡∏∑‡∏≠‡∏Å‡πÉ‡∏ä‡πâ‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤:', bold=True)
        _eq_para(
            f'D_{layer_no}(design)  =  {d_cm:.0f} cm  ({d_in:.2f} in)',
            indent_cm=2.5, bold=True, italic=False
        )

        # --- SN contribution ‡∏û‡∏£‡πâ‡∏≠‡∏°‡πÅ‡∏ó‡∏ô‡∏Ñ‡πà‡∏≤ ---
        p_sn2 = _para(indent_cm=1.5)
        _run(p_sn2, 'SN contribution:', bold=True)
        _eq_para(
            f'ŒîSN_{layer_no} = a_{layer_no} √ó D_{layer_no} √ó m_{layer_no}'
            f'  =  {a_i:.2f} √ó {d_in:.2f} √ó {m_i:.2f}  =  {sn_cont:.3f}',
            indent_cm=2.5, italic=True
        )
        _eq_para(
            f'Œ£SN  =  {sn_cum:.2f}',
            indent_cm=2.5, bold=True, italic=False
        )

        # --- ‡∏™‡∏ñ‡∏≤‡∏ô‡∏∞ ---
        status_txt  = '‚úì OK' if is_ok else '‚úó NG'
        status_note = (f'‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡πÄ‡∏û‡∏µ‡∏¢‡∏á‡∏û‡∏≠ ({d_cm:.0f} ‚â• {d_min_cm:.1f} cm)'
                       if is_ok else f'‡∏ï‡πâ‡∏≠‡∏á‡πÄ‡∏û‡∏¥‡πà‡∏°‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡∏≠‡∏µ‡∏Å {d_min_cm - d_cm:.1f} cm')
        p_st = _para(indent_cm=2.0)
        _run(p_st, f'‡∏™‡∏ñ‡∏≤‡∏ô‡∏∞:  {status_txt}  ‚Äî  {status_note}',
             bold=True, color=GREEN if is_ok else RED)

    # ‡∏™‡∏£‡∏∏‡∏õ‡∏ú‡∏•
    doc.add_paragraph()
    safety_margin = design_check.get('safety_margin', sn_prov - sn_req)
    p_sum = _para(indent_cm=1.0)
    _run(p_sum, '‡∏™‡∏£‡∏∏‡∏õ‡∏ú‡∏•‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö:', bold=True)
    p_sum2 = _para(indent_cm=2.0)
    _run(p_sum2,
         f'SN required  =  {sn_req:.2f}   |   '
         f'SN provided  =  {sn_prov:.2f}   |   '
         f'Safety Margin  =  {safety_margin:.2f}', bold=True)
    p_sum3 = _para(indent_cm=2.0)
    result_txt = '‚úì ‡∏ú‡πà‡∏≤‡∏ô‡πÄ‡∏Å‡∏ì‡∏ë‡πå (OK)' if design_check.get('passed') else '‚úó ‡πÑ‡∏°‡πà‡∏ú‡πà‡∏≤‡∏ô‡πÄ‡∏Å‡∏ì‡∏ë‡πå (NG)'
    _run(p_sum3, f'‡∏ú‡∏•‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö:  {result_txt}', bold=True,
         color=GREEN if design_check.get('passed') else RED)

    # ‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏™‡∏£‡∏∏‡∏õ‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á
    doc.add_paragraph()
    surf_name = short_material_name(calc_results['layers'][0]['material']) if calc_results.get('layers') else ''
    p_sf = _para(indent_cm=0, space_before=6)
    _run(p_sf, f'‡∏£‡∏π‡∏õ‡πÅ‡∏ö‡∏ö‡∏ó‡∏µ‡πà: {surf_name}', bold=True)

    structure_rows = _build_structure_rows(calc_results, inputs.get('CBR', 3.0))
    sum_tbl = doc.add_table(rows=1 + len(structure_rows), cols=3)
    sum_tbl.style = 'Table Grid'
    sum_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['‡∏•‡∏≥‡∏î‡∏±‡∏ö', '‡∏ä‡∏ô‡∏¥‡∏î‡∏ß‡∏±‡∏™‡∏î‡∏∏', '‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤ (‡∏ã‡∏°.)']):
        _tbl_cell(sum_tbl.rows[0].cells[j], h, bold=True, fill='BDD7EE')
    for i, (num, mat_name, thickness) in enumerate(structure_rows):
        row = sum_tbl.rows[i + 1]
        _tbl_cell(row.cells[0], str(num))
        _tbl_cell(row.cells[1], mat_name, align=WD_ALIGN_PARAGRAPH.LEFT)
        _tbl_cell(row.cells[2], thickness)

    # ‡∏£‡∏π‡∏õ‡∏ï‡∏±‡∏î‡∏Ç‡∏ß‡∏≤‡∏á + caption
    doc.add_paragraph()
    fig_bytes_intro = get_figure_as_bytes(fig)
    doc.add_picture(fig_bytes_intro, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _fig_caption(f'‡∏£‡∏π‡∏õ‡∏ó‡∏µ‡πà {fig_no}  {fig_cap}')

    # Footer
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(footer_p,
         '‡∏û‡∏±‡∏í‡∏ô‡∏≤‡πÇ‡∏î‡∏¢ ‡∏£‡∏®.‡∏î‡∏£.‡∏≠‡∏¥‡∏ó‡∏ò‡∏¥‡∏û‡∏• ‡∏°‡∏µ‡∏ú‡∏• // ‡∏†‡∏≤‡∏Ñ‡∏ß‡∏¥‡∏ä‡∏≤‡∏Ñ‡∏£‡∏∏‡∏®‡∏≤‡∏™‡∏ï‡∏£‡πå‡πÇ‡∏¢‡∏ò‡∏≤ // ‡∏°‡∏à‡∏û.',
         size=12, italic=True)

    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes


# ================================================================================
# MAIN APP
# ================================================================================

def main():
    # ========================================
    # HEADER
    # ========================================
    st.title("üõ£Ô∏è  Flexible Pavement Design (AASHTO 1993) v6")
    st.markdown("**‡πÅ‡∏≠‡∏õ‡∏û‡∏•‡∏¥‡πÄ‡∏Ñ‡∏ä‡∏±‡∏ô‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ó‡∏≤‡∏á‡πÅ‡∏ö‡∏ö‡∏¢‡∏∑‡∏î‡∏´‡∏¢‡∏∏‡πà‡∏ô ‡∏ï‡∏≤‡∏°‡∏ß‡∏¥‡∏ò‡∏µ‡∏Å‡∏≤‡∏£ AASHTO (1993) ‚Äî ‡∏°‡∏≤‡∏ï‡∏£‡∏ê‡∏≤‡∏ô‡∏Å‡∏£‡∏°‡∏ó‡∏≤‡∏á‡∏´‡∏•‡∏ß‡∏á**")

    # ========================================
    # SIDEBAR ‚Äî ‡∏Å‡∏£‡∏∞‡∏ä‡∏±‡∏ö: Project / Preset / JSON / ‡∏†‡∏≤‡∏©‡∏≤‡∏£‡∏π‡∏õ
    # ========================================
    with st.sidebar:
        st.header("üìã ‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡πÇ‡∏Ñ‡∏£‡∏á‡∏Å‡∏≤‡∏£")
        project_title = st.text_input(
            "‡∏ä‡∏∑‡πà‡∏≠‡πÇ‡∏Ñ‡∏£‡∏á‡∏Å‡∏≤‡∏£",
            value=st.session_state.get('input_project_title', "‡πÇ‡∏Ñ‡∏£‡∏á‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö‡∏ñ‡∏ô‡∏ô"),
            key="project_title_input"
        )

        st.markdown("---")
        st.header("üèóÔ∏è Preset ‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á ‡∏ó‡∏•.")
        preset_names = list(PRESET_STRUCTURES.keys())
        selected_preset = st.selectbox(
            "‡πÄ‡∏•‡∏∑‡∏≠‡∏Å‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏°‡∏≤‡∏ï‡∏£‡∏ê‡∏≤‡∏ô", options=preset_names, index=0,
            help="‡πÄ‡∏•‡∏∑‡∏≠‡∏Å‡πÄ‡∏û‡∏∑‡πà‡∏≠ Auto-fill ‡∏ß‡∏±‡∏™‡∏î‡∏∏‡πÅ‡∏•‡∏∞‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡πÄ‡∏£‡∏¥‡πà‡∏°‡∏ï‡πâ‡∏ô (‡πÅ‡∏Å‡πâ‡πÑ‡∏Ç‡πÄ‡∏û‡∏¥‡πà‡∏°‡πÄ‡∏ï‡∏¥‡∏°‡πÑ‡∏î‡πâ)"
        )
        if selected_preset != "--- ‡πÄ‡∏•‡∏∑‡∏≠‡∏Å‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏°‡∏≤‡∏ï‡∏£‡∏ê‡∏≤‡∏ô ---":
            preset = PRESET_STRUCTURES[selected_preset]
            if preset:
                st.info(f"üìã {preset['description']}")
                if st.button("‚úÖ ‡πÉ‡∏ä‡πâ‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ô‡∏µ‡πâ", type="primary"):
                    st.session_state['input_num_layers'] = preset['num_layers']
                    for i, layer in enumerate(preset['layers']):
                        st.session_state[f'layer{i+1}_mat'] = layer['material']
                        st.session_state[f'layer{i+1}_thick'] = layer['thickness_cm']
                        mat = MATERIALS[layer['material']]
                        st.session_state[f'layer{i+1}_a'] = mat['layer_coeff']
                        st.session_state[f'layer{i+1}_m'] = mat['drainage_coeff']
                    st.session_state['use_ac_sublayers'] = False
                    st.session_state['ac_sublayers'] = None
                    st.rerun()

        st.markdown("---")
        st.header("üíæ ‡∏ö‡∏±‡∏ô‡∏ó‡∏∂‡∏Å/‡πÇ‡∏´‡∏•‡∏î‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•")
        uploaded_json = st.file_uploader(
            "üìÇ ‡πÇ‡∏´‡∏•‡∏î‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏à‡∏≤‡∏Å‡πÑ‡∏ü‡∏•‡πå JSON", type=['json'],
            help="‡∏≠‡∏±‡∏õ‡πÇ‡∏´‡∏•‡∏î‡πÑ‡∏ü‡∏•‡πå JSON ‡∏ó‡∏µ‡πà‡∏ö‡∏±‡∏ô‡∏ó‡∏∂‡∏Å‡πÑ‡∏ß‡πâ‡∏Å‡πà‡∏≠‡∏ô‡∏´‡∏ô‡πâ‡∏≤"
        )
        if uploaded_json is not None:
            try:
                loaded_data = json.load(uploaded_json)
                file_id = f"{uploaded_json.name}_{uploaded_json.size}"
                if st.session_state.get('last_uploaded_file') != file_id:
                    st.session_state['last_uploaded_file'] = file_id
                    st.session_state['input_W18']         = loaded_data.get('W18', 5000000)
                    st.session_state['input_reliability'] = loaded_data.get('reliability', 90)
                    st.session_state['input_So']          = loaded_data.get('So', 0.45)
                    st.session_state['input_P0']          = loaded_data.get('P0', 4.2)
                    st.session_state['input_Pt']          = loaded_data.get('Pt', 2.5)
                    st.session_state['input_CBR']         = loaded_data.get('CBR', 5.0)
                    st.session_state['input_num_layers']  = loaded_data.get('num_layers', 4)
                    st.session_state['input_project_title'] = loaded_data.get('project_title', '‡πÇ‡∏Ñ‡∏£‡∏á‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö‡∏ñ‡∏ô‡∏ô')
                    rs = loaded_data.get('report_settings', {})
                    for key, default in [
                        ('section_number', '4.4'), ('table_number_inputs', '4-8'),
                        ('table_number_materials', '4-9'), ('figure_number', '4-8'),
                        ('section_title', '‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡∏•‡∏≤‡∏î‡∏¢‡∏≤‡∏á (Flexible Pavement)'),
                        ('table_caption_inputs', '‡∏Ñ‡πà‡∏≤‡∏û‡∏≤‡∏£‡∏≤‡∏°‡∏¥‡πÄ‡∏ï‡∏≠‡∏£‡πå‡∏ó‡∏µ‡πà‡πÉ‡∏ä‡πâ‡πÉ‡∏ô‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡∏¢‡∏∑‡∏î‡∏´‡∏¢‡∏∏‡πà‡∏ô'),
                        ('table_caption_materials', '‡∏Ñ‡πà‡∏≤‡∏™‡∏±‡∏°‡∏õ‡∏£‡∏∞‡∏™‡∏¥‡∏ó‡∏ò‡∏¥‡πå‡πÅ‡∏•‡∏∞‡∏Ñ‡πà‡∏≤‡πÇ‡∏°‡∏î‡∏π‡∏•‡∏±‡∏™‡∏Ç‡∏≠‡∏á‡∏ß‡∏±‡∏™‡∏î‡∏∏‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á'),
                        ('figure_caption', '‡∏£‡∏π‡∏õ‡∏ï‡∏±‡∏î‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏ó‡∏µ‡πà‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö'),
                    ]:
                        if key in rs:
                            st.session_state[f'rs_{key}'] = rs[key]
                    layers = loaded_data.get('layers', [])
                    for i, layer in enumerate(layers):
                        st.session_state[f'layer{i+1}_mat']   = layer.get('material', '')
                        st.session_state[f'layer{i+1}_thick'] = layer.get('thickness_cm', 15.0)
                        st.session_state[f'layer{i+1}_m']     = layer.get('drainage_coeff', 1.0)
                    st.success("‚úÖ ‡πÇ‡∏´‡∏•‡∏î‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏™‡∏≥‡πÄ‡∏£‡πá‡∏à!")
                    st.rerun()
            except Exception as e:
                st.error(f"‚ùå ‡πÑ‡∏°‡πà‡∏™‡∏≤‡∏°‡∏≤‡∏£‡∏ñ‡∏≠‡πà‡∏≤‡∏ô‡πÑ‡∏ü‡∏•‡πå‡πÑ‡∏î‡πâ: {e}")

        st.markdown("---")
        st.header("üñºÔ∏è ‡∏ï‡∏±‡πâ‡∏á‡∏Ñ‡πà‡∏≤‡∏£‡∏π‡∏õ‡∏†‡∏≤‡∏û")
        figure_language = st.radio(
            "‡∏†‡∏≤‡∏©‡∏≤‡πÉ‡∏ô‡∏£‡∏π‡∏õ‡∏†‡∏≤‡∏û", options=["English", "‡∏†‡∏≤‡∏©‡∏≤‡πÑ‡∏ó‡∏¢"], index=0,
            help="‡πÄ‡∏•‡∏∑‡∏≠‡∏Å‡∏†‡∏≤‡∏©‡∏≤‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡πÅ‡∏™‡∏î‡∏á‡πÉ‡∏ô‡∏£‡∏π‡∏õ‡∏ï‡∏±‡∏î‡∏Ç‡∏ß‡∏≤‡∏á"
        )

    # ========================================
    # MAIN CONTENT ‚Äî TABS
    # ========================================
    tab_input, tab_layers, tab_results, tab_report = st.tabs([
        "üìù ‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏ô‡∏≥‡πÄ‡∏Ç‡πâ‡∏≤", "üèóÔ∏è ‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á", "üìä ‡∏ú‡∏•‡∏•‡∏±‡∏û‡∏ò‡πå", "üìÑ ‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô"
    ])

    # ========================================
    # TAB 1: DESIGN INPUTS
    # ========================================
    with tab_input:
        st.header("üìù Design Inputs")
        col_t1, col_t2 = st.columns(2)

        with col_t1:
            st.subheader("1Ô∏è‚É£ Traffic & Reliability")
            W18 = st.number_input(
                "Design ESALs (W‚ÇÅ‚Çà)",
                min_value=100000, max_value=250000000,
                value=st.session_state.get('input_W18', 5000000),
                step=100000, format="%d",
                help="‡∏à‡∏≥‡∏ô‡∏ß‡∏ô 18-kip ESAL ‡∏ï‡∏•‡∏≠‡∏î‡∏≠‡∏≤‡∏¢‡∏∏‡∏Å‡∏≤‡∏£‡πÉ‡∏ä‡πâ‡∏á‡∏≤‡∏ô (‡∏™‡∏π‡∏á‡∏™‡∏∏‡∏î 250 ‡∏•‡πâ‡∏≤‡∏ô)",
                key="input_W18"
            )
            esal_million = W18 / 1000000
            st.markdown(
                f'<p style="color:#1E90FF;font-size:20px;font-weight:bold;">'
                f'üí° W‚ÇÅ‚Çà = {esal_million:,.2f} ‡∏•‡πâ‡∏≤‡∏ô ESALs</p>',
                unsafe_allow_html=True)

            reliability_options = list(RELIABILITY_ZR.keys())
            current_reliability = st.session_state.get('input_reliability', 90)
            default_reliability_idx = (reliability_options.index(current_reliability)
                                       if current_reliability in reliability_options
                                       else reliability_options.index(90))
            reliability = st.selectbox(
                "Reliability Level (R)", options=reliability_options,
                index=default_reliability_idx, key="input_reliability"
            )
            Zr = RELIABILITY_ZR[reliability]
            st.info(f"Z·µ£ = {Zr:.3f}")

            So = st.number_input(
                "Overall Standard Deviation (S‚Çí)",
                min_value=0.30, max_value=0.60,
                value=st.session_state.get('input_So', 0.45),
                step=0.01, format="%.2f", key="input_So"
            )

        with col_t2:
            st.subheader("2Ô∏è‚É£ Serviceability")
            col_p1, col_p2 = st.columns(2)
            with col_p1:
                P0 = st.number_input("P‚ÇÄ (Initial)", min_value=3.0, max_value=5.0,
                    value=st.session_state.get('input_P0', 4.2), step=0.1, key="input_P0")
            with col_p2:
                Pt = st.number_input("P‚Çú (Terminal)", min_value=1.5, max_value=3.5,
                    value=st.session_state.get('input_Pt', 2.5), step=0.1, key="input_Pt")
            delta_psi = P0 - Pt
            st.success(f"**ŒîPSI = {delta_psi:.1f}**")

            st.subheader("3Ô∏è‚É£ Subgrade (‡∏î‡∏¥‡∏ô‡πÄ‡∏î‡∏¥‡∏°/‡∏î‡∏¥‡∏ô‡∏ñ‡∏°)")
            CBR = st.number_input("CBR (%)", min_value=1.0, max_value=30.0,
                value=st.session_state.get('input_CBR', 5.0), step=0.5,
                help="‡∏Ñ‡πà‡∏≤ CBR ‡∏Ç‡∏≠‡∏á‡∏î‡∏¥‡∏ô‡πÄ‡∏î‡∏¥‡∏°‡∏´‡∏£‡∏∑‡∏≠‡∏î‡∏¥‡∏ô‡∏ñ‡∏°‡∏Ñ‡∏±‡∏ô‡∏ó‡∏≤‡∏á", key="input_CBR")
            Mr = int(1500 * CBR)
            st.info(f"**M·µ£ = 1,500 √ó CBR = 1,500 √ó {CBR:.1f} = {Mr:,} psi**")

        with st.expander("üìñ ‡∏ï‡∏≤‡∏£‡∏≤‡∏á Drainage Coefficient (AASHTO Table 2.4)"):
            st.markdown("**‡∏Ñ‡πà‡∏≤‡∏™‡∏±‡∏°‡∏õ‡∏£‡∏∞‡∏™‡∏¥‡∏ó‡∏ò‡∏¥‡πå‡∏Å‡∏≤‡∏£‡∏£‡∏∞‡∏ö‡∏≤‡∏¢‡∏ô‡πâ‡∏≥ (m·µ¢) ‚Äî AASHTO 1993 Table 2.4**")
            st.markdown("‡∏Ñ‡πà‡∏≤ default ‡∏Å‡∏£‡∏°‡∏ó‡∏≤‡∏á‡∏´‡∏•‡∏ß‡∏á = **1.0** (‡∏™‡∏†‡∏≤‡∏û‡∏Å‡∏≤‡∏£‡∏£‡∏∞‡∏ö‡∏≤‡∏¢‡∏ô‡πâ‡∏≥‡∏î‡∏µ)")
            drain_data = []
            for quality, info in DRAINAGE_TABLE.items():
                row = {"‡∏Ñ‡∏∏‡∏ì‡∏†‡∏≤‡∏û‡∏Å‡∏≤‡∏£‡∏£‡∏∞‡∏ö‡∏≤‡∏¢‡∏ô‡πâ‡∏≥": f"{quality} ‚Äî {info['description']}"}
                for pct, val in info['values'].items():
                    row[f"‡πÄ‡∏ß‡∏•‡∏≤‡∏≠‡∏¥‡πà‡∏°‡∏ï‡∏±‡∏ß {pct}"] = f"{val:.2f}"
                drain_data.append(row)
            st.table(drain_data)

    # ========================================
    # TAB 2: LAYER CONFIGURATION
    # ========================================
    with tab_layers:
        st.header("üèóÔ∏è Layer Configuration")
        num_layers = st.slider(
            "‡∏à‡∏≥‡∏ô‡∏ß‡∏ô‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á", min_value=2, max_value=6,
            value=st.session_state.get('input_num_layers', 4),
            help="‡πÄ‡∏•‡∏∑‡∏≠‡∏Å‡∏à‡∏≥‡∏ô‡∏ß‡∏ô‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á (2-6 ‡∏ä‡∏±‡πâ‡∏ô)", key="input_num_layers"
        )

        all_materials    = [m for m, p in MATERIALS.items() if p['layer_type'] != 'none']
        surface_materials = [m for m, p in MATERIALS.items() if p['layer_type'] == 'surface']

        layer_data = []
        status_placeholders = {}

        # ===== Global m panel =====
        # ‡∏´‡∏•‡∏±‡∏Å‡∏Å‡∏≤‡∏£: ‡πÉ‡∏ä‡πâ st.session_state flag "apply_global_m" 
        # ‡πÄ‡∏û‡∏∑‡πà‡∏≠‡πÅ‡∏¢‡∏Å logic ‡∏Å‡∏≤‡∏£ apply ‡∏≠‡∏≠‡∏Å‡∏à‡∏≤‡∏Å widget rendering
        # ‡πÑ‡∏°‡πà‡πÉ‡∏ä‡πâ rerun() ‡πÄ‡∏û‡∏∑‡πà‡∏≠‡∏´‡∏•‡∏µ‡∏Å‡πÄ‡∏•‡∏µ‡πà‡∏¢‡∏á‡∏Å‡∏≤‡∏£ reset widget values
        with st.container():
            st.markdown(
                '<div style="background:#EFF6FF;border:1.5px solid #3B82F6;border-radius:8px;'
                'padding:10px 16px 6px 16px;margin-bottom:12px;">'
                '<b style="color:#1D4ED8;">üîß ‡∏Å‡∏≥‡∏´‡∏ô‡∏î‡∏Ñ‡πà‡∏≤ m (Drainage Coefficient) ‡πÄ‡∏î‡∏µ‡∏¢‡∏ß‡∏Å‡∏±‡∏ô‡∏ó‡∏∏‡∏Å‡∏ä‡∏±‡πâ‡∏ô</b>'
                '</div>', unsafe_allow_html=True
            )
            gcol1, gcol2 = st.columns([2, 1])
            with gcol1:
                global_m = st.number_input(
                    "‡∏Ñ‡πà‡∏≤ m ‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡∏ó‡∏∏‡∏Å‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á",
                    min_value=0.40, max_value=1.50,
                    value=float(st.session_state.get('global_m_value', 1.00)),
                    step=0.05, format="%.2f",
                    help="‡∏Å‡∏£‡∏≠‡∏Å‡∏Ñ‡πà‡∏≤‡πÅ‡∏•‡πâ‡∏ß‡∏Å‡∏î '‡πÉ‡∏ä‡πâ‡πÄ‡∏´‡∏°‡∏∑‡∏≠‡∏ô‡∏Å‡∏±‡∏ô‡∏ó‡∏∏‡∏Å‡∏ä‡∏±‡πâ‡∏ô' ‚Äî ‡∏¢‡∏±‡∏á‡πÅ‡∏Å‡πâ‡πÑ‡∏Ç‡πÅ‡∏ï‡πà‡∏•‡∏∞‡∏ä‡∏±‡πâ‡∏ô‡πÑ‡∏î‡πâ‡∏†‡∏≤‡∏¢‡∏´‡∏•‡∏±‡∏á"
                )
                # ‡∏ö‡∏±‡∏ô‡∏ó‡∏∂‡∏Å‡∏Ñ‡πà‡∏≤‡∏•‡πà‡∏≤‡∏™‡∏∏‡∏î‡∏ó‡∏∏‡∏Å rerun (‡πÑ‡∏°‡πà‡πÉ‡∏ä‡πâ key ‡πÄ‡∏û‡∏∑‡πà‡∏≠‡∏´‡∏•‡∏µ‡∏Å‡πÄ‡∏•‡∏µ‡πà‡∏¢‡∏á widget state conflict)
                st.session_state['global_m_value'] = global_m
            with gcol2:
                st.markdown("<div style='height:28px'></div>", unsafe_allow_html=True)
                if st.button("‚úÖ ‡πÉ‡∏ä‡πâ‡πÄ‡∏´‡∏°‡∏∑‡∏≠‡∏ô‡∏Å‡∏±‡∏ô‡∏ó‡∏∏‡∏Å‡∏ä‡∏±‡πâ‡∏ô", type="primary", use_container_width=True):
                    nl = st.session_state.get('input_num_layers', 4)
                    for idx in range(1, nl + 1):
                        st.session_state[f'layer{idx}_m'] = global_m
                    st.toast(f"‚úÖ ‡∏ï‡∏±‡πâ‡∏á‡∏Ñ‡πà‡∏≤ m = {global_m:.2f} ‡πÉ‡∏´‡πâ‡∏ó‡∏∏‡∏Å‡∏ä‡∏±‡πâ‡∏ô‡πÅ‡∏•‡πâ‡∏ß", icon="‚úÖ")

        # ===== ‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏µ‡πà 1: ‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á =====
        st.subheader("üî∂ ‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏µ‡πà 1: ‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á (Surface)")

        layer1_mat_default = st.session_state.get('layer1_mat', surface_materials[0])
        layer1_mat_idx = (surface_materials.index(layer1_mat_default)
                         if layer1_mat_default in surface_materials else 0)
        layer1_mat = st.selectbox(
            "‡πÄ‡∏•‡∏∑‡∏≠‡∏Å‡∏ß‡∏±‡∏™‡∏î‡∏∏", options=surface_materials,
            index=layer1_mat_idx, key="layer1_mat"
        )

        mat_props_1 = MATERIALS[layer1_mat]
        default_a1  = mat_props_1['layer_coeff']
        default_m1  = mat_props_1['drainage_coeff']

        # ===== AC Sublayer ‚Äî UI ‡πÄ‡∏£‡∏µ‡∏¢‡∏ö‡∏á‡πà‡∏≤‡∏¢ =====
        use_sublayers = st.checkbox(
            "üìê ‡πÅ‡∏ö‡πà‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏¢‡πà‡∏≠‡∏¢‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á AC (Wearing / Binder / Base Course)",
            value=st.session_state.get('use_ac_sublayers', False),
            help="‡πÅ‡∏ö‡πà‡∏á‡∏ä‡∏±‡πâ‡∏ô AC ‡∏≠‡∏≠‡∏Å‡πÄ‡∏õ‡πá‡∏ô 3 ‡∏ä‡∏±‡πâ‡∏ô‡∏¢‡πà‡∏≠‡∏¢ ‡∏ï‡∏≤‡∏°‡∏°‡∏≤‡∏ï‡∏£‡∏ê‡∏≤‡∏ô‡∏Å‡∏£‡∏°‡∏ó‡∏≤‡∏á‡∏´‡∏•‡∏ß‡∏á",
            key="use_ac_sublayers"
        )

        if use_sublayers:
            st.info("üìã ‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡∏°‡∏≤‡∏ï‡∏£‡∏ê‡∏≤‡∏ô ‡∏ó‡∏•.: Wearing 40-70 ‡∏°‡∏°. / Binder 40-80 ‡∏°‡∏°. / Base 70-100 ‡∏°‡∏°.")

            # ‡∏ï‡∏≤‡∏£‡∏≤‡∏á 3 ‡∏Ñ‡∏≠‡∏•‡∏±‡∏°‡∏ô‡πå ‚Äî ‡πÅ‡∏ï‡πà‡∏•‡∏∞‡∏ä‡∏±‡πâ‡∏ô‡∏°‡∏µ‡πÅ‡∏Ñ‡πà 1 input
            sub_labels = list(DOH_THICKNESS_STANDARDS.keys())
            sub_keys   = ['wearing', 'binder', 'base']
            sub_defaults = [5.0, 7.0, 10.0]
            sub_results  = {}

            cols = st.columns(3)
            for idx, (col, key, label, defaults_mm, default_cm) in enumerate(zip(
                cols, sub_keys, sub_labels,
                list(DOH_THICKNESS_STANDARDS.values()),
                sub_defaults
            )):
                with col:
                    st.markdown(f"**{label}**")
                    std_options = ["‡∏Å‡∏≥‡∏´‡∏ô‡∏î‡πÄ‡∏≠‡∏á"] + (["‡πÑ‡∏°‡πà‡πÉ‡∏ä‡πâ"] if key == 'base' else []) + [f"{t} ‡∏°‡∏°." for t in defaults_mm if t > 0]
                    std_sel = st.selectbox(
                        "‡∏°‡∏≤‡∏ï‡∏£‡∏ê‡∏≤‡∏ô ‡∏ó‡∏•.", std_options, index=0, key=f"{key}_std_select"
                    )
                    if std_sel == "‡πÑ‡∏°‡πà‡πÉ‡∏ä‡πâ":
                        sub_results[key] = 0.0
                        st.metric("‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤", "0.0 cm")
                    elif std_sel != "‡∏Å‡∏≥‡∏´‡∏ô‡∏î‡πÄ‡∏≠‡∏á":
                        sub_results[key] = int(std_sel.replace(" ‡∏°‡∏°.", "")) / 10
                        st.metric("‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤", f"{sub_results[key]:.1f} cm")
                    else:
                        max_val = 15.0 if key != 'base' else 15.0
                        sub_results[key] = st.number_input(
                            "‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤ (cm)", 0.0, max_val,
                            value=st.session_state.get(f'{key}_thick_val', default_cm),
                            step=0.5, key=f"{key}_thick"
                        )

            layer1_thick = sum(sub_results.values())
            st.markdown(
                f'<p style="color:#1E90FF;font-size:18px;font-weight:bold;">'
                f'üìè ‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡∏£‡∏ß‡∏° AC = '
                f'{sub_results["wearing"]:.1f} + {sub_results["binder"]:.1f} + {sub_results["base"]:.1f}'
                f' = {layer1_thick:.1f} cm</p>',
                unsafe_allow_html=True)

            st.session_state['ac_sublayers'] = {
                'wearing': sub_results['wearing'],
                'binder':  sub_results['binder'],
                'base':    sub_results['base'],
                'total':   layer1_thick
            }

            col_am1, col_am2 = st.columns(2)
            with col_am1:
                st.markdown(f"a‚ÇÅ <span style='color:#1E90FF;font-size:12px;'>(default={default_a1:.2f})</span>",
                            unsafe_allow_html=True)
                layer1_a = st.number_input("a1", 0.10, 0.50,
                    value=st.session_state.get('layer1_a', default_a1), step=0.01,
                    key="layer1_a", label_visibility="collapsed")
            with col_am2:
                st.markdown(f"m‚ÇÅ <span style='color:#1E90FF;font-size:12px;'>(default={default_m1:.2f})</span>",
                            unsafe_allow_html=True)
                layer1_m = st.number_input("m1", 0.5, 1.5,
                    value=st.session_state.get('layer1_m', default_m1), step=0.05,
                    key="layer1_m", label_visibility="collapsed")
        else:
            st.session_state['ac_sublayers'] = None
            col_a, col_b, col_c = st.columns(3)
            with col_a:
                layer1_thick = st.number_input("‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤ (cm)", 1.0, 30.0,
                    value=st.session_state.get('layer1_thick', 5.0), step=1.0, key="layer1_thick")
            with col_b:
                st.markdown(f"a‚ÇÅ <span style='color:#1E90FF;font-size:12px;'>(default={default_a1:.2f})</span>",
                            unsafe_allow_html=True)
                layer1_a = st.number_input("a1", 0.10, 0.50,
                    value=st.session_state.get('layer1_a', default_a1), step=0.01,
                    key="layer1_a", label_visibility="collapsed")
            with col_c:
                st.markdown(f"m‚ÇÅ <span style='color:#1E90FF;font-size:12px;'>(default={default_m1:.2f})</span>",
                            unsafe_allow_html=True)
                layer1_m = st.number_input("m1", 0.5, 1.5,
                    value=st.session_state.get('layer1_m', default_m1), step=0.05,
                    key="layer1_m", label_visibility="collapsed")

        st.markdown(f'<p style="color:#1E90FF;font-size:14px;">E = {mat_props_1["mr_mpa"]:,} MPa</p>',
                    unsafe_allow_html=True)
        status_placeholders[1] = st.empty()

        layer_data.append({
            'material': layer1_mat,
            'thickness_cm': layer1_thick,
            'layer_coeff': layer1_a,
            'drainage_coeff': layer1_m
        })

        # ===== ‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏µ‡πà 2-6 =====
        default_materials = [
            "‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏´‡∏¥‡∏ô‡∏Ñ‡∏•‡∏∏‡∏Å‡∏õ‡∏£‡∏±‡∏ö‡∏õ‡∏£‡∏∏‡∏á‡∏Ñ‡∏∏‡∏ì‡∏†‡∏≤‡∏û‡∏î‡πâ‡∏ß‡∏¢‡∏õ‡∏π‡∏ô‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå (Cement Treated Base)",
            "‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏°‡∏ß‡∏•‡∏£‡∏ß‡∏° CBR 25%",
            "‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Ñ‡∏±‡∏î‡πÄ‡∏•‡∏∑‡∏≠‡∏Å ‡∏Å",
            "‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Ñ‡∏±‡∏î‡πÄ‡∏•‡∏∑‡∏≠‡∏Å ‡∏Å",
            "‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Ñ‡∏±‡∏î‡πÄ‡∏•‡∏∑‡∏≠‡∏Å ‡∏Å"
        ]
        default_thickness = [15.0, 15.0, 30.0, 30.0, 30.0]
        layer_icons = ['üî∑', 'üî∂', 'üü¢', 'üü°', 'üî¥']

        for i in range(2, num_layers + 1):
            st.markdown("---")
            st.subheader(f"{layer_icons[i-2]} ‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏µ‡πà {i}")

            layer_i_mat_default = st.session_state.get(f'layer{i}_mat', default_materials[i-2])
            if layer_i_mat_default in all_materials:
                default_idx = all_materials.index(layer_i_mat_default)
            else:
                default_idx = (all_materials.index(default_materials[i-2])
                              if default_materials[i-2] in all_materials else 0)

            layer_mat = st.selectbox(
                f"‡πÄ‡∏•‡∏∑‡∏≠‡∏Å‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏µ‡πà {i}", options=all_materials,
                index=min(default_idx, len(all_materials)-1), key=f"layer{i}_mat"
            )

            mat_props = MATERIALS[layer_mat]
            default_a = mat_props['layer_coeff']
            default_m = mat_props['drainage_coeff']

            prev_mat_key = f'layer{i}_prev_mat'
            if prev_mat_key not in st.session_state:
                st.session_state[prev_mat_key] = layer_mat
            if st.session_state[prev_mat_key] != layer_mat:
                st.session_state[f'layer{i}_a'] = default_a
                st.session_state[f'layer{i}_m'] = default_m
                st.session_state[prev_mat_key] = layer_mat

            col_c, col_d, col_e = st.columns(3)
            with col_c:
                layer_thick = st.number_input("‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤ (cm)", 1.0, 150.0,
                    value=st.session_state.get(f'layer{i}_thick', default_thickness[i-2]),
                    step=5.0, key=f"layer{i}_thick")
            with col_d:
                st.markdown(f"a{i} <span style='color:#1E90FF;font-size:12px;'>(default={default_a:.2f})</span>",
                            unsafe_allow_html=True)
                layer_a = st.number_input(f"a{i}", 0.01, 0.50,
                    value=st.session_state.get(f'layer{i}_a', default_a), step=0.01,
                    key=f"layer{i}_a", label_visibility="collapsed")
            with col_e:
                st.markdown(f"m{i} <span style='color:#1E90FF;font-size:12px;'>(default={default_m:.2f})</span>",
                            unsafe_allow_html=True)
                layer_m = st.number_input(f"m{i}", 0.5, 1.5,
                    value=st.session_state.get(f'layer{i}_m', default_m), step=0.05,
                    key=f"layer{i}_m", label_visibility="collapsed")

            st.markdown(f'<p style="color:#1E90FF;font-size:14px;">E = {mat_props["mr_mpa"]:,} MPa</p>',
                        unsafe_allow_html=True)
            status_placeholders[i] = st.empty()

            layer_data.append({
                'material': layer_mat,
                'thickness_cm': layer_thick,
                'layer_coeff': layer_a,
                'drainage_coeff': layer_m
            })

        # ===== ‡∏ê‡∏≤‡∏ô‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏ß‡∏±‡∏™‡∏î‡∏∏ (‡∏¢‡πâ‡∏≤‡∏¢‡∏°‡∏≤‡∏≠‡∏¢‡∏π‡πà‡πÉ‡∏ô Tab 2) =====
        st.markdown("---")
        with st.expander("üìö ‡∏ê‡∏≤‡∏ô‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏ß‡∏±‡∏™‡∏î‡∏∏ (‡∏°‡∏≤‡∏ï‡∏£‡∏ê‡∏≤‡∏ô ‡∏ó‡∏•.) ‚Äî ‡∏Ñ‡∏•‡∏¥‡∏Å‡πÄ‡∏û‡∏∑‡πà‡∏≠‡∏î‡∏π‡∏Ñ‡πà‡∏≤ ‡∏™‡∏õ‡∏™. ‡∏ó‡∏±‡πâ‡∏á‡∏´‡∏°‡∏î"):
            for mat_name, props in MATERIALS.items():
                if props['layer_coeff'] > 0:
                    st.markdown(f"**{mat_name}**")
                    st.markdown(f"- a = {props['layer_coeff']}, m = {props['drainage_coeff']}")
                    st.markdown(f"- MR = {props['mr_psi']:,} psi ({props['mr_mpa']:,} MPa)")
                    st.markdown("---")

    # ========================================
    # CALCULATION (‡∏ó‡∏≥‡∏á‡∏≤‡∏ô‡∏ó‡∏∏‡∏Å rerun ‡∏≠‡∏±‡∏ï‡πÇ‡∏ô‡∏°‡∏±‡∏ï‡∏¥)
    # ========================================
    inputs = {
        'W18': W18, 'reliability': reliability, 'Zr': Zr, 'So': So,
        'P0': P0, 'Pt': Pt, 'delta_psi': delta_psi, 'CBR': CBR, 'Mr': Mr
    }
    ac_sublayers = st.session_state.get('ac_sublayers', None)
    calc_results  = calculate_layer_thicknesses(W18, Zr, So, delta_psi, Mr, layer_data, ac_sublayers)
    design_check  = check_design(calc_results['total_sn_required'], calc_results['total_sn_provided'])

    # ===== QUICK SUMMARY BANNER ‚Äî ‡πÅ‡∏™‡∏î‡∏á‡πÄ‡∏´‡∏ô‡∏∑‡∏≠ tabs ‡∏ó‡∏∏‡∏Å‡∏Ñ‡∏£‡∏±‡πâ‡∏á =====
    if design_check['passed']:
        st.markdown(
            f"""<div style="background:#d4edda;border:2px solid #28a745;border-radius:10px;
            padding:14px 24px;text-align:center;margin:10px 0 18px 0;">
            <h3 style="color:#28a745;margin:0;">‚úÖ PASS ‚Äî ‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö‡∏ú‡πà‡∏≤‡∏ô‡πÄ‡∏Å‡∏ì‡∏ë‡πå</h3>
            <p style="font-size:17px;margin:6px 0;">
            SN<sub>provided</sub> = <b>{calc_results['total_sn_provided']:.2f}</b> &nbsp;‚â•&nbsp;
            SN<sub>required</sub> = <b>{calc_results['total_sn_required']:.2f}</b>
            &nbsp;|&nbsp; Safety Margin = <b>{design_check['safety_margin']:.2f}</b>
            &nbsp;|&nbsp; ‡πÇ‡∏Ñ‡∏£‡∏á‡∏Å‡∏≤‡∏£: <b>{project_title}</b>
            </p></div>""", unsafe_allow_html=True)
    else:
        st.markdown(
            f"""<div style="background:#f8d7da;border:2px solid #dc3545;border-radius:10px;
            padding:14px 24px;text-align:center;margin:10px 0 18px 0;">
            <h3 style="color:#dc3545;margin:0;">‚ùå FAIL ‚Äî ‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö‡πÑ‡∏°‡πà‡∏ú‡πà‡∏≤‡∏ô</h3>
            <p style="font-size:17px;margin:6px 0;">
            SN<sub>provided</sub> = <b>{calc_results['total_sn_provided']:.2f}</b> &nbsp;&lt;&nbsp;
            SN<sub>required</sub> = <b>{calc_results['total_sn_required']:.2f}</b>
            &nbsp;|&nbsp; ‡∏Ç‡∏≤‡∏î‡∏≠‡∏µ‡∏Å = <b>{abs(design_check['safety_margin']):.2f}</b>
            &nbsp;|&nbsp; ‡πÇ‡∏Ñ‡∏£‡∏á‡∏Å‡∏≤‡∏£: <b>{project_title}</b>
            </p></div>""", unsafe_allow_html=True)

    # Fill status placeholders in Tab 2
    for layer in calc_results['layers']:
        layer_no = layer['layer_no']
        if layer_no in status_placeholders:
            with status_placeholders[layer_no]:
                if layer['is_ok']:
                    st.success(f"‚úÖ ‡∏ú‡πà‡∏≤‡∏ô (‡∏ï‡πâ‡∏≠‡∏á‡∏Å‡∏≤‡∏£ ‚â• {layer['min_thickness_cm']:.1f} cm)")
                else:
                    shortage = layer['min_thickness_cm'] - layer['design_thickness_cm']
                    st.error(f"‚ùå ‡πÑ‡∏°‡πà‡∏ú‡πà‡∏≤‡∏ô (‡∏ï‡πâ‡∏≠‡∏á‡πÄ‡∏û‡∏¥‡πà‡∏°‡∏≠‡∏µ‡∏Å {shortage:.1f} cm)")

    # ========================================
    # TAB 3: RESULTS
    # ========================================
    with tab_results:
        # ===== Warnings =====
        warnings = calc_results.get('warnings', [])
        if warnings:
            for w in warnings:
                st.warning(w)

        # ===== W18 Supported metrics =====
        w18_supported = calculate_w18_supported(
            calc_results['total_sn_provided'], Zr, So, delta_psi, Mr)
        w18_supported_million = w18_supported / 1_000_000
        w18_diff_percent = ((w18_supported - W18) / W18) * 100

        mc1, mc2, mc3, mc4 = st.columns(4)
        with mc1:
            st.metric("W‚ÇÅ‚Çà ‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö", f"{W18/1e6:,.2f} ‡∏•‡πâ‡∏≤‡∏ô")
        with mc2:
            st.metric("W‚ÇÅ‚Çà ‡∏£‡∏≠‡∏á‡∏£‡∏±‡∏ö‡πÑ‡∏î‡πâ", f"{w18_supported_million:,.2f} ‡∏•‡πâ‡∏≤‡∏ô",
                      delta=f"{w18_diff_percent:+.1f}%",
                      delta_color="normal" if w18_diff_percent >= 0 else "inverse")
        with mc3:
            st.metric("SN Required", f"{calc_results['total_sn_required']:.2f}")
        with mc4:
            st.metric("SN Provided", f"{calc_results['total_sn_provided']:.2f}",
                      delta=f"{design_check['safety_margin']:+.2f}")

        st.markdown("---")

        # ===== STEP-BY-STEP CALCULATION =====
        st.subheader("üî¢ ‡∏Ç‡∏±‡πâ‡∏ô‡∏ï‡∏≠‡∏ô‡∏Å‡∏≤‡∏£‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡πÅ‡∏ï‡πà‡∏•‡∏∞‡∏ä‡∏±‡πâ‡∏ô")

        for layer in calc_results['layers']:
            layer_status = "‚úÖ" if layer['is_ok'] else "‚ùå"
            st.markdown(f"### {layer_status} ‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏µ‡πà {layer['layer_no']}: {short_material_name(layer['material'])}")

            # AC sublayer info
            layer_ac_sub = layer.get('ac_sublayers', None)
            if layer_ac_sub is not None and layer['layer_no'] == 1:
                st.info(
                    f"**üìê ‡πÅ‡∏ö‡πà‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏¢‡πà‡∏≠‡∏¢ AC:** "
                    f"‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á (Wearing) = {layer_ac_sub['wearing']:.1f} cm  |  "
                    f"‡∏£‡∏≠‡∏á‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á (Binder) = {layer_ac_sub['binder']:.1f} cm  |  "
                    f"‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á (Base) = {layer_ac_sub['base']:.1f} cm  |  "
                    f"**‡∏£‡∏ß‡∏° = {layer_ac_sub['total']:.1f} cm**"
                )

            col_a, col_b = st.columns([1, 1])
            with col_a:
                st.markdown("**‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏ß‡∏±‡∏™‡∏î‡∏∏:**")
                st.markdown(f"- E (MPa) = **{layer['mr_mpa']:,}**")
                st.markdown(f"- M·µ£ (psi) = **{layer['mr_psi']:,}**")
                st.markdown(f"- Layer Coefficient (a{layer['layer_no']}) = **{layer['a_i']:.2f}**")
                st.markdown(f"- Drain Coefficient (m{layer['layer_no']}) = **{layer['m_i']:.2f}**")

            with col_b:
                st.markdown("**‡∏à‡∏≤‡∏Å‡∏™‡∏°‡∏Å‡∏≤‡∏£ AASHTO:**")
                sn_at_layer = layer['sn_required_at_layer']
                st.latex(f"SN_{{{layer['layer_no']}}} = {sn_at_layer:.2f}")

            st.markdown("**‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á:**")
            if layer['layer_no'] == 1:
                st.latex(
                    f"D_{{1}} \\geq \\frac{{SN_{{1}}}}{{a_{{1}} \\times m_{{1}}}} = "
                    f"\\frac{{{sn_at_layer:.2f}}}{{{layer['a_i']:.2f} \\times {layer['m_i']:.2f}}} = "
                    f"{layer['min_thickness_inch']:.2f} \\text{{ ‡∏ô‡∏¥‡πâ‡∏ß}} = {layer['min_thickness_cm']:.1f} \\text{{ ‡∏ã‡∏°.}}")
            else:
                prev_sn = calc_results['layers'][layer['layer_no']-2]['cumulative_sn']
                st.latex(
                    f"D_{{{layer['layer_no']}}} \\geq "
                    f"\\frac{{SN_{{{layer['layer_no']}}} - SN_{{prev}}}}"
                    f"{{a_{{{layer['layer_no']}}} \\times m_{{{layer['layer_no']}}}}} = "
                    f"\\frac{{{sn_at_layer:.2f} - {prev_sn:.2f}}}"
                    f"{{{layer['a_i']:.2f} \\times {layer['m_i']:.2f}}} = "
                    f"{layer['min_thickness_inch']:.2f} \\text{{ ‡∏ô‡∏¥‡πâ‡∏ß}} = {layer['min_thickness_cm']:.1f} \\text{{ ‡∏ã‡∏°.}}")

            result_cols = st.columns(4)
            with result_cols[0]:
                st.metric("‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡∏Ç‡∏±‡πâ‡∏ô‡∏ï‡πà‡∏≥", f"{layer['min_thickness_cm']:.1f} cm")
            with result_cols[1]:
                st.metric("‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡∏ó‡∏µ‡πà‡πÄ‡∏•‡∏∑‡∏≠‡∏Å", f"{layer['design_thickness_cm']:.0f} cm",
                         delta=f"{layer['design_thickness_cm'] - layer['min_thickness_cm']:.1f} cm")
            with result_cols[2]:
                st.metric("SN contribution", f"{layer['sn_contribution']:.3f}")
            with result_cols[3]:
                st.metric("Cumulative SN", f"{layer['cumulative_sn']:.2f}")

            if layer['is_ok']:
                st.success(f"‚úÖ **OK** ‚Äî ‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡πÄ‡∏û‡∏µ‡∏¢‡∏á‡∏û‡∏≠ ({layer['design_thickness_cm']:.0f} ‚â• {layer['min_thickness_cm']:.1f} cm)")
            else:
                st.error(f"‚ùå **NG** ‚Äî ‡∏ï‡πâ‡∏≠‡∏á‡πÄ‡∏û‡∏¥‡πà‡∏°‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡∏≠‡∏µ‡∏Å {layer['min_thickness_cm'] - layer['design_thickness_cm']:.1f} cm")
            st.markdown("---")

        # ===== SN Summary Table =====
        with st.expander("üìã ‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏™‡∏£‡∏∏‡∏õ‡∏Å‡∏≤‡∏£‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì SN"):
            table_data = []
            for layer in calc_results['layers']:
                table_data.append({
                    '‡∏ä‡∏±‡πâ‡∏ô': layer['layer_no'],
                    '‡∏ß‡∏±‡∏™‡∏î‡∏∏': layer['short_name'],
                    'a·µ¢': layer['a_i'],
                    'D·µ¢ (cm)': layer['design_thickness_cm'],
                    'D·µ¢ (in)': layer['design_thickness_inch'],
                    'm·µ¢': layer['m_i'],
                    'E (MPa)': layer['mr_mpa'],
                    'SN contrib.': layer['sn_contribution'],
                    'SN cumul.': layer['cumulative_sn']
                })
            st.table(table_data)
            st.markdown(
                f"**‡∏™‡∏π‡∏ï‡∏£:** $SN = \\sum_{{i=1}}^{{n}} a_i \\times D_i \\times m_i$  |  "
                f"**SN_provided = {calc_results['total_sn_provided']:.2f}**  |  "
                f"**SN_required = {calc_results['total_sn_required']:.2f}**")

        # ===== PAVEMENT SECTION FIGURE =====
        st.subheader("üìê ‡∏†‡∏≤‡∏û‡∏ï‡∏±‡∏î‡∏Ç‡∏ß‡∏≤‡∏á‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ñ‡∏ô‡∏ô")
        fig_lang = 'th' if figure_language == "‡∏†‡∏≤‡∏©‡∏≤‡πÑ‡∏ó‡∏¢" else 'en'
        fig = plot_pavement_section(calc_results['layers'], Mr, CBR, lang=fig_lang)
        st.pyplot(fig)
        plt.close(fig)

        # ===== SENSITIVITY ANALYSIS =====
        st.markdown("---")
        st.subheader("üìà Sensitivity Analysis")
        st.caption("üî¥ Red dot = current design value  |  ‡∏Å‡∏£‡∏≤‡∏ü‡πÅ‡∏™‡∏î‡∏á‡∏ú‡∏•‡∏Å‡∏£‡∏∞‡∏ó‡∏ö‡∏Ç‡∏≠‡∏á‡∏ï‡∏±‡∏ß‡πÅ‡∏õ‡∏£‡∏´‡∏•‡∏±‡∏Å‡∏ï‡πà‡∏≠ SN_required")

        sens_col1, sens_col2 = st.columns(2)
        with sens_col1:
            fig_cbr = plot_sensitivity_cbr(W18, Zr, So, delta_psi, CBR)
            st.pyplot(fig_cbr)
            plt.close(fig_cbr)
        with sens_col2:
            fig_w18 = plot_sensitivity_w18(Zr, So, delta_psi, Mr, W18)
            st.pyplot(fig_w18)
            plt.close(fig_w18)

    # ========================================
    # TAB 4: REPORT & EXPORT
    # ========================================
    with tab_report:
        st.header("üìÑ ‡∏™‡πà‡∏á‡∏≠‡∏≠‡∏Å‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô")

        # ===== Report Settings =====
        st.markdown("### üìù ‡∏ï‡∏±‡πâ‡∏á‡∏Ñ‡πà‡∏≤‡∏´‡∏°‡∏≤‡∏¢‡πÄ‡∏•‡∏Ç‡∏´‡∏±‡∏ß‡∏Ç‡πâ‡∏≠‡πÅ‡∏•‡∏∞‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô Word")

        col_num1, col_num2, col_num3 = st.columns(3)
        with col_num1:
            rs_section_number = st.text_input(
                "‡πÄ‡∏•‡∏Ç‡∏´‡∏±‡∏ß‡∏Ç‡πâ‡∏≠",
                value=st.session_state.get('rs_section_number', '4.4'),
                key='rs_section_number'
            )
        with col_num2:
            rs_table_number_inputs = st.text_input(
                "‡πÄ‡∏•‡∏Ç‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏û‡∏≤‡∏£‡∏≤‡∏°‡∏¥‡πÄ‡∏ï‡∏≠‡∏£‡πå",
                value=st.session_state.get('rs_table_number_inputs', '4-8'),
                key='rs_table_number_inputs'
            )
        with col_num3:
            rs_table_number_materials = st.text_input(
                "‡πÄ‡∏•‡∏Ç‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏ß‡∏±‡∏™‡∏î‡∏∏",
                value=st.session_state.get('rs_table_number_materials', '4-9'),
                key='rs_table_number_materials'
            )

        col_num4, col_num5 = st.columns([1, 2])
        with col_num4:
            rs_figure_number = st.text_input(
                "‡πÄ‡∏•‡∏Ç‡∏£‡∏π‡∏õ",
                value=st.session_state.get('rs_figure_number', '4-8'),
                key='rs_figure_number'
            )
        with col_num5:
            rs_section_title = st.text_input(
                "‡∏ä‡∏∑‡πà‡∏≠‡∏´‡∏±‡∏ß‡∏Ç‡πâ‡∏≠",
                value=st.session_state.get('rs_section_title',
                      '‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡∏•‡∏≤‡∏î‡∏¢‡∏≤‡∏á (Flexible Pavement)'),
                key='rs_section_title'
            )

        col_cap1, col_cap2 = st.columns(2)
        with col_cap1:
            rs_table_caption_inputs = st.text_input(
                "‡∏Ñ‡∏≥‡∏ö‡∏£‡∏£‡∏¢‡∏≤‡∏¢‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏û‡∏≤‡∏£‡∏≤‡∏°‡∏¥‡πÄ‡∏ï‡∏≠‡∏£‡πå",
                value=st.session_state.get('rs_table_caption_inputs',
                      '‡∏Ñ‡πà‡∏≤‡∏û‡∏≤‡∏£‡∏≤‡∏°‡∏¥‡πÄ‡∏ï‡∏≠‡∏£‡πå‡∏ó‡∏µ‡πà‡πÉ‡∏ä‡πâ‡πÉ‡∏ô‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡∏¢‡∏∑‡∏î‡∏´‡∏¢‡∏∏‡πà‡∏ô'),
                key='rs_table_caption_inputs'
            )
        with col_cap2:
            rs_table_caption_materials = st.text_input(
                "‡∏Ñ‡∏≥‡∏ö‡∏£‡∏£‡∏¢‡∏≤‡∏¢‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏ß‡∏±‡∏™‡∏î‡∏∏",
                value=st.session_state.get('rs_table_caption_materials',
                      '‡∏Ñ‡πà‡∏≤‡∏™‡∏±‡∏°‡∏õ‡∏£‡∏∞‡∏™‡∏¥‡∏ó‡∏ò‡∏¥‡πå‡πÅ‡∏•‡∏∞‡∏Ñ‡πà‡∏≤‡πÇ‡∏°‡∏î‡∏π‡∏•‡∏±‡∏™‡∏Ç‡∏≠‡∏á‡∏ß‡∏±‡∏™‡∏î‡∏∏‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á'),
                key='rs_table_caption_materials'
            )

        rs_figure_caption = st.text_input(
            "‡∏Ñ‡∏≥‡∏ö‡∏£‡∏£‡∏¢‡∏≤‡∏¢‡∏£‡∏π‡∏õ",
            value=st.session_state.get('rs_figure_caption', '‡∏£‡∏π‡∏õ‡∏ï‡∏±‡∏î‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏ó‡∏µ‡πà‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö'),
            key='rs_figure_caption'
        )

        report_settings = {
            'section_number':          rs_section_number,
            'table_number_inputs':     rs_table_number_inputs,
            'table_number_materials':  rs_table_number_materials,
            'figure_number':           rs_figure_number,
            'section_title':           rs_section_title,
            'table_caption_inputs':    rs_table_caption_inputs,
            'table_caption_materials': rs_table_caption_materials,
            'figure_caption':          rs_figure_caption,
        }

        st.markdown("---")

        # ===== Preview ‡∏ö‡∏ó‡πÄ‡∏Å‡∏£‡∏¥‡πà‡∏ô‡∏ô‡∏≥ =====
        st.markdown("### üëÅÔ∏è Preview ‡∏ö‡∏ó‡πÄ‡∏Å‡∏£‡∏¥‡πà‡∏ô‡∏ô‡∏≥")

        total_thick_prev = sum(l['design_thickness_cm'] for l in calc_results['layers'])
        num_layers_prev  = len(calc_results['layers'])
        passed_prev      = '‡∏ú‡πà‡∏≤‡∏ô‡πÄ‡∏Å‡∏ì‡∏ë‡πå' if design_check['passed'] else '‡πÑ‡∏°‡πà‡∏ú‡πà‡∏≤‡∏ô‡πÄ‡∏Å‡∏ì‡∏ë‡πå'

        def hl_purple(val):
            return f'<span style="background-color:#D8B4FE;padding:1px 4px;border-radius:3px;font-weight:bold;">{val}</span>'
        def hl_yellow(val):
            return f'<span style="background-color:#FDE68A;padding:1px 4px;border-radius:3px;font-weight:bold;">{val}</span>'

        intro_html = f"""
        <div style="background:#f9f9f9;padding:15px 20px;border-radius:8px;border:1px solid #ddd;
                    font-family:'TH SarabunPSK',Sarabun,sans-serif;font-size:16px;line-height:1.9;">
            <p style="font-weight:bold;margin-bottom:5px;">
                {hl_yellow(rs_section_number)}&nbsp;&nbsp;{hl_yellow(rs_section_title)}
            </p>
            <p style="text-indent:40px;text-align:justify;">
                ‡∏ñ‡∏ô‡∏ô‡∏•‡∏≤‡∏î‡∏¢‡∏≤‡∏á‡∏ã‡∏∂‡πà‡∏á‡∏õ‡∏£‡∏∞‡∏Å‡∏≠‡∏ö‡∏î‡πâ‡∏ß‡∏¢‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏á‡∏≤‡∏ô‡∏ó‡∏≤‡∏á‡∏´‡∏•‡∏≤‡∏¢‡∏ä‡∏ô‡∏¥‡∏î ‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ñ‡∏ô‡∏ô‡πÅ‡∏ö‡∏ö‡∏¢‡∏∑‡∏î‡∏´‡∏¢‡∏∏‡πà‡∏ô (Flexible Pavement)
                ‡πÉ‡∏ä‡πâ‡∏ß‡∏¥‡∏ò‡∏µ AASHTO 1993 Guide for Design of Pavement Structures ‡πÇ‡∏î‡∏¢‡∏û‡∏¥‡∏à‡∏≤‡∏£‡∏ì‡∏≤‡∏õ‡∏±‡∏à‡∏à‡∏±‡∏¢‡∏î‡πâ‡∏≤‡∏ô‡∏õ‡∏£‡∏¥‡∏°‡∏≤‡∏ì‡∏à‡∏£‡∏≤‡∏à‡∏£‡∏™‡∏∞‡∏™‡∏° ESALs
                ‡∏Ñ‡∏ß‡∏≤‡∏°‡∏ô‡πà‡∏≤‡πÄ‡∏ä‡∏∑‡πà‡∏≠‡∏ñ‡∏∑‡∏≠ ‡πÅ‡∏•‡∏∞‡∏Ñ‡∏∏‡∏ì‡∏™‡∏°‡∏ö‡∏±‡∏ï‡∏¥‡∏Ç‡∏≠‡∏á‡∏î‡∏¥‡∏ô‡∏£‡∏≠‡∏á‡∏£‡∏±‡∏ö
                ‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡πÇ‡∏Ñ‡∏£‡∏á‡∏Å‡∏≤‡∏£‡∏ô‡∏µ‡πâ‡∏ó‡∏µ‡πà‡∏õ‡∏£‡∏∂‡∏Å‡∏©‡∏≤‡πÑ‡∏î‡πâ‡∏Å‡∏≥‡∏´‡∏ô‡∏î‡∏Ñ‡πà‡∏≤‡∏û‡∏≤‡∏£‡∏≤‡∏°‡∏¥‡πÄ‡∏ï‡∏≠‡∏£‡πå‡∏´‡∏•‡∏±‡∏Å‡πÉ‡∏ô‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö ‡πÑ‡∏î‡πâ‡πÅ‡∏Å‡πà
                ‡∏õ‡∏£‡∏¥‡∏°‡∏≤‡∏ì W&#8321;&#8328; = {hl_purple(f"{W18:,.0f}")} 18-kip ESALs
                ‡∏ó‡∏µ‡πà‡∏£‡∏∞‡∏î‡∏±‡∏ö‡∏Ñ‡∏ß‡∏≤‡∏°‡∏ô‡πà‡∏≤‡πÄ‡∏ä‡∏∑‡πà‡∏≠‡∏ñ‡∏∑‡∏≠ (Reliability) = {hl_purple(reliability)} %
                ‡πÇ‡∏î‡∏¢‡∏°‡∏µ‡∏î‡∏¥‡∏ô‡πÄ‡∏î‡∏¥‡∏°‡∏Ñ‡πà‡∏≤ CBR = {hl_purple(f"{CBR:.1f}")} % (M&#7523; = {hl_purple(f"{Mr:,.0f}")} psi)
                ‡∏ú‡∏•‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö‡πÑ‡∏î‡πâ‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á {hl_purple(num_layers_prev)} ‡∏ä‡∏±‡πâ‡∏ô
                ‡∏ó‡∏µ‡πà SN&#8203;_required = {hl_purple(f"{calc_results['total_sn_required']:.2f}")}
                ‡πÅ‡∏•‡∏∞ SN&#8203;_provided = {hl_purple(f"{calc_results['total_sn_provided']:.2f}")}
                ‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡∏£‡∏ß‡∏° {hl_purple(f"{total_thick_prev:.0f}")} ‡∏ã‡∏°.
                ‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö{hl_purple(passed_prev)}
                ‡∏î‡∏±‡∏á‡πÅ‡∏™‡∏î‡∏á‡∏ú‡∏•‡∏Å‡∏≤‡∏£‡∏ß‡∏¥‡πÄ‡∏Ñ‡∏£‡∏≤‡∏∞‡∏´‡πå‡πÉ‡∏ô‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏ó‡∏µ‡πà <b>{hl_yellow(rs_table_number_inputs)}</b>
                ‡πÅ‡∏•‡∏∞‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏ó‡∏µ‡πà <b>{hl_yellow(rs_table_number_materials)}</b>
                ‡πÅ‡∏•‡∏∞‡∏£‡∏π‡∏õ‡∏ó‡∏µ‡πà <b>{hl_yellow(rs_figure_number)}</b>
            </p>
        </div>
        """
        st.markdown(intro_html, unsafe_allow_html=True)
        st.caption("üü£ ‡∏™‡∏µ‡∏°‡πà‡∏ß‡∏á = ‡∏î‡∏∂‡∏á‡∏à‡∏≤‡∏Å‡∏ú‡∏•‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡∏≠‡∏±‡∏ï‡πÇ‡∏ô‡∏°‡∏±‡∏ï‡∏¥ | üü° ‡∏™‡∏µ‡πÄ‡∏´‡∏•‡∏∑‡∏≠‡∏á = ‡∏ú‡∏π‡πâ‡πÉ‡∏ä‡πâ‡∏Å‡∏£‡∏≠‡∏Å‡πÄ‡∏≠‡∏á")
        st.markdown("---")

        # ===== EXPORT BUTTONS ‚Äî ‡∏à‡∏±‡∏î‡πÉ‡∏´‡∏°‡πà‡πÄ‡∏õ‡πá‡∏ô 2 ‡πÅ‡∏ñ‡∏ß =====
        st.markdown("### üì• ‡∏î‡∏≤‡∏ß‡∏ô‡πå‡πÇ‡∏´‡∏•‡∏î‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô")

        # ‡πÅ‡∏ñ‡∏ß 1: Word reports
        col_r1, col_r2 = st.columns(2)
        with col_r1:
            if st.button("üìã ‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô‡πÅ‡∏ö‡∏ö‡∏ó‡∏µ‡πà‡∏õ‡∏£‡∏∂‡∏Å‡∏©‡∏≤", type="primary",
                         use_container_width=True,
                         help="‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô‡∏£‡∏π‡∏õ‡πÅ‡∏ö‡∏ö‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡∏£‡∏ß‡∏°‡∏Å‡∏±‡∏ö‡∏ö‡∏ó‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô‡∏≠‡∏∑‡πà‡∏ô"):
                with st.spinner("‡∏Å‡∏≥‡∏•‡∏±‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô..."):
                    fig_intro = plot_pavement_section(calc_results['layers'], Mr, CBR, lang='th')
                    doc_intro_bytes = create_word_report_intro(
                        project_title, inputs, calc_results, design_check, fig_intro, report_settings)
                    plt.close(fig_intro)
                st.download_button(
                    label="‚¨áÔ∏è ‡∏î‡∏≤‡∏ß‡∏ô‡πå‡πÇ‡∏´‡∏•‡∏î‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô‡πÅ‡∏ö‡∏ö‡∏ó‡∏µ‡πà‡∏õ‡∏£‡∏∂‡∏Å‡∏©‡∏≤ (.docx)",
                    data=doc_intro_bytes,
                    file_name=f"Flexible_Intro_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

        with col_r2:
            if st.button("üìù ‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô‡πÅ‡∏ö‡∏ö‡∏¢‡πà‡∏≠", use_container_width=True):
                with st.spinner("‡∏Å‡∏≥‡∏•‡∏±‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô..."):
                    fig_thai = plot_pavement_section(calc_results['layers'], Mr, CBR, lang='th')
                    doc_bytes = create_word_report(project_title, inputs, calc_results, design_check, fig_thai)
                    plt.close(fig_thai)
                st.download_button(
                    label="‚¨áÔ∏è ‡∏î‡∏≤‡∏ß‡∏ô‡πå‡πÇ‡∏´‡∏•‡∏î‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô‡πÅ‡∏ö‡∏ö‡∏¢‡πà‡∏≠ (.docx)",
                    data=doc_bytes,
                    file_name=f"AASHTO_Flexible_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

        # ‡πÅ‡∏ñ‡∏ß 2: PNG + JSON
        col_r3, col_r4 = st.columns(2)
        with col_r3:
            fig_export = plot_pavement_section(calc_results['layers'], Mr, CBR, lang=fig_lang)
            fig_bytes  = get_figure_as_bytes(fig_export)
            plt.close(fig_export)
            st.download_button(
                label="üì∏ ‡∏î‡∏≤‡∏ß‡∏ô‡πå‡πÇ‡∏´‡∏•‡∏î‡∏£‡∏π‡∏õ‡∏ï‡∏±‡∏î‡∏Ç‡∏ß‡∏≤‡∏á (.png)",
                data=fig_bytes,
                file_name=f"Pavement_Section_{datetime.now().strftime('%Y%m%d_%H%M')}.png",
                mime="image/png",
                use_container_width=True
            )

        with col_r4:
            export_data = {
                'project_title': project_title,
                'W18': W18, 'reliability': reliability, 'So': So,
                'P0': P0, 'Pt': Pt, 'CBR': CBR,
                'num_layers': num_layers,
                'layers': layer_data,
                'ac_sublayers': st.session_state.get('ac_sublayers', None),
                'report_settings': report_settings,
            }
            json_str = json.dumps(export_data, ensure_ascii=False, indent=2)
            st.download_button(
                label="üíæ ‡∏î‡∏≤‡∏ß‡∏ô‡πå‡πÇ‡∏´‡∏•‡∏î‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏• (.json)",
                data=json_str,
                file_name=f"Flexible_Input_{datetime.now().strftime('%Y%m%d_%H%M')}.json",
                mime="application/json",
                use_container_width=True
            )

        st.markdown("---")

        # ===== Summary Table =====
        st.subheader("üìä ‡∏™‡∏£‡∏∏‡∏õ‡∏ú‡∏•‡∏Å‡∏≤‡∏£‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö")
        summary_data = [
            ("‡∏ä‡∏∑‡πà‡∏≠‡πÇ‡∏Ñ‡∏£‡∏á‡∏Å‡∏≤‡∏£", project_title),
            ("W‚ÇÅ‚Çà (Design ESALs)", f"{W18:,.0f} ({W18/1e6:,.2f} ‡∏•‡πâ‡∏≤‡∏ô)"),
            ("Reliability", f"{reliability}%"),
            ("CBR", f"{CBR:.1f}%"),
            ("M·µ£ (Subgrade)", f"{Mr:,} psi"),
            ("SN Required", f"{calc_results['total_sn_required']:.2f}"),
            ("SN Provided", f"{calc_results['total_sn_provided']:.2f}"),
            ("Safety Margin", f"{design_check['safety_margin']:.2f}"),
            ("‡∏ú‡∏•‡∏Å‡∏≤‡∏£‡∏ï‡∏£‡∏ß‡∏à‡∏™‡∏≠‡∏ö", "‚úÖ PASS" if design_check['passed'] else "‚ùå FAIL"),
        ]
        st.table(summary_data)

    # ===== FOOTER =====
    st.markdown("---")
    st.markdown("""
    <div style='text-align:center;color:gray;'>
    <p>AASHTO 1993 Flexible Pavement Design Application v6.0</p>
    <p>‡∏û‡∏±‡∏í‡∏ô‡∏≤‡πÇ‡∏î‡∏¢ ‡∏£‡∏®.‡∏î‡∏£.‡∏≠‡∏¥‡∏ó‡∏ò‡∏¥‡∏û‡∏• ‡∏°‡∏µ‡∏ú‡∏• // ‡∏†‡∏≤‡∏Ñ‡∏ß‡∏¥‡∏ä‡∏≤‡∏Ñ‡∏£‡∏∏‡∏®‡∏≤‡∏™‡∏ï‡∏£‡πå‡πÇ‡∏¢‡∏ò‡∏≤ // ‡∏°‡∏à‡∏û.</p>
    </div>
    """, unsafe_allow_html=True)


# ================================================================================
# ENTRY POINT
# ================================================================================

if __name__ == "__main__":
    main()
