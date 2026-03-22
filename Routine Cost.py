import streamlit as st
import pandas as pd
from io import BytesIO
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
import json, hashlib
from datetime import datetime

st.set_page_config(
    page_title="ระบบคำนวณงานบำรุงปกติ",
    page_icon="🛣️",
    layout="wide",
)

# ─────────────────────────────────────────────
# LOOKUP TABLES
# ─────────────────────────────────────────────

X1_MAP = {
    "High Type (AC/PM บนหินคลุก)": 0.00,
    "Intermediate Type (AC/PM บน Stabilized)": 0.50,
    "Low Type (ST บน Soil-Aggregate)": 1.00,
}

X2_BREAKS = [(0,2,1.00),(2.01,3,0.75),(3.01,4,0.50),(4.01,5,0.25),(5.01,999,0.00)]

X3_LOWER = [0,501,601,701,801,901,1001,1101,1201,1301,1401,1501,1601,1701,1801,1901,2001,2201,2401,2601,2801,3001,3301,3601,3901,4201,4501,4801,5101,5401,5701]
X3_UPPER = [500,600,700,800,900,1000,1100,1200,1300,1400,1500,1600,1700,1800,1900,2000,2200,2400,2600,2800,3000,3300,3600,3900,4200,4500,4600,5100,5400,5700,999999]
X3_VAL   = [0,0.04,0.08,0.12,0.16,0.20,0.24,0.29,0.33,0.37,0.41,0.45,0.49,0.53,0.57,0.61,0.69,0.78,0.86,0.94,1.02,1.14,1.27,1.37,1.51,1.64,1.76,1.88,2.00,2.13,2.25]

X3_OPTIONS = {
    "0 – 500       (X3 = 0.00)": 0.00,
    "501 – 600     (X3 = 0.04)": 0.04,
    "601 – 700     (X3 = 0.08)": 0.08,
    "701 – 800     (X3 = 0.12)": 0.12,
    "801 – 900     (X3 = 0.16)": 0.16,
    "901 – 1,000   (X3 = 0.20)": 0.20,
    "1,001 – 1,100 (X3 = 0.24)": 0.24,
    "1,101 – 1,200 (X3 = 0.29)": 0.29,
    "1,201 – 1,300 (X3 = 0.33)": 0.33,
    "1,301 – 1,400 (X3 = 0.37)": 0.37,
    "1,401 – 1,500 (X3 = 0.41)": 0.41,
    "1,501 – 1,600 (X3 = 0.45)": 0.45,
    "1,601 – 1,700 (X3 = 0.49)": 0.49,
    "1,701 – 1,800 (X3 = 0.53)": 0.53,
    "1,801 – 1,900 (X3 = 0.57)": 0.57,
    "1,901 – 2,000 (X3 = 0.61)": 0.61,
    "2,001 – 2,200 (X3 = 0.69)": 0.69,
    "2,201 – 2,400 (X3 = 0.78)": 0.78,
    "2,401 – 2,600 (X3 = 0.86)": 0.86,
    "2,601 – 2,800 (X3 = 0.94)": 0.94,
    "2,801 – 3,000 (X3 = 1.02)": 1.02,
    "3,001 – 3,300 (X3 = 1.14)": 1.14,
    "3,301 – 3,600 (X3 = 1.27)": 1.27,
    "3,601 – 3,900 (X3 = 1.37)": 1.37,
    "3,901 – 4,200 (X3 = 1.51)": 1.51,
    "4,201 – 4,500 (X3 = 1.64)": 1.64,
    "4,501 – 4,800 (X3 = 1.76)": 1.76,
    "4,801 – 5,100 (X3 = 1.88)": 1.88,
    "5,101 – 5,400 (X3 = 2.00)": 2.00,
    "5,401 – 5,700 (X3 = 2.13)": 2.13,
    "5,701+         (X3 = 2.25)": 2.25,
}

X4_BREAKS = [(0,3,0.00),(4,4,0.20),(5,5,0.40),(6,6,0.60),(7,7,0.80),(8,8,1.00),(9,9,1.20),(10,10,1.40),(11,11,1.60),(12,99999,1.80)]
X5_BREAKS = [(0,5.49,0.00),(5.50,5.99,0.02),(6.00,6.49,0.05),(6.50,6.99,0.10),(7.00,9999,0.19)]

TERRAIN_MAP  = {"ที่ราบ (0-3%)": "P", "ลูกเนิน (3-5%)": "R", "ลูกเนินสลับเขา (5-7%)": "RM", "เขา (>7%)": "S"}
TERRAIN_KEYS = list(TERRAIN_MAP.keys())
X6_MAP  = {"P": 0.00, "R": 0.02, "RM": 0.04, "S": 0.07}
Y3_MAP  = {"P": 0.00, "R": 0.24, "RM": 0.36, "S": 0.48}
Y4_MAP  = {"P": 0.00, "R": 0.24, "RM": 0.36, "S": 0.48}
Y6_MAP  = {"P": 0.00, "R": 0.04, "RM": 0.08, "S": 0.12}
Y1_BREAKS = [(0,40,0.00),(40.01,60,0.10),(60.01,80,0.20),(80.01,9999,0.30)]
Y2_BREAKS = [(0,1.75,0.00),(1.76,2.00,0.10),(2.01,2.25,0.15),(2.26,9999,0.20)]
Y5_BREAKS = [(0,20.99,0.00),(21,25,0.02),(25.01,30,0.04),(30.01,9999,0.06)]

Z1_MAP   = {1:0.00,2:0.25,3:0.50,4:0.75,5:1.00,6:1.30,7:1.60,8:2.00}
Z2_BREAKS= [(0,2,1.00),(2.01,3,0.75),(3.01,4,0.50),(4.01,5,0.25),(5.01,999,0.00)]

Z3_LOWER = [0,1001,2001,3001,4001,5001,6001,7001,8001,9001,10001,15001]
Z3_UPPER = [1000,2000,3000,4000,5000,6000,7000,8000,9000,10000,15000,999999]
Z3_VAL   = [0,0.20,0.30,0.50,0.75,1.00,1.25,1.50,1.75,2.00,2.50,3.00]

Z3_OPTIONS = {
    "0 – 1,000        (Z3 = 0.00)": 0.00,
    "1,001 – 2,000    (Z3 = 0.20)": 0.20,
    "2,001 – 3,000    (Z3 = 0.30)": 0.30,
    "3,001 – 4,000    (Z3 = 0.50)": 0.50,
    "4,001 – 5,000    (Z3 = 0.75)": 0.75,
    "5,001 – 6,000    (Z3 = 1.00)": 1.00,
    "6,001 – 7,000    (Z3 = 1.25)": 1.25,
    "7,001 – 8,000    (Z3 = 1.50)": 1.50,
    "8,001 – 9,000    (Z3 = 1.75)": 1.75,
    "9,001 – 10,000   (Z3 = 2.00)": 2.00,
    "10,001 – 15,000  (Z3 = 2.50)": 2.50,
    "15,001+           (Z3 = 3.00)": 3.00,
}

Z4_BREAKS = [(0,6.49,0.00),(6.50,6.99,0.08),(7.00,9999,0.17)]
A1_BREAKS = [(0,100,0.00),(101,150,0.13),(151,200,0.24),(201,250,0.36),(251,300,0.47),(301,350,0.59),(351,400,0.71),(401,9999,0.95)]
A3_BREAKS = [(0,6.49,0.00),(6.50,7.49,0.17),(7.50,8.49,0.33),(8.50,9.49,0.55),(9.50,10.49,0.67),(10.50,11.49,0.84),(11.50,9999,1.00)]
B1_BREAKS_KB = [(0,20,0.00),(20.01,30,0.08),(30.01,40,0.13),(40.01,50,0.21),(50.01,9999,0.24)]
B2_MAP   = {"P": 0.05, "R": 0.13, "RM": 0.22, "S": 0.32}
B3_MAP   = {"P": 0.00, "R": 0.40, "RM": 0.60, "S": 0.80}
B4_BREAKS= [(0,20,0.02),(21,21,0.03),(22,22,0.10),(23,23,0.15),(24,24,0.20),(25,25,0.25),(26,26,0.30),(27,27,0.35),(28,28,0.40),(29,29,0.45),(30,9999,0.50)]

# ─────────────────────────────────────────────
# CALCULATION FUNCTIONS
# ─────────────────────────────────────────────

def lookup_range(value, breaks):
    for lo, hi, v in breaks:
        if lo <= value <= hi:
            return v
    return breaks[-1][2]

def lookup_list(value, lower, upper, vals):
    for i, (lo, hi) in enumerate(zip(lower, upper)):
        if lo <= value <= hi:
            return vals[i]
    return vals[-1]

def calc_Ka(x1, x2_cbr, x3_factor, x4_age, x5_width, x6_terrain,
            y1_row, y2_shoulder, y3_terrain, y4_terrain, y5_bridge, y6_terrain):
    X1 = x1
    X2 = lookup_range(x2_cbr, X2_BREAKS)
    X3 = x3_factor
    X4 = lookup_range(x4_age, X4_BREAKS)
    X5 = lookup_range(x5_width, X5_BREAKS)
    X6 = X6_MAP[x6_terrain]
    Y1 = lookup_range(y1_row, Y1_BREAKS)
    Y2 = lookup_range(y2_shoulder, Y2_BREAKS)
    Y3 = Y3_MAP[y3_terrain]
    Y4 = Y4_MAP[y4_terrain]
    Y5 = lookup_range(y5_bridge, Y5_BREAKS)
    Y6 = Y6_MAP[y6_terrain]
    Ka = 1 + 0.50*(X1+X2+X3+X4+X5+X6+Y1+Y2+Y3+Y4+Y5+Y6)
    return Ka, {"X1":X1,"X2":X2,"X3":X3,"X4":X4,"X5":X5,"X6":X6,
                "Y1":Y1,"Y2":Y2,"Y3":Y3,"Y4":Y4,"Y5":Y5,"Y6":Y6}

def calc_Kc(z1, z2_cbr, z3_factor, z4_width,
            y1_row, y2_shoulder, y3_terrain, y4_terrain, y5_bridge, y6_terrain):
    Z1 = Z1_MAP.get(z1, 0)
    Z2 = lookup_range(z2_cbr, Z2_BREAKS)
    Z3 = z3_factor
    Z4 = lookup_range(z4_width, Z4_BREAKS)
    Y1 = lookup_range(y1_row, Y1_BREAKS)
    Y2 = lookup_range(y2_shoulder, Y2_BREAKS)
    Y3 = Y3_MAP[y3_terrain]
    Y4 = Y4_MAP[y4_terrain]
    Y5 = lookup_range(y5_bridge, Y5_BREAKS)
    Y6 = Y6_MAP[y6_terrain]
    Kc = 1 + 0.50*(Z1+Z2+Z3+Z4+Y1+Y2+Y3+Y4+Y5+Y6)
    return Kc, {"Z1":Z1,"Z2":Z2,"Z3":Z3,"Z4":Z4,
                "Y1":Y1,"Y2":Y2,"Y3":Y3,"Y4":Y4,"Y5":Y5,"Y6":Y6}

def calc_Ks(a1_aadt, a3_width, b1_row, b2_terrain, b3_terrain, b4_bridge):
    A1 = lookup_range(a1_aadt, A1_BREAKS)
    A2 = 0.00
    A3 = lookup_range(a3_width, A3_BREAKS)
    B1 = lookup_range(b1_row, B1_BREAKS_KB)
    B2 = B2_MAP[b2_terrain]
    B3 = B3_MAP[b3_terrain]
    B4 = lookup_range(b4_bridge, B4_BREAKS)
    Ks = 1 + 0.70*(A1+A2+A3) + 0.30*(B1+B2+B3+B4)
    return Ks, {"A1":A1,"A2":A2,"A3":A3,"B1":B1,"B2":B2,"B3":B3,"B4":B4}

def calc_K_prime(K, warranty_years):
    if warranty_years == 0:   return K
    elif warranty_years == 1: return 0.5 * K
    else:                     return 0.25 * K

def calc_budget(dist_km, K, Km, N_std):
    return round(dist_km * K * Km * N_std / 100) * 100

def calc_workload(dist_equiv, K_prime):
    return round(dist_equiv * K_prime, 3)

# ─────────────────────────────────────────────
# SESSION STATE INIT
# ─────────────────────────────────────────────

def init_state():
    defaults = {
        "project_name": "โครงการบำรุงรักษาทางหลวง",
        "district": "",
        "year": "2569",
        "Na": 35000.0, "Ns": 6500.0, "Nc": 35000.0,
        "Km_a": 1.0,  "Km_s": 1.0,  "Km_c": 1.0,
        "rows_ac": [], "rows_cc": [], "rows_gr": [],
        "json_version": 0,
        "loaded_data": {},
        "loaded_json_hash": "",
    }
    for k, dv in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = dv

init_state()

# ─────────────────────────────────────────────
# JSON HELPERS (Streamlit_JSON_Load_Prompt)
# ─────────────────────────────────────────────

def build_json_data():
    return {
        "project_name": st.session_state["project_name"],
        "district":     st.session_state["district"],
        "year":         st.session_state["year"],
        "Na":  st.session_state["Na"],  "Ns":  st.session_state["Ns"],
        "Nc":  st.session_state["Nc"],
        "Km_a":st.session_state["Km_a"],"Km_s":st.session_state["Km_s"],
        "Km_c":st.session_state["Km_c"],
        "rows_ac": st.session_state["rows_ac"],
        "rows_cc": st.session_state["rows_cc"],
        "rows_gr": st.session_state["rows_gr"],
        "saved_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }

def get_default(key, fallback=None):
    """กฎที่ 3: ตรวจ loaded_data ก่อน return ค่า default เสมอ"""
    return st.session_state.get("loaded_data", {}).get(key, fallback)

# ─────────────────────────────────────────────
# WORD REPORT HELPERS (Universal_Word_Report_Prompt)
# ─────────────────────────────────────────────

def set_run_font(run, font="TH SarabunPSK", size=None, bold=False, italic=False):
    """CRITICAL: ตั้งค่า font ผ่าน XML w:cs ด้วยเสมอ เพื่อให้ภาษาไทยแสดงผลถูกต้อง"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    if size is None:
        size = Pt(14)
    run.font.name   = font
    run.font.size   = size
    run.font.bold   = bold
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"),    font)   # REQUIRED for Thai script

def add_thai_para(doc, text="", bold=False, italic=False,
                  first_indent=True, font="TH SarabunPSK", size=None):
    """Body paragraph พร้อม thaiDistribute justify (ป้องกันอักษรไทยกระจายห่าง)"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    if size is None:
        size = Pt(14)
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "thaiDistribute")
    pPr.append(jc)
    if first_indent:
        ind = OxmlElement("w:ind")
        ind.set(qn("w:firstLine"), "720")
        pPr.append(ind)
    if text:
        run = p.add_run(text)
        set_run_font(run, font=font, size=size, bold=bold, italic=italic)
    return p

def add_heading_word(doc, text, level=1):
    """Heading พร้อม override theme font เพื่อให้ภาษาไทยแสดงผลถูกต้อง"""
    from docx.shared import Pt
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=Pt(15), bold=True)
    return p

def add_table_word(doc, headers, rows, col_widths=None):
    """Table พร้อม Table Grid style, NEVER ใช้ cell.text = '...' โดยตรง"""
    from docx.shared import Pt, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT
    table = doc.add_table(rows=1, cols=len(headers))
    table.style     = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            table.columns[i].width = Cm(w)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(str(h))
        set_run_font(run, size=Pt(14), bold=True)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=Pt(14))
    return table

def add_total_row_word(table, values, bold=True):
    """เพิ่มแถวรวม (Total row) ที่ท้ายตาราง"""
    from docx.shared import Pt
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.paragraphs[0].clear()
        if val is not None:
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=Pt(14), bold=bold)
    return row

def add_mixed_para(doc, parts, first_indent=True):
    """Paragraph ที่มีหลาย run (บางส่วน bold) — parts = [(text, bold), ...]"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "thaiDistribute")
    pPr.append(jc)
    if first_indent:
        ind = OxmlElement("w:ind")
        ind.set(qn("w:firstLine"), "720")
        pPr.append(ind)
    for text, bold in parts:
        run = p.add_run(text)
        set_run_font(run, size=Pt(14), bold=bold)
    return p

# ─────────────────────────────────────────────
# WORD REPORT GENERATOR
# ─────────────────────────────────────────────

def generate_word_report(include_gravel=True):
    import io
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ss = st.session_state
    doc = Document()

    # ─── Section Counter (dynamic numbering) ────────────────────────────────
    # ใช้ class เพื่อให้ nested function แก้ไขค่าได้
    class SC:
        h1 = 0   # หัวข้อระดับ 1
        h2 = 0   # หัวข้อระดับ 2 (reset ทุกครั้งที่ h1 เพิ่ม)

    def next_h1(title):
        """เพิ่มเลข h1 แล้วสร้าง heading พร้อม label เช่น  '3.  สูตรการคำนวณ'"""
        SC.h1 += 1
        SC.h2 = 0
        add_heading_word(doc, f"{SC.h1}.  {title}", level=1)
        return SC.h1

    def next_h2(title):
        """เพิ่มเลข h2 แล้วสร้าง heading พร้อม label เช่น '3.2  ผิวคอนกรีต (Kc)'"""
        SC.h2 += 1
        add_heading_word(doc, f"{SC.h1}.{SC.h2}  {title}", level=2)
        return SC.h1, SC.h2

    # ─── ตั้งค่า Normal style ────────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name = "TH SarabunPSK"
    normal.font.size = Pt(14)

    # ขนาดหน้ากระดาษ A4
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ─── ปกรายงาน ───────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = title_p.add_run("รายการคำนวณงานบำรุงปกติ")
    set_run_font(run_t, size=Pt(18), bold=True)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = sub_p.add_run("เครดิต : รศ.ดร.อิทธิพล มีผล")
    set_run_font(run_s, size=Pt(14))

    doc.add_paragraph()

    # ─── บทเกริ่นนำ ─────────────────────────────────────────────────────────
    add_thai_para(doc,
        f"รายการคำนวณฉบับนี้จัดทำขึ้นเพื่อประกอบการพิจารณาจัดสรรงบประมาณงานบำรุงปกติ "
        f"สำหรับ{ss['project_name']} ในความรับผิดชอบของ{ss['district'] or 'หน่วยงานที่รับผิดชอบ'} "
        f"ประจำปีงบประมาณ พ.ศ. {ss['year']} โดยดำเนินการคำนวณตามหลักเกณฑ์และวิธีการที่กำหนดไว้ใน "
        f"คู่มือการคิดค่าปริมาณงานและงานบำรุงปกติ กองบำรุง กรมทางหลวง มกราคม พ.ศ. 2538"
    )
    add_thai_para(doc,
        "การคำนวณงบประมาณงานบำรุงปกติใช้วิธีการกำหนดค่าสัมประสิทธิ์ปรับแก้ (K) "
        "เพื่อสะท้อนลักษณะเฉพาะของแต่ละสายทาง ได้แก่ ประเภทและสภาพผิวทาง ความสามารถรับแรงของดินคันทาง "
        "ปริมาณจราจร อายุการใช้งาน ขนาดเรขาคณิตของถนน สภาพภูมิประเทศ รวมถึงองค์ประกอบของงานบำรุงรักษา "
        "ที่เกี่ยวเนื่อง ทั้งงานจราจรสงเคราะห์ ท่อระบายน้ำ สะพาน และงานทำความสะอาดทางระบายน้ำ "
        "โดยแบ่งการคำนวณออกเป็น 3 ประเภทตามลักษณะผิวทาง ได้แก่ "
        "ผิวแอสฟัลท์คอนกรีต (Ka) ผิวคอนกรีตซีเมนต์ (Kc) และผิวลูกรัง (Ks)"
    )
    add_thai_para(doc,
        "ผลการคำนวณที่ได้แสดงถึงงบประมาณที่ต้องการรายปี (บาท/ปี) และปริมาณงาน (Workload) "
        "ของแต่ละสายทาง ซึ่งนำไปใช้ประกอบการจัดทำแผนงานและของบประมาณในระดับแขวงการทาง "
        "ตามขั้นตอนของกรมทางหลวงต่อไป"
    )
    doc.add_paragraph()

    # ─── หัวข้อ 1: ข้อมูลโครงการ ────────────────────────────────────────────
    next_h1("ข้อมูลโครงการ")
    add_table_word(doc,
        headers=["รายการ", "ข้อมูล"],
        rows=[
            ["ชื่อโครงการ", ss["project_name"]],
            ["วันที่จัดทำ", datetime.now().strftime("%d/%m/%Y %H:%M")],
        ],
        col_widths=[5, 11],
    )
    doc.add_paragraph()

    # ─── หัวข้อ 2: อัตราค่าบำรุงมาตรฐาน ────────────────────────────────────
    next_h1("อัตราค่าบำรุงมาตรฐาน (N) และค่า Factor วัสดุ (Km)")
    add_table_word(doc,
        headers=["ประเภทผิวทาง", "N มาตรฐาน (บาท/กม./ปี)", "Km วัสดุ"],
        rows=[
            ["ผิวแอสฟัลท์ (Ka)", f"{ss['Na']:,.0f}", f"{ss['Km_a']:.3f}"],
            ["ผิวลูกรัง (Ks)",   f"{ss['Ns']:,.0f}", f"{ss['Km_s']:.3f}"],
            ["ผิวคอนกรีต (Kc)", f"{ss['Nc']:,.0f}", f"{ss['Km_c']:.3f}"],
        ],
        col_widths=[5, 6, 5],
    )
    doc.add_paragraph()

    # ─── หัวข้อ 3: สูตรการคำนวณและค่า Factor ───────────────────────────────
    next_h1("สูตรการคำนวณและค่า Factor")

    # 3.1  ผิวแอสฟัลท์ (Ka)
    next_h2("ผิวแอสฟัลท์ (Ka)")
    add_thai_para(doc, "Ka = 1 + 0.50 × (X1 + X2 + X3 + X4 + X5 + X6 + Y1 + Y2 + Y3 + Y4 + Y5 + Y6)",
                  first_indent=False)
    if ss["rows_ac"]:
        for row_ac in ss["rows_ac"]:
            route_lbl = f"{row_ac.get('ชื่อสายทาง','')}  (ตอน {row_ac.get('ตอนควบคุม','')})"
            add_thai_para(doc, route_lbl, bold=True, first_indent=False)
            add_table_word(doc,
                headers=["Factor", "คำอธิบาย", "ค่าที่ใช้"],
                rows=[
                    ["X1", "ลักษณะผิวทางและพื้นทาง",        f"{float(row_ac.get('X1',0)):.4f}"],
                    ["X2", "CBR ดินเดิม",                    f"{float(row_ac.get('X2',0)):.4f}"],
                    ["X3", "ปริมาณจราจร AADT",               f"{float(row_ac.get('X3',0)):.4f}"],
                    ["X4", "อายุบริการ",                      f"{float(row_ac.get('X4',0)):.4f}"],
                    ["X5", "ความกว้างผิวทาง",                 f"{float(row_ac.get('X5',0)):.4f}"],
                    ["X6", "ภูมิประเทศ",                      f"{float(row_ac.get('X6',0)):.4f}"],
                    ["Y1", "ความกว้างเขตทาง",                 f"{float(row_ac.get('Y1',0)):.4f}"],
                    ["Y2", "ไหล่ทางกว้างสุด 1 ข้าง",         f"{float(row_ac.get('Y2',0)):.4f}"],
                    ["Y3", "จราจรสงเคราะห์",                  f"{float(row_ac.get('Y3',0)):.4f}"],
                    ["Y4", "ท่อระบายน้ำ",                     f"{float(row_ac.get('Y4',0)):.4f}"],
                    ["Y5", "สะพาน",                           f"{float(row_ac.get('Y5',0)):.4f}"],
                    ["Y6", "ทำความสะอาดระบาย",                f"{float(row_ac.get('Y6',0)):.4f}"],
                    ["Ka", "ค่า K ผิวแอสฟัลท์",               f"{float(row_ac.get('K',0)):.4f}"],
                ],
                col_widths=[1.5, 6, 3],
            )
            doc.add_paragraph()
    doc.add_paragraph()

    # 3.2  ผิวคอนกรีต (Kc)
    next_h2("ผิวคอนกรีต (Kc)")
    add_thai_para(doc, "Kc = 1 + 0.50 × (Z1 + Z2 + Z3 + Z4 + Y1 + Y2 + Y3 + Y4 + Y5 + Y6)",
                  first_indent=False)
    if ss["rows_cc"]:
        for row_cc in ss["rows_cc"]:
            route_lbl = f"{row_cc.get('ชื่อสายทาง','')}  (ตอน {row_cc.get('ตอนควบคุม','')})"
            add_thai_para(doc, route_lbl, bold=True, first_indent=False)
            add_table_word(doc,
                headers=["Factor", "คำอธิบาย", "ค่าที่ใช้"],
                rows=[
                    ["Z1", "ดัชนีสภาพผิวทาง",              f"{float(row_cc.get('Z1',0)):.4f}"],
                    ["Z2", "CBR ดินคันทาง",                 f"{float(row_cc.get('Z2',0)):.4f}"],
                    ["Z3", "ปริมาณจราจร AADT",              f"{float(row_cc.get('Z3',0)):.4f}"],
                    ["Z4", "ความกว้างผิวทาง",                f"{float(row_cc.get('Z4',0)):.4f}"],
                    ["Y1", "ความกว้างเขตทาง",                f"{float(row_cc.get('Y1',0)):.4f}"],
                    ["Y2", "ไหล่ทางกว้างสุด 1 ข้าง",        f"{float(row_cc.get('Y2',0)):.4f}"],
                    ["Y3", "จราจรสงเคราะห์",                 f"{float(row_cc.get('Y3',0)):.4f}"],
                    ["Y4", "ท่อระบายน้ำ",                    f"{float(row_cc.get('Y4',0)):.4f}"],
                    ["Y5", "สะพาน",                          f"{float(row_cc.get('Y5',0)):.4f}"],
                    ["Y6", "ทำความสะอาดระบาย",               f"{float(row_cc.get('Y6',0)):.4f}"],
                    ["Kc", "ค่า K ผิวคอนกรีต",               f"{float(row_cc.get('K',0)):.4f}"],
                ],
                col_widths=[1.5, 6, 3],
            )
            doc.add_paragraph()
    doc.add_paragraph()

    # 3.3  ผิวลูกรัง (Ks) — แสดงเฉพาะเมื่อ include_gravel=True
    if include_gravel:
        next_h2("ผิวลูกรัง (Ks)")
        add_thai_para(doc, "Ks = 1 + 0.70 × (A1 + A2 + A3) + 0.30 × (B1 + B2 + B3 + B4)",
                      first_indent=False)
        if ss["rows_gr"]:
            for row_gr in ss["rows_gr"]:
                route_lbl = f"{row_gr.get('ชื่อสายทาง','')}  (ตอน {row_gr.get('ตอนควบคุม','')})"
                add_thai_para(doc, route_lbl, bold=True, first_indent=False)
                add_table_word(doc,
                    headers=["Factor", "คำอธิบาย", "ค่าที่ใช้"],
                    rows=[
                        ["A1", "ปริมาณจราจร AADT",             f"{float(row_gr.get('A1',0)):.4f}"],
                        ["A2", "ลมฟ้าอากาศ",                   f"{float(row_gr.get('A2',0)):.4f}"],
                        ["A3", "ความกว้างผิวทาง",               f"{float(row_gr.get('A3',0)):.4f}"],
                        ["B1", "ความกว้างเขตทาง",               f"{float(row_gr.get('B1',0)):.4f}"],
                        ["B2", "ภูมิประเทศ (จราจรสงเคราะห์)",  f"{float(row_gr.get('B2',0)):.4f}"],
                        ["B3", "ภูมิประเทศ (ท่อระบายน้ำ)",      f"{float(row_gr.get('B3',0)):.4f}"],
                        ["B4", "สะพาน",                         f"{float(row_gr.get('B4',0)):.4f}"],
                        ["Ks", "ค่า K ผิวลูกรัง",               f"{float(row_gr.get('K',0)):.4f}"],
                    ],
                    col_widths=[1.5, 6, 3],
                )
                doc.add_paragraph()
        doc.add_paragraph()

    # 3.x  การคำนวณงบประมาณงานบำรุงปกติรายปี
    next_h2("การคำนวณงบประมาณงานบำรุงปกติรายปี")
    add_thai_para(doc, "งบประมาณ (บาท/ปี)  =  ระยะทาง (กม.) × K × Km × N  (ปัดเป็นหลักร้อย)",
                  first_indent=False)
    add_thai_para(doc, "ระยะเทียบเท่า (กม.) =  ระยะจริง (กม.) × (จำนวนช่องจราจร / 2)",
                  first_indent=False)
    add_thai_para(doc, "Workload (หน่วย)    =  ระยะเทียบเท่า (กม.) × K'",
                  first_indent=False)

    # 3.x  การปรับ K' ตามช่วงประกัน
    next_h2("การปรับค่า K' ตามช่วงระยะเวลาประกัน")
    add_table_word(doc,
        headers=["เงื่อนไข", "สูตร"],
        rows=[
            ["ไม่มีประกัน",      "K' = K"],
            ["มีประกัน 1 ปี",   "K' = 0.50 × K"],
            ["มีประกัน > 1 ปี", "K' = 0.25 × K"],
        ],
        col_widths=[5, 11],
    )
    doc.add_paragraph()

    # ─── หัวข้อ 4+ : ตารางสายทางแต่ละประเภทผิวทาง ──────────────────────────
    # IMPORTANT: next_h1() เรียกเสมอ (ไม่ skip) เพื่อให้เลข section ต่อเนื่อง
    def write_surface_section(surf_label, rows_data):
        """สร้าง section ตารางสายทาง — next_h1() เรียกเสมอเพื่อให้เลขต่อเนื่อง"""
        next_h1(f"สายทางผิว{surf_label}")
        if not rows_data:
            add_thai_para(doc, "(ไม่มีข้อมูลสายทางประเภทนี้)", first_indent=False)
            doc.add_paragraph()
            return
        headers = ["ตอนควบคุม", "ชื่อสายทาง", "ระยะทาง\n(กม.)", "ช่องจราจร",
                   "ระยะเทียบเท่า\n(กม.)", "K", "ประกัน\n(ปี)", "K'",
                   "Workload\n(หน่วย)", "งบประมาณ\n(บาท/ปี)"]
        table_rows = []
        for r in rows_data:
            table_rows.append([
                r.get("ตอนควบคุม", ""),
                r.get("ชื่อสายทาง", ""),
                f"{r.get('ระยะทาง(กม.)', 0):.3f}",
                r.get("ช่องจราจร", 2),
                f"{r.get('ระยะเทียบเท่า(กม.)', 0):.3f}",
                f"{r.get('K', 0):.4f}",
                r.get("ประกัน(ปี)", 0),
                f"{r.get(chr(75) + chr(39), 0):.4f}",   # K'
                f"{r.get('Workload(หน่วย)', 0):.3f}",
                f"{r.get('งบประมาณ(บาท/ปี)', 0):,.0f}",
            ])
        tbl = add_table_word(doc, headers=headers, rows=table_rows,
                             col_widths=[2.5, 4, 2, 1.5, 2.5, 2, 1.5, 2, 2.5, 3])
        tot_dist  = sum(r.get("ระยะทาง(กม.)", 0)       for r in rows_data)
        tot_equiv = sum(r.get("ระยะเทียบเท่า(กม.)", 0) for r in rows_data)
        tot_wl    = sum(r.get("Workload(หน่วย)", 0)    for r in rows_data)
        tot_bud   = sum(r.get("งบประมาณ(บาท/ปี)", 0)  for r in rows_data)
        add_total_row_word(tbl, [
            "รวม", "", f"{tot_dist:.3f}", "",
            f"{tot_equiv:.3f}", "", "", "",
            f"{tot_wl:.3f}", f"{tot_bud:,.0f}",
        ])
        doc.add_paragraph()

    write_surface_section("แอสฟัลท์ (Ka)", ss["rows_ac"])
    write_surface_section("คอนกรีต (Kc)",  ss["rows_cc"])
    write_surface_section("ลูกรัง (Ks)",   ss["rows_gr"] if include_gravel else [])

    # ─── หัวข้อ x: สรุปผลการคำนวณ ──────────────────────────────────────────
    n_ac = len(ss["rows_ac"]); n_cc = len(ss["rows_cc"]); n_gr = len(ss["rows_gr"])
    bud_ac  = sum(r["งบประมาณ(บาท/ปี)"] for r in ss["rows_ac"])
    bud_cc  = sum(r["งบประมาณ(บาท/ปี)"] for r in ss["rows_cc"])
    bud_gr  = sum(r["งบประมาณ(บาท/ปี)"] for r in ss["rows_gr"])
    wl_ac   = sum(r["Workload(หน่วย)"]   for r in ss["rows_ac"])
    wl_cc   = sum(r["Workload(หน่วย)"]   for r in ss["rows_cc"])
    wl_gr   = sum(r["Workload(หน่วย)"]   for r in ss["rows_gr"])
    dist_ac = sum(r["ระยะทาง(กม.)"]      for r in ss["rows_ac"])
    dist_cc = sum(r["ระยะทาง(กม.)"]      for r in ss["rows_cc"])
    dist_gr = sum(r["ระยะทาง(กม.)"]      for r in ss["rows_gr"])

    def rpc(bud, dist): return f"{bud/dist:,.2f}" if dist > 0 else "-"
    sum_rows = [
        ["ผิวแอสฟัลท์ (Ka)", n_ac, f"{dist_ac:.3f}", f"{wl_ac:.3f}", f"{bud_ac:,.0f}", rpc(bud_ac, dist_ac)],
        ["ผิวคอนกรีต (Kc)",  n_cc, f"{dist_cc:.3f}", f"{wl_cc:.3f}", f"{bud_cc:,.0f}", rpc(bud_cc, dist_cc)],
    ]
    if include_gravel:
        sum_rows.append(["ผิวลูกรัง (Ks)", n_gr, f"{dist_gr:.3f}", f"{wl_gr:.3f}", f"{bud_gr:,.0f}", rpc(bud_gr, dist_gr)])
    bud_total  = bud_ac + bud_cc + (bud_gr if include_gravel else 0)
    dist_total = dist_ac + dist_cc + (dist_gr if include_gravel else 0)
    wl_total   = wl_ac + wl_cc + (wl_gr if include_gravel else 0)
    sum_rows.append(["รวมทุกประเภท", n_ac+n_cc+(n_gr if include_gravel else 0),
                     f"{dist_total:.3f}", f"{wl_total:.3f}",
                     f"{bud_total:,.0f}", rpc(bud_total, dist_total)])

    next_h1("สรุปผลการคำนวณ")
    tbl_sum = add_table_word(doc,
        headers=["ประเภทผิวทาง", "จำนวนสายทาง", "ระยะทาง (กม.)", "Workload (หน่วย)", "งบประมาณ (บาท/ปี)", "อัตรา (บาท/กม./ปี)"],
        rows=sum_rows[:-1],   # ข้อมูลปกติ
        col_widths=[4, 2.5, 3, 3.5, 3.5, 3.5],
    )
    # เพิ่มแถวรวมด้วย bold
    add_total_row_word(tbl_sum, list(sum_rows[-1]))
    doc.add_paragraph()

    # ─── หัวข้อ x: หมายเหตุ ─────────────────────────────────────────────────
    doc.add_paragraph()
    next_h1("หมายเหตุ")
    notes = [
        "การคำนวณอ้างอิงคู่มือการคิดค่าปริมาณงานและงานบำรุงปกติ กองบำรุง กรมทางหลวง มกราคม พ.ศ. 2538",
        "ค่า A2 (ลมฟ้าอากาศ) สำหรับผิวลูกรัง ใช้ค่า 0.00 เนื่องจากกรมทางหลวงยังอยู่ระหว่างการศึกษาเก็บสถิติ",
        "งบประมาณปัดเป็นหลักร้อย ตามแนวทางของกรมทางหลวง",
        "ระยะเทียบเท่า = ระยะจริง × (จำนวนช่องจราจร / 2)",
    ]
    for note in notes:
        p = add_thai_para(doc, first_indent=False)
        run_n = p.add_run(f"- {note}")
        set_run_font(run_n, size=Pt(13))

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream

# ─────────────────────────────────────────────
# WORD REPORT — แบบที่ปรึกษา (nested 3.5.x.x)
# ─────────────────────────────────────────────

def generate_word_report_consultant(include_gravel=True, base_sec="3.5"):
    """
    รายงานแบบที่ปรึกษา — section numbering เริ่มต้นที่ base_sec (default 3.5)
    โครงสร้าง:
      {base}     การคำนวณงบประมาณงานบำรุงปกติ  [h1-level heading]
      {base}.1   อัตราค่าบำรุงมาตรฐาน (N) และค่า Factor วัสดุ (Km)
      {base}.2   สูตรการคำนวณและค่า Factor
      {base}.2.1 ผิวแอสฟัลท์ (Ka)
      {base}.2.2 ผิวคอนกรีต (Kc)
      {base}.2.3 ผิวลูกรัง (Ks)           [เฉพาะ include_gravel=True]
      {base}.3   วิธีการคำนวณงบประมาณงานบำรุงปกติรายปี
      {base}.4   การปรับค่า K' ตามช่วงระยะเวลาประกัน
      {base}.5   สายทางผิวแอสฟัลท์ (Ka)
      {base}.6   สายทางผิวคอนกรีต (Kc)
      {base}.7   สายทางผิวลูกรัง (Ks)     [เรียกเสมอ เพื่อให้เลขต่อเนื่อง]
      {base}.8   สรุปผลการคำนวณ
    """
    import io
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ss  = st.session_state
    doc = Document()

    # ── Section Counter แบบ nested ──────────────────────────────────────────
    # base_parts = [3, 5]  →  prefix = "3.5"
    base_parts = [int(x) for x in base_sec.split(".")]

    class SC:
        lv1 = 0   # 3.5.X
        lv2 = 0   # 3.5.X.Y

    def label_lv0():
        """ป้าย heading หลัก เช่น '3.5' """
        return ".".join(str(x) for x in base_parts)

    def label_lv1():
        """ป้าย heading ระดับ 1 เช่น '3.5.1' """
        return ".".join(str(x) for x in base_parts) + f".{SC.lv1}"

    def label_lv2():
        """ป้าย heading ระดับ 2 เช่น '3.5.2.1' """
        return ".".join(str(x) for x in base_parts) + f".{SC.lv2_parent}.{SC.lv2}"

    class SC:
        lv1 = 0
        lv2 = 0
        lv2_parent = 0   # เก็บ lv1 ที่ lv2 อยู่ใต้

    def h0(title):
        """หัวข้อ base เช่น '3.5  การคำนวณงบประมาณ...' — level=1"""
        SC.lv1 = 0; SC.lv2 = 0
        add_heading_word(doc, f"{label_lv0()}  {title}", level=1)

    def h1(title):
        """หัวข้อระดับ 1 เช่น '3.5.1  อัตราค่าบำรุง...' — level=2"""
        SC.lv1 += 1; SC.lv2 = 0; SC.lv2_parent = SC.lv1
        add_heading_word(doc, f"{label_lv1()}  {title}", level=2)

    def h2(title):
        """หัวข้อระดับ 2 เช่น '3.5.2.1  ผิวแอสฟัลท์' — level=3"""
        SC.lv2 += 1
        add_heading_word(doc, f"{label_lv2()}  {title}", level=3)

    # ── ตั้งค่า Normal style ────────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name = "TH SarabunPSK"
    normal.font.size = Pt(14)
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.5)

    # ── ปกรายงาน ────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = title_p.add_run("รายการคำนวณงานบำรุงปกติ (แบบที่ปรึกษา)")
    set_run_font(run_t, size=Pt(18), bold=True)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = sub_p.add_run("เครดิต : รศ.ดร.อิทธิพล มีผล")
    set_run_font(run_s, size=Pt(14))

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # {base}  การคำนวณงบประมาณงานบำรุงปกติ
    # ════════════════════════════════════════════════════════════════════════
    h0("การคำนวณงบประมาณงานบำรุงปกติ")
    add_thai_para(doc,
        "การคำนวณงบประมาณงานบำรุงปกติใช้วิธีการกำหนดค่าสัมประสิทธิ์ปรับแก้ (K) "
        "เพื่อสะท้อนลักษณะเฉพาะของแต่ละสายทาง ได้แก่ ประเภทและสภาพผิวทาง "
        "ความสามารถรับแรงของดินคันทาง ปริมาณจราจร อายุการใช้งาน ขนาดเรขาคณิตของถนน "
        "สภาพภูมิประเทศ รวมถึงองค์ประกอบของงานบำรุงรักษาที่เกี่ยวเนื่อง "
        "ทั้งงานจราจรสงเคราะห์ ท่อระบายน้ำ สะพาน และงานทำความสะอาดทางระบายน้ำ "
        "โดยแบ่งการคำนวณออกเป็น 3 ประเภทตามลักษณะผิวทาง ได้แก่ "
        "ผิวแอสฟัลท์คอนกรีต (Ka) ผิวคอนกรีตซีเมนต์ (Kc) และผิวลูกรัง (Ks)"
    )
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # {base}.1  อัตราค่าบำรุงมาตรฐาน (N) และค่า Factor วัสดุ (Km)
    # ════════════════════════════════════════════════════════════════════════
    h1("อัตราค่าบำรุงมาตรฐาน (N) และค่า Factor วัสดุ (Km)")
    add_thai_para(doc,
        "อัตราค่าบำรุงมาตรฐาน (N) คือค่าใช้จ่ายพื้นฐานต่อกิโลเมตรต่อปีสำหรับผิวทางแต่ละประเภท "
        "ก่อนการปรับแก้ด้วยค่าสัมประสิทธิ์ K โดยมีค่า Factor วัสดุ (Km) ใช้ปรับตามราคาวัสดุปัจจุบัน "
        "ดังแสดงในตาราง"
    )
    add_table_word(doc,
        headers=["ประเภทผิวทาง", "N มาตรฐาน (บาท/กม./ปี)", "Km วัสดุ"],
        rows=[
            ["ผิวแอสฟัลท์ (Ka)", f"{ss['Na']:,.0f}", f"{ss['Km_a']:.3f}"],
            ["ผิวลูกรัง (Ks)",   f"{ss['Ns']:,.0f}", f"{ss['Km_s']:.3f}"],
            ["ผิวคอนกรีต (Kc)", f"{ss['Nc']:,.0f}", f"{ss['Km_c']:.3f}"],
        ],
        col_widths=[5, 6, 5],
    )
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # {base}.2  สูตรการคำนวณและค่า Factor
    # ════════════════════════════════════════════════════════════════════════
    h1("สูตรการคำนวณและค่า Factor")
    add_thai_para(doc,
        "ค่าสัมประสิทธิ์ปรับแก้ (K) คำนวณจากผลรวมของค่า Factor ที่สะท้อนลักษณะของสายทาง "
        "โดยแบ่งเป็น Factor X หรือ A (เกี่ยวกับผิวทางและปริมาณจราจร) และ Factor Y หรือ B "
        "(เกี่ยวกับเขตทาง ภูมิประเทศ และงานบำรุงรักษาที่เกี่ยวเนื่อง)"
    )
    doc.add_paragraph()

    # ── {base}.2.1  ผิวแอสฟัลท์ (Ka) ───────────────────────────────────────
    h2("ผิวแอสฟัลท์ (Ka)")
    add_thai_para(doc,
        "ค่าสัมประสิทธิ์ปรับแก้สำหรับผิวแอสฟัลท์คอนกรีต (Ka) คำนวณตามสมการ"
    )
    add_thai_para(doc, "Ka = 1 + 0.50 × (X1 + X2 + X3 + X4 + X5 + X6 + Y1 + Y2 + Y3 + Y4 + Y5 + Y6)",
                  first_indent=False)
    if ss["rows_ac"]:
        add_thai_para(doc, "ผลการคำนวณค่า Factor สำหรับสายทางผิวแอสฟัลท์แสดงดังนี้")
        for row_ac in ss["rows_ac"]:
            route_lbl = f"{row_ac.get('ชื่อสายทาง','')}  (ตอน {row_ac.get('ตอนควบคุม','')})"
            add_thai_para(doc, route_lbl, bold=True, first_indent=False)
            add_table_word(doc,
                headers=["Factor", "คำอธิบาย", "ค่าที่ใช้"],
                rows=[
                    ["X1","ลักษณะผิวทางและพื้นทาง",       f"{float(row_ac.get('X1',0)):.4f}"],
                    ["X2","CBR ดินเดิม",                   f"{float(row_ac.get('X2',0)):.4f}"],
                    ["X3","ปริมาณจราจร AADT",               f"{float(row_ac.get('X3',0)):.4f}"],
                    ["X4","อายุบริการ",                     f"{float(row_ac.get('X4',0)):.4f}"],
                    ["X5","ความกว้างผิวทาง",                f"{float(row_ac.get('X5',0)):.4f}"],
                    ["X6","ภูมิประเทศ",                     f"{float(row_ac.get('X6',0)):.4f}"],
                    ["Y1","ความกว้างเขตทาง",                f"{float(row_ac.get('Y1',0)):.4f}"],
                    ["Y2","ไหล่ทางกว้างสุด 1 ข้าง",        f"{float(row_ac.get('Y2',0)):.4f}"],
                    ["Y3","จราจรสงเคราะห์",                 f"{float(row_ac.get('Y3',0)):.4f}"],
                    ["Y4","ท่อระบายน้ำ",                    f"{float(row_ac.get('Y4',0)):.4f}"],
                    ["Y5","สะพาน",                          f"{float(row_ac.get('Y5',0)):.4f}"],
                    ["Y6","ทำความสะอาดระบาย",               f"{float(row_ac.get('Y6',0)):.4f}"],
                    ["Ka","ค่า K ผิวแอสฟัลท์",              f"{float(row_ac.get('K',0)):.4f}"],
                ],
                col_widths=[1.5, 6, 3],
            )
            doc.add_paragraph()
    doc.add_paragraph()

    # ── {base}.2.2  ผิวคอนกรีต (Kc) ────────────────────────────────────────
    h2("ผิวคอนกรีตซีเมนต์ (Kc)")
    add_thai_para(doc,
        "ค่าสัมประสิทธิ์ปรับแก้สำหรับผิวคอนกรีตซีเมนต์ (Kc) คำนวณตามสมการ"
    )
    add_thai_para(doc, "Kc = 1 + 0.50 × (Z1 + Z2 + Z3 + Z4 + Y1 + Y2 + Y3 + Y4 + Y5 + Y6)",
                  first_indent=False)
    if ss["rows_cc"]:
        add_thai_para(doc, "ผลการคำนวณค่า Factor สำหรับสายทางผิวคอนกรีตแสดงดังนี้")
        for row_cc in ss["rows_cc"]:
            route_lbl = f"{row_cc.get('ชื่อสายทาง','')}  (ตอน {row_cc.get('ตอนควบคุม','')})"
            add_thai_para(doc, route_lbl, bold=True, first_indent=False)
            add_table_word(doc,
                headers=["Factor", "คำอธิบาย", "ค่าที่ใช้"],
                rows=[
                    ["Z1","ดัชนีสภาพผิวทาง",               f"{float(row_cc.get('Z1',0)):.4f}"],
                    ["Z2","CBR ดินคันทาง",                  f"{float(row_cc.get('Z2',0)):.4f}"],
                    ["Z3","ปริมาณจราจร AADT",               f"{float(row_cc.get('Z3',0)):.4f}"],
                    ["Z4","ความกว้างผิวทาง",                 f"{float(row_cc.get('Z4',0)):.4f}"],
                    ["Y1","ความกว้างเขตทาง",                 f"{float(row_cc.get('Y1',0)):.4f}"],
                    ["Y2","ไหล่ทางกว้างสุด 1 ข้าง",         f"{float(row_cc.get('Y2',0)):.4f}"],
                    ["Y3","จราจรสงเคราะห์",                  f"{float(row_cc.get('Y3',0)):.4f}"],
                    ["Y4","ท่อระบายน้ำ",                     f"{float(row_cc.get('Y4',0)):.4f}"],
                    ["Y5","สะพาน",                           f"{float(row_cc.get('Y5',0)):.4f}"],
                    ["Y6","ทำความสะอาดระบาย",                f"{float(row_cc.get('Y6',0)):.4f}"],
                    ["Kc","ค่า K ผิวคอนกรีต",                f"{float(row_cc.get('K',0)):.4f}"],
                ],
                col_widths=[1.5, 6, 3],
            )
            doc.add_paragraph()
    doc.add_paragraph()

    # ── {base}.2.3  ผิวลูกรัง (Ks) ─────────────────────────────────────────
    if include_gravel:
        h2("ผิวทางลูกรัง (Ks)")
        add_thai_para(doc,
            "ค่าสัมประสิทธิ์ปรับแก้สำหรับผิวทางลูกรัง (Ks) คำนวณตามสมการ"
        )
        add_thai_para(doc, "Ks = 1 + 0.70 × (A1 + A2 + A3) + 0.30 × (B1 + B2 + B3 + B4)",
                      first_indent=False)
        add_thai_para(doc,
            "ทั้งนี้ ค่า A2 (ลมฟ้าอากาศ) ใช้ค่าเท่ากับ 0.00 "
            "เนื่องจากกรมทางหลวงยังอยู่ระหว่างการศึกษาเก็บสถิติข้อมูล"
        )
        if ss["rows_gr"]:
            add_thai_para(doc, "ผลการคำนวณค่า Factor สำหรับสายทางผิวลูกรังแสดงดังนี้")
            for row_gr in ss["rows_gr"]:
                route_lbl = f"{row_gr.get('ชื่อสายทาง','')}  (ตอน {row_gr.get('ตอนควบคุม','')})"
                add_thai_para(doc, route_lbl, bold=True, first_indent=False)
                add_table_word(doc,
                    headers=["Factor", "คำอธิบาย", "ค่าที่ใช้"],
                    rows=[
                        ["A1","ปริมาณจราจร AADT",            f"{float(row_gr.get('A1',0)):.4f}"],
                        ["A2","ลมฟ้าอากาศ",                  f"{float(row_gr.get('A2',0)):.4f}"],
                        ["A3","ความกว้างผิวทาง",              f"{float(row_gr.get('A3',0)):.4f}"],
                        ["B1","ความกว้างเขตทาง",              f"{float(row_gr.get('B1',0)):.4f}"],
                        ["B2","ภูมิประเทศ (จราจรสงเคราะห์)", f"{float(row_gr.get('B2',0)):.4f}"],
                        ["B3","ภูมิประเทศ (ท่อระบายน้ำ)",     f"{float(row_gr.get('B3',0)):.4f}"],
                        ["B4","สะพาน",                        f"{float(row_gr.get('B4',0)):.4f}"],
                        ["Ks","ค่า K ผิวลูกรัง",              f"{float(row_gr.get('K',0)):.4f}"],
                    ],
                    col_widths=[1.5, 6, 3],
                )
                doc.add_paragraph()
        doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # {base}.3  วิธีการคำนวณงบประมาณงานบำรุงปกติรายปี
    #           (รวมการปรับค่า K' ตามช่วงระยะเวลาประกันไว้ในหัวข้อเดียวกัน)
    # ════════════════════════════════════════════════════════════════════════
    h1("วิธีการคำนวณงบประมาณงานบำรุงปกติรายปี")
    add_thai_para(doc,
        "งบประมาณงานบำรุงปกติรายปีของแต่ละสายทางคำนวณจากผลคูณของระยะทาง "
        "ค่าสัมประสิทธิ์ปรับแก้ (K) ค่า Factor วัสดุ (Km) และอัตราค่าบำรุงมาตรฐาน (N) "
        "โดยปัดค่าที่ได้เป็นหลักร้อยบาทตามแนวทางของกรมทางหลวง ดังสมการ"
    )
    add_thai_para(doc, "งบประมาณ (บาท/ปี)  =  ระยะทาง (กม.) × K × Km × N  (ปัดเป็นหลักร้อย)",
                  first_indent=False)
    add_thai_para(doc,
        "ปริมาณงาน (Workload) คำนวณจากระยะเทียบเท่า (กม.) คูณด้วยค่า K' "
        "ซึ่งเป็น K ที่ปรับแล้วตามช่วงระยะเวลาประกัน โดยระยะเทียบเท่าคำนวณตามสมการ"
    )
    add_thai_para(doc, "ระยะเทียบเท่า (กม.) =  ระยะจริง (กม.) × (จำนวนช่องจราจร / 2)",
                  first_indent=False)
    add_thai_para(doc, "Workload (หน่วย)    =  ระยะเทียบเท่า (กม.) × K'",
                  first_indent=False)
    doc.add_paragraph()
    # ── การปรับค่า K' ตามช่วงระยะเวลาประกัน (รวมอยู่ใน 3.5.3) ─────────────
    add_thai_para(doc,
        "ในกรณีที่สายทางอยู่ในช่วงระยะเวลาประกันผลงาน ค่าสัมประสิทธิ์ที่ใช้คำนวณ Workload "
        "จะถูกลดทอนลงตามเงื่อนไขที่กำหนด เพื่อสะท้อนภาระงานที่ผู้รับจ้างยังรับผิดชอบอยู่ "
        "ดังแสดงในตาราง",
        bold=False,
    )
    add_table_word(doc,
        headers=["เงื่อนไข", "สูตร", "หมายเหตุ"],
        rows=[
            ["ไม่มีประกัน",      "K' = K",          "ใช้ค่า K เต็ม"],
            ["มีประกัน 1 ปี",   "K' = 0.50 × K",   "ลด 50%"],
            ["มีประกัน > 1 ปี", "K' = 0.25 × K",   "ลด 75%"],
        ],
        col_widths=[4.5, 5, 6.5],
    )
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # {base}.4  ผลการคำนวณงบประมาณและ Workload สำหรับผิวทางต่างๆ
    #           3.5.4.1  คำนวณงบประมาณผิวแอสฟัลท์ (Ka)
    #           3.5.4.2  คำนวณงบประมาณผิวคอนกรีต (Kc)
    #           3.5.4.3  คำนวณงบประมาณผิวลูกรัง (Ks)
    # ════════════════════════════════════════════════════════════════════════
    h1("ผลการคำนวณงบประมาณและ Workload สำหรับผิวทางต่างๆ")
    add_thai_para(doc,
        "ตารางต่อไปนี้แสดงผลการคำนวณงบประมาณรายปีและปริมาณงาน (Workload) "
        "สำหรับสายทางแต่ละประเภทผิวทาง โดยแยกตามผิวแอสฟัลท์คอนกรีต "
        "ผิวคอนกรีตซีเมนต์ และผิวลูกรัง"
    )
    doc.add_paragraph()

    def write_surface_section_con(budget_label, rows_data):
        """ตารางงบประมาณ — h2() เรียกเสมอ, ไม่มีคอลัมน์ตอนควบคุม, ไม่มีแถวรวม"""
        h2(f"คำนวณงบประมาณ{budget_label}")
        if not rows_data:
            add_thai_para(doc, "(ไม่มีข้อมูลสายทางประเภทนี้)", first_indent=False)
            doc.add_paragraph()
            return
        headers = ["ชื่อสายทาง", "ระยะทาง\n(กม.)", "ช่องจราจร",
                   "ระยะเทียบเท่า\n(กม.)", "K", "ประกัน\n(ปี)", "K'",
                   "Workload\n(หน่วย)", "งบประมาณ\n(บาท/ปี)"]
        table_rows = []
        for r in rows_data:
            table_rows.append([
                r.get("ชื่อสายทาง",""),
                f"{r.get('ระยะทาง(กม.)',0):.3f}",
                r.get("ช่องจราจร",2),
                f"{r.get('ระยะเทียบเท่า(กม.)',0):.3f}",
                f"{r.get('K',0):.4f}",
                r.get("ประกัน(ปี)",0),
                f"{r.get(chr(75)+chr(39),0):.4f}",
                f"{r.get('Workload(หน่วย)',0):.3f}",
                f"{r.get('งบประมาณ(บาท/ปี)',0):,.0f}",
            ])
        tbl = add_table_word(doc, headers=headers, rows=table_rows,
                             col_widths=[5, 2, 1.5, 2.5, 2, 1.5, 2, 2.5, 3])
        tot_dist  = sum(r.get("ระยะทาง(กม.)",0)        for r in rows_data)
        tot_equiv = sum(r.get("ระยะเทียบเท่า(กม.)",0)  for r in rows_data)
        tot_wl    = sum(r.get("Workload(หน่วย)",0)      for r in rows_data)
        tot_bud   = sum(r.get("งบประมาณ(บาท/ปี)",0)    for r in rows_data)
        add_total_row_word(tbl, [
            "รวม", f"{tot_dist:.3f}", "",
            f"{tot_equiv:.3f}", "", "", "",
            f"{tot_wl:.3f}", f"{tot_bud:,.0f}",
        ])
        doc.add_paragraph()

    write_surface_section_con("ผิวแอสฟัลท์ (Ka)", ss["rows_ac"])
    write_surface_section_con("ผิวคอนกรีต (Kc)",  ss["rows_cc"])
    write_surface_section_con("ผิวลูกรัง (Ks)",    ss["rows_gr"] if include_gravel else [])

    # ════════════════════════════════════════════════════════════════════════
    # {base}.5  สรุปผลการคำนวณ
    # ════════════════════════════════════════════════════════════════════════
    n_ac = len(ss["rows_ac"]); n_cc = len(ss["rows_cc"]); n_gr = len(ss["rows_gr"])
    bud_ac  = sum(r["งบประมาณ(บาท/ปี)"] for r in ss["rows_ac"])
    bud_cc  = sum(r["งบประมาณ(บาท/ปี)"] for r in ss["rows_cc"])
    bud_gr  = sum(r["งบประมาณ(บาท/ปี)"] for r in ss["rows_gr"])
    wl_ac   = sum(r["Workload(หน่วย)"]   for r in ss["rows_ac"])
    wl_cc   = sum(r["Workload(หน่วย)"]   for r in ss["rows_cc"])
    wl_gr   = sum(r["Workload(หน่วย)"]   for r in ss["rows_gr"])
    dist_ac = sum(r["ระยะทาง(กม.)"]      for r in ss["rows_ac"])
    dist_cc = sum(r["ระยะทาง(กม.)"]      for r in ss["rows_cc"])
    dist_gr = sum(r["ระยะทาง(กม.)"]      for r in ss["rows_gr"])

    def rpc(bud, dist): return f"{bud/dist:,.2f}" if dist > 0 else "-"

    bud_total  = bud_ac + bud_cc + (bud_gr if include_gravel else 0)
    dist_total = dist_ac + dist_cc + (dist_gr if include_gravel else 0)
    wl_total   = wl_ac + wl_cc + (wl_gr if include_gravel else 0)
    n_total    = n_ac + n_cc + (n_gr if include_gravel else 0)

    # ── ตารางสรุป: ไม่มี "จำนวนสายทาง" และไม่มีแถว "รวมทุกประเภท" ──
    sum_rows = [
        ["ผิวแอสฟัลท์ (Ka)", f"{dist_ac:.3f}", f"{wl_ac:.3f}", f"{bud_ac:,.0f}", rpc(bud_ac, dist_ac)],
        ["ผิวคอนกรีต (Kc)",  f"{dist_cc:.3f}", f"{wl_cc:.3f}", f"{bud_cc:,.0f}", rpc(bud_cc, dist_cc)],
    ]
    if include_gravel:
        sum_rows.append(["ผิวลูกรัง (Ks)", f"{dist_gr:.3f}", f"{wl_gr:.3f}", f"{bud_gr:,.0f}", rpc(bud_gr, dist_gr)])

    h1("สรุปผลการคำนวณ")
    # ── กำหนดป้ายชื่อประเภทผิวทางที่มีข้อมูล ──
    surf_types = []
    if n_ac > 0: surf_types.append("ผิวทางลาดยาง")
    if n_cc > 0: surf_types.append("ผิวทางคอนกรีต")
    if include_gravel and n_gr > 0: surf_types.append("ผิวทางลูกรัง")
    surf_label_str = "และ".join(surf_types) if surf_types else "ทุกประเภท"
    n_types = len(surf_types)

    add_thai_para(doc,
        f"จากการคำนวณงบประมาณงานบำรุงปกติสำหรับสายทางทั้งหมด {n_types} รูปแบบ "
        f"คือ{surf_label_str} รายละเอียดดังแสดงในตารางสรุป"
    )
    add_table_word(doc,
        headers=["ประเภทผิวทาง", "ระยะทาง (กม.)",
                 "Workload (หน่วย)", "งบประมาณ (บาท/ปี)", "อัตรา (บาท/กม./ปี)"],
        rows=sum_rows,
        col_widths=[4.5, 3, 3.5, 3.5, 3.5],
    )
    doc.add_paragraph()

    # ── หมายเหตุท้ายรายงาน ──────────────────────────────────────────────────
    doc.add_paragraph()
    add_heading_word(doc, "หมายเหตุ", level=2)
    notes = [
        "การคำนวณอ้างอิงคู่มือการคิดค่าปริมาณงานและงานบำรุงปกติ กองบำรุง กรมทางหลวง มกราคม พ.ศ. 2538",
        "ค่า A2 (ลมฟ้าอากาศ) สำหรับผิวลูกรัง ใช้ค่า 0.00 เนื่องจากกรมทางหลวงยังอยู่ระหว่างการศึกษาเก็บสถิติ",
        "งบประมาณปัดเป็นหลักร้อย ตามแนวทางของกรมทางหลวง",
        "ระยะเทียบเท่า = ระยะจริง × (จำนวนช่องจราจร / 2)",
    ]
    for note in notes:
        p = add_thai_para(doc, first_indent=False)
        run_n = p.add_run(f"- {note}")
        set_run_font(run_n, size=Pt(13))

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


# ─────────────────────────────────────────────
# CSS
# ─────────────────────────────────────────────

st.markdown("""
<style>
.main-title {font-size:1.5rem; font-weight:600; margin-bottom:0.2rem;}
.sub-title  {font-size:0.9rem; color:#666; margin-bottom:1.2rem;}
.k-value    {font-size:2rem; font-weight:700; color:#1a6b3c; text-align:center;}
.k-label    {font-size:0.8rem; color:#555; text-align:center; margin-top:-0.3rem;}
.budget-val {font-size:1.4rem; font-weight:700; color:#1565c0; text-align:center;}
.factor-item{display:flex; justify-content:space-between; padding:3px 0;
             border-bottom:1px solid #eee; font-size:0.82rem;}
.factor-name{color:#555;}.factor-val{font-weight:600;}
.note-box   {background:#fff8e1; border-left:4px solid #f9a825;
             padding:0.6rem 1rem; border-radius:4px; font-size:0.82rem; color:#5d4037;}
</style>
""", unsafe_allow_html=True)

st.markdown('<div class="main-title">🛣️ ระบบคำนวณงานบำรุงปกติ</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-title"> พัฒนาโดย รศ.ดร.อิทธิพล มีผล ภาควิชาครุศาสตร์โยธา มจพ.</div>', unsafe_allow_html=True)

tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "⚙️ ตั้งค่าโครงการ",
    "🟫 ผิวแอสฟัลท์ (Ka)",
    "🟩 ผิวคอนกรีต (Kc)",
    "🟨 ผิวลูกรัง (Ks)",
    "📊 สรุปรวม & Export",
])

# ─────────────────────────────────────────────
# TAB 1: PROJECT SETTINGS + JSON LOAD
# ─────────────────────────────────────────────

with tab1:

    # ── JSON Load Section (กฎที่ 1: md5 hash + json_version) ──
    with st.expander("💾 บันทึก / โหลดโครงการ (JSON)", expanded=False):
        col_save, col_load = st.columns(2)

        with col_save:
            st.markdown("**บันทึกโครงการ**")
            json_bytes = json.dumps(build_json_data(), ensure_ascii=False, indent=2).encode("utf-8")
            st.download_button(
                "📥 ดาวน์โหลด JSON",
                data=json_bytes,
                file_name=f"routine_{st.session_state['year']}_{st.session_state['project_name'].replace(' ','_')}.json",
                mime="application/json",
                key="json_dl_tab1",
            )

        with col_load:
            st.markdown("**โหลดโครงการ**")
            uploaded_json = st.file_uploader("เลือกไฟล์ JSON", type=["json"], key="upload_json_tab1")
            if uploaded_json is not None:
                file_bytes = uploaded_json.read()
                file_hash  = hashlib.md5(file_bytes).hexdigest()
                loaded_data = json.loads(file_bytes.decode("utf-8"))
                st.info(f"📌 ไฟล์: **{uploaded_json.name}**  |  บันทึกเมื่อ: {loaded_data.get('saved_at','?')}")
                if st.button("📥 นำเข้าข้อมูล", key="import_json_tab1"):
                    # กฎที่ 1: ป้องกัน load ซ้ำด้วย hash
                    if st.session_state.get("loaded_json_hash") != file_hash:
                        st.session_state["loaded_data"]      = loaded_data
                        st.session_state["loaded_json_hash"] = file_hash
                        # กฎที่ 1: เพิ่ม version → widget keys เปลี่ยน → อ่าน value= ใหม่
                        st.session_state["json_version"] = st.session_state.get("json_version", 0) + 1
                        # อัปเดต rows และ settings ตรง
                        for k in ["project_name","district","year",
                                  "Na","Ns","Nc","Km_a","Km_s","Km_c",
                                  "rows_ac","rows_cc","rows_gr"]:
                            if k in loaded_data:
                                st.session_state[k] = loaded_data[k]
                    st.rerun()

    st.markdown("---")

    # กฎที่ 2: ดึง version ก่อนสร้าง widget ทุกตัว
    v = st.session_state.get("json_version", 0)

    st.markdown("### ข้อมูลโครงการ")
    c1, c2, c3 = st.columns(3)
    with c1:
        pn = st.text_input("ชื่อโครงการ",
                           value=get_default("project_name", "โครงการบำรุงรักษาทางหลวง"),
                           key=f"project_name_v{v}")
        st.session_state["project_name"] = pn
    with c2:
        dt = st.text_input("แขวงการทาง / สำนักงาน",
                           value=get_default("district", ""),
                           key=f"district_v{v}")
        st.session_state["district"] = dt
    with c3:
        yr = st.text_input("ปีงบประมาณ (พ.ศ.)",
                           value=get_default("year", "2568"),
                           key=f"year_v{v}")
        st.session_state["year"] = yr

    st.markdown("---")
    st.markdown("### อัตราค่าบำรุงทางมาตรฐาน (N) และค่า Factor วัสดุ (Km)")
    st.markdown('<div class="note-box">💡 ค่าเริ่มต้นตามคู่มือกรมทางหลวง พ.ศ. 2538 — แก้ไขได้ตามปีงบประมาณปัจจุบัน</div>', unsafe_allow_html=True)
    st.markdown("")

    col_na, col_ns, col_nc = st.columns(3)
    with col_na:
        st.markdown("**ผิวแอสฟัลท์**")
        na = st.number_input("Na (บาท/กม./ปี)",
                             value=float(get_default("Na", 35000.0)),
                             min_value=0.0, step=500.0, format="%.0f",
                             key=f"Na_v{v}")
        st.session_state["Na"] = na
        kma = st.number_input("Km วัสดุ (ลาดยาง)",
                              value=float(get_default("Km_a", 1.0)),
                              min_value=0.01, step=0.01, format="%.3f",
                              key=f"Km_a_v{v}")
        st.session_state["Km_a"] = kma
    with col_ns:
        st.markdown("**ผิวลูกรัง**")
        ns = st.number_input("Ns (บาท/กม./ปี)",
                             value=float(get_default("Ns", 6500.0)),
                             min_value=0.0, step=500.0, format="%.0f",
                             key=f"Ns_v{v}")
        st.session_state["Ns"] = ns
        kms = st.number_input("Km วัสดุ (ลูกรัง)",
                              value=float(get_default("Km_s", 1.0)),
                              min_value=0.01, step=0.01, format="%.3f",
                              key=f"Km_s_v{v}")
        st.session_state["Km_s"] = kms
    with col_nc:
        st.markdown("**ผิวคอนกรีต**")
        nc = st.number_input("Nc (บาท/กม./ปี)",
                             value=float(get_default("Nc", 35000.0)),
                             min_value=0.0, step=500.0, format="%.0f",
                             key=f"Nc_v{v}")
        st.session_state["Nc"] = nc
        kmc = st.number_input("Km วัสดุ (คอนกรีต)",
                              value=float(get_default("Km_c", 1.0)),
                              min_value=0.01, step=0.01, format="%.3f",
                              key=f"Km_c_v{v}")
        st.session_state["Km_c"] = kmc

    st.markdown("---")
    st.markdown("### K' ปรับตามช่วงประกัน (Warranty Adjustment)")
    st.info("**K' = K** (ไม่มีประกัน)　|　**K' = 0.5K** (มีประกัน 1 ปี)　|　**K' = 0.25K** (มีประกัน > 1 ปี)\n\nWorkload = ระยะเทียบเท่า (กม.) × K'")

# ─────────────────────────────────────────────
# HELPER: Factor display + Y-factors form
# ─────────────────────────────────────────────

def show_factor_breakdown(factors, K, K_prime, budget, workload, surf_type):
    col_k, col_b, col_w = st.columns(3)
    with col_k:
        kp_label = f"  |  K' = {K_prime:.4f}" if K != K_prime else ""
        st.markdown(f'<div class="k-value">{K:.4f}</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="k-label">K {surf_type}{kp_label}</div>', unsafe_allow_html=True)
    with col_b:
        st.markdown(f'<div class="budget-val">{budget:,.0f}</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="k-label">งบประมาณ (บาท/ปี)</div>', unsafe_allow_html=True)
    with col_w:
        st.markdown(f'<div class="budget-val" style="color:#6a1b9a">{workload:.3f}</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="k-label">Workload (หน่วย)</div>', unsafe_allow_html=True)
    with st.expander("📋 รายละเอียด Factor"):
        fc1, fc2 = st.columns(2)
        items = list(factors.items())
        half = len(items)//2 + len(items)%2
        with fc1:
            for k, val in items[:half]:
                st.markdown(f'<div class="factor-item"><span class="factor-name">{k}</span><span class="factor-val">{val:.4f}</span></div>', unsafe_allow_html=True)
        with fc2:
            for k, val in items[half:]:
                st.markdown(f'<div class="factor-item"><span class="factor-name">{k}</span><span class="factor-val">{val:.4f}</span></div>', unsafe_allow_html=True)

def y_factors_form(prefix):
    """Factor Y สำหรับทั้ง Ka และ Kc"""
    c1, c2, c3 = st.columns(3)
    with c1:
        y1 = st.number_input("Y1 กว้างเขตทาง (ม.)",       value=40.0, min_value=0.0, step=5.0,  key=f"{prefix}_y1")
        y2 = st.number_input("Y2 ไหล่ทาง 1 ข้าง (ม.)",    value=2.50, min_value=0.0, step=0.25, key=f"{prefix}_y2")
    with c2:
        y3_lbl = st.selectbox("Y3 จราจรสงเคราะห์",  TERRAIN_KEYS, key=f"{prefix}_y3")
        y4_lbl = st.selectbox("Y4 ท่อระบายน้ำ",      TERRAIN_KEYS, key=f"{prefix}_y4")
    with c3:
        y5 = st.number_input("Y5 สะพาน (ม./กม.)",          value=0.0,  min_value=0.0, step=1.0,  key=f"{prefix}_y5")
        y6_lbl = st.selectbox("Y6 ทำความสะอาดระบาย", TERRAIN_KEYS, key=f"{prefix}_y6")
    return y1, y2, TERRAIN_MAP[y3_lbl], TERRAIN_MAP[y4_lbl], y5, TERRAIN_MAP[y6_lbl]

# ─────────────────────────────────────────────
# TAB 2: ASPHALT (Ka)
# ─────────────────────────────────────────────

with tab2:
    input_mode_ac = st.radio("วิธีป้อนข้อมูล",
                             ["✏️ กรอกทีละสายทาง", "📂 Upload Excel"],
                             horizontal=True, key="mode_ac")
    st.markdown("---")

    if input_mode_ac == "✏️ กรอกทีละสายทาง":
        st.markdown("#### ข้อมูลสายทาง")
        ca1, ca2, ca3, ca4 = st.columns(4)
        with ca1:
            route_id_ac   = st.text_input("ตอนควบคุม",  key="ac_route_id")
            route_name_ac = st.text_input("ชื่อสายทาง", key="ac_route_name")
        with ca2:
            km_s_ac = st.number_input("กม. เริ่มต้น", value=0.000, step=0.001, format="%.3f", key="ac_km_s")
            km_e_ac = st.number_input("กม. สิ้นสุด",  value=1.000, step=0.001, format="%.3f", key="ac_km_e")
        with ca3:
            lanes_ac   = st.number_input("ช่องจราจร",          value=2, min_value=1, step=1, key="ac_lanes")
            warranty_ac= st.number_input("มีประกันอีก (ปี)",   value=0, min_value=0, step=1, key="ac_warranty")
        with ca4:
            dist_ac = km_e_ac - km_s_ac
            dist_eq_ac = dist_ac * (lanes_ac / 2)
            st.metric("ระยะทาง (กม.)",       f"{dist_ac:.3f}")
            st.metric("ระยะเทียบเท่า (กม.)", f"{dist_eq_ac:.3f}")

        st.markdown("#### Factor X")
        cx1, cx2, cx3 = st.columns(3)
        with cx1:
            x1_lbl = st.selectbox("X1 ลักษณะผิว+พื้นทาง", list(X1_MAP.keys()), key="ac_x1")
            x1_val = X1_MAP[x1_lbl]
            st.caption(f"X1 = {x1_val:.2f}")
        with cx2:
            x2_cbr = st.number_input("X2 CBR ดินเดิม (%)", value=5.0, min_value=0.0, step=0.5, key="ac_x2")
            st.caption(f"X2 = {lookup_range(x2_cbr, X2_BREAKS):.2f}")
        with cx3:
            x3_lbl = st.selectbox("X3 AADT / 2 ช่อง (คัน/วัน)", list(X3_OPTIONS.keys()), index=5, key="ac_x3")
            x3_val = X3_OPTIONS[x3_lbl]

        cx4, cx5, cx6 = st.columns(3)
        with cx4:
            x4_age = st.number_input("X4 อายุบริการ (ปี)", value=5, min_value=0, step=1, key="ac_x4")
            st.caption(f"X4 = {lookup_range(x4_age, X4_BREAKS):.2f}")
        with cx5:
            x5_wid = st.number_input("X5 กว้างผิว / 2 ช่อง (ม.)", value=7.0, min_value=0.0, step=0.5, key="ac_x5")
            st.caption(f"X5 = {lookup_range(x5_wid, X5_BREAKS):.2f}")
        with cx6:
            x6_lbl = st.selectbox("X6 ภูมิประเทศ", TERRAIN_KEYS, key="ac_x6")
            st.caption(f"X6 = {X6_MAP[TERRAIN_MAP[x6_lbl]]:.2f}")

        st.markdown("#### Factor Y")
        y1_a, y2_a, y3_a, y4_a, y5_a, y6_a = y_factors_form("ac")

        Ka, fac_ac = calc_Ka(x1_val, x2_cbr, x3_val, x4_age, x5_wid,
                             TERRAIN_MAP[x6_lbl], y1_a, y2_a, y3_a, y4_a, y5_a, y6_a)
        Kap = calc_K_prime(Ka, warranty_ac)
        bud_ac_  = calc_budget(dist_ac, Ka, st.session_state["Km_a"], st.session_state["Na"])
        wl_ac_   = calc_workload(dist_eq_ac, Kap)

        st.markdown("---")
        st.markdown("#### ผลการคำนวณ")
        show_factor_breakdown(fac_ac, Ka, Kap, bud_ac_, wl_ac_, "แอสฟัลท์")

        if st.button("➕ เพิ่มสายทางนี้", type="primary", key="add_ac"):
            if dist_ac <= 0:
                st.error("ระยะทางต้องมากกว่า 0")
            else:
                st.session_state["rows_ac"].append({
                    "ตอนควบคุม": route_id_ac, "ชื่อสายทาง": route_name_ac,
                    "กม.เริ่ม": km_s_ac, "กม.สิ้นสุด": km_e_ac,
                    "ระยะทาง(กม.)": round(dist_ac,3), "ช่องจราจร": lanes_ac,
                    "ระยะเทียบเท่า(กม.)": round(dist_eq_ac,3),
                    **{k: float(val) for k, val in fac_ac.items()},
                    "K": round(Ka,4), "ประกัน(ปี)": warranty_ac,
                    "K'": round(Kap,4), "Workload(หน่วย)": wl_ac_,
                    "งบประมาณ(บาท/ปี)": bud_ac_,
                })
                st.success(f"เพิ่มสายทาง '{route_name_ac}' แล้ว")
    else:
        st.markdown('<div class="note-box">📌 คอลัมน์ที่ต้องการ: ตอนควบคุม, ชื่อสายทาง, ระยะจริง(กม.), ช่องจราจร, X1(h/i/l), X2(CBR%), X3(AADT), X4(อายุ), X5(กว้าง), X6(P/R/RM/S), Y1-Y6, ประกัน(ปี)</div>', unsafe_allow_html=True)
        uploaded_ac = st.file_uploader("เลือกไฟล์ Excel", type=["xlsx"], key="up_ac")
        if uploaded_ac:
            try:
                df_up = pd.read_excel(uploaded_ac, header=2)
                st.dataframe(df_up.head(10), use_container_width=True)
                if st.button("⚙️ คำนวณ Ka จากไฟล์", key="calc_up_ac"):
                    new_rows = []
                    for _, r in df_up.iterrows():
                        try:
                            dist = float(r.get("ระยะจริง\n(กม.)", 0) or 0)
                            if dist <= 0: continue
                            ln   = int(r.get("ช่อง\nจราจร", 2) or 2)
                            eq   = dist * (ln / 2)
                            x1v  = {"h":0.0,"i":0.5,"l":1.0}.get(str(r.get("X1","h")).lower(), 0.0)
                            cbr  = float(r.get("X2",5) or 5)
                            x3f  = lookup_list(int(r.get("X3",500) or 500), X3_LOWER, X3_UPPER, X3_VAL)
                            age  = int(r.get("X4",5) or 5)
                            wid  = float(r.get("X5",7) or 7)
                            x6t  = str(r.get("X6","P") or "P").upper()
                            y1v  = float(r.get("Y1",40) or 40)
                            y2v  = float(r.get("Y2",2.5) or 2.5)
                            y3t  = str(r.get("Y3","P") or "P").upper()
                            y4t  = str(r.get("Y4","P") or "P").upper()
                            y5v  = float(r.get("Y5",0) or 0)
                            y6t  = str(r.get("Y6","P") or "P").upper()
                            war  = int(r.get("ประกัน",0) or 0)
                            Ka_, f_ = calc_Ka(x1v,cbr,x3f,age,wid,x6t,y1v,y2v,y3t,y4t,y5v,y6t)
                            Kap_ = calc_K_prime(Ka_, war)
                            bud_ = calc_budget(dist, Ka_, st.session_state["Km_a"], st.session_state["Na"])
                            wl_  = calc_workload(eq, Kap_)
                            new_rows.append({"ตอนควบคุม":r.iloc[0],"ชื่อสายทาง":r.iloc[1],
                                "ระยะทาง(กม.)":round(dist,3),"ช่องจราจร":ln,
                                "ระยะเทียบเท่า(กม.)":round(eq,3),
                                "K":round(Ka_,4),"ประกัน(ปี)":war,
                                "K'":round(Kap_,4),"Workload(หน่วย)":wl_,"งบประมาณ(บาท/ปี)":bud_})
                        except Exception: continue
                    st.session_state["rows_ac"].extend(new_rows)
                    st.success(f"เพิ่ม {len(new_rows)} สายทางแล้ว")
            except Exception as e:
                st.error(f"อ่านไฟล์ไม่ได้: {e}")

    if st.session_state["rows_ac"]:
        st.markdown("---")
        st.markdown(f"#### 📋 ตารางผิวแอสฟัลท์ ({len(st.session_state['rows_ac'])} สายทาง)")
        # Header row
        hc = st.columns([2, 3, 2, 1.5, 1.5, 2, 2.5, 1.2])
        for txt in ["ตอนควบคุม","ชื่อสายทาง","ระยะทาง(กม.)","K","K'","Workload(หน่วย)","งบประมาณ(บาท/ปี)","ลบ"]:
            hc[["ตอนควบคุม","ชื่อสายทาง","ระยะทาง(กม.)","K","K'","Workload(หน่วย)","งบประมาณ(บาท/ปี)","ลบ"].index(txt)].markdown(f"**{txt}**")
        st.markdown('<hr style="margin:2px 0 6px 0">', unsafe_allow_html=True)
        for i, r in enumerate(st.session_state["rows_ac"]):
            rc = st.columns([2, 3, 2, 1.5, 1.5, 2, 2.5, 1.2])
            rc[0].write(r.get("ตอนควบคุม",""))
            rc[1].write(r.get("ชื่อสายทาง",""))
            rc[2].write(f"{r.get('ระยะทาง(กม.)',0):.3f}")
            rc[3].write(f"{r.get('K',0):.4f}")
            kp_val = r.get("K'",0)
            rc[4].write(f"{kp_val:.4f}")
            rc[5].write(f"{r.get('Workload(หน่วย)',0):.3f}")
            rc[6].write(f"{r.get('งบประมาณ(บาท/ปี)',0):,.0f}")
            if rc[7].button("🗑️", key=f"del_ac_{i}", help=f"ลบสายทาง #{i+1}"):
                st.session_state["rows_ac"].pop(i); st.rerun()
        st.markdown("")
        m1, m2, m3 = st.columns(3)
        m1.metric("งบประมาณรวม (บาท/ปี)", f"{sum(r['งบประมาณ(บาท/ปี)'] for r in st.session_state['rows_ac']):,.0f}")
        m2.metric("Workload รวม (หน่วย)",  f"{sum(r['Workload(หน่วย)']   for r in st.session_state['rows_ac']):.3f}")
        if m3.button("🗑️ ล้างทั้งหมด", key="clr_ac"):
            st.session_state["rows_ac"] = []; st.rerun()

# ─────────────────────────────────────────────
# TAB 3: CONCRETE (Kc)
# ─────────────────────────────────────────────

with tab3:
    input_mode_cc = st.radio("วิธีป้อนข้อมูล",
                             ["✏️ กรอกทีละสายทาง", "📂 Upload Excel (KC.xlsx Sheet C)"],
                             horizontal=True, key="mode_cc")
    st.markdown("---")

    if input_mode_cc == "✏️ กรอกทีละสายทาง":
        st.markdown("#### ข้อมูลสายทาง")
        cc1, cc2, cc3, cc4 = st.columns(4)
        with cc1:
            route_id_cc   = st.text_input("ตอนควบคุม",  key="cc_route_id")
            route_name_cc = st.text_input("ชื่อสายทาง", key="cc_route_name")
        with cc2:
            km_s_cc = st.number_input("กม. เริ่มต้น", value=0.000, step=0.001, format="%.3f", key="cc_km_s")
            km_e_cc = st.number_input("กม. สิ้นสุด",  value=1.000, step=0.001, format="%.3f", key="cc_km_e")
        with cc3:
            lanes_cc    = st.number_input("ช่องจราจร",        value=2, min_value=1, step=1, key="cc_lanes")
            warranty_cc = st.number_input("มีประกันอีก (ปี)", value=0, min_value=0, step=1, key="cc_warranty")
        with cc4:
            dist_cc = km_e_cc - km_s_cc
            dist_eq_cc = dist_cc * (lanes_cc / 2)
            st.metric("ระยะทาง (กม.)",       f"{dist_cc:.3f}")
            st.metric("ระยะเทียบเท่า (กม.)", f"{dist_eq_cc:.3f}")

        st.markdown("#### Factor Z")
        cz1, cz2, cz3, cz4 = st.columns(4)
        with cz1:
            z1_idx = st.selectbox("Z1 ดัชนีสภาพผิว (1-8)", list(range(1,9)), index=0, key="cc_z1")
            st.caption(f"Z1 = {Z1_MAP[z1_idx]:.2f}")
        with cz2:
            z2_cbr = st.number_input("Z2 CBR ดินคันทาง (%)", value=5.0, min_value=0.0, step=0.5, key="cc_z2")
            st.caption(f"Z2 = {lookup_range(z2_cbr, Z2_BREAKS):.2f}")
        with cz3:
            z3_lbl = st.selectbox("Z3 AADT / 2 ช่อง (คัน/วัน)", list(Z3_OPTIONS.keys()), index=4, key="cc_z3")
            z3_val = Z3_OPTIONS[z3_lbl]
        with cz4:
            z4_wid = st.number_input("Z4 กว้างผิว / 2 ช่อง (ม.)", value=7.0, min_value=0.0, step=0.5, key="cc_z4")
            st.caption(f"Z4 = {lookup_range(z4_wid, Z4_BREAKS):.2f}")

        st.markdown("#### Factor Y")
        y1_c, y2_c, y3_c, y4_c, y5_c, y6_c = y_factors_form("cc")

        Kc, fac_cc = calc_Kc(z1_idx, z2_cbr, z3_val, z4_wid,
                              y1_c, y2_c, y3_c, y4_c, y5_c, y6_c)
        Kcp = calc_K_prime(Kc, warranty_cc)
        bud_cc_  = calc_budget(dist_cc, Kc, st.session_state["Km_c"], st.session_state["Nc"])
        wl_cc_   = calc_workload(dist_eq_cc, Kcp)

        st.markdown("---")
        st.markdown("#### ผลการคำนวณ")
        show_factor_breakdown(fac_cc, Kc, Kcp, bud_cc_, wl_cc_, "คอนกรีต")

        if st.button("➕ เพิ่มสายทางนี้", type="primary", key="add_cc"):
            if dist_cc <= 0:
                st.error("ระยะทางต้องมากกว่า 0")
            else:
                st.session_state["rows_cc"].append({
                    "ตอนควบคุม": route_id_cc, "ชื่อสายทาง": route_name_cc,
                    "กม.เริ่ม": km_s_cc, "กม.สิ้นสุด": km_e_cc,
                    "ระยะทาง(กม.)": round(dist_cc,3), "ช่องจราจร": lanes_cc,
                    "ระยะเทียบเท่า(กม.)": round(dist_eq_cc,3),
                    **{k: float(val) for k, val in fac_cc.items()},
                    "K": round(Kc,4), "ประกัน(ปี)": warranty_cc,
                    "K'": round(Kcp,4), "Workload(หน่วย)": wl_cc_,
                    "งบประมาณ(บาท/ปี)": bud_cc_,
                })
                st.success(f"เพิ่มสายทาง '{route_name_cc}' แล้ว")
    else:
        st.markdown('<div class="note-box">📌 เลือก Sheet C ของ KC.xlsx — ดูโครงสร้างคอลัมน์จากไฟล์ต้นฉบับ</div>', unsafe_allow_html=True)
        uploaded_cc = st.file_uploader("เลือกไฟล์ Excel", type=["xlsx"], key="up_cc")
        if uploaded_cc:
            try:
                df_upc = pd.read_excel(uploaded_cc, sheet_name="C", header=2)
                st.dataframe(df_upc.head(10), use_container_width=True)
                if st.button("⚙️ คำนวณ Kc จากไฟล์", key="calc_up_cc"):
                    new_rows_c = []
                    for _, r in df_upc.iterrows():
                        try:
                            dist = float(r.iloc[5] if pd.notna(r.iloc[5]) else 0)
                            if dist <= 0: continue
                            ln   = int(r.iloc[6] if pd.notna(r.iloc[6]) else 2)
                            eq   = dist * (ln / 2)
                            z1i  = int(r.iloc[8]  if pd.notna(r.iloc[8])  else 1)
                            z2c  = float(r.iloc[9]  if pd.notna(r.iloc[9])  else 5)
                            z3a  = int(r.iloc[10] if pd.notna(r.iloc[10]) else 1000)
                            z3f  = lookup_list(z3a, Z3_LOWER, Z3_UPPER, Z3_VAL)
                            z4w  = float(r.iloc[11] if pd.notna(r.iloc[11]) else 7)
                            y1v  = float(r.iloc[12] if pd.notna(r.iloc[12]) else 40)
                            y2v  = float(r.iloc[13] if pd.notna(r.iloc[13]) else 2.5)
                            y3t  = str(r.iloc[14] if pd.notna(r.iloc[14]) else "P")
                            y4t  = str(r.iloc[15] if pd.notna(r.iloc[15]) else "P")
                            y5v  = float(r.iloc[16] if pd.notna(r.iloc[16]) else 0)
                            y6t  = str(r.iloc[17] if pd.notna(r.iloc[17]) else "P")
                            war  = int(r.iloc[21] if pd.notna(r.iloc[21]) else 0)
                            Kc_, f_ = calc_Kc(z1i,z2c,z3f,z4w,y1v,y2v,y3t,y4t,y5v,y6t)
                            Kcp_ = calc_K_prime(Kc_, war)
                            bud_ = calc_budget(dist, Kc_, st.session_state["Km_c"], st.session_state["Nc"])
                            wl_  = calc_workload(eq, Kcp_)
                            new_rows_c.append({"ตอนควบคุม":r.iloc[0],"ชื่อสายทาง":r.iloc[1],
                                "ระยะทาง(กม.)":round(dist,3),"ช่องจราจร":ln,
                                "ระยะเทียบเท่า(กม.)":round(eq,3),
                                "K":round(Kc_,4),"ประกัน(ปี)":war,
                                "K'":round(Kcp_,4),"Workload(หน่วย)":wl_,"งบประมาณ(บาท/ปี)":bud_})
                        except Exception: continue
                    st.session_state["rows_cc"].extend(new_rows_c)
                    st.success(f"เพิ่ม {len(new_rows_c)} สายทางแล้ว")
            except Exception as e:
                st.error(f"อ่านไฟล์ไม่ได้: {e}")

    if st.session_state["rows_cc"]:
        st.markdown("---")
        st.markdown(f"#### 📋 ตารางผิวคอนกรีต ({len(st.session_state['rows_cc'])} สายทาง)")
        hc2 = st.columns([2, 3, 2, 1.5, 1.5, 2, 2.5, 1.2])
        for txt in ["ตอนควบคุม","ชื่อสายทาง","ระยะทาง(กม.)","K","K'","Workload(หน่วย)","งบประมาณ(บาท/ปี)","ลบ"]:
            hc2[["ตอนควบคุม","ชื่อสายทาง","ระยะทาง(กม.)","K","K'","Workload(หน่วย)","งบประมาณ(บาท/ปี)","ลบ"].index(txt)].markdown(f"**{txt}**")
        st.markdown('<hr style="margin:2px 0 6px 0">', unsafe_allow_html=True)
        for i, r in enumerate(st.session_state["rows_cc"]):
            rc2 = st.columns([2, 3, 2, 1.5, 1.5, 2, 2.5, 1.2])
            rc2[0].write(r.get("ตอนควบคุม",""))
            rc2[1].write(r.get("ชื่อสายทาง",""))
            rc2[2].write(f"{r.get('ระยะทาง(กม.)',0):.3f}")
            rc2[3].write(f"{r.get('K',0):.4f}")
            kp_val = r.get("K'",0)
            rc2[4].write(f"{kp_val:.4f}")
            rc2[5].write(f"{r.get('Workload(หน่วย)',0):.3f}")
            rc2[6].write(f"{r.get('งบประมาณ(บาท/ปี)',0):,.0f}")
            if rc2[7].button("🗑️", key=f"del_cc_{i}", help=f"ลบสายทาง #{i+1}"):
                st.session_state["rows_cc"].pop(i); st.rerun()
        st.markdown("")
        m1c, m2c, m3c = st.columns(3)
        m1c.metric("งบประมาณรวม (บาท/ปี)", f"{sum(r['งบประมาณ(บาท/ปี)'] for r in st.session_state['rows_cc']):,.0f}")
        m2c.metric("Workload รวม (หน่วย)",  f"{sum(r['Workload(หน่วย)']   for r in st.session_state['rows_cc']):.3f}")
        if m3c.button("🗑️ ล้างทั้งหมด", key="clr_cc"):
            st.session_state["rows_cc"] = []; st.rerun()

# ─────────────────────────────────────────────
# TAB 4: GRAVEL (Ks)
# ─────────────────────────────────────────────

with tab4:
    st.markdown('<div class="note-box">⚠️ A2 (ลักษณะลมฟ้าอากาศ) ใช้ค่า 0.00 — กรมทางหลวงยังอยู่ระหว่างการศึกษาเก็บสถิติ</div>', unsafe_allow_html=True)

    st.checkbox(
        "📄 รวมผิวลูกรังในรายงาน Word",
        value=st.session_state.get("include_gravel_report", False),
        key="include_gravel_report",
        help="หากไม่ติ๊ก ผิวลูกรังจะไม่ปรากฏในรายงาน Word แต่ยังคำนวณและแสดงในตารางสรุปตามปกติ",
    )
    st.markdown("---")

    st.markdown("#### ข้อมูลสายทาง")
    cg1, cg2, cg3, cg4 = st.columns(4)
    with cg1:
        route_id_gr   = st.text_input("ตอนควบคุม",  key="gr_route_id")
        route_name_gr = st.text_input("ชื่อสายทาง", key="gr_route_name")
    with cg2:
        km_s_gr = st.number_input("กม. เริ่มต้น", value=0.000, step=0.001, format="%.3f", key="gr_km_s")
        km_e_gr = st.number_input("กม. สิ้นสุด",  value=1.000, step=0.001, format="%.3f", key="gr_km_e")
    with cg3:
        lanes_gr    = st.number_input("ช่องจราจร",        value=2, min_value=1, step=1, key="gr_lanes")
        warranty_gr = st.number_input("มีประกันอีก (ปี)", value=0, min_value=0, step=1, key="gr_warranty")
    with cg4:
        dist_gr = km_e_gr - km_s_gr
        dist_eq_gr = dist_gr * (lanes_gr / 2)
        st.metric("ระยะทาง (กม.)",       f"{dist_gr:.3f}")
        st.metric("ระยะเทียบเท่า (กม.)", f"{dist_eq_gr:.3f}")

    st.markdown("#### Factor A")
    cga1, cga2, cga3 = st.columns(3)
    with cga1:
        a1_aadt = st.number_input("A1 ADT (คัน/วัน)", value=300, min_value=0, step=10, key="gr_a1")
        st.caption(f"A1 = {lookup_range(a1_aadt, A1_BREAKS):.2f}")
    with cga2:
        st.metric("A2 ลมฟ้าอากาศ", "0.00")
        st.caption("ยังไม่มีข้อมูล — กรมทางหลวงกำลังศึกษา")
    with cga3:
        a3_wid = st.number_input("A3 กว้างคันทาง (ม.)", value=7.0, min_value=0.0, step=0.5, key="gr_a3")
        st.caption(f"A3 = {lookup_range(a3_wid, A3_BREAKS):.2f}")

    st.markdown("#### Factor B")
    cgb1, cgb2, cgb3, cgb4 = st.columns(4)
    with cgb1:
        b1_row = st.number_input("B1 กว้างเขตทาง (ม.)", value=30.0, min_value=0.0, step=5.0, key="gr_b1")
        st.caption(f"B1 = {lookup_range(b1_row, B1_BREAKS_KB):.2f}")
    with cgb2:
        b2_lbl = st.selectbox("B2 จราจรสงเคราะห์", TERRAIN_KEYS, key="gr_b2")
        st.caption(f"B2 = {B2_MAP[TERRAIN_MAP[b2_lbl]]:.2f}")
    with cgb3:
        b3_lbl = st.selectbox("B3 ระบายน้ำ", TERRAIN_KEYS, key="gr_b3")
        st.caption(f"B3 = {B3_MAP[TERRAIN_MAP[b3_lbl]]:.2f}")
    with cgb4:
        b4_br = st.number_input("B4 สะพาน (ม./กม.)", value=0.0, min_value=0.0, step=1.0, key="gr_b4")
        st.caption(f"B4 = {lookup_range(b4_br, B4_BREAKS):.2f}")

    Ks, fac_gr = calc_Ks(a1_aadt, a3_wid, b1_row, TERRAIN_MAP[b2_lbl], TERRAIN_MAP[b3_lbl], b4_br)
    Ksp = calc_K_prime(Ks, warranty_gr)
    bud_gr_  = calc_budget(dist_gr, Ks, st.session_state["Km_s"], st.session_state["Ns"])
    wl_gr_   = calc_workload(dist_eq_gr, Ksp)

    st.markdown("---")
    st.markdown("#### ผลการคำนวณ")
    show_factor_breakdown(fac_gr, Ks, Ksp, bud_gr_, wl_gr_, "ลูกรัง")

    if st.button("➕ เพิ่มสายทางนี้", type="primary", key="add_gr"):
        if dist_gr <= 0:
            st.error("ระยะทางต้องมากกว่า 0")
        else:
            st.session_state["rows_gr"].append({
                "ตอนควบคุม": route_id_gr, "ชื่อสายทาง": route_name_gr,
                "กม.เริ่ม": km_s_gr, "กม.สิ้นสุด": km_e_gr,
                "ระยะทาง(กม.)": round(dist_gr,3), "ช่องจราจร": lanes_gr,
                "ระยะเทียบเท่า(กม.)": round(dist_eq_gr,3),
                **{k: float(val) for k, val in fac_gr.items()},
                "K": round(Ks,4), "ประกัน(ปี)": warranty_gr,
                "K'": round(Ksp,4), "Workload(หน่วย)": wl_gr_,
                "งบประมาณ(บาท/ปี)": bud_gr_,
            })
            st.success(f"เพิ่มสายทาง '{route_name_gr}' แล้ว")

    if st.session_state["rows_gr"]:
        st.markdown("---")
        st.markdown(f"#### 📋 ตารางผิวลูกรัง ({len(st.session_state['rows_gr'])} สายทาง)")
        hc3 = st.columns([2, 3, 2, 1.5, 1.5, 2, 2.5, 1.2])
        for txt in ["ตอนควบคุม","ชื่อสายทาง","ระยะทาง(กม.)","K","K'","Workload(หน่วย)","งบประมาณ(บาท/ปี)","ลบ"]:
            hc3[["ตอนควบคุม","ชื่อสายทาง","ระยะทาง(กม.)","K","K'","Workload(หน่วย)","งบประมาณ(บาท/ปี)","ลบ"].index(txt)].markdown(f"**{txt}**")
        st.markdown('<hr style="margin:2px 0 6px 0">', unsafe_allow_html=True)
        for i, r in enumerate(st.session_state["rows_gr"]):
            rc3 = st.columns([2, 3, 2, 1.5, 1.5, 2, 2.5, 1.2])
            rc3[0].write(r.get("ตอนควบคุม",""))
            rc3[1].write(r.get("ชื่อสายทาง",""))
            rc3[2].write(f"{r.get('ระยะทาง(กม.)',0):.3f}")
            rc3[3].write(f"{r.get('K',0):.4f}")
            kp_val = r.get("K'",0)
            rc3[4].write(f"{kp_val:.4f}")
            rc3[5].write(f"{r.get('Workload(หน่วย)',0):.3f}")
            rc3[6].write(f"{r.get('งบประมาณ(บาท/ปี)',0):,.0f}")
            if rc3[7].button("🗑️", key=f"del_gr_{i}", help=f"ลบสายทาง #{i+1}"):
                st.session_state["rows_gr"].pop(i); st.rerun()
        st.markdown("")
        m1g, m2g, m3g = st.columns(3)
        m1g.metric("งบประมาณรวม (บาท/ปี)", f"{sum(r['งบประมาณ(บาท/ปี)'] for r in st.session_state['rows_gr']):,.0f}")
        m2g.metric("Workload รวม (หน่วย)",  f"{sum(r['Workload(หน่วย)']   for r in st.session_state['rows_gr']):.3f}")
        if m3g.button("🗑️ ล้างทั้งหมด", key="clr_gr"):
            st.session_state["rows_gr"] = []; st.rerun()

# ─────────────────────────────────────────────
# TAB 5: SUMMARY & EXPORT
# ─────────────────────────────────────────────

with tab5:
    st.markdown(f"### 📊 สรุปรวม — {st.session_state['project_name']}")
    st.caption(f"{st.session_state['district']}  |  ปีงบประมาณ {st.session_state['year']}")

    n_ac = len(st.session_state["rows_ac"])
    n_cc = len(st.session_state["rows_cc"])
    n_gr = len(st.session_state["rows_gr"])

    bud_ac  = sum(r["งบประมาณ(บาท/ปี)"] for r in st.session_state["rows_ac"])
    bud_cc  = sum(r["งบประมาณ(บาท/ปี)"] for r in st.session_state["rows_cc"])
    bud_gr  = sum(r["งบประมาณ(บาท/ปี)"] for r in st.session_state["rows_gr"])
    bud_total = bud_ac + bud_cc + bud_gr

    wl_ac  = sum(r["Workload(หน่วย)"] for r in st.session_state["rows_ac"])
    wl_cc  = sum(r["Workload(หน่วย)"] for r in st.session_state["rows_cc"])
    wl_gr  = sum(r["Workload(หน่วย)"] for r in st.session_state["rows_gr"])
    wl_total = wl_ac + wl_cc + wl_gr

    dist_ac_tot = sum(r["ระยะทาง(กม.)"] for r in st.session_state["rows_ac"])
    dist_cc_tot = sum(r["ระยะทาง(กม.)"] for r in st.session_state["rows_cc"])
    dist_gr_tot = sum(r["ระยะทาง(กม.)"] for r in st.session_state["rows_gr"])
    dist_total  = dist_ac_tot + dist_cc_tot + dist_gr_tot

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("ระยะทางรวม (กม.)",     f"{dist_total:.3f}")
    c2.metric("งบประมาณรวม (บาท/ปี)", f"{bud_total:,.0f}")
    c3.metric("Workload รวม (หน่วย)", f"{wl_total:.3f}")
    c4.metric("จำนวนสายทาง",          f"{n_ac+n_cc+n_gr}")

    st.markdown("---")
    st.markdown("#### สรุปแยกตามประเภทผิวทาง")
    st.dataframe(pd.DataFrame({
        "ประเภทผิวทาง":     ["🟫 แอสฟัลท์ (Ka)","🟩 คอนกรีต (Kc)","🟨 ลูกรัง (Ks)","**รวม**"],
        "จำนวนสายทาง":      [n_ac, n_cc, n_gr, n_ac+n_cc+n_gr],
        "ระยะทาง (กม.)":    [round(dist_ac_tot,3), round(dist_cc_tot,3), round(dist_gr_tot,3), round(dist_total,3)],
        "Workload (หน่วย)": [round(wl_ac,3), round(wl_cc,3), round(wl_gr,3), round(wl_total,3)],
        "งบประมาณ (บาท/ปี)":[f"{bud_ac:,.0f}", f"{bud_cc:,.0f}", f"{bud_gr:,.0f}", f"{bud_total:,.0f}"],
    }), use_container_width=True, hide_index=True)

    st.markdown("---")

    # ── 5 ปุ่ม Export ─────────────────────────
    col_e1, col_e2, col_e3, col_e4, col_e5 = st.columns(5)

    # ── Excel ──────────────────────────────────
    with col_e1:
        def generate_excel():
            wb = openpyxl.Workbook()
            hdr_font   = Font(name="TH SarabunPSK", bold=True, size=12, color="FFFFFF")
            data_font  = Font(name="TH SarabunPSK", size=12)
            title_font = Font(name="TH SarabunPSK", bold=True, size=14)
            center     = Alignment(horizontal="center", vertical="center", wrap_text=True)
            left_al    = Alignment(horizontal="left", vertical="center")
            fill_hdr   = PatternFill("solid", fgColor="1565C0")
            fill_tot   = PatternFill("solid", fgColor="FFF9C4")
            thin       = Side(style="thin", color="BBBBBB")
            border     = Border(left=thin, right=thin, top=thin, bottom=thin)

            def write_sheet(ws, rows, surf_type, factors_cols):
                ws.sheet_view.showGridLines = False
                ws.merge_cells("A1:M1")
                ws["A1"] = f"งานบำรุงปกติ — ผิว{surf_type}  |  {st.session_state['project_name']}  |  {st.session_state['district']}  |  ปีงบประมาณ {st.session_state['year']}"
                ws["A1"].font = title_font; ws["A1"].alignment = center
                headers = ["ตอนควบคุม","ชื่อสายทาง","กม.เริ่ม","กม.สิ้นสุด",
                           "ระยะทาง\n(กม.)","ช่องจราจร","ระยะเทียบเท่า\n(กม.)"] + \
                          factors_cols + ["K","ประกัน\n(ปี)","K'","Workload\n(หน่วย)","งบประมาณ\n(บาท/ปี)"]
                for ci, h in enumerate(headers, 1):
                    cell = ws.cell(row=2, column=ci, value=h)
                    cell.font = hdr_font; cell.fill = fill_hdr
                    cell.alignment = center; cell.border = border
                for ri, row in enumerate(rows, 3):
                    vals = [row.get("ตอนควบคุม",""), row.get("ชื่อสายทาง",""),
                            row.get("กม.เริ่ม",0), row.get("กม.สิ้นสุด",0),
                            row.get("ระยะทาง(กม.)",0), row.get("ช่องจราจร",2),
                            row.get("ระยะเทียบเท่า(กม.)",0)] + \
                           [float(row.get(fc,0)) for fc in factors_cols] + \
                           [row.get("K",0), row.get("ประกัน(ปี)",0),
                            row.get("K'",0), row.get("Workload(หน่วย)",0),
                            row.get("งบประมาณ(บาท/ปี)",0)]
                    for ci, val in enumerate(vals, 1):
                        cell = ws.cell(row=ri, column=ci, value=val)
                        cell.font = data_font; cell.border = border
                        cell.alignment = center if isinstance(val,(int,float)) else left_al
                last_col = len(headers)
                if rows:
                    tr = len(rows) + 3
                    for ci in range(1, last_col+1):
                        cell = ws.cell(row=tr, column=ci)
                        cell.fill = fill_tot; cell.border = border
                        cell.font = Font(name="TH SarabunPSK", bold=True, size=12)
                        cell.alignment = center
                    ws.cell(row=tr, column=1, value="รวม")
                    ws.cell(row=tr, column=5, value=sum(r.get("ระยะทาง(กม.)",0) for r in rows))
                    ws.cell(row=tr, column=7, value=sum(r.get("ระยะเทียบเท่า(กม.)",0) for r in rows))
                    ws.cell(row=tr, column=last_col-1, value=round(sum(r.get("Workload(หน่วย)",0) for r in rows),3))
                    ws.cell(row=tr, column=last_col,   value=sum(r.get("งบประมาณ(บาท/ปี)",0) for r in rows))
                    ws.cell(row=tr, column=last_col).number_format = "#,##0"
                ws.column_dimensions["A"].width = 14; ws.column_dimensions["B"].width = 24
                ws.column_dimensions["C"].width = 10; ws.column_dimensions["D"].width = 10
                for i in range(5, last_col+1):
                    ws.column_dimensions[get_column_letter(i)].width = 10
                ws.row_dimensions[1].height = 22; ws.row_dimensions[2].height = 36

            ws_ac = wb.active; ws_ac.title = "ผิวแอสฟัลท์"
            write_sheet(ws_ac, st.session_state["rows_ac"], "แอสฟัลท์",
                        ["X1","X2","X3","X4","X5","X6","Y1","Y2","Y3","Y4","Y5","Y6"])
            ws_cc = wb.create_sheet("ผิวคอนกรีต")
            write_sheet(ws_cc, st.session_state["rows_cc"], "คอนกรีต",
                        ["Z1","Z2","Z3","Z4","Y1","Y2","Y3","Y4","Y5","Y6"])
            ws_gr = wb.create_sheet("ผิวลูกรัง")
            write_sheet(ws_gr, st.session_state["rows_gr"], "ลูกรัง",
                        ["A1","A2","A3","B1","B2","B3","B4"])
            ws_sum = wb.create_sheet("สรุปรวม")
            ws_sum.sheet_view.showGridLines = False
            ws_sum.merge_cells("A1:E1")
            ws_sum["A1"] = f"สรุปงานบำรุงปกติ  |  {st.session_state['project_name']}  |  {st.session_state['district']}  |  ปีงบประมาณ {st.session_state['year']}"
            ws_sum["A1"].font = title_font; ws_sum["A1"].alignment = center
            sum_hdrs = ["ประเภทผิวทาง","จำนวนสายทาง","ระยะทาง (กม.)","Workload (หน่วย)","งบประมาณ (บาท/ปี)"]
            for ci, h in enumerate(sum_hdrs, 1):
                cell = ws_sum.cell(row=2, column=ci, value=h)
                cell.font = hdr_font; cell.fill = fill_hdr
                cell.alignment = center; cell.border = border
            sum_rows_d = [
                ("ผิวแอสฟัลท์ (Ka)", n_ac, round(dist_ac_tot,3), round(wl_ac,3), bud_ac),
                ("ผิวคอนกรีต (Kc)",  n_cc, round(dist_cc_tot,3), round(wl_cc,3), bud_cc),
                ("ผิวลูกรัง (Ks)",   n_gr, round(dist_gr_tot,3), round(wl_gr,3), bud_gr),
                ("รวม", n_ac+n_cc+n_gr, round(dist_total,3), round(wl_total,3), bud_total),
            ]
            for ri, row_data in enumerate(sum_rows_d, 3):
                is_tot = (ri == len(sum_rows_d)+2)
                for ci, val in enumerate(row_data, 1):
                    cell = ws_sum.cell(row=ri, column=ci, value=val)
                    cell.font = Font(name="TH SarabunPSK", bold=is_tot, size=12)
                    cell.fill = fill_tot if is_tot else PatternFill("solid", fgColor="E8F5E9")
                    cell.alignment = center; cell.border = border
                    if ci == 5: cell.number_format = "#,##0"
            for col, w in zip(["A","B","C","D","E"], [22,14,16,16,20]):
                ws_sum.column_dimensions[col].width = w
            buf = BytesIO(); wb.save(buf); buf.seek(0)
            return buf

        if bud_total > 0:
            excel_buf = generate_excel()
            st.download_button(
                "📊 Export Excel",
                data=excel_buf,
                file_name=f"routine_{st.session_state['year']}_{st.session_state['project_name'].replace(' ','_')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                type="primary",
            )
        else:
            st.info("ยังไม่มีข้อมูล")

    # ── Word Report (แบบย่อ) ──────────────────
    with col_e2:
        if bud_total > 0:
            if st.button("📄 รายงาน Word (แบบย่อ)", type="primary", key="make_word"):
                with st.spinner("กำลังสร้างรายงาน..."):
                    try:
                        include_gr = st.session_state.get("include_gravel_report", False)
                        word_buf = generate_word_report(include_gravel=include_gr)
                        st.download_button(
                            "⬇️ ดาวน์โหลด Word (แบบย่อ)",
                            data=word_buf,
                            file_name=f"routine_report_{st.session_state['year']}_{st.session_state['project_name'].replace(' ','_')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="dl_word",
                        )
                    except Exception as e:
                        st.error(f"สร้างรายงานไม่ได้: {e}")
        else:
            st.info("ยังไม่มีข้อมูล")

    # ── Word Report (แบบที่ปรึกษา) ────────────
    with col_e3:
        if bud_total > 0:
            base_sec_input = st.text_input(
                "Base Section (เช่น 3.5)",
                value="3.5",
                key="consultant_base_sec",
                help="ระบุเลข section หลักของรายงานที่ปรึกษา เช่น 3.5 หรือ 4.2"
            )
            if st.button("📋 รายงาน Word (ที่ปรึกษา)", type="primary", key="make_word_con"):
                with st.spinner("กำลังสร้างรายงานแบบที่ปรึกษา..."):
                    try:
                        include_gr = st.session_state.get("include_gravel_report", False)
                        word_buf_con = generate_word_report_consultant(
                            include_gravel=include_gr,
                            base_sec=base_sec_input.strip() or "3.5",
                        )
                        st.download_button(
                            "⬇️ ดาวน์โหลด Word (ที่ปรึกษา)",
                            data=word_buf_con,
                            file_name=f"routine_consultant_{st.session_state['year']}_{st.session_state['project_name'].replace(' ','_')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="dl_word_con",
                        )
                    except Exception as e:
                        st.error(f"สร้างรายงานไม่ได้: {e}")
        else:
            st.info("ยังไม่มีข้อมูล")

    # ── JSON Save ──────────────────────────────
    with col_e4:
        json_bytes5 = json.dumps(build_json_data(), ensure_ascii=False, indent=2).encode("utf-8")
        st.download_button(
            "💾 บันทึก JSON",
            data=json_bytes5,
            file_name=f"routine_{st.session_state['year']}_{st.session_state['project_name'].replace(' ','_')}.json",
            mime="application/json",
            key="json_dl_tab5",
        )

    # ── Send to LCCA ───────────────────────────
    with col_e5:
        if bud_total > 0:
            if st.button("📤 ส่งให้ LCCA", type="secondary", key="send_lcca"):
                rpc = bud_total / dist_total if dist_total > 0 else 0
                st.session_state["routine_to_lcca"] = {
                    "total_budget_per_year":        bud_total,
                    "total_distance_km":            round(dist_total, 3),
                    "routine_cost_per_km_per_year": round(rpc, 2),
                    "workload_total":               round(wl_total, 3),
                    "project": st.session_state["project_name"],
                    "year":    st.session_state["year"],
                }
                st.success(f"✅ Routine Cost = {rpc:,.2f} บาท/กม./ปี")
                st.json(st.session_state["routine_to_lcca"])
        else:
            st.info("ยังไม่มีข้อมูล")
