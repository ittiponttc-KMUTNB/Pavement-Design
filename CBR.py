import math
import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')
import json
import base64
import io

# Check if python-docx is available
try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Check if openpyxl is available (for template)
try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

st.set_page_config(
    page_title="CBR Percentile Analysis",
    page_icon="📊",
    layout="wide"
)

st.title("📊 การวิเคราะห์ค่า CBR ที่เปอร์เซ็นต์ไทล์")
st.markdown("### Subgrade CBR Analysis Tool")
st.markdown("---")

# Sample data (CBR values only)
sample_cbr = [14.8, 14.37, 5.31, 17.37, 5.48, 18.46, 4.85, 6.23,
              5.02, 10.78, 10.52, 14, 15.5, 8.7, 12.93, 8.19,
              8.1, 15.56, 16.88, 20.75, 20.3, 8, 7.84, 7.48,
              23.55, 8.92, 13.3, 13.5, 13.86, 7.18, 6.95, 5.8,
              6, 11.18, 9.69, 7.48]


# =============================================================================
# Function: Max Rank Percentile Calculation (Method 2 - ≥ approach)
# =============================================================================
def calc_max_rank_percentile(cbr_values_raw):
    """
    คำนวณ Percentile แบบ Max Rank (จำนวนที่มีค่า ≥ CBR นั้น / n × 100)
    
    Returns:
        cbr_sorted: array ของ CBR ทั้งหมด (sorted ascending)
        n: จำนวนตัวอย่างทั้งหมด
        unique_cbr: array ของ CBR ที่ไม่ซ้ำ (sorted ascending)
        unique_pct: array ของ Percentile สำหรับแต่ละ unique CBR
        full_table: list of dict สำหรับตารางแสดงผลเต็ม (ทุกแถว + เว้นว่างค่าซ้ำ)
    """
    cbr_sorted = np.sort(cbr_values_raw)
    n = len(cbr_sorted)
    
    # หาค่า unique CBR (sorted ascending)
    unique_cbr = np.unique(cbr_sorted)
    
    # คำนวณ Percentile = (จำนวนตัวอย่างที่ ≥ CBR) / n × 100
    unique_pct = np.array([
        np.sum(cbr_sorted >= cbr_val) / n * 100 
        for cbr_val in unique_cbr
    ])
    
    # สร้างตารางเต็ม: แสดงทุกแถว แต่ค่าซ้ำแสดง count+pct เฉพาะแถวแรกของกลุ่ม
    full_table = []
    seen = set()
    for i, cbr_val in enumerate(cbr_sorted):
        count_gte = int(np.sum(cbr_sorted >= cbr_val))
        pct_gte = count_gte / n * 100
        
        if cbr_val not in seen:
            # แถวแรกของกลุ่ม - แสดงค่า count และ percentile
            seen.add(cbr_val)
            full_table.append({
                'order': i + 1,
                'cbr': cbr_val,
                'count_gte': count_gte,
                'pct_gte': pct_gte,
                'show_pct': True
            })
        else:
            # แถวซ้ำ - เว้นว่างคอลัมน์ count และ percentile
            full_table.append({
                'order': i + 1,
                'cbr': cbr_val,
                'count_gte': count_gte,
                'pct_gte': pct_gte,
                'show_pct': False
            })
    
    return cbr_sorted, n, unique_cbr, unique_pct, full_table


# =============================================================================
# Function: Generate Excel Template
# =============================================================================
def generate_template_excel():
    """สร้างไฟล์ Excel template สำหรับกรอกค่า CBR"""
    if not OPENPYXL_AVAILABLE:
        return None
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "CBR Data"
    
    # Styles
    header_font = Font(name='TH SarabunPSK', size=14, bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
    header_align = Alignment(horizontal='center', vertical='center')
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    data_font = Font(name='TH SarabunPSK', size=14)
    data_align = Alignment(horizontal='center', vertical='center')
    note_font = Font(name='TH SarabunPSK', size=12, italic=True, color='808080')
    
    # Header
    ws['A1'] = 'CBR (%)'
    ws['A1'].font = header_font
    ws['A1'].fill = header_fill
    ws['A1'].alignment = header_align
    ws['A1'].border = thin_border
    ws.column_dimensions['A'].width = 15
    
    # Example data rows (2 rows as guide)
    for row_idx in range(2, 4):
        cell = ws.cell(row=row_idx, column=1)
        cell.font = data_font
        cell.alignment = data_align
        cell.border = thin_border
    ws['A2'] = 6.5
    ws['A3'] = 7.2
    
    # Empty rows for user to fill (up to row 102)
    for row_idx in range(4, 103):
        cell = ws.cell(row=row_idx, column=1)
        cell.font = data_font
        cell.alignment = data_align
        cell.border = thin_border
    
    # Notes in column C
    ws['C1'] = 'คำแนะนำ:'
    ws['C1'].font = Font(name='TH SarabunPSK', size=12, bold=True, color='4472C4')
    ws['C2'] = '• กรอกค่า CBR (%) ในคอลัมน์ A'
    ws['C2'].font = note_font
    ws['C3'] = '• ค่าซ้ำได้ ระบบจะจัดการให้อัตโนมัติ'
    ws['C3'].font = note_font
    ws['C4'] = '• ลบแถวตัวอย่าง (6.5, 7.2) แล้วกรอกข้อมูลจริง'
    ws['C4'].font = note_font
    ws['C5'] = '• บันทึกแล้ว Upload ในโปรแกรม'
    ws['C5'].font = note_font
    ws.column_dimensions['C'].width = 45
    
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# Sidebar for file upload
with st.sidebar:
    st.header("📁 อัปโหลดข้อมูล")
    
    # Upload JSON for settings
    st.markdown("#### 📂 โหลดการตั้งค่า")
    uploaded_json = st.file_uploader(
        "โหลดข้อมูลจากไฟล์ JSON",
        type=['json'],
        help="โหลดค่า Percentile และข้อมูล CBR จากไฟล์ JSON"
    )
    
    if uploaded_json is not None:
        try:
            loaded_data = json.load(uploaded_json)
            
            # ตรวจสอบว่าเป็นไฟล์ใหม่
            file_id = f"{uploaded_json.name}_{uploaded_json.size}"
            if st.session_state.get('last_uploaded_json') != file_id:
                st.session_state['last_uploaded_json'] = file_id
                
                # อัพเดท session_state
                if 'target_percentile' in loaded_data:
                    st.session_state['input_percentile'] = float(loaded_data['target_percentile'])
                if 'cbr_values' in loaded_data:
                    st.session_state['loaded_cbr_values'] = loaded_data['cbr_values']
                if 'use_sample' in loaded_data:
                    st.session_state['input_use_sample'] = loaded_data['use_sample']
                
                # โหลดค่า report settings ถ้ามี
                if 'report_settings' in loaded_data:
                    rs = loaded_data['report_settings']
                    if 'section_number' in rs:
                        st.session_state['input_section_number'] = rs['section_number']
                    if 'table_number' in rs:
                        st.session_state['input_table_number'] = rs['table_number']
                    if 'figure_number' in rs:
                        st.session_state['input_figure_number'] = rs['figure_number']
                    if 'section_title' in rs:
                        st.session_state['input_section_title'] = rs['section_title']
                    if 'table_caption' in rs:
                        st.session_state['input_table_caption'] = rs['table_caption']
                    if 'figure_caption' in rs:
                        st.session_state['input_figure_caption'] = rs['figure_caption']
                    if 'design_cbr' in rs:
                        st.session_state['input_design_cbr'] = rs['design_cbr']
                
                st.success("✅ โหลดการตั้งค่าสำเร็จ!")
                st.rerun()
                
        except Exception as e:
            st.error(f"❌ ไม่สามารถอ่านไฟล์ JSON ได้: {e}")
    
    st.markdown("---")
    
    # Upload Excel for CBR data
    st.markdown("#### 📊 อัปโหลดข้อมูล CBR")
    uploaded_file = st.file_uploader(
        "เลือกไฟล์ Excel (.xlsx)",
        type=['xlsx'],
        help="ไฟล์ควรมีคอลัมน์ CBR(%) เพียงคอลัมน์เดียว"
    )
    
    # Download Template button
    st.markdown("#### 📄 ดาวน์โหลด Template")
    template_buf = generate_template_excel()
    if template_buf is not None:
        st.download_button(
            label="📥 Download Template Excel",
            data=template_buf,
            file_name="CBR_Template.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            help="ดาวน์โหลดแบบฟอร์มสำหรับกรอกค่า CBR"
        )
    else:
        st.info("ต้องติดตั้ง openpyxl เพื่อสร้าง Template")
    
    st.markdown("---")
    st.markdown("### 📋 รูปแบบข้อมูลที่ต้องการ")
    st.markdown("""
    | CBR(%) |
    |--------|
    | 14.8   |
    | 14.37  |
    | 5.31   |
    | ...    |
    """)
    st.info("ระบบจะคำนวณ Percentile ให้อัตโนมัติ\n(รองรับค่า CBR ซ้ำ)")

# Process uploaded Excel file
if uploaded_file is not None:
    try:
        # Read Excel file
        df = pd.read_excel(uploaded_file)
        
        # Try to identify CBR column
        cbr_col = None
        
        for col in df.columns:
            col_lower = str(col).lower()
            if 'cbr' in col_lower:
                cbr_col = col
                break
        
        # If not found, use first column
        if cbr_col is None:
            cbr_col = df.columns[0]
        
        # Get CBR values
        cbr_values = pd.to_numeric(df[cbr_col], errors='coerce').dropna().tolist()
        
        st.success(f"✅ อ่านข้อมูลสำเร็จ: {len(cbr_values)} ตัวอย่าง")
        
    except Exception as e:
        st.error(f"❌ เกิดข้อผิดพลาด: {str(e)}")
        st.info("กรุณาตรวจสอบรูปแบบไฟล์ Excel")
        cbr_values = None

elif 'loaded_cbr_values' in st.session_state and st.session_state['loaded_cbr_values']:
    # Use CBR values from loaded JSON
    cbr_values = st.session_state['loaded_cbr_values']
    st.info(f"📌 ใช้ข้อมูลจากไฟล์ JSON: {len(cbr_values)} ตัวอย่าง")

else:
    st.info("📌 กรุณาอัปโหลดไฟล์ Excel หรือใช้ข้อมูลตัวอย่าง")
    
    default_use_sample = st.session_state.get('input_use_sample', True)
    use_sample = st.checkbox(
        "ใช้ข้อมูลตัวอย่าง", 
        value=default_use_sample,
        key="input_use_sample"
    )
    
    if use_sample:
        cbr_values = sample_cbr
    else:
        cbr_values = None

if cbr_values is not None and len(cbr_values) > 0:
    
    # =================================================================
    # Calculate using Max Rank method (Method 2 - ≥ approach)
    # =================================================================
    cbr_sorted, n, unique_cbr, unique_pct, full_table = calc_max_rank_percentile(cbr_values)
    
    # Create interpolation function using unique values only (no duplicates)
    # unique_pct is descending (high pct at low CBR), reverse for interp
    # np.interp requires xp to be increasing
    def f_interp(target_pct):
        """Interpolate CBR from target percentile using unique values"""
        # unique_pct is descending, unique_cbr is ascending
        # reverse both for np.interp (needs increasing xp)
        return np.interp(target_pct, unique_pct[::-1], unique_cbr[::-1])
    
    # Input percentile at the top
    st.markdown("### 🎯 กำหนดค่า Percentile")
    
    default_percentile = st.session_state.get('input_percentile', 90.0)
    target_percentile = st.number_input(
        "Percentile ที่ต้องการ (%)",
        min_value=0.0,
        max_value=100.0,
        value=default_percentile,
        step=1.0,
        help="ใส่ค่า Percentile ที่ต้องการหาค่า CBR (% ที่มีค่าเท่ากับหรือมากกว่า)",
        key="input_percentile"
    )
    
    # Calculate CBR at target percentile
    cbr_at_percentile = float(f_interp(target_percentile))
    st.session_state['cbr_design'] = cbr_at_percentile  # เก็บไว้ใช้ใน Odemark section
    
    st.markdown("---")
    
    # =====================================================================
    # Report Numbering Settings (ผู้ใช้กรอกเอง - สีเหลือง)
    # =====================================================================
    st.markdown("### 📝 ตั้งค่าหัวข้อและเลขที่ สำหรับรายงาน Word")
    
    col_num1, col_num2, col_num3 = st.columns(3)
    
    with col_num1:
        section_number = st.text_input(
            "เลขหัวข้อ",
            value=st.session_state.get('input_section_number', "4.3"),
            key="input_section_number",
            help="เช่น 4.3, 5.1"
        )
    
    with col_num2:
        table_number = st.text_input(
            "เลขตาราง",
            value=st.session_state.get('input_table_number', "4-7"),
            key="input_table_number",
            help="เช่น 4-7, 5-1"
        )
    
    with col_num3:
        figure_number = st.text_input(
            "เลขรูป",
            value=st.session_state.get('input_figure_number', "4-7"),
            key="input_figure_number",
            help="เช่น 4-7, 5-1"
        )
    
    section_title = st.text_input(
        "ชื่อหัวข้อ",
        value=st.session_state.get('input_section_title', "ข้อมูลความแข็งแรงของดินฐานรากบริเวณพื้นที่โครงการ"),
        key="input_section_title",
        help="ชื่อหัวข้อในรายงาน"
    )
    
    col_cap1, col_cap2 = st.columns(2)
    
    with col_cap1:
        table_caption = st.text_input(
            "คำบรรยายตาราง",
            value=st.session_state.get('input_table_caption', "ค่าเปอร์เซ็นต์ไทล์ และค่า CBR ของตัวอย่างดินฐานรากตามแนวสายทาง"),
            key="input_table_caption",
            help="คำบรรยายใต้ตาราง"
        )
    
    with col_cap2:
        figure_caption = st.text_input(
            "คำบรรยายรูป",
            value=st.session_state.get('input_figure_caption', "กราฟแสดงความสัมพันธ์ระหว่าง Percentile และ CBR ของดินฐานรากตามแนวสายทาง"),
            key="input_figure_caption",
            help="คำบรรยายใต้รูป"
        )
    
    design_cbr = st.number_input(
        "ค่า CBR ที่ใช้ในการออกแบบ (%)",
        min_value=0.0,
        max_value=100.0,
        value=st.session_state.get('input_design_cbr', 4.0),
        step=1.0,
        key="input_design_cbr",
        help="ค่า CBR ที่ที่ปรึกษาเลือกใช้ในการออกแบบโครงสร้างชั้นทาง"
    )
    
    # =====================================================================
    # Preview introduction paragraph
    # =====================================================================
    st.markdown("---")
    st.markdown("### 👁️ ตัวอย่างข้อความเกริ่นนำในรายงาน")
    
    # Build preview - สีม่วงดึงจากข้อมูล, สีเหลืองผู้ใช้กรอก
    intro_preview = (
        f'<div style="font-family: TH SarabunPSK, Tahoma, sans-serif; font-size: 15px; line-height: 1.8; '
        f'background-color: #f9f9f9; padding: 15px; border-radius: 8px; border: 1px solid #ddd;">'
        f'<p style="margin-bottom: 5px;"><b>{section_number} &nbsp;&nbsp;&nbsp; {section_title}</b></p>'
        f'<p style="text-indent: 40px; text-align: justify; text-justify: inter-character;">'
        f'ความแข็งแรงของดินฐานรากบริเวณโดยรอบพื้นที่โครงการ หรือกำลังรับน้ำหนักของดินพื้นทางเดิม '
        f'หรือพื้นทางเดิมสามารถประเมินจากรายงานสำรวจภูมิประเทศของดิน ซึ่งสามารถทำการทดสอบได้หลากหลายวิธี เช่น '
        f'Plate Bearing Test CBR Test หรือ Modulus of Subgrade Reaction สำหรับการออกแบบถนนลาดยางและคอนกรีตนั้นใช้ค่า CBR '
        f'ซึ่งนิยมใช้กันแพร่หลาย เมื่อกำหนดกำลังรับน้ำหนักของดินพื้นทางเดิม '
        f'โดยการเจาะสำรวจดินในสนามตามรายงานการสอบดินของห้องปฏิบัติการ เพื่อหาค่า CBR '
        f'ของดินพื้นทางเดินเพื่อเป็นข้อมูลในการออกแบบ ซึ่งผลการทดสอบค่า CBR ของดินฐานรากตามแนวสายทาง จำนวน '
        f'<span style="background-color: #D8B4FE; padding: 1px 4px; border-radius: 3px; font-weight: bold;">{n}</span> ตัวอย่าง '
        f'พบว่าที่เปอร์เซ็นต์ไทล์ ร้อยละ '
        f'<span style="background-color: #D8B4FE; padding: 1px 4px; border-radius: 3px; font-weight: bold;">{target_percentile:.0f}</span> '
        f'ของค่ากำลังที่พบเท่ากับ CBR เท่ากับ '
        f'<span style="background-color: #D8B4FE; padding: 1px 4px; border-radius: 3px; font-weight: bold;">{cbr_at_percentile:.1f}</span> % '
        f'อย่างไรก็ตาม ที่ปรึกษาเลือกค่า CBR เท่ากับ '
        f'<span style="background-color: #FDE68A; padding: 1px 4px; border-radius: 3px; font-weight: bold;">{int(design_cbr)}</span> % '
        f'มาใช้ในการออกแบบโครงสร้างชั้นทาง '
        f'ดังแสดงผลการวิเคราะห์ใน'
        f'<span style="background-color: #FDE68A; padding: 1px 4px; border-radius: 3px; font-weight: bold;">ตารางที่ {table_number}</span> '
        f'และ<span style="background-color: #FDE68A; padding: 1px 4px; border-radius: 3px; font-weight: bold;">รูปที่ {figure_number}</span></p>'
        f'</div>'
    )
    
    st.markdown(intro_preview, unsafe_allow_html=True)
    st.caption("🟣 สีม่วง = ดึงจากข้อมูลอัตโนมัติ | 🟡 สีเหลือง = ผู้ใช้กรอกเอง")
    
    st.markdown("---")
    
    # Graph section - full width (uses unique values only for clean curve)
    st.markdown("### 📈 กราฟ Percentile vs CBR")
    
    # Create figure
    fig = go.Figure()
    
    # Calculate axis ranges
    x_max = max(unique_cbr) * 1.1
    y_max = 100
    
    # Add main curve - plot unique points only (no vertical lines from duplicates)
    fig.add_trace(go.Scatter(
        x=unique_cbr,
        y=unique_pct,
        mode='lines+markers',
        name='CBR Distribution',
        line=dict(color='blue', width=2),
        marker=dict(size=6, symbol='x', color='black')
    ))
    
    # Add horizontal red dashed line at target percentile
    fig.add_trace(go.Scatter(
        x=[0, cbr_at_percentile],
        y=[target_percentile, target_percentile],
        mode='lines',
        name=f'Percentile {target_percentile}%',
        line=dict(color='red', width=2, dash='dash')
    ))
    
    # Add vertical red dashed line at CBR value
    fig.add_trace(go.Scatter(
        x=[cbr_at_percentile, cbr_at_percentile],
        y=[0, target_percentile],
        mode='lines',
        name=f'CBR = {cbr_at_percentile:.2f}%',
        line=dict(color='red', width=2, dash='dash')
    ))
    
    # Add annotation for CBR value
    fig.add_annotation(
        x=cbr_at_percentile,
        y=0,
        text=f"<b>{cbr_at_percentile:.2f}</b>",
        showarrow=True,
        arrowhead=2,
        arrowsize=1,
        arrowwidth=2,
        arrowcolor='red',
        ax=0,
        ay=40,
        font=dict(size=16, color='red')
    )
    
    # Border line width (consistent for all 4 sides)
    border_width = 4
    
    # Update layout - remove axis lines, we'll draw border using shapes
    fig.update_layout(
        xaxis_title="CBR (%)",
        yaxis_title="Percentile (%)",
        xaxis=dict(
            range=[0, x_max],
            gridcolor='lightgray',
            showgrid=True,
            showline=False,  # Disable built-in axis line
            zeroline=False,
            ticks='outside',
            tickwidth=1,
            tickcolor='black',
            ticklen=5,
        ),
        yaxis=dict(
            range=[0, y_max],
            gridcolor='lightgray',
            showgrid=True,
            showline=False,  # Disable built-in axis line
            zeroline=False,
            ticks='outside',
            tickwidth=1,
            tickcolor='black',
            ticklen=5,
        ),
        plot_bgcolor='white',
        width=600,
        height=600,
        showlegend=True,
        legend=dict(
            yanchor="top",
            y=0.99,
            xanchor="right",
            x=0.99,
            bgcolor='rgba(255,255,255,0.8)',
            bordercolor='black',
            borderwidth=1
        ),
        title=dict(
            text=f"ค่าร้อยละ CBR ที่เปอร์เซ็นต์ไทล์ ร้อยละ {target_percentile:.0f}",
            x=0.5,
            xanchor='center'
        ),
        margin=dict(l=70, r=70, t=70, b=70)
    )
    
    # Draw complete border using a rectangle shape (ensures all 4 corners connect)
    fig.add_shape(
        type="rect",
        x0=0, y0=0,
        x1=x_max, y1=y_max,
        line=dict(color="black", width=border_width),
        xref="x", yref="y"
    )
    
    # Center the chart
    col_left, col_chart, col_right = st.columns([1, 2, 1])
    with col_chart:
        st.plotly_chart(fig, use_container_width=False)
    
    # Results section - below the graph
    st.markdown("---")
    
    col_result, col_stat = st.columns(2)
    
    with col_result:
        st.markdown("### 📊 ผลการวิเคราะห์")
        st.metric(
            label=f"CBR ที่ Percentile {target_percentile}%",
            value=f"{cbr_at_percentile:.2f} %"
        )
    
    with col_stat:
        st.markdown("### 📋 สถิติข้อมูล CBR")
        st.write(f"**จำนวนตัวอย่าง:** {n}")
        st.write(f"**ค่าต่ำสุด:** {np.min(cbr_values):.2f} %")
        st.write(f"**ค่าสูงสุด:** {np.max(cbr_values):.2f} %")
        st.write(f"**ค่าเฉลี่ย:** {np.mean(cbr_values):.2f} %")
        st.write(f"**ส่วนเบี่ยงเบนมาตรฐาน:** {np.std(cbr_values):.2f} %")
    
    # Export section
    st.markdown("---")
    st.markdown("### 💾 บันทึกข้อมูล")
    
    col_json, col_word = st.columns(2)
    
    with col_json:
        # Prepare export data for JSON
        export_data = {
            'target_percentile': target_percentile,
            'cbr_at_percentile': round(cbr_at_percentile, 2),
            'cbr_values': [float(v) for v in cbr_values],
            'statistics': {
                'n_samples': n,
                'min': round(float(np.min(cbr_values)), 2),
                'max': round(float(np.max(cbr_values)), 2),
                'mean': round(float(np.mean(cbr_values)), 2),
                'std': round(float(np.std(cbr_values)), 2)
            },
            'report_settings': {
                'section_number': section_number,
                'table_number': table_number,
                'figure_number': figure_number,
                'section_title': section_title,
                'table_caption': table_caption,
                'figure_caption': figure_caption,
                'design_cbr': design_cbr,
            },
            'use_sample': st.session_state.get('input_use_sample', True)
        }
        
        json_str = json.dumps(export_data, ensure_ascii=False, indent=2)
        
        st.download_button(
            label="📥 Download JSON",
            data=json_str,
            file_name="cbr_percentile_data.json",
            mime="application/json",
            help="บันทึกข้อมูลและการตั้งค่าเป็นไฟล์ JSON"
        )
    
    with col_word:
        # Generate Word document using python-docx
        if DOCX_AVAILABLE:
            if st.button("📄 สร้างรายงาน Word", help="สร้างรายงานผลการวิเคราะห์เป็นไฟล์ Word"):
                try:
                    # Create Word document
                    doc = DocxDocument()
                    
                    # Set Thai font style
                    style = doc.styles['Normal']
                    style.font.name = 'TH SarabunPSK'
                    style.font.size = Pt(15)
                    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')
                    
                    # Helper function to set cell background color
                    def set_cell_bg(cell, color_hex):
                        """Set background color for a cell using tcPr"""
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        shading = OxmlElement('w:shd')
                        shading.set(qn('w:val'), 'clear')
                        shading.set(qn('w:color'), 'auto')
                        shading.set(qn('w:fill'), color_hex)
                        tcPr.append(shading)
                    
                    # =========================================================
                    # 1) Section heading: "4.3    ข้อมูลความแข็งแรง..."
                    # =========================================================
                    heading_para = doc.add_paragraph()
                    heading_run = heading_para.add_run(f'{section_number}\t{section_title}')
                    heading_run.font.name = 'TH SarabunPSK'
                    heading_run.font.size = Pt(15)
                    heading_run.font.bold = True
                    
                    # =========================================================
                    # 2) Introduction paragraph (เกริ่นนำ) - Thai Distributed, 15pt
                    # =========================================================
                    intro_para = doc.add_paragraph()
                    intro_para.paragraph_format.first_line_indent = Cm(1.25)
                    # Thai Distributed alignment via XML (thaiDistribute)
                    pPr = intro_para._element.get_or_add_pPr()
                    jc = OxmlElement('w:jc')
                    jc.set(qn('w:val'), 'thaiDistribute')
                    pPr.append(jc)
                    
                    # เลือก CBR ที่ใช้ออกแบบสำหรับรายงาน
                    _improve = st.session_state.get('improve_soil_check', False)
                    _ode = st.session_state.get('odemark_result')
                    if _improve and _ode:
                        cbr_for_report = _ode.get('cbr_eq_design', math.floor(_ode['cbr_eq']))
                    else:
                        cbr_for_report = math.floor(design_cbr)

                    # Build intro text with mixed formatting
                    intro_parts = [
                        ('ความแข็งแรงของดินฐานรากบริเวณโดยรอบพื้นที่โครงการ หรือกำลังรับน้ำหนักของดินพื้นทางเดิม '
                         'หรือพื้นทางเดิมสามารถประเมินจากรายงานสำรวจภูมิประเทศของดิน ซึ่งสามารถทำการทดสอบได้หลากหลายวิธี เช่น '
                         'Plate Bearing Test CBR Test หรือ Modulus of Subgrade Reaction สำหรับการออกแบบถนนลาดยางและคอนกรีตนั้นใช้ค่า CBR '
                         'ซึ่งนิยมใช้กันแพร่หลาย เมื่อกำหนดกำลังรับน้ำหนักของดินพื้นทางเดิม '
                         'โดยการเจาะสำรวจดินในสนามตามรายงานการสอบดินของห้องปฏิบัติการ เพื่อหาค่า CBR '
                         'ของดินพื้นทางเดินเพื่อเป็นข้อมูลในการออกแบบ ซึ่งผลการทดสอบค่า CBR ของดินฐานรากตามแนวสายทาง จำนวน ', False),
                        (f'{n}', True),
                        (' ตัวอย่าง พบว่าที่เปอร์เซ็นต์ไทล์ ร้อยละ ', False),
                        (f'{target_percentile:.0f}', True),
                        (' ของค่ากำลังที่พบเท่ากับ CBR เท่ากับ ', False),
                        (f'{cbr_at_percentile:.1f}', True),
                        (' % อย่างไรก็ตาม ที่ปรึกษาเลือกค่า CBR เท่ากับ ', False),
                        (f'{cbr_for_report}', True),
                        (' % มาใช้ในการออกแบบโครงสร้างชั้นทาง ดังแสดงผลการวิเคราะห์ใน', False),
                        (f'ตารางที่ {table_number}', True),
                        (' และ', False),
                        (f'รูปที่ {figure_number}', True),
                    ]
                    
                    for text, is_bold in intro_parts:
                        run = intro_para.add_run(text)
                        run.font.name = 'TH SarabunPSK'
                        run.font.size = Pt(15)
                        run.font.bold = is_bold
                    
                    doc.add_paragraph()  # spacing
                    
                    # =========================================================
                    # 3) TABLE (ตารางมาก่อนรูป) - ใช้ full_table (Max Rank)
                    # =========================================================
                    # Table caption above table
                    table_cap_para = doc.add_paragraph()
                    table_cap_run = table_cap_para.add_run(f'ตารางที่ {table_number} {table_caption}')
                    table_cap_run.font.name = 'TH SarabunPSK'
                    table_cap_run.font.size = Pt(15)
                    table_cap_run.font.bold = True
                    table_cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Calculate half point for splitting data
                    half_n = (n + 1) // 2
                    
                    # Create CBR data table with 8 columns
                    # CBR | จำนวนที่≥ | %ที่≥ | (เว้น) | CBR | จำนวนที่≥ | %ที่≥
                    cbr_table = doc.add_table(rows=half_n+1, cols=6)
                    cbr_table.style = 'Table Grid'
                    cbr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    # Header row
                    header_row = cbr_table.rows[0]
                    headers = ['CBR (%)', 'จำนวนที่≥', 'Percentile (%)', 'CBR (%)', 'จำนวนที่≥', 'Percentile (%)']
                    for j, header_text in enumerate(headers):
                        cell = header_row.cells[j]
                        run = cell.paragraphs[0].add_run(header_text)
                        run.font.name = 'TH SarabunPSK'
                        run.font.size = Pt(14)
                        run.font.bold = True
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        set_cell_bg(cell, 'D9E2F3')  # light blue header
                    
                    # Data rows - split into left and right halves
                    for i in range(half_n):
                        row = cbr_table.rows[i+1]
                        
                        # Left side data (first half)
                        left_idx = i
                        if left_idx < n:
                            ft = full_table[left_idx]
                            left_data = [
                                f'{ft["cbr"]:.2f}',
                                f'{ft["count_gte"]}' if ft['show_pct'] else '',
                                f'{ft["pct_gte"]:.1f}' if ft['show_pct'] else ''
                            ]
                        else:
                            left_data = ['', '', '']
                        
                        # Right side data (second half)
                        right_idx = i + half_n
                        if right_idx < n:
                            ft = full_table[right_idx]
                            right_data = [
                                f'{ft["cbr"]:.2f}',
                                f'{ft["count_gte"]}' if ft['show_pct'] else '',
                                f'{ft["pct_gte"]:.1f}' if ft['show_pct'] else ''
                            ]
                        else:
                            right_data = ['', '', '']
                        
                        # Fill left side (columns 0-2)
                        for j, val in enumerate(left_data):
                            cell = row.cells[j]
                            run = cell.paragraphs[0].add_run(val)
                            run.font.name = 'TH SarabunPSK'
                            run.font.size = Pt(14)
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Fill right side (columns 3-5)
                        for j, val in enumerate(right_data):
                            cell = row.cells[j+3]
                            run = cell.paragraphs[0].add_run(val)
                            run.font.name = 'TH SarabunPSK'
                            run.font.size = Pt(14)
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Set column widths for CBR table
                    for row in cbr_table.rows:
                        row.cells[0].width = Cm(2.0)
                        row.cells[1].width = Cm(2.5)
                        row.cells[2].width = Cm(3.5)
                        row.cells[3].width = Cm(2.0)
                        row.cells[4].width = Cm(2.5)
                        row.cells[5].width = Cm(3.5)
                    
                    doc.add_paragraph()  # spacing
                    
                    # =========================================================
                    # 4) Statistics summary table
                    # =========================================================
                    h_stat = doc.add_paragraph()
                    h_stat_run = h_stat.add_run('ผลการวิเคราะห์')
                    h_stat_run.font.name = 'TH SarabunPSK'
                    h_stat_run.font.size = Pt(15)
                    h_stat_run.font.bold = True
                    
                    # Create statistics table
                    table = doc.add_table(rows=7, cols=2)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    # Table data
                    table_data = [
                        ('รายการ', 'ค่า'),
                        ('จำนวนตัวอย่าง', f'{n}'),
                        ('ค่าต่ำสุด', f'{np.min(cbr_values):.2f} %'),
                        ('ค่าสูงสุด', f'{np.max(cbr_values):.2f} %'),
                        ('ค่าเฉลี่ย', f'{np.mean(cbr_values):.2f} %'),
                        ('ส่วนเบี่ยงเบนมาตรฐาน', f'{np.std(cbr_values):.2f} %'),
                        (f'CBR ที่ Percentile {target_percentile:.0f}%', f'{cbr_at_percentile:.2f} %')
                    ]
                    
                    for i, (col1, col2) in enumerate(table_data):
                        row = table.rows[i]
                        cell1 = row.cells[0]
                        cell2 = row.cells[1]
                        
                        run1 = cell1.paragraphs[0].add_run(col1)
                        run1.font.name = 'TH SarabunPSK'
                        run1.font.size = Pt(14)
                        if i == 0:
                            run1.font.bold = True
                            set_cell_bg(cell1, 'D9E2F3')
                            set_cell_bg(cell2, 'D9E2F3')
                        
                        run2 = cell2.paragraphs[0].add_run(col2)
                        run2.font.name = 'TH SarabunPSK'
                        run2.font.size = Pt(14)
                        if i == 0:
                            run2.font.bold = True
                        if i == 6:  # Last row - CBR result
                            run2.font.bold = True
                            run2.font.color.rgb = RGBColor(255, 0, 0)
                    
                    # Set column widths
                    for row in table.rows:
                        row.cells[0].width = Cm(12)
                        row.cells[1].width = Cm(4)
                    
                    doc.add_paragraph()  # spacing
                    
                    # =========================================================
                    # 5) FIGURE (รูปมาหลังตาราง) - ใช้ unique values
                    # =========================================================
                    # Create chart using matplotlib
                    fig_mpl, ax = plt.subplots(figsize=(6, 6))
                    
                    # Plot main curve using unique values only
                    ax.plot(unique_cbr, unique_pct, 'b-', linewidth=2, marker='x', 
                           markersize=6, markerfacecolor='black', markeredgecolor='black',
                           label='CBR Distribution')
                    
                    # Plot dashed lines
                    ax.plot([0, cbr_at_percentile], [target_percentile, target_percentile], 
                           'r--', linewidth=2, label=f'Percentile {target_percentile}%')
                    ax.plot([cbr_at_percentile, cbr_at_percentile], [0, target_percentile], 
                           'r--', linewidth=2, label=f'CBR = {cbr_at_percentile:.2f}%')
                    
                    # Annotation
                    ax.annotate(f'{cbr_at_percentile:.2f}', 
                               xy=(cbr_at_percentile, 0), 
                               xytext=(cbr_at_percentile, -8),
                               fontsize=12, color='red', fontweight='bold',
                               ha='center')
                    
                    ax.set_xlim(0, max(unique_cbr) * 1.1)
                    ax.set_ylim(0, 100)
                    ax.set_xlabel('CBR (%)', fontsize=12)
                    ax.set_ylabel('Percentile (%)', fontsize=12)
                    ax.set_title(f'CBR at Percentile {target_percentile:.0f}%', fontsize=14)
                    ax.legend(loc='upper right', fontsize=10)
                    ax.grid(True, alpha=0.3)
                    
                    # Set border
                    for spine in ax.spines.values():
                        spine.set_linewidth(2)
                        spine.set_color('black')
                    
                    plt.tight_layout()
                    
                    # Save chart to buffer
                    chart_buffer = io.BytesIO()
                    fig_mpl.savefig(chart_buffer, format='png', dpi=150, 
                                   bbox_inches='tight', facecolor='white', edgecolor='none')
                    chart_buffer.seek(0)
                    plt.close(fig_mpl)
                    
                    # Add chart image to document
                    chart_para = doc.add_paragraph()
                    chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    chart_run = chart_para.add_run()
                    chart_run.add_picture(chart_buffer, width=Cm(12))
                    
                    # Add figure caption
                    caption = doc.add_paragraph()
                    caption_run = caption.add_run(f'รูปที่ {figure_number} {figure_caption}')
                    caption_run.font.name = 'TH SarabunPSK'
                    caption_run.font.size = Pt(15)
                    caption_run.font.bold = True
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # =========================================================
                    # 6) Footer
                    # =========================================================
                    doc.add_paragraph()
                    doc.add_paragraph()
                    footer1 = doc.add_paragraph()
                    footer1_run = footer1.add_run('พัฒนาโดย รศ.ดร.อิทธิพล มีผล')
                    footer1_run.font.name = 'TH SarabunPSK'
                    footer1_run.font.size = Pt(14)
                    footer1_run.font.italic = True
                    footer1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    footer2 = doc.add_paragraph()
                    footer2_run = footer2.add_run('ภาควิชาครุศาสตร์โยธา คณะครุศาสตร์อุตสาหกรรม มจพ.')
                    footer2_run.font.name = 'TH SarabunPSK'
                    footer2_run.font.size = Pt(14)
                    footer2_run.font.italic = True
                    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # =========================================================
                    # 7) ส่วนปรับปรุงดินคันทาง (ถ้ามี)
                    # =========================================================
                    if (st.session_state.get('improve_soil_check')
                            and st.session_state.get('odemark_result')):

                        res = st.session_state['odemark_result']
                        MPA_PER_CBR_DOC = 1500 * 0.006895

                        doc.add_paragraph()
                        doc.add_paragraph()

                        # หัวข้อ
                        imp_heading = doc.add_paragraph()
                        imp_run = imp_heading.add_run("การปรับปรุงดินคันทาง (Subgrade Improvement)")
                        imp_run.font.name = 'TH SarabunPSK'
                        imp_run.font.size = Pt(15)
                        imp_run.font.bold = True

                        # ย่อหน้าอธิบาย
                        imp_intro = doc.add_paragraph()
                        imp_intro.paragraph_format.first_line_indent = Cm(1.25)
                        pPr2 = imp_intro._element.get_or_add_pPr()
                        jc2 = OxmlElement('w:jc')
                        jc2.set(qn('w:val'), 'thaiDistribute')
                        pPr2.append(jc2)
                        intro_run2 = imp_intro.add_run(
                            f"เนื่องจากค่า CBR ของดินเดิมที่ได้จากการวิเคราะห์ทางสถิติมีค่าต่ำหรือแหล่งวัสดุดินถมคันทางมีค่า CBR ต่ำ  "
                            f"จึงได้ทำการปรับปรุงดินคันทางโดยการใช้วัสดุคุณภาพดีปูทับ "
                            f"และคำนวณค่า CBR เทียบเท่า (CBR Equivalent) ด้วยวิธี Odemark (1974) "
                            f"โดยพิจารณาโครงสร้างดิน 2 ชั้น ดังนี้"
                        )
                        intro_run2.font.name = 'TH SarabunPSK'
                        intro_run2.font.size = Pt(15)

                        doc.add_paragraph()

                        # ตารางชั้นดิน
                        imp_tbl = doc.add_table(rows=3, cols=4)
                        imp_tbl.style = 'Table Grid'
                        imp_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

                        imp_headers = ['ชั้นดิน', 'ชนิดวัสดุ', 'ความหนา (ซม.)', 'MR (MPa)']
                        imp_data = [
                            ('ชั้นที่ 1 (วัสดุปรับปรุง)',
                             res['mat1'],
                             f"{res['h1_cm']:.1f}",
                             f"{res['mr1_mpa']:.1f}"),
                            ('ชั้นที่ 2 (ดินถมคันทางใหม่)',
                             f"ดินถมคันทาง CBR = {res['cbr2']:.1f} %",
                             f"{res['h2_cm']:.1f}",
                             f"{res['mr2_mpa']:.2f}"),
                        ]

                        for j, hdr in enumerate(imp_headers):
                            cell = imp_tbl.rows[0].cells[j]
                            r = cell.paragraphs[0].add_run(hdr)
                            r.font.name = 'TH SarabunPSK'
                            r.font.size = Pt(14)
                            r.font.bold = True
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            set_cell_bg(cell, 'D9E2F3')

                        for row_i, row_vals in enumerate(imp_data):
                            for col_j, val in enumerate(row_vals):
                                cell = imp_tbl.rows[row_i + 1].cells[col_j]
                                r = cell.paragraphs[0].add_run(val)
                                r.font.name = 'TH SarabunPSK'
                                r.font.size = Pt(14)
                                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        for row in imp_tbl.rows:
                            row.cells[0].width = Cm(4.0)
                            row.cells[1].width = Cm(7.0)
                            row.cells[2].width = Cm(2.5)
                            row.cells[3].width = Cm(2.5)

                        doc.add_paragraph()

                        # สูตรและการคำนวณ
                        calc_heading = doc.add_paragraph()
                        calc_r = calc_heading.add_run("วิธีการคำนวณ")
                        calc_r.font.name = 'TH SarabunPSK'
                        calc_r.font.size = Pt(15)
                        calc_r.font.bold = True

                        calc_lines = [
                            "สูตร Odemark (1974):  MR_eq = ( Σ(h_i × MR_i^(1/3)) / Σh_i )^3",
                            f"ชั้นที่ 1 : h = {res['h1_cm']:.1f} cm, MR = {res['mr1_mpa']:.1f} MPa, "
                            f"MR^(1/3) = {res['mr1_mpa']**(1/3):.4f}",
                            f"ชั้นที่ 2 : h = {res['h2_cm']:.1f} cm, MR = {res['mr2_mpa']:.2f} MPa, "
                            f"MR^(1/3) = {res['mr2_mpa']**(1/3):.4f}",
                            f"Σh = {res['sum_h']:.1f} cm",
                            f"Σ(h·MR^(1/3)) = {res['sum_hE13']:.4f}",
                            f"MR_eq = ({res['sum_hE13']:.4f} / {res['sum_h']:.1f})^3 = {res['mr_eq_mpa']:.2f} MPa",
                            f"CBR_equivalent = MR_eq / (1500 × 0.006895) = {res['cbr_eq']:.2f} %",
                        ]
                        for line in calc_lines:
                            p = doc.add_paragraph(style='Normal')
                            r2 = p.add_run(line)
                            r2.font.name = 'TH SarabunPSK'
                            r2.font.size = Pt(14)

                        doc.add_paragraph()

                        # สรุปผล
                        result_heading = doc.add_paragraph()
                        result_r = result_heading.add_run("ผลการคำนวณ")
                        result_r.font.name = 'TH SarabunPSK'
                        result_r.font.size = Pt(15)
                        result_r.font.bold = True

                        summary_tbl = doc.add_table(rows=3, cols=2)
                        summary_tbl.style = 'Table Grid'
                        summary_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

                        summary_data = [
                            ('รายการ', 'ค่า'),
                            ('MR equivalent', f'{res["mr_eq_mpa"]:.2f} MPa'),
                            ('CBR equivalent (จากการคำนวณ)', f'{res["cbr_eq"]:.2f} %'),
                            ('CBR equivalent (ใช้ออกแบบ)', f'{res["cbr_eq_design"]} %'),
                        ]
                        # ปรับจำนวนแถวในตารางให้ครบ
                        while len(summary_tbl.rows) < len(summary_data):
                            summary_tbl.add_row()
                        for row_i, (c1, c2) in enumerate(summary_data):
                            cell1 = summary_tbl.rows[row_i].cells[0]
                            cell2 = summary_tbl.rows[row_i].cells[1]
                            r1s = cell1.paragraphs[0].add_run(c1)
                            r2s = cell2.paragraphs[0].add_run(c2)
                            for r_s in (r1s, r2s):
                                r_s.font.name = 'TH SarabunPSK'
                                r_s.font.size = Pt(14)
                                if row_i == 0:
                                    r_s.font.bold = True
                            if row_i == 0:
                                set_cell_bg(cell1, 'D9E2F3')
                                set_cell_bg(cell2, 'D9E2F3')
                            if row_i == 3:  # แถว CBR ใช้ออกแบบ เน้นสีแดง
                                r2s.font.bold = True
                                r2s.font.color.rgb = RGBColor(255, 0, 0)

                        # ขยายตารางเต็มหน้ากระดาษ A4 (16 cm usable)
                        for row in summary_tbl.rows:
                            row.cells[0].width = Cm(12)
                            row.cells[1].width = Cm(4)

                    # Save to buffer
                    buffer = io.BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    
                    st.download_button(
                        label="📥 Download Word",
                        data=buffer,
                        file_name="cbr_percentile_report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    st.success("✅ สร้างรายงาน Word สำเร็จ!")
                    
                except Exception as e:
                    st.error(f"❌ ไม่สามารถสร้างรายงาน Word ได้: {e}")
        else:
            st.warning("⚠️ ต้องติดตั้ง python-docx เพื่อใช้งานฟีเจอร์นี้")
            st.code("pip install python-docx", language="bash")
    
    # Show data table (Max Rank format)
    st.markdown("---")
    st.markdown("### 📋 ตารางข้อมูล (เรียงตาม CBR)")
    
    # Create display table using full_table (Max Rank method)
    df_display = pd.DataFrame({
        'ลำดับ': [ft['order'] for ft in full_table],
        'CBR (%)': [ft['cbr'] for ft in full_table],
        'จำนวนที่≥': [ft['count_gte'] if ft['show_pct'] else None for ft in full_table],
        'Percentile (%)': [round(ft['pct_gte'], 1) if ft['show_pct'] else None for ft in full_table]
    })
    
    col_a, col_b = st.columns(2)
    
    with col_a:
        st.dataframe(
            df_display.head(len(df_display)//2 + 1),
            use_container_width=True,
            hide_index=True
        )
    
    with col_b:
        st.dataframe(
            df_display.tail(len(df_display)//2),
            use_container_width=True,
            hide_index=True
        )

# =====================================================================
# ODEMARK SECTION: ปรับปรุงดินคันทาง
# =====================================================================

# วัสดุที่เหมาะสมสำหรับปรับปรุงดินคันทาง
IMPROVE_MATERIAL_DB = {
    "หินคลุก CBR 80%":                     {"MR_default": 350},
    "รองพื้นทางวัสดุมวลรวม (CBR 25%)":    {"MR_default": 150},
    "วัสดุคัดเลือก ก":                     {"MR_default": 100},
}

MPA_PER_CBR = 1500 * 0.006895  # = 10.3425 MPa / %CBR

def update_imp_mr1():
    """อัปเดต MR เมื่อเปลี่ยนชนิดวัสดุชั้นที่ 1 และ clear ผลเก่า"""
    mat = st.session_state["imp_mat1"]
    st.session_state["imp_mr1"] = float(IMPROVE_MATERIAL_DB[mat]["MR_default"])
    st.session_state.pop('odemark_result', None)

def clear_odemark_result():
    """Clear ผลการคำนวณเมื่อ input เปลี่ยน"""
    st.session_state.pop('odemark_result', None)

if st.session_state.get('cbr_design') is not None:
    st.markdown("---")
    st.markdown("### 🔧 การปรับปรุงดินคันทาง")

    improve_soil = st.checkbox(
        "ต้องการปรับปรุงดินคันทาง (CBR ดินเดิมต่ำ)",
        key="improve_soil_check"
    )

    if improve_soil:
        st.markdown("#### กำหนดชั้นดิน 2 ชั้น")

        # ตั้งค่าเริ่มต้นครั้งแรก (ป้องกัน key error)
        if "imp_mat1" not in st.session_state:
            first_mat = list(IMPROVE_MATERIAL_DB.keys())[0]
            st.session_state["imp_mat1"] = first_mat
            st.session_state["imp_mr1"] = float(IMPROVE_MATERIAL_DB[first_mat]["MR_default"])

        # ── ชั้นที่ 1: วัสดุปรับปรุง ──────────────────────────────────
        st.markdown("**ชั้นที่ 1 — วัสดุปรับปรุง**")
        col_m1, col_h1, col_mr1 = st.columns(3)

        with col_m1:
            mat1 = st.selectbox(
                "ชนิดวัสดุ",
                list(IMPROVE_MATERIAL_DB.keys()),
                key="imp_mat1",
                on_change=update_imp_mr1
            )
        with col_h1:
            h1_cm = st.number_input(
                "ความหนา (ซม.)",
                min_value=1.0, max_value=150.0,
                value=30.0, step=5.0,
                key="imp_h1",
                on_change=clear_odemark_result
            )
        with col_mr1:
            mr1_mpa = st.number_input(
                "MR (MPa)",
                min_value=10.0, max_value=1000.0,
                step=10.0,
                key="imp_mr1",
                on_change=clear_odemark_result
            )

        # ── ชั้นที่ 2: ดินถมคันทางใหม่ ────────────────────────────────
        st.markdown("**ชั้นที่ 2 — ดินถมคันทางใหม่**")
        col_h2, col_cbr2 = st.columns(2)

        with col_h2:
            h2_cm = st.number_input(
                "ความหนา (ซม.)",
                min_value=1.0, max_value=300.0,
                value=50.0, step=5.0,
                key="imp_h2",
                on_change=clear_odemark_result
            )
        with col_cbr2:
            cbr2_input = st.number_input(
                "CBR ดินถมคันทางใหม่ (%)",
                min_value=0.1, max_value=100.0,
                value=10.0, step=1.0,
                key="imp_cbr2",
                on_change=clear_odemark_result
            )

        mr2_mpa = cbr2_input * MPA_PER_CBR
        st.caption(f"MR ดินถมคันทางใหม่ = {cbr2_input:.1f} × {MPA_PER_CBR:.4f} = **{mr2_mpa:.2f} MPa**")

        # ── คำนวณ Odemark ──────────────────────────────────────────────
        if st.button("คำนวณ CBR_equivalent (Odemark)", key="btn_odemark"):
            sum_h     = h1_cm + h2_cm
            sum_hE13  = h1_cm * (mr1_mpa ** (1/3)) + h2_cm * (mr2_mpa ** (1/3))
            mr_eq_mpa = (sum_hE13 / sum_h) ** 3
            cbr_eq    = mr_eq_mpa / MPA_PER_CBR

            st.session_state['odemark_result'] = {
                'mat1': mat1, 'h1_cm': h1_cm, 'mr1_mpa': mr1_mpa,
                'h2_cm': h2_cm, 'cbr2': cbr2_input, 'mr2_mpa': mr2_mpa,
                'sum_h': sum_h, 'sum_hE13': sum_hE13,
                'mr_eq_mpa': mr_eq_mpa, 'cbr_eq': cbr_eq,
                'cbr_eq_design': math.floor(cbr_eq)
            }

        # ── แสดงผล ────────────────────────────────────────────────────
        if st.session_state.get('odemark_result'):
            res = st.session_state['odemark_result']
            st.markdown("---")
            st.markdown("#### ผลการคำนวณ Odemark")

            col_r1, col_r2, col_r3 = st.columns(3)
            with col_r1:
                st.metric("MR equivalent", f"{res['mr_eq_mpa']:.2f} MPa")
            with col_r2:
                st.metric("CBR equivalent (จากการคำนวณ)", f"{res['cbr_eq']:.2f} %")
            with col_r3:
                st.metric("CBR equivalent (ใช้ออกแบบ)", f"{res.get('cbr_eq_design', math.floor(res['cbr_eq']))} %")

            st.info(
                f"**สรุป:** ใช้ CBR_equivalent = **{res.get('cbr_eq_design', math.floor(res['cbr_eq']))} %** "
                f"แทนค่า CBR ดินเดิม ({st.session_state['cbr_design']:.2f} %) "
                f"ในการออกแบบโครงสร้างชั้นทาง"
            )

            with st.expander("แสดงวิธีการคำนวณ"):
                st.latex(
                    r"MR_{eq} = \left(\frac{\sum h_i \cdot MR_i^{1/3}}{\sum h_i}\right)^3"
                )
                st.write(f"- ชั้นที่ 1 ({res['mat1']}): h = {res['h1_cm']:.1f} cm, MR = {res['mr1_mpa']:.1f} MPa, MR¹ᐟ³ = {res['mr1_mpa']**(1/3):.4f}")
                st.write(f"- ชั้นที่ 2 (ดินถมคันทางใหม่): h = {res['h2_cm']:.1f} cm, MR = {res['mr2_mpa']:.2f} MPa, MR¹ᐟ³ = {res['mr2_mpa']**(1/3):.4f}")
                st.write(f"- Σh = {res['sum_h']:.1f} cm")
                st.write(f"- Σ(h·MR¹ᐟ³) = {res['sum_hE13']:.4f}")
                st.write(f"- MR_eq = ({res['sum_hE13']:.4f} / {res['sum_h']:.1f})³ = {res['mr_eq_mpa']:.2f} MPa")
                st.write(f"- CBR_eq = {res['mr_eq_mpa']:.2f} / {MPA_PER_CBR:.4f} = **{res['cbr_eq']:.2f} %**")

# Footer
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: gray;'>
    <p>สำหรับการวิเคราะห์ค่า CBR ดินฐานรากตามแนวสายทาง</p>
    <p>พัฒนาโดย รศ.ดร.อิทธิพล มีผล // ภาควิชาครุศาสตร์โยธา // คณะครุศาสตร์อุตสาหกรรม // มจพ.</p>
</div>
""", unsafe_allow_html=True)
