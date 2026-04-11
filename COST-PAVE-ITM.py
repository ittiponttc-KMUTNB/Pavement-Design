"""
ระบบวิเคราะห์ค่าก่อสร้างโครงสร้างชั้นทาง
Version 6.2.4 - Refactored
พัฒนาโดย: รศ.ดร.อิทธิพล มีผล — KMUTNB
- render_layer_editor() และ render_joint_editor() ใช้ st.data_editor แทน number_input loop
- ตัด Tab รูปภาพออก
- รวม get_default_*_layers() เป็น get_default_layers(ptype)
- รวม get_price_from_library() เป็นจุดเดียว
- ตรวจ syntax ด้วย ast.parse() ก่อน deploy - แก้ไขบทเกริ่นนำ
"""

import ast
import io
import json
from datetime import datetime

import numpy as np
import pandas as pd
import streamlit as st

try:
    import plotly.graph_objects as go
    PLOTLY_AVAILABLE = True
except ImportError:
    PLOTLY_AVAILABLE = False

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

# ── ตรวจ syntax ──────────────────────────────────────────────────────────────
_src = open(__file__).read()
try:
    ast.parse(_src)
except SyntaxError as _e:
    raise SyntaxError(f"[ast.parse] Syntax error in {__file__}: {_e}") from _e

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="วิเคราะห์ค่าก่อสร้างโครงสร้างชั้นทาง",
    page_icon="🛣️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── CSS ───────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Sans+Thai:wght@300;400;500;600&family=Space+Grotesk:wght@400;500;600;700&display=swap');

html, body, [class*="css"] {
    font-family: 'IBM Plex Sans Thai', sans-serif;
}

/* ── Header ── */
.main-header {
    background: linear-gradient(135deg, #0f2942 0%, #1a4a7a 60%, #0d7377 100%);
    border-radius: 16px;
    padding: 2rem 2.5rem;
    margin-bottom: 1.5rem;
    position: relative;
    overflow: hidden;
}
.main-header::before {
    content: '';
    position: absolute;
    top: -40px; right: -40px;
    width: 200px; height: 200px;
    border-radius: 50%;
    background: rgba(255,255,255,0.04);
}
.main-header h1 {
    font-family: 'Space Grotesk', sans-serif;
    color: #ffffff;
    font-size: 1.75rem;
    font-weight: 700;
    margin: 0;
    letter-spacing: -0.5px;
}
.main-header p {
    color: rgba(255,255,255,0.65);
    font-size: 0.9rem;
    margin: 0.35rem 0 0 0;
}

/* ── Metric cards ── */
.metric-card {
    background: #ffffff;
    border: 1px solid #e8edf2;
    border-radius: 12px;
    padding: 1.1rem 1.3rem;
    box-shadow: 0 2px 8px rgba(0,0,0,0.05);
}
.metric-card .label {
    font-size: 0.78rem;
    color: #6b7a8d;
    font-weight: 500;
    text-transform: uppercase;
    letter-spacing: 0.5px;
    margin-bottom: 0.3rem;
}
.metric-card .value {
    font-family: 'Space Grotesk', sans-serif;
    font-size: 1.5rem;
    font-weight: 700;
    color: #0f2942;
}
.metric-card .unit {
    font-size: 0.8rem;
    color: #8a95a3;
    margin-left: 4px;
}

/* ── Section card ── */
.section-card {
    background: #f8fafc;
    border: 1px solid #e2e8f0;
    border-left: 4px solid #1a4a7a;
    border-radius: 10px;
    padding: 1rem 1.2rem;
    margin-bottom: 1rem;
}

/* ── Selectbox green ── */
div[data-baseweb="select"] > div {
    background-color: #f0faf4 !important;
    border-color: #52b788 !important;
    border-radius: 8px !important;
}
div[data-baseweb="select"] span { color: #1b5e20 !important; font-weight: 500 !important; }
div[data-baseweb="menu"] { background-color: #f1f8e9 !important; }
div[data-baseweb="menu"] li { background-color: #f1f8e9 !important; color: #1b5e20 !important; }
div[data-baseweb="menu"] li:hover { background-color: #c8e6c9 !important; }

/* ── Checkbox accent ── */
label[data-baseweb="checkbox"] span { color: #0f2942 !important; }

/* ── data_editor clean ── */
[data-testid="stDataFrame"] { border-radius: 10px; overflow: hidden; }

/* ── Tab styling ── */
button[data-baseweb="tab"] {
    font-family: 'Space Grotesk', sans-serif !important;
    font-weight: 600 !important;
    font-size: 0.9rem !important;
}

/* ── Footer ── */
.footer {
    text-align: center;
    color: #94a3b8;
    font-size: 0.8rem;
    padding: 1.5rem 0 0.5rem;
    border-top: 1px solid #e2e8f0;
    margin-top: 2rem;
}
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════════
# DEFAULT PRICE TABLES
# ═══════════════════════════════════════════════════════════════════════════════

# ราคา AC ต่อตัน (บาท/ตัน) — ใช้ density 2.4 คำนวณ price table
DEFAULT_AC_TON_PRICES: dict = {
    'PMA Wearing Course': 3100,
    'AC Wearing Course':  2973,
    'AC Binder Course':   2929,
    'AC Base Course':     1795,
}
DEFAULT_AC_DENSITY: float = 2.4  # ตัน/ลบ.ม.

# ราคา AC คำนวณจากราคาต่อตัน × density 2.4 × ความหนา
# PMA=3,100 | AC Wearing=2,973 | AC Binder=2,929 | AC Base=1,795 บาท/ตัน
DEFAULT_AC_PRICES: dict = {
    'PMA Wearing Course': {2.5:186, 3:223, 4:298, 5:372, 6:446, 7:521, 8:595, 9:670, 10:744},
    'AC Wearing Course':  {2.5:178, 3:214, 4:285, 5:357, 6:428, 7:499, 8:571, 9:642, 10:714},
    'AC Binder Course':   {2.5:176, 3:211, 4:281, 5:351, 6:422, 7:492, 8:562, 9:633, 10:703},
    'AC Base Course':     {2.5:108, 3:129, 4:172, 5:215, 6:258, 7:302, 8:345, 9:388, 10:431},
}

# ราคาคอนกรีต บาท/ลบ.ม. — calibrate จากข้อมูลจริง 28cm
# JPCP=2,732 | JRCP=3,077 | CRCP=3,663 บาท/ลบ.ม.
DEFAULT_CONCRETE_CUM_PRICES: dict = {
    'JPCP': 2732,
    'JRCP': 3077,
    'CRCP': 3663,
}

def _calc_concrete_prices(cum_prices: dict) -> dict:
    """คำนวณ concrete price table (บาท/ตร.ม.) จาก บาท/ลบ.ม. × ความหนา"""
    thicknesses = [20, 25, 28, 30, 32, 35]
    return {
        pt: {t: round(cum * t / 100, 0) for t in thicknesses}
        for pt, cum in cum_prices.items()
    }

DEFAULT_CONCRETE_PRICES: dict = _calc_concrete_prices(DEFAULT_CONCRETE_CUM_PRICES)
# JPCP: {20:546, 25:683, 28:765, 30:820, 32:874, 35:956}
# JRCP: {20:615, 25:769, 28:862, 30:923, 32:985, 35:1077}
# CRCP: {20:733, 25:916, 28:1026, 30:1099, 32:1172, 35:1282}

DEFAULT_BASE_PRICES: dict = {
    'Cement Treated Base (UCS 40 ksc)':                    1096,
    'Cement Modified Crushed Rock Base (UCS 24.5 ksc)':    919,  # calibrated
    'Crushed Rock Base Course':                             583,
    'Soil Cement Subbase (UCS 7 ksc)':                     854,
    'Soil Aggregate Subbase':                              375,
    'Selected Material A':                                 375,
    'Embankment':                                          352,
    'Sand Embankment':                                     220,
    'Prime Coat':                                          37.47,
    'Non Woven Geotextile':                                78,
    'Wire Mesh':                                           100,
    'Tack Coat':                                           20,
}

BASE_MATERIAL_LIST = [
    'Cement Treated Base (UCS 40 ksc)',
    'Cement Modified Crushed Rock Base (UCS 24.5 ksc)',
    'Crushed Rock Base Course',
    'Soil Cement Subbase (UCS 7 ksc)',
    'Soil Aggregate Subbase',
    'Selected Material A',
    'Embankment',
    'Sand Embankment',
]

WEARING_OPTIONS = ['AC Wearing Course', 'PMA Wearing Course']

# ═══════════════════════════════════════════════════════════════════════════════
# PRICE LIBRARY HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def get_price_library() -> dict:
    """ดึง price_library จาก session_state หรือใช้ default"""
    if 'price_library' in st.session_state:
        return st.session_state['price_library']
    return {
        'ac_prices':       DEFAULT_AC_PRICES,
        'concrete_prices': DEFAULT_CONCRETE_PRICES,
        'base_prices':     DEFAULT_BASE_PRICES,
    }


def lookup_price(name: str, thickness: float, ptype: str = 'AC') -> float:
    """ดึงราคา (บาท/ตร.ม.) จาก library ตามชื่อวัสดุและความหนา
    สำหรับ base materials คืนราคา บาท/ลบ.ม. (caller แปลงเองด้วย ×t/100)
    """
    lib = get_price_library()
    n = name.lower()

    def _nearest(d: dict, t: float) -> float:
        if not d:
            return 0.0
        if t in d:
            return float(d[t])
        return float(d[min(d.keys(), key=lambda x: abs(x - t))])

    # AC surface
    if 'pma' in n:
        return _nearest(lib['ac_prices'].get('PMA Wearing Course', {}), thickness)
    if 'wearing' in n:
        return _nearest(lib['ac_prices'].get('AC Wearing Course', {}), thickness)
    if 'binder' in n:
        return _nearest(lib['ac_prices'].get('AC Binder Course', {}), thickness)
    if ('asphalt' in n and 'base' in n) or 'ac base' in n or 'interlayer' in n:
        return _nearest(lib['ac_prices'].get('AC Base Course', {}), thickness)

    # Concrete — คำนวณจาก บาท/ลบ.ม. × ความหนา (แม่นยำกว่า nearest lookup)
    _conc_cum_lib = st.session_state.get('concrete_cum_prices', DEFAULT_CONCRETE_CUM_PRICES)
    for ct in ('jpcp', 'jrcp', 'crcp'):
        if ct in n:
            cum = float(_conc_cum_lib.get(ct.upper(), DEFAULT_CONCRETE_CUM_PRICES.get(ct.upper(), 0)))
            return round(cum * thickness / 100, 2) if cum > 0 else _nearest(lib['concrete_prices'].get(ct.upper(), {}), thickness)
    # fallback by ptype
    if ptype in ('JPCP', 'JRCP', 'CRCP') and ('concrete' in n or 'ksc' in n or '350' in n or 'slab' in n):
        cum = float(_conc_cum_lib.get(ptype, DEFAULT_CONCRETE_CUM_PRICES.get(ptype, 0)))
        return round(cum * thickness / 100, 2) if cum > 0 else _nearest(lib['concrete_prices'].get(ptype, {}), thickness)

    # Base materials — คืนราคา/ลบ.ม.
    if 'tack' in n:
        return float(lib['base_prices'].get('Tack Coat', 20))
    if 'prime' in n:
        return float(lib['base_prices'].get('Prime Coat', 37.47))
    if 'geotextile' in n:
        return float(lib['base_prices'].get('Non Woven Geotextile', 78))
    if 'wire' in n:
        return float(lib['base_prices'].get('Wire Mesh', 100))

    # exact match ก่อน แล้วค่อย partial (ป้องกัน 'embankment' match 'sand embankment' ผิด)
    for key in BASE_MATERIAL_LIST:
        if key.lower() == n:
            return float(lib['base_prices'].get(key, 0))
    for key in BASE_MATERIAL_LIST:
        if key.lower() in n and len(key) > 5:
            return float(lib['base_prices'].get(key, 0))

    # fallback: ค้นใน base_prices โดยตรง
    for key, val in lib['base_prices'].items():
        if key.lower() == n:
            return float(val)

    return 0.0


# ═══════════════════════════════════════════════════════════════════════════════
# DEFAULT LAYERS / JOINTS
# ═══════════════════════════════════════════════════════════════════════════════

def get_default_layers(ptype: str, area_per_km: float = 22000) -> list:
    """คืน default layers list ตาม ptype: AC | JPCP | JRCP | CRCP"""
    lib = get_price_library()
    ac = lib['ac_prices']
    cp = lib['concrete_prices']
    bp = lib['base_prices']

    def ap(mat, t):  # ac price
        return ac.get(mat, {}).get(t, list(ac.get(mat, {0: 0}).values())[0] if ac.get(mat) else 0)

    if ptype == 'AC':
        return [
            {'name': 'AC Wearing Course',  'thickness': 7,  'unit': 'cm',    'quantity': area_per_km,   'qty_unit': 'sq.m', 'unit_cost': ap('AC Wearing Course', 7)},
            {'name': 'AC Binder Course',   'thickness': 7,  'unit': 'cm',    'quantity': area_per_km,   'qty_unit': 'sq.m', 'unit_cost': ap('AC Binder Course', 7)},
            {'name': 'AC Base Course',     'thickness': 10, 'unit': 'cm',    'quantity': area_per_km,   'qty_unit': 'sq.m', 'unit_cost': ap('AC Base Course', 10)},
            {'name': 'Tack Coat',          'thickness': 1,  'unit': 'Layer', 'quantity': area_per_km*2, 'qty_unit': 'sq.m', 'unit_cost': float(bp.get('Tack Coat', 20))},
            {'name': 'Prime Coat',         'thickness': 1,  'unit': 'Layer', 'quantity': area_per_km,   'qty_unit': 'sq.m', 'unit_cost': float(bp.get('Prime Coat', 37.47))},
        ]
    if ptype == 'JPCP':
        return [
            {'name': 'Concrete Slab (JPCP)', 'thickness': 28, 'unit': 'cm', 'quantity': area_per_km, 'qty_unit': 'sq.m', 'unit_cost': cp.get('JPCP', {}).get(28, 1000)},
        ]
    if ptype == 'JRCP':
        return [
            {'name': 'Concrete Slab (JRCP)', 'thickness': 28, 'unit': 'cm', 'quantity': area_per_km, 'qty_unit': 'sq.m', 'unit_cost': cp.get('JRCP', {}).get(28, 1002)},
        ]
    if ptype == 'CRCP':
        return [
            {'name': 'Concrete Slab (CRCP)', 'thickness': 25, 'unit': 'cm', 'quantity': area_per_km, 'qty_unit': 'sq.m', 'unit_cost': cp.get('CRCP', {}).get(25, 1245)},
        ]
    return []


def get_default_base_layers(ptype: str, area_per_km: float = 22000) -> list:
    """คืน default base layers list"""
    lib = get_price_library()
    bp = lib['base_prices']

    def b(mat, t):
        return float(bp.get(mat, 0)) * t / 100

    if ptype == 'AC':
        return [
            {'name': 'Crushed Rock Base Course', 'thickness': 20, 'unit': 'cm', 'quantity': area_per_km, 'qty_unit': 'sq.m',
             'unit_cost': b('Crushed Rock Base Course', 20), 'cost_cum': float(bp.get('Crushed Rock Base Course', 583))},
            {'name': 'Soil Aggregate Subbase',   'thickness': 30, 'unit': 'cm', 'quantity': area_per_km, 'qty_unit': 'sq.m',
             'unit_cost': b('Soil Aggregate Subbase', 30), 'cost_cum': float(bp.get('Soil Aggregate Subbase', 375))},
            {'name': 'Sand Embankment',           'thickness': 40, 'unit': 'cm', 'quantity': area_per_km, 'qty_unit': 'sq.m',
             'unit_cost': b('Sand Embankment', 40), 'cost_cum': float(bp.get('Sand Embankment', 220))},
        ]
    # Concrete types — JPCP/JRCP/CRCP ใช้ default เดียวกัน
    # (ผู้ใช้คัดลอกจาก JPCP ไปยัง JRCP/CRCP ได้ด้วยปุ่ม)
    return [
        {'name': 'Cement Modified Crushed Rock Base (UCS 24.5 ksc)', 'thickness': 20, 'unit': 'cm', 'quantity': area_per_km, 'qty_unit': 'sq.m',
         'unit_cost': b('Cement Modified Crushed Rock Base (UCS 24.5 ksc)', 20),
         'cost_cum': float(bp.get('Cement Modified Crushed Rock Base (UCS 24.5 ksc)', 864))},
        {'name': 'Soil Aggregate Subbase', 'thickness': 20, 'unit': 'cm', 'quantity': area_per_km, 'qty_unit': 'sq.m',
         'unit_cost': b('Soil Aggregate Subbase', 20),
         'cost_cum': float(bp.get('Soil Aggregate Subbase', 375))},
        {'name': 'Sand Embankment', 'thickness': 50, 'unit': 'cm', 'quantity': area_per_km, 'qty_unit': 'sq.m',
         'unit_cost': b('Sand Embankment', 50),
         'cost_cum': float(bp.get('Sand Embankment', 220))},
    ]


def get_default_joints(ptype: str, area_per_km: float = 22000, road_length: float = 1.0) -> list:
    """คืน default joints list ตาม ptype"""
    width_m = area_per_km / 1000
    if ptype == 'JPCP':
        trans_qty = (road_length * 1000 / 4) * width_m
        long_qty  = road_length * 1000
        return [
            {'name': 'Transverse Joint @4m', 'quantity': trans_qty, 'qty_unit': 'm', 'unit_cost': 430},
            {'name': 'Longitudinal Joint',   'quantity': long_qty,  'qty_unit': 'm', 'unit_cost': 120},
        ]
    if ptype == 'JRCP':
        trans_qty = (road_length * 1000 / 10) * width_m
        long_qty  = road_length * 1000
        return [
            {'name': 'Transverse Joint @10m', 'quantity': trans_qty, 'qty_unit': 'm', 'unit_cost': 430},
            {'name': 'Longitudinal Joint',    'quantity': long_qty,  'qty_unit': 'm', 'unit_cost': 120},
        ]
    if ptype == 'CRCP':
        long_qty = road_length * 1000
        return [
            {'name': 'Longitudinal Steel (CRCP)', 'quantity': long_qty, 'qty_unit': 'm', 'unit_cost': 200},
            {'name': 'Transverse Joint (End)',     'quantity': 0,        'qty_unit': 'm', 'unit_cost': 500},
        ]
    return []


# ═══════════════════════════════════════════════════════════════════════════════
# CALCULATE FUNCTIONS (ไม่แก้ logic เดิม)
# ═══════════════════════════════════════════════════════════════════════════════

BASE_KEYWORDS = ['crushed rock', 'soil aggregate', 'soil cement', 'cement modified',
                 'cement treated', 'selected material', 'embankment', 'sand embankment']


def calculate_layer_cost(layers: list, road_length_km: float = 1.0) -> tuple:
    """คำนวณค่าก่อสร้างจากชั้นโครงสร้าง — ไม่แก้ logic เดิม"""
    total = 0.0
    details = []
    for layer in layers:
        qty_raw   = float(layer['quantity'])
        unit_cost = float(layer['unit_cost'])
        cost      = qty_raw * unit_cost
        total    += cost

        name_lower = layer['name'].lower()
        is_base = any(kw in name_lower for kw in BASE_KEYWORDS)

        if is_base:
            thick_cm = float(layer.get('thickness', 1))
            u = layer.get('unit', 'cm').lower()
            if layer.get('cost_cum'):
                price_cum = float(layer['cost_cum'])
            elif thick_cm > 0 and u in ('cm', 'ซม.', 'ซ.ม.'):
                price_cum = unit_cost / (thick_cm / 100) if thick_cm > 0 else unit_cost
            else:
                price_cum = unit_cost
            qty_display = qty_raw * thick_cm / 100 if (thick_cm > 0 and u in ('cm', 'ซม.', 'ซ.ม.')) else qty_raw
            display_unit       = 'ลบ.ม.'
            display_price_str  = f"{price_cum:,.0f}"
            display_price_label = 'บาท/ลบ.ม.'
            qty_show = qty_display
        else:
            qty_show            = qty_raw
            display_unit        = 'ตร.ม.'
            display_price_str   = f"{unit_cost:,.0f}"
            display_price_label = 'บาท/ตร.ม.'

        details.append({
            'รายการ':              layer['name'],
            'ความหนา':             f"{layer['thickness']} {layer['unit']}",
            'ปริมาณ':              qty_show,
            'หน่วย':               display_unit,
            'ราคา/หน่วย':          unit_cost,
            'ราคา/หน่วย (แสดง)':   display_price_str,
            'หน่วยราคา':            display_price_label,
            'มูลค่า (บาท)':        cost,
        })
    return total, details


def calculate_joint_cost(joints: list, road_length_km: float = 1.0, include_joints: bool = True) -> tuple:
    """คำนวณค่ารอยต่อ — ไม่แก้ logic เดิม"""
    total = 0.0
    details = []
    for joint in joints:
        qty  = float(joint['quantity'])
        cost = qty * float(joint['unit_cost']) if include_joints else 0.0
        total += cost
        unit_th = 'ม.' if joint.get('qty_unit', 'm') == 'm' else joint.get('qty_unit', 'm')
        details.append({
            'รายการ':             joint['name'],
            'ความหนา':            '-',
            'ปริมาณ':             qty,
            'หน่วย':              unit_th,
            'ราคา/หน่วย':         float(joint['unit_cost']),
            'ราคา/หน่วย (แสดง)':  f"{float(joint['unit_cost']):,.0f}",
            'หน่วยราคา':           'บาท/ม.',
            'มูลค่า (บาท)':       cost,
        })
    return total, details


# ═══════════════════════════════════════════════════════════════════════════════
# RENDER LAYER EDITOR — ใช้ st.data_editor
# ═══════════════════════════════════════════════════════════════════════════════

def _auto_price_surface(name: str, thick: float, ptype: str) -> float:
    """คืนราคา บาท/ตร.ม. สำหรับ surface layer"""
    return lookup_price(name, thick, ptype)


def _auto_price_base(name: str, thick: float) -> tuple:
    """คืน (cost_per_sqm, cost_cum) สำหรับ base layer"""
    cost_cum = lookup_price(name, thick)  # บาท/ลบ.ม.
    cost_sqm = cost_cum * thick / 100 if thick > 0 else 0.0
    return cost_sqm, cost_cum


def render_layer_editor(
    ptype: str,
    key_prefix: str,
    total_width: float,
    road_length: float,
    v: int = 0,
) -> list:
    """
    แสดง editor โครงสร้างชั้นทางด้วย st.data_editor
    Pattern ที่ถูกต้อง:
      - เก็บ list of dict ใน session_state[sk_surf_rows] (surf) / sk_base_rows (base)
      - ส่ง key เดิมเข้า data_editor ทุก run
      - อ่านผลจาก session_state[editor_key] หลัง render (Streamlit อัพเดทให้อัตโนมัติ)
      - ไม่ overwrite session_state[sk_*_data] ซ้ำหลัง init ครั้งแรก
    """
    area_per_km = total_width * 1000
    proj_area   = area_per_km * road_length
    lib         = get_price_library()
    is_concrete = ptype in ('JPCP', 'JRCP', 'CRCP')

    updated_layers: list = []

    # ══════════════════════════════════════════════════════════════
    # SECTION A: ผิวทาง
    # AC: Wearing (radio) + Binder (fixed) + Base (checkbox) + Tack/Prime auto
    # Concrete: Slab (thickness input) เท่านั้น
    # ══════════════════════════════════════════════════════════════
    st.markdown('<div class="section-card"><b>🏗️ ผิวทาง</b> &nbsp;<span style="color:#6b7a8d;font-size:0.85rem">(บาท/ตร.ม.) — ราคาดึงจาก Library อัตโนมัติ แก้ราคาได้ที่ Tab 💰 ราคาวัสดุ</span></div>', unsafe_allow_html=True)

    def _get_price(name, thick):
        """ดึงราคาจาก library และ reset เมื่อ thickness หรือ price library เปลี่ยน"""
        pk           = f"{key_prefix}_p_{name.replace(' ','_')}_v{v}"
        prev_thick_k = f"{key_prefix}_pt_{name.replace(' ','_')}_v{v}"
        prev_ver_k   = f"{key_prefix}_pver_{name.replace(' ','_')}_v{v}"
        prev_t   = st.session_state.get(prev_thick_k, thick)
        # version = id ของ concrete_cum_prices + ac_ton_prices รวมกัน
        cur_ver  = id(st.session_state.get('concrete_cum_prices', {})) +                    id(st.session_state.get('ac_ton_prices', {})) +                    id(st.session_state.get('price_library', {}))
        prev_ver = st.session_state.get(prev_ver_k, None)
        if pk not in st.session_state or prev_t != thick or prev_ver != cur_ver:
            lib_p = lookup_price(name, thick, ptype)
            st.session_state[pk]           = float(lib_p) if lib_p > 0 else 0.0
            st.session_state[prev_thick_k] = thick
            st.session_state[prev_ver_k]   = cur_ver
        return pk

    ac_layer_count = 0

    if not is_concrete:
        _lib2 = get_price_library()

        # ── Header row ────────────────────────────────────────────
        hc = st.columns([3.5, 1.2, 1.8])
        hc[0].markdown("<span style='color:#6b7a8d;font-size:0.82rem;font-weight:600'>รายการ</span>", unsafe_allow_html=True)
        hc[1].markdown("<span style='color:#6b7a8d;font-size:0.82rem;font-weight:600'>หนา (cm)</span>", unsafe_allow_html=True)
        hc[2].markdown("<div style='color:#6b7a8d;font-size:0.82rem;font-weight:600;text-align:right'>ราคา (บาท/ตร.ม.)</div>", unsafe_allow_html=True)

        # ── Row 1: Wearing Course (radio AC / PMA) ────────────────
        _wr_key = f"{key_prefix}_wearing_type_v{v}"
        if _wr_key not in st.session_state:
            st.session_state[_wr_key] = 'AC Wearing Course'
        r1 = st.columns([3.5, 1.2, 1.8])
        with r1[0]:
            wearing_type = st.radio(
                "Wearing", ['AC Wearing Course', 'PMA Wearing Course'],
                index=0 if st.session_state[_wr_key] == 'AC Wearing Course' else 1,
                horizontal=True, key=f"{key_prefix}_wearing_radio_v{v}",
                label_visibility="collapsed",
            )
            st.session_state[_wr_key] = wearing_type
        with r1[1]:
            _wt_key = f"{key_prefix}_sthick_w_v{v}"
            if _wt_key not in st.session_state:
                st.session_state[_wt_key] = 5.0
            st.number_input("หนา W", min_value=0.5, max_value=20.0,
                step=0.5, format="%.1f", key=_wt_key, label_visibility="collapsed")
        wearing_thick = float(st.session_state[_wt_key])
        with r1[2]:
            _wp = _get_price(wearing_type, wearing_thick)
            wearing_price = float(st.session_state[_wp])
            st.markdown(f"<div style='padding:8px 0;font-weight:700;color:#0f2942;text-align:right'>{wearing_price:,.2f}</div>", unsafe_allow_html=True)
        ac_layer_count += 1
        updated_layers.append({
            'name': wearing_type, 'thickness': wearing_thick, 'unit': 'cm',
            'quantity': proj_area, 'qty_unit': 'sq.m',
            'unit_cost': wearing_price, 'cost_per_sqm': wearing_price,
        })

        # ── Row 2: Binder Course (fixed) ──────────────────────────
        r2 = st.columns([3.5, 1.2, 1.8])
        with r2[0]:
            st.markdown("<div style='padding:8px 0;font-weight:600'>AC Binder Course</div>", unsafe_allow_html=True)
        with r2[1]:
            _bt_key = f"{key_prefix}_sthick_b_v{v}"
            if _bt_key not in st.session_state:
                st.session_state[_bt_key] = 5.0
            st.number_input("หนา B", min_value=0.5, max_value=20.0,
                step=0.5, format="%.1f", key=_bt_key, label_visibility="collapsed")
        binder_thick = float(st.session_state[_bt_key])
        with r2[2]:
            _bp2 = _get_price('AC Binder Course', binder_thick)
            binder_price = float(st.session_state[_bp2])
            st.markdown(f"<div style='padding:8px 0;font-weight:700;color:#0f2942;text-align:right'>{binder_price:,.2f}</div>", unsafe_allow_html=True)
        ac_layer_count += 1
        updated_layers.append({
            'name': 'AC Binder Course', 'thickness': binder_thick, 'unit': 'cm',
            'quantity': proj_area, 'qty_unit': 'sq.m',
            'unit_cost': binder_price, 'cost_per_sqm': binder_price,
        })

        # ── Row 3: Base Course (checkbox) ─────────────────────────
        r3 = st.columns([3.5, 1.2, 1.8])
        with r3[0]:
            use_base = st.checkbox("AC Base Course", value=True,
                key=f"{key_prefix}_use_base_v{v}")
        if use_base:
            with r3[1]:
                _basethick_key = f"{key_prefix}_sthick_base_v{v}"
                if _basethick_key not in st.session_state:
                    st.session_state[_basethick_key] = 8.0
                st.number_input("หนา Base", min_value=0.5, max_value=30.0,
                    step=0.5, format="%.1f", key=_basethick_key, label_visibility="collapsed")
            base_thick = float(st.session_state[_basethick_key])
            with r3[2]:
                _basep = _get_price('AC Base Course', base_thick)
                base_price = float(st.session_state[_basep])
                st.markdown(f"<div style='padding:8px 0;font-weight:700;color:#0f2942;text-align:right'>{base_price:,.2f}</div>", unsafe_allow_html=True)
            ac_layer_count += 1
            updated_layers.append({
                'name': 'AC Base Course', 'thickness': base_thick, 'unit': 'cm',
                'quantity': proj_area, 'qty_unit': 'sq.m',
                'unit_cost': base_price, 'cost_per_sqm': base_price,
            })

        # ── Row 4: Tack Coat (auto qty) ───────────────────────────
        tack_times = max(ac_layer_count - 1, 1)
        tack_qty   = proj_area * tack_times
        _tk_key     = f"{key_prefix}_p_tack_v{v}"
        _tk_ver_key = f"{key_prefix}_p_tack_ver_v{v}"
        _lib_ver    = id(st.session_state.get('price_library', {}))
        if _tk_key not in st.session_state or st.session_state.get(_tk_ver_key) != _lib_ver:
            st.session_state[_tk_key]     = float(_lib2['base_prices'].get('Tack Coat', 20))
            st.session_state[_tk_ver_key] = _lib_ver
        r4 = st.columns([3.5, 1.2, 1.8])
        with r4[0]:
            st.markdown(
                f"<div style='padding:8px 0;color:#0f2942'>"
                f"Tack Coat "
                f"<span style='color:#6b7a8d;font-size:0.82rem'>"
                f"({tack_times} ครั้ง × {proj_area:,.0f} = {tack_qty:,.0f} ตร.ม.)</span></div>",
                unsafe_allow_html=True
            )
        with r4[1]:
            st.markdown(f"<div style='padding:8px 0;color:#94a3b8;font-size:0.85rem'>auto</div>", unsafe_allow_html=True)
        with r4[2]:
            tack_price = float(st.session_state[_tk_key])
            st.markdown(f"<div style='padding:8px 0;font-weight:700;color:#0f2942;text-align:right'>{tack_price:,.2f}</div>", unsafe_allow_html=True)
        updated_layers.append({
            'name': 'Tack Coat', 'thickness': 1, 'unit': 'Layer',
            'quantity': tack_qty, 'qty_unit': 'sq.m',
            'unit_cost': tack_price, 'cost_per_sqm': tack_price,
        })

        # ── Row 5: Prime Coat (auto qty) ──────────────────────────
        _pck     = f"{key_prefix}_p_primecoat_v{v}"
        _pck_ver = f"{key_prefix}_p_primecoat_ver_v{v}"
        if _pck not in st.session_state or st.session_state.get(_pck_ver) != _lib_ver:
            st.session_state[_pck]     = float(_lib2['base_prices'].get('Prime Coat', 37.47))
            st.session_state[_pck_ver] = _lib_ver
        r5 = st.columns([3.5, 1.2, 1.8])
        with r5[0]:
            st.markdown(
                f"<div style='padding:8px 0;color:#0f2942'>"
                f"Prime Coat "
                f"<span style='color:#6b7a8d;font-size:0.82rem'>({proj_area:,.0f} ตร.ม.)</span></div>",
                unsafe_allow_html=True
            )
        with r5[1]:
            st.markdown(f"<div style='padding:8px 0;color:#94a3b8;font-size:0.85rem'>auto</div>", unsafe_allow_html=True)
        with r5[2]:
            prime_price = float(st.session_state[_pck])
            st.markdown(f"<div style='padding:8px 0;font-weight:700;color:#0f2942;text-align:right'>{prime_price:,.2f}</div>", unsafe_allow_html=True)
        updated_layers.append({
            'name': 'Prime Coat', 'thickness': 1, 'unit': 'Layer',
            'quantity': proj_area, 'qty_unit': 'sq.m',
            'unit_cost': prime_price, 'cost_per_sqm': prime_price,
        })

    else:
        # ── Concrete slab ─────────────────────────────────────────
        slab_name = f'Concrete Slab ({ptype})'
        # Header
        _sh = st.columns([2, 1, 1.5])
        _sh[0].markdown("<span style='color:#6b7a8d;font-size:0.82rem;font-weight:600'>รายการ</span>", unsafe_allow_html=True)
        _sh[1].markdown("<span style='color:#6b7a8d;font-size:0.82rem;font-weight:600'>หนา (cm)</span>", unsafe_allow_html=True)
        _sh[2].markdown("<div style='color:#6b7a8d;font-size:0.82rem;font-weight:600;text-align:right'>ราคา (บาท/ตร.ม.)</div>", unsafe_allow_html=True)
        c1, c2, c3 = st.columns([2, 1, 1.5])
        with c1:
            st.markdown(f"**{slab_name}**")
        with c2:
            _slab_thick_key = f"{key_prefix}_sthick_slab_v{v}"
            _default_slab_t = 28.0 if ptype in ('JPCP','JRCP') else 25.0
            if _slab_thick_key not in st.session_state:
                st.session_state[_slab_thick_key] = _default_slab_t
            st.number_input("ความหนาแผ่น (cm)", min_value=15.0, max_value=50.0,
                step=1.0, format="%.0f", key=_slab_thick_key, label_visibility="collapsed")
        slab_thick = float(st.session_state[_slab_thick_key])
        with c3:
            _slabp = _get_price(slab_name, slab_thick)
            slab_price = float(st.session_state[_slabp])
            st.markdown(f"<div style='padding:8px 0;font-weight:700;color:#0f2942;text-align:right'>{slab_price:,.2f}</div>", unsafe_allow_html=True)
        updated_layers.append({
            'name': slab_name, 'thickness': slab_thick, 'unit': 'cm',
            'quantity': proj_area, 'qty_unit': 'sq.m',
            'unit_cost': slab_price, 'cost_per_sqm': slab_price,
        })

    # ══════════════════════════════════════════════════════════════
    # SECTION B: Checkboxes วัสดุประกอบ (คอนกรีตทุกประเภท)
    # AC Interlayer / Prime Coat / Non Woven Geotextile / Wire Mesh
    # ══════════════════════════════════════════════════════════════
    if is_concrete:
        st.markdown("---")
        st.markdown("**🔧 วัสดุประกอบ** — ติ๊กเลือกและปรับแก้ได้")

        # แถวที่ 1: AC Interlayer + Prime Coat
        col_cb1, col_cb2 = st.columns(2)
        with col_cb1:
            use_acil = st.checkbox("AC Interlayer รองใต้แผ่นคอนกรีต", value=True,
                key=f"{key_prefix}_use_acil_v{v}", help="ชั้น AC รองใต้แผ่น ทั่วไป 5 cm")
        with col_cb2:
            use_pc = st.checkbox("Prime Coat", value=True,
                key=f"{key_prefix}_use_pc_v{v}", help="ราดบน Base ก่อนปู AC Interlayer (บาท/ตร.ม.)")

        # แถวที่ 2: Non Woven Geotextile + Wire Mesh (Wire Mesh เฉพาะ JRCP/CRCP)
        col_cb3, col_cb4 = st.columns(2)
        with col_cb3:
            use_geo = st.checkbox("Non Woven Geotextile", value=True,
                key=f"{key_prefix}_use_geo_v{v}", help="แผ่น Geotextile รองใต้แผ่นคอนกรีต")
        with col_cb4:
            use_wire = st.checkbox(
                "Wire Mesh" if ptype != 'JPCP' else "Wire Mesh (ไม่ใช้กับ JPCP)",
                value=(ptype != 'JPCP'),
                disabled=(ptype == 'JPCP'),
                key=f"{key_prefix}_use_wire_v{v}",
                help="ตะแกรงเหล็กในแผ่นคอนกรีต (JRCP/CRCP เท่านั้น)"
            )

        # ── AC Interlayer ──
        if use_acil:
            c1, c2 = st.columns([2, 2])
            _acil_thick_key = f"{key_prefix}_acil_thick_v{v}"
            _acil_price_key = f"{key_prefix}_acil_price_v{v}"
            if _acil_thick_key not in st.session_state:
                st.session_state[_acil_thick_key] = 5.0
            with c1:
                st.number_input("ความหนา AC Interlayer (cm)",
                    min_value=1.0, max_value=15.0, step=1.0,
                    key=_acil_thick_key)
            acil_thick = float(st.session_state[_acil_thick_key])
            # ดึงราคาจาก library ตามความหนาปัจจุบัน
            _acil_lib = lookup_price('AC Binder Course', acil_thick, ptype)
            if _acil_lib == 0:
                _acil_lib = 251.0
            # reset ราคาเมื่อความหนาเปลี่ยน
            _prev_acil_thick = st.session_state.get(f"{key_prefix}_acil_thick_prev_v{v}", acil_thick)
            if acil_thick != _prev_acil_thick or _acil_price_key not in st.session_state:
                st.session_state[_acil_price_key] = float(_acil_lib)
            st.session_state[f"{key_prefix}_acil_thick_prev_v{v}"] = acil_thick
            with c2:
                acil_price = float(st.session_state[_acil_price_key])
                st.markdown(f"AC Interlayer {acil_thick:.0f} cm → **{acil_price:,.2f}** บาท/ตร.ม.")
            updated_layers.append({
                'name': f'AC Interlayer ({acil_thick:.0f} cm)',
                'thickness': acil_thick, 'unit': 'cm',
                'quantity': proj_area, 'qty_unit': 'sq.m',
                'unit_cost': acil_price, 'cost_per_sqm': acil_price,
            })

        # ── Prime Coat ──
        if use_pc:
            _pc_key     = f"{key_prefix}_pc_price_v{v}"
            _pc_ver_key = f"{key_prefix}_pc_price_ver_v{v}"
            _pc_lib     = float(lib['base_prices'].get('Prime Coat', 37.47))
            _lib_ver2   = id(st.session_state.get('price_library', {}))
            if _pc_key not in st.session_state or st.session_state.get(_pc_ver_key) != _lib_ver2:
                st.session_state[_pc_key]     = _pc_lib
                st.session_state[_pc_ver_key] = _lib_ver2
            c1, c2 = st.columns([3, 1])
            with c1:
                st.caption("Prime Coat — ราดบน Base Course ก่อนปู AC Interlayer")
            with c2:
                pc_price = float(st.session_state[_pc_key])
                st.markdown(f"<div style='padding:6px 0;font-weight:700;color:#0f2942;text-align:right'>{pc_price:,.2f}</div>", unsafe_allow_html=True)
            updated_layers.append({
                'name': 'Prime Coat', 'thickness': 1, 'unit': 'Layer',
                'quantity': proj_area, 'qty_unit': 'sq.m',
                'unit_cost': pc_price, 'cost_per_sqm': pc_price,
            })

        # ── Non Woven Geotextile ──
        if use_geo:
            _geo_key     = f"{key_prefix}_geo_price_v{v}"
            _geo_ver_key = f"{key_prefix}_geo_price_ver_v{v}"
            _geo_lib     = float(lib['base_prices'].get('Non Woven Geotextile', 78))
            _lib_ver2    = id(st.session_state.get('price_library', {}))
            if _geo_key not in st.session_state or st.session_state.get(_geo_ver_key) != _lib_ver2:
                st.session_state[_geo_key]     = _geo_lib
                st.session_state[_geo_ver_key] = _lib_ver2
            c1, c2 = st.columns([3, 1])
            with c1:
                st.caption("Non Woven Geotextile — รองใต้แผ่นคอนกรีต")
            with c2:
                geo_price = float(st.session_state[_geo_key])
                st.markdown(f"<div style='padding:6px 0;font-weight:700;color:#0f2942;text-align:right'>{geo_price:,.2f}</div>", unsafe_allow_html=True)
            updated_layers.append({
                'name': 'Non Woven Geotextile', 'thickness': 1, 'unit': 'ชั้น',
                'quantity': proj_area, 'qty_unit': 'sq.m',
                'unit_cost': geo_price, 'cost_per_sqm': geo_price,
            })

        # ── Wire Mesh (JRCP/CRCP เท่านั้น) ──
        if use_wire and ptype != 'JPCP':
            _wire_key     = f"{key_prefix}_wire_price_v{v}"
            _wire_ver_key = f"{key_prefix}_wire_price_ver_v{v}"
            _wire_lib     = float(lib['base_prices'].get('Wire Mesh', 100))
            _lib_ver2     = id(st.session_state.get('price_library', {}))
            if _wire_key not in st.session_state or st.session_state.get(_wire_ver_key) != _lib_ver2:
                st.session_state[_wire_key]     = _wire_lib
                st.session_state[_wire_ver_key] = _lib_ver2
            c1, c2 = st.columns([3, 1])
            with c1:
                st.caption("Wire Mesh — ตะแกรงเหล็กในแผ่นคอนกรีต")
            with c2:
                wire_price = float(st.session_state[_wire_key])
                st.markdown(f"<div style='padding:6px 0;font-weight:700;color:#0f2942;text-align:right'>{wire_price:,.2f}</div>", unsafe_allow_html=True)
            updated_layers.append({
                'name': 'Wire Mesh', 'thickness': 1, 'unit': 'ชั้น',
                'quantity': proj_area, 'qty_unit': 'sq.m',
                'unit_cost': wire_price, 'cost_per_sqm': wire_price,
            })

    # ══════════════════════════════════════════════════════════════
    # SECTION C: พื้นทาง / รองพื้นทาง
    # ใช้ selectbox + number_input (ไม่มี reset ปัญหา)
    # ══════════════════════════════════════════════════════════════
    st.markdown("---")
    st.caption("💡 ราคาดึงจาก Library อัตโนมัติ — แก้ราคาได้ที่ Tab **💰 ราคาวัสดุ**")

    # ── Header + ปุ่มคัดลอกจาก JPCP ──────────────────────────────
    sk_base_rows  = f"{key_prefix}_base_rows_v{v}"
    sk_price_ver  = f"{key_prefix}_base_price_ver_v{v}"
    sk_copy_flag  = f"{key_prefix}_do_copy_base_v{v}"
    cur_price_ver = id(st.session_state.get('price_library', {}))

    hcol1, hcol2 = st.columns([3, 1])
    with hcol1:
        st.markdown(
            '<div class="section-card"><b>🪨 พื้นทาง / รองพื้นทาง</b> ' +
            '<span style="color:#6b7a8d;font-size:0.85rem">บาท/ลบ.ม. × ความหนา = บาท/ตร.ม.</span></div>',
            unsafe_allow_html=True
        )
    with hcol2:
        if is_concrete and ptype != 'JPCP':
            if st.button("📋 คัดลอก Base จาก JPCP",
                         key=f"{key_prefix}_copy_base_v{v}",
                         use_container_width=True, type="secondary"):
                st.session_state[sk_copy_flag] = True

    # ── ตรวจ copy flag — ทำก่อน init เพื่อ override _cur_rows ในรัน เดียวกัน ──
    if st.session_state.get(sk_copy_flag):
        jpcp_sk = f"jpcp_base_rows_v{v}"
        if jpcp_sk in st.session_state:
            copied = [dict(r) for r in st.session_state[jpcp_sk]]
            st.session_state[sk_base_rows] = copied
            # เพิ่ม copy_version เพื่อเปลี่ยน widget keys ทั้งหมด
            _cv_key = f"{key_prefix}_base_cv"
            st.session_state[_cv_key] = st.session_state.get(_cv_key, 0) + 1
            st.session_state[sk_copy_flag] = False
            st.success("✅ คัดลอก Base จาก JPCP สำเร็จ")
        else:
            st.warning("⚠️ ยังไม่มีข้อมูล Base ของ JPCP — กรุณาตั้งค่า JPCP ก่อน")
            st.session_state[sk_copy_flag] = False

    # ── init base rows ─────────────────────────────────────────────
    if sk_base_rows not in st.session_state:
        _def_base = get_default_base_layers(ptype, area_per_km)
        st.session_state[sk_base_rows] = [
            {'name': r['name'], 'thickness': r['thickness'], 'cost_cum': r['cost_cum']}
            for r in _def_base
        ]
        st.session_state[sk_price_ver] = cur_price_ver
    elif st.session_state.get(sk_price_ver) != cur_price_ver:
        _lib_fresh = {m: lookup_price(m, 20) for m in BASE_MATERIAL_LIST}
        st.session_state[sk_base_rows] = [
            {'name': r['name'], 'thickness': r['thickness'],
             'cost_cum': _lib_fresh.get(r['name'], lookup_price(r['name'], 20))}
            for r in st.session_state[sk_base_rows]
        ]
        st.session_state[sk_price_ver] = cur_price_ver

    # jpcp_base_rows จะถูกบันทึกหลัง render loop (จาก new_rows ที่ถูกต้อง)

    # cv = copy version — เปลี่ยนทุกครั้งที่ copy เพื่อ force widget key ใหม่
    _cv    = st.session_state.get(f"{key_prefix}_base_cv", 0)
    _cv_sf = f"cv{_cv}"   # suffix สำหรับ widget keys

    # ── จำนวนชั้น ──────────────────────────────────────────────────
    _cur_rows = st.session_state[sk_base_rows]
    _nb_key   = f"{key_prefix}_num_base_{_cv_sf}_v{v}"
    if _nb_key not in st.session_state:
        st.session_state[_nb_key] = len(_cur_rows)
    st.number_input(
        "จำนวนชั้นพื้นทาง/รองพื้นทาง",
        min_value=0, max_value=8, step=1, key=_nb_key,
    )
    num_base = int(st.session_state[_nb_key])

    # ── Header columns ─────────────────────────────────────────────
    _lib_cum = {m: lookup_price(m, 20) for m in BASE_MATERIAL_LIST}
    _base_cols_ratio = [3, 1.2, 1.5, 1.5]
    hdr = st.columns(_base_cols_ratio)
    _hdr_style = "color:#6b7a8d;font-size:0.82rem;font-weight:600"
    hdr[0].markdown(f"<span style='{_hdr_style}'>วัสดุ</span>", unsafe_allow_html=True)
    hdr[1].markdown(f"<span style='{_hdr_style}'>หนา (cm)</span>", unsafe_allow_html=True)
    hdr[2].markdown(f"<span style='{_hdr_style};display:block;text-align:right'>ราคา (บาท/ลบ.ม.)</span>", unsafe_allow_html=True)
    hdr[3].markdown(f"<span style='{_hdr_style};display:block;text-align:right'>ราคา (บาท/ตร.ม.)</span>", unsafe_allow_html=True)

    # ── rows ───────────────────────────────────────────────────────
    # sk_prev_names เก็บชื่อวัสดุของ run ก่อนหน้า เพื่อ detect การเปลี่ยนวัสดุ
    sk_prev_names = f"{key_prefix}_prev_names_v{v}"
    prev_names = st.session_state.get(sk_prev_names, {})

    new_rows = []
    for i in range(num_base):
        # ดึงค่าเดิมถ้ามี
        prev = _cur_rows[i] if i < len(_cur_rows) else {
            'name': BASE_MATERIAL_LIST[0], 'thickness': 20.0,
            'cost_cum': _lib_cum.get(BASE_MATERIAL_LIST[0], 0)
        }
        prev_name  = str(prev.get('name', BASE_MATERIAL_LIST[0]))
        prev_thick = float(prev.get('thickness', 20.0) or 20.0)
        prev_cum   = float(prev.get('cost_cum', 0) or 0)

        if prev_cum == 0:
            prev_cum = _lib_cum.get(prev_name, 0)

        cols = st.columns(_base_cols_ratio)

        with cols[0]:
            try:
                name_idx = BASE_MATERIAL_LIST.index(prev_name)
            except ValueError:
                name_idx = 0
            sel_name = st.selectbox(
                "วัสดุ", BASE_MATERIAL_LIST, index=name_idx,
                key=f"{key_prefix}_bname_{i}_{_cv_sf}_v{v}",
                label_visibility="collapsed",
            )

        with cols[1]:
            bthick_key = f"{key_prefix}_bthick_{i}_{_cv_sf}_v{v}"
            if bthick_key not in st.session_state:
                st.session_state[bthick_key] = float(prev_thick)
            st.number_input(
                "หนา", min_value=0.0, step=5.0, format="%.0f",
                key=bthick_key, label_visibility="collapsed",
            )
            sel_thick = float(st.session_state.get(bthick_key, prev_thick))

        # ตรวจว่า name เพิ่งเปลี่ยนใน run นี้ไหม
        last_rendered_name = prev_names.get(i, prev_name)
        name_just_changed  = (sel_name != last_rendered_name)

        if name_just_changed:
            lib_cum = _lib_cum.get(sel_name, 0)
            wkey = f"{key_prefix}_bcum_{i}_{_cv_sf}_v{v}"
            if wkey in st.session_state:
                del st.session_state[wkey]
            prev_cum = lib_cum

        bcum_key = f"{key_prefix}_bcum_{i}_{_cv_sf}_v{v}"
        if bcum_key not in st.session_state:
            st.session_state[bcum_key] = float(prev_cum)

        # ราคา บาท/ลบ.ม. — read-only ดึงจาก Library (แก้ได้ที่ Tab 💰 ราคาวัสดุ)
        sel_cum = float(st.session_state.get(bcum_key, prev_cum))
        with cols[2]:
            st.markdown(
                f'<div style="padding:8px 0;font-weight:700;color:#0f2942;'
                f'text-align:right">{sel_cum:,.0f}</div>',
                unsafe_allow_html=True
            )
        cost_sqm = sel_cum * sel_thick / 100 if sel_thick > 0 else 0.0

        with cols[3]:
            st.markdown(
                f'<div style="padding:8px 0;font-weight:700;color:#0f2942;text-align:right">{cost_sqm:,.2f}</div>',
                unsafe_allow_html=True
            )

        new_rows.append({'name': sel_name, 'thickness': sel_thick, 'cost_cum': sel_cum})

        if sel_thick > 0 and sel_name:
            updated_layers.append({
                'name':         sel_name,
                'thickness':    sel_thick,
                'unit':         'cm',
                'quantity':     proj_area,
                'qty_unit':     'sq.m',
                'unit_cost':    cost_sqm,
                'cost_per_sqm': cost_sqm,
                'cost_cum':     sel_cum,
            })

    # บันทึก rows กลับ session_state
    st.session_state[sk_base_rows] = new_rows
    # บันทึก prev_names สำหรับ run ถัดไป
    st.session_state[sk_prev_names] = {i: r['name'] for i, r in enumerate(new_rows)}
    # บันทึก JPCP base rows จาก new_rows (ค่าที่ user แก้จริงใน run นี้)
    if ptype == 'JPCP':
        st.session_state[f"jpcp_base_rows_v{v}"] = [dict(r) for r in new_rows]

    return updated_layers


# ═══════════════════════════════════════════════════════════════════════════════
# RENDER JOINT EDITOR — ใช้ st.data_editor
# ═══════════════════════════════════════════════════════════════════════════════

def render_joint_editor(
    ptype: str,
    key_prefix: str,
    area_per_km: float,
    road_length: float,
    v: int = 0,
) -> tuple:
    """
    แสดง editor รอยต่อด้วย st.data_editor
    คืน (updated_joints, include_joints)
    """
    width_m    = area_per_km / 1000
    lane_w     = st.session_state.get('project_info', {}).get('lane_width', 3.5)
    if not lane_w or lane_w <= 0:
        lane_w = 3.5
    total_area = area_per_km * road_length

    # spacing label
    if ptype == 'JPCP':
        joint_label = 'Transverse Joint @4m'
        spacing = 4
    else:
        joint_label = 'Transverse Joint @10m'
        spacing = 10

    # init default เพียงครั้งเดียว — คำนวณ qty อัตโนมัติแล้วใส่ใน init
    sk_joint_init = f"{key_prefix}_joint_init_v{v}"

    if ptype in ('JPCP', 'JRCP'):
        auto_trans = (road_length * 1000 / spacing) * width_m
        auto_long  = max(1, round(width_m / lane_w) - 1) * road_length * 1000
    else:
        auto_trans = 0.0
        auto_long  = max(1, round(width_m / lane_w) - 1) * road_length * 1000

    if sk_joint_init not in st.session_state:
        defaults_j = get_default_joints(ptype, area_per_km, road_length)
        rows_init = []
        for j in defaults_j:
            name = j['name']
            qty  = j['quantity']
            if 'transverse' in name.lower():
                qty = auto_trans if ptype in ('JPCP', 'JRCP') else qty
            elif 'longitudinal' in name.lower() or 'steel' in name.lower():
                qty = auto_long
            rows_init.append({'name': name, 'quantity': qty, 'unit_cost': float(j['unit_cost'])})
        st.session_state[sk_joint_init] = pd.DataFrame(rows_init)

    st.markdown("---")
    col_h1, col_h2 = st.columns([3, 1])
    with col_h1:
        if ptype == 'CRCP':
            st.markdown('<div class="section-card"><b>⛓️ Longitudinal Steel & Transverse Joint (CRCP)</b></div>', unsafe_allow_html=True)
        else:
            st.markdown(f'<div class="section-card"><b>🔗 รอยต่อ (Joints) — {ptype} ระยะ {spacing} ม.</b></div>', unsafe_allow_html=True)
    with col_h2:
        _cb_label = "รวมราคา Steel & Joints" if ptype == 'CRCP' else "รวมราคา Joints"
        include_joints = st.checkbox(_cb_label, value=True, key=f"{key_prefix}_include_joints_v{v}")

    ek_joint = f"{key_prefix}_joint_editor_v{v}"

    st.data_editor(
        st.session_state[sk_joint_init],
        column_config={
            'name':      st.column_config.TextColumn('รายการ', width='large'),
            'quantity':  st.column_config.NumberColumn('ปริมาณ (ม.)', min_value=0.0, step=100.0, format='%.0f'),
            'unit_cost': st.column_config.NumberColumn('ราคา/ม. (บาท)', min_value=0.0, step=10.0, format='%.0f'),
        },
        num_rows='dynamic',
        use_container_width=True,
        key=ek_joint,
        hide_index=True,
        on_change=None,
    )

    # อ่านผล joint editor
    _joint_state = st.session_state.get(ek_joint, {})
    edited_joint = st.session_state[sk_joint_init].copy()
    if isinstance(_joint_state, dict):
        for idx_str, changes in _joint_state.get("edited_rows", {}).items():
            idx = int(idx_str)
            if idx < len(edited_joint):
                for col, val in changes.items():
                    edited_joint.at[idx, col] = val
        for new_row in _joint_state.get("added_rows", []):
            edited_joint = pd.concat(
                [edited_joint, pd.DataFrame([new_row])], ignore_index=True
            )
        del_idxs = _joint_state.get("deleted_rows", [])
        if del_idxs:
            edited_joint = edited_joint.drop(index=del_idxs).reset_index(drop=True)

    # แสดงราคา/ตร.ม. ใต้ตาราง
    joint_total_cost = 0.0
    updated_joints = []
    for _, row in edited_joint.iterrows():
        name = str(row.get('name', '') or '')
        qty  = float(row.get('quantity', 0) or 0)
        uc   = float(row.get('unit_cost', 0) or 0)
        cpsqm = (qty * uc / total_area) if total_area > 0 else 0.0
        joint_total_cost += qty * uc
        updated_joints.append({
            'name': name, 'quantity': qty,
            'qty_unit': 'm', 'unit_cost': uc,
            'cost_per_sqm': cpsqm,
        })

    if total_area > 0:
        st.caption(f"รวม Joints = **{joint_total_cost/total_area:,.2f}** บาท/ตร.ม. | **{joint_total_cost/1e6:,.3f}** ล้านบาท/โครงการ")

    return updated_joints, include_joints


# ═══════════════════════════════════════════════════════════════════════════════
# WORD REPORT
# ═══════════════════════════════════════════════════════════════════════════════

def _set_run_font(run, size: int = 16, bold: bool = False, italic: bool = False):
    run.font.name = 'TH SarabunPSK'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    for attr in ('w:eastAsia', 'w:ascii', 'w:hAnsi'):
        rFonts.set(qn(attr), 'TH SarabunPSK')


def generate_word_report(
    project_info: dict,
    all_details: dict,
    report_type: str = 'materials',
    chapter_num: str = '4',
    section_start: str = '4.7',
    intro_text: str = '',
) -> 'Document':
    """สร้างรายงาน Word — report_type: 'materials' | 'consultant'"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx ไม่สามารถใช้งานได้")

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'TH SarabunPSK'
    style.font.size = Pt(16)

    length = float(project_info.get('length', 1))

    def add_heading(text, size=16, bold=True, underline=False, space_before=6, space_after=3):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.space_after  = Pt(space_after)
        run = para.add_run(text)
        _set_run_font(run, size=size, bold=bold)
        run.underline = underline
        return para

    if report_type == 'consultant':
        add_heading(f"{section_start} รายงานวัสดุและราคาโครงสร้างชั้นทาง",
                    size=18, bold=True, underline=True, space_before=12, space_after=6)
        add_heading(f"{section_start}.1 ข้อมูลของถนน", size=16, bold=True, underline=True, space_before=8)
        if intro_text:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(1.0)
            _set_run_font(p.add_run(intro_text), size=16)
    else:
        doc.add_heading('รายงานวัสดุและราคาโครงสร้างชั้นทาง', 0)
        doc.add_heading('1. ข้อมูลโครงการ', level=1)

    fields = [
        ("ชื่อโครงการ",      project_info.get('name', '-')),
        ("ความยาวถนน",      f"{length:.2f} กม."),
        ("ความกว้างรวม",    f"{project_info.get('total_width', 0):.2f} ม."),
        ("จำนวนช่องจราจร",  f"{project_info.get('num_lanes', 4)} ช่อง"),
        ("อายุออกแบบ",      f"{project_info.get('design_life', 20)} ปี"),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.0)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r1 = p.add_run(f"{label}: ")
        _set_run_font(r1, size=16, bold=True)
        r2 = p.add_run(value)
        _set_run_font(r2, size=16)

    sec_detail = f"{section_start}.2" if report_type == 'consultant' else '2.'
    add_heading(f"{sec_detail} รายละเอียดวัสดุและราคา", size=16, bold=True, underline=True, space_before=10)

    summary_data = []
    for ptype, data in all_details.items():
        sname   = data.get('name', ptype)
        details = data.get('details', [])

        add_heading(f"ผิวทางประเภท {sname}", size=16, bold=True, space_before=6, space_after=2)
        if data.get('name_detail', '') and data['name_detail'] != sname:
            p_sub = doc.add_paragraph()
            r_sub = p_sub.add_run(data['name_detail'])
            _set_run_font(r_sub, size=14, italic=True)

        # กรอง detail ที่มูลค่า = 0 ออก (เช่น Joint ที่ปิด)
        details = [d for d in details if float(d.get('มูลค่า (บาท)', 0)) != 0]

        if details:
            table = doc.add_table(rows=len(details) + 2, cols=5)
            table.style = 'Table Grid'
            col_widths = [Cm(6.5), Cm(2.5), Cm(1.8), Cm(3.5), Cm(3.5)]
            for row in table.rows:
                for idx, cell in enumerate(row.cells):
                    cell.width = col_widths[idx]

            headers = ['รายการ', 'ปริมาณ', 'หน่วย', 'ราคา/หน่วย (บาท)', 'มูลค่า (บาท)']
            for j, h in enumerate(headers):
                p_h = table.rows[0].cells[j].paragraphs[0]
                p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_run_font(p_h.add_run(h), size=15, bold=True)

            subtotal = 0.0
            for i, d in enumerate(details):
                rc = table.rows[i + 1].cells
                _set_run_font(rc[0].paragraphs[0].add_run(str(d['รายการ'])), size=15)
                rc[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _set_run_font(rc[1].paragraphs[0].add_run(f"{d['ปริมาณ']:,.0f}"), size=15)
                rc[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_run_font(rc[2].paragraphs[0].add_run(d.get('หน่วย', 'ตร.ม.')), size=15)
                rc[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _set_run_font(rc[3].paragraphs[0].add_run(d.get('ราคา/หน่วย (แสดง)', '')), size=15)
                _set_run_font(rc[3].paragraphs[0].add_run(f" ({d.get('หน่วยราคา', 'บาท/ตร.ม.')})"), size=12)
                rc[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _set_run_font(rc[4].paragraphs[0].add_run(f"{d['มูลค่า (บาท)']:,.0f}"), size=15)
                subtotal += d['มูลค่า (บาท)']

            # แถวรวม
            lr = table.rows[len(details) + 1]
            lr.cells[0].merge(lr.cells[3])
            lr.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_run_font(lr.cells[0].paragraphs[0].add_run(f"รวม {sname}"), size=15, bold=True)
            lr.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_run_font(lr.cells[4].paragraphs[0].add_run(f"{subtotal:,.0f}"), size=15, bold=True)

            doc.add_paragraph()
            summary_data.append({
                'name':              sname,
                'total_value':       subtotal,
                'cost_per_km':       data.get('cost_per_km', 0),
                'cost_sqm':          data.get('cost_sqm', 0),
            })

    sec_sum = f"{section_start}.3" if report_type == 'consultant' else '3.'
    add_heading(f"{sec_sum} สรุปค่าใช้จ่าย", size=16, bold=True, underline=True, space_before=10)

    if summary_data:
        sum_table = doc.add_table(rows=len(summary_data) + 1, cols=4)
        sum_table.style = 'Table Grid'
        hdrs = ['ชนิดโครงสร้าง', 'มูลค่ารวม/กม. (บาท)', 'ราคา/กม. (ล้านบาท)', 'ราคา/ตร.ม. (บาท)']
        for j, h in enumerate(hdrs):
            ph = sum_table.rows[0].cells[j].paragraphs[0]
            ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_run_font(ph.add_run(h), size=15, bold=True)
        for i, item in enumerate(summary_data):
            tpk = item['total_value'] / length if length > 0 else 0
            _sum_vals   = [item['name'], f"{tpk:,.0f}", f"{item['cost_per_km']:.3f}", f"{item['cost_sqm']:,.2f}"]
            _sum_aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT,
                           WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT]
            for j, (val, align) in enumerate(zip(_sum_vals, _sum_aligns)):
                _cell = sum_table.rows[i+1].cells[j]
                _cell.text = ''
                _p = _cell.paragraphs[0]
                _p.alignment = align
                _set_run_font(_p.add_run(val), size=15)

    doc.add_paragraph()
    p_date = doc.add_paragraph()
    _set_run_font(p_date.add_run(f"รายงานสร้างเมื่อ: {datetime.now().strftime('%d/%m/%Y %H:%M')}"), size=14)
    p_by = doc.add_paragraph()
    _set_run_font(p_by.add_run("พัฒนาโดย รศ.ดร.อิทธิพล มีผล — KMUTNB"), size=14)
    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# EXCEL TEMPLATE
# ═══════════════════════════════════════════════════════════════════════════════

def generate_excel_template() -> bytes:
    """สร้าง Excel template
    AC_Prices: Material | Price (Baht/ton)  ← ใช้ราคาต่อตันเท่านั้น
    Concrete_Prices: Type | 25cm ... 35cm   (บาท/ตร.ม.)
    Base_Materials: Material | Price (Baht/cu.m)
    """
    lib = get_price_library()

    # AC_Prices — เฉพาะ บาท/ตัน (กรอกง่าย โปรแกรมคำนวณเอง)
    ac_ton = st.session_state.get('ac_ton_prices', DEFAULT_AC_TON_PRICES)
    ac_rows = [
        {'Material': mat, 'Price (Baht/ton)': float(ac_ton.get(mat, 0))}
        for mat in DEFAULT_AC_TON_PRICES.keys()
    ]

    # Concrete_Prices — บาท/ลบ.ม. (กรอกง่าย โปรแกรมคำนวณเอง)
    _conc_cum_tpl = st.session_state.get('concrete_cum_prices', DEFAULT_CONCRETE_CUM_PRICES)
    conc_rows = [
        {'ประเภท': ct, 'บาท/ลบ.ม.': float(_conc_cum_tpl.get(ct, DEFAULT_CONCRETE_CUM_PRICES.get(ct, 0)))}
        for ct in ['JPCP', 'JRCP', 'CRCP']
    ]

    # Base_Materials — แยก บาท/ลบ.ม. และ บาท/ตร.ม.
    SQMKEYS = {'Prime Coat', 'Non Woven Geotextile', 'Wire Mesh', 'Tack Coat'}
    base_rows = []
    for k, v in lib['base_prices'].items():
        unit = 'Baht/sq.m' if k in SQMKEYS else 'Baht/cu.m'
        base_rows.append({'Material': k, 'Price': v, 'Unit': unit})

    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        pd.DataFrame(ac_rows).to_excel(writer, sheet_name='AC_Prices', index=False)
        pd.DataFrame(conc_rows).to_excel(writer, sheet_name='Concrete_Prices', index=False)
        pd.DataFrame(base_rows).to_excel(writer, sheet_name='Base_Materials', index=False)
    output.seek(0)
    return output.getvalue()


def _calc_ac_prices_from_ton(ton_prices: dict, density: float = DEFAULT_AC_DENSITY) -> dict:
    """คำนวณ AC price table (บาท/ตร.ม.) จากราคา/ตัน × density × ความหนา"""
    thicknesses = [2.5, 3, 4, 5, 6, 7, 8, 9, 10]
    result = {}
    for mat, ton_p in ton_prices.items():
        if ton_p > 0:
            result[mat] = {t: round(ton_p * density * t / 100, 2) for t in thicknesses}
        else:
            # ถ้าราคา = 0 → ใช้ default
            result[mat] = dict(DEFAULT_AC_PRICES.get(mat, {}))
    return result


def load_excel_price_library(uploaded_file) -> dict:
    """อ่าน Excel → dict price_library
    AC_Prices: อ่าน Price (Baht/ton) → คำนวณ price table อัตโนมัติ
    Concrete_Prices: อ่าน บาท/ตร.ม. ตามความหนา
    Base_Materials: อ่านราคา/หน่วย
    """
    ac_df   = pd.read_excel(uploaded_file, sheet_name='AC_Prices')
    conc_df = pd.read_excel(uploaded_file, sheet_name='Concrete_Prices')
    base_df = pd.read_excel(uploaded_file, sheet_name='Base_Materials')

    # ── AC: อ่าน บาท/ตัน → คำนวณ price table ──────────────────
    ac_ton_prices: dict = dict(DEFAULT_AC_TON_PRICES)  # fallback
    if 'Price (Baht/ton)' in ac_df.columns:
        for _, row in ac_df.iterrows():
            try:
                mat = str(row['Material'])
                val = row['Price (Baht/ton)']
                if pd.notna(mat) and pd.notna(val) and float(val) > 0:
                    ac_ton_prices[mat] = float(val)
            except (ValueError, TypeError):
                pass
        # บันทึก ton prices ไว้ใน session_state ด้วย (สำหรับ UI Tab 2)
        st.session_state['ac_ton_prices'] = ac_ton_prices
    density = st.session_state.get('tab2_density', DEFAULT_AC_DENSITY)
    ac_prices = _calc_ac_prices_from_ton(ac_ton_prices, density)
    # เติม key ที่ขาด
    for mat, dp in DEFAULT_AC_PRICES.items():
        if mat not in ac_prices:
            ac_prices[mat] = dict(dp)

    # ── Concrete: อ่าน บาท/ลบ.ม. → คำนวณ price table ──────────
    conc_cum: dict = dict(DEFAULT_CONCRETE_CUM_PRICES)
    cum_col = 'บาท/ลบ.ม.' if 'บาท/ลบ.ม.' in conc_df.columns else None
    price_col_conc = 'Price (Baht/cu.m)' if 'Price (Baht/cu.m)' in conc_df.columns else None
    type_col = 'Type' if 'Type' in conc_df.columns else 'ประเภท'

    for _, row in conc_df.iterrows():
        try:
            ct = str(row[type_col])
            col_to_use = cum_col or price_col_conc
            if col_to_use and pd.notna(row.get(col_to_use, None)):
                val = float(row[col_to_use])
                if val > 0:
                    conc_cum[ct] = val
        except (ValueError, TypeError, KeyError):
            pass
    st.session_state['concrete_cum_prices'] = conc_cum
    conc_prices = _calc_concrete_prices(conc_cum)

    # ── Base Materials: อ่านราคา/หน่วย ────────────────────────
    base_prices: dict = dict(DEFAULT_BASE_PRICES)
    price_col = 'Price' if 'Price' in base_df.columns else 'Price (Baht/cu.m)'
    for _, row in base_df.iterrows():
        try:
            mat = str(row['Material'])
            val = row[price_col]
            if pd.notna(mat) and pd.notna(val):
                base_prices[mat] = float(val)
        except (ValueError, TypeError):
            pass

    return {'ac_prices': ac_prices, 'concrete_prices': conc_prices, 'base_prices': base_prices}


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN APP
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    # ── Header ────────────────────────────────────────────────────────────────
    st.markdown("""
    <div class="main-header">
        <h1>🛣️ ระบบวิเคราะห์ค่าก่อสร้างโครงสร้างชั้นทาง</h1>
        <p>ตามแนวทาง AASHTO 1993 &nbsp;|&nbsp; รองรับ AC · JPCP · JRCP · CRCP</p>
    </div>
    """, unsafe_allow_html=True)

    # ══════════════════════════════════════════════════════════════
    # SIDEBAR
    # ══════════════════════════════════════════════════════════════
    with st.sidebar:
        st.markdown("## 📋 ข้อมูลโครงการ")

        # ── Price Library (Excel) ──
        with st.expander("💰 Price Library (Excel)", expanded=False):
            uploaded_excel = st.file_uploader(
                "Upload Excel", type=['xlsx', 'xls'],
                key="sidebar_price_excel",
                help="ดาวน์โหลด Template จาก Tab 2 ก่อน แล้วแก้ราคา แล้ว Upload กลับ"
            )
            if uploaded_excel is not None:
                try:
                    lib = load_excel_price_library(uploaded_excel)
                    st.session_state['price_library'] = lib
                    st.success("✅ โหลด Price Library สำเร็จ")
                    st.caption(f"AC: {len(lib['ac_prices'])} ประเภท | Concrete: {len(lib['concrete_prices'])} ประเภท")
                except Exception as e:
                    st.error(f"❌ อ่านไฟล์ไม่สำเร็จ: {e}")

        st.divider()

        # ── Load Project (JSON) ──
        with st.expander("📂 โหลดโครงการ (JSON)", expanded=False):
            uploaded_json = st.file_uploader(
                "Upload JSON", type=['json'],
                key="sidebar_upload_json",
            )
            if uploaded_json is not None:
                try:
                    import hashlib
                    fb = uploaded_json.read()
                    fh = hashlib.md5(fb).hexdigest()
                    loaded = json.loads(fb.decode('utf-8'))
                    if 'project_info' in loaded:
                        st.info(f"📌 {loaded['project_info'].get('name', '-')}")
                        st.caption(f"บันทึกเมื่อ: {loaded.get('saved_at', '-')}")
                    if st.button("📥 นำเข้าข้อมูล", key="import_json_btn"):
                        if st.session_state.get('loaded_json_hash') != fh:
                            st.session_state['loaded_project'] = loaded
                            st.session_state['loaded_json_hash'] = fh
                            new_v = st.session_state.get('json_version', 0) + 1
                            st.session_state['json_version'] = new_v
                            # ล้าง data_editor state ทั้งหมด
                            keys_to_clear = [k for k in st.session_state
                                             if any(p in k for p in ['_surf_init_', '_base_init_', '_joint_init_',
                                                                       '_surf_editor_', '_base_editor_', '_joint_editor_'])]
                            for k in keys_to_clear:
                                del st.session_state[k]
                        st.rerun()
                except Exception as e:
                    st.error(f"❌ ไม่สามารถอ่านไฟล์: {e}")

        st.divider()

        # ── Project Info ──
        lp   = st.session_state.get('loaded_project', {})
        li   = lp.get('project_info', {})
        v_sb = st.session_state.get('json_version', 0)

        project_name = st.text_input(
            "ชื่อโครงการ",
            value=li.get('name', 'โครงการก่อสร้างทางหลวง'),
            key=f"sb_pname_v{v_sb}"
        )
        road_length = st.number_input(
            "ความยาวถนน (กม.)",
            value=float(li.get('length', 1.0)),
            min_value=0.1, step=0.1,
            key=f"sb_length_v{v_sb}"
        )
        design_life = st.number_input(
            "อายุออกแบบ (ปี)",
            value=int(li.get('design_life', 20)),
            min_value=1, max_value=50, step=1,
            key=f"sb_dlife_v{v_sb}"
        )

        st.divider()
        st.markdown("**📐 ขนาดถนน**")

        lane_width = st.number_input(
            "ความกว้างช่องจราจร (ม.)",
            value=float(li.get('lane_width', 3.5)),
            min_value=2.5, max_value=4.5, step=0.25,
            key=f"sb_lw_v{v_sb}"
        )
        lpo = [2, 3, 4]
        lp_def = li.get('num_lanes', 4) // 2
        lp_idx = lpo.index(lp_def) if lp_def in lpo else 0
        lanes_per_dir = st.selectbox(
            "ช่องจราจร/ทิศทาง",
            options=lpo, index=lp_idx,
            key=f"sb_lpd_v{v_sb}"
        )
        num_lanes = lanes_per_dir * 2

        shoulder_l = st.number_input(
            "ไหล่ทางซ้าย (ม.)",
            value=float(li.get('shoulder_left', 2.5)),
            min_value=0.0, max_value=4.0, step=0.25,
            key=f"sb_sl_v{v_sb}"
        )
        shoulder_r = st.number_input(
            "ไหล่ทางขวา (ม.)",
            value=float(li.get('shoulder_right', 1.5)),
            min_value=0.0, max_value=4.0, step=0.25,
            key=f"sb_sr_v{v_sb}"
        )

        road_surface_width = lane_width * num_lanes
        total_shoulders    = (shoulder_l + shoulder_r) * 2
        total_width        = road_surface_width + total_shoulders
        area_per_km        = total_width * 1000

        st.info(
            f"📏 ช่องรวม: **{num_lanes}** ช่อง\n"
            f"📏 ผิวจราจร: **{road_surface_width:.2f}** ม.\n"
            f"📏 ไหล่ทาง: **{total_shoulders:.2f}** ม.\n"
            f"📏 กว้างรวม: **{total_width:.2f}** ม."
        )

    project_info = {
        'name':         project_name,
        'length':       road_length,
        'design_life':  design_life,
        'lane_width':   lane_width,
        'num_lanes':    num_lanes,
        'shoulder_left':  shoulder_l,
        'shoulder_right': shoulder_r,
        'total_width':  total_width,
    }
    st.session_state['project_info'] = project_info

    v = st.session_state.get('json_version', 0)

    # ══════════════════════════════════════════════════════════════
    # TABS
    # ══════════════════════════════════════════════════════════════
    tab1, tab2, tab3 = st.tabs([
        "🏗️ กำหนดหน้าตัด",
        "💰 ราคาวัสดุ",
        "📊 สรุปต้นทุน & รายงาน",
    ])

    all_results: dict = {}   # เก็บผลคำนวณจาก Tab 1 ส่งต่อ Tab 3

    # ══════════════════════════════════════════════════════════════
    # TAB 1 — Section Setup
    # ══════════════════════════════════════════════════════════════
    with tab1:
        st.markdown(f"**โครงการ:** {project_name} &nbsp;|&nbsp; **ความยาว:** {road_length:.2f} กม. &nbsp;|&nbsp; **กว้างรวม:** {total_width:.2f} ม.")

        sub_ac, sub_jpcp, sub_jrcp, sub_crcp = st.tabs(["🛣️ AC", "🏗️ JPCP", "🏗️ JRCP", "🏗️ CRCP"])

        for sub_tab, ptype in [(sub_ac, 'AC'), (sub_jpcp, 'JPCP'), (sub_jrcp, 'JRCP'), (sub_crcp, 'CRCP')]:
            with sub_tab:
                kp = ptype.lower()
                layers = render_layer_editor(ptype, kp, total_width, road_length, v=v)

                joints = []
                include_joints = True
                if ptype in ('JPCP', 'JRCP', 'CRCP'):
                    joints, include_joints = render_joint_editor(ptype, kp, area_per_km, road_length, v=v)

                # คำนวณ
                layer_cost, layer_details = calculate_layer_cost(layers, road_length)
                joint_cost, joint_details = calculate_joint_cost(joints, road_length, include_joints)
                total_cost = layer_cost + joint_cost

                total_area = area_per_km * road_length
                cost_sqm   = total_cost / total_area if total_area > 0 else 0
                cost_per_km_m = total_cost / road_length / 1e6 if road_length > 0 else 0

                # Metric
                st.markdown("---")
                mc1, mc2, mc3 = st.columns(3)
                with mc1:
                    st.markdown(f"""<div class="metric-card">
                        <div class="label">💰 ราคา/ตร.ม.</div>
                        <div class="value">{cost_sqm:,.2f}<span class="unit">บาท</span></div>
                    </div>""", unsafe_allow_html=True)
                with mc2:
                    st.markdown(f"""<div class="metric-card">
                        <div class="label">📏 ราคา/กม.</div>
                        <div class="value">{cost_per_km_m:,.3f}<span class="unit">ล้านบาท</span></div>
                    </div>""", unsafe_allow_html=True)
                with mc3:
                    st.markdown(f"""<div class="metric-card">
                        <div class="label">⏱️ อายุออกแบบ</div>
                        <div class="value">{design_life}<span class="unit">ปี</span></div>
                    </div>""", unsafe_allow_html=True)

                # เก็บผลสำหรับ Tab 3
                all_results[ptype] = {
                    'name':        ptype,
                    'name_detail': f"{ptype} — {project_name}",
                    'layers':      layers,
                    'joints':      joints,
                    'details':     layer_details + joint_details,
                    'cost_total':  total_cost,
                    'cost_sqm':    cost_sqm,
                    'cost_per_km': cost_per_km_m,
                    'include_joints': include_joints,
                }

        # บันทึก all_results ไว้ใน session state สำหรับ Tab 3
        st.session_state['all_results'] = all_results

    # ══════════════════════════════════════════════════════════════
    # TAB 2 — Price Library
    # ══════════════════════════════════════════════════════════════
    with tab2:
        st.header("💰 ตารางราคาวัสดุ")
        st.info("💡 แก้ราคาได้โดยตรงในตาราง หรือคำนวณจากราคาต่อตัน หรือ Upload Excel ใน Sidebar")

        lib = get_price_library()

        col_dl, col_up = st.columns([1, 2])
        with col_dl:
            tpl = generate_excel_template()
            st.download_button(
                "⬇️ Download Template Excel",
                data=tpl,
                file_name="price_library_template.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

        # ══════════════════════════════════════════════════════════
        # คำนวณราคา AC จากราคาต่อตัน (UI แนวตั้ง + preview)
        # ══════════════════════════════════════════════════════════
        with st.expander("🧮 คำนวณราคา AC จากราคาต่อตัน", expanded=True):
            st.caption("ราคา (บาท/ตร.ม.) = ราคา (บาท/ตัน) × density × ความหนา (m)")
            _thicknesses = [2.5, 3, 4, 5, 6, 7, 8, 9, 10]

            _dc1, _dc2 = st.columns([1, 3])
            with _dc1:
                _density_key = "tab2_density"
                if _density_key not in st.session_state:
                    st.session_state[_density_key] = DEFAULT_AC_DENSITY
                st.number_input(
                    "Density (ตัน/ลบ.ม.)",
                    min_value=2.0, max_value=2.6, step=0.05,
                    format="%.2f", key=_density_key
                )
            _density = float(st.session_state[_density_key])

            # ── ตารางแนวตั้ง: วัสดุ | บาท/ตัน | preview บาท/ตร.ม. ──
            _ac_order = ['PMA Wearing Course', 'AC Wearing Course',
                         'AC Binder Course', 'AC Base Course']
            _ton_defaults = dict(DEFAULT_AC_TON_PRICES)
            _saved_ton = st.session_state.get('ac_ton_prices', _ton_defaults)

            # preview thickness
            _pv1, _pv2 = st.columns([3, 1])
            with _pv2:
                _prev_thick_key = "tab2_preview_thick"
                if _prev_thick_key not in st.session_state:
                    st.session_state[_prev_thick_key] = 5.0
                st.number_input("ดูตัวอย่างที่หนา (cm)",
                    min_value=2.5, max_value=10.0, step=0.5, format="%.1f",
                    key=_prev_thick_key)
            _prev_thick = float(st.session_state[_prev_thick_key])
            with _pv1:
                st.markdown(f"<span style='color:#6b7a8d;font-size:0.85rem'>ตัวอย่างราคาที่ความหนา <b>{_prev_thick:.1f} cm</b></span>",
                    unsafe_allow_html=True)

            # header
            _gh = st.columns([3, 1.5, 1.5])
            _gh[0].markdown("<span style='color:#6b7a8d;font-size:0.82rem;font-weight:600'>วัสดุ</span>", unsafe_allow_html=True)
            _gh[1].markdown("<span style='color:#6b7a8d;font-size:0.82rem;font-weight:600'>บาท/ตัน</span>", unsafe_allow_html=True)
            _gh[2].markdown(f"<span style='color:#6b7a8d;font-size:0.82rem;font-weight:600'>บาท/ตร.ม. ({_prev_thick:.1f}cm)</span>", unsafe_allow_html=True)

            _ton_prices = {}
            for mat in _ac_order:
                _tkey = f"tab2_ton_{mat.replace(' ','_')}"
                if _tkey not in st.session_state:
                    st.session_state[_tkey] = float(_saved_ton.get(mat, _ton_defaults.get(mat, 0)))
                _gr = st.columns([3, 1.5, 1.5])
                with _gr[0]:
                    st.markdown(f"<div style='padding:8px 0;font-weight:500'>{mat}</div>", unsafe_allow_html=True)
                with _gr[1]:
                    st.number_input(f"ton_{mat}", min_value=0.0, step=50.0,
                        format="%.0f", key=_tkey, label_visibility="collapsed")
                ton_p = float(st.session_state[_tkey])
                _ton_prices[mat] = ton_p
                preview_price = round(ton_p * _density * _prev_thick / 100, 2) if ton_p > 0 else 0.0
                with _gr[2]:
                    color = '#0f2942' if preview_price > 0 else '#94a3b8'
                    st.markdown(
                        f"<div style='padding:8px 0;font-weight:600;color:{color}'>"
                        f"{'—' if preview_price == 0 else f'{preview_price:,.2f}'}</div>",
                        unsafe_allow_html=True
                    )

            if st.button("🔄 คำนวณและอัพเดท Price Table AC", type="primary", key="tab2_calc_ton"):
                _new_ac_prices = _calc_ac_prices_from_ton(_ton_prices, _density)
                _cur_lib = get_price_library()
                _cur_lib['ac_prices'] = _new_ac_prices
                st.session_state['price_library'] = _cur_lib
                st.session_state['ac_ton_prices'] = dict(_ton_prices)
                # reset sprice เพื่อให้ Tab 1 ดึงราคาใหม่
                for k in list(st.session_state.keys()):
                    if '_sprice_' in k or '_p_AC' in k or '_p_PMA' in k:
                        del st.session_state[k]
                st.success(f"✅ อัพเดทราคา AC สำเร็จ (density={_density:.2f})")
                st.rerun()

        st.subheader("🛣️ ราคา AC — บาท/ตัน และ บาท/ตร.ม. ทุกความหนา")
        st.caption("แก้ **บาท/ตัน** ได้โดยตรง — ราคา/ตร.ม. คำนวณอัตโนมัติ (read-only)")

        _saved_ton2 = st.session_state.get('ac_ton_prices', DEFAULT_AC_TON_PRICES)
        _den2       = float(st.session_state.get('tab2_density', DEFAULT_AC_DENSITY))
        _ac_order2  = ['PMA Wearing Course', 'AC Wearing Course', 'AC Binder Course', 'AC Base Course']
        _thk_list   = [2.5, 3, 4, 5, 6, 7, 8, 9, 10]

        # สร้าง DataFrame: วัสดุ | บาท/ตัน | 2.5cm ... 10cm
        ac_ton_rows = []
        for mat in _ac_order2:
            ton_p = float(_saved_ton2.get(mat, DEFAULT_AC_TON_PRICES.get(mat, 0)))
            row = {'วัสดุ': mat, 'บาท/ตัน': ton_p}
            for t in _thk_list:
                row[f"{t}cm"] = round(ton_p * _den2 * t / 100, 0) if ton_p > 0 else 0.0
            ac_ton_rows.append(row)

        # column_config: บาท/ตัน แก้ได้, ความหนาทั้งหมด read-only
        _ac_col_cfg = {
            'วัสดุ':    st.column_config.TextColumn('วัสดุ', width='medium', disabled=True),
            'บาท/ตัน': st.column_config.NumberColumn('บาท/ตัน', min_value=0.0, step=50.0, format='%.0f', width='small'),
        }
        for t in _thk_list:
            _ac_col_cfg[f"{t}cm"] = st.column_config.NumberColumn(
                f"{t}cm", format='%.0f', disabled=True, width='small'
            )

        ac_edited = st.data_editor(
            pd.DataFrame(ac_ton_rows),
            column_config=_ac_col_cfg,
            use_container_width=True,
            hide_index=True,
            key="tab2_ac_editor",
        )

        st.subheader("🏗️ ราคาคอนกรีต — บาท/ลบ.ม. และ บาท/ตร.ม. ทุกความหนา")
        st.caption("แก้ **บาท/ลบ.ม.** ได้โดยตรง — ราคา/ตร.ม. คำนวณอัตโนมัติ (read-only)")

        _conc_thk   = [20, 25, 28, 30, 32, 35]
        _conc_order = ['JPCP', 'JRCP', 'CRCP']

        # init concrete_cum_prices ด้วย default ถ้ายังไม่มี
        if 'concrete_cum_prices' not in st.session_state:
            st.session_state['concrete_cum_prices'] = dict(DEFAULT_CONCRETE_CUM_PRICES)
        _conc_cum = st.session_state['concrete_cum_prices']

        # สร้าง DataFrame แบบ dynamic จาก session_state (ไม่ lock)
        # ใช้ number_input แยกต่อ row แทน data_editor เพื่อหลีกเลี่ยง lock
        _cp_col_cfg = {
            'ประเภท':    st.column_config.TextColumn('ประเภท', width='small', disabled=True),
            'บาท/ลบ.ม.': st.column_config.NumberColumn('บาท/ลบ.ม.', min_value=0.0, step=50.0, format='%.0f', width='small'),
        }
        for t in _conc_thk:
            _cp_col_cfg[f"{t}cm"] = st.column_config.NumberColumn(
                f"{t}cm", format='%.0f', disabled=True, width='small'
            )

        # Header
        _cph = st.columns([1.2, 1.2] + [1]*len(_conc_thk))
        _cph[0].markdown("<span style='color:#6b7a8d;font-size:0.82rem;font-weight:600'>ประเภท</span>", unsafe_allow_html=True)
        _cph[1].markdown("<span style='color:#6b7a8d;font-size:0.82rem;font-weight:600'>บาท/ลบ.ม.</span>", unsafe_allow_html=True)
        for j, t in enumerate(_conc_thk):
            _cph[2+j].markdown(f"<span style='color:#6b7a8d;font-size:0.82rem;font-weight:600'>{t}cm</span>", unsafe_allow_html=True)

        # detect การเปลี่ยน concrete_cum → reset slab price keys
        _prev_conc_ver_k = "tab2_conc_ver"
        _cur_conc_ver    = str({ct: st.session_state.get(f'tab2_conc_cum_{ct}', 
                           DEFAULT_CONCRETE_CUM_PRICES.get(ct,0)) for ct in _conc_order})
        if st.session_state.get(_prev_conc_ver_k) != _cur_conc_ver:
            for k in list(st.session_state.keys()):
                if '_p_Concrete' in k or '_pver_Concrete' in k                         or '_p_350' in k or '_pver_350' in k:
                    del st.session_state[k]
            st.session_state[_prev_conc_ver_k] = _cur_conc_ver

        cp_edited_rows = []
        for ct in _conc_order:
            _cum_key = f"tab2_conc_cum_{ct}"
            if _cum_key not in st.session_state:
                st.session_state[_cum_key] = float(_conc_cum.get(ct, DEFAULT_CONCRETE_CUM_PRICES.get(ct, 0)))
            _cpr = st.columns([1.2, 1.2] + [1]*len(_conc_thk))
            with _cpr[0]:
                st.markdown(f"<div style='padding:8px 0;font-weight:600'>{ct}</div>", unsafe_allow_html=True)
            with _cpr[1]:
                st.number_input(f"cum_{ct}", min_value=0.0, step=50.0, format="%.0f",
                    key=_cum_key, label_visibility="collapsed")
            cum_val = float(st.session_state[_cum_key])
            for j, t in enumerate(_conc_thk):
                sqm_val = round(cum_val * t / 100, 0)
                _cpr[2+j].markdown(
                    f"<div style='padding:8px 4px;font-size:0.9rem;color:#0f2942'>{sqm_val:,.0f}</div>",
                    unsafe_allow_html=True
                )
            cp_edited_rows.append({'ประเภท': ct, 'บาท/ลบ.ม.': cum_val})

        cp_edited = pd.DataFrame(cp_edited_rows)

        # แยก base_prices เป็น 2 กลุ่ม
        SQMKEYS = {'Prime Coat', 'Non Woven Geotextile', 'Wire Mesh', 'Tack Coat'}
        bp_cum_rows = [{'วัสดุ': k, 'ราคา (บาท/ลบ.ม.)': v}
                       for k, v in lib['base_prices'].items() if k not in SQMKEYS]
        bp_sqm_rows = [{'วัสดุ': k, 'ราคา (บาท/ตร.ม.)': v}
                       for k, v in lib['base_prices'].items() if k in SQMKEYS]

        st.subheader("🪨 ราคาวัสดุพื้นทาง / รองพื้นทาง (บาท/ลบ.ม.)")
        bp_cum_edited = st.data_editor(
            pd.DataFrame(bp_cum_rows),
            column_config={
                'วัสดุ':              st.column_config.TextColumn('วัสดุ', width='large', disabled=True),
                'ราคา (บาท/ลบ.ม.)': st.column_config.NumberColumn('ราคา (บาท/ลบ.ม.)', min_value=0.0, step=10.0, format='%.2f'),
            },
            use_container_width=True,
            hide_index=True,
            key="tab2_bp_cum_editor",
        )

        st.subheader("🧴 ราคาวัสดุผิว / อุปกรณ์ (บาท/ตร.ม.)")
        st.caption("Prime Coat, Tack Coat, Non Woven Geotextile, Wire Mesh — คิดตามพื้นที่ ไม่ใช่ปริมาตร")
        bp_sqm_edited = st.data_editor(
            pd.DataFrame(bp_sqm_rows),
            column_config={
                'วัสดุ':              st.column_config.TextColumn('วัสดุ', width='large', disabled=True),
                'ราคา (บาท/ตร.ม.)': st.column_config.NumberColumn('ราคา (บาท/ตร.ม.)', min_value=0.0, step=1.0, format='%.2f'),
            },
            use_container_width=True,
            hide_index=True,
            key="tab2_bp_sqm_editor",
        )

        if st.button("💾 บันทึกราคาที่แก้ไขลง Library", type="primary"):
            # AC: อ่าน บาท/ตัน → คำนวณ price table
            _new_ton: dict = {}
            for _, row in ac_edited.iterrows():
                mat = str(row['วัสดุ'])
                val = row.get('บาท/ตัน', row.get('ราคา (บาท/ตัน)', 0))
                if pd.notna(val):
                    _new_ton[mat] = float(val)
            _den = float(st.session_state.get('tab2_density', DEFAULT_AC_DENSITY))
            new_ac = _calc_ac_prices_from_ton(_new_ton, _den)
            st.session_state['ac_ton_prices'] = _new_ton

            # Concrete: อ่าน บาท/ลบ.ม. → คำนวณ price table
            _new_conc_cum: dict = {}
            for _, row in cp_edited.iterrows():
                ct = str(row['ประเภท'])
                val = row.get('บาท/ลบ.ม.', 0)
                if pd.notna(val) and float(val) > 0:
                    _new_conc_cum[ct] = float(val)
            st.session_state['concrete_cum_prices'] = _new_conc_cum
            new_cp = _calc_concrete_prices(_new_conc_cum)
            # ล้าง sprice/pver keys เพื่อให้ Tab 1 ดึงราคาใหม่
            for k in list(st.session_state.keys()):
                if '_sprice_' in k or '_p_Concrete' in k or '_pver_Concrete' in k                         or '_p_350' in k or '_pver_350' in k                         or '_slabp_' in k or 'sthick_slab' in k:
                    del st.session_state[k]

            # Base
            new_bp: dict = {}
            for _, row in bp_cum_edited.iterrows():
                new_bp[str(row['วัสดุ'])] = float(row['ราคา (บาท/ลบ.ม.)'])
            for _, row in bp_sqm_edited.iterrows():
                new_bp[str(row['วัสดุ'])] = float(row['ราคา (บาท/ตร.ม.)'])

            st.session_state['price_library'] = {
                'ac_prices': new_ac,
                'concrete_prices': new_cp,
                'base_prices': new_bp,
            }
            # reset sprice เพื่อให้ Tab 1 ดึงราคาใหม่
            for k in list(st.session_state.keys()):
                if '_sprice_' in k or '_p_AC' in k or '_p_PMA' in k:
                    del st.session_state[k]
            st.success("✅ อัพเดท Price Library สำเร็จ — กลับไป Tab 1 เพื่อดูผล")

    # ══════════════════════════════════════════════════════════════
    # TAB 3 — Cost Summary + Report
    # ══════════════════════════════════════════════════════════════
    with tab3:
        st.header("📊 สรุปต้นทุนและรายงาน")

        ar = st.session_state.get('all_results', all_results)

        if not ar:
            st.info("กรุณากำหนดหน้าตัดใน Tab 1 ก่อน")
        else:
            # ── Summary metrics ──
            ptypes_list = list(ar.keys())
            cols_m = st.columns(len(ptypes_list))
            for i, pt in enumerate(ptypes_list):
                r = ar[pt]
                with cols_m[i]:
                    st.markdown(f"""<div class="metric-card">
                        <div class="label">{pt}</div>
                        <div class="value">{r['cost_sqm']:,.0f}<span class="unit">บาท/ตร.ม.</span></div>
                        <div class="label" style="margin-top:6px">{r['cost_per_km']:.3f} ล้านบาท/กม.</div>
                    </div>""", unsafe_allow_html=True)

            st.markdown("---")

            # ── Plotly bar chart ──
            if PLOTLY_AVAILABLE:
                fig = go.Figure()
                fig.add_trace(go.Bar(
                    x=[ar[pt]['name'] for pt in ptypes_list],
                    y=[ar[pt]['cost_sqm'] for pt in ptypes_list],
                    marker_color=['#1a4a7a', '#0d7377', '#14a085', '#52b788'],
                    text=[f"{ar[pt]['cost_sqm']:,.0f}" for pt in ptypes_list],
                    textposition='outside',
                    name='ราคา (บาท/ตร.ม.)',
                ))
                fig.update_layout(
                    title='เปรียบเทียบราคา/ตร.ม. ตามประเภทโครงสร้าง',
                    yaxis_title='บาท/ตร.ม.',
                    plot_bgcolor='white',
                    paper_bgcolor='white',
                    font=dict(family='IBM Plex Sans Thai, sans-serif'),
                    height=380,
                )
                st.plotly_chart(fig, use_container_width=True)

            # ── ตารางสรุปรายละเอียด ──
            st.subheader("📋 รายละเอียดต้นทุนแต่ละประเภท")
            for pt in ptypes_list:
                r = ar[pt]
                with st.expander(f"🔍 {pt} — {r['cost_sqm']:,.2f} บาท/ตร.ม.", expanded=False):
                    if r['details']:
                        st.dataframe(
                            pd.DataFrame(r['details'])[['รายการ', 'ปริมาณ', 'หน่วย', 'ราคา/หน่วย (แสดง)', 'หน่วยราคา', 'มูลค่า (บาท)']],
                            use_container_width=True, hide_index=True
                        )

            st.markdown("---")
            st.subheader("💾 บันทึกและส่งออก")

            col_s1, col_s2, col_s3 = st.columns(3)

            # ── Save JSON ──
            with col_s1:
                construction_out: dict = {}
                for pt, r in ar.items():
                    construction_out[pt] = {
                        'layers': r.get('layers', []),
                        'joints': r.get('joints', []),
                        'include_joints': r.get('include_joints', True),
                    }
                save_data = {
                    'saved_at':    datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                    'project_info': project_info,
                    'construction': construction_out,
                }
                json_bytes = json.dumps(save_data, ensure_ascii=False, indent=2, default=str).encode('utf-8')
                fname_json = f"{project_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.json"
                st.download_button(
                    "📥 บันทึก JSON",
                    data=json_bytes,
                    file_name=fname_json,
                    mime="application/json",
                    use_container_width=True,
                )

            # ── Word Materials ──
            with col_s2:
                if DOCX_AVAILABLE:
                    if st.button("📄 Word แบบวัสดุ+ราคา", use_container_width=True):
                        try:
                            doc = generate_word_report(project_info, ar, report_type='materials')
                            buf = io.BytesIO()
                            doc.save(buf)
                            buf.seek(0)
                            fname_w = f"{project_name.replace(' ', '_')}_Materials_{datetime.now().strftime('%Y%m%d')}.docx"
                            st.download_button(
                                "⬇️ ดาวน์โหลด Word (วัสดุ)",
                                data=buf.getvalue(),
                                file_name=fname_w,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key="dl_word_mat",
                            )
                        except Exception as e:
                            st.error(f"สร้างรายงานไม่สำเร็จ: {e}")
                else:
                    st.warning("python-docx ไม่พร้อมใช้งาน")

            # ── Word Consultant ──
            with col_s3:
                if DOCX_AVAILABLE:
                    with st.expander("⚙️ ตั้งค่ารายงานที่ปรึกษา"):
                        sec_start = st.text_input("หมายเลขหัวข้อ", value="4.7", key="cons_sec")

                        # สร้างบทเกริ่นนำอัตโนมัติจากข้อมูลโครงการ
                        _pi = project_info
                        _nl = _pi.get('num_lanes', 4)
                        _tw = _pi.get('total_width', 0)
                        _ln = _pi.get('length', 1.0)
                        _auto_intro = (
                            f"รายงานวิเคราะห์ต้นทุนค่าก่อสร้างโครงสร้างชั้นทางฉบับนี้ "
                            f"จัดทำขึ้นเพื่อเปรียบเทียบทางเลือกโครงสร้างชั้นทางประเภทต่าง ๆ "
                            f"สำหรับถนน {_nl} ช่องจราจร "
                            f"ความกว้างรวม {_tw:.2f} เมตร "
                            f"ระยะทาง {_ln:.2f} กิโลเมตร "
                            f"โดยครอบคลุมทั้งผิวทางแอสฟัลต์คอนกรีต (AC) "
                            f"และผิวทางคอนกรีตซีเมนต์ (JPCP, JRCP, CRCP) "
                            f"การวิเคราะห์อ้างอิงราคาวัสดุและค่าก่อสร้างตามมาตรฐานกรมบัญชีกลาง "
                            f"เพื่อใช้เป็นข้อมูลประกอบการตัดสินใจเลือกโครงสร้างชั้นทาง"
                            f"ที่เหมาะสมกับสภาพโครงการ"
                        )
                        # key ผูกกับข้อมูลโครงการ → เปลี่ยนอัตโนมัติเมื่อ project_info เปลี่ยน
                        _intro_key = f"cons_intro_{_nl}_{_tw:.1f}_{_ln:.2f}"
                        if _intro_key not in st.session_state:
                            st.session_state[_intro_key] = _auto_intro
                        intro_txt = st.text_area(
                            "บทเกริ่นนำ (แก้ไขได้)",
                            height=120, key=_intro_key
                        )
                    if st.button("📑 Word แบบที่ปรึกษา", use_container_width=True):
                        try:
                            doc = generate_word_report(
                                project_info, ar,
                                report_type='consultant',
                                section_start=sec_start,
                                intro_text=intro_txt,
                            )
                            buf = io.BytesIO()
                            doc.save(buf)
                            buf.seek(0)
                            fname_c = f"{project_name.replace(' ', '_')}_Consultant_{datetime.now().strftime('%Y%m%d')}.docx"
                            st.download_button(
                                "⬇️ ดาวน์โหลด Word (ที่ปรึกษา)",
                                data=buf.getvalue(),
                                file_name=fname_c,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key="dl_word_cons",
                            )
                        except Exception as e:
                            st.error(f"สร้างรายงานไม่สำเร็จ: {e}")
                else:
                    st.warning("python-docx ไม่พร้อมใช้งาน")

    # ── Footer ──
    st.markdown("""
    <div class="footer">
        <b>รศ.ดร.อิทธิพล มีผล</b><br>
        ภาควิชาครุศาสตร์โยธา คณะครุศาสตร์อุตสาหกรรม มจพ.<br>
        <span style="color:#b0bec5">Pavement Structure Cost Analysis v6.0</span>
    </div>
    """, unsafe_allow_html=True)


if __name__ == "__main__":
    main()
