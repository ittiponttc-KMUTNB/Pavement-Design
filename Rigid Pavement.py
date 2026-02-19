"""
โปรแกรมออกแบบและตรวจสอบความหนาถนนคอนกรีต (Rigid Pavement)
ตามวิธี AASHTO 1993
รองรับทั้ง JPCP (Jointed Plain Concrete Pavement) และ CRCP (Continuously Reinforced Concrete Pavement)

รวมโปรแกรม:
1. การหาค่า k-value และปรับแก้ Loss of Support (LS) จาก Nomograph
2. การคำนวณความหนาถนนคอนกรีตตาม AASHTO 1993
การปรับปรุง
1. ย้ายรูปแบบการตรวจสอบความหนาผิวทางไปcolumn 2 ฝั่งขวา
พัฒนาสำหรับใช้ในการเรียนการสอน
ภาควิชาครุศาสตร์โยธา มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ
"""

import streamlit as st
import math
from io import BytesIO
from datetime import datetime
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from PIL import Image, ImageDraw
import io
import json
import pandas as pd

# ============================================================
# ค่าคงที่และตารางอ้างอิง AASHTO 1993
# ============================================================

ZR_TABLE = {
    50: -0.000, 60: -0.253, 70: -0.524, 75: -0.674, 80: -0.841, 85: -1.037,
    90: -1.282, 91: -1.340, 92: -1.405, 93: -1.476, 94: -1.555, 95: -1.645,
    96: -1.751, 97: -1.881, 98: -2.054, 99: -2.327
}

J_VALUES = {"JRCP": 2.8, "JPCP": 2.8, "JRCP/JPCP": 2.8, "CRCP": 2.5}
CD_DEFAULT = 1.0

MATERIAL_MODULUS = {
    "รองผิวทางคอนกรีตด้วย AC": 2500, "รองผิวทางคอนกรีตด้วย PMA(AC)": 3700,
    "พื้นทางซีเมนต์ CTB": 1200, "หินคลุกผสมซีเมนต์ UCS 24.5 ksc": 850,
    "หินคลุก CBR 80%": 350, "ดินซีเมนต์ UCS 17.5 ksc": 350,
    "วัสดุหมุนเวียน (Recycling)": 850, "รองพื้นทางวัสดุมวลรวม CBR 25%": 150,
    "วัสดุคัดเลือก ก": 100, "ดินถมคันทาง / ดินเดิม": 100, "กำหนดเอง...": 100,
}

LS_PRESETS = {
    0.0: (138, 715, 753, 84), 0.5: (129, 728, 908, 0), 1.0: (150, 718, 903, 84),
    1.5: (153, 721, 928, 138), 2.0: (164, 718, 929, 220), 3.0: (212, 719, 929, 328)
}

# ============================================================
# ฟังก์ชันการคำนวณ
# ============================================================

def convert_cube_to_cylinder(fc_cube_ksc):
    return 0.8 * fc_cube_ksc

def calculate_concrete_modulus(fc_cylinder_ksc):
    fc_psi = fc_cylinder_ksc * 14.223
    return 57000 * math.sqrt(fc_psi)

def estimate_modulus_of_rupture(fc_cylinder_ksc):
    fc_psi = fc_cylinder_ksc * 14.223
    return 10.0 * math.sqrt(fc_psi)

def get_zr_value(reliability):
    return ZR_TABLE.get(int(reliability), -1.282)

def calculate_aashto_rigid_w18(d_inch, delta_psi, pt, zr, so, sc_psi, cd, j, ec_psi, k_pci):
    term1 = zr * so
    term2 = 7.35 * math.log10(d_inch + 1) - 0.06
    numerator3 = math.log10(delta_psi / (4.5 - 1.5))
    denominator3 = 1 + (1.624e7 / ((d_inch + 1) ** 8.46))
    term3 = numerator3 / denominator3
    d_power = d_inch ** 0.75
    numerator4 = sc_psi * cd * (d_power - 1.132)
    ec_k_ratio = ec_psi / k_pci
    denominator4 = 215.63 * j * (d_power - 18.42 / (ec_k_ratio ** 0.25))
    if numerator4 <= 0 or denominator4 <= 0:
        return (float('-inf'), 0)
    inner_term = numerator4 / denominator4
    if inner_term <= 0:
        return (float('-inf'), 0)
    term4 = (4.22 - 0.32 * pt) * math.log10(inner_term)
    log10_w18 = term1 + term2 + term3 + term4
    w18 = 10 ** log10_w18
    return (log10_w18, w18)

def check_design(w18_required, w18_capacity):
    ratio = w18_capacity / w18_required if w18_required > 0 else float('inf')
    return (w18_capacity >= w18_required, ratio)

def draw_arrow_fixed(draw, start, end, color, width=4, arrow_size=15):
    draw.line([start, end], fill=color, width=width)
    dx = end[0] - start[0]
    dy = end[1] - start[1]
    length = math.sqrt(dx*dx + dy*dy)
    if length > 0:
        dx /= length
        dy /= length
        px, py = -dy, dx
        x3, y3 = end[0], end[1]
        base_x = end[0] - arrow_size * dx
        base_y = end[1] - arrow_size * dy
        x4 = base_x + arrow_size * 0.5 * px
        y4 = base_y + arrow_size * 0.5 * py
        x5 = base_x - arrow_size * 0.5 * px
        y5 = base_y - arrow_size * 0.5 * py
        draw.polygon([(x3, y3), (x4, y4), (x5, y5)], fill=color)

# ============================================================
# ฟังก์ชันสร้างรูปโครงสร้างชั้นทาง
# ============================================================

def create_pavement_structure_figure(layers_data, concrete_thickness_cm=None):
    THAI_TO_ENG = {
        "รองผิวทางคอนกรีตด้วย AC": "AC Interlayer", "รองผิวทางคอนกรีตด้วย PMA(AC)": "PMA Interlayer",
        "พื้นทางซีเมนต์ CTB": "Cement Treated Base", "หินคลุกผสมซีเมนต์ UCS 24.5 ksc": "Mod.Crushed Rock ",
        "หินคลุก CBR 80%": "Crushed Rock Base", "ดินซีเมนต์ UCS 17.5 ksc": "Soil Cement",
        "วัสดุหมุนเวียน (Recycling)": "Recycled Material", "รองพื้นทางวัสดุมวลรวม CBR 25%": "Aggregate Subbase",
        "วัสดุคัดเลือก ก": "Selected Material", "ดินถมคันทาง / ดินเดิม": "Subgrade",
        "กำหนดเอง...": "Custom Material", "แผ่นคอนกรีต": "Concrete Slab", "Concrete Slab": "Concrete Slab",
    }
    LAYER_COLORS = {
        "รองผิวทางคอนกรีตด้วย AC": "#2C3E50", "รองผิวทางคอนกรีตด้วย PMA(AC)": "#1A252F",
        "พื้นทางซีเมนต์ CTB": "#7F8C8D", "หินคลุกผสมซีเมนต์ UCS 24.5 ksc": "#95A5A6",
        "หินคลุก CBR 80%": "#BDC3C7", "ดินซีเมนต์ UCS 17.5 ksc": "#AAB7B8",
        "วัสดุหมุนเวียน (Recycling)": "#85929E", "รองพื้นทางวัสดุมวลรวม CBR 25%": "#FFCC99",
        "วัสดุคัดเลือก ก": "#E8DAEF", "ดินถมคันทาง / ดินเดิม": "#F5CBA7",
        "กำหนดเอง...": "#FADBD8", "Concrete Slab": "#808080",
    }
    
    valid_layers = [l for l in layers_data if l.get("thickness_cm", 0) > 0]
    all_layers = []
    if concrete_thickness_cm and concrete_thickness_cm > 0:
        all_layers.append({"name": "Concrete Slab", "thickness_cm": concrete_thickness_cm, "E_MPa": None})
    all_layers.extend(valid_layers)
    if not all_layers:
        return None
    
    total_thickness = sum(l.get("thickness_cm", 0) for l in all_layers)
    min_display_height = 8
    fig, ax = plt.subplots(figsize=(12, 8))
    width, x_center = 3, 6
    x_start = x_center - width / 2
    display_heights = [max(l.get("thickness_cm", 0), min_display_height) for l in all_layers]
    total_display = sum(display_heights)
    y_current = total_display
    
    for i, layer in enumerate(all_layers):
        thickness = layer.get("thickness_cm", 0)
        name = layer.get("name", f"Layer {i+1}")
        e_mpa = layer.get("E_MPa", None)
        display_h = display_heights[i]
        if thickness <= 0:
            continue
        color = LAYER_COLORS.get(name, "#CCCCCC")
        hatch_pattern = '///' if name == "วัสดุหมุนเวียน (Recycling)" else None
        y_bottom = y_current - display_h
        rect = patches.Rectangle((x_start, y_bottom), width, display_h, linewidth=2,
                                  edgecolor='black', facecolor=color, hatch=hatch_pattern)
        ax.add_patch(rect)
        y_center_pos = y_bottom + display_h / 2
        display_name = THAI_TO_ENG.get(name, name)
        is_dark = name in ["รองผิวทางคอนกรีตด้วย AC", "รองผิวทางคอนกรีตด้วย PMA(AC)", "Concrete Slab",
                          "พื้นทางซีเมนต์ CTB", "หินคลุกผสมซีเมนต์ UCS 24.5 ksc", "วัสดุหมุนเวียน (Recycling)"]
        text_color = 'white' if is_dark else 'black'
        ax.text(x_center, y_center_pos, f"{thickness} cm", ha='center', va='center', fontsize=16, fontweight='bold', color=text_color)
        ax.text(x_start - 0.5, y_center_pos, display_name, ha='right', va='center', fontsize=14, fontweight='bold', color='black')
        if e_mpa:
            ax.text(x_start + width + 0.5, y_center_pos, f"E = {e_mpa:,} MPa", ha='left', va='center', fontsize=12, color='#0066CC')
        y_current = y_bottom
    
    ax.annotate('', xy=(x_start + width + 3.5, total_display), xytext=(x_start + width + 3.5, 0),
                arrowprops=dict(arrowstyle='<->', color='red', lw=2))
    ax.text(x_start + width + 4, total_display / 2, f"Total\n{total_thickness} cm", ha='left', va='center', fontsize=14, color='red', fontweight='bold')
    margin = 10
    ax.set_xlim(0, 14)
    ax.set_ylim(-margin, total_display + margin)
    ax.axis('off')
    ax.set_title('Pavement Structure', fontsize=20, fontweight='bold', pad=20)
    ax.text(x_center, -margin + 4, f"Total Pavement Thickness: {total_thickness} cm", ha='center', va='center', fontsize=15, fontweight='bold',
            bbox=dict(boxstyle='round', facecolor='lightyellow', alpha=0.9, edgecolor='orange'))
    plt.tight_layout()
    return fig

def save_figure_to_bytes(fig):
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white', edgecolor='none')
    buf.seek(0)
    return buf

# ============================================================
# ฟังก์ชัน Save/Load JSON
# ============================================================

def save_project_to_json(project_data):
    json_str = json.dumps(project_data, ensure_ascii=False, indent=2)
    return json_str.encode('utf-8')

def load_project_from_json(uploaded_file):
    try:
        content = uploaded_file.read()
        return json.loads(content.decode('utf-8'))
    except Exception as e:
        st.error(f"เกิดข้อผิดพลาดในการอ่านไฟล์: {str(e)}")
        return None

def collect_design_data(project_name, pavement_type, num_layers, layers_data, w18_design, pt, reliability, so,
                        k_eff, ls_value, fc_cube, sc, j_value, cd, d_cm_selected, cbr_value,
                        mr_val=0, esb_val=0, dsb_val=0, k_inf_val=0, ls_select=0, k_corrected=0):
    return {
        "version": "1.0",
        "save_date": datetime.now().isoformat(),
        "project_info": {"project_name": project_name, "pavement_type": pavement_type},
        "layers": {"num_layers": num_layers, "layers_data": layers_data},
        "design_parameters": {
            "w18_design": w18_design, "pt": pt, "reliability": reliability, "so": so,
            "k_eff": k_eff, "ls_value": ls_value, "fc_cube": fc_cube, "sc": sc,
            "j_value": j_value, "cd": cd, "d_cm_selected": d_cm_selected
        },
        "subgrade": {"cbr_value": cbr_value},
        "nomograph": {"mr_val": mr_val, "esb_val": esb_val, "dsb_val": dsb_val,
                      "k_inf_val": k_inf_val, "ls_select": ls_select, "k_corrected": k_corrected}
    }

# ============================================================
# ฟังก์ชันสร้างรายงาน Word
# ============================================================

def create_word_report(pavement_type, inputs, calculated_values, comparison_results, selected_d_cm,
                       main_result, layers_data=None, project_name="", structure_figure=None,
                       subgrade_info=None, e_equivalent_psi=0):
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        st.error("กรุณาติดตั้ง python-docx: pip install python-docx")
        return None
    
    selected_d_inch = round(selected_d_cm / 2.54)
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'TH SarabunPSK'
    style.font.size = Pt(15)
    
    title = doc.add_heading('รายการคำนวณออกแบบความหนาถนนคอนกรีต', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('ตามวิธี AASHTO 1993').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('1. ข้อมูลทั่วไป', level=1)
    if project_name:
        doc.add_paragraph(f'ชื่อโครงการ: {project_name}')
    doc.add_paragraph(f'ประเภทถนน: {pavement_type}')
    doc.add_paragraph(f'วันที่คำนวณ: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    
    doc.add_heading('2. ชั้นโครงสร้างทาง', level=1)
    table_layers = doc.add_table(rows=1, cols=4)
    table_layers.style = 'Table Grid'
    hdr = table_layers.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'ลำดับ', 'ชนิดวัสดุ', 'ความหนา (ซม.)', 'Modulus E (MPa)'
    row = table_layers.add_row().cells
    row[0].text, row[1].text, row[2].text, row[3].text = '1', f'ผิวทางคอนกรีต {pavement_type}', f'{selected_d_cm}', '-'
    layer_count = 1
    if layers_data:
        for i, layer in enumerate(layers_data):
            layer_count += 1
            row = table_layers.add_row().cells
            row[0].text = str(layer_count)
            row[1].text = layer.get('name', f'Layer {i+1}')
            row[2].text = f"{layer.get('thickness_cm', 0)}"
            row[3].text = f"{layer.get('E_MPa', 0):,}"
    if subgrade_info:
        layer_count += 1
        row = table_layers.add_row().cells
        row[0].text = str(layer_count)
        row[1].text = 'ดินคันทาง'
        row[2].text = f"CBR {subgrade_info.get('cbr', 0)} %"
        row[3].text = f"{subgrade_info.get('mr_mpa', 0):.0f} ({subgrade_info.get('mr_psi', 0):,.0f} psi)"
    
    if structure_figure:
        doc.add_paragraph('รูปตัดโครงสร้างชั้นทาง:')
        img_buf = BytesIO()
        structure_figure.savefig(img_buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
        img_buf.seek(0)
        doc.add_picture(img_buf, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('3. ข้อมูลนำเข้า', level=1)
    table1 = doc.add_table(rows=1, cols=4)
    table1.style = 'Table Grid'
    hdr = table1.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'พารามิเตอร์', 'สัญลักษณ์', 'ค่า', 'หน่วย'
    input_data = [
        ('ESAL ออกแบบ', 'W₁₈', f"{inputs['w18_design']:,.0f}", 'ESALs'),
        ('Terminal Serviceability', 'Pt', f"{inputs['pt']:.1f}", '-'),
        ('Reliability', 'R', f"{inputs['reliability']:.0f}", '%'),
        ('Standard Deviation', 'So', f"{inputs['so']:.2f}", '-'),
        ('Modulus of Subgrade Reaction', 'k_eff', f"{inputs['k_eff']:,.0f}", 'pci'),
        ('Loss of Support', 'LS', f"{inputs.get('ls', 1.0):.1f}", '-'),
        ('กำลังคอนกรีต', "f'c", f"{inputs['fc_cube']:.0f} Cube", 'ksc'),
        ('Modulus of Rupture', 'Sc', f"{inputs['sc']:.0f}", 'psi'),
        ('Load Transfer Coefficient', 'J', f"{inputs['j']:.1f}", '-'),
        ('Drainage Coefficient', 'Cd', f"{inputs['cd']:.1f}", '-'),
    ]
    for param, symbol, value, unit in input_data:
        row = table1.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text = param, symbol, value, unit
    
    doc.add_heading('4. ค่าที่คำนวณได้', level=1)
    table2 = doc.add_table(rows=1, cols=4)
    table2.style = 'Table Grid'
    hdr = table2.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'พารามิเตอร์', 'สัญลักษณ์', 'ค่า', 'หน่วย'
    calc_data = [
        ('Modulus of Elasticity', 'Ec', f"{calculated_values['ec']:,.0f}", 'psi'),
        ('Standard Normal Deviate', 'ZR', f"{calculated_values['zr']:.3f}", '-'),
        ('การสูญเสีย Serviceability', 'ΔPSI', f"{calculated_values['delta_psi']:.1f}", '-'),
    ]
    for param, symbol, value, unit in calc_data:
        row = table2.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text = param, symbol, value, unit
    
    # --------------------------------------------------------
    # 5. สมการออกแบบ AASHTO 1993
    # --------------------------------------------------------
    doc.add_heading('5. สมการออกแบบ AASHTO 1993', level=1)

    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def add_equation_line(document, parts):
        """
        เพิ่มย่อหน้าที่ประกอบด้วย runs หลายส่วน
        parts = list of (text, bold, italic, subscript, superscript)
        """
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for text, bold, italic, is_sub, is_sup in parts:
            run = p.add_run(text)
            run.bold = bold
            run.italic = italic
            run.font.name = 'Times New Roman'
            run.font.size = Pt(15)
            if is_sub or is_sup:
                rPr = run._r.get_or_add_rPr()
                vertAlign = OxmlElement('w:vertAlign')
                vertAlign.set(qn('w:val'), 'subscript' if is_sub else 'superscript')
                rPr.append(vertAlign)
        return p

    def set_paragraph_indent(para, left_twips=360):
        pPr = para._p.get_or_add_pPr()
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(left_twips))
        pPr.append(ind)

    # คำอธิบาย
    p_desc = doc.add_paragraph('สมการหลักที่ใช้ในการออกแบบความหนาถนนคอนกรีตตาม AASHTO 1993 มีดังนี้:')
    p_desc.runs[0].font.name = 'TH SarabunPSK'
    p_desc.runs[0].font.size = Pt(15)

    # บรรทัดที่ 1: log10(W18) = ZR x So + 7.35 x log10(D+1) - 0.06
    line1_parts = [
        ('log', False, False, False, False),
        ('10', False, False, True, False),
        ('(W', False, False, False, False),
        ('18', False, False, True, False),
        (') = Z', False, False, False, False),
        ('R', False, False, True, False),
        (' x S', False, False, False, False),
        ('o', False, False, True, False),
        (' + 7.35 x log', False, False, False, False),
        ('10', False, False, True, False),
        ('(D+1) - 0.06', False, False, False, False),
    ]
    p1 = add_equation_line(doc, line1_parts)
    set_paragraph_indent(p1, 360)

    # บรรทัดที่ 2: + log10(ΔPSI/(4.5-1.5)) / (1 + 1.624x10^7/(D+1)^8.46)
    line2_parts = [
        ('        + log', False, False, False, False),
        ('10', False, False, True, False),
        ('(\u0394PSI/(4.5-1.5)) / (1 + 1.624\u00d710', False, False, False, False),
        ('7', False, False, False, True),
        ('/(D+1)', False, False, False, False),
        ('8.46', False, False, False, True),
        (')', False, False, False, False),
    ]
    p2 = add_equation_line(doc, line2_parts)
    set_paragraph_indent(p2, 360)

    # บรรทัดที่ 3: + (4.22 - 0.32xPt) x log10([ScxCdx(D^0.75-1.132)/(215.63xJx(D^0.75-18.42/(Ec/k)^0.25))])
    line3_parts = [
        ('        + (4.22 - 0.32\u00d7P', False, False, False, False),
        ('t', False, False, True, False),
        (') \u00d7 log', False, False, False, False),
        ('10', False, False, True, False),
        ('[(S', False, False, False, False),
        ('c', False, False, True, False),
        ('\u00d7C', False, False, False, False),
        ('d', False, False, True, False),
        ('\u00d7(D', False, False, False, False),
        ('0.75', False, False, False, True),
        ('-1.132))/(215.63\u00d7J\u00d7(D', False, False, False, False),
        ('0.75', False, False, False, True),
        (' - 18.42/(E', False, False, False, False),
        ('c', False, False, True, False),
        ('/k)', False, False, False, False),
        ('0.25', False, False, False, True),
        (')]', False, False, False, False),
    ]
    p3 = add_equation_line(doc, line3_parts)
    set_paragraph_indent(p3, 360)

    # ตารางสัญลักษณ์
    doc.add_paragraph()
    p_sym = doc.add_paragraph('โดยที่:')
    p_sym.runs[0].font.name = 'TH SarabunPSK'
    p_sym.runs[0].font.size = Pt(15)

    tbl_sym = doc.add_table(rows=1, cols=3)
    tbl_sym.style = 'Table Grid'
    hdr_sym = tbl_sym.rows[0].cells
    hdr_sym[0].text = 'สัญลักษณ์'
    hdr_sym[1].text = 'ความหมาย'
    hdr_sym[2].text = 'หน่วย'
    for cell in hdr_sym:
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.name = 'TH SarabunPSK'
        run.font.size = Pt(15)

    symbol_data = [
        ('W\u2081\u2088',        'จำนวนแกนเดี่ยว 18 kip ที่รองรับได้',     'ESALs'),
        ('Z\u1D3F',              'Standard Normal Deviate ที่ความเชื่อมั่น R', '-'),
        ('S\u2092',              'Overall Standard Deviation',               '-'),
        ('D',                    'ความหนาแผ่นคอนกรีต',                        'นิ้ว'),
        ('\u0394PSI',            'การสูญเสีย Serviceability (4.5 - P\u209C)',  '-'),
        ('P\u209C',              'Terminal Serviceability Index',             '-'),
        ('S\u1D9C',              'Modulus of Rupture ของคอนกรีต',            'psi'),
        ('C\u1D48',              'Drainage Coefficient',                      '-'),
        ('J',                    'Load Transfer Coefficient',                '-'),
        ('E\u1D9C',              'Modulus of Elasticity ของคอนกรีต',         'psi'),
        ('k',                    'Modulus of Subgrade Reaction',             'pci'),
    ]
    for sym, meaning, unit in symbol_data:
        row_s = tbl_sym.add_row().cells
        row_s[0].text = sym
        row_s[1].text = meaning
        row_s[2].text = unit
        for cell in row_s:
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'TH SarabunPSK'
            run.font.size = Pt(15)

    doc.add_paragraph()

    # --------------------------------------------------------
    # 6. ผลการเปรียบเทียบความหนา (เดิมคือหัวข้อ 5)
    # --------------------------------------------------------
    doc.add_heading('6. ผลการเปรียบเทียบความหนา', level=1)
    table3 = doc.add_table(rows=1, cols=6)
    table3.style = 'Table Grid'
    hdr = table3.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = 'D (ซม.)', 'D (นิ้ว)', 'log₁₀(W₁₈)'
    hdr[3].text, hdr[4].text, hdr[5].text = 'W₁₈ รองรับได้', 'อัตราส่วน', 'ผล'
    for r in comparison_results:
        row = table3.add_row().cells
        row[0].text = f"{r['d_cm']:.0f}"
        row[1].text = f"{r['d_inch']:.0f}"
        row[2].text = f"{r['log_w18']:.4f}"
        row[3].text = f"{r['w18']:,.0f}"
        row[4].text = f"{r['ratio']:.2f}"
        row[5].text = "ผ่าน ✓" if r['passed'] else "ไม่ผ่าน ✗"

    doc.add_heading('7. สรุปผล', level=1)
    passed, ratio = main_result
    w18_cap = None
    for r in comparison_results:
        if r['d_cm'] == selected_d_cm:
            w18_cap = r['w18']
            break
    e_eq_mpa = e_equivalent_psi / 145.038 if e_equivalent_psi > 0 else 0
    doc.add_paragraph(f"ความหนาที่เลือก: {selected_d_cm:.0f} ซม. ({selected_d_inch:.0f} นิ้ว)")
    doc.add_paragraph(f"ESAL ที่ต้องการ: {inputs['w18_design']:,.0f} ESALs")
    if w18_cap:
        doc.add_paragraph(f"ESAL ที่รองรับได้: {w18_cap:,.0f} ESALs")
    doc.add_paragraph(f"อัตราส่วน: {ratio:.2f}")
    doc.add_paragraph(f"ผลการตรวจสอบ: {'ผ่านเกณฑ์ ✓' if passed else 'ไม่ผ่านเกณฑ์ ✗'}")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_word_report_nomograph(params, img1_bytes, img2_bytes=None):
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        return None, "ไม่พบ library python-docx"
    
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'TH SarabunPSK'
    style.font.size = Pt(15)
    
    title = doc.add_heading('รายการคำนวณ Corrected Modulus of Subgrade Reaction', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'วันที่: {datetime.now().strftime("%d/%m/%Y %H:%M")}').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_heading('ส่วนที่ 1: การหาค่า Composite Modulus (k∞)', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = 'พารามิเตอร์', 'ค่า', 'หน่วย'
    for h in hdr:
        h.paragraphs[0].runs[0].bold = True
    data1 = [
        ('Roadbed Soil Resilient Modulus (MR)', f"{params.get('MR', 0):,.0f}", 'psi'),
        ('Subbase Elastic Modulus (ESB)', f"{params.get('ESB', 0):,.0f}", 'psi'),
        ('Subbase Thickness (DSB)', f"{params.get('DSB', 0):.1f}", 'inches'),
        ('Composite Modulus (k∞)', f"{params.get('k_inf', 0):,.0f}", 'pci'),
    ]
    for p, v, u in data1:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text = p, v, u
    if img1_bytes:
        doc.add_paragraph()
        doc.add_picture(io.BytesIO(img1_bytes), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    doc.add_heading('ส่วนที่ 2: การปรับแก้ค่า Loss of Support (LS)', level=1)
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    hdr2 = table2.rows[0].cells
    hdr2[0].text, hdr2[1].text, hdr2[2].text = 'พารามิเตอร์', 'ค่า', 'หน่วย'
    for h in hdr2:
        h.paragraphs[0].runs[0].bold = True
    data2 = [
        ('Effective Modulus (k) - จากส่วนที่ 1', f"{params.get('k_inf', 0):,.0f}", 'pci'),
        ('Loss of Support Factor (LS)', f"{params.get('LS_factor', 0):.1f}", '-'),
        ('Corrected Modulus (k)', f"{params.get('k_corrected', 0):,.0f}", 'pci'),
    ]
    for p, v, u in data2:
        row = table2.add_row().cells
        row[0].text, row[1].text, row[2].text = p, v, u
    if img2_bytes:
        doc.add_paragraph()
        doc.add_picture(io.BytesIO(img2_bytes), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph("Reference: AASHTO Guide for Design of Pavement Structures 1993").style = 'List Bullet'
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer, None

# ============================================================
# ฟังก์ชันสร้างรายงาน Word ฉบับสมบูรณ์ (พร้อมบทเกริ่นนำ + เลขหัวข้อยืดหยุ่น)
# ============================================================

DEFAULT_INTRO_TEXT = (
    "การออกแบบความหนาแผ่นคอนกรีตตามแนวทางของ AASHTO 1993 จำเป็นต้องอาศัยสมเหตุสมผลที่"
    "พัฒนามาจากผลการทดสอบ AASHO Road Test ซึ่งสะท้อนพฤติกรรมการรับน้ำหนักและการเสื่อมสภาพของแผ่น"
    "คอนกรีตภายใต้สภาพการใช้งานจริง สมการดังกล่าวรวมปัจจัยสำคัญหลายด้าน ทั้งด้านปริมาณจราจร ความ"
    "น่าเชื่อถือของการออกแบบ คุณสมบัติวัสดุ และสภาพชั้นรองรับ เพื่อให้สามารถประเมินความหนาที่เหมาะสม"
    "สำหรับรองรับปริมาณจราจรตลอดอายุโครงการได้อย่างแม่นยำ สมการหลักที่ใช้ในการออกแบบความหนาถนน"
    "คอนกรีตตาม AASHTO 1993 มีดังนี้"
)

DEFAULT_SUMMARY_TEXT = (
    "จากการคำนวณตามวิธีของ AASHTO 1993 ผิวทางคอนกรีต (Concrete Pavement) สามารถสรุปรูปแบบของ"
    "โครงสร้างชั้นทางที่ออกแบบได้ดังแสดงในตารางและรูปด้านล่าง"
)

def _get_font_name():
    return 'TH SarabunPSK'

def _heading_num(prefix, sub=None):
    """สร้างเลขหัวข้อ เช่น prefix='4.5' sub=1 -> '4.5.1'"""
    if sub is None:
        return prefix
    return f"{prefix}.{sub}"

def _setup_doc_styles(doc):
    """ตั้งค่า Font, Page A4, Margin"""
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    font_name = _get_font_name()
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(15)
    # ตั้งค่า font สำหรับ East Asian
    rPr = style.element.get_or_add_rPr()

    # ตั้ง page A4 + margin
    section = doc.sections[0]
    section.page_width  = int(21.0 * 914400 / 25.4 * 914400 / 914400)   # 21 cm
    section.page_height = int(29.7 * 914400 / 25.4 * 914400 / 914400)   # 29.7 cm
    from docx.shared import Cm
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)

def _add_heading(doc, text, level=1):
    """หัวข้อแบบ bold + underline (ตามภาพตัวอย่าง) ไม่ใช้ Heading style"""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = _get_font_name()
    run.font.size = Pt(15)
    run.bold = True
    run.underline = (level <= 2)   # underline สำหรับ level 1-2, level 3 bold อย่างเดียว
    return p

def _add_para(doc, text, bold=False, italic=False, indent_cm=0, justify=True):
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = _get_font_name()
    run.font.size = Pt(15)
    if indent_cm > 0:
        p.paragraph_format.left_indent = Cm(indent_cm)
    return p

def _add_equation_section(doc):
    """สมการ AASHTO 1993 — Times New Roman 12pt พร้อม subscript/superscript และตารางสัญลักษณ์"""
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    EQ_FONT = 'Times New Roman'
    EQ_SIZE = Pt(11)
    TH_FONT = _get_font_name()
    TH_SIZE = Pt(15)

    def _eq_run(p, text, sub=False, sup=False, bold=False):
        run = p.add_run(text)
        run.font.name = EQ_FONT
        run.font.size = EQ_SIZE
        run.bold = bold
        if sub or sup:
            rPr = run._r.get_or_add_rPr()
            va = OxmlElement('w:vertAlign')
            va.set(qn('w:val'), 'subscript' if sub else 'superscript')
            rPr.append(va)
        return run

    def eq_line():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.space_after = Pt(2)
        return p

    # บรรทัด 1: log10(W18) = ZR × So + 7.35 × log10(D+1) - 0.06
    p1 = eq_line()
    _eq_run(p1, 'log'); _eq_run(p1, '10', sub=True)
    _eq_run(p1, '(W');  _eq_run(p1, '18', sub=True)
    _eq_run(p1, ') = Z'); _eq_run(p1, 'R', sub=True)
    _eq_run(p1, ' \u00d7 S'); _eq_run(p1, 'o', sub=True)
    _eq_run(p1, ' + 7.35 \u00d7 log'); _eq_run(p1, '10', sub=True)
    _eq_run(p1, '(D+1) \u2212 0.06')

    # บรรทัด 2: + log10(ΔPSI/(4.5-1.5)) / (1 + 1.624×10^7/(D+1)^8.46)
    p2 = eq_line()
    _eq_run(p2, '        + log'); _eq_run(p2, '10', sub=True)
    _eq_run(p2, '(\u0394PSI / (4.5 \u2212 1.5)) / (1 + 1.624\u00d710')
    _eq_run(p2, '7', sup=True)
    _eq_run(p2, ' / (D+1)'); _eq_run(p2, '8.46', sup=True); _eq_run(p2, ')')

    # บรรทัด 3: + (4.22 - 0.32×Pt) × log10([Sc×Cd×(D^0.75-1.132)] / [...])
    p3 = eq_line()
    _eq_run(p3, '        + (4.22 \u2212 0.32\u00d7P'); _eq_run(p3, 't', sub=True)
    _eq_run(p3, ') \u00d7 log'); _eq_run(p3, '10', sub=True)
    _eq_run(p3, ' [(S'); _eq_run(p3, 'c', sub=True)
    _eq_run(p3, '\u00d7C'); _eq_run(p3, 'd', sub=True)
    _eq_run(p3, '\u00d7(D'); _eq_run(p3, '0.75', sup=True)
    _eq_run(p3, '\u22121.132)) / (215.63\u00d7J\u00d7(D'); _eq_run(p3, '0.75', sup=True)
    _eq_run(p3, ' \u2212 18.42 / (E'); _eq_run(p3, 'c', sub=True)
    _eq_run(p3, '/k)'); _eq_run(p3, '0.25', sup=True); _eq_run(p3, ')]')

    doc.add_paragraph()

    # "โดยที่:" — TH SarabunPSK 15pt
    p_by = doc.add_paragraph()
    r_by = p_by.add_run('โดยที่:')
    r_by.font.name = TH_FONT; r_by.font.size = TH_SIZE

    # ตารางสัญลักษณ์
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE

    HEADER_BG = 'BDD7EE'
    col_w_sym = [1396, 6281, 1395]

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    def _sym_cell(cell, text, bold=False, font=TH_FONT, fsize=TH_SIZE, bg=None,
                  align=WD_ALIGN_PARAGRAPH.LEFT):
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(text)
        run.font.name = font; run.font.size = fsize; run.bold = bold
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        tcMar = _OE('w:tcMar')
        for side in ['top','bottom','left','right']:
            m = _OE(f'w:{side}'); m.set(_qn('w:w'),'80'); m.set(_qn('w:type'),'dxa')
            tcMar.append(m)
        tcPr.append(tcMar)
        if bg:
            shd = _OE('w:shd'); shd.set(_qn('w:val'),'clear')
            shd.set(_qn('w:color'),'auto'); shd.set(_qn('w:fill'), bg)
            tcPr.append(shd)

    def _set_sym_widths(row):
        for i, cell in enumerate(row.cells):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            tcW = _OE('w:tcW')
            tcW.set(_qn('w:w'), str(col_w_sym[i])); tcW.set(_qn('w:type'),'dxa')
            tcPr.append(tcW)

    hdr = tbl.rows[0]; _set_sym_widths(hdr)
    _sym_cell(hdr.cells[0], 'สัญลักษณ์', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=HEADER_BG)
    _sym_cell(hdr.cells[1], 'ความหมาย',  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=HEADER_BG)
    _sym_cell(hdr.cells[2], 'หน่วย',     bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=HEADER_BG)

    symbols = [
        ('W₁₈',  'จำนวนแกนเดี่ยว 18 kip ที่รองรับได้',              'ESALs'),
        ('ZR',   'Standard Normal Deviate ที่ความเชื่อมั่น R',       '-'),
        ('So',   'Overall Standard Deviation',                        '-'),
        ('D',    'ความหนาแผ่นคอนกรีต',                                'นิ้ว'),
        ('ΔPSI', 'การสูญเสีย Serviceability (4.5 − Pt)',             '-'),
        ('Pt',   'Terminal Serviceability Index',                     '-'),
        ('Sc',   'Modulus of Rupture ของคอนกรีต',                    'psi'),
        ('Cd',   'Drainage Coefficient',                              '-'),
        ('J',    'Load Transfer Coefficient',                         '-'),
        ('Ec',   'Modulus of Elasticity ของคอนกรีต',                 'psi'),
        ('k',    'Modulus of Subgrade Reaction',                      'pci'),
    ]
    for sym, meaning, unit in symbols:
        row = tbl.add_row(); _set_sym_widths(row)
        # สัญลักษณ์ใช้ Times New Roman, ความหมาย/หน่วยใช้ TH SarabunPSK
        _sym_cell(row.cells[0], sym,     font=EQ_FONT, fsize=EQ_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER)
        _sym_cell(row.cells[1], meaning, font=TH_FONT, fsize=TH_SIZE)
        _sym_cell(row.cells[2], unit,    font=EQ_FONT, fsize=EQ_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

def _fmt_layer_name(name: str) -> str:
    """แทน 'CBR xx%' ด้วย 'CBR ≥ xx%' ในชื่อชั้นวัสดุ"""
    import re
    # จับรูปแบบ CBR ตามด้วยช่องว่าง+ตัวเลข+% เช่น CBR 80%, CBR 25%
    return re.sub(r'CBR\s+(\d+\.?\d*)\s*%', r'CBR \u2265 \1%', name)

def _add_layer_table(doc, layers_data, d_cm, pavement_type, fig_caption="",
                     cbr_subgrade=3.0, show_figure=False):
    """ตารางชั้นโครงสร้างทาง รูปแบบตามภาพ:
    คอลัมน์: ลำดับ | ชนิดวัสดุ | ความหนา (ซม.) | Modulus E (MPa)
    Header สีฟ้าอ่อน, แถวข้อมูล justify ซ้าย, ตัวเลข center
    """
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    HEADER_BG = "BDD7EE"
    FONT = _get_font_name()
    FS = Pt(15)

    def _sc(cell, text, bold=False,
            align=WD_ALIGN_PARAGRAPH.LEFT, bg=None):
        """set cell content"""
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = FS
        run.bold = bold
        # cell padding
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for side in ['top', 'bottom', 'left', 'right']:
            m = OxmlElement(f'w:{side}')
            m.set(qn('w:w'), '80')
            m.set(qn('w:type'), 'dxa')
            tcMar.append(m)
        tcPr.append(tcMar)
        if bg:
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), bg)
            tcPr.append(shd)

    # กว้างคอลัมน์ (DXA): ลำดับ | ชนิดวัสดุ | ความหนา | Modulus E
    col_w = [756, 4536, 1728, 2052]
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # ตั้งความกว้างตาราง
    from docx.oxml import OxmlElement as _OE
    tbl_xml = tbl._tbl
    tbl_pr = tbl_xml.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = _OE('w:tblPr')
        tbl_xml.insert(0, tbl_pr)
    tbl_w = OxmlElement('w:tblW')
    tbl_w.set(qn('w:w'), str(sum(col_w)))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_pr.append(tbl_w)

    def _set_col_widths(row):
        for i, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(col_w[i]))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

    # Header
    hdr = tbl.rows[0]
    _set_col_widths(hdr)
    _sc(hdr.cells[0], 'ลำดับ',          bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=HEADER_BG)
    _sc(hdr.cells[1], 'ชนิดวัสดุ',      bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=HEADER_BG)
    _sc(hdr.cells[2], 'ความหนา (ซม.)',  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=HEADER_BG)
    _sc(hdr.cells[3], 'Modulus E (MPa)', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=HEADER_BG)

    # แถวที่ 1: คอนกรีต
    row = tbl.add_row(); _set_col_widths(row)
    _sc(row.cells[0], '1',  align=WD_ALIGN_PARAGRAPH.CENTER)
    _sc(row.cells[1], f'ผิวทางคอนกรีต {pavement_type}')
    _sc(row.cells[2], str(d_cm), align=WD_ALIGN_PARAGRAPH.CENTER)
    _sc(row.cells[3], '-',       align=WD_ALIGN_PARAGRAPH.CENTER)

    # แถวชั้นวัสดุ
    row_num = 2
    for layer in layers_data:
        thick = layer.get('thickness_cm', 0)
        if thick <= 0:
            continue
        e_mpa = layer.get('E_MPa', 0)
        row = tbl.add_row(); _set_col_widths(row)
        _sc(row.cells[0], str(row_num), align=WD_ALIGN_PARAGRAPH.CENTER)
        _sc(row.cells[1], _fmt_layer_name(layer.get('name', '')))
        _sc(row.cells[2], str(thick),   align=WD_ALIGN_PARAGRAPH.CENTER)
        _sc(row.cells[3], f"{e_mpa:,}" if e_mpa > 0 else '-',
            align=WD_ALIGN_PARAGRAPH.CENTER)
        row_num += 1

    # แถว Subgrade
    row = tbl.add_row(); _set_col_widths(row)
    _sc(row.cells[0], str(row_num),          align=WD_ALIGN_PARAGRAPH.CENTER)
    _sc(row.cells[1], 'ดินคันทาง')
    mr_psi = int(1500 * cbr_subgrade if cbr_subgrade < 10 else 1000 + 555 * cbr_subgrade)
    _sc(row.cells[2], f'CBR \u2265 {cbr_subgrade:.1f} %', align=WD_ALIGN_PARAGRAPH.CENTER)
    _sc(row.cells[3], f'{mr_psi:,} ({mr_psi:,} psi)', align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # รูปตัดขวาง (แสดงเฉพาะเมื่อ show_figure=True)
    if show_figure:
        fig = create_pavement_structure_figure(layers_data, d_cm)
        if fig:
            img_buf = BytesIO()
            fig.savefig(img_buf, format='png', dpi=150,
                        bbox_inches='tight', facecolor='white')
            img_buf.seek(0)
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(img_buf, width=Inches(4.5))
            plt.close(fig)

        if fig_caption:
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_cap = p_cap.add_run(fig_caption)
            run_cap.font.name = FONT
            run_cap.font.size = FS
            run_cap.bold = True

    doc.add_paragraph()

def _add_kvalue_section(doc, params, img1_bytes=None, img2_bytes=None,
                        fig_prefix='4-', fig_num_start=4):
    """การคำนวณ k-value (Nomograph) พร้อม caption ใต้รูป"""
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    FONT = _get_font_name()
    FS   = Pt(15)
    HEADER_BG = 'BDD7EE'

    def _sc(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, bg=None):
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        cell.text = ''
        p = cell.paragraphs[0]; p.alignment = align
        run = p.add_run(text)
        run.font.name = FONT; run.font.size = FS; run.bold = bold
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for side in ['top','bottom','left','right']:
            m = OxmlElement(f'w:{side}')
            m.set(qn('w:w'),'80'); m.set(qn('w:type'),'dxa')
            tcMar.append(m)
        tcPr.append(tcMar)
        if bg:
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
            shd.set(qn('w:fill'), bg); tcPr.append(shd)

    def _add_fig_caption(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = FONT; run.font.size = FS
        run.bold = True; run.underline = True

    # ── Step 1: Composite k∞ ──────────────────────────────────────────
    col_w1 = [5772, 1924, 1376]   # sum = 9072
    _add_para(doc, 'ขั้นตอนที่ 1: หาค่า Composite Modulus of Subgrade Reaction (k∞)', bold=True)
    tbl1 = doc.add_table(rows=1, cols=3)
    tbl1.style = 'Table Grid'
    tbl1.alignment = WD_TABLE_ALIGNMENT.LEFT

    def _set_w(row, widths):
        from docx.oxml.ns import qn; from docx.oxml import OxmlElement
        for i, cell in enumerate(row.cells):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(widths[i])); tcW.set(qn('w:type'),'dxa')
            tcPr.append(tcW)

    hdr = tbl1.rows[0]; _set_w(hdr, col_w1)
    _sc(hdr.cells[0], 'พารามิเตอร์', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=HEADER_BG)
    _sc(hdr.cells[1], 'ค่า',         bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=HEADER_BG)
    _sc(hdr.cells[2], 'หน่วย',       bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=HEADER_BG)

    for p_name, val, unit in [
        ('Roadbed Soil Resilient Modulus (MR)', f"{params.get('MR',0):,.0f}", 'psi'),
        ('Subbase Elastic Modulus (ESB)',        f"{params.get('ESB',0):,.0f}",'psi'),
        ('Subbase Thickness (DSB)',              f"{params.get('DSB',0):.1f}", 'inches'),
        ('Composite Modulus k∞',                f"{params.get('k_inf',0):,.0f}",'pci'),
    ]:
        row = tbl1.add_row(); _set_w(row, col_w1)
        _sc(row.cells[0], p_name)
        _sc(row.cells[1], val,  align=WD_ALIGN_PARAGRAPH.CENTER)
        _sc(row.cells[2], unit, align=WD_ALIGN_PARAGRAPH.CENTER)

    if img1_bytes:
        doc.add_paragraph()
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(io.BytesIO(img1_bytes), width=Inches(5.0))
        _add_fig_caption(
            f'รูปที่ {fig_prefix}{fig_num_start}  '
            f'ค่า Composite Modulus of Subgrade Reaction, k\u221e (pci)'
        )

    doc.add_paragraph()

    # ── Step 2: Loss of Support ───────────────────────────────────────
    col_w2 = [5772, 1924, 1376]   # sum = 9072
    _add_para(doc, 'ขั้นตอนที่ 2: ปรับแก้ค่า Loss of Support (LS)', bold=True)
    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = 'Table Grid'
    tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr2 = tbl2.rows[0]; _set_w(hdr2, col_w2)
    _sc(hdr2.cells[0], 'พารามิเตอร์', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=HEADER_BG)
    _sc(hdr2.cells[1], 'ค่า',         bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=HEADER_BG)
    _sc(hdr2.cells[2], 'หน่วย',       bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg=HEADER_BG)

    for p_name, val, unit in [
        ('Effective Modulus k∞ (จาก Step 1)',  f"{params.get('k_inf',0):,.0f}",      'pci'),
        ('Loss of Support Factor (LS)',          f"{params.get('LS_factor',0):.1f}",  '-'),
        ('Corrected Modulus k (ที่ใช้ออกแบบ)', f"{params.get('k_corrected',0):,.0f}",'pci'),
    ]:
        row = tbl2.add_row(); _set_w(row, col_w2)
        _sc(row.cells[0], p_name)
        _sc(row.cells[1], val,  align=WD_ALIGN_PARAGRAPH.CENTER)
        _sc(row.cells[2], unit, align=WD_ALIGN_PARAGRAPH.CENTER)

    if img2_bytes:
        doc.add_paragraph()
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(io.BytesIO(img2_bytes), width=Inches(5.0))
        _add_fig_caption(
            f'รูปที่ {fig_prefix}{fig_num_start+1}  '
            f'การปรับแก้ค่า Modulus of Subgrade Reaction ประสิทธิผล '
            f'เนื่องจากการสูญเสียฐานรองรับ'
        )

    doc.add_paragraph()

def _add_design_result_section(doc, inputs, calculated_values, comparison_results,
                                selected_d_cm, main_result, layers_data, subgrade_info):
    """ตารางผลการคำนวณออกแบบ — รูปแบบตามภาพตัวอย่าง"""
    from docx.shared import Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    HEADER_BG = "BDD7EE"
    FONT = _get_font_name()
    FS = Pt(15)

    def _sc(cell, text, bold=False,
            align=WD_ALIGN_PARAGRAPH.LEFT, bg=None):
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = FS
        run.bold = bold
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for side in ['top', 'bottom', 'left', 'right']:
            m = OxmlElement(f'w:{side}')
            m.set(qn('w:w'), '80')
            m.set(qn('w:type'), 'dxa')
            tcMar.append(m)
        tcPr.append(tcMar)
        if bg:
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), bg)
            tcPr.append(shd)

    def _set_col_widths(row, widths):
        for i, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(widths[i]))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

    # ── หัวข้อ "ข้อมูลนำเข้าการออกแบบ:" (bold underline) ─────────────
    p_lbl = doc.add_paragraph()
    run_lbl = p_lbl.add_run('ข้อมูลนำเข้าการออกแบบ:')
    run_lbl.font.name = FONT
    run_lbl.font.size = FS
    run_lbl.bold = True
    run_lbl.underline = True

    # คอลัมน์: พารามิเตอร์ | สัญลักษณ์ | ค่า | หน่วย
    col_w_in = [3923, 1471, 2207, 1471]
    tbl_in = doc.add_table(rows=1, cols=4)
    tbl_in.style = 'Table Grid'
    tbl_in.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = tbl_in.rows[0]
    _set_col_widths(hdr, col_w_in)
    for i, t in enumerate(['พารามิเตอร์', 'สัญลักษณ์', 'ค่า', 'หน่วย']):
        _sc(hdr.cells[i], t, bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER, bg=HEADER_BG)

    delta_psi = calculated_values.get('delta_psi', 4.5 - inputs['pt'])
    zr        = calculated_values.get('zr', -1.282)
    ec        = calculated_values.get('ec', 0)

    input_rows = [
        ('ESAL ออกแบบ',               'W₁₈',   f"{inputs['w18_design']:,.0f}",      'ESALs'),
        ('Terminal Serviceability',    'Pt',    f"{inputs['pt']:.1f}",                '-'),
        ('การสูญเสีย Serviceability',  'ΔPSI',  f"{delta_psi:.1f}",                  '-'),
        ('Reliability',                'R',     f"{inputs['reliability']:.0f}",       '%'),
        ('Standard Normal Deviate',    'ZR',    f"{zr:.3f}",                          '-'),
        ('Standard Deviation',         'So',    f"{inputs['so']:.2f}",                '-'),
        ('Modulus of Subgrade Reaction','k_eff', f"{inputs['k_eff']:,.0f}",           'pci'),
        ('Loss of Support',            'LS',    f"{inputs.get('ls', 1.0):.1f}",       '-'),
        ('กำลังคอนกรีต',               "f'c",   f"{inputs['fc_cube']:.0f} Cube",     'ksc'),
        ('Modulus of Elasticity',      'Ec',    f"{ec:,.0f}",                         'psi'),
        ('Modulus of Rupture',         'Sc',    f"{inputs['sc']:.0f}",                'psi'),
        ('Load Transfer Coefficient',  'J',     f"{inputs['j']:.1f}",                 '-'),
        ('Drainage Coefficient',       'Cd',    f"{inputs['cd']:.2f}",                '-'),
    ]
    for row_data in input_rows:
        row = tbl_in.add_row()
        _set_col_widths(row, col_w_in)
        _sc(row.cells[0], row_data[0])
        _sc(row.cells[1], row_data[1], align=WD_ALIGN_PARAGRAPH.CENTER)
        _sc(row.cells[2], row_data[2], align=WD_ALIGN_PARAGRAPH.CENTER)
        _sc(row.cells[3], row_data[3], align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # ── ตารางผลการตรวจสอบความหนา ──────────────────────────────────────
    p_lbl2 = doc.add_paragraph()
    run_lbl2 = p_lbl2.add_run('ผลการตรวจสอบความหนาแผ่นคอนกรีต:')
    run_lbl2.font.name = FONT
    run_lbl2.font.size = FS
    run_lbl2.bold = True
    run_lbl2.underline = True

    col_w_res = [1188, 1188, 1620, 2052, 1512, 1512]
    tbl_res = doc.add_table(rows=1, cols=6)
    tbl_res.style = 'Table Grid'
    tbl_res.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr2 = tbl_res.rows[0]
    _set_col_widths(hdr2, col_w_res)
    for i, t in enumerate(['D (ซม.)', 'D (นิ้ว)', 'log₁₀(W₁₈)',
                            'W₁₈ รองรับได้', 'อัตราส่วน', 'ผล']):
        _sc(hdr2.cells[i], t, bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER, bg=HEADER_BG)

    for r in comparison_results:
        is_sel = (r['d_cm'] == selected_d_cm)
        bg_row = 'FFFFAA' if is_sel else None
        bg_res = 'CCFFCC' if r['passed'] else 'FFCCCC'
        row = tbl_res.add_row()
        _set_col_widths(row, col_w_res)
        _sc(row.cells[0], f"{r['d_cm']:.0f}",    bold=is_sel,
            align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg_row)
        _sc(row.cells[1], f"{r['d_inch']:.0f}",
            align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg_row)
        _sc(row.cells[2], f"{r['log_w18']:.4f}",
            align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg_row)
        _sc(row.cells[3], f"{r['w18']:,.0f}",
            align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg_row)
        _sc(row.cells[4], f"{r['ratio']:.2f}",
            align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg_row)
        _sc(row.cells[5], "ผ่าน ✓" if r['passed'] else "ไม่ผ่าน ✗",
            align=WD_ALIGN_PARAGRAPH.CENTER, bg=bg_res)

    doc.add_paragraph()

    # ── สรุปผล ────────────────────────────────────────────────────────
    passed, ratio = main_result
    sel_inch = round(selected_d_cm / 2.54)
    w18_cap  = next((r['w18'] for r in comparison_results
                     if r['d_cm'] == selected_d_cm), 0)

    p_lbl3 = doc.add_paragraph()
    run_lbl3 = p_lbl3.add_run('สรุปผลการออกแบบ:')
    run_lbl3.font.name = FONT
    run_lbl3.font.size = FS
    run_lbl3.bold = True
    run_lbl3.underline = True

    for item in [
        f"ความหนาที่เลือก : {selected_d_cm:.0f} ซม. ({sel_inch:.0f} นิ้ว)",
        f"ESAL ที่ต้องการ  : {inputs['w18_design']:,.0f} ESALs",
        f"ESAL ที่รองรับได้ : {w18_cap:,.0f} ESALs",
        f"อัตราส่วน        : {ratio:.2f}",
        f"ผลการตรวจสอบ  : {'✅ ผ่านเกณฑ์' if passed else '❌ ไม่ผ่านเกณฑ์'}",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(36)
        run = p.add_run(item)
        run.font.name = FONT
        run.font.size = FS

    doc.add_paragraph()


def _add_summary_layer_table(doc, layers_data, d_cm, pavement_type,
                              fig_caption="", cbr_subgrade=3.0):
    """
    ตารางสรุปชั้นทางสำหรับหัวข้อ 4.6  —  รูปแบบภาพ 1
    โครงสร้าง:
      ┌──────┬──────────────────────────┬──────────────┐
      │ ลำดับ│       ชนิดวัสดุ          │ ความหนา (ซม.)│  ← Header สีฟ้า
      ├──────┼──────────────────────────┼──────────────┤
      │  1   │ ผิวทางคอนกรีต JPCP      │      28      │
      │  2   │ ชื่อชั้น                 │       5      │
      │  ... │ ...                      │     ...      │
      ├──────┴──────────────────────────┴──────────────┤
      │         รูปตัดขวาง (merge 3 col)               │
      ├──────┬──────────────────────────┬──────────────┤
      │  N   │ ดินคันทาง               │  CBR x.x %   │
      └──────┴──────────────────────────┴──────────────┘
      Caption: รูปที่ X-X  โครงสร้างชั้นทาง... (bold underline center)
    """
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    HEADER_BG = 'BDD7EE'
    FONT  = _get_font_name()
    FS    = Pt(15)
    col_w = [934, 6004, 2134]   # ลำดับ | ชนิดวัสดุ | ความหนา  (9072 DXA = เต็มหน้า)

    def _qset(el, attr, val):
        el.set(qn(attr), val)

    def _set_widths(row):
        for i, cell in enumerate(row.cells):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            _qset(tcW, 'w:w', str(col_w[i])); _qset(tcW, 'w:type', 'dxa')
            tcPr.append(tcW)

    def _cell_margin(cell):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for side in ['top','bottom','left','right']:
            m = OxmlElement(f'w:{side}')
            _qset(m, 'w:w', '80'); _qset(m, 'w:type', 'dxa')
            tcMar.append(m)
        tcPr.append(tcMar)

    def _bg(cell, color):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        _qset(shd, 'w:val', 'clear'); _qset(shd, 'w:color', 'auto')
        _qset(shd, 'w:fill', color); tcPr.append(shd)

    def _sc(cell, text, bold=False,
            align=WD_ALIGN_PARAGRAPH.LEFT, bg_color=None):
        cell.text = ''
        p = cell.paragraphs[0]; p.alignment = align
        run = p.add_run(text)
        run.font.name = FONT; run.font.size = FS; run.bold = bold
        _cell_margin(cell)
        if bg_color: _bg(cell, bg_color)

    def _merge_row_3col(tbl):
        """เพิ่มแถวและ merge 3 คอลัมน์"""
        row = tbl.add_row()
        a, b, c = row.cells
        a.merge(c)
        return row

    # ── สร้างตาราง ──────────────────────────────────────────────────────
    # นับจำนวนชั้นที่มีความหนา > 0
    valid_layers = [l for l in layers_data if l.get('thickness_cm', 0) > 0]
    # แถว: header + คอนกรีต + ชั้นวัสดุ + merge(รูป) + subgrade
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # ── Header ──────────────────────────────────────────────────────────
    hdr = tbl.rows[0]; _set_widths(hdr)
    _sc(hdr.cells[0], 'ลำดับ',         bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER, bg_color=HEADER_BG)
    _sc(hdr.cells[1], 'ชนิดวัสดุ',    bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER, bg_color=HEADER_BG)
    _sc(hdr.cells[2], 'ความหนา (ซม.)', bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER, bg_color=HEADER_BG)

    # ── แถวคอนกรีต ──────────────────────────────────────────────────────
    row = tbl.add_row(); _set_widths(row)
    _sc(row.cells[0], '1',  align=WD_ALIGN_PARAGRAPH.CENTER)
    _sc(row.cells[1], f'ผิวทางคอนกรีต {pavement_type}')
    _sc(row.cells[2], str(d_cm), align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── แถวชั้นวัสดุ ─────────────────────────────────────────────────────
    row_num = 2
    for layer in valid_layers:
        row = tbl.add_row(); _set_widths(row)
        _sc(row.cells[0], str(row_num), align=WD_ALIGN_PARAGRAPH.CENTER)
        _sc(row.cells[1], _fmt_layer_name(layer.get('name', '')))
        _sc(row.cells[2], str(layer.get('thickness_cm', 0)),
            align=WD_ALIGN_PARAGRAPH.CENTER)
        row_num += 1

    # ── แถว merge — รูปตัดขวาง ──────────────────────────────────────────
    fig = create_pavement_structure_figure(valid_layers, d_cm)
    merged_row = _merge_row_3col(tbl)
    merged_cell = merged_row.cells[0]
    # ตั้ง width ของ merged cell = ผลรวมทั้งหมด
    tc = merged_cell._tc; tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    _qset(tcW, 'w:w', str(sum(col_w))); _qset(tcW, 'w:type', 'dxa')
    tcPr.append(tcW)
    _cell_margin(merged_cell)

    p_fig = merged_cell.paragraphs[0]
    p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if fig:
        img_buf = BytesIO()
        fig.savefig(img_buf, format='png', dpi=150,
                    bbox_inches='tight', facecolor='white')
        img_buf.seek(0)
        p_fig.add_run().add_picture(img_buf, width=Inches(4.2))
        plt.close(fig)

    # ── แถว Subgrade ────────────────────────────────────────────────────
    row = tbl.add_row(); _set_widths(row)
    _sc(row.cells[0], str(row_num), align=WD_ALIGN_PARAGRAPH.CENTER)
    _sc(row.cells[1], 'ดินคันทาง')
    _sc(row.cells[2], f'CBR \u2265 {cbr_subgrade:.1f} %',
        align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Caption ──────────────────────────────────────────────────────────
    if fig_caption:
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cap = p_cap.add_run(fig_caption)
        run_cap.font.name = FONT; run_cap.font.size = FS
        run_cap.bold = True; run_cap.underline = True

    doc.add_paragraph()


def create_full_word_report(
    # ข้อมูลหัวข้อ
    section_prefix,        # เช่น "4.5"
    fig_prefix,            # เช่น "4-"
    fig_start_num,         # เช่น 5
    intro_text,            # บทเกริ่นนำ
    summary_text,          # บทสรุป

    # ข้อมูลโครงการ
    project_name,
    pavement_type,

    # ข้อมูล JPCP/JRCP
    include_jpcp,
    jpcp_layers_data,
    jpcp_d_cm,
    jpcp_inputs,
    jpcp_calc,
    jpcp_comparison,
    jpcp_result,
    jpcp_subgrade,
    jpcp_nomo_params,
    img1_bytes_jpcp,
    img2_bytes_jpcp,

    # ข้อมูล CRCP
    include_crcp,
    crcp_layers_data,
    crcp_d_cm,
    crcp_inputs,
    crcp_calc,
    crcp_comparison,
    crcp_result,
    crcp_subgrade,
    crcp_nomo_params,
    img1_bytes_crcp,
    img2_bytes_crcp,

    # ตัวเลือกเพิ่มเติม
    include_summary_section,
):
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return None, "กรุณาติดตั้ง python-docx: pip install python-docx"

    doc = Document()
    _setup_doc_styles(doc)

    fig_counter = [fig_start_num]

    def next_fig_num():
        n = fig_counter[0]
        fig_counter[0] += 1
        return n

    # ── หน้าปก ──────────────────────────────────────────────────────────
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run('รายการคำนวณออกแบบ\nผิวทางคอนกรีต')
    run_title.font.name = _get_font_name()
    run_title.font.size = Pt(20)
    run_title.bold = True

    doc.add_paragraph()
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run('ตามวิธี AASHTO 1993')
    run_sub.font.name = _get_font_name()
    run_sub.font.size = Pt(16)

    if project_name:
        doc.add_paragraph()
        p_proj = doc.add_paragraph()
        p_proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_proj.add_run(f'โครงการ: {project_name}')
        r.font.name = _get_font_name()
        r.font.size = Pt(15)

    doc.add_paragraph()
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_date.add_run(f'วันที่: {datetime.now().strftime("%d/%m/%Y")}')
    r.font.name = _get_font_name()
    r.font.size = Pt(15)

    doc.add_page_break()

    # ── หัวข้อ X.X  การออกแบบผิวทางคอนกรีต ────────────────────────────
    h_main = _heading_num(section_prefix)
    _add_heading(doc, f'{h_main}  การออกแบบผิวทางคอนกรีต', level=1)

    # บทเกริ่นนำ
    _add_para(doc, intro_text, indent_cm=0)
    doc.add_paragraph()

    # สมการ
    _add_equation_section(doc)

    # คำอธิบายประเภทถนน
    _add_para(doc, (
        'โดยมาตรฐานการออกแบบตามวิธี AASHTO 1993 ได้แนกโครงสร้างทางคอนกรีตออกเป็นหลายรูปแบบตาม'
        'ลักษณะการควบคุมความแตกร้าวและการถ่ายแรงระหว่างแผ่นคอนกรีต แต่ละประเภทมีแนวคิดการออกแบบ'
        'และยุทธวิธีดำเนินโครงการสร้างที่แตกต่างกัน โครงสร้างทางคอนกรีต 3 ประเภทหลักสำหรับการคำนวณ'
        ' ได้แก่ Jointed Plain Concrete Pavement (JPCP), Jointed Reinforced Concrete Pavement (JRCP) '
        'และ Continuously Reinforced Concrete Pavement (CRCP)'
    ))
    doc.add_paragraph()

    # ── JPCP/JRCP ────────────────────────────────────────────────────────
    if include_jpcp:
        h_jpcp_layer = _heading_num(section_prefix, 1)
        _add_heading(doc, f'{h_jpcp_layer}  ชั้นโครงสร้างทางคอนกรีตประเภท JPCP/JRCP', level=2)
        fig_n = next_fig_num()
        caption = f'รูปที่ {fig_prefix}{fig_n}  โครงสร้างชั้นทางผิวทางคอนกรีต แบบ JPCP/JRCP'
        _add_layer_table(doc, jpcp_layers_data, jpcp_d_cm, pavement_type,
                         fig_caption=caption,
                         cbr_subgrade=jpcp_subgrade.get('cbr', 3.0))

        h_jpcp_k = _heading_num(section_prefix, 2)
        _add_heading(doc, f'{h_jpcp_k}  การคำนวณ Corrected Modulus of Subgrade Reaction (k-value) สำหรับ JPCP/JRCP', level=2)
        k_fig_n = next_fig_num()
        _add_kvalue_section(doc, jpcp_nomo_params, img1_bytes_jpcp, img2_bytes_jpcp,
                            fig_prefix=fig_prefix, fig_num_start=k_fig_n)
        fig_counter[0] += 1   # นับรูปที่ 2 ของ nomograph (LS)

        # ผลการออกแบบ JPCP
        _add_heading(doc, f'ผลการออกแบบความหนาผิวทางคอนกรีต JPCP/JRCP', level=3)
        _add_design_result_section(doc, jpcp_inputs, jpcp_calc, jpcp_comparison,
                                   jpcp_d_cm, jpcp_result, jpcp_layers_data, jpcp_subgrade)

    # ── CRCP ─────────────────────────────────────────────────────────────
    if include_crcp:
        sub_offset = 2 if include_jpcp else 0
        h_crcp_layer = _heading_num(section_prefix, sub_offset + 1)
        _add_heading(doc, f'{h_crcp_layer}  ชั้นโครงสร้างทางคอนกรีตประเภท CRCP', level=2)
        fig_n = next_fig_num()
        caption = f'รูปที่ {fig_prefix}{fig_n}  โครงสร้างชั้นทางผิวทางคอนกรีต แบบ CRCP'
        _add_layer_table(doc, crcp_layers_data, crcp_d_cm, 'CRCP',
                         fig_caption=caption,
                         cbr_subgrade=crcp_subgrade.get('cbr', 3.0))

        h_crcp_k = _heading_num(section_prefix, sub_offset + 2)
        _add_heading(doc, f'{h_crcp_k}  การคำนวณ Corrected Modulus of Subgrade Reaction (k-value) สำหรับ CRCP', level=2)
        k_fig_n2 = next_fig_num()
        _add_kvalue_section(doc, crcp_nomo_params, img1_bytes_crcp, img2_bytes_crcp,
                            fig_prefix=fig_prefix, fig_num_start=k_fig_n2)
        fig_counter[0] += 1   # นับรูปที่ 2 ของ nomograph CRCP

        _add_heading(doc, f'ผลการออกแบบความหนาผิวทางคอนกรีต CRCP', level=3)
        _add_design_result_section(doc, crcp_inputs, crcp_calc, crcp_comparison,
                                   crcp_d_cm, crcp_result, crcp_layers_data, crcp_subgrade)

    # ── หัวข้อ X.X+1  สรุปโครงสร้างชั้นทาง ────────────────────────────
    if include_summary_section:
        doc.add_page_break()
        parts = section_prefix.split('.')
        try:
            parts[-1] = str(int(parts[-1]) + 1)
            h_summary = '.'.join(parts)
        except Exception:
            h_summary = section_prefix + '_สรุป'

        _add_heading(doc, f'{h_summary}  สรุปโครงสร้างชั้นทางที่ออกแบบด้วยวิธี AASHTO 1993', level=1)
        _add_para(doc, summary_text)
        doc.add_paragraph()

        if include_jpcp:
            fig_n = next_fig_num()
            _add_para(doc, f'รูปแบบที่ 1: ผิวทางคอนกรีต แบบ JPCP/JRCP  (รูปที่ {fig_prefix}{fig_n})', bold=True)
            _add_summary_layer_table(
                doc, jpcp_layers_data, jpcp_d_cm, pavement_type,
                fig_caption=f'รูปที่ {fig_prefix}{fig_n}  โครงสร้างชั้นทางรูปแบบที่ 1 ผิวทางคอนกรีต แบบ JPCP/JRCP',
                cbr_subgrade=jpcp_subgrade.get('cbr', 3.0))

        if include_crcp:
            fig_n = next_fig_num()
            _add_para(doc, f'รูปแบบที่ 2: ผิวทางคอนกรีต แบบ CRCP  (รูปที่ {fig_prefix}{fig_n})', bold=True)
            _add_summary_layer_table(
                doc, crcp_layers_data, crcp_d_cm, 'CRCP',
                fig_caption=f'รูปที่ {fig_prefix}{fig_n}  โครงสร้างชั้นทางรูปแบบที่ 2 ผิวทางคอนกรีต แบบ CRCP',
                cbr_subgrade=crcp_subgrade.get('cbr', 3.0))

    # ── เอกสารอ้างอิง ────────────────────────────────────────────────────
    doc.add_paragraph()
    _add_para(doc, 'เอกสารอ้างอิง', bold=True)
    _add_para(doc, 'AASHTO Guide for Design of Pavement Structures 1993. American Association of State Highway and Transportation Officials, Washington, D.C.')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer, None


# ============================================================
# Main Application
# ============================================================

def main():
    st.set_page_config(page_title="AASHTO 1993 Rigid Pavement Design", page_icon="🛣️", layout="wide")
    st.title("🛣️ Rigid Pavement Design Calculator (AASHTO 1993)")
    st.markdown("**โปรแกรมออกแบบความหนาถนนคอนกรีต และหาค่า k-value พร้อมปรับแก้ Loss of Support**")
    
    # Initialize Session State
    for key, val in [('k_inf_result', 500), ('img1_bytes', None), ('img2_bytes', None), ('last_uploaded_file', None)]:
        if key not in st.session_state:
            st.session_state[key] = val
    
    # Sidebar: JSON Save/Load
    with st.sidebar:
        st.header("📁 จัดการโปรเจกต์")
        st.subheader("📂 โหลดไฟล์โปรเจกต์")
        uploaded_json = st.file_uploader("อัปโหลดไฟล์ .json", type=['json'], key='json_uploader')
        
        if uploaded_json is not None:
            try:
                # ตรวจสอบว่าเป็นไฟล์ใหม่หรือไม่
                file_id = f"{uploaded_json.name}_{uploaded_json.size}"
                if st.session_state.get('last_uploaded_file') != file_id:
                    st.session_state['last_uploaded_file'] = file_id
                    
                    # โหลดข้อมูลจาก JSON
                    loaded = load_project_from_json(uploaded_json)
                    if loaded:
                        # อัพเดท session_state สำหรับทุก input field
                        
                        # Project Info
                        st.session_state['calc_project_name'] = loaded.get('project_info', {}).get('project_name', '')
                        st.session_state['calc_pave_type'] = loaded.get('project_info', {}).get('pavement_type', 'JPCP')
                        
                        # Layers
                        st.session_state['calc_num_layers'] = loaded.get('layers', {}).get('num_layers', 5)
                        layers_data = loaded.get('layers', {}).get('layers_data', [])
                        for i, layer in enumerate(layers_data):
                            st.session_state[f'calc_layer_name_{i}'] = layer.get('name', '')
                            st.session_state[f'calc_layer_thick_{i}'] = layer.get('thickness_cm', 0)
                            layer_name = layer.get('name', '')
                            st.session_state[f'calc_layer_E_{i}_{layer_name}'] = layer.get('E_MPa', 100)
                        
                        # Design Parameters
                        dp = loaded.get('design_parameters', {})
                        st.session_state['calc_w18'] = dp.get('w18_design', 500000)
                        st.session_state['calc_pt'] = dp.get('pt', 2.0)
                        st.session_state['calc_reliability'] = dp.get('reliability', 90)
                        st.session_state['calc_so'] = dp.get('so', 0.35)
                        st.session_state['calc_k_eff'] = dp.get('k_eff', 200)
                        st.session_state['calc_ls'] = dp.get('ls_value', 1.0)
                        st.session_state['calc_fc'] = dp.get('fc_cube', 350)
                        st.session_state['calc_sc'] = dp.get('sc', 600)
                        st.session_state['calc_j'] = dp.get('j_value', 2.8)
                        st.session_state['calc_cd'] = dp.get('cd', 1.0)
                        st.session_state['calc_d'] = dp.get('d_cm_selected', 30)
                        
                        # Subgrade
                        st.session_state['calc_cbr'] = loaded.get('subgrade', {}).get('cbr_value', 4.0)
                        
                        # Nomograph
                        nomo = loaded.get('nomograph', {})
                        st.session_state['nomo_mr'] = nomo.get('mr_val', 7000)
                        st.session_state['nomo_esb'] = nomo.get('esb_val', 50000)
                        st.session_state['nomo_dsb'] = nomo.get('dsb_val', 6.0)
                        st.session_state['nomo_k_inf'] = nomo.get('k_inf_val', 400)
                        st.session_state['k_inf_result'] = nomo.get('k_inf_val', 400)
                        st.session_state['ls_select_box'] = nomo.get('ls_select', 1.0)
                        st.session_state['k_corr_input'] = nomo.get('k_corrected', 300)
                        
                        st.success("✅ โหลดข้อมูลสำเร็จ!")
                        st.rerun()
            except Exception as e:
                st.error(f"❌ ไม่สามารถอ่านไฟล์ได้: {e}")
        
        # แสดงสถานะโปรเจกต์ที่โหลด
        if st.session_state.get('calc_project_name'):
            st.info(f"📌 โปรเจกต์: {st.session_state.get('calc_project_name', 'ไม่ระบุ')}")
            if st.button("🗑️ ล้างข้อมูลที่โหลด"):
                # ล้าง session_state ทั้งหมด
                keys_to_clear = [key for key in st.session_state.keys() if key.startswith(('calc_', 'nomo_', 'ls_select', 'k_corr', 'k_inf'))]
                for key in keys_to_clear:
                    del st.session_state[key]
                st.session_state['last_uploaded_file'] = None
                st.rerun()
        st.markdown("---")
    
    
    # Define Tabs
    tab1, tab2, tab3, tab4, tab5, tab_report = st.tabs([
        "🔢 AASHTO Calculator", "📊 Nomograph: Composite k∞", "📉 Nomograph: Loss of Support",
        "💾 บันทึกโปรเจกต์", "📋 คู่มือการใช้งาน", "📄 สร้างรายงาน"
    ])
    
    # =========================================================
    # TAB 1: AASHTO Calculator
    # =========================================================
    with tab1:
        st.header("1️⃣ การออกแบบความหนาถนนคอนกรีต (AASHTO 1993)")
        col1, col2 = st.columns([1, 1])
        
        with col1:
            st.subheader("📥 ข้อมูลนำเข้า (Input)")
            project_name = st.text_input("🏗️ ชื่อโครงการ", value=st.session_state.get('calc_project_name', ''), key="calc_project_name")
            st.markdown("---")
            
            pave_options = list(J_VALUES.keys())
            current_pave_type = st.session_state.get('calc_pave_type', 'JPCP')
            default_pave_idx = pave_options.index(current_pave_type) if current_pave_type in pave_options else 1
            pavement_type = st.selectbox("ประเภทผิวทางคอนกรีต", pave_options, index=default_pave_idx, key="calc_pave_type")
            st.markdown("---")
            
            st.subheader("🔶 ชั้นโครงสร้างทาง")
            material_options = list(MATERIAL_MODULUS.keys())
            num_layers = st.slider("จำนวนชั้นวัสดุ", 1, 6, st.session_state.get('calc_num_layers', 5), key="calc_num_layers")
            
            default_layers = [
                {"name": "รองผิวทางคอนกรีตด้วย AC", "thickness_cm": 5},
                {"name": "พื้นทางซีเมนต์ CTB", "thickness_cm": 20},
                {"name": "หินคลุก CBR 80%", "thickness_cm": 15},
                {"name": "รองพื้นทางวัสดุมวลรวม CBR 25%", "thickness_cm": 25},
                {"name": "วัสดุคัดเลือก ก", "thickness_cm": 30},
                {"name": "ดินถมคันทาง / ดินเดิม", "thickness_cm": 0},
            ]
            
            layers_data = []
            for i in range(num_layers):
                st.markdown(f"**ชั้นที่ {i+1}**")
                col_a, col_b, col_c = st.columns([2, 1, 1])
                
                # ใช้ค่าจาก session_state หรือ default
                def_name = st.session_state.get(f'calc_layer_name_{i}', default_layers[i]["name"] if i < len(default_layers) else "กำหนดเอง...")
                def_thick = st.session_state.get(f'calc_layer_thick_{i}', default_layers[i]["thickness_cm"] if i < len(default_layers) else 20)
                def_idx = material_options.index(def_name) if def_name in material_options else len(material_options) - 1
                
                with col_a:
                    layer_name = st.selectbox("เลือกวัสดุ", material_options, index=def_idx, key=f"calc_layer_name_{i}")
                with col_b:
                    layer_thickness = st.number_input("ความหนา (ซม.)", 0, 100, def_thick, key=f"calc_layer_thick_{i}")
                rec_mod = MATERIAL_MODULUS.get(layer_name, 100)
                def_E = st.session_state.get(f'calc_layer_E_{i}_{layer_name}', rec_mod)
                with col_c:
                    layer_modulus = st.number_input("E (MPa)", 10, 10000, def_E, key=f"calc_layer_E_{i}_{layer_name}")
                layers_data.append({"name": layer_name, "thickness_cm": layer_thickness, "E_MPa": layer_modulus})
            
            total_layer_cm = sum(l['thickness_cm'] for l in layers_data)
            st.markdown(f"**รวมความหนา {total_layer_cm:.0f} ซม. ({round(total_layer_cm/2.54)} นิ้ว)**")
            
            # คำนวณ E_equivalent
            valid_layers = [l for l in layers_data if l['thickness_cm'] > 0 and l['E_MPa'] > 0]
            if valid_layers:
                sum_h_e_cbrt = sum(l['thickness_cm'] * (l['E_MPa'] ** (1/3)) for l in valid_layers)
                total_valid_cm = sum(l['thickness_cm'] for l in valid_layers)
                e_eq_mpa = (sum_h_e_cbrt / total_valid_cm) ** 3 if total_valid_cm > 0 else 0
                e_eq_psi = e_eq_mpa * 145.038
                st.info(f"โมดูลัสเทียบเท่า (E_equivalent) = **{e_eq_psi:,.0f} psi** ({e_eq_mpa:.1f} MPa)")
            st.markdown("---")
            
            st.subheader("1️⃣ ปริมาณจราจร 🚛 ")
            with st.expander("📊 ตัวช่วยประมาณ ESAL ตามประเภทถนน", expanded=False):
                st.markdown("""
                | ประเภทถนน | ESAL (ล้าน) |
                |-----------|-------------|
                | ทางหลวงพิเศษระหว่างเมือง | 50-200 |
                | ทางหลวงแผ่นดินสายหลัก | 20-80 |
                | ทางหลวงแผ่นดินสายรอง | 5-30 |
                | ถนนในเมือง | 1-10 |
                """)
            w18_design = st.number_input("ESAL ที่ต้องการรองรับ (W₁₈)", 10000, 500000000, st.session_state.get('calc_w18', 500000), 100000, key="calc_w18")
            esal_million = w18_design / 1_000_000
            st.info(f"**{esal_million:.2f} ล้าน ESALs**")
            st.markdown("---")
            
            st.subheader("2️⃣ Serviceability📉")
            pt = st.slider("Terminal Serviceability (Pt)", 1.5, 3.0, st.session_state.get('calc_pt', 2.0), 0.1, key="calc_pt")
            delta_psi = 4.5 - pt
            st.info(f"ΔPSI = 4.5 - {pt:.1f} = **{delta_psi:.1f}**")
            st.markdown("---")
            
            st.subheader("3️⃣ ความเชื่อมั่น📈")
            reliability = st.select_slider("Reliability (R)", [80, 85, 90, 95], st.session_state.get('calc_reliability', 90), key="calc_reliability")
            zr = get_zr_value(reliability)
            st.info(f"ZR = **{zr:.3f}**")
            so = st.number_input("Standard Deviation (So)", 0.30, 0.45, st.session_state.get('calc_so', 0.35), 0.01, "%.2f", key="calc_so")
            st.markdown("---")
            
            st.subheader("4️⃣ คุณสมบัติดินคันทาง")
            cbr_value = st.number_input("ค่า CBR (%)", 1.0, 100.0, st.session_state.get('calc_cbr', 4.0), 0.5, key="calc_cbr")
            mr_subgrade_psi = 1500 * cbr_value if cbr_value < 10 else 1000 + 555 * cbr_value
            mr_subgrade_mpa = mr_subgrade_psi / 145.038
            st.info(f"M_R = {mr_subgrade_psi:,.0f} psi ({mr_subgrade_mpa:.0f} MPa)")
            
            k_eff = st.number_input("Effective k (pci)", 50, 1000, st.session_state.get('calc_k_eff', 200), 25, key="calc_k_eff")
        
            with st.expander("📊 ตารางค่า Loss of Support แนะนำ (AASHTO 1993)"):
                st.markdown("""
                | ประเภทวัสดุ | Loss of Support (LS) |
                |------------|---------------------|
                | Cement Treated Granular Base | 0.0 - 1.0 |
                | Cement Aggregate Mixtures | 0.0 - 1.0 |
                | Asphalt Treated Base | 0.0 - 1.0 |
                | Bituminous Stabilized Mixtures | 0.0 - 1.0 |
                | Lime Stabilized | 1.0 - 3.0 |
                | Unbound Granular Materials | 1.0 - 3.0 |
                | Fine Grained or Natural Subgrade | 2.0 - 3.0 |
            
                **หมายเหตุ:** ค่า LS ใช้ปรับลดค่า k_eff เพื่อคำนึงถึงการสูญเสียการรองรับจากการกัดเซาะ
                """)
            ls_value = st.number_input("Loss of Support (LS)", 0.0, 3.0, st.session_state.get('calc_ls', 1.0), 0.5, "%.1f", key="calc_ls")
            st.markdown("---")
            
            st.subheader("5️⃣ คุณสมบัติคอนกรีต")
            fc_cube = st.number_input("กำลังอัด Cube (ksc)", 200, 600, st.session_state.get('calc_fc', 350), 10, key="calc_fc")
            fc_cylinder = convert_cube_to_cylinder(fc_cube)
            ec = calculate_concrete_modulus(fc_cylinder)
            st.info(f"f'c (Cyl) = **{fc_cylinder:.0f} ksc** | Ec = **{ec:,.0f} psi**")
            sc_auto = estimate_modulus_of_rupture(fc_cylinder)
            sc = st.number_input("Modulus of Rupture (Sc) psi", 400, 1000, st.session_state.get('calc_sc', int(sc_auto)), 10, key="calc_sc")
            st.markdown("---")
            
            st.subheader("6️⃣ Load Transfer🔗 และ Drainage💧")
            st.caption(f"ค่าแนะนำสำหรับ {pavement_type}: **J = {J_VALUES[pavement_type]}**")
            with st.expander("📊 ตารางค่า Load Transfer Coefficient (J)", expanded=False):
                st.markdown("""
                | ประเภทถนน | J (AC Shoulder_Yes) | J (AC Shoulder_No) | J (Tied P.C.C_Yes) | J (Tied P.C.C_No) |
                |-----------|---------------------|--------------------|--------------------|-------------------|
                | 1. JRCP/JPCP | 3.2 | 3.8-4.4 | 2.5-3.1 (Mid 2.8) | 3.6-4.2 |
                | 2. CRCP | 2.9-3.2 | N/A | 2.3-2.9 (Mid 2.5) | N/A |
                
                **หมายเหตุ:** ค่า J ต่ำ = การถ่ายแรงดี = รองรับ ESAL ได้มากขึ้น
                
                ค่า J สามารถปรับได้ตามเงื่อนไข:
                - มี Dowel Bar: ลดลง 0.2-0.3
                - มี Tied Shoulder: ลดลง 0.2
                - ไม่มี Dowel Bar: เพิ่มขึ้น 0.5-1.0
                """)
            j_auto = J_VALUES[pavement_type]
            j_value = st.number_input("Load Transfer (J)", 2.0, 4.5, st.session_state.get('calc_j', j_auto), 0.1, "%.1f", key="calc_j")
            cd = st.number_input("Drainage (Cd)", 0.7, 1.3, st.session_state.get('calc_cd', 1.0), 0.05, "%.2f", key="calc_cd")

            
        with col2:
            st.subheader("7️⃣ 👷 🚧 ความหนาที่ตรวจสอบ")
            st.caption("ความหนาผิวทางคอนกรีต D (ซม.)")
            d_cm_selected = st.slider("", 20, 40, st.session_state.get('calc_d', 30), 1, key="calc_d", label_visibility="collapsed")
            d_inch_selected = round(d_cm_selected / 2.54)
            st.success(f"**ความหนาผิวทางคอนกรีต D = {d_cm_selected} ซม. ≈ {d_inch_selected} นิ้ว**")
            st.markdown("---")
            st.subheader(f"🎯 ผลการตรวจสอบ D = {d_cm_selected} ซม.")
            log_w18_sel, w18_sel = calculate_aashto_rigid_w18(d_inch_selected, delta_psi, pt, zr, so, sc, cd, j_value, ec, k_eff)
            passed_sel, ratio_sel = check_design(w18_design, w18_sel)
            
            col_a, col_b = st.columns(2)
            with col_a:
                st.metric("log₁₀(W₁₈)", f"{log_w18_sel:.4f}")
                st.metric("W₁₈ รองรับได้", f"{w18_sel:,.0f}", f"{w18_sel - w18_design:+,.0f}")
            with col_b:
                st.metric("W₁₈ ที่ต้องการ", f"{w18_design:,.0f}")
                st.metric("อัตราส่วน", f"{ratio_sel:.2f}")
            
            if passed_sel:
                st.success(f"✅ **ผ่านเกณฑ์** อัตราส่วน = {ratio_sel:.2f}")
            else:
                st.error(f"❌ **ไม่ผ่านเกณฑ์** อัตราส่วน = {ratio_sel:.2f}")
            
            st.markdown("---")
            st.subheader("📊 ผลการคำนวณ")
            comparison_results = []
            thicknesses_cm = [20, 22, 25, 28, 30, 32, 35, 38, 40]

            for d_cm in thicknesses_cm:
                d_inch = round(d_cm / 2.54)
                log_w18, w18_capacity = calculate_aashto_rigid_w18(d_inch, delta_psi, pt, zr, so, sc, cd, j_value, ec, k_eff)
                passed, ratio = check_design(w18_design, w18_capacity)
                comparison_results.append({'d_cm': d_cm, 'd_inch': d_inch, 'log_w18': log_w18, 'w18': w18_capacity, 'passed': passed, 'ratio': ratio})
            
            df = pd.DataFrame([{
                'D (ซม.)': r['d_cm'], 'D (นิ้ว)': r['d_inch'], 'log₁₀(W₁₈)': f"{r['log_w18']:.4f}",
                'W₁₈ รองรับได้': f"{r['w18']:,.0f}", 'อัตราส่วน': f"{r['ratio']:.2f}", 'ผล': "✅" if r['passed'] else "❌"
            } for r in comparison_results])
            st.dataframe(df, use_container_width=True, hide_index=True)
            
            st.markdown("---")
           
            fig_structure = create_pavement_structure_figure(layers_data, d_cm_selected)
            if fig_structure:
                st.pyplot(fig_structure)
                img_buf = save_figure_to_bytes(fig_structure)
                st.download_button("📥 ดาวน์โหลดรูปโครงสร้าง", img_buf, f"pavement_structure_{datetime.now().strftime('%Y%m%d_%H%M')}.png", "image/png")
                plt.close(fig_structure)
            
            st.markdown("---")
            if st.button("📥 สร้างรายงาน Word", type="primary"):
                with st.spinner("กำลังสร้างรายงาน..."):
                    inputs_dict = {'w18_design': w18_design, 'pt': pt, 'reliability': reliability, 'so': so,
                                   'k_eff': k_eff, 'ls': ls_value, 'fc_cube': fc_cube, 'sc': sc, 'j': j_value, 'cd': cd}
                    calc_dict = {'fc_cylinder': fc_cylinder, 'ec': ec, 'zr': zr, 'delta_psi': delta_psi}
                    subgrade_info = {'cbr': cbr_value, 'mr_psi': mr_subgrade_psi, 'mr_mpa': mr_subgrade_mpa}
                    fig_report = create_pavement_structure_figure(layers_data, d_cm_selected)
                    
                    total_cm = sum(l['thickness_cm'] for l in layers_data)
                    sum_h_e_cbrt = sum(l['thickness_cm'] * (l['E_MPa'] ** (1/3)) for l in layers_data if l['thickness_cm'] > 0 and l['E_MPa'] > 0)
                    e_eq_mpa = (sum_h_e_cbrt / total_cm) ** 3 if total_cm > 0 else 0
                    e_eq_psi = e_eq_mpa * 145.038
                    
                    buffer = create_word_report(pavement_type, inputs_dict, calc_dict, comparison_results, d_cm_selected,
                                                (passed_sel, ratio_sel), layers_data, project_name, fig_report, subgrade_info, e_eq_psi)
                    if fig_report:
                        plt.close(fig_report)
                    if buffer:
                        st.download_button("⬇️ ดาวน์โหลดรายงาน (.docx)", buffer, f"AASHTO_Design_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                                           "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    
    # =========================================================
    # TAB 2: Composite Modulus (Nomograph)
    # =========================================================
    with tab2:
        st.header("2️⃣ หาค่า Composite Modulus of Subgrade Reaction (k∞)")
        uploaded_file = st.file_uploader("📂 อัปโหลดภาพ Figure 3.3 (Composite k)", type=['png', 'jpg', 'jpeg'], key='uploader_1')
        
        if uploaded_file is not None:
            image = Image.open(uploaded_file).convert("RGB")
            width, height = image.size
            img_draw = image.copy()
            draw = ImageDraw.Draw(img_draw)
            
            col_ctrl, col_img = st.columns([1, 2])
            with col_ctrl:
                st.subheader("⚙️ ปรับเส้นอ่านค่า")
                with st.expander("1. เส้น Turning Line (เขียว)", expanded=True):
                    gx1 = st.slider("X เริ่ม", 0, width, 411, key="gx1")
                    gy1 = st.slider("Y เริ่ม", 0, height, 339, key="gy1")
                    gx2 = st.slider("X จบ", 0, width, 470, key="gx2")
                    gy2 = st.slider("Y จบ", 0, height, 397, key="gy2")
                    draw.line([(gx1, gy1), (gx2, gy2)], fill="green", width=5)
                    slope_green = (gy2 - gy1) / (gx2 - gx1) if (gx2 - gx1) != 0 else 0
                
                with st.expander("2. พารามิเตอร์ (ส้ม/แดง/น้ำเงิน)", expanded=True):
                    start_x = st.slider("ตำแหน่งแกน D_sb (ซ้าย)", 0, width, int(width*0.15), key="s1_sx")
                    stop_y_esb = st.slider("ระดับค่า ESB (บน)", 0, height, int(height*0.10), key="s1_sy_esb")
                    stop_y_mr = st.slider("ระดับค่า MR (ล่าง)", 0, height, int(height*0.55), key="s1_sy_mr")
                    constrained_x = int(gx1 + (stop_y_mr - gy1) / slope_green) if slope_green != 0 else gx1
                
                lw = 4
                draw_arrow_fixed(draw, (start_x, stop_y_esb), (constrained_x, stop_y_esb), "orange", lw)
                draw_arrow_fixed(draw, (start_x, stop_y_esb), (start_x, stop_y_mr), "red", lw)
                draw_arrow_fixed(draw, (start_x, stop_y_mr), (constrained_x, stop_y_mr), "darkblue", lw)
                draw_arrow_fixed(draw, (constrained_x, stop_y_mr), (constrained_x, stop_y_esb), "blue", lw)
                r = 8
                draw.ellipse([(constrained_x-r, stop_y_mr-r), (constrained_x+r, stop_y_mr+r)], fill="black", outline="white")
                
                st.markdown("---")
                st.subheader("📝 บันทึกค่าที่อ่านได้")
                mr_val = st.number_input("MR (psi)", value=st.session_state.get('nomo_mr', 7000), step=500, key="nomo_mr")
                esb_val = st.number_input("ESB (psi)", value=st.session_state.get('nomo_esb', 50000), step=1000, key="nomo_esb")
                dsb_val = st.number_input("DSB (inches)", value=st.session_state.get('nomo_dsb', 6.0), step=0.5, key="nomo_dsb")
                k_inf_val = st.number_input("ค่า k∞ ที่อ่านได้ (pci)", value=st.session_state.get('nomo_k_inf', 400), step=10, key="nomo_k_inf")
                st.session_state.k_inf_result = k_inf_val
                
                buf = io.BytesIO()
                img_draw.save(buf, format='PNG')
                st.session_state.img1_bytes = buf.getvalue()
            
            with col_img:
                st.image(img_draw, caption="Step 1: Nomograph Analysis", use_container_width=True)
    
    # =========================================================
    # TAB 3: Loss of Support (Nomograph)
    # =========================================================
    with tab3:
        st.header("3️⃣ ปรับแก้ Loss of Support (LS)")
        st.info("ใช้กราฟ Figure 3.4 เพื่อปรับค่า k∞ กรณีที่มีการสูญเสียการรองรับ (LS > 0)")
        uploaded_file_2 = st.file_uploader("📂 อัปโหลดภาพ Figure 3.4 (LS Correction)", type=['png', 'jpg', 'jpeg'], key='uploader_2')
        
        if uploaded_file_2 is not None:
            img2 = Image.open(uploaded_file_2).convert("RGB")
            w2, h2 = img2.size
            img2_draw = img2.copy()
            draw2 = ImageDraw.Draw(img2_draw)
            
            col_ctrl2, col_img2 = st.columns([1, 2])
            with col_ctrl2:
                st.subheader("⚙️ กำหนดเส้นกราฟ")
                st.write("#### 1. เลือกค่า LS (เส้นแดง)")
                ls_options = [0.0, 0.5, 1.0, 1.5, 2.0, 3.0]
                current_ls = st.session_state.get('ls_select_box', 1.0)
                default_ls_idx = ls_options.index(current_ls) if current_ls in ls_options else 2
                ls_select = st.selectbox("เลือกค่า LS", ls_options, index=default_ls_idx, key="ls_select_box")
                
                if 'last_ls_select' not in st.session_state or st.session_state.last_ls_select != ls_select:
                    st.session_state.last_ls_select = ls_select
                    coords = LS_PRESETS.get(ls_select, (150, 718, 903, 84))
                    st.session_state['_ls_x1'], st.session_state['_ls_y1'] = coords[0], coords[1]
                    st.session_state['_ls_x2'], st.session_state['_ls_y2'] = coords[2], coords[3]
                
                with st.expander("ปรับแต่งตำแหน่งเส้น LS ละเอียด", expanded=False):
                    ls_x1 = st.slider("จุดเริ่ม X", -100, w2+100, key="_ls_x1")
                    ls_y1 = st.slider("จุดเริ่ม Y", -100, h2+100, key="_ls_y1")
                    ls_x2 = st.slider("จุดจบ X", -100, w2+100, key="_ls_x2")
                    ls_y2 = st.slider("จุดจบ Y", -100, h2+100, key="_ls_y2")
                
                draw2.line([(ls_x1, ls_y1), (ls_x2, ls_y2)], fill="red", width=6)
                m_red = (ls_y2 - ls_y1) / (ls_x2 - ls_x1) if ls_x2 - ls_x1 != 0 else None
                c_red = ls_y1 - m_red * ls_x1 if m_red else 0
                
                st.markdown("---")
                st.write("#### 2. ค่า k และขอบเขตแกน (เส้นเขียว)")
                with st.expander("📍 ตั้งค่าตำแหน่งแกนกราฟ", expanded=True):
                    col_b1, col_b2 = st.columns(2)
                    with col_b1:
                        axis_left_x = st.number_input("ตำแหน่งแกน Y (ซ้ายสุด)", value=100, step=5, key="axis_left")
                    with col_b2:
                        axis_bottom_y = st.number_input("ตำแหน่งแกน X (ล่างสุด)", value=h2-50, step=5, key="axis_bottom")
                
                st.caption(f"ค่า k จาก Step 1 คือ: {st.session_state.k_inf_result} pci")
                k_input_x = st.slider("ตำแหน่ง k บนแกน X", 0, w2, int(w2*0.5), key="k_pos_x")
                intersect_y = int(m_red * k_input_x + c_red) if m_red else h2//2
                
                draw2.line([(k_input_x, axis_bottom_y), (k_input_x, intersect_y)], fill="springgreen", width=5)
                draw_arrow_fixed(draw2, (k_input_x, intersect_y), (axis_left_x, intersect_y), "springgreen", width=5)
                draw2.ellipse([(k_input_x-8, intersect_y-8), (k_input_x+8, intersect_y+8)], fill="black", outline="white", width=2)
                
                st.markdown("---")
                st.subheader("📝 บันทึกผลลัพธ์")
                k_corrected = st.number_input("Corrected k (pci)", value=st.session_state.get('k_corr_input', st.session_state.k_inf_result - 100), step=10, key="k_corr_input")
                
                buf2 = io.BytesIO()
                img2_draw.save(buf2, format='PNG')
                st.session_state.img2_bytes = buf2.getvalue()
                
                st.markdown("---")
                params = {
                    'MR': st.session_state.get('nomo_mr', 7000),
                    'ESB': st.session_state.get('nomo_esb', 50000),
                    'DSB': st.session_state.get('nomo_dsb', 6.0),
                    'k_inf': st.session_state.k_inf_result,
                    'LS_factor': ls_select,
                    'k_corrected': k_corrected
                }
                if st.button("📄 สร้างรายงาน Nomograph (Word)", key="btn_nomo_report"):
                    with st.spinner("กำลังสร้างรายงาน..."):
                        doc_file, err = generate_word_report_nomograph(params, st.session_state.get('img1_bytes'), st.session_state.get('img2_bytes'))
                        if err:
                            st.error(err)
                        else:
                            st.download_button("📥 ดาวน์โหลด Word Report", doc_file, f"AASHTO_Nomograph_{datetime.now().strftime('%Y%m%d')}.docx",
                                               "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            
            with col_img2:
                st.image(img2_draw, caption=f"Step 2: LS Correction (LS={ls_select})", use_container_width=True)
        else:
            st.info("👆 กรุณาอัปโหลดภาพ Figure 3.4 เพื่อเริ่มใช้งาน")
    
    # =========================================================
    # TAB 4: Save Project
    # =========================================================
    with tab4:
        st.header("💾 บันทึกโปรเจกต์")
        st.info("บันทึกข้อมูลทั้งหมดเป็นไฟล์ JSON เพื่อโหลดกลับมาแก้ไขภายหลัง")
        
        if st.button("💾 สร้างไฟล์บันทึก", type="primary"):
            project_data = collect_design_data(
                project_name=st.session_state.get('calc_project_name', ''),
                pavement_type=st.session_state.get('calc_pave_type', 'JPCP'),
                num_layers=st.session_state.get('calc_num_layers', 5),
                layers_data=[{"name": st.session_state.get(f'calc_layer_name_{i}', ''),
                              "thickness_cm": st.session_state.get(f'calc_layer_thick_{i}', 0),
                              "E_MPa": st.session_state.get(f'calc_layer_E_{i}_{st.session_state.get(f"calc_layer_name_{i}", "")}', 100)}
                             for i in range(st.session_state.get('calc_num_layers', 5))],
                w18_design=st.session_state.get('calc_w18', 500000),
                pt=st.session_state.get('calc_pt', 2.0),
                reliability=st.session_state.get('calc_reliability', 90),
                so=st.session_state.get('calc_so', 0.35),
                k_eff=st.session_state.get('calc_k_eff', 200),
                ls_value=st.session_state.get('calc_ls', 1.0),
                fc_cube=st.session_state.get('calc_fc', 350),
                sc=st.session_state.get('calc_sc', 600),
                j_value=st.session_state.get('calc_j', 2.8),
                cd=st.session_state.get('calc_cd', 1.0),
                d_cm_selected=st.session_state.get('calc_d', 30),
                cbr_value=st.session_state.get('calc_cbr', 4.0),
                mr_val=st.session_state.get('nomo_mr', 7000),
                esb_val=st.session_state.get('nomo_esb', 50000),
                dsb_val=st.session_state.get('nomo_dsb', 6.0),
                k_inf_val=st.session_state.get('nomo_k_inf', 400),
                ls_select=st.session_state.get('ls_select_box', 1.0),
                k_corrected=st.session_state.get('k_corr_input', 300)
            )
            json_bytes = save_project_to_json(project_data)
            proj_name = project_data['project_info']['project_name'] or 'AASHTO_Project'
            st.download_button("📥 ดาวน์โหลดไฟล์ JSON", json_bytes, f"{proj_name}_{datetime.now().strftime('%Y%m%d_%H%M')}.json", "application/json")
            st.success("สร้างไฟล์บันทึกสำเร็จ!")
    
    # =========================================================
    # TAB 5: User Guide
    # =========================================================
    with tab5:
        st.header("📋 คู่มือการใช้งาน")
        st.markdown("""
        ### 🔢 Tab 1: AASHTO Calculator
        1. กรอกข้อมูลโครงการและชั้นโครงสร้างทาง
        2. ระบุ ESAL, Serviceability, Reliability
        3. ระบุคุณสมบัติดินและคอนกรีต
        4. เลือกความหนาที่ต้องการตรวจสอบ
        5. ดูผลการคำนวณและสร้างรายงาน
        
        ### 📊 Tab 2: Nomograph - Composite k∞
        1. อัปโหลดรูป **Figure 3.3**
        2. ปรับ **Turning Line (เส้นเขียว)** ให้ตรงกับเส้นบนกราฟ
        3. ปรับตำแหน่งลูกศรสีแดง/ส้ม ให้ตรงกับค่า **MR** และ **ESB**
        4. บันทึกค่า k∞ ที่อ่านได้
        
        ### 📉 Tab 3: Nomograph - Loss of Support
        1. อัปโหลดรูป **Figure 3.4**
        2. เลือกค่า **LS** จากตัวเลือก
        3. ตั้งค่าตำแหน่งแกนกราฟ
        4. เลื่อน Slider ตำแหน่ง k บนแกน X
        5. อ่านค่า Corrected k และบันทึก
        
        ### 💾 Tab 4: บันทึกโปรเจกต์
        - กดปุ่ม **สร้างไฟล์บันทึก** เพื่อบันทึกข้อมูลทั้งหมดเป็น JSON
        - ไฟล์ JSON สามารถอัปโหลดกลับมาได้ที่ **Sidebar**
        
        ---
        **Reference:** AASHTO Guide for Design of Pavement Structures 1993
        """)
    
    # =========================================================
    # TAB REPORT: สร้างรายงาน Word ฉบับสมบูรณ์
    # =========================================================
    with tab_report:
        st.header("📄 สร้างรายงาน Word ฉบับสมบูรณ์")
        st.info("รายงานครบถ้วน: บทเกริ่นนำ + สมการ + ชั้นโครงสร้างทาง + k-value + สรุป (ไฟล์เดียว)")

        col_cfg, col_preview = st.columns([1, 1])

        with col_cfg:
            st.subheader("⚙️ ตั้งค่ารายงาน")

            with st.expander("🔢 เลขหัวข้อและเลขรูป", expanded=True):
                rpt_prefix = st.text_input(
                    "Prefix หัวข้อหลัก (เช่น 4.5)",
                    value=st.session_state.get('rpt_prefix', '4.5'),
                    key='rpt_prefix',
                    help="ระบบจะสร้าง 4.5.1, 4.5.2 ... อัตโนมัติ"
                )
                col_fig1, col_fig2 = st.columns(2)
                with col_fig1:
                    rpt_fig_prefix = st.text_input(
                        "Prefix เลขรูป (เช่น 4-)",
                        value=st.session_state.get('rpt_fig_prefix', '4-'),
                        key='rpt_fig_prefix'
                    )
                with col_fig2:
                    rpt_fig_start = st.number_input(
                        "เริ่มที่รูปที่",
                        min_value=1, max_value=99,
                        value=st.session_state.get('rpt_fig_start', 5),
                        step=1, key='rpt_fig_start'
                    )
                st.caption(f"ตัวอย่าง: รูปที่ {rpt_fig_prefix}{rpt_fig_start}, {rpt_fig_prefix}{rpt_fig_start+1} ...")

            with st.expander("📝 บทเกริ่นนำ", expanded=True):
                rpt_intro = st.text_area(
                    "เนื้อหาบทเกริ่นนำ (แก้ไขได้)",
                    value=st.session_state.get('rpt_intro', DEFAULT_INTRO_TEXT),
                    height=180,
                    key='rpt_intro'
                )

            with st.expander("📋 บทสรุป (หัวข้อสรุปโครงสร้างชั้นทาง)", expanded=False):
                rpt_summary_text = st.text_area(
                    "เนื้อหาบทสรุป",
                    value=st.session_state.get('rpt_summary_text', DEFAULT_SUMMARY_TEXT),
                    height=100,
                    key='rpt_summary_text'
                )

            st.markdown("---")
            st.subheader("📑 เลือกเนื้อหาที่รวมในรายงาน")

            rpt_include_jpcp = st.checkbox(
                "✅ รวม JPCP/JRCP (จากข้อมูล Tab 1)",
                value=st.session_state.get('rpt_include_jpcp', True),
                key='rpt_include_jpcp'
            )
            rpt_include_crcp = st.checkbox(
                "⬜ รวม CRCP (ต้องกรอกข้อมูลแยก)",
                value=st.session_state.get('rpt_include_crcp', False),
                key='rpt_include_crcp'
            )
            rpt_include_summary = st.checkbox(
                "✅ รวมหัวข้อสรุปโครงสร้างชั้นทาง",
                value=st.session_state.get('rpt_include_summary', True),
                key='rpt_include_summary'
            )

            # ข้อมูล CRCP แยก (ถ้าเลือก)
            if rpt_include_crcp:
                st.markdown("---")
                st.subheader("🔧 ข้อมูล CRCP (แยกจาก JPCP)")
                st.caption("กรอกเฉพาะค่าที่แตกต่างจาก JPCP — ค่าอื่นใช้ร่วมกัน")

                col_c1, col_c2 = st.columns(2)
                with col_c1:
                    crcp_d_manual = st.number_input(
                        "ความหนา CRCP (ซม.)", 20, 40,
                        value=st.session_state.get('rpt_crcp_d', 28),
                        key='rpt_crcp_d'
                    )
                    crcp_j_manual = st.number_input(
                        "Load Transfer J (CRCP)", 2.0, 4.5,
                        value=st.session_state.get('rpt_crcp_j', 2.5),
                        step=0.1, format="%.1f", key='rpt_crcp_j',
                        help="ค่าแนะนำ CRCP = 2.5 (มี Tied shoulder)"
                    )
                    # Sc ใช้ร่วมกับ JPCP — แสดงค่าอย่างเดียว
                    crcp_sc_use = st.session_state.get('calc_sc', 600)
                    st.caption("Modulus of Rupture Sc (psi)")
                    st.info(f"**{crcp_sc_use}** psi  *(ใช้ร่วมกับ JPCP)*", icon="📌")
                with col_c2:
                    crcp_k_manual = st.number_input(
                        "k_eff CRCP (pci)", 50, 1000,
                        value=st.session_state.get('rpt_crcp_k', 200),
                        step=25, key='rpt_crcp_k'
                    )
                    crcp_cd_manual = st.number_input(
                        "Drainage Cd (CRCP)", 0.7, 1.3,
                        value=st.session_state.get('rpt_crcp_cd',
                              st.session_state.get('calc_cd', 1.0)),
                        step=0.05, format="%.2f", key='rpt_crcp_cd'
                    )
                    # CBR ใช้ร่วมกับ JPCP — แสดงค่าอย่างเดียว
                    crcp_cbr_use = st.session_state.get('calc_cbr', 4.0)
                    st.caption("CBR ดินคันทาง (%)")
                    st.info(f"**{crcp_cbr_use:.1f}** %  *(ใช้ร่วมกับ JPCP)*", icon="📌")
                # แสดงสรุปค่า
                crcp_sc_use  = st.session_state.get('calc_sc', 600)
                crcp_cbr_use = st.session_state.get('calc_cbr', 4.0)
                st.caption(
                    f"📊 CRCP: D={st.session_state.get('rpt_crcp_d',28)} ซม. | "
                    f"J={st.session_state.get('rpt_crcp_j',2.5):.1f} | "
                    f"Sc={crcp_sc_use} psi | "
                    f"k={st.session_state.get('rpt_crcp_k',200)} pci | "
                    f"Cd={st.session_state.get('rpt_crcp_cd',1.0):.2f} | "
                    f"CBR={crcp_cbr_use:.1f}%"
                )

        with col_preview:
            st.subheader("👁️ ตัวอย่างโครงสร้างรายงาน")
            prev_lines = [
                f"📄 **หน้าปก**",
                f"────────────────────────────",
                f"**{rpt_prefix}**  การออกแบบผิวทางคอนกรีต",
                f"   *(บทเกริ่นนำ + สมการ AASHTO 1993)*",
                f"",
            ]
            sub_n = 1
            if rpt_include_jpcp:
                prev_lines += [
                    f"**{rpt_prefix}.{sub_n}**  ชั้นโครงสร้างทาง JPCP/JRCP",
                    f"   *(รูปที่ {rpt_fig_prefix}{rpt_fig_start})*",
                    f"**{rpt_prefix}.{sub_n+1}**  k-value สำหรับ JPCP/JRCP",
                    f"   *(Nomograph + ตาราง k_eff + ผลการออกแบบ)*",
                    f"",
                ]
                sub_n += 2
            if rpt_include_crcp:
                prev_lines += [
                    f"**{rpt_prefix}.{sub_n}**  ชั้นโครงสร้างทาง CRCP",
                    f"   *(รูปที่ {rpt_fig_prefix}{rpt_fig_start + (2 if rpt_include_jpcp else 0)})*",
                    f"**{rpt_prefix}.{sub_n+1}**  k-value สำหรับ CRCP",
                    f"   *(Nomograph + ตาราง k_eff + ผลการออกแบบ)*",
                    f"",
                ]
                sub_n += 2
            if rpt_include_summary:
                try:
                    parts = rpt_prefix.split('.')
                    parts[-1] = str(int(parts[-1]) + 1)
                    h_sum = '.'.join(parts)
                except Exception:
                    h_sum = rpt_prefix + '_สรุป'
                prev_lines += [
                    f"**{h_sum}**  สรุปโครงสร้างชั้นทาง AASHTO 1993",
                    f"   *(ตาราง + รูปตัดขวาง รูปแบบที่ 1-2)*",
                ]
            st.markdown('\n'.join(prev_lines))

            st.markdown("---")
            st.caption("🔴 หมายเหตุ: รายงานใช้ข้อมูลจาก Tab 1 (AASHTO Calculator) และ Tab 2-3 (Nomograph)")

        st.markdown("---")

        # ── ปุ่มสร้างรายงาน ──────────────────────────────────────────────
        if st.button("📄 สร้างรายงาน Word (ฉบับสมบูรณ์)", type="primary", use_container_width=True):
            with st.spinner("กำลังสร้างรายงาน..."):
                # รวบรวมข้อมูลจาก session_state (Tab 1)
                proj_name_r   = st.session_state.get('calc_project_name', '')
                pave_type_r   = st.session_state.get('calc_pave_type', 'JPCP')
                num_layers_r  = st.session_state.get('calc_num_layers', 5)
                layers_r      = [
                    {
                        "name": st.session_state.get(f'calc_layer_name_{i}', ''),
                        "thickness_cm": st.session_state.get(f'calc_layer_thick_{i}', 0),
                        "E_MPa": st.session_state.get(
                            f'calc_layer_E_{i}_{st.session_state.get(f"calc_layer_name_{i}", "")}', 100
                        )
                    }
                    for i in range(num_layers_r)
                ]
                w18_r    = st.session_state.get('calc_w18', 500000)
                pt_r     = st.session_state.get('calc_pt', 2.0)
                rel_r    = st.session_state.get('calc_reliability', 90)
                so_r     = st.session_state.get('calc_so', 0.35)
                k_eff_r  = st.session_state.get('calc_k_eff', 200)
                ls_r     = st.session_state.get('calc_ls', 1.0)
                fc_r     = st.session_state.get('calc_fc', 350)
                sc_r     = st.session_state.get('calc_sc', 600)
                j_r      = st.session_state.get('calc_j', 2.8)
                cd_r     = st.session_state.get('calc_cd', 1.0)
                d_r      = st.session_state.get('calc_d', 30)
                cbr_r    = st.session_state.get('calc_cbr', 4.0)

                # คำนวณค่า
                fc_cyl_r = convert_cube_to_cylinder(fc_r)
                ec_r     = calculate_concrete_modulus(fc_cyl_r)
                zr_r     = get_zr_value(rel_r)
                dpsi_r   = 4.5 - pt_r
                mr_r     = 1500 * cbr_r if cbr_r < 10 else 1000 + 555 * cbr_r

                inputs_r = {
                    'w18_design': w18_r, 'pt': pt_r, 'reliability': rel_r, 'so': so_r,
                    'k_eff': k_eff_r, 'ls': ls_r, 'fc_cube': fc_r, 'sc': sc_r, 'j': j_r, 'cd': cd_r
                }
                calc_r = {'fc_cylinder': fc_cyl_r, 'ec': ec_r, 'zr': zr_r, 'delta_psi': dpsi_r}
                subgrade_r = {'cbr': cbr_r, 'mr_psi': mr_r, 'mr_mpa': mr_r / 145.038}

                # คำนวณตารางเปรียบเทียบ
                thicknesses_cm = [20, 22, 25, 28, 30, 32, 35, 38, 40]
                comparison_r = []
                for d_cm in thicknesses_cm:
                    d_inch = round(d_cm / 2.54)
                    log_w18, w18_cap = calculate_aashto_rigid_w18(
                        d_inch, dpsi_r, pt_r, zr_r, so_r, sc_r, cd_r, j_r, ec_r, k_eff_r
                    )
                    passed, ratio = check_design(w18_r, w18_cap)
                    comparison_r.append({
                        'd_cm': d_cm, 'd_inch': d_inch, 'log_w18': log_w18,
                        'w18': w18_cap, 'passed': passed, 'ratio': ratio
                    })

                d_inch_sel = round(d_r / 2.54)
                log_w18_sel, w18_sel = calculate_aashto_rigid_w18(
                    d_inch_sel, dpsi_r, pt_r, zr_r, so_r, sc_r, cd_r, j_r, ec_r, k_eff_r
                )
                passed_sel, ratio_sel = check_design(w18_r, w18_sel)
                main_result_r = (passed_sel, ratio_sel)

                # Nomograph params
                nomo_r = {
                    'MR':          st.session_state.get('nomo_mr', 7000),
                    'ESB':         st.session_state.get('nomo_esb', 50000),
                    'DSB':         st.session_state.get('nomo_dsb', 6.0),
                    'k_inf':       st.session_state.get('k_inf_result', 400),
                    'LS_factor':   st.session_state.get('ls_select_box', 1.0),
                    'k_corrected': st.session_state.get('k_corr_input', 300),
                }

                # ── ข้อมูล CRCP (ค่าแยกจาก JPCP) ──────────────────────────
                crcp_d_use   = st.session_state.get('rpt_crcp_d', 28)
                crcp_k_use   = st.session_state.get('rpt_crcp_k', 200)
                crcp_j_use   = st.session_state.get('rpt_crcp_j', 2.5)
                crcp_cd_use  = st.session_state.get('rpt_crcp_cd', cd_r)
                crcp_sc_use  = sc_r        # ใช้ร่วมกับ JPCP
                crcp_cbr_use = cbr_r       # ใช้ร่วมกับ JPCP
                crcp_mr_use  = 1500 * crcp_cbr_use if crcp_cbr_use < 10 else 1000 + 555 * crcp_cbr_use

                crcp_inputs  = {
                    **inputs_r,
                    'k_eff': crcp_k_use,
                    'j':     crcp_j_use,
                    'sc':    crcp_sc_use,
                    'cd':    crcp_cd_use,
                    'ls':    inputs_r.get('ls', 1.0),
                }
                crcp_ec      = calculate_concrete_modulus(convert_cube_to_cylinder(fc_r))
                crcp_comp    = []
                for d_cm in thicknesses_cm:
                    d_inch = round(d_cm / 2.54)
                    log_w18, w18_cap = calculate_aashto_rigid_w18(
                        d_inch, dpsi_r, pt_r, zr_r, so_r,
                        crcp_sc_use, crcp_cd_use, crcp_j_use, crcp_ec, crcp_k_use
                    )
                    passed, ratio = check_design(w18_r, w18_cap)
                    crcp_comp.append({
                        'd_cm': d_cm, 'd_inch': d_inch, 'log_w18': log_w18,
                        'w18': w18_cap, 'passed': passed, 'ratio': ratio
                    })
                d_inch_crcp  = round(crcp_d_use / 2.54)
                lw_crcp, w18_crcp = calculate_aashto_rigid_w18(
                    d_inch_crcp, dpsi_r, pt_r, zr_r, so_r,
                    crcp_sc_use, crcp_cd_use, crcp_j_use, crcp_ec, crcp_k_use
                )
                passed_crcp, ratio_crcp = check_design(w18_r, w18_crcp)
                subgrade_crcp = {'cbr': crcp_cbr_use, 'mr_psi': crcp_mr_use, 'mr_mpa': crcp_mr_use / 145.038}

                try:
                    buf, err = create_full_word_report(
                        section_prefix    = st.session_state.get('rpt_prefix', '4.5'),
                        fig_prefix        = st.session_state.get('rpt_fig_prefix', '4-'),
                        fig_start_num     = int(st.session_state.get('rpt_fig_start', 5)),
                        intro_text        = st.session_state.get('rpt_intro', DEFAULT_INTRO_TEXT),
                        summary_text      = st.session_state.get('rpt_summary_text', DEFAULT_SUMMARY_TEXT),
                        project_name      = proj_name_r,
                        pavement_type     = pave_type_r,
                        include_jpcp      = st.session_state.get('rpt_include_jpcp', True),
                        jpcp_layers_data  = layers_r,
                        jpcp_d_cm         = d_r,
                        jpcp_inputs       = inputs_r,
                        jpcp_calc         = calc_r,
                        jpcp_comparison   = comparison_r,
                        jpcp_result       = main_result_r,
                        jpcp_subgrade     = subgrade_r,
                        jpcp_nomo_params  = nomo_r,
                        img1_bytes_jpcp   = st.session_state.get('img1_bytes'),
                        img2_bytes_jpcp   = st.session_state.get('img2_bytes'),
                        include_crcp      = st.session_state.get('rpt_include_crcp', False),
                        crcp_layers_data  = layers_r,
                        crcp_d_cm         = crcp_d_use,
                        crcp_inputs       = crcp_inputs,
                        crcp_calc         = {**calc_r, 'ec': crcp_ec},
                        crcp_comparison   = crcp_comp,
                        crcp_result       = (passed_crcp, ratio_crcp),
                        crcp_subgrade     = subgrade_crcp,
                        crcp_nomo_params  = nomo_r,
                        img1_bytes_crcp   = st.session_state.get('img1_bytes'),
                        img2_bytes_crcp   = st.session_state.get('img2_bytes'),
                        include_summary_section = st.session_state.get('rpt_include_summary', True),
                    )
                    if err:
                        st.error(f"❌ ข้อผิดพลาด: {err}")
                    elif buf:
                        filename = f"AASHTO_Report_{proj_name_r or 'Project'}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                        st.success("✅ สร้างรายงานสำเร็จ!")
                        st.download_button(
                            "⬇️ ดาวน์โหลดรายงาน Word (ฉบับสมบูรณ์)",
                            buf, filename,
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                except Exception as ex:
                    st.error(f"❌ เกิดข้อผิดพลาด: {ex}")
                    import traceback
                    st.code(traceback.format_exc())

    st.markdown("---")
    st.caption("พัฒนาโดย: รศ.ดร.อิทธิพล มีผล // ภาควิชาครุศาสตร์โยธา // มจพ.")

if __name__ == "__main__":
    main()
