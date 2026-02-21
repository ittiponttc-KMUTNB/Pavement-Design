"""
‡πÅ‡∏≠‡∏õ‡∏û‡∏•‡∏¥‡πÄ‡∏Ñ‡∏ä‡∏±‡∏ô‡∏ß‡∏¥‡πÄ‡∏Ñ‡∏£‡∏≤‡∏∞‡∏´‡πå‡∏Ñ‡∏ß‡∏≤‡∏°‡∏Ñ‡∏∏‡πâ‡∏°‡∏Ñ‡πà‡∏≤‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á (AASHTO 1993)
Version 5.0 - Simplified: ‡∏ß‡∏±‡∏™‡∏î‡∏∏‡πÅ‡∏•‡∏∞‡∏£‡∏≤‡∏Ñ‡∏≤ (‡πÑ‡∏°‡πà‡∏°‡∏µ NPV)
‡∏û‡∏±‡∏í‡∏ô‡∏≤‡πÇ‡∏î‡∏¢: Claude AI ‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö ‡∏≠.‡∏≠‡∏¥‡∏ó‡∏ò‡∏¥‡∏û‡∏• - KMUTNB
"""

import streamlit as st
import pandas as pd
import numpy as np
import json
from datetime import datetime
import io

# Import with error handling
try:
    import plotly.graph_objects as go
    from plotly.subplots import make_subplots
    PLOTLY_AVAILABLE = True
except ImportError:
    PLOTLY_AVAILABLE = False
    st.warning("‚ö†Ô∏è Plotly ‡πÑ‡∏°‡πà‡∏™‡∏≤‡∏°‡∏≤‡∏£‡∏ñ‡πÉ‡∏ä‡πâ‡∏á‡∏≤‡∏ô‡πÑ‡∏î‡πâ ‡∏Å‡∏£‡∏≤‡∏ü‡∏ö‡∏≤‡∏á‡∏™‡πà‡∏ß‡∏ô‡∏≠‡∏≤‡∏à‡πÑ‡∏°‡πà‡πÅ‡∏™‡∏î‡∏á")

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    st.warning("‚ö†Ô∏è python-docx ‡πÑ‡∏°‡πà‡∏™‡∏≤‡∏°‡∏≤‡∏£‡∏ñ‡πÉ‡∏ä‡πâ‡∏á‡∏≤‡∏ô‡πÑ‡∏î‡πâ ‡∏Å‡∏≤‡∏£‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô Word ‡∏≠‡∏≤‡∏à‡πÑ‡∏°‡πà‡∏ó‡∏≥‡∏á‡∏≤‡∏ô")

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    st.warning("‚ö†Ô∏è openpyxl ‡πÑ‡∏°‡πà‡∏™‡∏≤‡∏°‡∏≤‡∏£‡∏ñ‡πÉ‡∏ä‡πâ‡∏á‡∏≤‡∏ô‡πÑ‡∏î‡πâ ‡∏Å‡∏≤‡∏£ Upload/Download Excel ‡∏≠‡∏≤‡∏à‡πÑ‡∏°‡πà‡∏ó‡∏≥‡∏á‡∏≤‡∏ô")

# ‡∏ï‡∏±‡πâ‡∏á‡∏Ñ‡πà‡∏≤‡∏´‡∏ô‡πâ‡∏≤‡πÄ‡∏ß‡πá‡∏ö
st.set_page_config(
    page_title="‡∏ß‡∏¥‡πÄ‡∏Ñ‡∏£‡∏≤‡∏∞‡∏´‡πå‡∏Ñ‡πà‡∏≤‡∏Å‡πà‡∏≠‡∏™‡∏£‡πâ‡∏≤‡∏á‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á",
    page_icon="üõ£Ô∏è",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS
st.markdown("""
<style>
    .main-header {
        font-size: 2rem;
        font-weight: bold;
        color: #1E3A5F;
        text-align: center;
        padding: 1rem;
        background: linear-gradient(90deg, #E8F4FD, #D1E9FA);
        border-radius: 10px;
        margin-bottom: 2rem;
    }
    .cost-box {
        background: #f0f8ff;
        padding: 10px;
        border-radius: 8px;
        border-left: 4px solid #2E86AB;
        margin: 10px 0;
    }
</style>
""", unsafe_allow_html=True)


# ===== Library ‡∏£‡∏≤‡∏Ñ‡∏≤‡∏ß‡∏±‡∏™‡∏î‡∏∏ (Price Library) =====
# ‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏à‡∏≤‡∏Å‡πÑ‡∏ü‡∏•‡πå ‡∏£‡∏≤‡∏Ñ‡∏≤‡πÄ‡∏õ‡∏£‡∏µ‡∏¢‡∏ö‡πÄ‡∏ó‡∏µ‡∏¢‡∏ö‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á

# ‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏£‡∏≤‡∏Ñ‡∏≤‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á AC (‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.) ‡∏ï‡∏≤‡∏°‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤
AC_PRICE_TABLE = {
    'PMA Wearing Course': {
        2.5: 170, 3: 203, 4: 268, 5: 333, 6: 406, 7: 471, 8: 536, 9: 601, 10: 667
    },
    'AC Wearing Course': {
        2.5: 128, 3: 152, 4: 202, 5: 250, 6: 306, 7: 355, 8: 403, 9: 452, 10: 502
    },
    'AC Binder Course': {
        2.5: 129, 3: 154, 4: 202, 5: 251, 6: 308, 7: 356, 8: 405, 9: 454, 10: 503
    },
    'AC Base Course': {
        2.5: 129, 3: 154, 4: 202, 5: 251, 6: 308, 7: 356, 8: 405, 9: 454, 10: 503
    },
}

# ‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏£‡∏≤‡∏Ñ‡∏≤‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï (‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.) ‡∏ï‡∏≤‡∏°‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤
CONCRETE_PRICE_TABLE = {
    'JRCP': {25: 924, 28: 1002, 32: 1106, 35: 1184},
    'JPCP': {25: 928, 28: 1000, 32: 1095, 35: 1167},
    'CRCP': {25: 1245, 28: 1358, 32: 1509, 35: 1622},
}

# ‡∏£‡∏≤‡∏Ñ‡∏≤‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï (‡πÑ‡∏°‡πà‡∏£‡∏ß‡∏° Joint)
CONCRETE_EXCL_JOINT = {
    'JRCP': 830,
    'JPCP': 764,
    'CRCP': 1204,
}

# ‡∏£‡∏≤‡∏Ñ‡∏≤‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á/‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á (‡∏ö‡∏≤‡∏ó/‡∏•‡∏ö.‡∏°.)
BASE_MATERIAL_PRICES = {
    'Crushed Rock Base Course': 583,
    'Cement Modified Crushed Rock Base (UCS 24.5 ksc)': 864,
    'Cement Treated Base (UCS 40 ksc)': 1096,
    'Soil Aggregate Subbase': 375,
    'Soil Cement Subbase (UCS 7 ksc)': 854,
    'Selected Material A': 375,
}

# Library ‡∏ß‡∏±‡∏™‡∏î‡∏∏ (‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö UI)
MATERIAL_LIBRARY = {
    '‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á': {
        '‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡∏•‡∏≤‡∏î‡∏¢‡∏≤‡∏á AC': {'unit_cost': 480, 'cost_unit': '‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.'},
        '‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡∏•‡∏≤‡∏î‡∏¢‡∏≤‡∏á PMA': {'unit_cost': 550, 'cost_unit': '‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.'},
        '‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï 350 Ksc.': {'unit_cost': 800, 'cost_unit': '‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.'},
        '‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï 350 Ksc.': {'unit_cost': 850, 'cost_unit': '‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.'},
    },
    '‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á': {
        'Crushed Rock Base Course': {'unit_cost': 583, 'cost_unit': '‡∏ö‡∏≤‡∏ó/‡∏•‡∏ö.‡∏°.'},
        'Cement Modified Crushed Rock Base (UCS 24.5 ksc)': {'unit_cost': 864, 'cost_unit': '‡∏ö‡∏≤‡∏ó/‡∏•‡∏ö.‡∏°.'},
        'Cement Treated Base (UCS 40 ksc)': {'unit_cost': 1096, 'cost_unit': '‡∏ö‡∏≤‡∏ó/‡∏•‡∏ö.‡∏°.'},
        'Soil Cement Subbase (UCS 7 ksc)': {'unit_cost': 854, 'cost_unit': '‡∏ö‡∏≤‡∏ó/‡∏•‡∏ö.‡∏°.'},
    },
    '‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á': {
        'Soil Aggregate Subbase': {'unit_cost': 375, 'cost_unit': '‡∏ö‡∏≤‡∏ó/‡∏•‡∏ö.‡∏°.'},
        'Selected Material A': {'unit_cost': 375, 'cost_unit': '‡∏ö‡∏≤‡∏ó/‡∏•‡∏ö.‡∏°.'},
    },
    '‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏≠‡∏∑‡πà‡∏ô‡πÜ': {
        'Tack Coat': {'unit_cost': 20, 'cost_unit': '‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.'},
        'Prime Coat': {'unit_cost': 30, 'cost_unit': '‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.'},
        'Non Woven Geotextile': {'unit_cost': 78, 'cost_unit': '‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.'},
    },
}

# ===== ‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡πÄ‡∏£‡∏¥‡πà‡∏°‡∏ï‡πâ‡∏ô‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á =====

def _parse_json_details_to_layers(details):
    """‡πÅ‡∏õ‡∏•‡∏á JSON details ‚Üí (layers, joints) format ‡∏ó‡∏µ‡πà app ‡πÉ‡∏ä‡πâ‡∏†‡∏≤‡∏¢‡πÉ‡∏ô"""
    layers, joints = [], []
    # ‡∏ä‡∏∑‡πà‡∏≠‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á/‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏ó‡∏µ‡πà‡∏£‡∏≤‡∏Ñ‡∏≤‡πÄ‡∏õ‡πá‡∏ô ‡∏ö‡∏≤‡∏ó/‡∏•‡∏ö.‡∏°.
    BASE_KEYWORDS = ['crushed rock', 'soil aggregate', 'soil cement', 'cement modified',
                     'cement treated', 'selected material', 'sand embankment']
    for item in details:
        name = item.get('‡∏£‡∏≤‡∏¢‡∏Å‡∏≤‡∏£', '')
        unit_raw = item.get('‡∏´‡∏ô‡πà‡∏ß‡∏¢', '‡∏ï‡∏£.‡∏°.')
        qty = item.get('‡∏õ‡∏£‡∏¥‡∏°‡∏≤‡∏ì', 22000)
        unit_cost = item.get('‡∏£‡∏≤‡∏Ñ‡∏≤/‡∏´‡∏ô‡πà‡∏ß‡∏¢', 0)
        if 'Joint' in name or unit_raw == 'm':
            joints.append({'name': name, 'quantity': qty, 'qty_unit': 'm', 'unit_cost': unit_cost})
            continue
        thick_str = str(item.get('‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤', '1'))
        try:
            parts = thick_str.split()
            thick_val = float(parts[0])
            unit_val = parts[1] if len(parts) > 1 else 'cm'
        except:
            thick_val = 1.0
            unit_val = 'cm'
        # ‡∏Å‡∏≥‡∏´‡∏ô‡∏î qty_unit ‡∏ï‡∏≤‡∏°‡∏ä‡∏ô‡∏¥‡∏î‡∏ß‡∏±‡∏™‡∏î‡∏∏ ‡πÑ‡∏°‡πà‡πÉ‡∏ä‡πà‡∏ï‡∏≤‡∏°‡∏´‡∏ô‡πà‡∏ß‡∏¢‡πÉ‡∏ô JSON
        # (JSON ‡∏ö‡∏±‡∏ô‡∏ó‡∏∂‡∏Å‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡πÄ‡∏õ‡πá‡∏ô ‡∏ï‡∏£.‡∏°. ‡πÅ‡∏ï‡πà app ‡πÉ‡∏ä‡πâ sq.m ‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡∏ó‡∏∏‡∏Å‡∏≠‡∏¢‡πà‡∏≤‡∏á)
        name_lower = name.lower()
        is_base_material = any(kw in name_lower for kw in BASE_KEYWORDS)
        qty_unit = 'cu.m' if is_base_material else 'sq.m'
        layers.append({
            'name': name, 'thickness': thick_val, 'unit': unit_val,
            'quantity': qty, 'qty_unit': qty_unit, 'unit_cost': unit_cost,
        })
    return layers, joints


def get_default_ac1_layers():
    """AC1: ‡πÅ‡∏≠‡∏™‡∏ü‡∏±‡∏•‡∏ï‡πå‡∏ö‡∏ô‡∏´‡∏¥‡∏ô‡∏Ñ‡∏•‡∏∏‡∏Å (‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏ó‡∏µ‡πà 5.3-18)"""
    _d = st.session_state.get('loaded_project', {}).get('construction', {}).get('AC1', {})
    if _d.get('details'):
        layers, _ = _parse_json_details_to_layers(_d['details'])
        if layers: return layers
    return [
        {'name': 'Wearing Course', 'thickness': 7, 'unit': 'cm', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 480},
        {'name': 'Binder Course', 'thickness': 7, 'unit': 'cm', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 480},
        {'name': 'Asphalt Base Course', 'thickness': 10, 'unit': 'cm', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 600},
        {'name': 'Tack Coat', 'thickness': 2, 'unit': 'Layer', 'quantity': 44000, 'qty_unit': 'sq.m', 'unit_cost': 20},
        {'name': 'Prime Coat', 'thickness': 1, 'unit': 'Layer', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 30},
        {'name': 'Crushed Rock Base', 'thickness': 20, 'unit': 'cm', 'quantity': 4400, 'qty_unit': 'cu.m', 'unit_cost': 714},
        {'name': 'Soil Aggregate Subbase', 'thickness': 30, 'unit': 'cm', 'quantity': 6600, 'qty_unit': 'cu.m', 'unit_cost': 714},
        {'name': 'Sand Embankment', 'thickness': 40, 'unit': 'cm', 'quantity': 8800, 'qty_unit': 'cu.m', 'unit_cost': 361},
    ]

def get_default_ac2_layers():
    """AC2: ‡πÅ‡∏≠‡∏™‡∏ü‡∏±‡∏•‡∏ï‡πå‡∏ö‡∏ô‡∏´‡∏¥‡∏ô‡∏Ñ‡∏•‡∏∏‡∏Å‡∏ú‡∏™‡∏°‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå (‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏ó‡∏µ‡πà 5.3-20)"""
    _d = st.session_state.get('loaded_project', {}).get('construction', {}).get('AC2', {})
    if _d.get('details'):
        layers, _ = _parse_json_details_to_layers(_d['details'])
        if layers: return layers
    return [
        {'name': 'Wearing Course', 'thickness': 5, 'unit': 'cm', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 400},
        {'name': 'Binder Course', 'thickness': 5, 'unit': 'cm', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 400},
        {'name': 'Tack Coat', 'thickness': 1, 'unit': 'Layer', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 20},
        {'name': 'Prime Coat', 'thickness': 1, 'unit': 'Layer', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 30},
        {'name': 'Cement Modified Crushed Rock', 'thickness': 20, 'unit': 'cm', 'quantity': 4400, 'qty_unit': 'cu.m', 'unit_cost': 914},
        {'name': 'Soil Aggregate Subbase', 'thickness': 20, 'unit': 'cm', 'quantity': 4400, 'qty_unit': 'cu.m', 'unit_cost': 714},
        {'name': 'Sand Embankment', 'thickness': 30, 'unit': 'cm', 'quantity': 6600, 'qty_unit': 'cu.m', 'unit_cost': 361},
    ]

def get_default_jrcp1_layers():
    """JPCP/JRCP (1): ‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï‡∏ö‡∏ô‡∏î‡∏¥‡∏ô‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå (‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏ó‡∏µ‡πà 5.3-22)"""
    _d = st.session_state.get('loaded_project', {}).get('construction', {}).get('JRCP1', {})
    if _d.get('details'):
        layers, _ = _parse_json_details_to_layers(_d['details'])
        if layers: return layers
    return [
        {'name': '350 Ksc. Cubic Type Concrete', 'thickness': 28, 'unit': 'cm', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 800},
        {'name': 'Non Woven Geotextile', 'thickness': 1, 'unit': '‡∏ä‡∏±‡πâ‡∏ô', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 78},
        {'name': 'Soil Cement Base', 'thickness': 20, 'unit': 'cm', 'quantity': 4400, 'qty_unit': 'cu.m', 'unit_cost': 621},
        {'name': 'Sand Embankment', 'thickness': 60, 'unit': 'cm', 'quantity': 13200, 'qty_unit': 'cu.m', 'unit_cost': 361},
    ]

def get_default_jrcp1_joints():
    """‡∏£‡∏≠‡∏¢‡∏ï‡πà‡∏≠‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö JRCP1 - ‡∏õ‡∏£‡∏¥‡∏°‡∏≤‡∏ì‡∏ï‡πà‡∏≠ 1 ‡∏Å‡∏°."""
    _d = st.session_state.get('loaded_project', {}).get('construction', {}).get('JRCP1', {})
    if _d.get('details'):
        _, joints = _parse_json_details_to_layers(_d['details'])
        if joints: return joints
    return [
        {'name': 'Transverse Joint @10m', 'quantity': 2200, 'qty_unit': 'm', 'unit_cost': 430},
        {'name': 'Longitudinal Joint', 'quantity': 4000, 'qty_unit': 'm', 'unit_cost': 120},
    ]

def get_default_jrcp2_layers():
    """JPCP/JRCP (2): ‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï‡∏ö‡∏ô‡∏´‡∏¥‡∏ô‡∏Ñ‡∏•‡∏∏‡∏Å‡∏ú‡∏™‡∏°‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå (‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏ó‡∏µ‡πà 5.3-24)"""
    _d = st.session_state.get('loaded_project', {}).get('construction', {}).get('JRCP2', {})
    if _d.get('details'):
        layers, _ = _parse_json_details_to_layers(_d['details'])
        if layers: return layers
    return [
        {'name': '350 Ksc. Cubic Type Concrete', 'thickness': 28, 'unit': 'cm', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 800},
        {'name': 'Non Woven Geotextile', 'thickness': 1, 'unit': '‡∏ä‡∏±‡πâ‡∏ô', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 78},
        {'name': 'Cement Modified Crushed Rock', 'thickness': 20, 'unit': 'cm', 'quantity': 4400, 'qty_unit': 'cu.m', 'unit_cost': 914},
        {'name': 'Sand Embankment', 'thickness': 50, 'unit': 'cm', 'quantity': 11000, 'qty_unit': 'cu.m', 'unit_cost': 361},
    ]

def get_default_jrcp2_joints():
    """‡∏£‡∏≠‡∏¢‡∏ï‡πà‡∏≠‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö JRCP2 - ‡∏õ‡∏£‡∏¥‡∏°‡∏≤‡∏ì‡∏ï‡πà‡∏≠ 1 ‡∏Å‡∏°."""
    _d = st.session_state.get('loaded_project', {}).get('construction', {}).get('JRCP2', {})
    if _d.get('details'):
        _, joints = _parse_json_details_to_layers(_d['details'])
        if joints: return joints
    return [
        {'name': 'Transverse Joint @10m', 'quantity': 2200, 'qty_unit': 'm', 'unit_cost': 430},
        {'name': 'Longitudinal Joint', 'quantity': 4000, 'qty_unit': 'm', 'unit_cost': 120},
    ]

def get_default_crcp1_layers():
    """CRCP1: ‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï‡πÄ‡∏™‡∏£‡∏¥‡∏°‡πÄ‡∏´‡∏•‡πá‡∏Å‡∏ï‡πà‡∏≠‡πÄ‡∏ô‡∏∑‡πà‡∏≠‡∏á‡∏ö‡∏ô‡∏î‡∏¥‡∏ô‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå"""
    _d = st.session_state.get('loaded_project', {}).get('construction', {}).get('CRCP1', {})
    if _d.get('details'):
        layers, _ = _parse_json_details_to_layers(_d['details'])
        if layers: return layers
    return [
        {'name': '350 Ksc. Cubic Type Concrete', 'thickness': 25, 'unit': 'cm', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 850},
        {'name': 'Steel Reinforcement', 'thickness': 1, 'unit': '‡∏ä‡∏±‡πâ‡∏ô', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 150},
        {'name': 'Non Woven Geotextile', 'thickness': 1, 'unit': '‡∏ä‡∏±‡πâ‡∏ô', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 78},
        {'name': 'Soil Cement Base', 'thickness': 15, 'unit': 'cm', 'quantity': 3300, 'qty_unit': 'cu.m', 'unit_cost': 621},
        {'name': 'Sand Embankment', 'thickness': 50, 'unit': 'cm', 'quantity': 11000, 'qty_unit': 'cu.m', 'unit_cost': 361},
    ]

def get_default_crcp2_layers():
    """CRCP2: ‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï‡πÄ‡∏™‡∏£‡∏¥‡∏°‡πÄ‡∏´‡∏•‡πá‡∏Å‡∏ï‡πà‡∏≠‡πÄ‡∏ô‡∏∑‡πà‡∏≠‡∏á‡∏ö‡∏ô‡∏´‡∏¥‡∏ô‡∏Ñ‡∏•‡∏∏‡∏Å‡∏ú‡∏™‡∏°‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå"""
    _d = st.session_state.get('loaded_project', {}).get('construction', {}).get('CRCP2', {})
    if _d.get('details'):
        layers, _ = _parse_json_details_to_layers(_d['details'])
        if layers: return layers
    return [
        {'name': '350 Ksc. Cubic Type Concrete', 'thickness': 25, 'unit': 'cm', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 850},
        {'name': 'Steel Reinforcement', 'thickness': 1, 'unit': '‡∏ä‡∏±‡πâ‡∏ô', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 150},
        {'name': 'Non Woven Geotextile', 'thickness': 1, 'unit': '‡∏ä‡∏±‡πâ‡∏ô', 'quantity': 22000, 'qty_unit': 'sq.m', 'unit_cost': 78},
        {'name': 'Cement Modified Crushed Rock', 'thickness': 15, 'unit': 'cm', 'quantity': 3300, 'qty_unit': 'cu.m', 'unit_cost': 914},
        {'name': 'Sand Embankment', 'thickness': 40, 'unit': 'cm', 'quantity': 8800, 'qty_unit': 'cu.m', 'unit_cost': 361},
    ]


def calculate_quantity(thickness_cm, width_m, length_km, qty_unit):
    """‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡∏õ‡∏£‡∏¥‡∏°‡∏≤‡∏ì‡∏à‡∏≤‡∏Å‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤ ‡∏Ñ‡∏ß‡∏≤‡∏°‡∏Å‡∏ß‡πâ‡∏≤‡∏á ‡πÅ‡∏•‡∏∞‡∏Ñ‡∏ß‡∏≤‡∏°‡∏¢‡∏≤‡∏ß"""
    area = width_m * length_km * 1000  # ‡∏ï‡∏£.‡∏°.
    if qty_unit == 'sq.m':
        return area
    elif qty_unit == 'cu.m':
        return area * thickness_cm / 100  # ‡∏•‡∏ö.‡∏°.
    return area


def calculate_layer_cost(layers, road_length_km=1.0):
    """‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡∏Ñ‡πà‡∏≤‡∏Å‡πà‡∏≠‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏à‡∏≤‡∏Å‡∏ä‡∏±‡πâ‡∏ô‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á
    ‡∏£‡∏≤‡∏Ñ‡∏≤‡∏ó‡∏±‡πâ‡∏á‡∏´‡∏°‡∏î‡πÄ‡∏õ‡πá‡∏ô ‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°. √ó ‡∏õ‡∏£‡∏¥‡∏°‡∏≤‡∏ì (‡∏ï‡∏£.‡∏°.)
    """
    total = 0
    details = []
    
    for layer in layers:
        # ‡∏õ‡∏£‡∏¥‡∏°‡∏≤‡∏ì‡πÄ‡∏õ‡πá‡∏ô ‡∏ï‡∏£.‡∏°. ‡πÅ‡∏•‡πâ‡∏ß (‡πÑ‡∏°‡πà‡∏ï‡πâ‡∏≠‡∏á‡∏Ñ‡∏π‡∏ì road_length ‡∏≠‡∏µ‡∏Å ‡πÄ‡∏û‡∏£‡∏≤‡∏∞‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡πÑ‡∏ß‡πâ‡πÅ‡∏•‡πâ‡∏ß)
        qty = layer['quantity']
        # ‡∏£‡∏≤‡∏Ñ‡∏≤‡πÄ‡∏õ‡πá‡∏ô ‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.
        cost = qty * layer['unit_cost']
        total += cost
        
        details.append({
            '‡∏£‡∏≤‡∏¢‡∏Å‡∏≤‡∏£': layer['name'],
            '‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤': f"{layer['thickness']} {layer['unit']}",
            '‡∏õ‡∏£‡∏¥‡∏°‡∏≤‡∏ì': qty,
            '‡∏´‡∏ô‡πà‡∏ß‡∏¢': '‡∏ï‡∏£.‡∏°.',
            '‡∏£‡∏≤‡∏Ñ‡∏≤/‡∏´‡∏ô‡πà‡∏ß‡∏¢': layer['unit_cost'],
            '‡∏°‡∏π‡∏•‡∏Ñ‡πà‡∏≤ (‡∏ö‡∏≤‡∏ó)': cost
        })
    
    return total, details


def calculate_joint_cost(joints, road_length_km=1.0):
    """‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡∏Ñ‡πà‡∏≤‡∏£‡∏≠‡∏¢‡∏ï‡πà‡∏≠"""
    total = 0
    details = []
    
    for joint in joints:
        qty = joint['quantity'] * road_length_km
        cost = qty * joint['unit_cost']
        total += cost
        
        details.append({
            '‡∏£‡∏≤‡∏¢‡∏Å‡∏≤‡∏£': joint['name'],
            '‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤': '-',
            '‡∏õ‡∏£‡∏¥‡∏°‡∏≤‡∏ì': qty,
            '‡∏´‡∏ô‡πà‡∏ß‡∏¢': joint['qty_unit'],
            '‡∏£‡∏≤‡∏Ñ‡∏≤/‡∏´‡∏ô‡πà‡∏ß‡∏¢': joint['unit_cost'],
            '‡∏°‡∏π‡∏•‡∏Ñ‡πà‡∏≤ (‡∏ö‡∏≤‡∏ó)': cost
        })
    
    return total, details



def get_price_from_library(layer_name, thickness):
    """‡∏î‡∏∂‡∏á‡∏£‡∏≤‡∏Ñ‡∏≤‡∏à‡∏≤‡∏Å Library ‡∏ï‡∏≤‡∏°‡∏ä‡∏∑‡πà‡∏≠‡πÅ‡∏•‡∏∞‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤"""
    if 'price_library' not in st.session_state:
        return None
    
    lib = st.session_state['price_library']
    name_lower = layer_name.lower()
    
    # AC Prices
    if 'pma' in name_lower and 'wearing' in name_lower:
        return lib['ac_prices'].get('PMA Wearing Course', {}).get(thickness)
    elif 'wearing' in name_lower:
        return lib['ac_prices'].get('AC Wearing Course', {}).get(thickness)
    elif 'binder' in name_lower:
        return lib['ac_prices'].get('AC Binder Course', {}).get(thickness)
    elif 'asphalt' in name_lower and 'base' in name_lower:
        return lib['ac_prices'].get('AC Base Course', {}).get(thickness)
    
    # Concrete Prices
    elif 'jrcp' in name_lower or ('concrete' in name_lower and 'jrcp' in str(thickness)):
        return lib['concrete_prices'].get('JRCP', {}).get(int(thickness))
    elif 'jpcp' in name_lower:
        return lib['concrete_prices'].get('JPCP', {}).get(int(thickness))
    elif 'crcp' in name_lower:
        return lib['concrete_prices'].get('CRCP', {}).get(int(thickness))
    
    # Base Material Prices
    elif 'crushed rock' in name_lower and 'cement' not in name_lower:
        return lib['base_prices'].get('Crushed Rock Base Course')
    elif 'cement modified' in name_lower or 'cmcr' in name_lower:
        return lib['base_prices'].get('Cement Modified Crushed Rock Base (UCS 24.5 ksc)')
    elif 'cement treated' in name_lower or 'ctb' in name_lower:
        return lib['base_prices'].get('Cement Treated Base (UCS 40 ksc)')
    elif 'soil aggregate' in name_lower:
        return lib['base_prices'].get('Soil Aggregate Subbase')
    elif 'soil cement' in name_lower:
        return lib['base_prices'].get('Soil Cement Subbase (UCS 7 ksc)')
    elif 'selected' in name_lower:
        return lib['base_prices'].get('Selected Material A')
    
    return None


def render_layer_editor(layers, key_prefix, total_width, road_length, v=0):
    """‡πÅ‡∏™‡∏î‡∏á UI ‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡πÅ‡∏Å‡πâ‡πÑ‡∏Ç‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á ‡∏û‡∏£‡πâ‡∏≠‡∏°‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡∏õ‡∏£‡∏¥‡∏°‡∏≤‡∏ì‡∏≠‡∏±‡∏ï‡πÇ‡∏ô‡∏°‡∏±‡∏ï‡∏¥
    ‡∏£‡∏≤‡∏Ñ‡∏≤‡∏ó‡∏±‡πâ‡∏á‡∏´‡∏°‡∏î‡πÅ‡∏™‡∏î‡∏á‡πÄ‡∏õ‡πá‡∏ô ‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.
    v = json_version ‡πÄ‡∏û‡∏∑‡πà‡∏≠ force refresh ‡πÄ‡∏°‡∏∑‡πà‡∏≠ load JSON ‡πÉ‡∏´‡∏°‡πà
    """
    updated_layers = []
    
    # ‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏µ‡πà‡∏ï‡πà‡∏≠ ‡∏Å‡∏°.
    # total_width ‡∏£‡∏ß‡∏°‡∏ó‡∏±‡πâ‡∏á 2 ‡∏ó‡∏¥‡∏®‡∏ó‡∏≤‡∏á‡πÑ‡∏ß‡πâ‡πÅ‡∏•‡πâ‡∏ß (num_lanes = lanes_per_direction * 2)
    area_per_km = total_width * 1000  # ‡∏ï‡∏£.‡∏°./‡∏Å‡∏°.
    
    # ‡πÅ‡∏¢‡∏Å layers ‡πÄ‡∏õ‡πá‡∏ô‡∏Å‡∏•‡∏∏‡πà‡∏°
    surface_layers = []
    base_layers = []
    
    for layer in layers:
        name_lower = layer['name'].lower()
        if any(x in name_lower for x in [
            'wearing', 'binder', 'asphalt', 'concrete', 'tack', 'prime',
            'geotextile', 'steel',
            'ac base', 'ac interlayer', 'interlayer',   # ‡∏ä‡∏∑‡πà‡∏≠‡∏à‡∏≤‡∏Å JSON save
            'ac wearing', 'ac binder',
        ]):
            surface_layers.append(layer)
        else:
            base_layers.append(layer)
    
    # ===== ‡∏™‡πà‡∏ß‡∏ô‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á =====
    st.markdown("**‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á** (‡∏´‡∏ô‡πà‡∏ß‡∏¢: ‡∏ï‡∏£.‡∏°.)")
    cols = st.columns([3, 1, 1.5])
    cols[0].markdown("‡∏£‡∏≤‡∏¢‡∏Å‡∏≤‡∏£")
    cols[1].markdown("‡∏´‡∏ô‡∏≤ (cm)")
    cols[2].markdown("‡∏£‡∏≤‡∏Ñ‡∏≤ (‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.)")
    
    # ‡∏ï‡∏±‡∏ß‡πÄ‡∏•‡∏∑‡∏≠‡∏Å‡∏ß‡∏±‡∏™‡∏î‡∏∏
    wearing_options = ['AC Wearing Course', 'PMA Wearing Course']
    binder_options = ['AC Binder Course']
    base_options = ['AC Base Course']
    concrete_options = ['JPCP', 'JRCP', 'CRCP']
    
    for i, layer in enumerate(surface_layers):
        cols = st.columns([3, 1, 1.5])
        name_lower = layer['name'].lower()
        
        # ‡∏Å‡∏≥‡∏´‡∏ô‡∏î‡∏ß‡πà‡∏≤‡πÄ‡∏õ‡πá‡∏ô‡∏ä‡∏±‡πâ‡∏ô‡πÑ‡∏´‡∏ô
        is_wearing = 'wearing' in name_lower
        is_binder = 'binder' in name_lower
        is_ac_base = ('asphalt' in name_lower and 'base' in name_lower) or \
                     ('ac base' in name_lower) or \
                     ('interlayer' in name_lower)
        is_concrete = 'concrete' in name_lower or 'ksc' in name_lower
        
        with cols[0]:
            if is_wearing:
                # Dropdown ‡πÄ‡∏•‡∏∑‡∏≠‡∏Å PMA ‡∏´‡∏£‡∏∑‡∏≠ AC Wearing
                default_idx = 1 if 'pma' in name_lower else 0
                selected_material = st.selectbox(
                    "‡∏ß‡∏±‡∏™‡∏î‡∏∏", wearing_options, index=default_idx,
                    key=f"{key_prefix}_mat_{i}_v{v}", label_visibility="collapsed"
                )
            elif is_binder:
                selected_material = st.selectbox(
                    "‡∏ß‡∏±‡∏™‡∏î‡∏∏", binder_options, index=0,
                    key=f"{key_prefix}_mat_{i}_v{v}", label_visibility="collapsed"
                )
            elif is_ac_base:
                selected_material = st.selectbox(
                    "‡∏ß‡∏±‡∏™‡∏î‡∏∏", base_options, index=0,
                    key=f"{key_prefix}_mat_{i}_v{v}", label_visibility="collapsed"
                )
            elif is_concrete:
                # Dropdown ‡πÄ‡∏•‡∏∑‡∏≠‡∏Å JPCP, JRCP, CRCP
                # ‡∏≠‡πà‡∏≤‡∏ô type ‡∏à‡∏≤‡∏Å‡∏ä‡∏∑‡πà‡∏≠ layer ‡∏ó‡∏µ‡πà‡∏°‡∏≤‡∏à‡∏≤‡∏Å JSON ‡∏Å‡πà‡∏≠‡∏ô ‡πÄ‡∏ä‡πà‡∏ô "350 Ksc. Cubic Type Concrete (JPCP)"
                name_upper = layer['name'].upper()
                if 'JPCP' in name_upper:
                    default_idx = 0
                elif 'JRCP' in name_upper:
                    default_idx = 1
                elif 'CRCP' in name_upper:
                    default_idx = 2
                elif 'jrcp' in key_prefix:
                    default_idx = 1
                elif 'crcp' in key_prefix:
                    default_idx = 2
                else:
                    default_idx = 0  # JPCP
                selected_type = st.selectbox(
                    "‡∏ä‡∏ô‡∏¥‡∏î", concrete_options, index=default_idx,
                    key=f"{key_prefix}_ctype_{i}_v{v}", label_visibility="collapsed"
                )
                selected_material = f"350 Ksc. Cubic Type Concrete ({selected_type})"
            else:
                st.text(layer['name'])
                selected_material = layer['name']
        
        with cols[1]:
            thick = st.number_input("‡∏´‡∏ô‡∏≤", value=float(layer['thickness']),
                key=f"{key_prefix}_st_{i}_v{v}", label_visibility="collapsed", min_value=0.0, step=1.0)
        
        # ‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡∏õ‡∏£‡∏¥‡∏°‡∏≤‡∏ì‡∏≠‡∏±‡∏ï‡πÇ‡∏ô‡∏°‡∏±‡∏ï‡∏¥ (‡∏ï‡∏£.‡∏°.)
        auto_qty = area_per_km * road_length
        
        # ‡∏î‡∏∂‡∏á‡∏£‡∏≤‡∏Ñ‡∏≤‡∏à‡∏≤‡∏Å Library (‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.) ‡∏ï‡∏≤‡∏°‡∏ß‡∏±‡∏™‡∏î‡∏∏‡πÅ‡∏•‡∏∞‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡∏ó‡∏µ‡πà‡πÄ‡∏•‡∏∑‡∏≠‡∏Å
        lib_price = None
        if 'price_library' in st.session_state:
            lib = st.session_state['price_library']
            
            if is_wearing:
                prices = lib['ac_prices'].get(selected_material, {})
                lib_price = prices.get(thick)
                if lib_price is None and prices:
                    closest = min(prices.keys(), key=lambda x: abs(x - thick))
                    lib_price = prices.get(closest)
            elif is_binder:
                prices = lib['ac_prices'].get('AC Binder Course', {})
                lib_price = prices.get(thick)
                if lib_price is None and prices:
                    closest = min(prices.keys(), key=lambda x: abs(x - thick))
                    lib_price = prices.get(closest)
            elif is_ac_base:
                prices = lib['ac_prices'].get('AC Base Course', {})
                lib_price = prices.get(thick)
                if lib_price is None and prices:
                    closest = min(prices.keys(), key=lambda x: abs(x - thick))
                    lib_price = prices.get(closest)
            elif is_concrete:
                # ‡∏î‡∏∂‡∏á‡∏£‡∏≤‡∏Ñ‡∏≤‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï‡∏à‡∏≤‡∏Å Library
                concrete_type = selected_type if 'selected_type' in dir() else 'JPCP'
                prices = lib['concrete_prices'].get(concrete_type, {})
                lib_price = prices.get(int(thick))
                if lib_price is None and prices:
                    closest = min(prices.keys(), key=lambda x: abs(x - thick))
                    lib_price = prices.get(closest)
        
        # ‡πÉ‡∏ä‡πâ‡∏£‡∏≤‡∏Ñ‡∏≤‡∏à‡∏≤‡∏Å Library ‡∏´‡∏£‡∏∑‡∏≠‡∏Ñ‡πà‡∏≤ default
        default_cost = lib_price if lib_price else layer['unit_cost']
        
        with cols[2]:
            st.markdown(f"**{default_cost:,.2f}**")
        
        # ‡πÄ‡∏Å‡πá‡∏ö‡∏ä‡∏∑‡πà‡∏≠‡∏ó‡∏µ‡πà‡∏ñ‡∏π‡∏Å‡∏ï‡πâ‡∏≠‡∏á
        if is_concrete:
            final_name = selected_material
        elif is_wearing or is_binder or is_ac_base:
            final_name = selected_material
        else:
            final_name = layer['name']
        
        updated_layers.append({
            'name': final_name, 'thickness': thick, 'unit': layer['unit'],
            'quantity': auto_qty, 'qty_unit': 'sq.m', 'unit_cost': default_cost,
            'cost_per_sqm': default_cost
        })
    
    # ===== ‡∏™‡πà‡∏ß‡∏ô‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á/‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á =====
    st.markdown("---")
    st.markdown("**‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á/‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á** (‡∏£‡∏≤‡∏Ñ‡∏≤‡πÅ‡∏™‡∏î‡∏á‡πÄ‡∏õ‡πá‡∏ô ‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.)")
    
    # ‡∏ï‡∏£‡∏ß‡∏à‡∏™‡∏≠‡∏ö‡∏ß‡πà‡∏≤‡πÄ‡∏õ‡πá‡∏ô JRCP ‡∏´‡∏£‡∏∑‡∏≠ CRCP ‡∏´‡∏£‡∏∑‡∏≠‡πÑ‡∏°‡πà (‡πÄ‡∏û‡∏∑‡πà‡∏≠‡πÄ‡∏û‡∏¥‡πà‡∏° AC Interlayer)
    is_concrete_pavement = any(x in key_prefix.lower() for x in ['jrcp', 'crcp'])
    
    # Library ‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á (‡∏î‡∏∂‡∏á‡∏à‡∏≤‡∏Å session_state ‡∏´‡∏£‡∏∑‡∏≠‡πÉ‡∏ä‡πâ‡∏Ñ‡πà‡∏≤ default)
    # ‡∏£‡∏≤‡∏Ñ‡∏≤‡πÉ‡∏ô Library ‡πÄ‡∏õ‡πá‡∏ô ‡∏ö‡∏≤‡∏ó/‡∏•‡∏ö.‡∏°. ‡∏¢‡∏Å‡πÄ‡∏ß‡πâ‡∏ô AC Interlayer ‡πÄ‡∏õ‡πá‡∏ô ‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.
    if 'price_library' in st.session_state:
        base_lib = st.session_state['price_library']['base_prices']
        ac_lib = st.session_state['price_library']['ac_prices']
        
        base_materials = {}
        
        # ‡πÄ‡∏û‡∏¥‡πà‡∏° AC Interlayer ‡πÄ‡∏â‡∏û‡∏≤‡∏∞ JRCP ‡πÅ‡∏•‡∏∞ CRCP
        if is_concrete_pavement:
            base_materials['AC Interlayer (5 cm)'] = {'unit_cost_cum': ac_lib.get('AC Base Course', {}).get(5, 251), 'is_ac': True, 'default_thick': 5}
        
        # ‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏õ‡∏Å‡∏ï‡∏¥
        base_materials.update({
            'Crushed Rock Base Course': {'unit_cost_cum': base_lib.get('Crushed Rock Base Course', 583), 'is_ac': False},
            'Cement Modified Crushed Rock Base (UCS 24.5 ksc)': {'unit_cost_cum': base_lib.get('Cement Modified Crushed Rock Base (UCS 24.5 ksc)', 864), 'is_ac': False},
            'Cement Treated Base (UCS 40 ksc)': {'unit_cost_cum': base_lib.get('Cement Treated Base (UCS 40 ksc)', 1096), 'is_ac': False},
            'Soil Cement Subbase (UCS 7 ksc)': {'unit_cost_cum': base_lib.get('Soil Cement Subbase (UCS 7 ksc)', 854), 'is_ac': False},
            'Soil Aggregate Subbase': {'unit_cost_cum': base_lib.get('Soil Aggregate Subbase', 375), 'is_ac': False},
            'Selected Material A': {'unit_cost_cum': base_lib.get('Selected Material A', 375), 'is_ac': False},
        })
    else:
        base_materials = {}
        
        if is_concrete_pavement:
            base_materials['AC Interlayer (5 cm)'] = {'unit_cost_cum': 251, 'is_ac': True, 'default_thick': 5}
        
        base_materials.update({
            'Crushed Rock Base Course': {'unit_cost_cum': 583, 'is_ac': False},
            'Cement Modified Crushed Rock Base (UCS 24.5 ksc)': {'unit_cost_cum': 864, 'is_ac': False},
            'Cement Treated Base (UCS 40 ksc)': {'unit_cost_cum': 1096, 'is_ac': False},
            'Soil Cement Subbase (UCS 7 ksc)': {'unit_cost_cum': 854, 'is_ac': False},
            'Soil Aggregate Subbase': {'unit_cost_cum': 375, 'is_ac': False},
            'Selected Material A': {'unit_cost_cum': 375, 'is_ac': False},
        })
    material_names = list(base_materials.keys())
    
    # ‡∏à‡∏≥‡∏ô‡∏ß‡∏ô‡∏ä‡∏±‡πâ‡∏ô‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á (‡∏™‡∏π‡∏á‡∏™‡∏∏‡∏î 5 ‡∏ä‡∏±‡πâ‡∏ô)
    num_base_default = len(base_layers) if len(base_layers) > 0 else 0
    num_base = st.number_input("‡∏à‡∏≥‡∏ô‡∏ß‡∏ô‡∏ä‡∏±‡πâ‡∏ô‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á/‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á", value=num_base_default, 
                                min_value=0, max_value=5, key=f"{key_prefix}_num_base_v{v}")
    
    cols = st.columns([3, 1, 1.2, 1.2, 1.2])
    cols[0].markdown("‡∏ß‡∏±‡∏™‡∏î‡∏∏")
    cols[1].markdown("‡∏´‡∏ô‡∏≤ (cm)")
    cols[2].markdown("‡∏õ‡∏£‡∏¥‡∏°‡∏≤‡∏ì (‡∏ï‡∏£.‡∏°.)")
    cols[3].markdown("‡∏£‡∏≤‡∏Ñ‡∏≤ (‡∏ö‡∏≤‡∏ó/‡∏•‡∏ö.‡∏°.)")
    cols[4].markdown("‡∏£‡∏≤‡∏Ñ‡∏≤ (‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.)")
    
    for i in range(int(num_base)):
        cols = st.columns([3, 1, 1.2, 1.2, 1.2])
        
        # ‡∏Ñ‡πà‡∏≤ default
        if i < len(base_layers):
            default_name = base_layers[i]['name']
            default_thick = base_layers[i]['thickness']
        else:
            default_name = material_names[0]
            default_thick = 20.0
        
        # ‡∏´‡∏≤ index ‡∏Ç‡∏≠‡∏á‡∏ß‡∏±‡∏™‡∏î‡∏∏ default ‚Äî ‡∏£‡∏≠‡∏á‡∏£‡∏±‡∏ö‡∏ó‡∏±‡πâ‡∏á‡∏ä‡∏∑‡πà‡∏≠‡∏ï‡∏£‡∏á‡πÅ‡∏•‡∏∞ partial match ‡∏à‡∏≤‡∏Å JSON
        try:
            default_idx = material_names.index(default_name)
        except ValueError:
            # ‡∏•‡∏≠‡∏á partial match (‡∏ä‡∏∑‡πà‡∏≠‡∏à‡∏≤‡∏Å JSON ‡∏≠‡∏≤‡∏à‡∏¢‡∏≤‡∏ß‡∏Å‡∏ß‡πà‡∏≤)
            default_idx = 0
            dn_lower = default_name.lower()
            for mi, mn in enumerate(material_names):
                if mn.lower() in dn_lower or dn_lower in mn.lower():
                    default_idx = mi
                    break
        
        with cols[0]:
            selected = st.selectbox("‡∏ß‡∏±‡∏™‡∏î‡∏∏", material_names, index=default_idx,
                key=f"{key_prefix}_bm_{i}_v{v}", label_visibility="collapsed")
        with cols[1]:
            # AC Interlayer ‡πÉ‡∏ä‡πâ‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡∏Ñ‡∏á‡∏ó‡∏µ‡πà‡∏à‡∏≤‡∏Å Library
            if base_materials[selected].get('is_ac', False):
                default_thick_val = base_materials[selected].get('default_thick', 5)
                thick = st.number_input("‡∏´‡∏ô‡∏≤", value=float(default_thick_val),
                    key=f"{key_prefix}_bt_{i}_v{v}", label_visibility="collapsed", min_value=0.0, step=1.0)
            else:
                thick = st.number_input("‡∏´‡∏ô‡∏≤", value=float(default_thick),
                    key=f"{key_prefix}_bt_{i}_v{v}", label_visibility="collapsed", min_value=0.0, step=5.0)
        
        # ‡∏õ‡∏£‡∏¥‡∏°‡∏≤‡∏ì = ‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏µ‡πà (‡∏ï‡∏£.‡∏°.) - ‡πÑ‡∏°‡πà‡πÉ‡∏ä‡πà ‡∏•‡∏ö.‡∏°. ‡∏≠‡∏µ‡∏Å‡∏ï‡πà‡∏≠‡πÑ‡∏õ
        auto_qty = area_per_km * road_length
        
        # ‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡∏£‡∏≤‡∏Ñ‡∏≤
        if base_materials[selected].get('is_ac', False):
            # AC Interlayer: ‡∏£‡∏≤‡∏Ñ‡∏≤‡πÄ‡∏õ‡πá‡∏ô ‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°. ‡∏≠‡∏¢‡∏π‡πà‡πÅ‡∏•‡πâ‡∏ß (‡∏î‡∏∂‡∏á‡∏à‡∏≤‡∏Å AC Library ‡∏ï‡∏≤‡∏°‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤)
            if 'price_library' in st.session_state:
                ac_prices = st.session_state['price_library']['ac_prices'].get('AC Base Course', {})
                cost_per_sqm = ac_prices.get(thick, 0)
                if cost_per_sqm == 0 and ac_prices:
                    closest = min(ac_prices.keys(), key=lambda x: abs(x - thick))
                    cost_per_sqm = ac_prices.get(closest, 251)
            else:
                cost_per_sqm = 251  # default 5cm
            lib_cost_cum = cost_per_sqm  # ‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö AC ‡πÄ‡∏Å‡πá‡∏ö‡∏£‡∏≤‡∏Ñ‡∏≤‡∏ï‡∏£‡∏á ‡πÑ‡∏°‡πà‡πÉ‡∏ä‡πà ‡∏•‡∏ö.‡∏°.
        else:
            # ‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡∏õ‡∏Å‡∏ï‡∏¥: ‡πÅ‡∏õ‡∏•‡∏á‡∏£‡∏≤‡∏Ñ‡∏≤ ‡∏ö‡∏≤‡∏ó/‡∏•‡∏ö.‡∏°. ‚Üí ‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.
            lib_cost_cum = base_materials[selected]['unit_cost_cum']  # ‡∏ö‡∏≤‡∏ó/‡∏•‡∏ö.‡∏°.
            cost_per_sqm = lib_cost_cum * thick / 100  # ‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.
        
        with cols[2]:
            st.text(f"{auto_qty:,.0f}")
        with cols[3]:
            # ‡πÅ‡∏™‡∏î‡∏á‡∏£‡∏≤‡∏Ñ‡∏≤ ‡∏ö‡∏≤‡∏ó/‡∏•‡∏ö.‡∏°. (‡∏î‡∏∂‡∏á‡∏à‡∏≤‡∏Å Library)
            if base_materials[selected].get('is_ac', False):
                st.markdown("**-**")  # AC ‡πÑ‡∏°‡πà‡∏°‡∏µ‡∏´‡∏ô‡πà‡∏ß‡∏¢ ‡∏•‡∏ö.‡∏°.
            else:
                st.markdown(f"**{lib_cost_cum:,.2f}**")
        with cols[4]:
            # ‡πÅ‡∏™‡∏î‡∏á‡∏£‡∏≤‡∏Ñ‡∏≤‡∏ó‡∏µ‡πà‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡πÅ‡∏•‡πâ‡∏ß ‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°. (‡∏≠‡∏±‡∏û‡πÄ‡∏î‡∏ó‡∏ï‡∏≤‡∏°‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤‡∏≠‡∏±‡∏ï‡πÇ‡∏ô‡∏°‡∏±‡∏ï‡∏¥)
            st.markdown(f"**{cost_per_sqm:,.2f}**")
        
        updated_layers.append({
            'name': selected, 'thickness': thick, 'unit': 'cm',
            'quantity': auto_qty, 'qty_unit': 'sq.m', 'unit_cost': cost_per_sqm,
            'cost_per_sqm': cost_per_sqm,  # ‡∏£‡∏≤‡∏Ñ‡∏≤‡∏ï‡πà‡∏≠ ‡∏ï‡∏£.‡∏°.
            'cost_cum': lib_cost_cum  # ‡πÄ‡∏Å‡πá‡∏ö‡∏£‡∏≤‡∏Ñ‡∏≤ ‡∏•‡∏ö.‡∏°. ‡πÑ‡∏ß‡πâ‡∏≠‡πâ‡∏≤‡∏á‡∏≠‡∏¥‡∏á
        })
    
    return updated_layers


def render_joint_editor(joints, key_prefix, area_per_km, road_length, v=0):
    """‡πÅ‡∏™‡∏î‡∏á UI ‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡πÅ‡∏Å‡πâ‡πÑ‡∏Ç‡∏£‡∏≠‡∏¢‡∏ï‡πà‡∏≠ ‡∏û‡∏£‡πâ‡∏≠‡∏°‡πÅ‡∏™‡∏î‡∏á‡∏£‡∏≤‡∏Ñ‡∏≤ ‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°."""
    st.markdown("---")
    
    # Checkbox ‡πÄ‡∏•‡∏∑‡∏≠‡∏Å‡∏£‡∏ß‡∏°/‡πÑ‡∏°‡πà‡∏£‡∏ß‡∏° Joints
    col_header = st.columns([3, 1])
    with col_header[0]:
        st.markdown("**‡∏£‡∏≠‡∏¢‡∏ï‡πà‡∏≠ (Joints)**")
    with col_header[1]:
        include_joints = st.checkbox("‡∏£‡∏ß‡∏°‡∏£‡∏≤‡∏Ñ‡∏≤ Joints", value=True, key=f"{key_prefix}_include_joints_v{v}")
    
    cols = st.columns([3, 1.5, 1.5, 1.5])
    cols[0].markdown("‡∏£‡∏≤‡∏¢‡∏Å‡∏≤‡∏£")
    cols[1].markdown("‡∏õ‡∏£‡∏¥‡∏°‡∏≤‡∏ì (m)")
    cols[2].markdown("‡∏£‡∏≤‡∏Ñ‡∏≤/‡∏´‡∏ô‡πà‡∏ß‡∏¢")
    cols[3].markdown("‡∏£‡∏≤‡∏Ñ‡∏≤ (‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.)")
    
    updated_joints = []
    total_area = area_per_km * road_length
    
    for i, joint in enumerate(joints):
        cols = st.columns([3, 1.5, 1.5, 1.5])
        
        with cols[0]:
            st.text(joint['name'])
        
        with cols[1]:
            qty = st.number_input(
                "‡∏õ‡∏£‡∏¥‡∏°‡∏≤‡∏ì (m)", value=float(joint['quantity']),
                key=f"{key_prefix}_jq_{i}_v{v}", label_visibility="collapsed",
                min_value=0.0, step=100.0
            )
        
        with cols[2]:
            cost = st.number_input(
                "‡∏£‡∏≤‡∏Ñ‡∏≤/‡∏°.", value=float(joint['unit_cost']),
                key=f"{key_prefix}_jc_{i}_v{v}", label_visibility="collapsed",
                min_value=0.0, step=10.0
            )
        
        # ‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡∏£‡∏≤‡∏Ñ‡∏≤ ‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.
        joint_total = qty * cost
        cost_per_sqm = joint_total / total_area if total_area > 0 else 0
        
        with cols[3]:
            st.markdown(f"**{cost_per_sqm:.2f}**")
        
        updated_joints.append({
            'name': joint['name'],
            'quantity': qty,
            'qty_unit': joint['qty_unit'],
            'unit_cost': cost,
            'cost_per_sqm': cost_per_sqm
        })
    
    return updated_joints, include_joints


def generate_word_report_table(project_info, structure_type, structure_name, cbr, layers, joints, road_length):
    """‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô Word ‡∏£‡∏π‡∏õ‡πÅ‡∏ö‡∏ö‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏Ñ‡πà‡∏≤‡∏Å‡πà‡∏≠‡∏™‡∏£‡πâ‡∏≤‡∏á (‡∏ï‡∏≤‡∏°‡∏ï‡∏±‡∏ß‡∏≠‡∏¢‡πà‡∏≤‡∏á‡πÉ‡∏ô‡πÄ‡∏≠‡∏Å‡∏™‡∏≤‡∏£)"""
    doc = Document()
    
    # ‡∏ï‡∏±‡πâ‡∏á‡∏Ñ‡πà‡∏≤ font
    style = doc.styles['Normal']
    style.font.name = 'TH SarabunPSK'
    style.font.size = Pt(14)
    
    # ‡∏´‡∏±‡∏ß‡∏Ç‡πâ‡∏≠
    title = doc.add_paragraph()
    title_run = title.add_run('‡∏£‡∏≤‡∏Ñ‡∏≤‡∏Ñ‡πà‡∏≤‡∏Å‡πà‡∏≠‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏Ç‡∏≠‡∏á‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á' + structure_name)
    title_run.bold = True
    title_run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏• CBR
    info_text = f"‡∏ú‡∏¥‡∏ß‡∏à‡∏£‡∏≤‡∏à‡∏£{structure_type} ‡∏Å‡∏£‡∏ì‡∏µ‡∏ä‡∏±‡πâ‡∏ô‡∏î‡∏¥‡∏ô‡πÄ‡∏î‡∏¥‡∏°‡∏°‡∏µ‡∏Ñ‡πà‡∏≤ CBR = {cbr}%"
    doc.add_paragraph(info_text).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ‡πÅ‡∏¢‡∏Å layers ‡πÄ‡∏õ‡πá‡∏ô‡∏Å‡∏•‡∏∏‡πà‡∏°
    surface_layers = []
    base_layers = []
    for layer in layers:
        name_lower = layer['name'].lower()
        if any(x in name_lower for x in ['wearing', 'binder', 'asphalt', 'concrete', 'tack', 'prime', 'geotextile', 'steel', 'ac base', 'ac interlayer', 'interlayer']):
            surface_layers.append(layer)
        else:
            base_layers.append(layer)
    
    # ‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡∏à‡∏≥‡∏ô‡∏ß‡∏ô‡πÅ‡∏ñ‡∏ß
    num_rows = 2 + len(surface_layers) + 1  # header + ‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á header + items + ‡∏£‡∏ß‡∏°1
    if joints:
        num_rows += 1 + len(joints) + 1  # ‡∏£‡∏≠‡∏¢‡∏ï‡πà‡∏≠ header + items + ‡∏£‡∏ß‡∏°2
    num_rows += 1 + len(base_layers) + 1  # ‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á header + items + ‡∏£‡∏ß‡∏°3
    num_rows += 2  # ‡∏£‡∏ß‡∏°‡∏ó‡∏±‡πâ‡∏á‡∏´‡∏°‡∏î + ‡∏™‡∏£‡∏∏‡∏õ
    
    table = doc.add_table(rows=num_rows, cols=7)
    table.style = 'Table Grid'
    
    # Header
    headers = ['‡∏•‡∏≥‡∏î‡∏±‡∏ö', '‡∏Ñ‡πà‡∏≤‡πÉ‡∏ä‡πâ‡∏à‡πà‡∏≤‡∏¢‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡∏ß‡∏±‡∏™‡∏î‡∏∏', '‡∏£‡∏≤‡∏¢‡∏•‡∏∞‡πÄ‡∏≠‡∏µ‡∏¢‡∏î‡∏´‡∏ô‡πà‡∏ß‡∏¢', '‡∏õ‡∏£‡∏¥‡∏°‡∏≤‡∏ì‡∏ï‡πà‡∏≠', '‡∏´‡∏ô‡πà‡∏ß‡∏¢', '‡∏£‡∏≤‡∏Ñ‡∏≤‡∏ï‡πà‡∏≠‡∏´‡∏ô‡πà‡∏ß‡∏¢\n(‡∏ö‡∏≤‡∏ó/‡∏´‡∏ô‡πà‡∏ß‡∏¢)', '‡∏°‡∏π‡∏•‡∏Ñ‡πà‡∏≤\n(‡∏ö‡∏≤‡∏ó)']
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    
    row_idx = 1
    running_total = 0
    
    # ‡∏Å‡∏•‡∏∏‡πà‡∏° 1: ‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á
    table.rows[row_idx].cells[0].text = '1'
    table.rows[row_idx].cells[1].text = '‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á'
    row_idx += 1
    
    surface_total = 0
    for i, layer in enumerate(surface_layers, 1):
        qty = layer['quantity'] * road_length
        cost = qty * layer['unit_cost']
        table.rows[row_idx].cells[0].text = f'1.{i}'
        table.rows[row_idx].cells[1].text = layer['name']
        table.rows[row_idx].cells[2].text = f"{layer['thickness']} {layer['unit']}"
        table.rows[row_idx].cells[3].text = f"{qty:,.0f}"
        table.rows[row_idx].cells[4].text = layer['qty_unit']
        table.rows[row_idx].cells[5].text = f"{layer['unit_cost']:,.0f}"
        table.rows[row_idx].cells[6].text = f"{cost:,.0f}"
        surface_total += cost
        row_idx += 1
    
    table.rows[row_idx].cells[1].text = '‡∏£‡∏ß‡∏° 1'
    table.rows[row_idx].cells[6].text = f"{surface_total:,.0f}"
    running_total += surface_total
    row_idx += 1
    
    # ‡∏Å‡∏•‡∏∏‡πà‡∏° 2: ‡∏£‡∏≠‡∏¢‡∏ï‡πà‡∏≠
    joint_total = 0
    if joints:
        table.rows[row_idx].cells[0].text = '2'
        table.rows[row_idx].cells[1].text = '‡∏£‡∏≠‡∏¢‡∏ï‡πà‡∏≠'
        row_idx += 1
        
        for i, joint in enumerate(joints, 1):
            qty = joint['quantity'] * road_length
            cost = qty * joint['unit_cost']
            table.rows[row_idx].cells[0].text = f'2.{i}'
            table.rows[row_idx].cells[1].text = joint['name']
            table.rows[row_idx].cells[3].text = f"{qty:,.0f}"
            table.rows[row_idx].cells[4].text = joint['qty_unit']
            table.rows[row_idx].cells[5].text = f"{joint['unit_cost']:,.0f}"
            table.rows[row_idx].cells[6].text = f"{cost:,.0f}"
            joint_total += cost
            row_idx += 1
        
        table.rows[row_idx].cells[1].text = '‡∏£‡∏ß‡∏° 2'
        table.rows[row_idx].cells[6].text = f"{joint_total:,.0f}"
        running_total += joint_total
        row_idx += 1
        group_num = 3
    else:
        group_num = 2
    
    # ‡∏Å‡∏•‡∏∏‡πà‡∏° 3: ‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡πÅ‡∏•‡∏∞‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á
    table.rows[row_idx].cells[0].text = str(group_num)
    table.rows[row_idx].cells[1].text = '‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡πÅ‡∏•‡∏∞‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á'
    row_idx += 1
    
    base_total = 0
    for i, layer in enumerate(base_layers, 1):
        qty = layer['quantity'] * road_length
        cost = qty * layer['unit_cost']
        table.rows[row_idx].cells[0].text = f'{group_num}.{i}'
        table.rows[row_idx].cells[1].text = layer['name']
        table.rows[row_idx].cells[2].text = f"{layer['thickness']} {layer['unit']}"
        table.rows[row_idx].cells[3].text = f"{qty:,.0f}"
        table.rows[row_idx].cells[4].text = layer['qty_unit']
        table.rows[row_idx].cells[5].text = f"{layer['unit_cost']:,.0f}"
        table.rows[row_idx].cells[6].text = f"{cost:,.0f}"
        base_total += cost
        row_idx += 1
    
    table.rows[row_idx].cells[1].text = f'‡∏£‡∏ß‡∏° {group_num}'
    table.rows[row_idx].cells[6].text = f"{base_total:,.0f}"
    running_total += base_total
    row_idx += 1
    
    # ‡∏£‡∏ß‡∏°‡∏ó‡∏±‡πâ‡∏á‡∏´‡∏°‡∏î
    sum_text = '‡∏£‡∏ß‡∏° 1+2+3' if joints else '‡∏£‡∏ß‡∏° 1+2'
    table.rows[row_idx].cells[1].text = sum_text
    table.rows[row_idx].cells[3].text = f"{running_total:,.0f}"
    table.rows[row_idx].cells[6].text = '‡∏ö‡∏≤‡∏ó'
    row_idx += 1
    
    # ‡∏™‡∏£‡∏∏‡∏õ‡∏£‡∏≤‡∏Ñ‡∏≤‡∏ï‡πà‡∏≠‡∏Å‡∏¥‡πÇ‡∏•‡πÄ‡∏°‡∏ï‡∏£
    cost_per_km = running_total / road_length / 1_000_000
    table.rows[row_idx].cells[1].text = '‡∏™‡∏£‡∏∏‡∏õ‡∏£‡∏≤‡∏Ñ‡∏≤‡∏ï‡πà‡∏≠‡∏Å‡∏¥‡πÇ‡∏•‡πÄ‡∏°‡∏ï‡∏£‡πÉ‡∏ô2‡∏ó‡∏¥‡∏®‡∏ó‡∏≤‡∏á'
    table.rows[row_idx].cells[3].text = f"{cost_per_km:.2f}"
    table.rows[row_idx].cells[6].text = '‡∏•‡πâ‡∏≤‡∏ô‡∏ö‡∏≤‡∏ó'
    
    # Footer
    doc.add_paragraph()
    lane_width = project_info.get('lane_width', 3.5)
    shoulder_left = project_info.get('shoulder_left', 2.5)
    shoulder_right = project_info.get('shoulder_right', 1.5)
    total_width = project_info.get('total_width', 11.0)
    
    doc.add_paragraph(f"‡∏´‡∏°‡∏≤‡∏¢‡πÄ‡∏´‡∏ï‡∏∏: ‡∏Ñ‡∏ß‡∏≤‡∏°‡∏Å‡∏ß‡πâ‡∏≤‡∏á‡∏ä‡πà‡∏≠‡∏á‡∏à‡∏£‡∏≤‡∏à‡∏£ {lane_width} ‡∏°. ‡πÑ‡∏´‡∏•‡πà‡∏ó‡∏≤‡∏á‡∏ã‡πâ‡∏≤‡∏¢ {shoulder_left} ‡∏°. ‡πÑ‡∏´‡∏•‡πà‡∏ó‡∏≤‡∏á‡∏Ç‡∏ß‡∏≤ {shoulder_right} ‡∏°.")
    doc.add_paragraph(f"‡∏£‡∏ß‡∏°‡∏ó‡∏±‡πâ‡∏á‡∏™‡∏¥‡πâ‡∏ô‡∏Ñ‡∏ß‡∏≤‡∏°‡∏Å‡∏ß‡πâ‡∏≤‡∏á‡∏ñ‡∏ô‡∏ô {total_width} ‡∏°. (‡∏ä‡πà‡∏≠‡∏á‡∏•‡∏∞ {lane_width} ‡∏°.) ‡∏¢‡∏≤‡∏ß {road_length} ‡∏Å‡∏¥‡πÇ‡∏•‡πÄ‡∏°‡∏ï‡∏£")
    doc.add_paragraph(f"‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô‡∏™‡∏£‡πâ‡∏≤‡∏á‡πÄ‡∏°‡∏∑‡πà‡∏≠: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    
    return doc


def generate_word_report_materials_only(project_info, all_details):
    """‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô Word - ‡πÄ‡∏â‡∏û‡∏≤‡∏∞‡∏ß‡∏±‡∏™‡∏î‡∏∏‡πÅ‡∏•‡∏∞‡∏£‡∏≤‡∏Ñ‡∏≤ (‡πÑ‡∏°‡πà‡∏°‡∏µ NPV) ‡∏û‡∏£‡πâ‡∏≠‡∏°‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏™‡∏£‡∏∏‡∏õ‡πÅ‡∏¢‡∏Å‡∏ä‡∏ô‡∏¥‡∏î"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx ‡πÑ‡∏°‡πà‡∏™‡∏≤‡∏°‡∏≤‡∏£‡∏ñ‡πÉ‡∏ä‡πâ‡∏á‡∏≤‡∏ô‡πÑ‡∏î‡πâ")
    
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'TH SarabunPSK'
    style.font.size = Pt(16)
    
    doc.add_heading('‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô‡∏ß‡∏±‡∏™‡∏î‡∏∏‡πÅ‡∏•‡∏∞‡∏£‡∏≤‡∏Ñ‡∏≤‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á', 0)
    
    doc.add_heading('1. ‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡πÇ‡∏Ñ‡∏£‡∏á‡∏Å‡∏≤‡∏£', level=1)
    doc.add_paragraph(f"‡∏ä‡∏∑‡πà‡∏≠‡πÇ‡∏Ñ‡∏£‡∏á‡∏Å‡∏≤‡∏£: {project_info.get('name', '-')}")
    doc.add_paragraph(f"‡∏Ñ‡∏ß‡∏≤‡∏°‡∏¢‡∏≤‡∏ß‡∏ñ‡∏ô‡∏ô: {project_info.get('length', 1):.2f} ‡∏Å‡∏°.")
    doc.add_paragraph(f"‡∏Ñ‡∏ß‡∏≤‡∏°‡∏Å‡∏ß‡πâ‡∏≤‡∏á‡∏£‡∏ß‡∏°: {project_info.get('total_width', 0):.2f} ‡∏°.")
    doc.add_paragraph(f"‡∏à‡∏≥‡∏ô‡∏ß‡∏ô‡∏ä‡πà‡∏≠‡∏á‡∏à‡∏£‡∏≤‡∏à‡∏£: {project_info.get('num_lanes', 2)} ‡∏ä‡πà‡∏≠‡∏á")
    
    doc.add_heading('2. ‡∏£‡∏≤‡∏¢‡∏•‡∏∞‡πÄ‡∏≠‡∏µ‡∏¢‡∏î‡∏ß‡∏±‡∏™‡∏î‡∏∏‡πÅ‡∏•‡∏∞‡∏£‡∏≤‡∏Ñ‡∏≤', level=1)
    
    # ‡πÄ‡∏Å‡πá‡∏ö‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏™‡∏£‡∏∏‡∏õ
    summary_data = []
    length = project_info.get('length', 1)
    
    for ptype, data in all_details.items():
        structure_name = data.get('name', ptype)
        details = data.get('details', [])
        
        doc.add_heading(structure_name, level=2)
        if details:
            table = doc.add_table(rows=len(details)+1, cols=4)
            table.style = 'Table Grid'
            
            # Header
            headers = ['‡∏£‡∏≤‡∏¢‡∏Å‡∏≤‡∏£', '‡∏õ‡∏£‡∏¥‡∏°‡∏≤‡∏ì', '‡∏£‡∏≤‡∏Ñ‡∏≤/‡∏´‡∏ô‡πà‡∏ß‡∏¢ (‡∏ö‡∏≤‡∏ó)', '‡∏°‡∏π‡∏•‡∏Ñ‡πà‡∏≤ (‡∏ö‡∏≤‡∏ó)']
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = h
                # ‡∏ó‡∏≥‡πÉ‡∏´‡πâ header ‡πÄ‡∏õ‡πá‡∏ô‡∏ï‡∏±‡∏ß‡∏´‡∏ô‡∏≤
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Data rows
            subtotal = 0
            for i, d in enumerate(details):
                table.rows[i+1].cells[0].text = str(d['‡∏£‡∏≤‡∏¢‡∏Å‡∏≤‡∏£'])
                table.rows[i+1].cells[1].text = f"{d['‡∏õ‡∏£‡∏¥‡∏°‡∏≤‡∏ì']:,.0f} {d['‡∏´‡∏ô‡πà‡∏ß‡∏¢']}"
                table.rows[i+1].cells[2].text = f"{d['‡∏£‡∏≤‡∏Ñ‡∏≤/‡∏´‡∏ô‡πà‡∏ß‡∏¢']:,.0f}"
                table.rows[i+1].cells[3].text = f"{d['‡∏°‡∏π‡∏•‡∏Ñ‡πà‡∏≤ (‡∏ö‡∏≤‡∏ó)']:,.0f}"
                subtotal += d['‡∏°‡∏π‡∏•‡∏Ñ‡πà‡∏≤ (‡∏ö‡∏≤‡∏ó)']
            
            # ‡πÅ‡∏™‡∏î‡∏á‡∏¢‡∏≠‡∏î‡∏£‡∏ß‡∏°‡∏¢‡πà‡∏≠‡∏¢
            doc.add_paragraph(f"‡∏£‡∏ß‡∏° {structure_name}: {subtotal:,.0f} ‡∏ö‡∏≤‡∏ó", style='Intense Quote')
            doc.add_paragraph()
            
            # ‡πÄ‡∏Å‡πá‡∏ö‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏™‡∏£‡∏∏‡∏õ
            cost_per_km_million = data.get('cost_per_km', 0)  # ‡∏•‡πâ‡∏≤‡∏ô‡∏ö‡∏≤‡∏ó/‡∏Å‡∏°.
            cost_per_km_baht = cost_per_km_million * 1_000_000  # ‡∏ö‡∏≤‡∏ó/‡∏Å‡∏°.
            cost_per_sqm = data.get('cost_sqm', 0)  # ‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.
            total_value = subtotal  # ‡∏°‡∏π‡∏•‡∏Ñ‡πà‡∏≤‡∏£‡∏ß‡∏° (‡∏ö‡∏≤‡∏ó)
            
            summary_data.append({
                'name': structure_name,
                'total_value': total_value,
                'cost_per_km_million': cost_per_km_million,
                'cost_per_sqm': cost_per_sqm
            })
    
    # ‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏™‡∏£‡∏∏‡∏õ‡πÅ‡∏¢‡∏Å‡πÅ‡∏ï‡πà‡∏•‡∏∞‡∏ä‡∏ô‡∏¥‡∏î
    doc.add_heading('3. ‡∏™‡∏£‡∏∏‡∏õ‡∏Ñ‡πà‡∏≤‡πÉ‡∏ä‡πâ‡∏à‡πà‡∏≤‡∏¢', level=1)
    
    if summary_data:
        table = doc.add_table(rows=len(summary_data)+1, cols=4)
        table.style = 'Table Grid'
        
        # Header
        headers = ['‡∏ä‡∏ô‡∏¥‡∏î‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á', '‡∏°‡∏π‡∏•‡∏Ñ‡πà‡∏≤‡∏£‡∏ß‡∏°/‡∏Å‡∏°. (‡∏ö‡∏≤‡∏ó)', '‡∏£‡∏≤‡∏Ñ‡∏≤/‡∏Å‡∏°. (‡∏•‡πâ‡∏≤‡∏ô‡∏ö‡∏≤‡∏ó)', '‡∏£‡∏≤‡∏Ñ‡∏≤/‡∏ï‡∏£.‡∏°. (‡∏ö‡∏≤‡∏ó)']
        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = h
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Data rows
        for i, item in enumerate(summary_data):
            # ‡∏°‡∏π‡∏•‡∏Ñ‡πà‡∏≤‡∏£‡∏ß‡∏°‡∏ï‡πà‡∏≠ ‡∏Å‡∏°. = ‡∏°‡∏π‡∏•‡∏Ñ‡πà‡∏≤‡∏£‡∏ß‡∏° / ‡∏Ñ‡∏ß‡∏≤‡∏°‡∏¢‡∏≤‡∏ß
            total_per_km = item['total_value'] / length if length > 0 else 0
            
            table.rows[i+1].cells[0].text = item['name']
            table.rows[i+1].cells[1].text = f"{total_per_km:,.0f}"
            table.rows[i+1].cells[2].text = f"{item['cost_per_km_million']:.2f}"
            table.rows[i+1].cells[3].text = f"{item['cost_per_sqm']:,.2f}"
    
    doc.add_paragraph()
    doc.add_paragraph(f"‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô‡∏™‡∏£‡πâ‡∏≤‡∏á‡πÄ‡∏°‡∏∑‡πà‡∏≠: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    
    return doc


# ===== Main Application =====

@st.cache_data
def generate_excel_template():
    """‡∏™‡∏£‡πâ‡∏≤‡∏á Excel Template ‡πÅ‡∏•‡∏∞ cache ‡πÑ‡∏ß‡πâ‡πÄ‡∏û‡∏∑‡πà‡∏≠‡∏õ‡∏£‡∏∞‡∏™‡∏¥‡∏ó‡∏ò‡∏¥‡∏†‡∏≤‡∏û"""
    template_data = {
        'AC_Prices': pd.DataFrame({
            'Material': ['PMA Wearing Course', 'AC Wearing Course', 'AC Binder Course', 'AC Base Course'],
            '2.5cm': [170, 128, 129, 129],
            '3cm': [203, 152, 154, 154],
            '4cm': [268, 202, 202, 202],
            '5cm': [333, 250, 251, 251],
            '6cm': [406, 306, 308, 308],
            '7cm': [471, 355, 356, 356],
            '8cm': [536, 403, 405, 405],
            '9cm': [601, 452, 454, 454],
            '10cm': [667, 502, 503, 503],
        }),
        'Concrete_Prices': pd.DataFrame({
            'Type': ['JRCP', 'JPCP', 'CRCP'],
            '25cm': [924, 928, 1245],
            '28cm': [1002, 1000, 1358],
            '32cm': [1106, 1095, 1509],
            '35cm': [1184, 1167, 1622],
        }),
        'Base_Materials': pd.DataFrame({
            'Material': list(BASE_MATERIAL_PRICES.keys()),
            'Price (Baht/cu.m)': list(BASE_MATERIAL_PRICES.values()),
        })
    }
    
    # ‡∏™‡∏£‡πâ‡∏≤‡∏á Excel file
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        template_data['AC_Prices'].to_excel(writer, sheet_name='AC_Prices', index=False)
        template_data['Concrete_Prices'].to_excel(writer, sheet_name='Concrete_Prices', index=False)
        template_data['Base_Materials'].to_excel(writer, sheet_name='Base_Materials', index=False)
    output.seek(0)
    
    return output.getvalue()


def main():
    st.markdown('<div class="main-header">üõ£Ô∏è ‡∏£‡∏∞‡∏ö‡∏ö‡∏ß‡∏¥‡πÄ‡∏Ñ‡∏£‡∏≤‡∏∞‡∏´‡πå‡∏Ñ‡πà‡∏≤‡∏Å‡πà‡∏≠‡∏™‡∏£‡πâ‡∏≤‡∏á‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á</div>', unsafe_allow_html=True)
    st.markdown("##### ‡∏ï‡∏≤‡∏°‡πÅ‡∏ô‡∏ß‡∏ó‡∏≤‡∏á AASHTO 1993 - ‡∏£‡∏≠‡∏á‡∏£‡∏±‡∏ö AC, JPCP/JRCP, CRCP")
    st.markdown("""
    <div style='text-align: center; color: #666; font-size: 0.9rem; margin-top: -10px; margin-bottom: 20px;'>
        ‡∏û‡∏±‡∏í‡∏ô‡∏≤‡πÇ‡∏î‡∏¢ <b>‡∏£‡∏®.‡∏î‡∏£.‡∏≠‡∏¥‡∏ó‡∏ò‡∏¥‡∏û‡∏• ‡∏°‡∏µ‡∏ú‡∏•</b><br>
        ‡∏†‡∏≤‡∏Ñ‡∏ß‡∏¥‡∏ä‡∏≤‡∏Ñ‡∏£‡∏∏‡∏®‡∏≤‡∏™‡∏ï‡∏£‡πå‡πÇ‡∏¢‡∏ò‡∏≤ ‡∏Ñ‡∏ì‡∏∞‡∏Ñ‡∏£‡∏∏‡∏®‡∏≤‡∏™‡∏ï‡∏£‡πå‡∏≠‡∏∏‡∏ï‡∏™‡∏≤‡∏´‡∏Å‡∏£‡∏£‡∏°<br>
        ‡∏°‡∏´‡∏≤‡∏ß‡∏¥‡∏ó‡∏¢‡∏≤‡∏•‡∏±‡∏¢‡πÄ‡∏ó‡∏Ñ‡πÇ‡∏ô‡πÇ‡∏•‡∏¢‡∏µ‡∏û‡∏£‡∏∞‡∏à‡∏≠‡∏°‡πÄ‡∏Å‡∏•‡πâ‡∏≤‡∏û‡∏£‡∏∞‡∏ô‡∏Ñ‡∏£‡πÄ‡∏´‡∏ô‡∏∑‡∏≠ (‡∏°‡∏à‡∏û.)
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar
    with st.sidebar:
        st.header("üìã ‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡πÇ‡∏Ñ‡∏£‡∏á‡∏Å‡∏≤‡∏£")
        
        # ===== Upload JSON =====
        st.subheader("üìÇ ‡πÇ‡∏´‡∏•‡∏î‡πÇ‡∏Ñ‡∏£‡∏á‡∏Å‡∏≤‡∏£")
        uploaded_json = st.file_uploader(
            "‡πÄ‡∏•‡∏∑‡∏≠‡∏Å‡πÑ‡∏ü‡∏•‡πå JSON",
            type=['json'],
            help="‡∏≠‡∏±‡∏û‡πÇ‡∏´‡∏•‡∏î‡πÑ‡∏ü‡∏•‡πå‡πÇ‡∏Ñ‡∏£‡∏á‡∏Å‡∏≤‡∏£‡∏ó‡∏µ‡πà‡∏ö‡∏±‡∏ô‡∏ó‡∏∂‡∏Å‡πÑ‡∏ß‡πâ",
            key="upload_json"
        )
        
        if uploaded_json is not None:
            try:
                import hashlib
                file_bytes = uploaded_json.read()
                file_hash = hashlib.md5(file_bytes).hexdigest()
                loaded_data = json.loads(file_bytes.decode('utf-8'))
                st.success("‚úÖ ‡πÇ‡∏´‡∏•‡∏î‡πÑ‡∏ü‡∏•‡πå‡∏™‡∏≥‡πÄ‡∏£‡πá‡∏à!")
                
                # ‡πÅ‡∏™‡∏î‡∏á‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏ó‡∏µ‡πà‡πÇ‡∏´‡∏•‡∏î
                if 'project_info' in loaded_data:
                    st.info(f"üìå ‡πÇ‡∏Ñ‡∏£‡∏á‡∏Å‡∏≤‡∏£: {loaded_data['project_info'].get('name', '-')}")
                    st.info(f"üìÖ ‡∏ö‡∏±‡∏ô‡∏ó‡∏∂‡∏Å‡πÄ‡∏°‡∏∑‡πà‡∏≠: {loaded_data.get('saved_at', '-')}")
                
                # ‡πÄ‡∏Å‡πá‡∏ö‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡πÉ‡∏ô session_state
                if st.button("üì• ‡∏ô‡∏≥‡πÄ‡∏Ç‡πâ‡∏≤‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•", key="import_json"):
                    if 'project_info' in loaded_data:
                        # ‡∏õ‡πâ‡∏≠‡∏á‡∏Å‡∏±‡∏ô load ‡∏ã‡πâ‡∏≥‡∏î‡πâ‡∏ß‡∏¢ hash
                        if st.session_state.get('loaded_json_hash') != file_hash:
                            st.session_state['loaded_project'] = loaded_data
                            st.session_state['loaded_json_hash'] = file_hash
                            # ‡πÄ‡∏û‡∏¥‡πà‡∏° version ‚Üí widget keys ‡πÄ‡∏õ‡∏•‡∏µ‡πà‡∏¢‡∏ô ‚Üí Streamlit ‡∏≠‡πà‡∏≤‡∏ô value= ‡πÉ‡∏´‡∏°‡πà
                            st.session_state['json_version'] = st.session_state.get('json_version', 0) + 1
                        st.rerun()
            except Exception as e:
                st.error(f"‚ùå ‡πÑ‡∏°‡πà‡∏™‡∏≤‡∏°‡∏≤‡∏£‡∏ñ‡∏≠‡πà‡∏≤‡∏ô‡πÑ‡∏ü‡∏•‡πå‡πÑ‡∏î‡πâ: {e}")
        
        st.divider()
        
        # ===== Upload Price Library Excel =====
        st.subheader("üí∞ Price Library (Excel)")
        st.caption("‡∏≠‡∏±‡∏û‡πÇ‡∏´‡∏•‡∏î Excel ‡πÄ‡∏û‡∏∑‡πà‡∏≠‡πÅ‡∏ó‡∏ô‡∏ó‡∏µ‡πà‡∏£‡∏≤‡∏Ñ‡∏≤ Default")
        
        uploaded_price_excel = st.file_uploader(
            "‡πÄ‡∏•‡∏∑‡∏≠‡∏Å‡πÑ‡∏ü‡∏•‡πå Excel Price Library",
            type=['xlsx', 'xls'],
            help="‡∏≠‡∏±‡∏û‡πÇ‡∏´‡∏•‡∏î‡πÑ‡∏ü‡∏•‡πå‡∏£‡∏≤‡∏Ñ‡∏≤‡∏ó‡∏µ‡πà Download ‡∏à‡∏≤‡∏Å Tab 1",
            key="sidebar_upload_price"
        )
        
        # ‡πÇ‡∏´‡∏•‡∏î‡∏£‡∏≤‡∏Ñ‡∏≤‡∏à‡∏≤‡∏Å Excel ‡∏ñ‡πâ‡∏≤‡∏°‡∏µ upload
        if uploaded_price_excel is not None:
            try:
                # ‡∏≠‡πà‡∏≤‡∏ô‡πÑ‡∏ü‡∏•‡πå Excel
                ac_df = pd.read_excel(uploaded_price_excel, sheet_name='AC_Prices')
                concrete_df = pd.read_excel(uploaded_price_excel, sheet_name='Concrete_Prices')
                base_df = pd.read_excel(uploaded_price_excel, sheet_name='Base_Materials')
                
                # ‡πÅ‡∏õ‡∏•‡∏á‡πÄ‡∏õ‡πá‡∏ô dictionary format
                uploaded_ac_prices = {}
                for _, row in ac_df.iterrows():
                    material = row['Material']
                    prices = {}
                    for col in ac_df.columns[1:]:
                        thickness = float(col.replace('cm', ''))
                        prices[thickness] = float(row[col])
                    uploaded_ac_prices[material] = prices
                
                uploaded_concrete_prices = {}
                for _, row in concrete_df.iterrows():
                    conc_type = row['Type']
                    prices = {}
                    for col in concrete_df.columns[1:]:
                        thickness = int(col.replace('cm', ''))
                        prices[thickness] = float(row[col])
                    uploaded_concrete_prices[conc_type] = prices
                
                uploaded_base_prices = {}
                for _, row in base_df.iterrows():
                    uploaded_base_prices[row['Material']] = float(row['Price (Baht/cu.m)'])
                
                # ‡πÄ‡∏Å‡πá‡∏ö‡πÉ‡∏ô session_state
                st.session_state['uploaded_price_library'] = {
                    'ac_prices': uploaded_ac_prices,
                    'concrete_prices': uploaded_concrete_prices,
                    'base_prices': uploaded_base_prices,
                }
                
                # ‡πÄ‡∏û‡∏¥‡πà‡∏° version ‡πÄ‡∏û‡∏∑‡πà‡∏≠‡∏ö‡∏±‡∏á‡∏Ñ‡∏±‡∏ö‡πÉ‡∏´‡πâ widget keys ‡πÄ‡∏õ‡∏•‡∏µ‡πà‡∏¢‡∏ô
                import hashlib
                file_hash = hashlib.md5(uploaded_price_excel.getvalue()).hexdigest()[:8]
                st.session_state['price_upload_version'] = file_hash
                
                st.success("‚úÖ ‡πÇ‡∏´‡∏•‡∏î Price Library ‡∏™‡∏≥‡πÄ‡∏£‡πá‡∏à!")
                st.caption(f"üìä {len(uploaded_ac_prices)} AC types, {len(uploaded_concrete_prices)} Concrete types")
                
                # ‡πÅ‡∏™‡∏î‡∏á‡∏ï‡∏±‡∏ß‡∏≠‡∏¢‡πà‡∏≤‡∏á‡∏£‡∏≤‡∏Ñ‡∏≤‡∏ó‡∏µ‡πà‡∏≠‡πà‡∏≤‡∏ô‡πÑ‡∏î‡πâ
                with st.expander("üîç ‡∏ï‡∏±‡∏ß‡∏≠‡∏¢‡πà‡∏≤‡∏á‡∏£‡∏≤‡∏Ñ‡∏≤‡∏ó‡∏µ‡πà‡∏≠‡πà‡∏≤‡∏ô‡πÑ‡∏î‡πâ"):
                    st.write("**AC Wearing Course (7cm):**", uploaded_ac_prices.get('AC Wearing Course', {}).get(7.0, 'N/A'))
                    st.write("**JPCP (25cm):**", uploaded_concrete_prices.get('JPCP', {}).get(25, 'N/A'))
                    st.write("**Crushed Rock:**", uploaded_base_prices.get('Crushed Rock Base Course', 'N/A'))
                
            except Exception as e:
                st.error(f"‚ùå ‡∏≠‡πà‡∏≤‡∏ô‡πÑ‡∏ü‡∏•‡πå‡πÑ‡∏°‡πà‡∏™‡∏≥‡πÄ‡∏£‡πá‡∏à: {str(e)}")
        
        st.divider()
        
        # ‡∏ï‡∏£‡∏ß‡∏à‡∏™‡∏≠‡∏ö‡∏ß‡πà‡∏≤‡∏°‡∏µ‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏ó‡∏µ‡πà‡πÇ‡∏´‡∏•‡∏î‡∏°‡∏≤‡∏´‡∏£‡∏∑‡∏≠‡πÑ‡∏°‡πà
        loaded_project = st.session_state.get('loaded_project', {})
        loaded_info = loaded_project.get('project_info', {})
        
        project_name = st.text_input("‡∏ä‡∏∑‡πà‡∏≠‡πÇ‡∏Ñ‡∏£‡∏á‡∏Å‡∏≤‡∏£", value=loaded_info.get('name', "‡πÇ‡∏Ñ‡∏£‡∏á‡∏Å‡∏≤‡∏£‡∏Å‡πà‡∏≠‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ó‡∏≤‡∏á‡∏´‡∏•‡∏ß‡∏á"))
        road_length = st.number_input("‡∏Ñ‡∏ß‡∏≤‡∏°‡∏¢‡∏≤‡∏ß‡∏ñ‡∏ô‡∏ô (‡∏Å‡∏°.)", value=loaded_info.get('length', 1.0), min_value=0.1, step=0.1)
        
        st.divider()
        st.header("üìê ‡∏Ç‡∏ô‡∏≤‡∏î‡∏ñ‡∏ô‡∏ô")
        lane_width = st.number_input("‡∏Ñ‡∏ß‡∏≤‡∏°‡∏Å‡∏ß‡πâ‡∏≤‡∏á‡∏ä‡πà‡∏≠‡∏á‡∏à‡∏£‡∏≤‡∏à‡∏£ (‡∏°.)", value=loaded_info.get('lane_width', 3.5), min_value=2.5, max_value=4.0, step=0.25)
        
        # ‡∏´‡∏≤ index ‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö num_lanes_per_direction
        # ‡πÅ‡∏õ‡∏•‡∏á‡∏à‡∏≤‡∏Å num_lanes ‡πÄ‡∏î‡∏¥‡∏° (‡∏£‡∏ß‡∏° 2 ‡∏ó‡∏¥‡∏®‡∏ó‡∏≤‡∏á) ‡πÄ‡∏õ‡πá‡∏ô lanes_per_direction
        default_lanes_total = loaded_info.get('num_lanes', 4)  # default 4 (2 ‡∏ä‡πà‡∏≠‡∏á/‡∏ó‡∏¥‡∏®‡∏ó‡∏≤‡∏á)
        default_lanes_per_dir = default_lanes_total // 2
        
        lanes_per_dir_options = [2, 3, 4]
        lanes_per_dir_idx = lanes_per_dir_options.index(default_lanes_per_dir) if default_lanes_per_dir in lanes_per_dir_options else 0
        lanes_per_direction = st.selectbox("‡∏à‡∏≥‡∏ô‡∏ß‡∏ô‡∏ä‡πà‡∏≠‡∏á‡∏ï‡πà‡∏≠‡∏ó‡∏¥‡∏®‡∏ó‡∏≤‡∏á (‡πÄ‡∏•‡∏ô/‡∏ó‡∏¥‡∏®‡∏ó‡∏≤‡∏á)", options=lanes_per_dir_options, index=lanes_per_dir_idx)
        
        # ‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡∏à‡∏≥‡∏ô‡∏ß‡∏ô‡∏ä‡πà‡∏≠‡∏á‡∏£‡∏ß‡∏° (‡∏Ñ‡∏π‡∏ì 2 ‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö 2 ‡∏ó‡∏¥‡∏®‡∏ó‡∏≤‡∏á)
        num_lanes = lanes_per_direction * 2
        
        shoulder_left = st.number_input("‡πÑ‡∏´‡∏•‡πà‡∏ó‡∏≤‡∏á‡∏ã‡πâ‡∏≤‡∏¢ (‡∏°.)", value=loaded_info.get('shoulder_left', 2.5), min_value=0.0, max_value=3.5, step=0.25)
        shoulder_right = st.number_input("‡πÑ‡∏´‡∏•‡πà‡∏ó‡∏≤‡∏á‡∏Ç‡∏ß‡∏≤ (‡∏°.)", value=loaded_info.get('shoulder_right', 1.5), min_value=0.0, max_value=3.5, step=0.25)
        
        # ‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡∏Ñ‡∏ß‡∏≤‡∏°‡∏Å‡∏ß‡πâ‡∏≤‡∏á‡∏£‡∏ß‡∏°
        # ‡∏Ñ‡∏ß‡∏≤‡∏°‡∏Å‡∏ß‡πâ‡∏≤‡∏á‡∏ú‡∏¥‡∏ß‡∏à‡∏£‡∏≤‡∏à‡∏£ = ‡∏ä‡πà‡∏≠‡∏á‡∏à‡∏£‡∏≤‡∏à‡∏£ √ó ‡∏à‡∏≥‡∏ô‡∏ß‡∏ô‡∏ä‡πà‡∏≠‡∏á (‡∏£‡∏ß‡∏° 2 ‡∏ó‡∏¥‡∏®‡∏ó‡∏≤‡∏á)
        # ‡πÅ‡∏ï‡πà‡∏•‡∏∞‡∏ó‡∏¥‡∏®‡∏ó‡∏≤‡∏á‡∏°‡∏µ‡πÑ‡∏´‡∏•‡πà‡∏ó‡∏≤‡∏á‡∏ã‡πâ‡∏≤‡∏¢ + ‡πÑ‡∏´‡∏•‡πà‡∏ó‡∏≤‡∏á‡∏Ç‡∏ß‡∏≤
        # ‡∏Ñ‡∏ß‡∏≤‡∏°‡∏Å‡∏ß‡πâ‡∏≤‡∏á‡∏£‡∏ß‡∏° = ‡∏ú‡∏¥‡∏ß‡∏à‡∏£‡∏≤‡∏à‡∏£ + (‡πÑ‡∏´‡∏•‡πà‡∏ó‡∏≤‡∏á‡∏ã‡πâ‡∏≤‡∏¢ + ‡πÑ‡∏´‡∏•‡πà‡∏ó‡∏≤‡∏á‡∏Ç‡∏ß‡∏≤) √ó 2 ‡∏ó‡∏¥‡∏®‡∏ó‡∏≤‡∏á
        road_surface_width = lane_width * num_lanes
        total_shoulders = (shoulder_left + shoulder_right) * 2  # ‡πÑ‡∏´‡∏•‡πà‡∏ó‡∏≤‡∏á‡∏£‡∏ß‡∏° 2 ‡∏ó‡∏¥‡∏®‡∏ó‡∏≤‡∏á
        total_width = road_surface_width + total_shoulders
        st.info(f"üìè ‡∏à‡∏≥‡∏ô‡∏ß‡∏ô‡∏ä‡πà‡∏≠‡∏á‡∏£‡∏ß‡∏° (2 ‡∏ó‡∏¥‡∏®‡∏ó‡∏≤‡∏á): {num_lanes} ‡∏ä‡πà‡∏≠‡∏á\nüìè ‡∏Ñ‡∏ß‡∏≤‡∏°‡∏Å‡∏ß‡πâ‡∏≤‡∏á‡∏ú‡∏¥‡∏ß‡∏à‡∏£‡∏≤‡∏à‡∏£: {road_surface_width:.2f} ‡∏°.\nüìè ‡∏Ñ‡∏ß‡∏≤‡∏°‡∏Å‡∏ß‡πâ‡∏≤‡∏á‡πÑ‡∏´‡∏•‡πà‡∏ó‡∏≤‡∏á (2 ‡∏ó‡∏¥‡∏®‡∏ó‡∏≤‡∏á): {total_shoulders:.2f} ‡∏°.\nüìè ‡∏Ñ‡∏ß‡∏≤‡∏°‡∏Å‡∏ß‡πâ‡∏≤‡∏á‡∏£‡∏ß‡∏°: {total_width:.2f} ‡∏°.")
    
    # ‡πÄ‡∏Å‡πá‡∏ö‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡πÇ‡∏Ñ‡∏£‡∏á‡∏Å‡∏≤‡∏£
    project_info = {
        'name': project_name,
        'length': road_length,
        'lane_width': lane_width,
        'shoulder_left': shoulder_left,
        'shoulder_right': shoulder_right,
        'num_lanes': num_lanes,
        'total_width': total_width
    }
    
    # ‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏µ‡πà‡∏ï‡πà‡∏≠ ‡∏Å‡∏°. (‡πÉ‡∏ä‡πâ‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡∏õ‡∏£‡∏¥‡∏°‡∏≤‡∏ì)
    area_per_km = total_width * 1000  # ‡∏ï‡∏£.‡∏°./‡∏Å‡∏°.
    
    # Tabs
    # Tabs - ‡∏•‡∏ö Tab 3-5 ‡∏≠‡∏≠‡∏Å‡πÅ‡∏•‡πâ‡∏ß
    tab1, tab2, tab3, tab4 = st.tabs([
        "üìä Library ‡∏£‡∏≤‡∏Ñ‡∏≤", 
        "üèóÔ∏è ‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á", 
        "üìÑ ‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô",
        "üì∑ ‡∏ß‡∏¥‡πÄ‡∏Ñ‡∏£‡∏≤‡∏∞‡∏´‡πå‡∏à‡∏≤‡∏Å‡∏£‡∏π‡∏õ‡∏†‡∏≤‡∏û"
    ])
    
    # ===== Tab 1: Library ‡∏£‡∏≤‡∏Ñ‡∏≤ =====
    with tab1:
        st.header("üìä ‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏£‡∏≤‡∏Ñ‡∏≤‡πÄ‡∏õ‡∏£‡∏µ‡∏¢‡∏ö‡πÄ‡∏ó‡∏µ‡∏¢‡∏ö‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á")
        st.info("üí° ‡∏™‡∏≤‡∏°‡∏≤‡∏£‡∏ñ‡∏õ‡∏£‡∏±‡∏ö‡πÄ‡∏õ‡∏•‡∏µ‡πà‡∏¢‡∏ô‡∏£‡∏≤‡∏Ñ‡∏≤‡πÑ‡∏î‡πâ‡∏ï‡∏≤‡∏°‡∏ï‡πâ‡∏≠‡∏á‡∏Å‡∏≤‡∏£ ‡∏´‡∏£‡∏∑‡∏≠ Upload ‡πÑ‡∏ü‡∏•‡πå Excel ‡πÉ‡∏ô **Sidebar** ‡πÄ‡∏û‡∏∑‡πà‡∏≠‡∏≠‡∏±‡∏û‡πÄ‡∏î‡∏ó‡∏£‡∏≤‡∏Ñ‡∏≤‡∏ó‡∏±‡πâ‡∏á‡∏´‡∏°‡∏î")
        
        # ===== Download Template =====
        st.subheader("üì• ‡∏î‡∏≤‡∏ß‡∏ô‡πå‡πÇ‡∏´‡∏•‡∏î Template Excel")
        
        col1, col2, col3 = st.columns([2, 1, 2])
        with col2:
            # ‡πÉ‡∏ä‡πâ cached template (‡πÑ‡∏°‡πà‡∏ï‡πâ‡∏≠‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡πÉ‡∏´‡∏°‡πà‡∏ó‡∏∏‡∏Å‡∏Ñ‡∏£‡∏±‡πâ‡∏á rerun)
            template_bytes = generate_excel_template()
            
            st.download_button(
                label="‚¨áÔ∏è ‡∏î‡∏≤‡∏ß‡∏ô‡πå‡πÇ‡∏´‡∏•‡∏î Template",
                data=template_bytes,
                file_name=f"Price_Library_Template_{datetime.now().strftime('%Y%m%d')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )
        
        st.caption("üìå ‡∏î‡∏≤‡∏ß‡∏ô‡πå‡πÇ‡∏´‡∏•‡∏î Template ‚Üí ‡πÅ‡∏Å‡πâ‡πÑ‡∏Ç‡∏£‡∏≤‡∏Ñ‡∏≤ ‚Üí Upload ‡πÉ‡∏ô Sidebar ‡∏î‡πâ‡∏≤‡∏ô‡∏ã‡πâ‡∏≤‡∏¢")
        
        st.divider()
        
        # ‡πÄ‡∏Å‡πá‡∏ö‡∏£‡∏≤‡∏Ñ‡∏≤‡πÉ‡∏ô session state
        # ‡∏ñ‡πâ‡∏≤‡∏°‡∏µ uploaded_price_library ‡πÉ‡∏´‡πâ‡πÉ‡∏ä‡πâ‡πÅ‡∏ó‡∏ô (‡∏ó‡∏∏‡∏Å‡∏Ñ‡∏£‡∏±‡πâ‡∏á‡∏ó‡∏µ‡πà rerun)
        if 'uploaded_price_library' in st.session_state:
            st.session_state['price_library'] = st.session_state['uploaded_price_library'].copy()
        elif 'price_library' not in st.session_state:
            # ‡πÑ‡∏°‡πà‡∏°‡∏µ‡∏ó‡∏±‡πâ‡∏á uploaded ‡πÅ‡∏•‡∏∞ price_library ‚Üí ‡πÉ‡∏ä‡πâ default
            st.session_state['price_library'] = {
                'ac_prices': {
                    'PMA Wearing Course': dict(AC_PRICE_TABLE['PMA Wearing Course']),
                    'AC Wearing Course': dict(AC_PRICE_TABLE['AC Wearing Course']),
                    'AC Binder Course': dict(AC_PRICE_TABLE['AC Binder Course']),
                    'AC Base Course': dict(AC_PRICE_TABLE['AC Base Course']),
                },
                'concrete_prices': {
                    'JRCP': dict(CONCRETE_PRICE_TABLE['JRCP']),
                    'JPCP': dict(CONCRETE_PRICE_TABLE['JPCP']),
                    'CRCP': dict(CONCRETE_PRICE_TABLE['CRCP']),
                },
                'base_prices': dict(BASE_MATERIAL_PRICES),
            }
        
        # Debug info - ‡πÅ‡∏™‡∏î‡∏á‡∏ó‡∏µ‡πà‡∏°‡∏≤‡∏Ç‡∏≠‡∏á‡∏£‡∏≤‡∏Ñ‡∏≤
        if 'uploaded_price_library' in st.session_state:
            st.info("üìã **‡∏Å‡∏≥‡∏•‡∏±‡∏á‡πÉ‡∏ä‡πâ‡∏£‡∏≤‡∏Ñ‡∏≤‡∏à‡∏≤‡∏Å‡πÑ‡∏ü‡∏•‡πå Excel ‡∏ó‡∏µ‡πà Upload ‡πÉ‡∏ô Sidebar**")
            # ‡πÅ‡∏™‡∏î‡∏á‡∏ï‡∏±‡∏ß‡∏≠‡∏¢‡πà‡∏≤‡∏á
            ac_7 = st.session_state['price_library']['ac_prices'].get('AC Wearing Course', {}).get(7.0, 'N/A')
            st.caption(f"‡∏ï‡∏±‡∏ß‡∏≠‡∏¢‡πà‡∏≤‡∏á: AC Wearing Course 7cm = {ac_7} ‡∏ö‡∏≤‡∏ó")
        else:
            st.caption("üí° ‡∏Å‡∏≥‡∏•‡∏±‡∏á‡πÉ‡∏ä‡πâ‡∏£‡∏≤‡∏Ñ‡∏≤ Default (Upload Excel ‡πÉ‡∏ô Sidebar ‡πÄ‡∏û‡∏∑‡πà‡∏≠‡πÄ‡∏õ‡∏•‡∏µ‡πà‡∏¢‡∏ô‡∏£‡∏≤‡∏Ñ‡∏≤)")
        
        # ===== ‡∏™‡πà‡∏ß‡∏ô‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á AC =====
        st.subheader("üîµ ‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á Asphalt Concrete (‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.)")
        
        # ‡πÉ‡∏ä‡πâ version ‡∏à‡∏≤‡∏Å upload ‡πÄ‡∏û‡∏∑‡πà‡∏≠ force refresh widgets
        upload_version = st.session_state.get('price_upload_version', 'default')
        
        ac_cols = st.columns(4)
        ac_types = ['PMA Wearing Course', 'AC Wearing Course', 'AC Binder Course', 'AC Base Course']
        thicknesses = [2.5, 3, 4, 5, 6, 7, 8, 9, 10]
        
        for col_idx, ac_type in enumerate(ac_types):
            with ac_cols[col_idx]:
                st.markdown(f"**{ac_type}**")
                for thk in thicknesses:
                    # ‡∏î‡∏∂‡∏á‡∏£‡∏≤‡∏Ñ‡∏≤‡∏à‡∏≤‡∏Å session_state (‡∏ñ‡πâ‡∏≤ upload ‡∏°‡∏≤‡∏à‡∏∞‡πÄ‡∏õ‡πá‡∏ô‡∏£‡∏≤‡∏Ñ‡∏≤‡πÉ‡∏´‡∏°‡πà)
                    current_price = st.session_state['price_library']['ac_prices'][ac_type].get(thk, 0)
                    price = st.number_input(
                        f"{thk} cm", 
                        value=float(current_price),
                        key=f"ac_{ac_type}_{thk}_{upload_version}",
                        step=10.0,
                        label_visibility="visible"
                    )
                    st.session_state['price_library']['ac_prices'][ac_type][thk] = price
        
        st.divider()
        
        # ===== ‡∏™‡πà‡∏ß‡∏ô‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï =====
        st.subheader("üü† ‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï (‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.)")
        
        conc_cols = st.columns(3)
        conc_types = ['JRCP', 'JPCP', 'CRCP']
        conc_thicknesses = [25, 28, 32, 35]
        
        for col_idx, conc_type in enumerate(conc_types):
            with conc_cols[col_idx]:
                st.markdown(f"**{conc_type}**")
                for thk in conc_thicknesses:
                    # ‡∏î‡∏∂‡∏á‡∏£‡∏≤‡∏Ñ‡∏≤‡∏à‡∏≤‡∏Å session_state
                    current_price = st.session_state['price_library']['concrete_prices'][conc_type].get(thk, 0)
                    price = st.number_input(
                        f"{thk} cm", 
                        value=float(current_price),
                        key=f"conc_{conc_type}_{thk}_{upload_version}",
                        step=10.0
                    )
                    st.session_state['price_library']['concrete_prices'][conc_type][thk] = price
                
                # ‡∏£‡∏≤‡∏Ñ‡∏≤‡πÑ‡∏°‡πà‡∏£‡∏ß‡∏° Joint
                st.markdown("---")
                excl_price = st.number_input(
                    f"{conc_type} (excl. Joint)",
                    value=float(CONCRETE_EXCL_JOINT[conc_type]),
                    key=f"conc_excl_{conc_type}_{upload_version}",
                    step=10.0
                )
        
        st.divider()
        
        # ===== ‡∏™‡πà‡∏ß‡∏ô‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á/‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á =====
        st.subheader("üü§ ‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á/‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á (‡∏ö‡∏≤‡∏ó/‡∏•‡∏ö.‡∏°.)")
        
        base_cols = st.columns(3)
        base_materials_list = list(BASE_MATERIAL_PRICES.keys())
        
        for i, mat in enumerate(base_materials_list):
            with base_cols[i % 3]:
                # ‡∏î‡∏∂‡∏á‡∏£‡∏≤‡∏Ñ‡∏≤‡∏à‡∏≤‡∏Å session_state
                current_price = st.session_state['price_library']['base_prices'].get(mat, 0)
                price = st.number_input(
                    mat,
                    value=float(current_price),
                    key=f"base_{mat}_{upload_version}",
                    step=10.0
                )
                st.session_state['price_library']['base_prices'][mat] = price
        
        # ===== ‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Å‡∏≥‡∏´‡∏ô‡∏î‡πÄ‡∏≠‡∏á =====
        st.markdown("---")
        st.markdown("**‚ú® ‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Å‡∏≥‡∏´‡∏ô‡∏î‡πÄ‡∏≠‡∏á** (‡πÄ‡∏û‡∏¥‡πà‡∏°‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Ç‡∏≠‡∏á‡∏Ñ‡∏∏‡∏ì‡πÄ‡∏≠‡∏á)")
        
        custom_cols = st.columns(3)
        
        for i in range(1, 4):  # ‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Å‡∏≥‡∏´‡∏ô‡∏î‡πÄ‡∏≠‡∏á 1, 2, 3
            with custom_cols[i-1]:
                # ‡∏Å‡∏≥‡∏´‡∏ô‡∏î key ‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏Å‡∏≥‡∏´‡∏ô‡∏î‡πÄ‡∏≠‡∏á
                custom_key = f"custom_material_{i}"
                
                # ‡∏î‡∏∂‡∏á‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏à‡∏≤‡∏Å session_state ‡∏ñ‡πâ‡∏≤‡∏°‡∏µ
                if 'custom_materials' not in st.session_state:
                    st.session_state['custom_materials'] = {}
                
                existing_data = st.session_state['custom_materials'].get(custom_key, {'name': '', 'price': 0.0})
                
                # ‡∏ä‡∏∑‡πà‡∏≠‡∏ß‡∏±‡∏™‡∏î‡∏∏
                material_name = st.text_input(
                    f"‡∏ä‡∏∑‡πà‡∏≠‡∏ß‡∏±‡∏™‡∏î‡∏∏ {i}",
                    value=existing_data['name'],
                    key=f"custom_name_{i}_{upload_version}",
                    placeholder=f"‡∏£‡∏∞‡∏ö‡∏∏‡∏ä‡∏∑‡πà‡∏≠‡∏ß‡∏±‡∏™‡∏î‡∏∏ {i}..."
                )
                
                # ‡∏£‡∏≤‡∏Ñ‡∏≤ (‡πÅ‡∏™‡∏î‡∏á‡πÄ‡∏â‡∏û‡∏≤‡∏∞‡πÄ‡∏°‡∏∑‡πà‡∏≠‡∏°‡∏µ‡∏ä‡∏∑‡πà‡∏≠)
                if material_name:
                    material_price = st.number_input(
                        f"‡∏£‡∏≤‡∏Ñ‡∏≤ (‡∏ö‡∏≤‡∏ó/‡∏•‡∏ö.‡∏°.)",
                        value=float(existing_data['price']),
                        key=f"custom_price_{i}_{upload_version}",
                        step=10.0,
                        min_value=0.0
                    )
                    
                    # ‡∏ö‡∏±‡∏ô‡∏ó‡∏∂‡∏Å‡πÉ‡∏ô session_state
                    st.session_state['custom_materials'][custom_key] = {
                        'name': material_name,
                        'price': material_price
                    }
                    
                    # ‡πÄ‡∏û‡∏¥‡πà‡∏°‡πÄ‡∏Ç‡πâ‡∏≤ price_library ‡∏î‡πâ‡∏ß‡∏¢ (‡πÄ‡∏û‡∏∑‡πà‡∏≠‡πÉ‡∏ä‡πâ‡πÉ‡∏ô Tab 2)
                    st.session_state['price_library']['base_prices'][material_name] = material_price
                else:
                    # ‡∏ñ‡πâ‡∏≤‡πÑ‡∏°‡πà‡∏°‡∏µ‡∏ä‡∏∑‡πà‡∏≠ ‡∏•‡∏ö‡∏≠‡∏≠‡∏Å‡∏à‡∏≤‡∏Å custom_materials
                    if custom_key in st.session_state['custom_materials']:
                        old_name = st.session_state['custom_materials'][custom_key]['name']
                        if old_name in st.session_state['price_library']['base_prices']:
                            del st.session_state['price_library']['base_prices'][old_name]
                        del st.session_state['custom_materials'][custom_key]
        
        st.divider()
        
        # ===== ‡∏õ‡∏∏‡πà‡∏°‡∏î‡∏≤‡∏ß‡∏ô‡πå‡πÇ‡∏´‡∏•‡∏î =====
        st.subheader("üì• ‡∏î‡∏≤‡∏ß‡∏ô‡πå‡πÇ‡∏´‡∏•‡∏î‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏£‡∏≤‡∏Ñ‡∏≤")
        
        col_dl1, col_dl2 = st.columns(2)
        
        with col_dl1:
            if st.button("üìä ‡∏™‡∏£‡πâ‡∏≤‡∏á‡πÑ‡∏ü‡∏•‡πå Excel", key="btn_excel_price", use_container_width=True):
                # ‡∏™‡∏£‡πâ‡∏≤‡∏á Excel
                output = io.BytesIO()
                with pd.ExcelWriter(output, engine='openpyxl') as writer:
                    # Sheet 1: AC Prices
                    ac_data = []
                    for ac_type in ac_types:
                        for thk in thicknesses:
                            ac_data.append({
                                '‡∏õ‡∏£‡∏∞‡πÄ‡∏†‡∏ó': ac_type,
                                '‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤ (cm)': thk,
                                '‡∏£‡∏≤‡∏Ñ‡∏≤ (‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.)': st.session_state['price_library']['ac_prices'][ac_type][thk]
                            })
                    pd.DataFrame(ac_data).to_excel(writer, sheet_name='AC Prices', index=False)
                    
                    # Sheet 2: Concrete Prices
                    conc_data = []
                    for conc_type in conc_types:
                        for thk in conc_thicknesses:
                            conc_data.append({
                                '‡∏õ‡∏£‡∏∞‡πÄ‡∏†‡∏ó': conc_type,
                                '‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤ (cm)': thk,
                                '‡∏£‡∏≤‡∏Ñ‡∏≤ (‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.)': st.session_state['price_library']['concrete_prices'][conc_type][thk]
                            })
                    pd.DataFrame(conc_data).to_excel(writer, sheet_name='Concrete Prices', index=False)
                    
                    # Sheet 3: Base Material Prices
                    base_data = [{'‡∏ß‡∏±‡∏™‡∏î‡∏∏': k, '‡∏£‡∏≤‡∏Ñ‡∏≤ (‡∏ö‡∏≤‡∏ó/‡∏•‡∏ö.‡∏°.)': v} for k, v in st.session_state['price_library']['base_prices'].items()]
                    pd.DataFrame(base_data).to_excel(writer, sheet_name='Base Materials', index=False)
                
                output.seek(0)
                st.download_button(
                    label="‚¨áÔ∏è Download Excel",
                    data=output,
                    file_name="‡∏£‡∏≤‡∏Ñ‡∏≤‡πÄ‡∏õ‡∏£‡∏µ‡∏¢‡∏ö‡πÄ‡∏ó‡∏µ‡∏¢‡∏ö‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
        
        with col_dl2:
            if st.button("üìÑ ‡∏™‡∏£‡πâ‡∏≤‡∏á‡πÑ‡∏ü‡∏•‡πå Word", key="btn_word_price", use_container_width=True):
                doc = Document()
                doc.add_heading('‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏£‡∏≤‡∏Ñ‡∏≤‡πÄ‡∏õ‡∏£‡∏µ‡∏¢‡∏ö‡πÄ‡∏ó‡∏µ‡∏¢‡∏ö‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á', 0)
                
                # AC Table
                doc.add_heading('1. ‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á Asphalt Concrete (‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.)', level=1)
                table = doc.add_table(rows=len(thicknesses)+1, cols=5)
                table.style = 'Table Grid'
                headers = ['‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤ (cm)'] + ac_types
                for j, h in enumerate(headers):
                    table.rows[0].cells[j].text = h
                for i, thk in enumerate(thicknesses):
                    table.rows[i+1].cells[0].text = str(thk)
                    for j, ac_type in enumerate(ac_types):
                        table.rows[i+1].cells[j+1].text = f"{st.session_state['price_library']['ac_prices'][ac_type][thk]:,.0f}"
                
                # Concrete Table
                doc.add_heading('2. ‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï (‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.)', level=1)
                table = doc.add_table(rows=len(conc_thicknesses)+1, cols=4)
                table.style = 'Table Grid'
                headers = ['‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤ (cm)'] + conc_types
                for j, h in enumerate(headers):
                    table.rows[0].cells[j].text = h
                for i, thk in enumerate(conc_thicknesses):
                    table.rows[i+1].cells[0].text = str(thk)
                    for j, conc_type in enumerate(conc_types):
                        table.rows[i+1].cells[j+1].text = f"{st.session_state['price_library']['concrete_prices'][conc_type][thk]:,.0f}"
                
                # Base Material Table
                doc.add_heading('3. ‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á/‡∏£‡∏≠‡∏á‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á (‡∏ö‡∏≤‡∏ó/‡∏•‡∏ö.‡∏°.)', level=1)
                table = doc.add_table(rows=len(base_materials_list)+1, cols=2)
                table.style = 'Table Grid'
                table.rows[0].cells[0].text = '‡∏ß‡∏±‡∏™‡∏î‡∏∏'
                table.rows[0].cells[1].text = '‡∏£‡∏≤‡∏Ñ‡∏≤ (‡∏ö‡∏≤‡∏ó/‡∏•‡∏ö.‡∏°.)'
                for i, mat in enumerate(base_materials_list):
                    table.rows[i+1].cells[0].text = mat
                    table.rows[i+1].cells[1].text = f"{st.session_state['price_library']['base_prices'][mat]:,.0f}"
                
                doc_output = io.BytesIO()
                doc.save(doc_output)
                doc_output.seek(0)
                st.download_button(
                    label="‚¨áÔ∏è Download Word",
                    data=doc_output,
                    file_name="‡∏£‡∏≤‡∏Ñ‡∏≤‡πÄ‡∏õ‡∏£‡∏µ‡∏¢‡∏ö‡πÄ‡∏ó‡∏µ‡∏¢‡∏ö‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    
    # ===== Tab 2: ‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á =====
    # ===== Tab 2: ‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á =====
    with tab2:
        st.header("‡∏Å‡∏≥‡∏´‡∏ô‡∏î‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á")
        st.info("üí° ‡πÅ‡∏Å‡πâ‡πÑ‡∏Ç‡∏ä‡∏∑‡πà‡∏≠ ‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤ ‡πÅ‡∏•‡∏∞‡∏£‡∏≤‡∏Ñ‡∏≤‡∏ï‡πà‡∏≠‡∏´‡∏ô‡πà‡∏ß‡∏¢‡πÑ‡∏î‡πâ‡∏ï‡∏≤‡∏°‡∏ï‡πâ‡∏≠‡∏á‡∏Å‡∏≤‡∏£ | ‚úÖ ‡πÄ‡∏•‡∏∑‡∏≠‡∏Å‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ó‡∏µ‡πà‡∏ï‡πâ‡∏≠‡∏á‡∏Å‡∏≤‡∏£‡πÅ‡∏™‡∏î‡∏á‡πÉ‡∏ô‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô")
        
        # version ‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö widget keys ‚Äî ‡πÄ‡∏û‡∏¥‡πà‡∏°‡∏ó‡∏∏‡∏Å‡∏Ñ‡∏£‡∏±‡πâ‡∏á‡∏ó‡∏µ‡πà load JSON ‡πÉ‡∏´‡∏°‡πà
        v = st.session_state.get('json_version', 0)
        
        # ‡πÅ‡∏à‡πâ‡∏á‡πÄ‡∏ï‡∏∑‡∏≠‡∏ô‡πÄ‡∏°‡∏∑‡πà‡∏≠‡πÇ‡∏´‡∏•‡∏î JSON
        if st.session_state.get('loaded_project'):
            loaded_name = st.session_state['loaded_project'].get('project_info', {}).get('name', '-')
            st.success(f"‚úÖ ‡∏Å‡∏≥‡∏•‡∏±‡∏á‡πÉ‡∏ä‡πâ‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏à‡∏≤‡∏Å: **{loaded_name}**")
        
        # ‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏µ‡πà‡∏ï‡πà‡∏≠ ‡∏Å‡∏°.
        # total_width ‡∏£‡∏ß‡∏°‡∏ó‡∏±‡πâ‡∏á 2 ‡∏ó‡∏¥‡∏®‡∏ó‡∏≤‡∏á‡πÑ‡∏ß‡πâ‡πÅ‡∏•‡πâ‡∏ß (num_lanes = lanes_per_direction * 2)
        area_per_km = total_width * 1000  # ‡∏ï‡∏£.‡∏°./‡∏Å‡∏°.
        
        # ===== AC Pavement =====
        st.subheader("üîµ ‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡πÅ‡∏≠‡∏™‡∏ü‡∏±‡∏•‡∏ï‡πå‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï (AC)")
        col1, col2 = st.columns(2)
        
        with col1:
            ac1_show = st.checkbox("‡πÅ‡∏™‡∏î‡∏á‡πÉ‡∏ô‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô", value=True, key="ac1_show")
            ac1_name = st.text_input("‡∏ä‡∏∑‡πà‡∏≠‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á AC1", value="AC1: ‡πÅ‡∏≠‡∏™‡∏ü‡∏±‡∏•‡∏ï‡πå‡∏ö‡∏ô‡∏´‡∏¥‡∏ô‡∏Ñ‡∏•‡∏∏‡∏Å", key="ac1_name")
            with st.expander(f"‚óè {ac1_name}", expanded=True):
                ac1_layers = render_layer_editor(get_default_ac1_layers(), "ac1", total_width, road_length, v=v)
                ac1_cost, ac1_details = calculate_layer_cost(ac1_layers, road_length)
                ac1_cost_per_km = ac1_cost / road_length / 1_000_000
                ac1_cost_per_sqm = ac1_cost / (area_per_km * road_length)
                st.markdown(f'<div class="cost-box">üí∞ <b>‡∏Ñ‡πà‡∏≤‡∏Å‡πà‡∏≠‡∏™‡∏£‡πâ‡∏≤‡∏á:</b> {ac1_cost_per_km:.2f} ‡∏•‡πâ‡∏≤‡∏ô‡∏ö‡∏≤‡∏ó/‡∏Å‡∏°.</div>', unsafe_allow_html=True)
                st.markdown(f'<div class="cost-box">üí∞ <b>‡∏Ñ‡πà‡∏≤‡∏Å‡πà‡∏≠‡∏™‡∏£‡πâ‡∏≤‡∏á:</b> {ac1_cost_per_sqm:.2f} ‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.</div>', unsafe_allow_html=True)
        
        with col2:
            ac2_show = st.checkbox("‡πÅ‡∏™‡∏î‡∏á‡πÉ‡∏ô‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô", value=True, key="ac2_show")
            ac2_name = st.text_input("‡∏ä‡∏∑‡πà‡∏≠‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á AC2", value="AC2: ‡πÅ‡∏≠‡∏™‡∏ü‡∏±‡∏•‡∏ï‡πå‡∏ö‡∏ô‡∏´‡∏¥‡∏ô‡∏Ñ‡∏•‡∏∏‡∏Å‡∏ú‡∏™‡∏°‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå", key="ac2_name")
            with st.expander(f"‚óè {ac2_name}", expanded=True):
                ac2_layers = render_layer_editor(get_default_ac2_layers(), "ac2", total_width, road_length, v=v)
                ac2_cost, ac2_details = calculate_layer_cost(ac2_layers, road_length)
                ac2_cost_per_km = ac2_cost / road_length / 1_000_000
                ac2_cost_per_sqm = ac2_cost / (area_per_km * road_length)
                st.markdown(f'<div class="cost-box">üí∞ <b>‡∏Ñ‡πà‡∏≤‡∏Å‡πà‡∏≠‡∏™‡∏£‡πâ‡∏≤‡∏á:</b> {ac2_cost_per_km:.2f} ‡∏•‡πâ‡∏≤‡∏ô‡∏ö‡∏≤‡∏ó/‡∏Å‡∏°.</div>', unsafe_allow_html=True)
                st.markdown(f'<div class="cost-box">üí∞ <b>‡∏Ñ‡πà‡∏≤‡∏Å‡πà‡∏≠‡∏™‡∏£‡πâ‡∏≤‡∏á:</b> {ac2_cost_per_sqm:.2f} ‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.</div>', unsafe_allow_html=True)
        
        # ===== JRCP/JPCP =====
        st.subheader("üü† ‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï‡πÄ‡∏™‡∏£‡∏¥‡∏°‡πÄ‡∏´‡∏•‡πá‡∏Å (JRCP/JPCP)")
        col3, col4 = st.columns(2)
        
        with col3:
            jrcp1_show = st.checkbox("‡πÅ‡∏™‡∏î‡∏á‡πÉ‡∏ô‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô", value=True, key="jrcp1_show")
            jrcp1_name = st.text_input("‡∏ä‡∏∑‡πà‡∏≠‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á JPCP/JRCP (1)", value="JPCP/JRCP (1): ‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï‡∏ö‡∏ô‡∏î‡∏¥‡∏ô‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå", key="jrcp1_name")
            with st.expander(f"‚óè {jrcp1_name}", expanded=True):
                jrcp1_layers = render_layer_editor(get_default_jrcp1_layers(), "jrcp1", total_width, road_length, v=v)
                jrcp1_layer_cost, jrcp1_layer_details = calculate_layer_cost(jrcp1_layers, road_length)
                jrcp1_joints, jrcp1_include_joints = render_joint_editor(get_default_jrcp1_joints(), "jrcp1", area_per_km, road_length, v=v)
                jrcp1_joint_cost, jrcp1_joint_details = calculate_joint_cost(jrcp1_joints, road_length)
                
                # ‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì ‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°. ‡∏ï‡∏≤‡∏° checkbox
                jrcp1_joints_sqm = sum(j.get('cost_per_sqm', 0) for j in jrcp1_joints)
                if jrcp1_include_joints:
                    jrcp1_total = jrcp1_layer_cost + jrcp1_joint_cost
                    jrcp1_cost_per_sqm = jrcp1_layer_cost / (area_per_km * road_length) + jrcp1_joints_sqm
                    joints_note = "(‡∏£‡∏ß‡∏° Joints)"
                else:
                    jrcp1_total = jrcp1_layer_cost + jrcp1_joint_cost  # ‡∏•‡πâ‡∏≤‡∏ô‡∏ö‡∏≤‡∏ó/‡∏Å‡∏°. ‡∏¢‡∏±‡∏á‡∏£‡∏ß‡∏° Joints
                    jrcp1_cost_per_sqm = jrcp1_layer_cost / (area_per_km * road_length)
                    joints_note = "(‡πÑ‡∏°‡πà‡∏£‡∏ß‡∏° Joints)"
                
                jrcp1_cost_per_km = jrcp1_total / road_length / 1_000_000
                jrcp1_details = jrcp1_layer_details + jrcp1_joint_details
                st.markdown(f'<div class="cost-box">üí∞ <b>‡∏Ñ‡πà‡∏≤‡∏Å‡πà‡∏≠‡∏™‡∏£‡πâ‡∏≤‡∏á:</b> {jrcp1_cost_per_km:.2f} ‡∏•‡πâ‡∏≤‡∏ô‡∏ö‡∏≤‡∏ó/‡∏Å‡∏°.</div>', unsafe_allow_html=True)
                st.markdown(f'<div class="cost-box">üí∞ <b>‡∏Ñ‡πà‡∏≤‡∏Å‡πà‡∏≠‡∏™‡∏£‡πâ‡∏≤‡∏á:</b> {jrcp1_cost_per_sqm:.2f} ‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°. {joints_note}</div>', unsafe_allow_html=True)
        
        with col4:
            jrcp2_show = st.checkbox("‡πÅ‡∏™‡∏î‡∏á‡πÉ‡∏ô‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô", value=True, key="jrcp2_show")
            jrcp2_name = st.text_input("‡∏ä‡∏∑‡πà‡∏≠‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á JPCP/JRCP (2)", value="JPCP/JRCP (2): ‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï‡∏ö‡∏ô‡∏´‡∏¥‡∏ô‡∏Ñ‡∏•‡∏∏‡∏Å‡∏ú‡∏™‡∏°‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå", key="jrcp2_name")
            with st.expander(f"‚óè {jrcp2_name}", expanded=True):
                jrcp2_layers = render_layer_editor(get_default_jrcp2_layers(), "jrcp2", total_width, road_length, v=v)
                jrcp2_layer_cost, jrcp2_layer_details = calculate_layer_cost(jrcp2_layers, road_length)
                jrcp2_joints, jrcp2_include_joints = render_joint_editor(get_default_jrcp2_joints(), "jrcp2", area_per_km, road_length, v=v)
                jrcp2_joint_cost, jrcp2_joint_details = calculate_joint_cost(jrcp2_joints, road_length)
                
                # ‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì ‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°. ‡∏ï‡∏≤‡∏° checkbox
                jrcp2_joints_sqm = sum(j.get('cost_per_sqm', 0) for j in jrcp2_joints)
                if jrcp2_include_joints:
                    jrcp2_total = jrcp2_layer_cost + jrcp2_joint_cost
                    jrcp2_cost_per_sqm = jrcp2_layer_cost / (area_per_km * road_length) + jrcp2_joints_sqm
                    joints_note2 = "(‡∏£‡∏ß‡∏° Joints)"
                else:
                    jrcp2_total = jrcp2_layer_cost + jrcp2_joint_cost
                    jrcp2_cost_per_sqm = jrcp2_layer_cost / (area_per_km * road_length)
                    joints_note2 = "(‡πÑ‡∏°‡πà‡∏£‡∏ß‡∏° Joints)"
                
                jrcp2_cost_per_km = jrcp2_total / road_length / 1_000_000
                jrcp2_details = jrcp2_layer_details + jrcp2_joint_details
                st.markdown(f'<div class="cost-box">üí∞ <b>‡∏Ñ‡πà‡∏≤‡∏Å‡πà‡∏≠‡∏™‡∏£‡πâ‡∏≤‡∏á:</b> {jrcp2_cost_per_km:.2f} ‡∏•‡πâ‡∏≤‡∏ô‡∏ö‡∏≤‡∏ó/‡∏Å‡∏°.</div>', unsafe_allow_html=True)
                st.markdown(f'<div class="cost-box">üí∞ <b>‡∏Ñ‡πà‡∏≤‡∏Å‡πà‡∏≠‡∏™‡∏£‡πâ‡∏≤‡∏á:</b> {jrcp2_cost_per_sqm:.2f} ‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°. {joints_note2}</div>', unsafe_allow_html=True)
        
        # ===== CRCP =====
        st.subheader("üî¥ ‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï‡πÄ‡∏™‡∏£‡∏¥‡∏°‡πÄ‡∏´‡∏•‡πá‡∏Å‡∏ï‡πà‡∏≠‡πÄ‡∏ô‡∏∑‡πà‡∏≠‡∏á (CRCP)")
        col5, col6 = st.columns(2)
        
        with col5:
            crcp1_show = st.checkbox("‡πÅ‡∏™‡∏î‡∏á‡πÉ‡∏ô‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô", value=True, key="crcp1_show")
            crcp1_name = st.text_input("‡∏ä‡∏∑‡πà‡∏≠‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á CRCP1", value="CRCP1: ‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï‡πÄ‡∏™‡∏£‡∏¥‡∏°‡πÄ‡∏´‡∏•‡πá‡∏Å‡∏ï‡πà‡∏≠‡πÄ‡∏ô‡∏∑‡πà‡∏≠‡∏á‡∏ö‡∏ô‡∏î‡∏¥‡∏ô‡∏ã‡∏µ‡πÄ‡∏°‡∏ô‡∏ï‡πå", key="crcp1_name")
            with st.expander(f"‚óè {crcp1_name}", expanded=True):
                crcp1_layers = render_layer_editor(get_default_crcp1_layers(), "crcp1", total_width, road_length, v=v)
                crcp1_cost, crcp1_details = calculate_layer_cost(crcp1_layers, road_length)
                crcp1_cost_per_km = crcp1_cost / road_length / 1_000_000
                crcp1_cost_per_sqm = crcp1_cost / (area_per_km * road_length)
                st.markdown(f'<div class="cost-box">üí∞ <b>‡∏Ñ‡πà‡∏≤‡∏Å‡πà‡∏≠‡∏™‡∏£‡πâ‡∏≤‡∏á:</b> {crcp1_cost_per_km:.2f} ‡∏•‡πâ‡∏≤‡∏ô‡∏ö‡∏≤‡∏ó/‡∏Å‡∏°.</div>', unsafe_allow_html=True)
                st.markdown(f'<div class="cost-box">üí∞ <b>‡∏Ñ‡πà‡∏≤‡∏Å‡πà‡∏≠‡∏™‡∏£‡πâ‡∏≤‡∏á:</b> {crcp1_cost_per_sqm:.2f} ‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.</div>', unsafe_allow_html=True)
        
        with col6:
            crcp2_show = st.checkbox("‡πÅ‡∏™‡∏î‡∏á‡πÉ‡∏ô‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô", value=True, key="crcp2_show")
            crcp2_name = st.text_input("‡∏ä‡∏∑‡πà‡∏≠‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á CRCP2", value="CRCP2: ‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï‡πÄ‡∏™‡∏£‡∏¥‡∏°‡πÄ‡∏´‡∏•‡πá‡∏Å‡∏ï‡πà‡∏≠‡πÄ‡∏ô‡∏∑‡πà‡∏≠‡∏á‡∏ö‡∏ô CMCR", key="crcp2_name")
            with st.expander(f"‚óè {crcp2_name}", expanded=True):
                crcp2_layers = render_layer_editor(get_default_crcp2_layers(), "crcp2", total_width, road_length, v=v)
                crcp2_cost, crcp2_details = calculate_layer_cost(crcp2_layers, road_length)
                crcp2_cost_per_km = crcp2_cost / road_length / 1_000_000
                crcp2_cost_per_sqm = crcp2_cost / (area_per_km * road_length)
                st.markdown(f'<div class="cost-box">üí∞ <b>‡∏Ñ‡πà‡∏≤‡∏Å‡πà‡∏≠‡∏™‡∏£‡πâ‡∏≤‡∏á:</b> {crcp2_cost_per_km:.2f} ‡∏•‡πâ‡∏≤‡∏ô‡∏ö‡∏≤‡∏ó/‡∏Å‡∏°.</div>', unsafe_allow_html=True)
                st.markdown(f'<div class="cost-box">üí∞ <b>‡∏Ñ‡πà‡∏≤‡∏Å‡πà‡∏≠‡∏™‡∏£‡πâ‡∏≤‡∏á:</b> {crcp2_cost_per_sqm:.2f} ‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.</div>', unsafe_allow_html=True)
        
        # Store in session state
        st.session_state['construction'] = {
            'AC1': {'name': ac1_name, 'cost': ac1_cost_per_km, 'cost_sqm': ac1_cost_per_sqm, 'details': ac1_details, 'layers': ac1_layers, 'joints': None, 'show': ac1_show},
            'AC2': {'name': ac2_name, 'cost': ac2_cost_per_km, 'cost_sqm': ac2_cost_per_sqm, 'details': ac2_details, 'layers': ac2_layers, 'joints': None, 'show': ac2_show},
            'JRCP1': {'name': jrcp1_name, 'cost': jrcp1_cost_per_km, 'cost_sqm': jrcp1_cost_per_sqm, 'details': jrcp1_details, 'layers': jrcp1_layers, 'joints': jrcp1_joints, 'show': jrcp1_show},
            'JRCP2': {'name': jrcp2_name, 'cost': jrcp2_cost_per_km, 'cost_sqm': jrcp2_cost_per_sqm, 'details': jrcp2_details, 'layers': jrcp2_layers, 'joints': jrcp2_joints, 'show': jrcp2_show},
            'CRCP1': {'name': crcp1_name, 'cost': crcp1_cost_per_km, 'cost_sqm': crcp1_cost_per_sqm, 'details': crcp1_details, 'layers': crcp1_layers, 'joints': None, 'show': crcp1_show},
            'CRCP2': {'name': crcp2_name, 'cost': crcp2_cost_per_km, 'cost_sqm': crcp2_cost_per_sqm, 'details': crcp2_details, 'layers': crcp2_layers, 'joints': None, 'show': crcp2_show},
        }
        st.session_state['project_info'] = project_info
        st.session_state['area_per_km'] = area_per_km
        
        # ===== Summary Tables =====
        st.divider()
        st.subheader("üìä ‡∏™‡∏£‡∏∏‡∏õ‡∏Ñ‡πà‡∏≤‡∏Å‡πà‡∏≠‡∏™‡∏£‡πâ‡∏≤‡∏á")
        
        # ‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏™‡∏£‡∏∏‡∏õ‡∏£‡∏ß‡∏°
        all_structures = [
            ('AC1', ac1_name, ac1_cost_per_km, ac1_cost_per_sqm, 20, ac1_show),
            ('AC2', ac2_name, ac2_cost_per_km, ac2_cost_per_sqm, 20, ac2_show),
            ('JRCP1', jrcp1_name, jrcp1_cost_per_km, jrcp1_cost_per_sqm, 25, jrcp1_show),
            ('JRCP2', jrcp2_name, jrcp2_cost_per_km, jrcp2_cost_per_sqm, 25, jrcp2_show),
            ('CRCP1', crcp1_name, crcp1_cost_per_km, crcp1_cost_per_sqm, 30, crcp1_show),
            ('CRCP2', crcp2_name, crcp2_cost_per_km, crcp2_cost_per_sqm, 30, crcp2_show),
        ]
        
        summary_data = []
        for key, name, cost_km, cost_sqm, life, show in all_structures:
            summary_data.append({
                '‡∏£‡∏´‡∏±‡∏™': key,
                '‡∏õ‡∏£‡∏∞‡πÄ‡∏†‡∏ó': name,
                '‡∏Ñ‡πà‡∏≤‡∏Å‡πà‡∏≠‡∏™‡∏£‡πâ‡∏≤‡∏á (‡∏•‡πâ‡∏≤‡∏ô‡∏ö‡∏≤‡∏ó/‡∏Å‡∏°.)': cost_km,
                '‡∏Ñ‡πà‡∏≤‡∏Å‡πà‡∏≠‡∏™‡∏£‡πâ‡∏≤‡∏á (‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.)': cost_sqm,
                '‡∏≠‡∏≤‡∏¢‡∏∏‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö (‡∏õ‡∏µ)': life,
                '‡πÅ‡∏™‡∏î‡∏á‡πÉ‡∏ô‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô': '‚úÖ' if show else '‚ùå'
            })
        
        summary_df = pd.DataFrame(summary_data)
        st.dataframe(
            summary_df.style.format({
                '‡∏Ñ‡πà‡∏≤‡∏Å‡πà‡∏≠‡∏™‡∏£‡πâ‡∏≤‡∏á (‡∏•‡πâ‡∏≤‡∏ô‡∏ö‡∏≤‡∏ó/‡∏Å‡∏°.)': '{:.2f}',
                '‡∏Ñ‡πà‡∏≤‡∏Å‡πà‡∏≠‡∏™‡∏£‡πâ‡∏≤‡∏á (‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.)': '{:.2f}'
            }),
            use_container_width=True,
            hide_index=True
        )
        
        # ===== ‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏™‡∏£‡∏∏‡∏õ‡∏£‡∏≤‡∏Ñ‡∏≤‡∏•‡∏∞‡πÄ‡∏≠‡∏µ‡∏¢‡∏î‡πÅ‡∏ï‡πà‡∏•‡∏∞‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á =====
        st.divider()
        st.subheader("üìã ‡∏£‡∏≤‡∏¢‡∏•‡∏∞‡πÄ‡∏≠‡∏µ‡∏¢‡∏î‡∏£‡∏≤‡∏Ñ‡∏≤‡πÅ‡∏ï‡πà‡∏•‡∏∞‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á")
        
        selected_structure = st.selectbox(
            "‡πÄ‡∏•‡∏∑‡∏≠‡∏Å‡∏î‡∏π‡∏£‡∏≤‡∏¢‡∏•‡∏∞‡πÄ‡∏≠‡∏µ‡∏¢‡∏î",
            options=['AC1', 'AC2', 'JRCP1', 'JRCP2', 'CRCP1', 'CRCP2'],
            format_func=lambda x: st.session_state['construction'][x]['name']
        )
        
        if selected_structure:
            struct = st.session_state['construction'][selected_structure]
            layers = struct['layers']
            joints = struct.get('joints')
            
            # ‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ï‡∏≤‡∏£‡∏≤‡∏á‡∏£‡∏≤‡∏¢‡∏•‡∏∞‡πÄ‡∏≠‡∏µ‡∏¢‡∏î
            detail_data = []
            total_cost = 0
            
            # ‡∏™‡πà‡∏ß‡∏ô‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á
            st.markdown(f"**{struct['name']}**")
            
            for i, layer in enumerate(layers):
                layer_cost = layer['quantity'] * layer['unit_cost']
                total_cost += layer_cost
                detail_data.append({
                    '‡∏•‡∏≥‡∏î‡∏±‡∏ö': i + 1,
                    '‡∏£‡∏≤‡∏¢‡∏Å‡∏≤‡∏£': layer['name'],
                    '‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤': f"{layer['thickness']} {layer['unit']}",
                    '‡∏õ‡∏£‡∏¥‡∏°‡∏≤‡∏ì (‡∏ï‡∏£.‡∏°.)': f"{layer['quantity']:,.0f}",
                    '‡∏£‡∏≤‡∏Ñ‡∏≤ (‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.)': f"{layer['unit_cost']:,.2f}",
                    '‡∏°‡∏π‡∏•‡∏Ñ‡πà‡∏≤ (‡∏ö‡∏≤‡∏ó)': f"{layer_cost:,.0f}"
                })
            
            # ‡∏™‡πà‡∏ß‡∏ô Joints (‡∏ñ‡πâ‡∏≤‡∏°‡∏µ)
            if joints:
                for j, joint in enumerate(joints):
                    joint_cost = joint['quantity'] * joint['unit_cost']
                    total_cost += joint_cost
                    detail_data.append({
                        '‡∏•‡∏≥‡∏î‡∏±‡∏ö': len(layers) + j + 1,
                        '‡∏£‡∏≤‡∏¢‡∏Å‡∏≤‡∏£': joint['name'],
                        '‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤': '-',
                        '‡∏õ‡∏£‡∏¥‡∏°‡∏≤‡∏ì (‡∏ï‡∏£.‡∏°.)': f"{joint['quantity']:,.0f}",
                        '‡∏£‡∏≤‡∏Ñ‡∏≤ (‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.)': f"{joint['unit_cost']:,.2f}",
                        '‡∏°‡∏π‡∏•‡∏Ñ‡πà‡∏≤ (‡∏ö‡∏≤‡∏ó)': f"{joint_cost:,.0f}"
                    })
            
            detail_df = pd.DataFrame(detail_data)
            st.dataframe(detail_df, use_container_width=True, hide_index=True)
            
            # ‡πÅ‡∏™‡∏î‡∏á‡∏£‡∏≤‡∏Ñ‡∏≤‡∏£‡∏ß‡∏°
            area_km = st.session_state.get('area_per_km', 22000) * road_length
            cost_per_sqm = total_cost / area_km if area_km > 0 else 0
            
            col_sum1, col_sum2, col_sum3, col_sum4 = st.columns(4)
            with col_sum1:
                st.metric("üí∞ ‡∏£‡∏≤‡∏Ñ‡∏≤‡∏£‡∏ß‡∏°", f"{total_cost:,.0f} ‡∏ö‡∏≤‡∏ó")
            with col_sum2:
                st.metric("üìè ‡∏£‡∏≤‡∏Ñ‡∏≤‡∏ï‡πà‡∏≠ ‡∏Å‡∏°.", f"{total_cost/road_length:,.0f} ‡∏ö‡∏≤‡∏ó/‡∏Å‡∏°.")
            with col_sum3:
                st.metric("üìä ‡∏•‡πâ‡∏≤‡∏ô‡∏ö‡∏≤‡∏ó/‡∏Å‡∏°.", f"{total_cost/road_length/1_000_000:.2f}")
            with col_sum4:
                st.metric("üìê ‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.", f"{cost_per_sqm:.2f}")
    
    # ===== Tab 3: ‡∏Ñ‡πà‡∏≤‡∏ö‡∏≥‡∏£‡∏∏‡∏á‡∏£‡∏±‡∏Å‡∏©‡∏≤ =====
    
    # ===== Tab 3: ‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô (‡πÄ‡∏î‡∏¥‡∏° Tab 6) =====
    with tab3:
        st.header("üìÑ ‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô")
        st.info("üí° ‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô‡∏à‡∏∞‡πÅ‡∏™‡∏î‡∏á‡πÄ‡∏â‡∏û‡∏≤‡∏∞‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏ß‡∏±‡∏™‡∏î‡∏∏‡πÅ‡∏•‡∏∞‡∏£‡∏≤‡∏Ñ‡∏≤ (‡πÑ‡∏°‡πà‡∏£‡∏ß‡∏° NPV)")
        
        # ‡∏ï‡∏£‡∏ß‡∏à‡∏™‡∏≠‡∏ö‡∏ß‡πà‡∏≤‡∏°‡∏µ‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏• construction ‡∏´‡∏£‡∏∑‡∏≠‡πÑ‡∏°‡πà
        if 'construction' in st.session_state and st.session_state['construction']:
            constr = st.session_state.get('construction', {})
            
            # ‡∏Å‡∏£‡∏≠‡∏á‡πÄ‡∏â‡∏û‡∏≤‡∏∞‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ó‡∏µ‡πà‡πÄ‡∏•‡∏∑‡∏≠‡∏Å "‡πÅ‡∏™‡∏î‡∏á‡πÉ‡∏ô‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô"
            all_details = {}
            structure_costs = {}  # ‡πÄ‡∏Å‡πá‡∏ö‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏• cost ‡πÅ‡∏¢‡∏Å
            for k, v in constr.items():
                if v.get('show', True):  # ‡πÄ‡∏â‡∏û‡∏≤‡∏∞‡∏ó‡∏µ‡πà tick ‡πÅ‡∏™‡∏î‡∏á‡πÉ‡∏ô‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô
                    all_details[k] = {
                        'name': v.get('name', k),
                        'details': v.get('details', []),
                        'cost_per_km': v.get('cost', 0),  # ‡∏•‡πâ‡∏≤‡∏ô‡∏ö‡∏≤‡∏ó/‡∏Å‡∏°.
                        'cost_sqm': v.get('cost_sqm', 0)   # ‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.
                    }
            
            # ‡∏ï‡∏£‡∏ß‡∏à‡∏™‡∏≠‡∏ö‡∏ß‡πà‡∏≤‡∏°‡∏µ‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏ó‡∏µ‡πà‡∏à‡∏∞‡πÅ‡∏™‡∏î‡∏á‡∏´‡∏£‡∏∑‡∏≠‡πÑ‡∏°‡πà
            if not all_details:
                st.warning("‚ö†Ô∏è ‡∏Å‡∏£‡∏∏‡∏ì‡∏≤‡πÄ‡∏•‡∏∑‡∏≠‡∏Å‡∏≠‡∏¢‡πà‡∏≤‡∏á‡∏ô‡πâ‡∏≠‡∏¢ 1 ‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ó‡∏µ‡πà‡∏ï‡πâ‡∏≠‡∏á‡∏Å‡∏≤‡∏£‡πÅ‡∏™‡∏î‡∏á‡πÉ‡∏ô‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô (tick ‚úÖ ‡πÅ‡∏™‡∏î‡∏á‡πÉ‡∏ô‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô ‡πÉ‡∏ô Tab 2)")
            else:
                # ‡πÅ‡∏™‡∏î‡∏á‡∏ï‡∏±‡∏ß‡∏≠‡∏¢‡πà‡∏≤‡∏á‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏ó‡∏µ‡πà‡∏à‡∏∞‡∏≠‡∏≠‡∏Å‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô
                st.subheader("üìä ‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏ó‡∏µ‡πà‡∏à‡∏∞‡∏£‡∏ß‡∏°‡πÉ‡∏ô‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô")
                
                for ptype, data in all_details.items():
                    if data['details']:
                        # ‡πÅ‡∏™‡∏î‡∏á‡∏ä‡∏∑‡πà‡∏≠‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á
                        structure_name = data['name']
                        with st.expander(f"üîç {structure_name}"):
                            df_preview = pd.DataFrame(data['details'])
                            st.dataframe(df_preview, use_container_width=True, hide_index=True)
                
                st.divider()
                
                # ‡∏õ‡∏∏‡πà‡∏°‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô
                c1, c2 = st.columns(2)
                
                with c1:
                    if not DOCX_AVAILABLE:
                        st.warning("‚ö†Ô∏è ‡πÑ‡∏°‡πà‡∏™‡∏≤‡∏°‡∏≤‡∏£‡∏ñ‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô Word ‡πÑ‡∏î‡πâ ‡πÄ‡∏ô‡∏∑‡πà‡∏≠‡∏á‡∏à‡∏≤‡∏Å python-docx ‡πÑ‡∏°‡πà‡∏™‡∏≤‡∏°‡∏≤‡∏£‡∏ñ‡πÉ‡∏ä‡πâ‡∏á‡∏≤‡∏ô‡πÑ‡∏î‡πâ")
                    elif st.button("üìÑ ‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô Word", type="primary", use_container_width=True):
                        try:
                            doc = generate_word_report_materials_only(
                                st.session_state['project_info'],
                                all_details
                            )
                            
                            buf = io.BytesIO()
                            doc.save(buf)
                            buf.seek(0)
                            
                            st.download_button("‚¨áÔ∏è ‡∏î‡∏≤‡∏ß‡∏ô‡πå‡πÇ‡∏´‡∏•‡∏î Word", data=buf,
                                               file_name=f"Materials_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                               use_container_width=True)
                            st.success("‚úÖ ‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô‡∏™‡∏≥‡πÄ‡∏£‡πá‡∏à!")
                        except Exception as e:
                            st.error(f"‚ùå ‡πÄ‡∏Å‡∏¥‡∏î‡∏Ç‡πâ‡∏≠‡∏ú‡∏¥‡∏î‡∏û‡∏•‡∏≤‡∏î‡πÉ‡∏ô‡∏Å‡∏≤‡∏£‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏£‡∏≤‡∏¢‡∏á‡∏≤‡∏ô: {str(e)}")
                
                with c2:
                    if st.button("üíæ ‡∏ö‡∏±‡∏ô‡∏ó‡∏∂‡∏Å‡πÇ‡∏Ñ‡∏£‡∏á‡∏Å‡∏≤‡∏£ (JSON)", use_container_width=True):
                        data = {
                            'project_info': st.session_state['project_info'],
                            'construction': {
                                k: {
                                    'cost': v.get('cost', 0),
                                    'details': v.get('details', [])
                                } for k, v in st.session_state.get('construction', {}).items()
                            },
                            'saved_at': datetime.now().isoformat()
                        }
                        st.download_button("‚¨áÔ∏è ‡∏î‡∏≤‡∏ß‡∏ô‡πå‡πÇ‡∏´‡∏•‡∏î JSON", data=json.dumps(data, ensure_ascii=False, indent=2),
                                           file_name=f"Project_{datetime.now().strftime('%Y%m%d_%H%M')}.json",
                                           mime="application/json",
                                           use_container_width=True)
                        st.success("‚úÖ ‡∏ö‡∏±‡∏ô‡∏ó‡∏∂‡∏Å‡πÇ‡∏Ñ‡∏£‡∏á‡∏Å‡∏≤‡∏£‡∏™‡∏≥‡πÄ‡∏£‡πá‡∏à!")
        else:
            st.warning("‚ö†Ô∏è ‡∏Å‡∏£‡∏∏‡∏ì‡∏≤‡πÄ‡∏û‡∏¥‡πà‡∏°‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á‡πÉ‡∏ô Tab 2 ‡∏Å‡πà‡∏≠‡∏ô")
    
    # ===== Tab 7: ‡∏ß‡∏¥‡πÄ‡∏Ñ‡∏£‡∏≤‡∏∞‡∏´‡πå‡∏à‡∏≤‡∏Å‡∏£‡∏π‡∏õ‡∏†‡∏≤‡∏û =====
    
    # ===== Tab 4: ‡∏ß‡∏¥‡πÄ‡∏Ñ‡∏£‡∏≤‡∏∞‡∏´‡πå‡∏à‡∏≤‡∏Å‡∏£‡∏π‡∏õ‡∏†‡∏≤‡∏û (‡πÄ‡∏î‡∏¥‡∏° Tab 7) =====
    with tab4:
        st.info("üí° Upload ‡∏£‡∏π‡∏õ‡∏†‡∏≤‡∏û‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á ‡πÅ‡∏•‡πâ‡∏ß‡∏£‡∏∞‡∏ö‡∏ö‡∏à‡∏∞‡∏ß‡∏¥‡πÄ‡∏Ñ‡∏£‡∏≤‡∏∞‡∏´‡πå‡πÅ‡∏•‡∏∞‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡∏£‡∏≤‡∏Ñ‡∏≤‡πÉ‡∏´‡πâ‡∏≠‡∏±‡∏ï‡πÇ‡∏ô‡∏°‡∏±‡∏ï‡∏¥")
        
        # Upload ‡∏£‡∏π‡∏õ‡∏†‡∏≤‡∏û
        uploaded_image = st.file_uploader(
            "‡πÄ‡∏•‡∏∑‡∏≠‡∏Å‡∏£‡∏π‡∏õ‡∏†‡∏≤‡∏û‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á",
            type=['png', 'jpg', 'jpeg'],
            help="‡∏£‡∏≠‡∏á‡∏£‡∏±‡∏ö‡πÑ‡∏ü‡∏•‡πå PNG, JPG, JPEG"
        )
        
        if uploaded_image is not None:
            col_img, col_result = st.columns([1, 1])
            
            with col_img:
                st.subheader("üñºÔ∏è ‡∏£‡∏π‡∏õ‡∏†‡∏≤‡∏û‡∏ó‡∏µ‡πà Upload")
                st.image(uploaded_image, use_container_width=True)
            
            with col_result:
                st.subheader("üìã ‡∏Å‡∏£‡∏≠‡∏Å‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ä‡∏±‡πâ‡∏ô‡∏ó‡∏≤‡∏á")
                st.markdown("‡∏Å‡∏£‡∏∏‡∏ì‡∏≤‡∏ï‡∏£‡∏ß‡∏à‡∏™‡∏≠‡∏ö‡πÅ‡∏•‡∏∞‡πÅ‡∏Å‡πâ‡πÑ‡∏Ç‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏ó‡∏µ‡πà‡∏≠‡πà‡∏≤‡∏ô‡∏à‡∏≤‡∏Å‡∏£‡∏π‡∏õ‡∏†‡∏≤‡∏û")
                
                # ‡πÄ‡∏•‡∏∑‡∏≠‡∏Å‡∏õ‡∏£‡∏∞‡πÄ‡∏†‡∏ó‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á
                structure_type = st.selectbox(
                    "‡∏õ‡∏£‡∏∞‡πÄ‡∏†‡∏ó‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á",
                    options=['AC Pavement', 'JPCP', 'JRCP', 'CRCP'],
                    key="img_structure_type"
                )
                
                # ‡∏Å‡∏≥‡∏´‡∏ô‡∏î‡∏à‡∏≥‡∏ô‡∏ß‡∏ô‡∏ä‡∏±‡πâ‡∏ô
                num_layers = st.number_input(
                    "‡∏à‡∏≥‡∏ô‡∏ß‡∏ô‡∏ä‡∏±‡πâ‡∏ô‡πÇ‡∏Ñ‡∏£‡∏á‡∏™‡∏£‡πâ‡∏≤‡∏á",
                    min_value=1, max_value=10, value=6,
                    key="img_num_layers"
                )
                
                st.divider()
                
                # ‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏ó‡∏µ‡πà‡πÄ‡∏•‡∏∑‡∏≠‡∏Å‡πÑ‡∏î‡πâ
                surface_materials = {
                    'AC Pavement': ['AC Wearing Course', 'PMA Wearing Course', 'AC Binder Course', 'AC Base Course', 'Tack Coat', 'Prime Coat'],
                    'JPCP': ['Concrete Slab (JPCP)', 'AC Interlayer', 'Non Woven Geotextile'],
                    'JRCP': ['Concrete Slab (JRCP)', 'AC Interlayer', 'Non Woven Geotextile'],
                    'CRCP': ['Concrete Slab (CRCP)', 'AC Interlayer', 'Steel Reinforcement', 'Non Woven Geotextile'],
                }
                
                base_materials = [
                    'Cement Treated Base (UCS 40 ksc)',
                    'Cement Modified Crushed Rock Base (UCS 24.5 ksc)',
                    'Crushed Rock Base Course',
                    'Soil Cement Subbase (UCS 7 ksc)',
                    'Soil Aggregate Subbase',
                    'Selected Material A',
                ]
                
                all_materials = surface_materials.get(structure_type, []) + base_materials
                
                # ‡πÄ‡∏Å‡πá‡∏ö‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏ä‡∏±‡πâ‡∏ô
                if 'img_layers' not in st.session_state:
                    st.session_state['img_layers'] = []
                
                img_layers = []
                total_cost_sqm = 0
                
                st.markdown("**‡∏£‡∏≤‡∏¢‡∏•‡∏∞‡πÄ‡∏≠‡∏µ‡∏¢‡∏î‡πÅ‡∏ï‡πà‡∏•‡∏∞‡∏ä‡∏±‡πâ‡∏ô:**")
                
                # Header
                cols_h = st.columns([3, 1.5, 2])
                cols_h[0].markdown("**‡∏ß‡∏±‡∏™‡∏î‡∏∏**")
                cols_h[1].markdown("**‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤ (cm)**")
                cols_h[2].markdown("**‡∏£‡∏≤‡∏Ñ‡∏≤ (‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.)**")
                
                for i in range(int(num_layers)):
                    cols = st.columns([3, 1.5, 2])
                    
                    with cols[0]:
                        # Default values ‡∏ï‡∏≤‡∏°‡∏•‡∏≥‡∏î‡∏±‡∏ö
                        default_materials = {
                            'AC Pavement': ['AC Wearing Course', 'AC Binder Course', 'AC Base Course', 'Cement Treated Base (UCS 40 ksc)', 'Soil Aggregate Subbase', 'Selected Material A'],
                            'JPCP': ['Concrete Slab (JPCP)', 'AC Interlayer', 'Cement Treated Base (UCS 40 ksc)', 'Crushed Rock Base Course', 'Soil Aggregate Subbase', 'Selected Material A'],
                            'JRCP': ['Concrete Slab (JRCP)', 'AC Interlayer', 'Cement Treated Base (UCS 40 ksc)', 'Crushed Rock Base Course', 'Soil Aggregate Subbase', 'Selected Material A'],
                            'CRCP': ['Concrete Slab (CRCP)', 'AC Interlayer', 'Cement Treated Base (UCS 40 ksc)', 'Crushed Rock Base Course', 'Soil Aggregate Subbase', 'Selected Material A'],
                        }
                        default_list = default_materials.get(structure_type, all_materials)
                        default_idx = i if i < len(default_list) else 0
                        default_mat = default_list[default_idx] if default_idx < len(default_list) else all_materials[0]
                        
                        try:
                            mat_idx = all_materials.index(default_mat)
                        except:
                            mat_idx = 0
                        
                        material = st.selectbox(
                            f"‡∏ß‡∏±‡∏™‡∏î‡∏∏‡∏ä‡∏±‡πâ‡∏ô {i+1}",
                            options=all_materials,
                            index=mat_idx,
                            key=f"img_mat_{i}",
                            label_visibility="collapsed"
                        )
                    
                    with cols[1]:
                        # Default thickness
                        default_thicknesses = {
                            'AC Pavement': [5, 7, 8, 20, 25, 30],
                            'JPCP': [30, 5, 20, 15, 25, 30],
                            'JRCP': [30, 5, 20, 15, 25, 30],
                            'CRCP': [30, 5, 20, 15, 25, 30],
                        }
                        default_thick_list = default_thicknesses.get(structure_type, [20]*10)
                        default_thick = default_thick_list[i] if i < len(default_thick_list) else 20
                        
                        thickness = st.number_input(
                            f"‡∏´‡∏ô‡∏≤ {i+1}",
                            min_value=0.0, max_value=100.0,
                            value=float(default_thick),
                            step=1.0,
                            key=f"img_thick_{i}",
                            label_visibility="collapsed"
                        )
                    
                    # ‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡∏£‡∏≤‡∏Ñ‡∏≤
                    price_sqm = 0
                    mat_lower = material.lower()
                    
                    if 'price_library' in st.session_state:
                        lib = st.session_state['price_library']
                        
                        # ‡∏ú‡∏¥‡∏ß‡∏ó‡∏≤‡∏á AC
                        if 'ac wearing' in mat_lower:
                            prices = lib['ac_prices'].get('AC Wearing Course', {})
                            price_sqm = prices.get(thickness, 0)
                            if price_sqm == 0 and prices:
                                closest = min(prices.keys(), key=lambda x: abs(x - thickness))
                                price_sqm = prices.get(closest, 0)
                        elif 'pma' in mat_lower:
                            prices = lib['ac_prices'].get('PMA Wearing Course', {})
                            price_sqm = prices.get(thickness, 0)
                            if price_sqm == 0 and prices:
                                closest = min(prices.keys(), key=lambda x: abs(x - thickness))
                                price_sqm = prices.get(closest, 0)
                        elif 'binder' in mat_lower:
                            prices = lib['ac_prices'].get('AC Binder Course', {})
                            price_sqm = prices.get(thickness, 0)
                            if price_sqm == 0 and prices:
                                closest = min(prices.keys(), key=lambda x: abs(x - thickness))
                                price_sqm = prices.get(closest, 0)
                        elif 'ac base' in mat_lower or 'ac interlayer' in mat_lower:
                            prices = lib['ac_prices'].get('AC Base Course', {})
                            price_sqm = prices.get(thickness, 0)
                            if price_sqm == 0 and prices:
                                closest = min(prices.keys(), key=lambda x: abs(x - thickness))
                                price_sqm = prices.get(closest, 0)
                        elif 'tack' in mat_lower:
                            price_sqm = 20
                        elif 'prime' in mat_lower:
                            price_sqm = 30
                        elif 'geotextile' in mat_lower:
                            price_sqm = 78
                        elif 'steel' in mat_lower:
                            price_sqm = 200
                        # ‡∏Ñ‡∏≠‡∏ô‡∏Å‡∏£‡∏µ‡∏ï
                        elif 'concrete' in mat_lower or 'slab' in mat_lower:
                            if 'jpcp' in mat_lower:
                                prices = lib['concrete_prices'].get('JPCP', {})
                            elif 'jrcp' in mat_lower:
                                prices = lib['concrete_prices'].get('JRCP', {})
                            elif 'crcp' in mat_lower:
                                prices = lib['concrete_prices'].get('CRCP', {})
                            else:
                                prices = lib['concrete_prices'].get('JPCP', {})
                            
                            price_sqm = prices.get(int(thickness), 0)
                            if price_sqm == 0 and prices:
                                closest = min(prices.keys(), key=lambda x: abs(x - thickness))
                                price_sqm = prices.get(closest, 0)
                        # ‡∏û‡∏∑‡πâ‡∏ô‡∏ó‡∏≤‡∏á (‡∏ö‡∏≤‡∏ó/‡∏•‡∏ö.‡∏°. ‚Üí ‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.)
                        elif 'cement treated' in mat_lower or 'ctb' in mat_lower:
                            base_price = lib['base_prices'].get('Cement Treated Base (UCS 40 ksc)', 1096)
                            price_sqm = base_price * thickness / 100
                        elif 'cement modified' in mat_lower or 'cmcr' in mat_lower:
                            base_price = lib['base_prices'].get('Cement Modified Crushed Rock Base (UCS 24.5 ksc)', 864)
                            price_sqm = base_price * thickness / 100
                        elif 'crushed rock' in mat_lower:
                            base_price = lib['base_prices'].get('Crushed Rock Base Course', 583)
                            price_sqm = base_price * thickness / 100
                        elif 'soil cement' in mat_lower:
                            base_price = lib['base_prices'].get('Soil Cement Subbase (UCS 7 ksc)', 854)
                            price_sqm = base_price * thickness / 100
                        elif 'soil aggregate' in mat_lower or 'aggregate subbase' in mat_lower:
                            base_price = lib['base_prices'].get('Soil Aggregate Subbase', 375)
                            price_sqm = base_price * thickness / 100
                        elif 'selected' in mat_lower:
                            base_price = lib['base_prices'].get('Selected Material A', 375)
                            price_sqm = base_price * thickness / 100
                    
                    with cols[2]:
                        st.markdown(f"**{price_sqm:,.2f}**")
                    
                    total_cost_sqm += price_sqm
                    img_layers.append({
                        'material': material,
                        'thickness': thickness,
                        'price_sqm': price_sqm
                    })
                
                st.session_state['img_layers'] = img_layers
        
        # ‡πÅ‡∏™‡∏î‡∏á‡∏ú‡∏•‡∏™‡∏£‡∏∏‡∏õ
        if uploaded_image is not None and 'img_layers' in st.session_state and st.session_state['img_layers']:
            st.divider()
            st.subheader("üìä ‡∏™‡∏£‡∏∏‡∏õ‡∏ú‡∏•‡∏Å‡∏≤‡∏£‡∏ß‡∏¥‡πÄ‡∏Ñ‡∏£‡∏≤‡∏∞‡∏´‡πå")
            
            img_layers = st.session_state['img_layers']
            total_cost_sqm = sum(layer['price_sqm'] for layer in img_layers)
            
            # ‡πÅ‡∏™‡∏î‡∏á‡∏ï‡∏≤‡∏£‡∏≤‡∏á
            summary_data = []
            for i, layer in enumerate(img_layers):
                summary_data.append({
                    '‡∏•‡∏≥‡∏î‡∏±‡∏ö': i + 1,
                    '‡∏ß‡∏±‡∏™‡∏î‡∏∏': layer['material'],
                    '‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏ô‡∏≤ (cm)': layer['thickness'],
                    '‡∏£‡∏≤‡∏Ñ‡∏≤ (‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.)': f"{layer['price_sqm']:,.2f}"
                })
            
            summary_df = pd.DataFrame(summary_data)
            st.dataframe(summary_df, use_container_width=True, hide_index=True)
            
            # Metrics
            col_m1, col_m2, col_m3 = st.columns(3)
            
            with col_m1:
                st.metric("üí∞ ‡∏£‡∏≤‡∏Ñ‡∏≤‡∏£‡∏ß‡∏°", f"{total_cost_sqm:,.2f} ‡∏ö‡∏≤‡∏ó/‡∏ï‡∏£.‡∏°.")
            
            with col_m2:
                # ‡∏Ñ‡∏≥‡∏ô‡∏ß‡∏ì‡∏ï‡πà‡∏≠ ‡∏Å‡∏°. (‡∏™‡∏°‡∏°‡∏ï‡∏¥ 22,000 ‡∏ï‡∏£.‡∏°./‡∏Å‡∏°.)
                area_km = st.session_state.get('area_per_km', 22000)
                cost_per_km = total_cost_sqm * area_km / 1_000_000
                st.metric("üìè ‡∏£‡∏≤‡∏Ñ‡∏≤‡∏ï‡πà‡∏≠ ‡∏Å‡∏°.", f"{cost_per_km:,.2f} ‡∏•‡πâ‡∏≤‡∏ô‡∏ö‡∏≤‡∏ó/‡∏Å‡∏°.")
            
            with col_m3:
                structure_type = st.session_state.get('img_structure_type', 'JPCP')
                if 'AC' in structure_type:
                    design_life = 20
                elif 'CRCP' in structure_type:
                    design_life = 30
                else:
                    design_life = 25
                st.metric("‚è±Ô∏è ‡∏≠‡∏≤‡∏¢‡∏∏‡∏≠‡∏≠‡∏Å‡πÅ‡∏ö‡∏ö", f"{design_life} ‡∏õ‡∏µ")
    
    # ‡πÄ‡∏Ñ‡∏£‡∏î‡∏¥‡∏ï (footer)
    st.divider()
    st.markdown("""
    <div style='text-align: center; color: #888; font-size: 0.85rem; padding: 20px;'>
        <b>‡∏û‡∏±‡∏í‡∏ô‡∏≤‡πÇ‡∏î‡∏¢</b><br>
        ‡∏£‡∏®.‡∏î‡∏£.‡∏≠‡∏¥‡∏ó‡∏ò‡∏¥‡∏û‡∏• ‡∏°‡∏µ‡∏ú‡∏•<br>
        ‡∏†‡∏≤‡∏Ñ‡∏ß‡∏¥‡∏ä‡∏≤‡∏Ñ‡∏£‡∏∏‡∏®‡∏≤‡∏™‡∏ï‡∏£‡πå‡πÇ‡∏¢‡∏ò‡∏≤ ‡∏Ñ‡∏ì‡∏∞‡∏Ñ‡∏£‡∏∏‡∏®‡∏≤‡∏™‡∏ï‡∏£‡πå‡∏≠‡∏∏‡∏ï‡∏™‡∏≤‡∏´‡∏Å‡∏£‡∏£‡∏°<br>
        ‡∏°‡∏´‡∏≤‡∏ß‡∏¥‡∏ó‡∏¢‡∏≤‡∏•‡∏±‡∏¢‡πÄ‡∏ó‡∏Ñ‡πÇ‡∏ô‡πÇ‡∏•‡∏¢‡∏µ‡∏û‡∏£‡∏∞‡∏à‡∏≠‡∏°‡πÄ‡∏Å‡∏•‡πâ‡∏≤‡∏û‡∏£‡∏∞‡∏ô‡∏Ñ‡∏£‡πÄ‡∏´‡∏ô‡∏∑‡∏≠ (‡∏°‡∏à‡∏û.)<br>
        <small style='color: #aaa;'>Pavement Structure Cost Analysis System v5.0</small>
    </div>
    """, unsafe_allow_html=True)


if __name__ == "__main__":
    main()
