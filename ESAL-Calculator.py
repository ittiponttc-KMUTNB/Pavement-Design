"""
ESAL Calculator - AASHTO 1993 (Version 4.0)
โปรแกรมคำนวณปริมาณเพลาเดี่ยวมาตรฐานเทียบเท่า (Equivalent Single Axle Load)
สำหรับผิวทาง Rigid Pavement และ Flexible Pavement
ตามมาตรฐาน AASHTO Guide for Design of Pavement Structures (1993)

Features V3:
- รองรับ Export Excel และ Word ในรูปแบบมาตรฐาน
- Save/Load Project สำหรับแก้ไขภายหลัง
- คำนวณ ACC. ESAL (สะสม)
- Export Word รวม Flexible + Rigid ในไฟล์เดียว (ตามรูปแบบรายงานมาตรฐาน)
- ระบบหมายเลขหัวข้อ/ตาราง แบบ Auto-increment
- บทเกริ่นนำแยก Flexible / Rigid พร้อม Preview
- Font TH SarabunPSK 15pt

พัฒนาโดย: รศ.ดร.อิทธิพล มีผล ภาควิชาครุศาสตร์โยธา มจพ.
"""

import streamlit as st
import pandas as pd
import numpy as np
from io import BytesIO
import json
import re
from datetime import datetime
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils.dataframe import dataframe_to_rows

# ============================================================
# ข้อมูลรถบรรทุก 6 ชนิดตามกรมทางหลวงประเทศไทย
# TRUCKS  : ใช้แสดงชื่อ/รายละเอียดใน UI และรายงาน
# VEHICLE_AXLES : (load_ton, L2_code, count) ต่อเพลา
#   L2_code : 1 = single axle, 2 = tandem axle
#   count   : จำนวนซ้ำ (ปกติ 1 เสมอในชุดข้อมูลนี้)
# ============================================================
TRUCKS = {
    'MB':  {'desc': 'Medium Bus (รถโดยสารขนาดกลาง)'},
    'HB':  {'desc': 'Heavy Bus (รถโดยสารขนาดใหญ่)'},
    'MT':  {'desc': 'Medium Truck (รถบรรทุกขนาดกลาง)'},
    'HT':  {'desc': 'Heavy Truck (รถบรรทุกขนาดใหญ่)'},
    'TR':  {'desc': 'Full Trailer (รถพ่วง)'},
    'STR': {'desc': 'Semi-Trailer (รถกึ่งพ่วง)'},
}

# โครงสร้างเพลาต่อประเภทรถ: list of (load_ton, L2_code, count)
# L2_code: 1 = single axle, 2 = tandem axle
VEHICLE_AXLES = {
    'MB':  [(4,  1, 1), (11, 1, 1)],
    'HB':  [(5,  1, 1), (20, 2, 1)],
    'MT':  [(4,  1, 1), (11, 1, 1)],
    'HT':  [(5,  1, 1), (20, 2, 1)],
    'TR':  [(5,  1, 1), (20, 2, 1), (11, 1, 1), (11, 1, 1)],
    'STR': [(5,  1, 1), (20, 2, 1), (20, 2, 1)],
}

TON_TO_KIP = 2.2046  # 1 metric ton = 2.2046 kips


# ============================================================
# Engine: สมการ AASHTO 1993 Appendix D — คำนวณ EALF จริง
# ไม่ใช้ lookup table → รองรับ SN ทศนิยม และ D ทุกค่าได้
# ============================================================

def ealf_flexible(L1_ton, L2, SN, pt):
    """คำนวณ Equivalent Axle Load Factor สำหรับ Flexible Pavement
    ตามสมการ AASHTO 1993 Appendix D

    Parameters
    ----------
    L1_ton : น้ำหนักเพลา (metric ton)
    L2     : axle code (1 = single, 2 = tandem)
    SN     : Structural Number (รับทศนิยมได้)
    pt     : Terminal Serviceability Index
    """
    import math
    L1  = L1_ton * TON_TO_KIP          # แปลงเป็น kips
    Gt  = math.log10((4.2 - pt) / (4.2 - 1.5))
    Bx  = 0.40 + 0.081 * (L1 + L2) ** 3.23 / ((SN + 1) ** 5.19 * L2 ** 3.23)
    B18 = 0.40 + 0.081 * (18 + 1)  ** 3.23 / ((SN + 1) ** 5.19 * 1.0  ** 3.23)
    return 10 ** (4.79 * math.log10(L1 + L2) - 4.33 * math.log10(L2)
                  - 4.79 * math.log10(19) + Gt * (1 / B18 - 1 / Bx))


def ealf_rigid(L1_ton, L2, D_in, pt):
    """คำนวณ Equivalent Axle Load Factor สำหรับ Rigid Pavement
    ตามสมการ AASHTO 1993 Appendix D

    Parameters
    ----------
    L1_ton : น้ำหนักเพลา (metric ton)
    L2     : axle code (1 = single, 2 = tandem)
    D_in   : ความหนาแผ่นคอนกรีต (นิ้ว, integer)
    pt     : Terminal Serviceability Index
    """
    import math
    L1  = L1_ton * TON_TO_KIP
    Gt  = math.log10((4.5 - pt) / (4.5 - 1.5))
    Bx  = 1.0 + 3.63 * (L1 + L2) ** 5.20 / ((D_in + 1) ** 8.46 * L2 ** 3.52)
    B18 = 1.0 + 3.63 * (18 + 1)  ** 5.20 / ((D_in + 1) ** 8.46 * 1.0  ** 3.52)
    return 10 ** (4.62 * math.log10(L1 + L2) - 3.28 * math.log10(L2)
                  - 4.62 * math.log10(19) + Gt * (1 / B18 - 1 / Bx))


def get_default_truck_factor(truck_code, pavement_type, pt, param):
    """คำนวณ Truck Factor ตามสมการ AASHTO 1993 Appendix D
    (แทนที่ lookup table เดิม — รองรับ SN ทศนิยม และ D ทุกค่า)

    Parameters
    ----------
    truck_code    : 'MB' | 'HB' | 'MT' | 'HT' | 'TR' | 'STR'
    pavement_type : 'rigid' | 'flexible'
    pt            : Terminal serviceability index
    param         : Rigid → D นิ้ว (integer) | Flexible → SN (ทศนิยมได้)
    """
    axles = VEHICLE_AXLES[truck_code]
    if pavement_type == 'flexible':
        return sum(ealf_flexible(L1, L2, param, pt) * cnt for L1, L2, cnt in axles)
    else:
        return sum(ealf_rigid(L1, L2, param, pt) * cnt for L1, L2, cnt in axles)


import logging
logger = logging.getLogger(__name__)


def validate_traffic_df(df):
    """ตรวจสอบความถูกต้องของข้อมูลปริมาณจราจรก่อนคำนวณ ESAL

    Raises
    ------
    ValueError  : ถ้า column ขาด, ค่าติดลบ, หรือ datatype ผิด
    """
    required_cols = list(TRUCKS.keys())  # ['MB','HB','MT','HT','TR','STR']

    # 1. ตรวจ column ครบหรือไม่
    missing = [c for c in required_cols if c not in df.columns]
    if missing:
        raise ValueError(f"ข้อมูลปริมาณจราจรขาด column: {', '.join(missing)}")

    # 2. ตรวจ datatype — ต้องเป็นตัวเลข
    for col in required_cols:
        if not pd.api.types.is_numeric_dtype(df[col]):
            raise ValueError(f"Column '{col}' ต้องเป็นตัวเลข (พบ {df[col].dtype})")

    # 3. ตรวจค่าติดลบ
    neg_cols = [c for c in required_cols if (df[c] < 0).any()]
    if neg_cols:
        raise ValueError(f"ค่า AADT ติดลบใน column: {', '.join(neg_cols)}")

    # 4. ตรวจ NaN
    nan_cols = [c for c in required_cols if df[c].isna().any()]
    if nan_cols:
        raise ValueError(f"พบค่า NaN ใน column: {', '.join(nan_cols)}")


def calculate_esal_with_acc(traffic_df, truck_factors, lane_factor=0.9, direction_factor=0.5):
    """คำนวณ ESAL และ Accumulated ESAL จากข้อมูลปริมาณจราจร

    Parameters
    ----------
    traffic_df       : DataFrame ที่มี column Year + รหัสรถ 6 ชนิด
    truck_factors    : dict {truck_code: TF_value}
    lane_factor      : Lane Distribution Factor (default 0.9)
    direction_factor : Directional Distribution Factor (default 0.5)

    Returns
    -------
    (results_df, total_acc_esal)
    """
    # ตรวจสอบข้อมูลก่อนคำนวณ
    validate_traffic_df(traffic_df)

    results  = []
    acc_esal = 0.0

    for idx, row in traffic_df.iterrows():
        year     = row.get('Year', idx + 1)
        year_data = {'Year': int(year)}

        total_aadt = 0
        for code in TRUCKS.keys():
            aadt = int(row[code])
            year_data[code] = aadt
            total_aadt += aadt
        year_data['AADT'] = total_aadt

        year_esal = 0.0
        for code in TRUCKS.keys():
            aadt      = row[code]
            tf        = truck_factors[code]
            year_esal += aadt * tf * lane_factor * direction_factor * 365

        year_data['ESAL']     = int(round(year_esal))
        acc_esal             += year_esal
        year_data['ACC_ESAL'] = int(round(acc_esal))
        results.append(year_data)

    return pd.DataFrame(results), int(round(acc_esal))


def create_template():
    """สร้าง Template Excel สำหรับอัพโหลดข้อมูล"""
    base = {'MB': 120, 'HB': 60, 'MT': 250, 'HT': 180, 'TR': 100, 'STR': 120}
    growth_rate = 1.045
    
    data = {'Year': list(range(1, 21))}
    for code in base.keys():
        data[code] = [int(base[code] * (growth_rate ** i)) for i in range(20)]
    
    return pd.DataFrame(data)


def create_excel_report(results_df, pavement_type, pt, param, lane_factor, direction_factor, 
                       total_esal, truck_factors, num_years):
    """สร้างรายงาน Excel ในรูปแบบมาตรฐาน"""
    wb = Workbook()
    ws = wb.active
    ws.title = "ESAL Report"
    
    # Styles
    header_font = Font(bold=True, size=14)
    title_font = Font(bold=True, size=16)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    header_fill = PatternFill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')
    center_align = Alignment(horizontal='center', vertical='center')
    right_align = Alignment(horizontal='right', vertical='center')
    
    pavement_text = "Rigid Pavement" if pavement_type == 'rigid' else "Flexible Pavement"
    ws.merge_cells('A1:I1')
    ws['A1'] = f"ปริมาณเพลามาตรฐาน (ESALs) ระยะเวลาออกแบบ {num_years} ปี"
    ws['A1'].font = title_font
    ws['A1'].alignment = center_align
    
    ws.merge_cells('A2:I2')
    ws['A2'] = f"ผิวทางแบบ{'แข็ง' if pavement_type == 'rigid' else 'ยืดหยุ่น'} ({pavement_text})"
    ws['A2'].font = header_font
    ws['A2'].alignment = center_align
    
    param_label = f"D = {param}" if pavement_type == 'rigid' else f"SN = {param}"
    params_data = [
        ('รายการ', 'ค่า'),
        ('ประเภทผิวทาง', pavement_text),
        ('pt', str(pt)),
        ('พารามิเตอร์', param_label),
        ('Lane Factor', str(lane_factor)),
        ('Direction Factor', str(direction_factor)),
        ('ESAL รวม', f"{total_esal:,}"),
        ('จำนวนปี', str(num_years))
    ]
    
    for i, (label, value) in enumerate(params_data):
        row = 4 + i
        ws[f'A{row}'] = label
        ws[f'B{row}'] = value
        ws[f'A{row}'].border = border
        ws[f'B{row}'].border = border
        if i == 0:
            ws[f'A{row}'].fill = header_fill
            ws[f'B{row}'].fill = header_fill
            ws[f'A{row}'].font = Font(bold=True)
            ws[f'B{row}'].font = Font(bold=True)
    
    ws['D4'] = 'รหัส'
    ws['E4'] = 'ประเภท'
    ws['F4'] = 'Truck Factor'
    for col in ['D', 'E', 'F']:
        ws[f'{col}4'].fill = header_fill
        ws[f'{col}4'].font = Font(bold=True)
        ws[f'{col}4'].border = border
        ws[f'{col}4'].alignment = center_align
    
    for i, code in enumerate(TRUCKS.keys()):
        row = 5 + i
        ws[f'D{row}'] = code
        ws[f'E{row}'] = TRUCKS[code]['desc']
        ws[f'F{row}'] = f"{truck_factors[code]:.3f}"
        ws[f'D{row}'].border = border
        ws[f'E{row}'].border = border
        ws[f'F{row}'].border = border
        ws[f'D{row}'].alignment = center_align
        ws[f'F{row}'].alignment = right_align
    
    start_row = 14
    ws[f'I{start_row-1}'] = 'แสดงปริมาณสะสม'
    ws[f'I{start_row-1}'].font = Font(italic=True, size=9)
    
    headers = ['Year', 'MB', 'HB', 'MT', 'HT', 'TR', 'STR', 'AADT', 'ESAL', 'ACC. ESAL']
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=start_row, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = Font(bold=True)
        cell.border = border
        cell.alignment = center_align
    
    for row_idx, row_data in results_df.iterrows():
        excel_row = start_row + 1 + row_idx
        for col_idx, header in enumerate(headers, 1):
            if header == 'ACC. ESAL':
                value = row_data.get('ACC_ESAL', 0)
            else:
                value = row_data.get(header, 0)
            
            cell = ws.cell(row=excel_row, column=col_idx, value=value)
            cell.border = border
            
            if header in ['ESAL', 'ACC. ESAL', 'AADT']:
                cell.number_format = '#,##0'
                cell.alignment = right_align
            elif header == 'Year':
                cell.alignment = center_align
            else:
                cell.alignment = right_align
    
    ws.column_dimensions['A'].width = 18
    ws.column_dimensions['B'].width = 15
    ws.column_dimensions['C'].width = 3
    ws.column_dimensions['D'].width = 8
    ws.column_dimensions['E'].width = 35
    ws.column_dimensions['F'].width = 14
    for col in ['G', 'H', 'I', 'J']:
        ws.column_dimensions[col].width = 14
    
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return output


# ============================================================
# Auto-increment table number helper
# ============================================================
def increment_table_number(base_number, offset):
    """
    เพิ่มเลขตารางอัตโนมัติ โดยบวก offset เข้ากับเลขท้ายสุด
    รองรับรูปแบบ:
      "4-1"   + 1 -> "4-2"
      "3.2-1" + 1 -> "3.2-2"   (X.Y-Z)
      "4.1"   + 1 -> "4.2"
      "1"     + 2 -> "3"
    """
    s = base_number.strip()

    # รูปแบบ "X.Y-Z" เช่น "3.2-1" — ต้องตรวจก่อน "X-Y"
    match = re.match(r'^(\d+\.\d+)-(\d+)$', s)
    if match:
        return f"{match.group(1)}-{int(match.group(2)) + offset}"

    # รูปแบบ "X-Y" เช่น "4-1"
    match = re.match(r'^(\d+)-(\d+)$', s)
    if match:
        return f"{match.group(1)}-{int(match.group(2)) + offset}"

    # รูปแบบ "X.Y" เช่น "4.1"
    match = re.match(r'^(\d+)\.(\d+)$', s)
    if match:
        return f"{match.group(1)}.{int(match.group(2)) + offset}"

    # รูปแบบตัวเลขเดียว เช่น "1"
    match = re.match(r'^(\d+)$', s)
    if match:
        return str(int(match.group(1)) + offset)

    return f"{base_number}+{offset}"  # fallback


# ============================================================
# Word Report Generation (python-docx)
# ============================================================
def create_word_report_single(results_df, pavement_type, pt, param, lane_factor, direction_factor,
                              total_esal, truck_factors, num_years, report_settings=None):
    """สร้างรายงาน Word สำหรับผิวทางประเภทเดียว (Flexible หรือ Rigid)"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import nsdecls, qn
        from docx.oxml import parse_xml, OxmlElement
    except ImportError:
        return None
    
    doc = Document()
    FONT_NAME = 'TH SarabunPSK'
    FONT_SIZE = 15
    TABLE_FONT_SIZE = 14
    
    # ตั้งค่า Normal style
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(FONT_SIZE)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    
    # ตั้งค่าหน้ากระดาษ A4
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    
    # กำหนดค่า default report settings
    if report_settings is None:
        report_settings = {}
    
    if pavement_type == 'flexible':
        section_num = report_settings.get('flex_section_number', '4.2.2')
        table_start = report_settings.get('flex_table_start', '4-1')
    else:
        section_num = report_settings.get('rigid_section_number', '4.2.3')
        table_start = report_settings.get('rigid_table_start', '4-4')
    
    tbl_param = table_start
    tbl_tf = increment_table_number(table_start, 1)
    tbl_esal = increment_table_number(table_start, 2)
    
    pavement_text = "Rigid Pavement" if pavement_type == 'rigid' else "Flexible Pavement"
    pavement_thai = "แบบแข็ง" if pavement_type == 'rigid' else "ยืดหยุ่น"
    param_label = f"D = {param} นิ้ว" if pavement_type == 'rigid' else f"SN = {param}"
    
    _build_section(doc, pavement_type, pavement_text, pavement_thai, section_num,
                   num_years, param_label, pt, lane_factor, direction_factor,
                   total_esal, truck_factors, results_df,
                   tbl_param, tbl_tf, tbl_esal,
                   FONT_NAME, FONT_SIZE, TABLE_FONT_SIZE)
    
    # Footer
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("พัฒนาเพื่อการเรียนการสอนโดย รศ.ดร.อิทธิพล มีผล ภาควิชาครุศาสตร์โยธา มจพ.")
    run.font.name = FONT_NAME
    run.font.size = Pt(14)
    run.italic = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def create_word_report_combined(traffic_df, flex_params, rigid_params, report_settings=None):
    """
    สร้างรายงาน Word รวม Flexible + Rigid ในไฟล์เดียว (รูปแบบ multi-param)
    flex_params  : dict with keys: pt, param_list, lane_factor, direction_factor
    rigid_params : dict with keys: pt, param_list, lane_factor, direction_factor
    backward-compat: รับ param (int/float เดี่ยว) แปลงเป็น list อัตโนมัติ
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return None

    if report_settings is None:
        report_settings = {}

    # ── backward-compat: param เดี่ยว → list ──
    def to_list(p):
        if isinstance(p, list):
            return p
        return [p]

    fp_list = to_list(flex_params.get('param_list',  flex_params.get('param', 7)))
    rp_list = to_list(rigid_params.get('param_list', rigid_params.get('param', 12)))

    # ── คำนวณ multi_results สำหรับ Flexible ──
    fp = flex_params
    flex_multi = {}
    for p in fp_list:
        tf_p = {code: get_default_truck_factor(code, 'flexible', fp['pt'], p)
                for code in TRUCKS.keys()}
        r_df, t_esal = calculate_esal_with_acc(
            traffic_df, tf_p, fp['lane_factor'], fp['direction_factor']
        )
        flex_multi[p] = (r_df, t_esal)

    # ── คำนวณ multi_results สำหรับ Rigid ──
    rp = rigid_params
    rigid_multi = {}
    for p in rp_list:
        tf_p = {code: get_default_truck_factor(code, 'rigid', rp['pt'], p)
                for code in TRUCKS.keys()}
        r_df, t_esal = calculate_esal_with_acc(
            traffic_df, tf_p, rp['lane_factor'], rp['direction_factor']
        )
        rigid_multi[p] = (r_df, t_esal)

    # ── settings แยก Flex / Rigid ──
    flex_settings = {
        'flex_section_number': report_settings.get('flex_section_number', '4.2.2'),
        'flex_table_start':    report_settings.get('flex_table_start',    '4-1'),
    }
    rigid_settings = {
        'rigid_section_number': report_settings.get('rigid_section_number', '4.2.3'),
        'rigid_table_start':    report_settings.get('rigid_table_start',    '4-4'),
    }

    # ── สร้าง Word Flex ก่อน ──
    buf_flex = create_word_report_multi(
        traffic_df, 'flexible', fp['pt'], fp_list,
        fp['lane_factor'], fp['direction_factor'],
        flex_multi, flex_settings
    )
    if buf_flex is None:
        return None

    # ── merge: โหลด flex doc แล้วต่อ rigid เข้าไป ──
    buf_flex.seek(0)
    doc = Document(buf_flex)

    # page break
    p_br = doc.add_paragraph()
    run  = p_br.add_run()
    br   = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)

    # ── สร้าง Word Rigid ──
    buf_rigid = create_word_report_multi(
        traffic_df, 'rigid', rp['pt'], rp_list,
        rp['lane_factor'], rp['direction_factor'],
        rigid_multi, rigid_settings
    )
    if buf_rigid is None:
        return None

    # ── copy body elements จาก rigid → doc ──
    buf_rigid.seek(0)
    doc_rigid = Document(buf_rigid)
    for elem in doc_rigid.element.body:
        import copy
        doc.element.body.append(copy.deepcopy(elem))

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output



# ===== 4.5 สมการคำนวณ ESAL (ตามรูปแบบรายงานที่ปรึกษา) =====
def _add_esal_formula_block(doc, num_years, FONT_NAME, FONT_SIZE):
    """
    เพิ่มสมการ ESAL (W18) ลงใน Word document
    รูปแบบตามรายงานที่ปรึกษากรมทางหลวง:
      W_{18,j} = { Σ_{i=1}^{6} A_i } × TF × L × D × 365
      W_{18}   = Σ_{j=1}^{k=20} W_{18,j}
    ตามด้วยตาราง "โดยที่" ไม่มีเส้น border
    """
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    EQ_FONT  = 'Times New Roman'
    EQ_SIZE  = 11
    TH_FONT  = FONT_NAME
    TH_SIZE  = FONT_SIZE

    # ── helpers ──────────────────────────────────────────────
    def _r(para, text, fname=EQ_FONT, fsize=EQ_SIZE, bold=False, italic=True):
        """เพิ่ม run พร้อม font"""
        run = para.add_run(text)
        run.font.name  = fname
        run.font.size  = Pt(fsize)
        run.bold       = bold
        run.italic     = italic
        run._element.rPr.rFonts.set(qn('w:eastAsia'), fname)
        return run

    def _sub(para, text, fname=EQ_FONT, fsize=EQ_SIZE - 1):
        run = para.add_run(text)
        run.font.name      = fname
        run.font.size      = Pt(fsize)
        run.font.subscript = True
        run.italic         = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), fname)
        return run

    def _sup(para, text, fname=EQ_FONT, fsize=EQ_SIZE - 1):
        run = para.add_run(text)
        run.font.name        = fname
        run.font.size        = Pt(fsize)
        run.font.superscript = True
        run.italic           = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), fname)
        return run

    def _th(para, text, fname=None, fsize=None, bold=False):
        """Thai run — ไม่ italic"""
        fn = fname or TH_FONT
        fs = fsize or TH_SIZE
        run = para.add_run(text)
        run.font.name  = fn
        run.font.size  = Pt(fs)
        run.bold       = bold
        run.italic     = False
        run._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
        return run

    def _thai_justify(para):
        pPr = para._element.get_or_add_pPr()
        jc  = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'thaiDistribute')
        pPr.append(jc)

    def _remove_tbl_border(tbl):
        tblPr     = tbl._tbl.tblPr
        tblBorder = OxmlElement('w:tblBorders')
        for side in ('top','left','bottom','right','insideH','insideV'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'none')
            b.set(qn('w:sz'),  '0')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'auto')
            tblBorder.append(b)
        tblPr.append(tblBorder)

    def _remove_cell_border(cell):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorder = OxmlElement('w:tcBorders')
        for side in ('top','left','bottom','right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'none')
            b.set(qn('w:sz'),  '0')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'auto')
            tcBorder.append(b)
        tcPr.append(tcBorder)

    # ── ย่อหน้าเกริ่นนำ (ตามภาพ) ────────────────────────────
    # "โดยที่ค่าปริมาณเพลามาตรฐาน Equivalent Single Axle Load 18 kips (W₁₈)
    #  หมายถึง ปริมาณการจราจรของรถบรรทุกมาตรฐานที่วิ่งผ่านช่องจราจรออกแบบ
    #  (Design Lane) ในช่วงระยะเวลาออกแบบ (Design Period) โดยสามารถคำนวณจากสมการ"
    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.first_line_indent = Cm(1.25)
    p_intro.paragraph_format.space_after = Pt(6)
    _thai_justify(p_intro)

    # "โดยที่ค่าปริมาณเพลามาตรฐาน "
    _th(p_intro, 'โดยที่ค่าปริมาณเพลามาตรฐาน ')

    # "Equivalent Single Axle Load 18 kips (W" — TH font แต่ขึ้นต้นด้วย Latin
    _th(p_intro, 'Equivalent Single Axle Load 18 kips (W')

    # "18" subscript
    r_sub = p_intro.add_run('18')
    r_sub.font.name      = TH_FONT
    r_sub.font.size      = Pt(TH_SIZE)
    r_sub.font.subscript = True
    r_sub.italic         = False
    r_sub._element.rPr.rFonts.set(qn('w:eastAsia'), TH_FONT)

    # ") หมายถึง ปริมาณการจราจรของรถบรรทุกมาตรฐานที่วิ่งผ่านช่องจราจรออกแบบ ("
    _th(p_intro, ') หมายถึง ปริมาณการจราจรของรถบรรทุกมาตรฐานที่วิ่งผ่านช่องจราจรออกแบบ (')

    # "Design Lane" — bold ตามภาพ
    _th(p_intro, 'Design Lane', bold=True)

    # ") ในช่วงระยะเวลาออกแบบ ("
    _th(p_intro, ') ในช่วงระยะเวลาออกแบบ (')

    # "Design Period" — bold ตามภาพ
    _th(p_intro, 'Design Period', bold=True)

    # ") โดยสามารถคำนวณจากสมการ"
    _th(p_intro, ') โดยสามารถคำนวณจากสมการ')

    # ── สมการที่ 1 ───────────────────────────────────────────
    # W_{18,j}  =  { Σ_{i=1}^{6} A_i }  ×  TF  ×  L  ×  D  ×  365
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(6)
    p1.paragraph_format.space_after  = Pt(2)

    _r(p1, 'W')
    _sub(p1, '18')
    _r(p1, 'j')   # subscript j (italic ปกติ)

    _r(p1, '  =  ', italic=False)

    # วงเล็บปีกกา { Σ A_i }  — ใช้ { } ธรรมดา (ตามรูปแบบรายงาน)
    _r(p1, '{', fname=EQ_FONT, fsize=EQ_SIZE + 2, italic=False)
    _sup(p1, '6 ', fsize=EQ_SIZE - 2)
    _r(p1, '\u03A3', fsize=EQ_SIZE + 2, italic=False)                 # Σ
    _sub(p1, 'i=1', fsize=EQ_SIZE - 2)
    _r(p1, ' A', italic=True)
    _sub(p1, 'i')
    _r(p1, '}', fname=EQ_FONT, fsize=EQ_SIZE + 2, italic=False)

    _r(p1, '  \u00D7  TF  \u00D7  L  \u00D7  D  \u00D7  365',
       italic=False)

    # ── สมการที่ 2 ───────────────────────────────────────────
    # W_{18}  =  Σ_{j=1}^{k=20} W_{18,j}
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after  = Pt(10)

    _r(p2, 'W')
    _sub(p2, '18')

    _r(p2, '  =  ', italic=False)

    _sup(p2, f'k={num_years} ', fsize=EQ_SIZE - 2)
    _r(p2, '\u03A3', fsize=EQ_SIZE + 2, italic=False)
    _sub(p2, 'j=1', fsize=EQ_SIZE - 2)

    _r(p2, '  W')
    _sub(p2, '18')
    _r(p2, 'j', italic=True)

    # ── ส่วน "โดยที่" — ตารางไม่มี border ────────────────────
    p_label = doc.add_paragraph()
    p_label.paragraph_format.space_after = Pt(2)
    _th(p_label, 'โดยที่')

    # คอลัมน์: [สัญลักษณ์] [=] [คำอธิบาย]
    # ใช้ตาราง 3 คอลัมน์, ไม่มี border
    legend = [
        # (sym_runs, desc_text)
        # sym_runs = list of (text, is_subscript, is_superscript)
        ([('W', False, False), ('18', True, False)],
         f'ผลรวมปริมาณเพลาเดี่ยวมาตรฐานออกแบบขนาด 18 kip ถึงอายุออกแบบ (k) เท่ากับ {num_years} ปี'),
        ([('W', False, False), ('18', True, False), ('j', True, False)],
         'ผลรวมปริมาณเพลาเดี่ยวมาตรฐานออกแบบขนาด 18 kip ณ ปีใดๆ (j)'),
        ([('TF', False, False)],
         'Truck Factor พิจารณาตามข้อมูลด้านชั่งน้ำหนักบริเวณพื้นที่โครงการ หรือใกล้เคียง'),
        ([('A', False, False), ('i', True, False)],
         'ปริมาณรถบรรทุกประเภทที่ i หน่วย คัน/วัน จากข้อมูลการคาดการณ์ปริมาณจราจรของโครงการ'),
        ([('L', False, False)],
         'Lane Distribution Factor'),
        ([('D', False, False)],
         'Directional Distribution Factor'),
    ]

    leg_tbl = doc.add_table(rows=len(legend), cols=3)
    leg_tbl.style = 'Table Grid'
    _remove_tbl_border(leg_tbl)

    # กำหนด indent ซ้ายของตาราง (~2 cm)
    tblPr = leg_tbl._tbl.tblPr
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'),    '1134')   # ~2 cm  (1 cm = 567 twips)
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)

    # กำหนดความกว้างคอลัมน์: สัญลักษณ์ | = | คำอธิบาย
    # รวม ~8200 twips (ภายใน margin ของหน้า A4)
    # col0=1100 (สัญลักษณ์), col1=400 (=), col2=6700 (คำอธิบาย บรรทัดเดียว)
    col_widths = ['1100', '400', '6700']
    tblGrid = OxmlElement('w:tblGrid')
    for w in col_widths:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), w)
        tblGrid.append(gc)
    leg_tbl._tbl.insert(0, tblGrid)

    for row_idx, (sym_runs, desc) in enumerate(legend):
        row  = leg_tbl.rows[row_idx]
        c0, c1, c2 = row.cells[0], row.cells[1], row.cells[2]

        # ลบ border ทุก cell
        for cell in (c0, c1, c2):
            _remove_cell_border(cell)
            # ตั้ง vertical align ให้ชิดบน
            tcPr = cell._tc.get_or_add_tcPr()
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'top')
            tcPr.append(vAlign)

        # col 0 — สัญลักษณ์ (italic Times New Roman)
        p_sym = c0.paragraphs[0]
        p_sym.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for (txt, is_sub, is_sup) in sym_runs:
            run = p_sym.add_run(txt)
            run.font.name   = EQ_FONT
            run.font.size   = Pt(EQ_SIZE)
            run.italic      = True
            run.font.subscript   = is_sub
            run.font.superscript = is_sup
            run._element.rPr.rFonts.set(qn('w:eastAsia'), EQ_FONT)

        # col 1 — เครื่องหมาย "="
        p_eq = c1.paragraphs[0]
        p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_eq = p_eq.add_run('=')
        r_eq.font.name  = TH_FONT
        r_eq.font.size  = Pt(TH_SIZE)
        r_eq.italic     = False
        r_eq._element.rPr.rFonts.set(qn('w:eastAsia'), TH_FONT)

        # col 2 — คำอธิบาย (TH SarabunPSK)
        p_desc = c2.paragraphs[0]
        _thai_justify(p_desc)
        r_desc = p_desc.add_run(desc)
        r_desc.font.name  = TH_FONT
        r_desc.font.size  = Pt(TH_SIZE)
        r_desc.italic     = False
        r_desc._element.rPr.rFonts.set(qn('w:eastAsia'), TH_FONT)

    doc.add_paragraph()   # blank line after legend

def _build_section(doc, pavement_type, pavement_text, pavement_thai, section_num,
                   num_years, param_label, pt, lane_factor, direction_factor,
                   total_esal, truck_factors, results_df,
                   tbl_param, tbl_tf, tbl_esal,
                   FONT_NAME, FONT_SIZE, TABLE_FONT_SIZE):
    """สร้างเนื้อหาหนึ่ง section (Flexible หรือ Rigid) ใน Word document"""
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import OxmlElement, parse_xml
    
    def set_run(run, font_name=FONT_NAME, font_size=FONT_SIZE, bold=False):
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    def set_cell_font(cell, font_name=FONT_NAME, font_size=TABLE_FONT_SIZE, bold=False, align=None):
        for paragraph in cell.paragraphs:
            if align:
                paragraph.alignment = align
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run.bold = bold
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    def set_cell_shading(cell, color="D9E2F3"):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    
    def set_thai_distribute(paragraph):
        pPr = paragraph._element.get_or_add_pPr()
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'thaiDistribute')
        pPr.append(jc)
    
    # ===== 1. หัวข้อ =====
    heading_para = doc.add_paragraph()
    heading_para.paragraph_format.space_after = Pt(6)
    
    # เลขหัวข้อ (tab) ชื่อหัวข้อ
    run = heading_para.add_run(f"{section_num}\t")
    set_run(run, font_size=FONT_SIZE, bold=True)
    
    run = heading_para.add_run(f"ปริมาณเพลามาตรฐาน (ESALs) ระยะเวลาออกแบบ {num_years} ปี ผิวทาง {pavement_text}")
    set_run(run, font_size=FONT_SIZE, bold=True)
    
    # ===== 2. บทเกริ่นนำ =====
    intro_para = doc.add_paragraph()
    intro_para.paragraph_format.first_line_indent = Cm(1.25)
    intro_para.paragraph_format.space_after = Pt(6)
    set_thai_distribute(intro_para)
    
    if pavement_type == 'flexible':
        intro_parts = [
            ("ในการคำนวณปริมาณเพลามาตรฐาน สำหรับผิวทางยืดหยุ่น ที่ปรึกษาได้กำหนดค่าพารามิเตอร์ต่าง ๆ และค่า ", False),
            ("Truck Factor ", False),
            ("ของรถบรรทุกหนัก ที่ใช้สำหรับการคำนวณ ดังแสดงในตารางที่ ", False),
            (f"{tbl_param}", False),
            (" และ ", False),
            (f"{tbl_tf}", False),
            (f" ดังนั้นค่าปริมาณเพลามาตรฐาน สำหรับผิวทางยืดหยุ่น ที่ระยะเวลาออกแบบ {num_years} ปี แสดงดังตารางที่ ", False),
            (f"{tbl_esal}", False),
        ]
    else:
        intro_parts = [
            ("ในการคำนวณปริมาณเพลามาตรฐานสำหรับผิวทางแบบแข็งหรือผิวทางคอนกรีต โดยที่ปรึกษาได้กำหนดค่าพารามิเตอร์ต่าง ๆ และค่า ", False),
            ("Truck Factor ", False),
            ("ของรถบรรทุกหนัก ที่ใช้สำหรับการคำนวณ ดังแสดงในตารางที่ ", False),
            (f"{tbl_param}", False),
            (" และ ", False),
            (f"{tbl_tf}", False),
            (f" ดังนั้นค่าปริมาณเพลามาตรฐาน สำหรับผิวทางแบบแข็ง ที่ระยะเวลาออกแบบ {num_years} ปี แสดงดังตารางที่ ", False),
            (f"{tbl_esal}", False),
        ]
    
    for text, is_bold in intro_parts:
        run = intro_para.add_run(text)
        set_run(run, font_size=FONT_SIZE, bold=is_bold)
    
    doc.add_paragraph()  # blank line
    
    # ===== 3. ตารางที่ X-1: ค่าพารามิเตอร์ =====
    cap1 = doc.add_paragraph()
    cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap1.paragraph_format.space_after = Pt(3)
    
    run = cap1.add_run(f"ตารางที่ {tbl_param} ")
    set_run(run, font_size=FONT_SIZE, bold=True)
    run = cap1.add_run("ค่าพารามิเตอร์ต่าง ๆ ที่ใช้สำหรับการคำนวณ")
    set_run(run, font_size=FONT_SIZE, bold=False)
    
    param_data = [
        ('รายการ', 'ค่า'),
        ('ประเภทผิวทาง', pavement_text),
        ('pt', str(pt)),
        ('พารามิเตอร์', param_label),
        ('Lane Factor', str(lane_factor)),
        ('Direction Factor', str(direction_factor)),
        ('ESAL รวม', f"{total_esal:,}"),
        ('จำนวนปี', str(num_years))
    ]
    
    param_table = doc.add_table(rows=len(param_data), cols=2)
    param_table.style = 'Table Grid'
    param_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, (label, value) in enumerate(param_data):
        row = param_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        is_header = (i == 0)
        set_cell_font(row.cells[0], font_size=TABLE_FONT_SIZE, bold=is_header)
        set_cell_font(row.cells[1], font_size=TABLE_FONT_SIZE, bold=is_header)
        if is_header:
            set_cell_shading(row.cells[0])
            set_cell_shading(row.cells[1])
    
    doc.add_paragraph()
    
    # ===== 4. ตารางที่ X-2: Truck Factor =====
    cap2 = doc.add_paragraph()
    cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap2.paragraph_format.space_after = Pt(3)
    
    run = cap2.add_run(f"ตารางที่ {tbl_tf} ")
    set_run(run, font_size=FONT_SIZE, bold=True)
    run = cap2.add_run("ค่า Truck Factor ของรถบรรทุกหนัก")
    set_run(run, font_size=FONT_SIZE, bold=False)
    
    tf_table = doc.add_table(rows=7, cols=3)
    tf_table.style = 'Table Grid'
    tf_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr = tf_table.rows[0]
    for j, h in enumerate(['รหัส', 'ประเภท', 'Truck Factor']):
        hdr.cells[j].text = h
        set_cell_font(hdr.cells[j], font_size=TABLE_FONT_SIZE, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(hdr.cells[j])
    
    for i, code in enumerate(TRUCKS.keys()):
        row = tf_table.rows[i + 1]
        row.cells[0].text = code
        row.cells[1].text = TRUCKS[code]['desc']
        row.cells[2].text = f"{truck_factors[code]:.3f}"
        for cell in row.cells:
            set_cell_font(cell, font_size=TABLE_FONT_SIZE, bold=False)
        set_cell_font(row.cells[0], font_size=TABLE_FONT_SIZE, bold=False,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_font(row.cells[2], font_size=TABLE_FONT_SIZE, bold=False,
                      align=WD_ALIGN_PARAGRAPH.RIGHT)
    
    doc.add_paragraph()


    _add_esal_formula_block(doc, num_years, FONT_NAME, FONT_SIZE)

    _add_esal_formula_block(doc, num_years, FONT_NAME, FONT_SIZE)

    # ===== 5. ตารางที่ X-3: ESAL =====
    cap3 = doc.add_paragraph()
    cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap3.paragraph_format.space_after = Pt(3)
    
    run = cap3.add_run(f"ตารางที่ {tbl_esal} ")
    set_run(run, font_size=FONT_SIZE, bold=True)
    run = cap3.add_run(f"ค่าปริมาณเพลามาตรฐาน สำหรับผิวทาง{pavement_thai} ที่ระยะเวลาออกแบบ {num_years} ปี")
    set_run(run, font_size=FONT_SIZE, bold=False)
    
    headers = ['Year', 'MB', 'HB', 'MT', 'HT', 'TR', 'STR', 'AADT', 'ESAL', 'ACC. ESAL']
    esal_table = doc.add_table(rows=len(results_df) + 1, cols=len(headers))
    esal_table.style = 'Table Grid'
    esal_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr = esal_table.rows[0]
    for j, header in enumerate(headers):
        hdr.cells[j].text = header
        for paragraph in hdr.cells[j].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_font(hdr.cells[j], font_size=12, bold=True)
        set_cell_shading(hdr.cells[j])
    
    for i, row_data in results_df.iterrows():
        row = esal_table.rows[i + 1]
        for j, header in enumerate(headers):
            if header == 'ACC. ESAL':
                value = row_data.get('ACC_ESAL', 0)
            else:
                value = row_data.get(header, 0)
            
            if header in ['ESAL', 'ACC. ESAL', 'AADT']:
                row.cells[j].text = f"{int(value):,}"
            else:
                row.cells[j].text = str(int(value))
            
            for paragraph in row.cells[j].paragraphs:
                if header == 'Year':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_cell_font(row.cells[j], font_size=12, bold=False)



def create_word_report_multi(traffic_df, pavement_type, pt, param_list,
                             lane_factor, direction_factor,
                             multi_results, report_settings=None):
    """
    สร้างรายงาน Word แบบ multi-param (3 ค่า D หรือ SN)
    โครงสร้าง 4 ตาราง:
      tbl_param  — พารามิเตอร์
      tbl_tf     — Truck Factor (3 คอลัมน์)
      tbl_traf   — ปริมาณจราจรรายปี (Year, MB-STR, ADTT)
      tbl_esal   — ESAL รายปี + footer ACC.ESAL (Year, ADTT, ESAL×3)

    multi_results : dict {param_val: (results_df, total_esal)}
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import nsdecls, qn
        from docx.oxml import parse_xml, OxmlElement
    except ImportError:
        return None

    doc      = Document()
    FN       = 'TH SarabunPSK'
    FS       = 15
    TFS      = 14
    TFS_SM   = 12   # font ตารางข้อมูลรายปี

    # ── ตั้งค่า Normal style ──
    style = doc.styles['Normal']
    style.font.name = FN
    style.font.size = Pt(FS)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FN)

    # ── A4 margins ──
    sec = doc.sections[0]
    sec.page_width  = Cm(21.0);  sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.0);   sec.right_margin  = Cm(2.0)
    sec.top_margin  = Cm(2.0);   sec.bottom_margin = Cm(2.0)

    if report_settings is None:
        report_settings = {}

    num_years     = len(traffic_df)
    pavement_text = "Rigid Pavement" if pavement_type == 'rigid' else "Flexible Pavement"
    pavement_thai = "แบบแข็ง"       if pavement_type == 'rigid' else "ยืดหยุ่น"

    # เลขหัวข้อและตาราง
    if pavement_type == 'flexible':
        section_num = report_settings.get('flex_section_number', '4.2.2')
        tbl_base    = report_settings.get('flex_table_start',    '4-1')
    else:
        section_num = report_settings.get('rigid_section_number', '4.2.3')
        tbl_base    = report_settings.get('rigid_table_start',    '4-1')

    tbl_param = tbl_base
    tbl_tf    = increment_table_number(tbl_base, 1)
    tbl_traf  = increment_table_number(tbl_base, 2)
    tbl_esal  = increment_table_number(tbl_base, 3)

    # label ต่อ param
    _D_CM_RPT = {10:25, 11:28, 12:30, 13:32, 14:35, 15:38, 16:40}
    def p_lbl(p):
        if pavement_type == 'rigid':
            cm = _D_CM_RPT.get(p, '')
            return f'D = {p} นิ้ว ({cm} cm)' if cm else f'D = {p} นิ้ว'
        return f'SN = {p}'
    def p_col(p):
        if pavement_type == 'rigid':
            cm = _D_CM_RPT.get(p, '')
            return f'D={p}" ({cm} cm)' if cm else f'D={p}"'
        return f'SN={p}'

    # ── helpers ──
    def set_run(run, bold=False, sz=FS):
        run.font.name = FN;  run.font.size = Pt(sz);  run.bold = bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FN)

    def set_cell(cell, bold=False, align=None, sz=TFS, shading=None):
        for para in cell.paragraphs:
            if align: para.alignment = align
            for r in para.runs:
                r.font.name = FN;  r.font.size = Pt(sz);  r.bold = bold
                r._element.rPr.rFonts.set(qn('w:eastAsia'), FN)
        if shading:
            xml = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shading}"/>')
            cell._tc.get_or_add_tcPr().append(xml)

    def add_caption(num_str, title_str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"ตารางที่ {num_str} "); set_run(r1, bold=True)
        r2 = p.add_run(title_str);              set_run(r2, bold=False)

    def thai_distribute(para):
        pPr = para._element.get_or_add_pPr()
        jc  = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'thaiDistribute')
        pPr.append(jc)

    HDR_COLOR  = "D9E2F3"
    FOOT_COLOR = "C6EFCE"   # สีเขียวอ่อน สำหรับ footer ACC.ESAL

    # ══════════════════════════════════════════════════════
    # 1. หัวข้อ
    # ══════════════════════════════════════════════════════
    h = doc.add_paragraph()
    h.paragraph_format.space_after = Pt(6)
    r = h.add_run(f"{section_num}	"); set_run(r, bold=True)
    r = h.add_run(f"ปริมาณเพลามาตรฐาน (ESALs) ระยะเวลาออกแบบ {num_years} ปี ผิวทาง {pavement_text}")
    set_run(r, bold=True)

    # ── บทเกริ่นนำ ──
    intro = doc.add_paragraph()
    intro.paragraph_format.first_line_indent = Cm(1.25)
    intro.paragraph_format.space_after = Pt(6)
    thai_distribute(intro)
    if pavement_type == 'flexible':
        txt = (f"ในการคำนวณปริมาณเพลามาตรฐาน สำหรับผิวทางยืดหยุ่น ที่ปรึกษาได้กำหนดค่าพารามิเตอร์ต่าง ๆ "
               f"และค่า Truck Factor ของรถบรรทุกหนัก ที่ใช้สำหรับการคำนวณ ดังแสดงในตารางที่ {tbl_param} และ {tbl_tf} "
               f"ดังนั้นค่าปริมาณเพลามาตรฐาน สำหรับผิวทางยืดหยุ่น ที่ระยะเวลาออกแบบ {num_years} ปี "
               f"แสดงดังตารางที่ {tbl_traf} และ {tbl_esal}")
    else:
        txt = (f"ในการคำนวณปริมาณเพลามาตรฐานสำหรับผิวทางแบบแข็งหรือผิวทางคอนกรีต "
               f"โดยที่ปรึกษาได้กำหนดค่าพารามิเตอร์ต่าง ๆ และค่า Truck Factor ของรถบรรทุกหนัก "
               f"ที่ใช้สำหรับการคำนวณ ดังแสดงในตารางที่ {tbl_param} และ {tbl_tf} "
               f"ดังนั้นค่าปริมาณเพลามาตรฐาน สำหรับผิวทางแบบแข็ง ที่ระยะเวลาออกแบบ {num_years} ปี "
               f"แสดงดังตารางที่ {tbl_traf} และ {tbl_esal}")
    intro.add_run(txt).font.name = FN
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════
    # ตารางที่ tbl_param — พารามิเตอร์
    # ══════════════════════════════════════════════════════
    add_caption(tbl_param, "ค่าพารามิเตอร์ต่าง ๆ ที่ใช้สำหรับการคำนวณ")
    params_str = " | ".join(p_lbl(p) for p in param_list)
    param_data = [
        ("รายการ",        "ค่า"),
        ("ประเภทผิวทาง",  pavement_text),
        ("pt",            str(pt)),
        ("พารามิเตอร์",   params_str),
        ("Lane Factor",   str(lane_factor)),
        ("Direction Factor", str(direction_factor)),
        ("จำนวนปี",       str(num_years)),
    ]
    t1 = doc.add_table(rows=len(param_data), cols=2)
    t1.style = 'Table Grid';  t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (lbl, val) in enumerate(param_data):
        t1.rows[i].cells[0].text = lbl
        t1.rows[i].cells[1].text = val
        sh = HDR_COLOR if i == 0 else None
        set_cell(t1.rows[i].cells[0], bold=(i==0), sz=TFS, shading=sh)
        set_cell(t1.rows[i].cells[1], bold=(i==0), sz=TFS, shading=sh)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════
    # ตารางที่ tbl_tf — Truck Factor (3 คอลัมน์ TF)
    # คำนวณ TF ใหม่จาก pt ที่รับมาโดยตรง (ป้องกัน pt เก่าจาก multi_results)
    # ══════════════════════════════════════════════════════
    add_caption(tbl_tf, "ค่า Truck Factor ของรถบรรทุกหนัก")
    tf_headers = ["รหัส", "ประเภท"] + [p_col(p) for p in param_list]
    n_tf_cols  = len(tf_headers)
    t2 = doc.add_table(rows=len(TRUCKS)+1, cols=n_tf_cols)
    t2.style = 'Table Grid';  t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    for j, h in enumerate(tf_headers):
        t2.rows[0].cells[j].text = h
        set_cell(t2.rows[0].cells[j], bold=True, sz=TFS,
                 align=WD_ALIGN_PARAGRAPH.CENTER, shading=HDR_COLOR)
    # data rows — คำนวณ TF จาก pt argument โดยตรง
    for i, code in enumerate(TRUCKS.keys()):
        row = t2.rows[i+1]
        row.cells[0].text = code
        row.cells[1].text = TRUCKS[code]['desc']
        set_cell(row.cells[0], sz=TFS, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(row.cells[1], sz=TFS)
        for k, p in enumerate(param_list):
            tf_val = get_default_truck_factor(code, pavement_type, float(pt), p)
            row.cells[2+k].text = f"{tf_val:.3f}"
            set_cell(row.cells[2+k], sz=TFS, align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_paragraph()

    # ── คำนวณ multi_results ใหม่จาก pt ที่รับมา (ป้องกันค่าเก่า) ──
    multi_results_word = {}
    for p in param_list:
        tf_p = {code: get_default_truck_factor(code, pavement_type, float(pt), p)
                for code in TRUCKS.keys()}
        r_df, t_esal = calculate_esal_with_acc(
            traffic_df, tf_p, lane_factor, direction_factor
        )
        multi_results_word[p] = (r_df, t_esal)
    # ใช้ multi_results_word แทน multi_results ที่รับมา
    multi_results = multi_results_word

    # ══════════════════════════════════════════════════════
    # สมการ W18 (คงเดิม — เรียก _add_esal_formula_block)
    # ══════════════════════════════════════════════════════
    _add_esal_formula_block(doc, num_years, FN, FS)

    # ══════════════════════════════════════════════════════
    # ตารางที่ tbl_traf — ปริมาณจราจรรายปี
    # ══════════════════════════════════════════════════════
    add_caption(tbl_traf,
                f"ปริมาณจราจรรายปี ระยะเวลาออกแบบ {num_years} ปี")
    traf_headers = ["Year", "MB", "HB", "MT", "HT", "TR", "STR", "ADTT"]
    t3 = doc.add_table(rows=num_years+1, cols=len(traf_headers))
    t3.style = 'Table Grid';  t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(traf_headers):
        t3.rows[0].cells[j].text = h
        set_cell(t3.rows[0].cells[j], bold=True, sz=TFS_SM,
                 align=WD_ALIGN_PARAGRAPH.CENTER, shading=HDR_COLOR)
    ref_df = list(multi_results.values())[0][0]   # ใช้ results_df ของ param แรก
    for i, row_data in ref_df.iterrows():
        row = t3.rows[i+1]
        adtt = int(sum(row_data.get(c, 0) for c in ['MB','HB','MT','HT','TR','STR']))
        vals = [int(row_data.get('Year', i+1)),
                int(row_data.get('MB',0)), int(row_data.get('HB',0)),
                int(row_data.get('MT',0)), int(row_data.get('HT',0)),
                int(row_data.get('TR',0)), int(row_data.get('STR',0)),
                adtt]
        for j, v in enumerate(vals):
            row.cells[j].text = f"{v:,}" if j > 0 else str(v)
            align = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            set_cell(row.cells[j], sz=TFS_SM, align=align)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════
    # ตารางที่ tbl_esal — ESAL รายปี + footer ACC.ESAL
    # ══════════════════════════════════════════════════════
    add_caption(tbl_esal,
                f"ค่าปริมาณเพลามาตรฐาน สำหรับผิวทาง{pavement_thai} "
                f"ที่ระยะเวลาออกแบบ {num_years} ปี")
    esal_headers = ["Year", "ADTT"] + [f"ESAL ({p_col(p)})" for p in param_list]
    n_ec = len(esal_headers)
    t4 = doc.add_table(rows=num_years+2, cols=n_ec)  # +2 = header + footer
    t4.style = 'Table Grid';  t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for j, h in enumerate(esal_headers):
        t4.rows[0].cells[j].text = h
        set_cell(t4.rows[0].cells[j], bold=True, sz=TFS_SM,
                 align=WD_ALIGN_PARAGRAPH.CENTER, shading=HDR_COLOR)
    # data rows
    for i, row_data in ref_df.iterrows():
        row = t4.rows[i+1]
        adtt = int(sum(row_data.get(c,0) for c in ['MB','HB','MT','HT','TR','STR']))
        row.cells[0].text = str(int(row_data.get('Year', i+1)))
        set_cell(row.cells[0], sz=TFS_SM, align=WD_ALIGN_PARAGRAPH.CENTER)
        row.cells[1].text = f"{adtt:,}"
        set_cell(row.cells[1], sz=TFS_SM, align=WD_ALIGN_PARAGRAPH.RIGHT)
        for k, p in enumerate(param_list):
            r_df, _ = multi_results[p]
            esal_val = int(r_df.iloc[i]['ESAL'])
            row.cells[2+k].text = f"{esal_val:,}"
            set_cell(row.cells[2+k], sz=TFS_SM, align=WD_ALIGN_PARAGRAPH.RIGHT)
    # footer row — ACC.ESAL
    foot = t4.rows[num_years+1]
    foot.cells[0].text = "ACC. ESAL"
    set_cell(foot.cells[0], bold=True, sz=TFS_SM,
             align=WD_ALIGN_PARAGRAPH.CENTER, shading=FOOT_COLOR)
    foot.cells[1].text = ""
    set_cell(foot.cells[1], sz=TFS_SM, shading=FOOT_COLOR)
    for k, p in enumerate(param_list):
        _, total = multi_results[p]
        foot.cells[2+k].text = f"{total:,}"
        set_cell(foot.cells[2+k], bold=True, sz=TFS_SM,
                 align=WD_ALIGN_PARAGRAPH.RIGHT, shading=FOOT_COLOR)
    doc.add_paragraph()

    # ── Footer ──
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("พัฒนาเพื่อการเรียนการสอนโดย รศ.ดร.อิทธิพล มีผล ภาควิชาครุศาสตร์โยธา มจพ.")
    r.font.name = FN;  r.font.size = Pt(14);  r.italic = True
    r._element.rPr.rFonts.set(qn('w:eastAsia'), FN)

    out = __import__('io').BytesIO()
    doc.save(out);  out.seek(0)
    return out

def save_project(pavement_type, pt, param, lane_factor, direction_factor, truck_factors, traffic_df,
                 report_settings=None, comb_rigid_params=None):
    """บันทึก Project เป็น JSON
    comb_rigid_params = dict ของ rigid params สำหรับรายงานรวม (เฉพาะกรณี flexible)
    """
    project = {
        'version': '3.0',
        'created': datetime.now().isoformat(),
        'pavement_type': pavement_type,
        'pt': pt,
        'param': param,
        'lane_factor': lane_factor,
        'direction_factor': direction_factor,
        'truck_factors': truck_factors,
        'traffic_data': traffic_df.to_dict('records'),
    }
    if report_settings:
        project['report_settings'] = report_settings
    # บันทึก rigid params สำหรับรายงานรวม (เฉพาะกรณี flexible)
    if comb_rigid_params:
        project['comb_rigid_params'] = comb_rigid_params
    return json.dumps(project, ensure_ascii=False, indent=2)


def load_project(json_data):
    """โหลด Project จาก JSON"""
    try:
        project = json.loads(json_data)
        return project
    except (json.JSONDecodeError, KeyError, TypeError) as e:
        logger.error(f"load_project error: {e}")
        return None


def get_all_truck_factors_table(pavement_type, pt):
    """สร้างตาราง Truck Factor ทั้งหมดสำหรับแสดงผล
    คำนวณจากสมการ AASHTO 1993 โดยตรง (ไม่ใช้ lookup table)
    """
    data = []
    params      = [10, 11, 12, 13, 14, 15, 16] if pavement_type == 'rigid' else [4, 5, 6, 7, 8, 9]
    param_label = 'D' if pavement_type == 'rigid'  else 'SN'

    for code in TRUCKS.keys():
        row = {'ประเภท': code, 'รายละเอียด': TRUCKS[code]['desc']}
        for p in params:
            col_name = f'{param_label}={p}"' if pavement_type == 'rigid' else f'{param_label}={p}'
            row[col_name] = f"{get_default_truck_factor(code, pavement_type, pt, p):.3f}"
        data.append(row)

    return pd.DataFrame(data)


# ============================================================
# Preview HTML for intro paragraph
# ============================================================
def generate_intro_preview_html(pavement_type, num_years, tbl_param, tbl_tf, tbl_esal, section_num):
    """สร้าง HTML preview ของบทเกริ่นนำ พร้อม highlight สี"""
    
    PURPLE = "background-color: #D8B4FE; padding: 1px 4px; border-radius: 3px; font-weight: bold;"
    YELLOW = "background-color: #FDE68A; padding: 1px 4px; border-radius: 3px; font-weight: bold;"
    
    if pavement_type == 'flexible':
        pavement_thai = "ยืดหยุ่น"
        intro_html = (
            f'<span style="{YELLOW}">{section_num}</span>&nbsp;&nbsp;'
            f'<b>ปริมาณเพลามาตรฐาน (ESALs) ระยะเวลาออกแบบ '
            f'<span style="{PURPLE}">{num_years}</span> ปี ผิวทาง Flexible Pavement</b>'
            f'<br><br>'
            f'<p style="text-indent: 40px; text-align: justify; text-justify: inter-character; margin: 0;">'
            f'ในการคำนวณปริมาณเพลามาตรฐาน สำหรับผิวทาง{pavement_thai} '
            f'ที่ปรึกษาได้กำหนดค่าพารามิเตอร์ต่าง ๆ และค่า Truck Factor '
            f'ของรถบรรทุกหนัก ที่ใช้สำหรับการคำนวณ ดังแสดงในตารางที่ '
            f'<span style="{YELLOW}">{tbl_param}</span> และ '
            f'<span style="{YELLOW}">{tbl_tf}</span> '
            f'ดังนั้นค่าปริมาณเพลามาตรฐาน สำหรับผิวทาง{pavement_thai} '
            f'ที่ระยะเวลาออกแบบ <span style="{PURPLE}">{num_years}</span> ปี '
            f'แสดงดังตารางที่ <span style="{YELLOW}">{tbl_esal}</span></p>'
        )
    else:
        pavement_thai = "แบบแข็ง"
        intro_html = (
            f'<span style="{YELLOW}">{section_num}</span>&nbsp;&nbsp;'
            f'<b>ปริมาณเพลามาตรฐาน (ESALs) ระยะเวลาออกแบบ '
            f'<span style="{PURPLE}">{num_years}</span> ปี ผิวทาง Rigid Pavement</b>'
            f'<br><br>'
            f'<p style="text-indent: 40px; text-align: justify; text-justify: inter-character; margin: 0;">'
            f'ในการคำนวณปริมาณเพลามาตรฐานสำหรับผิวทางแบบแข็งหรือผิวทางคอนกรีต '
            f'โดยที่ปรึกษาได้กำหนดค่าพารามิเตอร์ต่าง ๆ และค่า Truck Factor '
            f'ของรถบรรทุกหนัก ที่ใช้สำหรับการคำนวณ ดังแสดงในตารางที่ '
            f'<span style="{YELLOW}">{tbl_param}</span> และ '
            f'<span style="{YELLOW}">{tbl_tf}</span> '
            f'ดังนั้นค่าปริมาณเพลามาตรฐาน สำหรับผิวทาง{pavement_thai} '
            f'ที่ระยะเวลาออกแบบ <span style="{PURPLE}">{num_years}</span> ปี '
            f'แสดงดังตารางที่ <span style="{YELLOW}">{tbl_esal}</span></p>'
        )
    
    return f'''
    <div style="background: #f9f9f9; padding: 15px; border-radius: 8px; border: 1px solid #ddd;
                font-family: 'TH SarabunPSK', sans-serif; font-size: 15px; line-height: 1.8;">
        {intro_html}
    </div>
    '''


# ============================================================
# Streamlit App
# ============================================================
def main():
    st.set_page_config(
        page_title="ESAL Calculator - AASHTO 1993 v3.0",
        page_icon="🛣️",
        layout="wide"
    )
    
    st.markdown("""
    <style>
    .main-header { font-size: 2.5rem; font-weight: bold; color: #1E3A5F; text-align: center; margin-bottom: 0.5rem; }
    .sub-header { font-size: 1.2rem; color: #4A6FA5; text-align: center; margin-bottom: 2rem; }
    .metric-box { background: linear-gradient(135deg, #1E3A5F 0%, #4A6FA5 100%); padding: 1rem; border-radius: 10px; color: white; text-align: center; margin: 0.5rem 0; min-width: 0; }
    .metric-value { font-size: 1.6rem; font-weight: bold; white-space: nowrap; overflow: hidden; text-overflow: ellipsis; }
    .metric-label { font-size: 0.85rem; opacity: 0.9; }
    </style>
    """, unsafe_allow_html=True)
    
    st.markdown('<p class="main-header">🛣️ ESAL Calculator v3.0</p>', unsafe_allow_html=True)
    st.markdown('<p class="sub-header">คำนวณปริมาณเพลาเดี่ยวมาตรฐานเทียบเท่า ตามมาตรฐาน AASHTO 1993</p>', unsafe_allow_html=True)
    
    # Initialize session state
    if 'traffic_df' not in st.session_state:
        st.session_state['traffic_df'] = None
    
    # Sidebar
    with st.sidebar:
        st.header("⚙️ พารามิเตอร์การคำนวณ")
        
        # Project Load/Save
        st.subheader("📁 Project")
        
        uploaded_project = st.file_uploader("📥 โหลด Project", type=['json'], key='load_project')
        if uploaded_project is not None:
            try:
                file_id = f"{uploaded_project.name}_{uploaded_project.size}"
                if st.session_state.get('last_uploaded_file') != file_id:
                    st.session_state['last_uploaded_file'] = file_id
                    
                    project = load_project(uploaded_project.read().decode('utf-8'))
                    if project:
                        st.session_state['input_pavement_type'] = project.get('pavement_type', 'rigid')
                        st.session_state['input_pt'] = project.get('pt', 2.5)
                        st.session_state['input_param'] = project.get('param', 12)
                        st.session_state['input_lane_factor'] = project.get('lane_factor', 0.9)
                        st.session_state['input_direction_factor'] = project.get('direction_factor', 0.5)
                        st.session_state['loaded_tf'] = project.get('truck_factors', {})

                        # โหลด comb_rigid_params (สำหรับ flexible project)
                        crp = project.get('comb_rigid_params', {})
                        if crp:
                            st.session_state['comb_rigid_pt']   = crp.get('pt',   2.5)
                            st.session_state['comb_rigid_d']    = crp.get('param', 13)
                            st.session_state['comb_rigid_lane'] = crp.get('lane_factor',      0.9)
                            st.session_state['comb_rigid_dir']  = crp.get('direction_factor', 0.5)
                        
                        # โหลด report_settings
                        rs = project.get('report_settings', {})
                        if rs:
                            for key, val in rs.items():
                                st.session_state[f'input_{key}'] = val
                        
                        loaded_traffic = project.get('traffic_data', None)
                        if loaded_traffic:
                            st.session_state['traffic_df'] = pd.DataFrame(loaded_traffic)
                        
                        st.success("✅ โหลด Project สำเร็จ!")
                        st.rerun()
                    else:
                        st.error("❌ ไม่สามารถอ่านไฟล์ได้")
            except Exception as e:
                st.error(f"❌ เกิดข้อผิดพลาด: {e}")
        
        default_pavement = st.session_state.get('input_pavement_type', 'rigid')
        default_pt = st.session_state.get('input_pt', 2.5)
        default_param = st.session_state.get('input_param', 12 if default_pavement == 'rigid' else 7)
        default_lane = st.session_state.get('input_lane_factor', 0.9)
        default_dir = st.session_state.get('input_direction_factor', 0.5)
        loaded_tf = st.session_state.get('loaded_tf', {})
        
        st.divider()
        
        pavement_options = ['rigid', 'flexible']
        pavement_idx = pavement_options.index(default_pavement) if default_pavement in pavement_options else 0
        pavement_type = st.selectbox(
            "ประเภทผิวทาง",
            options=pavement_options,
            index=pavement_idx,
            format_func=lambda x: '🧱 Rigid Pavement (คอนกรีต)' if x == 'rigid' else '🛤️ Flexible Pavement (ลาดยาง)',
            key="input_pavement_type"
        )
        
        pt_options = [2.0, 2.5, 3.0]
        # แปลง default_pt เป็น float ก่อนเปรียบเทียบ ป้องกัน int vs float mismatch
        try:
            _default_pt_f = float(default_pt)
            pt_idx = next((i for i, v in enumerate(pt_options) if abs(v - _default_pt_f) < 1e-9), 1)
        except (TypeError, ValueError):
            pt_idx = 1
        pt = st.selectbox(
            "Terminal Serviceability (pt)",
            options=pt_options,
            index=pt_idx,
            format_func=lambda x: f"pt = {x}",
            key="input_pt"
        )
        
        if pavement_type == 'rigid':
            # ── Multi-select D: เลือกได้ 3 ค่า จาก 10–16 นิ้ว ──
            rigid_d_options = [10, 11, 12, 13, 14, 15, 16]
            # ใช้ saved_params_rigid เก็บค่าแยกจาก widget key
            # ป้องกัน reset เมื่อสลับ pavement_type
            if 'saved_params_rigid' not in st.session_state:
                st.session_state['saved_params_rigid'] = [10, 11, 12, 13]
            _saved_r = st.session_state['saved_params_rigid']
            # backward-compat: JSON เก่าบันทึก param เดี่ยว
            if isinstance(_saved_r, int):
                _saved_r = [_saved_r]
            _saved_r = [d for d in _saved_r if d in rigid_d_options] or [11, 12, 13]
            _D_CM = {10:25, 11:28, 12:30, 13:32, 14:35, 15:38, 16:40}
            params_selected = st.multiselect(
                "ความหนาพื้นคอนกรีต D (เลือก 4 ค่า)",
                options=rigid_d_options,
                default=_saved_r,
                format_func=lambda x: f"D = {x} นิ้ว ({_D_CM.get(x,'')} cm)",
                key="input_params_rigid",
            )
            # บันทึกทันทีที่เปลี่ยน
            st.session_state['saved_params_rigid'] = params_selected
            if len(params_selected) == 0:
                st.warning("⚠️ กรุณาเลือกอย่างน้อย 1 ค่า")
                params_selected = [12]
            elif len(params_selected) > 4:
                st.warning("⚠️ เลือกได้สูงสุด 4 ค่า — ใช้ 4 ค่าแรก")
                params_selected = params_selected[:4]
            param        = params_selected[0]          # ค่าแรก (ใช้กับ TF sidebar + combined report)
            param_list   = params_selected             # list ใช้คำนวณ multi
            param_label  = ", ".join(f'D={d}"' for d in param_list)
        else:
            # ── 3 number_input SN ทศนิยม ──
            st.caption("กำหนด Structural Number (SN) 3 ค่า")
            # ใช้ saved_sn_list เก็บค่าแยกจาก widget key
            if 'saved_sn_list' not in st.session_state:
                st.session_state['saved_sn_list'] = [6.5, 7.0, 7.5]
            _saved_sn = st.session_state['saved_sn_list']
            if isinstance(_saved_sn, (int, float)):
                _saved_sn = [float(_saved_sn), float(_saved_sn)+0.5, float(_saved_sn)+1.0]
            while len(_saved_sn) < 3:
                _saved_sn.append(_saved_sn[-1] + 0.5)
            sn_cols = st.columns(3)
            sn_vals = []
            for i, col in enumerate(sn_cols):
                with col:
                    v = st.number_input(f"SN {i+1}", value=float(_saved_sn[i]),
                                        min_value=1.0, max_value=20.0, step=0.1,
                                        format="%.1f", key=f"input_sn_{i}")
                    sn_vals.append(round(v, 2))
            # บันทึกทันทีที่เปลี่ยน
            st.session_state['saved_sn_list'] = sn_vals
            st.session_state['input_sn_list'] = sn_vals
            param        = sn_vals[0]                  # ค่าแรก (ใช้กับ TF sidebar + combined report)
            param_list   = sn_vals                     # list ใช้คำนวณ multi
            param_label  = ", ".join(f'SN={s}' for s in param_list)
        
        st.divider()
        
        st.subheader("🚗 ค่าสัดส่วน")
        lane_factor = st.slider(
            "Lane Distribution Factor", 
            0.1, 1.0, 
            value=st.session_state.get('input_lane_factor', default_lane), 
            step=0.05,
            key="input_lane_factor"
        )
        direction_factor = st.slider(
            "Directional Distribution Factor", 
            0.1, 1.0, 
            value=st.session_state.get('input_direction_factor', default_dir), 
            step=0.05,
            key="input_direction_factor"
        )
        
        st.divider()
        
        st.subheader("🚛 ค่า Truck Factor")
        
        tf_key = f"tf_{pavement_type}_{pt}_{param}"
        
        if loaded_tf and tf_key not in st.session_state:
            st.session_state[tf_key] = {}
            for code in TRUCKS.keys():
                if code in loaded_tf:
                    st.session_state[tf_key][code] = loaded_tf[code]
                else:
                    st.session_state[tf_key][code] = get_default_truck_factor(code, pavement_type, pt, param)
        elif tf_key not in st.session_state:
            st.session_state[tf_key] = {}
            for code in TRUCKS.keys():
                st.session_state[tf_key][code] = get_default_truck_factor(code, pavement_type, pt, param)
        
        if st.button("🔄 Reset เป็นค่า Default", use_container_width=True):
            for code in TRUCKS.keys():
                st.session_state[tf_key][code] = get_default_truck_factor(code, pavement_type, pt, param)
            st.rerun()
        
        st.caption("กรอกค่า Truck Factor (แก้ไขได้)")
        
        truck_factors = {}
        for code in TRUCKS.keys():
            default_val = get_default_truck_factor(code, pavement_type, pt, param)
            current_val = st.session_state[tf_key].get(code, default_val)
            
            new_val = st.number_input(
                f"{code}",
                min_value=0.0,
                max_value=50.0,
                value=float(current_val),
                step=0.0001,
                format="%.3f",
                key=f"input_{tf_key}_{code}",
                help=f"{TRUCKS[code]['desc']} | Default: {default_val:.3f}"
            )
            
            st.session_state[tf_key][code] = new_val
            truck_factors[code] = new_val
        
        st.divider()
        
        st.subheader("📥 ดาวน์โหลด Template")
        template_df = create_template()
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            template_df.to_excel(writer, index=False, sheet_name='Traffic Data')
        st.download_button(
            label="📄 ดาวน์โหลด Template Excel",
            data=output.getvalue(),
            file_name="traffic_template.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )
    
    # Main Content
    tab1, tab2, tab3 = st.tabs(["📊 คำนวณ ESAL", "🚛 ข้อมูล Truck Factor", "📘 คู่มือ"])
    
    with tab1:
        col1, col2 = st.columns([1, 2])
        
        with col1:
            st.subheader("📤 อัพโหลดข้อมูลปริมาณจราจร")
            
            uploaded_file = st.file_uploader(
                "เลือกไฟล์ Excel",
                type=['xlsx', 'xls'],
                help="อัพโหลดไฟล์ Excel (หน่วย: คัน/วัน)"
            )
            
            if 'use_sample' not in st.session_state:
                st.session_state['use_sample'] = False
            
            if uploaded_file is not None:
                try:
                    traffic_df = pd.read_excel(uploaded_file)
                    st.session_state['traffic_df'] = traffic_df
                    st.success("✅ อัพโหลดสำเร็จ!")
                    st.session_state['use_sample'] = False
                except Exception as e:
                    st.error(f"❌ เกิดข้อผิดพลาด: {e}")
                    traffic_df = st.session_state.get('traffic_df', None)
            elif st.session_state.get('traffic_df') is not None:
                traffic_df = st.session_state['traffic_df']
            else:
                st.info("📌 อัพโหลดไฟล์ Excel หรือใช้ข้อมูลตัวอย่าง")
                
                if st.button("🔄 ใช้ข้อมูลตัวอย่าง", use_container_width=True):
                    st.session_state['use_sample'] = True
                    st.session_state['traffic_df'] = create_template()
                
                traffic_df = st.session_state.get('traffic_df', None)
            
            if traffic_df is not None:
                st.write("**ข้อมูลปริมาณจราจร (คัน/วัน):**")
                st.dataframe(traffic_df, use_container_width=True, height=350)
        
        with col2:
            num_years_disp = len(traffic_df) if traffic_df is not None else 0
            suffix = f" — ระยะเวลาออกแบบ {num_years_disp} ปี" if num_years_disp > 0 else ""
            st.subheader(f"📈 ผลการคำนวณ ESAL{suffix}")
            
            if traffic_df is not None:
                # ── คำนวณ multi-param: loop ทุก D หรือ SN ──
                multi_results = {}   # {param_val: (results_df, total_esal)}
                for p in param_list:
                    tf_p = {code: get_default_truck_factor(code, pavement_type, pt, p)
                            for code in TRUCKS.keys()}
                    r_df, t_esal = calculate_esal_with_acc(
                        traffic_df, tf_p, lane_factor, direction_factor
                    )
                    multi_results[p] = (r_df, t_esal)

                # ใช้ค่าแรกสำหรับ metric หลัก
                results_df, total_esal = multi_results[param_list[0]]

                # Metrics — ประเภทผิวทาง + ESAL 3 ค่า (ตัดจำนวนปีออก ย้ายไปหัวข้อแล้ว)
                n_params = len(param_list)
                metric_cols = st.columns(n_params + 1)
                with metric_cols[0]:
                    pv_lbl = "Rigid" if pavement_type == 'rigid' else "Flexible"
                    st.markdown(f"""
                    <div class="metric-box">
                        <div class="metric-value">{pv_lbl}</div>
                        <div class="metric-label">ประเภทผิวทาง</div>
                    </div>""", unsafe_allow_html=True)
                for i, p in enumerate(param_list):
                    _, t = multi_results[p]
                    p_lbl = f'D={p}"' if pavement_type == 'rigid' else f"SN={p}"
                    with metric_cols[1 + i]:
                        st.markdown(f"""
                        <div class="metric-box">
                            <div class="metric-value">{t:,}</div>
                            <div class="metric-label">ESAL – {p_lbl}</div>
                        </div>""", unsafe_allow_html=True)

                st.divider()

                # Truck Factor Table — แสดงทุก param
                st.write("**🚛 ค่า Truck Factor ตาม AASHTO 1993 สมการจริง:**")
                tf_display = []
                for code in TRUCKS.keys():
                    row = {'รหัส': code, 'ประเภท': TRUCKS[code]['desc']}
                    for p in param_list:
                        p_lbl = f'D={p}"' if pavement_type == 'rigid' else f"SN={p}"
                        row[p_lbl] = f"{get_default_truck_factor(code, pavement_type, pt, p):.3f}"
                    tf_display.append(row)
                st.dataframe(pd.DataFrame(tf_display), use_container_width=True, hide_index=True)

                st.divider()

                # ESAL Results Table — รวม ESAL ทุก param ในตารางเดียว
                st.write("**📊 ESAL รายปี (ทุก param):**")
                # สร้าง display_df รวม: Year, MB–STR, ADTT, ESAL(p1), ACC(p1), ...
                base_df = results_df[['Year','MB','HB','MT','HT','TR','STR']].copy()
                base_df['ADTT'] = base_df[['MB','HB','MT','HT','TR','STR']].sum(axis=1)
                for p in param_list:
                    r_df, _ = multi_results[p]
                    p_lbl = f'D={p}"' if pavement_type == 'rigid' else f"SN={p}"
                    base_df[f'ESAL({p_lbl})']     = r_df['ESAL'].values
                    base_df[f'ACC.ESAL({p_lbl})'] = r_df['ACC_ESAL'].values
                base_df = base_df.rename(columns={'Year': 'ปีที่'})
                st.dataframe(base_df, use_container_width=True, height=400)
                
                st.divider()
                
                # ============================================================
                # ส่วนตั้งค่ารายงาน Word (Expander)
                # ============================================================
                with st.expander("📝 ตั้งค่ารายงาน Word", expanded=False):
                    st.markdown("#### หมายเลขหัวข้อและตาราง")
                    st.caption("กรอกเลขตารางเริ่มต้น ระบบจะเพิ่ม +1, +2 อัตโนมัติ (เช่น 4-1 → 4-2 → 4-3)")
                    
                    col_f1, col_f2, col_f3, col_f4 = st.columns(4)
                    
                    with col_f1:
                        flex_section_number = st.text_input(
                            "🛤️ เลขหัวข้อ Flexible",
                            value=st.session_state.get('input_flex_section_number', "4.2.2"),
                            key="input_flex_section_number"
                        )
                    with col_f2:
                        flex_table_start = st.text_input(
                            "🛤️ เลขตารางเริ่มต้น Flexible",
                            value=st.session_state.get('input_flex_table_start', "4-1"),
                            key="input_flex_table_start"
                        )
                    with col_f3:
                        rigid_section_number = st.text_input(
                            "🧱 เลขหัวข้อ Rigid",
                            value=st.session_state.get('input_rigid_section_number', "4.2.3"),
                            key="input_rigid_section_number"
                        )
                    with col_f4:
                        rigid_table_start = st.text_input(
                            "🧱 เลขตารางเริ่มต้น Rigid",
                            value=st.session_state.get('input_rigid_table_start', "4-4"),
                            key="input_rigid_table_start"
                        )
                    
                    # แสดงสรุปหมายเลขตาราง
                    col_sum1, col_sum2 = st.columns(2)
                    with col_sum1:
                        st.info(
                            f"**Flexible:** ตารางที่ {flex_table_start} (พารามิเตอร์), "
                            f"{increment_table_number(flex_table_start, 1)} (Truck Factor), "
                            f"{increment_table_number(flex_table_start, 2)} (จราจร), "
                            f"{increment_table_number(flex_table_start, 3)} (ESAL)"
                        )
                    with col_sum2:
                        st.info(
                            f"**Rigid:** ตารางที่ {rigid_table_start} (พารามิเตอร์), "
                            f"{increment_table_number(rigid_table_start, 1)} (Truck Factor), "
                            f"{increment_table_number(rigid_table_start, 2)} (จราจร), "
                            f"{increment_table_number(rigid_table_start, 3)} (ESAL)"
                        )
                    
                    st.markdown("---")
                    
                    # ===== พารามิเตอร์สำหรับ Combined Report =====
                    st.markdown("#### ⚙️ พารามิเตอร์สำหรับรายงานรวม (Flexible + Rigid)")
                    st.caption("กำหนดพารามิเตอร์ของผิวทางอีกประเภทหนึ่ง สำหรับ export รายงานรวม")
                    
                    if pavement_type == 'rigid':
                        # ตั้งอยู่ Rigid → กรอก Flexible SN 3 ช่อง
                        st.markdown("**🛤️ พารามิเตอร์ Flexible Pavement (สำหรับรายงานรวม)**")
                        col_p1, col_p2, col_p3 = st.columns([1, 2, 1])
                        with col_p1:
                            _cf_pt_def = st.session_state.get('comb_flex_pt', pt)
                            try:
                                _cf_pt_idx = [2.0, 2.5, 3.0].index(float(_cf_pt_def))
                            except (ValueError, TypeError):
                                _cf_pt_idx = [2.0, 2.5, 3.0].index(pt) if pt in [2.0, 2.5, 3.0] else 1
                            comb_flex_pt = st.selectbox("pt (Flexible)", [2.0, 2.5, 3.0],
                                                        index=_cf_pt_idx, key="comb_flex_pt")
                        with col_p2:
                            st.caption("SN (3 ค่า)")
                            _cf_sn_def = st.session_state.get('comb_flex_sn_list', [6.5, 7.0, 7.5])
                            if isinstance(_cf_sn_def, (int, float)):
                                _cf_sn_def = [float(_cf_sn_def)] * 3
                            while len(_cf_sn_def) < 3:
                                _cf_sn_def.append(_cf_sn_def[-1] + 0.5)
                            _cf_sn_cols = st.columns(3)
                            comb_flex_sn_list = []
                            for _i, _c in enumerate(_cf_sn_cols):
                                with _c:
                                    comb_flex_sn_list.append(round(st.number_input(
                                        f"SN {_i+1}", value=float(_cf_sn_def[_i]),
                                        min_value=1.0, max_value=20.0, step=0.1,
                                        format="%.1f", key=f"comb_flex_sn_{_i}"), 2))
                            st.session_state['comb_flex_sn_list'] = comb_flex_sn_list
                        with col_p3:
                            comb_flex_lane = st.number_input("Lane Factor", 0.1, 1.0,
                                                              value=lane_factor, step=0.05,
                                                              key="comb_flex_lane")
                            comb_flex_dir  = st.number_input("Direction Factor", 0.5, 1.0,
                                                              value=direction_factor, step=0.1,
                                                              key="comb_flex_dir")
                    else:
                        # ตั้งอยู่ Flexible → กรอก Rigid D multiselect
                        st.markdown("**🧱 พารามิเตอร์ Rigid Pavement (สำหรับรายงานรวม)**")
                        col_p1, col_p2, col_p3 = st.columns([1, 2, 1])
                        with col_p1:
                            _cr_pt_def = st.session_state.get('comb_rigid_pt', pt)
                            try:
                                _cr_pt_idx = [2.0, 2.5, 3.0].index(float(_cr_pt_def))
                            except (ValueError, TypeError):
                                _cr_pt_idx = [2.0, 2.5, 3.0].index(pt) if pt in [2.0, 2.5, 3.0] else 1
                            comb_rigid_pt = st.selectbox("pt (Rigid)", [2.0, 2.5, 3.0],
                                                          index=_cr_pt_idx, key="comb_rigid_pt")
                        with col_p2:
                            _cr_d_def = st.session_state.get('comb_rigid_d_list', [11, 12, 13])
                            if isinstance(_cr_d_def, int):
                                _cr_d_def = [_cr_d_def]
                            _cr_d_def = [d for d in _cr_d_def if d in [10,11,12,13,14,15,16]] or [11,12,13]
                            _D_CM2 = {10:25, 11:28, 12:30, 13:32, 14:35, 15:38, 16:40}
                            _cr_sel = st.multiselect(
                                "D (นิ้ว) — เลือก 3 ค่า",
                                options=[10,11,12,13,14,15,16],
                                default=_cr_d_def,
                                format_func=lambda x: f"D={x}\" ({_D_CM2.get(x,'')} cm)",
                                key="comb_rigid_d_list"
                            )
                            if len(_cr_sel) == 0:
                                _cr_sel = [12]
                            elif len(_cr_sel) > 3:
                                _cr_sel = _cr_sel[:3]
                            comb_rigid_d_list = _cr_sel
                        with col_p3:
                            comb_rigid_lane = st.number_input("Lane Factor", 0.1, 1.0,
                                                               value=lane_factor, step=0.05,
                                                               key="comb_rigid_lane")
                            comb_rigid_dir  = st.number_input("Direction Factor", 0.5, 1.0,
                                                               value=direction_factor, step=0.1,
                                                               key="comb_rigid_dir")
                    
                    st.markdown("---")
                    
                    # ===== Preview บทเกริ่นนำ =====
                    st.markdown("#### 👁️ ตัวอย่างบทเกริ่นนำ (Preview)")
                    
                    num_years = len(traffic_df)
                    
                    flex_tbl_param = flex_table_start
                    flex_tbl_tf = increment_table_number(flex_table_start, 1)
                    flex_tbl_esal = increment_table_number(flex_table_start, 2)
                    
                    rigid_tbl_param = rigid_table_start
                    rigid_tbl_tf = increment_table_number(rigid_table_start, 1)
                    rigid_tbl_esal = increment_table_number(rigid_table_start, 2)
                    
                    col_prev1, col_prev2 = st.columns(2)
                    
                    with col_prev1:
                        st.markdown("**🛤️ Flexible Pavement**")
                        html_flex = generate_intro_preview_html(
                            'flexible', num_years,
                            flex_tbl_param, flex_tbl_tf, flex_tbl_esal,
                            flex_section_number
                        )
                        st.markdown(html_flex, unsafe_allow_html=True)
                    
                    with col_prev2:
                        st.markdown("**🧱 Rigid Pavement**")
                        html_rigid = generate_intro_preview_html(
                            'rigid', num_years,
                            rigid_tbl_param, rigid_tbl_tf, rigid_tbl_esal,
                            rigid_section_number
                        )
                        st.markdown(html_rigid, unsafe_allow_html=True)
                    
                    st.caption("🟣 สีม่วง = ดึงจากข้อมูลอัตโนมัติ | 🟡 สีเหลือง = ผู้ใช้กรอกเอง")
                
                st.divider()
                
                # ============================================================
                # Download buttons
                # ============================================================
                st.write("**📥 ดาวน์โหลดรายงาน:**")
                
                # Collect report settings
                report_settings = {
                    'flex_section_number': st.session_state.get('input_flex_section_number', '4.2.2'),
                    'flex_table_start': st.session_state.get('input_flex_table_start', '4-1'),
                    'rigid_section_number': st.session_state.get('input_rigid_section_number', '4.2.3'),
                    'rigid_table_start': st.session_state.get('input_rigid_table_start', '4-4'),
                }
                
                col_dl1, col_dl2, col_dl3, col_dl4 = st.columns(4)
                
                with col_dl1:
                    excel_report = create_excel_report(
                        results_df, pavement_type, pt, param, lane_factor, direction_factor,
                        total_esal, truck_factors, len(traffic_df)
                    )
                    st.download_button(
                        label="📊 Excel (ปัจจุบัน)",
                        data=excel_report.getvalue(),
                        file_name=f"ESAL_Report_{pavement_type}_{param}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True
                    )
                
                with col_dl2:
                    try:
                        pv_label = "Rigid" if pavement_type == 'rigid' else "Flexible"
                        word_report = create_word_report_multi(
                            traffic_df, pavement_type, pt, param_list,
                            lane_factor, direction_factor,
                            multi_results, report_settings
                        )
                        if word_report:
                            param_str = "_".join(str(p) for p in param_list)
                            st.download_button(
                                label=f"📝 Word ({pv_label})",
                                data=word_report.getvalue(),
                                file_name=f"ESAL_Report_{pavement_type}_{param_str}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                        else:
                            st.warning("⚠️ กรุณาติดตั้ง python-docx")
                    except Exception as e:
                        st.error(f"❌ Word Report: {e}")
                
                with col_dl3:
                    # Combined Word Report
                    try:
                        if pavement_type == 'rigid':
                            # กำลังอยู่ Rigid → flex ใช้ SN list จาก combined settings
                            _cf_pt      = st.session_state.get('comb_flex_pt', 2.5)
                            _cf_sn_list = st.session_state.get('comb_flex_sn_list', [6.5, 7.0, 7.5])
                            if isinstance(_cf_sn_list, (int, float)):
                                _cf_sn_list = [float(_cf_sn_list)]
                            flex_params_comb = {
                                'pt': _cf_pt,
                                'param_list': _cf_sn_list,
                                'lane_factor': st.session_state.get('comb_flex_lane', lane_factor),
                                'direction_factor': st.session_state.get('comb_flex_dir', direction_factor),
                            }
                            rigid_params_comb = {
                                'pt': pt,
                                'param_list': param_list,
                                'lane_factor': lane_factor,
                                'direction_factor': direction_factor,
                            }
                        else:
                            # กำลังอยู่ Flexible → rigid ใช้ D list จาก combined settings
                            _cr_pt     = st.session_state.get('comb_rigid_pt', pt)
                            _cr_d_list = st.session_state.get('comb_rigid_d_list', [11, 12, 13])
                            if isinstance(_cr_d_list, int):
                                _cr_d_list = [_cr_d_list]
                            flex_params_comb = {
                                'pt': pt,
                                'param_list': param_list,
                                'lane_factor': lane_factor,
                                'direction_factor': direction_factor,
                            }
                            rigid_params_comb = {
                                'pt': _cr_pt,
                                'param_list': _cr_d_list,
                                'lane_factor': st.session_state.get('comb_rigid_lane', lane_factor),
                                'direction_factor': st.session_state.get('comb_rigid_dir', direction_factor),
                            }
                        
                        word_combined = create_word_report_combined(
                            traffic_df, flex_params_comb, rigid_params_comb, report_settings
                        )
                        if word_combined:
                            st.download_button(
                                label="📝 Word (รวม Flex+Rigid)",
                                data=word_combined.getvalue(),
                                file_name="ESAL_Report_Combined.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                        else:
                            st.warning("⚠️ กรุณาติดตั้ง python-docx")
                    except Exception as e:
                        st.error(f"❌ เกิดข้อผิดพลาด: {e}")
                
                with col_dl4:
                    # รวบรวม comb_rigid_params เฉพาะกรณี flexible
                    _crp = None
                    if pavement_type == 'flexible':
                        _crp = {
                            'pt':               st.session_state.get('comb_rigid_pt',  2.5),
                            'param':            st.session_state.get('comb_rigid_d',   13),
                            'lane_factor':      st.session_state.get('comb_rigid_lane', lane_factor),
                            'direction_factor': st.session_state.get('comb_rigid_dir',  direction_factor),
                        }
                    project_json = save_project(
                        pavement_type, pt, param, lane_factor, direction_factor,
                        truck_factors, traffic_df, report_settings,
                        comb_rigid_params=_crp
                    )
                    st.download_button(
                        label="💾 บันทึก Project",
                        data=project_json,
                        file_name=f"ESAL_Project_{pavement_type}_{param}.json",
                        mime="application/json",
                        use_container_width=True
                    )
            else:
                st.warning("⚠️ กรุณาอัพโหลดข้อมูลหรือใช้ข้อมูลตัวอย่าง")
    
    with tab2:
        st.subheader("🚛 ข้อมูลรถบรรทุก 6 ประเภทตามกรมทางหลวง")
        
        # แปลง VEHICLE_AXLES → ชื่อเพลาสำหรับแสดงใน Tab 2
        _AXLE_TYPE_LABEL = {1: 'Single', 2: 'Tandem'}
        _AXLE_NAMES = ['เพลาหน้า', 'เพลาหลัง', 'เพลาพ่วงหน้า', 'เพลาพ่วงหลัง']
        truck_details = []
        for code, truck in TRUCKS.items():
            axle_info = []
            for i, (load_ton, l2, _cnt) in enumerate(VEHICLE_AXLES[code]):
                name = _AXLE_NAMES[i] if i < len(_AXLE_NAMES) else f'เพลาที่ {i+1}'
                axle_info.append(f"{name}: {load_ton} ตัน ({_AXLE_TYPE_LABEL[l2]})")
            truck_details.append({'รหัส': code, 'ประเภท': truck['desc'], 'ข้อมูลเพลา': ' | '.join(axle_info)})
        
        st.dataframe(pd.DataFrame(truck_details), use_container_width=True, hide_index=True)
        
        st.divider()
        st.subheader("📊 ตาราง Truck Factor (ค่า Default ตาม AASHTO 1993)")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.write("**🧱 Rigid Pavement (pt = 2.0)**")
            st.dataframe(get_all_truck_factors_table('rigid', 2.0), use_container_width=True, hide_index=True)
            
            st.write("**🧱 Rigid Pavement (pt = 2.5)**")
            st.dataframe(get_all_truck_factors_table('rigid', 2.5), use_container_width=True, hide_index=True)
            
            st.write("**🧱 Rigid Pavement (pt = 3.0)**")
            st.dataframe(get_all_truck_factors_table('rigid', 3.0), use_container_width=True, hide_index=True)
        
        with col2:
            st.write("**🛤️ Flexible Pavement (pt = 2.0)**")
            st.dataframe(get_all_truck_factors_table('flexible', 2.0), use_container_width=True, hide_index=True)
            
            st.write("**🛤️ Flexible Pavement (pt = 2.5)**")
            st.dataframe(get_all_truck_factors_table('flexible', 2.5), use_container_width=True, hide_index=True)
            
            st.write("**🛤️ Flexible Pavement (pt = 3.0)**")
            st.dataframe(get_all_truck_factors_table('flexible', 3.0), use_container_width=True, hide_index=True)
    
    with tab3:
        st.subheader("📘 คู่มือการใช้งาน")
        
        st.markdown("""
        ### 1️⃣ เตรียมไฟล์ Excel
        
        | คอลัมน์ | คำอธิบาย |
        |---------|----------|
        | `Year` | ปีที่ (1, 2, 3, ... n) |
        | `MB` | Medium Bus (คัน/วัน) |
        | `HB` | Heavy Bus (คัน/วัน) |
        | `MT` | Medium Truck (คัน/วัน) |
        | `HT` | Heavy Truck (คัน/วัน) |
        | `STR` | Semi-Trailer (คัน/วัน) |
        | `TR` | Full Trailer (คัน/วัน) |
        
        ### 2️⃣ ตั้งค่าพารามิเตอร์
        
        - **Rigid:** D = 10-16 นิ้ว
        - **Flexible:** SN = 4-9
        - **pt:** 2.0, 2.5 หรือ 3.0
        
        ### 3️⃣ ฟีเจอร์ในเวอร์ชัน 3.0
        
        - **ACC. ESAL:** แสดงค่า ESAL สะสม
        - **Export Excel:** รายงานในรูปแบบมาตรฐาน
        - **Export Word (แยก/รวม):** รายงานสำหรับเอกสาร
        - **ระบบหมายเลขตาราง:** Auto-increment กรอกเลขเริ่มต้น ระบบเพิ่มให้อัตโนมัติ
        - **บทเกริ่นนำ:** แยก Flexible / Rigid พร้อม Preview
        - **Save/Load Project:** บันทึกค่าทั้งหมดรวม report settings
        
        ### 4️⃣ สูตรคำนวณ ESAL
        """)
        
        st.latex(r'ESAL = \sum_{i=1}^{n} \sum_{j=1}^{6} (ADT_{ij} \times TF_j \times LF \times DF \times 365)')
        
        st.markdown("""
        ### 5️⃣ การ Export รายงาน Word
        
        | ปุ่ม | คำอธิบาย |
        |------|----------|
        | 📝 Word (Flexible/Rigid) | รายงานเฉพาะประเภทที่กำลังคำนวณ |
        | 📝 Word (รวม Flex+Rigid) | รายงานรวมทั้ง 2 ประเภทในไฟล์เดียว |
        
        กำหนดเลขหัวข้อและเลขตารางได้ที่ **📝 ตั้งค่ารายงาน Word** (Expander)
        
        ### 📚 อ้างอิง
        - AASHTO Guide for Design of Pavement Structures (1993)
        - กรมทางหลวง กระทรวงคมนาคม
        """)
    
    st.divider()
    st.markdown("""
    <div style="text-align: center; color: #888;">
        พัฒนาเพื่อการเรียนการสอนโดย รศ.ดร.อิทธิพล มีผล ภาควิชาครุศาสตร์โยธา มจพ. | ESAL Calculator v3.0
    </div>
    """, unsafe_allow_html=True)


if __name__ == "__main__":
    main()
