"""
cross_section_tab.py  v3
แก้ไข:
  1. ลูกศรม่วง → ชี้กลาง layer (y_mid)
  2. เส้นขอบชั้น → ดำบาง แทนเส้นเขียว
  3. Thickness label: ใน layer ถ้าสูงพอ / leader ออกข้างถ้าบาง
  4. Caption ภาษาไทย (DejaVu Sans fallback)
  5. Download PDF
รศ.ดร.อิทธิพล มีผล | KMUTNB
"""

import io
import ast
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patheffects as pe
from matplotlib.patches import Polygon
import numpy as np
import streamlit as st

try:
    from docx import Document as WordDoc
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

# ══════════════════════════════════════════════════════════════
# MATERIAL COLOR PRESETS
# ══════════════════════════════════════════════════════════════
MATERIAL_COLORS = {
    'concrete'   : '#A0A0A0',
    'jpcp'       : '#A0A0A0',
    'jrcp'       : '#A0A0A0',
    'crcp'       : '#A0A0A0',
    'pcc'        : '#A0A0A0',
    'slab'       : '#A0A0A0',
    'asphalt'    : '#2C2C2C',
    'wearing'    : '#2C2C2C',
    'binder'     : '#2C2C2C',
    'hma'        : '#2C2C2C',
    'interlayer' : '#2C2C2C',
    'modified'   : '#6B9E9E',
    'cement'     : '#6B9E9E',
    'cmcr'       : '#6B9E9E',
    'stabilized' : '#6B9E9E',
    'crushed'    : '#9DB8A8',
    'base'       : '#9DB8A8',
    'granular'   : '#9DB8A8',
    'subbase'    : '#C8D8B8',
    'aggregate'  : '#C8D8B8',
    'soil'       : '#C8D8B8',
    'sand'       : '#D4C090',
    'default'    : '#CCCCCC',
}

SKIP_KEYWORDS = [
    'tack coat', 'prime coat', 'wire mesh',
    'geotextile', 'joint', 'steel', 'dowel',
    'tie bar', 'transverse', 'longitudinal steel',
]

def get_color(name: str, override: str = None) -> str:
    if override:
        return override
    n = name.lower()
    for k, v in MATERIAL_COLORS.items():
        if k in n:
            return v
    return MATERIAL_COLORS['default']

def has_interlayer(layers_m: list) -> bool:
    return any('interlayer' in lyr['name'].lower() for lyr in layers_m)

def should_skip(name: str, unit: str) -> bool:
    if unit.lower() not in ('cm', 'm'):
        return True
    return any(kw in name.lower() for kw in SKIP_KEYWORDS)

def filter_layers(raw_layers: list) -> list:
    result = []
    for lyr in raw_layers:
        name  = lyr.get('name', '')
        thick = float(lyr.get('thickness', 0) or 0)
        unit  = lyr.get('unit', 'cm')
        if should_skip(name, unit):
            continue
        if thick <= 0:
            continue
        thick_m = thick / 100.0 if unit.lower() == 'cm' else thick
        if thick_m > 0:
            result.append({'name': name, 'thickness': thick_m})
    return result

# ══════════════════════════════════════════════════════════════
# DRAWING FUNCTION
# ══════════════════════════════════════════════════════════════
def draw_cross_section(
    layers_m: list,
    verge_w: float = 0.5,
    shoulder_w: float = 1.5,
    carriageway_w: float = 7.5,
    embankment_cbr: float = 3.0,
    fig_no: str = "3.9-1",
    color_overrides: dict = None,
    show_prime_coat: bool = True,
) -> plt.Figure:

    color_overrides = color_overrides or {}

    # ── scale ──────────────────────────────────────────────
    S_h = 20.0
    S_v = 80.0

    road_w    = (verge_w + shoulder_w + carriageway_w) * S_h
    x_left    = 0.0
    x_right   = road_w
    x_verge_r = x_left    + verge_w    * S_h
    x_shld_r  = x_verge_r + shoulder_w * S_h

    y_surface = 0.0
    y_layers  = [y_surface]
    for lyr in layers_m:
        y_layers.append(y_layers[-1] - lyr['thickness'] * S_v)
    y_pave_bot  = y_layers[-1]
    y_scar_bot  = y_pave_bot - 0.10 * S_v
    y_emb_bot   = y_scar_bot - 5.0
    total_v     = abs(y_scar_bot - y_surface)
    slope_run   = 1.5 * total_v
    x_slope_far = x_left - slope_run

    def x_on_slope(y):
        t = (y - y_emb_bot) / (y_surface - y_emb_bot)
        return x_slope_far + t * (x_left - x_slope_far)

    # ── figsize ─────────────────────────────────────────────
    leader_space = road_w * 0.55
    y_bot_lim    = y_emb_bot - leader_space
    y_top_lim    = y_surface + road_w * 0.25
    x_range      = (x_right + 2.0) - (x_slope_far - 4.0)
    y_range      = y_top_lim - y_bot_lim
    BASE_W       = 13.0
    fig_h        = BASE_W * (y_range / x_range)

    fig, ax = plt.subplots(figsize=(BASE_W, fig_h))
    ax.set_aspect('equal')
    ax.axis('off')
    fig.patch.set_facecolor('white')

    BLACK  = '#000000'
    GREY   = '#555555'
    YELLOW = '#F9A825'
    PURPLE = '#6A0DAD'
    FONT   = 'DejaVu Sans'
    LINE_W = 4.5
    lw_dim = 1.0
    lw_ldr = 1.5
    lw_edge= 0.9
    TICK_H = road_w * 0.020

    # auto font size dimension
    min_span = min(verge_w, shoulder_w) * S_h
    FS_DIM   = max(7.0, min(12.0, min_span * 0.55))
    FS_LDR   = 10
    FS_THICK = 9.5   # thickness label

    # x position ของ thickness label (ขวาของ pavement)
    x_thick_line = x_right + road_w * 0.02   # จุดต่อเส้น leader
    x_thick_text = x_right + road_w * 0.025  # จุดเริ่ม text

    # ── Embankment fill ──────────────────────────────────────
    emb_pts = np.array([
        [x_on_slope(y_surface), y_surface],
        [x_right,               y_surface],
        [x_right,               y_emb_bot],
        [x_slope_far,           y_emb_bot],
    ])
    ax.add_patch(Polygon(emb_pts, closed=True,
                         facecolor='#E8DCC8', edgecolor='none',
                         alpha=0.6, zorder=2))

    # ── Pavement layers + thickness labels ───────────────────
    for i, lyr in enumerate(layers_m):
        y_top = y_layers[i]
        y_bot = y_layers[i + 1]
        y_mid = (y_top + y_bot) / 2

        pts = np.array([
            [x_on_slope(y_top), y_top],
            [x_right,           y_top],
            [x_right,           y_bot],
            [x_on_slope(y_bot), y_bot],
        ])
        color = color_overrides.get(lyr['name'], get_color(lyr['name']))
        ax.add_patch(Polygon(pts, closed=True,
                             facecolor=color, edgecolor=BLACK,
                             linewidth=lw_edge, zorder=5))

        # ── Thickness label — leader ดำออกขวาทุกชั้น ────────
        thick_cm   = lyr['thickness'] * 100
        thick_text = f"{thick_cm:.0f} cm"

        # เส้น leader สั้น: x_right → x_thick_line
        ax.plot([x_right, x_thick_line], [y_mid, y_mid],
                color=BLACK, lw=0.9, zorder=22)
        # text
        ax.text(x_thick_text, y_mid, thick_text,
                ha='left', va='center',
                fontsize=FS_THICK, fontname=FONT,
                fontweight='bold', color=BLACK,
                zorder=22)

    # ── Existing ground (wavy yellow dashed) ─────────────────
    np.random.seed(42)
    x_wave    = np.linspace(x_slope_far - 25.0, x_right, 500)
    amplitude = 1.8
    period    = road_w / 1.8
    y_sine    = amplitude * np.sin(2*np.pi*(x_wave - x_wave[0])/period)
    noise_raw = np.random.randn(len(x_wave)) * 0.5
    kernel    = np.ones(20) / 20
    noise     = np.convolve(noise_raw, kernel, mode='same')
    y_wave    = y_emb_bot + y_sine + noise
    ax.plot(x_wave, y_wave,
            color=YELLOW, lw=3.0, linestyle='--', dashes=(10,5),
            zorder=10, solid_capstyle='round')

    # ── Dimension lines บน ───────────────────────────────────
    y_brk1 = y_surface + road_w * 0.10
    y_brk2 = y_surface + road_w * 0.20

    def dim_line(x_l, x_r, y, label, fs=FS_DIM):
        span   = x_r - x_l
        fs_use = max(6.5, min(fs, span * 0.55))
        for xv in [x_l, x_r]:
            ax.plot([xv, xv], [y - TICK_H*0.5, y + TICK_H*0.7],
                    color=BLACK, lw=lw_dim, zorder=12)
        ax.annotate('', xy=(x_r, y), xytext=(x_l, y),
                    arrowprops=dict(arrowstyle='<->', color=BLACK,
                                    lw=lw_dim, mutation_scale=12), zorder=12)
        ax.text((x_l+x_r)/2, y + TICK_H*0.85, label,
                ha='center', va='bottom', fontsize=fs_use,
                fontname=FONT, fontweight='bold', color=BLACK)

    for xv in [x_left, x_verge_r, x_shld_r, x_right]:
        ax.plot([xv, xv], [y_surface, y_brk2 + TICK_H*2],
                color=GREY, lw=0.6, linestyle=':', zorder=3)

    dim_line(x_left,    x_verge_r, y_brk1, 'VERGE')
    dim_line(x_verge_r, x_shld_r,  y_brk1, 'SHOULDER')
    dim_line(x_shld_r,  x_right,   y_brk1, 'CARRIAGEWAY')
    if show_prime_coat:
        dim_line(x_verge_r, x_right, y_brk2, 'PRIME COAT')

    # ── 3 colored lines ──────────────────────────────────────
    # Blue: slope
    ax.plot([x_slope_far, x_left], [y_emb_bot, y_surface],
            color='#1565C0', lw=LINE_W, zorder=20, solid_capstyle='round')
    # Red: top surface
    ax.plot([x_on_slope(y_surface), x_right], [y_surface, y_surface],
            color='#D32F2F', lw=LINE_W, zorder=21, solid_capstyle='round')
    # ดำบาง: layer boundaries (แทนเขียว)
    for i in range(1, len(y_layers)):
        y_line  = y_layers[i]
        x_start = x_on_slope(y_line)
        ax.plot([x_start, x_right], [y_line, y_line],
                color=BLACK, lw=1.2, zorder=19, solid_capstyle='round')

    # ── L-shape leaders + text ───────────────────────────────
    emb_label = f"EARTH EMBANKMENT, CBR {embankment_cbr:.0f}% (MIN)"
    all_items = []
    for i, lyr in enumerate(layers_m):
        # ลูกศรชี้กลาง layer (y_mid)
        y_mid = (y_layers[i] + y_layers[i+1]) / 2
        all_items.append((lyr['name'].upper(), y_mid))
    y_emb_mid = (y_scar_bot + y_emb_bot) / 2
    all_items.append((emb_label, y_emb_mid))

    n_all     = len(all_items)
    x_center  = (x_left + x_right) / 2
    step_x    = road_w / (n_all + 1)
    x_leaders = [x_center - i * step_x for i in range(n_all)]

    y_leader_base = y_emb_bot - road_w * 0.08
    y_step        = road_w * 0.085
    # ขยับ x_text_end ออกเพื่อไม่ทับ thickness labels
    x_text_end = x_right + road_w * 0.18

    for i, (label, y_target) in enumerate(all_items):
        xd      = x_leaders[i]
        y_horiz = y_leader_base - i * y_step

        # เส้นตั้ง zorder=25
        ax.plot([xd, xd], [y_target, y_horiz],
                color=PURPLE, lw=lw_ldr, zorder=25)

        # arrowhead ชี้กลาง layer zorder=30
        ax.annotate('', xy=(xd, y_target),
                    xytext=(xd, y_target - road_w*0.04),
                    arrowprops=dict(arrowstyle='->', color=PURPLE,
                                    lw=lw_ldr, mutation_scale=12),
                    zorder=30)

        # เส้นนอน + dot
        ax.plot([xd, x_text_end], [y_horiz, y_horiz],
                color=PURPLE, lw=lw_ldr, zorder=25)
        ax.plot(xd, y_horiz, 'o', color=PURPLE, ms=3.5, zorder=26)

        # text
        ax.text(x_text_end + road_w*0.01, y_horiz, label,
                ha='left', va='center', fontsize=FS_LDR,
                fontname=FONT, fontweight='bold', color=BLACK)

    # ── Caption (ใช้ ASCII เพื่อ compatibility) ──────────────
    ax.text((x_left + x_right)/2,
            y_bot_lim + road_w * 0.04,
            f"Figure {fig_no}  Pavement Structure Cross-Section",
            ha='center', va='bottom', fontsize=FS_LDR,
            fontname=FONT, fontstyle='italic', color=BLACK)

    ax.set_xlim(x_slope_far - 4.0, x_right + road_w * 0.65)
    ax.set_ylim(y_bot_lim, y_top_lim)
    plt.tight_layout(pad=0.3)
    return fig

# ══════════════════════════════════════════════════════════════
# WORD REPORT
# ══════════════════════════════════════════════════════════════
def generate_word_report(
    fig, ptype, layers_m,
    intro_para1, intro_para2,
    fig_no, section_no, project_name,
) -> bytes:
    if not DOCX_OK:
        raise ImportError("python-docx not installed")

    doc = WordDoc()
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.0)

    def set_thai_font(run, size_pt=15, bold=False):
        run.font.name = 'TH SarabunPSK'
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        for attr in ('w:ascii','w:hAnsi','w:cs','w:eastAsia'):
            rFonts.set(qn(attr), 'TH SarabunPSK')
        rPr.insert(0, rFonts)

    # Heading
    h = doc.add_heading('', level=3)
    run = h.add_run(f"{section_no}  สรุปผลการวิเคราะห์")
    set_thai_font(run, 16, True)

    # 2 paragraphs
    for txt in [intro_para1, intro_para2]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(6)
        set_thai_font(p.add_run(txt), 15)

    # Figure
    buf_img = io.BytesIO()
    fig.savefig(buf_img, format='png', dpi=180,
                bbox_inches='tight', facecolor='white')
    buf_img.seek(0)
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(buf_img, width=Cm(15.0))

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_thai_font(p_cap.add_run(
        f"รูปที่ {fig_no}  รูปแบบโครงสร้างชั้นทาง ({ptype})"), 14)

    # Table
    doc.add_paragraph()
    set_thai_font(doc.add_paragraph().add_run(
        f"โครงสร้างชั้นทาง {ptype}"), 15, True)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    for cell, txt in zip(tbl.rows[0].cells,
                         ['ชั้นวัสดุ', 'ความหนา (ซม.)']):
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_thai_font(cell.paragraphs[0].add_run(txt), 14, True)

    for lyr in layers_m:
        row = tbl.add_row().cells
        set_thai_font(row[0].paragraphs[0].add_run(lyr['name']), 14)
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_thai_font(row[1].paragraphs[0].add_run(
            f"{lyr['thickness']*100:.0f}"), 14)

    buf_doc = io.BytesIO()
    doc.save(buf_doc)
    buf_doc.seek(0)
    return buf_doc.getvalue()

# ══════════════════════════════════════════════════════════════
# STREAMLIT UI
# ══════════════════════════════════════════════════════════════
def render_cross_section_tab(ss: dict = None, project_name: str = ""):
    if ss is None:
        ss = st.session_state

    st.markdown("### 🖼️ Cross-Section Drawing")

    cs_ar = ss.get('cs_all_results', {})

    # ── เลือก ptype ─────────────────────────────────────────
    best_ptype = None
    if ss.get('_lc_sum') is not None:
        sdf = ss.get('_lc_sum')
        if sdf is not None and 'NPV (ล้านบาท/กม.)' in sdf.columns:
            idx = sdf['NPV (ล้านบาท/กม.)'].idxmin()
            if idx is not None:
                # ประเภทผิวทาง = 'JPCP','JRCP','CRCP','AC'
                raw = sdf.loc[idx, 'ประเภทผิวทาง']
                # map ชื่อเต็ม → key ที่ใช้ใน cs_all_results
                for key in ['JPCP','JRCP','CRCP','AC']:
                    if key in str(raw).upper():
                        best_ptype = key
                        break

    mode = st.radio("เลือก pavement type",
                    ["🔄 ใช้ผล LCCA (NPV ต่ำสุด)", "✏️ เลือกเอง"],
                    horizontal=True, key="cs_mode_radio")

    available_types = [pt for pt in ['AC','JPCP','JRCP','CRCP']
                       if cs_ar.get(pt, {}).get('layers')]

    if mode == "🔄 ใช้ผล LCCA (NPV ต่ำสุด)":
        ptype = best_ptype or (available_types[0] if available_types else 'JPCP')
        st.success(f"✅ ประเภทที่เลือก: **{ptype}**")
    else:
        opts  = available_types or ['JPCP','JRCP','CRCP','AC']
        ptype = st.selectbox("เลือกประเภทผิวทาง", opts, key="cs_ptype_sel")

    # ── layers ──────────────────────────────────────────────
    raw_layers = cs_ar.get(ptype, {}).get('layers', [])
    layers_m   = filter_layers(raw_layers)

    if not layers_m:
        defaults = {
            'JPCP': [
                {'name':'Concrete Slab (JPCP)',              'thickness':0.28},
                {'name':'AC Interlayer',                     'thickness':0.05},
                {'name':'Cement Modified Crushed Rock Base', 'thickness':0.20},
                {'name':'Soil Aggregate Subbase',            'thickness':0.15},
            ],
            'AC': [
                {'name':'AC Wearing Course',                 'thickness':0.07},
                {'name':'AC Binder Course',                  'thickness':0.07},
                {'name':'AC Base Course',                    'thickness':0.10},
                {'name':'Crushed Rock Base Course',          'thickness':0.20},
                {'name':'Soil Aggregate Subbase',            'thickness':0.30},
            ],
        }
        layers_m = defaults.get(ptype, defaults['JPCP'])
        st.info("ℹ️ ใช้ค่า default — กรุณารันผ่าน PaveCost_LCCA เพื่อใช้ข้อมูลจริง")

    show_prime = has_interlayer(layers_m)

    # ── แสดง layers ─────────────────────────────────────────
    with st.expander("📋 Layers ที่จะวาด", expanded=False):
        for lyr in layers_m:
            st.write(f"• **{lyr['name']}** — {lyr['thickness']*100:.1f} cm")

    # ── Color pickers ────────────────────────────────────────
    st.markdown("**🎨 สีแต่ละชั้นทาง**")
    color_overrides = {}
    cols = st.columns(min(len(layers_m), 4))
    for i, lyr in enumerate(layers_m):
        key = f"cs_color_{i}_{lyr['name'][:10].replace(' ','_')}"
        with cols[i % len(cols)]:
            picked = st.color_picker(
                lyr['name'][:22],
                value=ss.get(key, get_color(lyr['name'])),
                key=key)
            color_overrides[lyr['name']] = picked

    st.divider()

    # ── Parameters ──────────────────────────────────────────
    col1, col2, col3 = st.columns(3)
    with col1:
        verge_w = st.number_input(
            "Verge (ม.)", 0.0, 3.0,
            float(ss.get('cs_verge_w', 0.5)), 0.25, key="cs_verge_w")
        shoulder_w = st.number_input(
            "Shoulder (ม.)", 0.0, 5.0,
            float(ss.get('cs_shoulder_w', 1.5)), 0.25, key="cs_shoulder_w")
    with col2:
        carriageway_w = st.number_input(
            "Carriageway (ม.)", 3.0, 20.0,
            float(ss.get('cs_carry_w', 7.5)), 0.5, key="cs_carry_w")
        embankment_cbr = st.number_input(
            "CBR Embankment (%)", 1.0, 20.0,
            float(ss.get('cs_emb_cbr', 3.0)), 0.5, key="cs_emb_cbr")
    with col3:
        fig_no     = st.text_input("เลขรูปที่",      ss.get('cs_fig_no','3.9-1'),  key="cs_fig_no")
        section_no = st.text_input("หมายเลขหัวข้อ", ss.get('cs_sec_no','3.9.8'),  key="cs_sec_no")

    st.divider()

    # ── Word inputs ──────────────────────────────────────────
    st.markdown("**📝 บทเกริ่นนำ Word Report**")
    ptype_name_map = {
        'JPCP': 'JPCP (คอนกรีตซีเมนต์แบบไม่มีเหล็กเสริม)',
        'JRCP': 'JRCP (คอนกรีตซีเมนต์แบบมีเหล็กตาข่าย)',
        'CRCP': 'CRCP (คอนกรีตซีเมนต์แบบเสริมเหล็กต่อเนื่อง)',
        'AC'  : 'แอสฟัลต์คอนกรีต (AC)',
    }
    ptype_full = ptype_name_map.get(ptype, ptype)
    default_p1 = (
        f"จากการวิเคราะห์ Life Cycle Cost Analysis โครงสร้างผิวทางทั้ง 4 ประเภท "
        f"พบว่า {ptype_full} เป็นทางเลือกที่มีความคุ้มค่าทางเศรษฐศาสตร์สูงสุด "
        f"จึงเป็นผิวทางที่มีความเหมาะสมมากที่สุดสำหรับใช้เป็นผิวทางถนนโครงการ"
    )
    default_p2 = (
        f"ในส่วนของการออกแบบผิวทางบนผิวทางเดิม/คันทางเดิม ปรึกษาได้พิจารณาผล"
        f"การทดสอบความแข็งแรงโครงสร้างชั้นทางเดิม และผลการทดสอบ "
        f"Falling Weight Deflectometer (FWD) หากพบว่ามีความแข็งแรงต่ำ "
        f"ควรดำเนินการ Reconstruction โครงสร้างชั้นทางเดิมก่อนที่จะก่อสร้าง"
        f"ผิวทางใหม่ โดยสรุปรูปแบบดังรูปที่ {fig_no}"
    )
    intro_p1 = st.text_area("ย่อหน้าที่ 1",
                             value=ss.get('cs_intro_p1', default_p1),
                             height=100, key="cs_intro_p1")
    intro_p2 = st.text_area("ย่อหน้าที่ 2",
                             value=ss.get('cs_intro_p2', default_p2),
                             height=120, key="cs_intro_p2")

    st.divider()

    # ── Generate ─────────────────────────────────────────────
    if st.button("🖼️ Generate Cross-Section", type="primary",
                 use_container_width=True, key="cs_gen_btn"):
        with st.spinner("กำลังวาดภาพ..."):
            fig = draw_cross_section(
                layers_m       = layers_m,
                verge_w        = verge_w,
                shoulder_w     = shoulder_w,
                carriageway_w  = carriageway_w,
                embankment_cbr = embankment_cbr,
                fig_no         = fig_no,
                color_overrides= color_overrides,
                show_prime_coat= show_prime,
            )
            # แปลง Figure → bytes ทันที (ปลอดภัยกว่าเก็บ fig object)
            buf_png = io.BytesIO()
            fig.savefig(buf_png, format='png', dpi=180,
                        bbox_inches='tight', facecolor='white')
            buf_png.seek(0)
            buf_pdf = io.BytesIO()
            fig.savefig(buf_pdf, format='pdf',
                        bbox_inches='tight', facecolor='white')
            buf_pdf.seek(0)

            ss['cs_last_img_bytes'] = buf_png.getvalue()  # เก็บ bytes ปลอดภัย 100%
            ss['cs_last_pdf_bytes'] = buf_pdf.getvalue()
            ss['cs_last_fig']       = fig                 # เก็บไว้แสดงบน UI เท่านั้น
            ss['cs_last_layers']    = layers_m
            ss['cs_last_ptype']     = ptype
            ss['cs_last_fig_no']    = fig_no
            plt.close(fig)                                 # free memory
        st.success("✅ วาดภาพสำเร็จ — รูปพร้อมแทรกใน Word Report แล้ว")

    # ── แสดงภาพ + downloads ──────────────────────────────────
    if 'cs_last_img_bytes' in ss:
        st.image(ss['cs_last_img_bytes'], use_column_width=True)

        dl1, dl2, dl3 = st.columns(3)

        with dl1:
            st.download_button("⬇️ PNG",
                                data=ss['cs_last_img_bytes'],
                                file_name=f"CrossSection_{ss.get('cs_last_ptype','')}.png",
                                mime="image/png", key="cs_dl_png")

        with dl2:
            st.download_button("⬇️ PDF",
                                data=ss['cs_last_pdf_bytes'],
                                file_name=f"CrossSection_{ss.get('cs_last_ptype','')}.pdf",
                                mime="application/pdf", key="cs_dl_pdf")

        # Word
        if DOCX_OK:
            with dl3:
                if st.button("📄 Word Report", key="cs_word_btn",
                             use_container_width=True):
                    with st.spinner("สร้าง Word..."):
                        try:
                            word_bytes = generate_word_report(
                                fig=ss['cs_last_fig'],
                                ptype=ss['cs_last_ptype'],
                                layers_m=ss['cs_last_layers'],
                                intro_para1=intro_p1,
                                intro_para2=intro_p2,
                                fig_no=fig_no,
                                section_no=section_no,
                                project_name=project_name,
                            )
                            st.download_button(
                                "⬇️ Download Word",
                                data=word_bytes,
                                file_name=f"CrossSection_{ptype}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key="cs_dl_word")
                            st.success("✅ Word พร้อม")
                        except Exception as e:
                            st.error(f"Error: {e}")
        else:
            with dl3:
                st.warning("pip install python-docx")


# ══════════════════════════════════════════════════════════════
# STANDALONE
# ══════════════════════════════════════════════════════════════
if __name__ == "__main__":
    _src = open(__file__, encoding="utf-8").read()
    ast.parse(_src)
    print("Syntax OK")

    st.set_page_config(page_title="Cross-Section", page_icon="🛣️", layout="wide")
    st.markdown("## 🛣️ Cross-Section Drawing")
    render_cross_section_tab(project_name="Test Project")
