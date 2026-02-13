import streamlit as st
import re
import io
import os
import copy
import zipfile
from docx import Document
from docx.shared import Inches, Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from PIL import Image
import tempfile
import shutil

# ─── Page Config ───
st.set_page_config(
    page_title="รวมรายงานโครงสร้างชั้นทาง",
    page_icon="📄",
    layout="wide"
)

# ─── Custom CSS ───
st.markdown("""
<style>
    .main-title { 
        text-align: center; font-size: 2rem; font-weight: bold; 
        color: #1E3A5F; margin-bottom: 0.5rem; 
    }
    .sub-title { 
        text-align: center; font-size: 1rem; color: #666; margin-bottom: 2rem; 
    }
    .section-card {
        background: #f8f9fa; border-radius: 8px; padding: 1rem;
        border-left: 4px solid #1E3A5F; margin-bottom: 0.5rem;
    }
    .stButton > button {
        background-color: #1E3A5F; color: white; border-radius: 8px;
        padding: 0.5rem 2rem; font-size: 1rem;
    }
</style>
""", unsafe_allow_html=True)

st.markdown('<div class="main-title">📄 ระบบรวมรายงาน Word</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-title">รวมไฟล์ Word แยกหัวข้อ พร้อม Upload รูปภาพ และแก้ไขหมายเลขบท/หัวข้อ</div>', unsafe_allow_html=True)

# ─── Initialize Session State ───
if "sections" not in st.session_state:
    st.session_state.sections = []
if "images" not in st.session_state:
    st.session_state.images = {}
if "chapter_num" not in st.session_state:
    st.session_state.chapter_num = 4
if "start_section_num" not in st.session_state:
    st.session_state.start_section_num = 1


def extract_docx_content(uploaded_file):
    """Extract content info from a docx file."""
    doc = Document(io.BytesIO(uploaded_file.read()))
    uploaded_file.seek(0)
    
    headings = []
    para_count = 0
    table_count = len(doc.tables)
    image_count = 0
    
    for para in doc.paragraphs:
        para_count += 1
        if para.style.name.startswith("Heading"):
            headings.append({
                "text": para.text,
                "level": para.style.name
            })
        # Count images
        for run in para.runs:
            if run._element.findall(qn('w:drawing')):
                image_count += 1
    
    return {
        "headings": headings,
        "para_count": para_count,
        "table_count": table_count,
        "image_count": image_count
    }


def renumber_headings_in_doc(doc, chapter_num, section_start, section_mapping):
    """
    Renumber headings in a document.
    section_mapping: dict mapping old section numbers to new ones
    """
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            original_text = para.text.strip()
            # Match patterns like "4.1", "4.1.1", "4.2.3", or just numbers at start
            # Pattern: digits.digits or digits.digits.digits etc.
            match = re.match(r'^(\d+(?:\.\d+)*)\s*(.*)', original_text)
            if match:
                old_num = match.group(1)
                rest_text = match.group(2)
                
                parts = old_num.split('.')
                if len(parts) >= 1:
                    parts[0] = str(chapter_num)
                
                if len(parts) >= 2 and old_num in section_mapping:
                    new_num = section_mapping[old_num]
                else:
                    # Apply offset to second level
                    if len(parts) >= 2:
                        try:
                            old_second = int(parts[1])
                            parts[1] = str(old_second + section_start - 1)
                        except ValueError:
                            pass
                    new_num = '.'.join(parts)
                
                new_text = f"{new_num} {rest_text}" if rest_text else new_num
                
                # Clear and rewrite
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = new_text
                else:
                    para.add_run(new_text)


def renumber_figures_tables(doc, chapter_num):
    """Renumber figure and table references."""
    fig_counter = 0
    tbl_counter = 0
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Match "รูปที่ X-Y" or "รูปที่ X.Y"
        fig_match = re.match(r'^(รูปที่\s*)\d+[-\.]\d+(.*)', text)
        if fig_match:
            fig_counter += 1
            new_text = f"{fig_match.group(1)}{chapter_num}-{fig_counter}{fig_match.group(2)}"
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = new_text
            else:
                para.add_run(new_text)
        
        # Match "ตารางที่ X-Y" or "ตารางที่ X.Y"
        tbl_match = re.match(r'^(ตารางที่\s*)\d+[-\.]\d+(.*)', text)
        if tbl_match:
            tbl_counter += 1
            new_text = f"{tbl_match.group(1)}{chapter_num}-{tbl_counter}{tbl_match.group(2)}"
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = new_text
            else:
                para.add_run(new_text)
    
    return fig_counter, tbl_counter


def insert_image_after_paragraph(doc, para_index, image_bytes, caption_text, width_cm=15):
    """Insert an image after a specific paragraph."""
    img = Image.open(io.BytesIO(image_bytes))
    img_width, img_height = img.size
    aspect = img_height / img_width
    
    width = Cm(width_cm)
    height = Cm(width_cm * aspect)
    
    # Create image paragraph
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(io.BytesIO(image_bytes), width=width)
    
    # Create caption paragraph  
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_para.add_run(caption_text)
    cap_run.bold = True
    cap_run.font.size = Pt(11)
    
    # Move paragraphs to correct position
    target_para = doc.paragraphs[para_index]
    target_element = target_para._element
    
    # Insert caption after target, then image after target (so image comes first)
    target_element.addnext(cap_para._element)
    target_element.addnext(img_para._element)
    
    return img_para, cap_para


def merge_documents(sections_data, chapter_num, images_dict, renumber_options):
    """Merge multiple Word documents into one."""
    if not sections_data:
        return None
    
    # Load first document as base
    first_data = sections_data[0]
    merged_doc = Document(io.BytesIO(first_data["file_bytes"]))
    
    # Process renumbering on first doc if needed
    if renumber_options.get("renumber_headings", False):
        section_mapping = renumber_options.get("section_mapping", {})
        renumber_headings_in_doc(merged_doc, chapter_num, 1, section_mapping)
    
    if renumber_options.get("renumber_figures", False):
        renumber_figures_tables(merged_doc, chapter_num)
    
    # Append remaining documents
    for i, section_data in enumerate(sections_data[1:], start=1):
        # Add page break before new section
        merged_doc.add_page_break()
        
        # Load section document
        section_doc = Document(io.BytesIO(section_data["file_bytes"]))
        
        # Renumber if needed
        if renumber_options.get("renumber_headings", False):
            section_mapping = renumber_options.get("section_mapping", {})
            renumber_headings_in_doc(section_doc, chapter_num, 1, section_mapping)
        
        if renumber_options.get("renumber_figures", False):
            renumber_figures_tables(section_doc, chapter_num)
        
        # Copy paragraphs
        for para in section_doc.paragraphs:
            new_para = merged_doc.add_paragraph()
            new_para.style = merged_doc.styles[para.style.name] if para.style.name in [s.name for s in merged_doc.styles] else merged_doc.styles['Normal']
            new_para.alignment = para.alignment
            
            # Copy paragraph formatting
            if para.paragraph_format.space_before:
                new_para.paragraph_format.space_before = para.paragraph_format.space_before
            if para.paragraph_format.space_after:
                new_para.paragraph_format.space_after = para.paragraph_format.space_after
            if para.paragraph_format.line_spacing:
                new_para.paragraph_format.line_spacing = para.paragraph_format.line_spacing
            
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                # Copy run formatting
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                if run.font.size:
                    new_run.font.size = run.font.size
                if run.font.name:
                    new_run.font.name = run.font.name
                if run.font.color and run.font.color.rgb:
                    new_run.font.color.rgb = run.font.color.rgb
                
                # Copy images in runs
                drawings = run._element.findall(qn('w:drawing'))
                if drawings:
                    for drawing in drawings:
                        new_run._element.append(copy.deepcopy(drawing))
        
        # Copy tables
        for table in section_doc.tables:
            # Add a blank paragraph before table for spacing
            merged_doc.add_paragraph()
            new_table = merged_doc.add_table(rows=0, cols=len(table.columns))
            new_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Copy table style if available
            try:
                new_table.style = table.style
            except:
                pass
            
            for row in table.rows:
                new_row = new_table.add_row()
                for j, cell in enumerate(row.cells):
                    if j < len(new_row.cells):
                        new_row.cells[j].text = cell.text
                        # Copy cell paragraph formatting
                        for k, para in enumerate(cell.paragraphs):
                            if k < len(new_row.cells[j].paragraphs):
                                target_para = new_row.cells[j].paragraphs[k]
                                target_para.alignment = para.alignment
                                for run in target_para.runs:
                                    if para.runs:
                                        src_run = para.runs[0]
                                        run.bold = src_run.bold
                                        if src_run.font.size:
                                            run.font.size = src_run.font.size
    
    # Insert images at specified positions
    for img_key, img_data in images_dict.items():
        if img_data.get("bytes") and img_data.get("insert_after_text"):
            target_text = img_data["insert_after_text"]
            caption = img_data.get("caption", f"รูปที่ {chapter_num}-{img_key}")
            width = img_data.get("width_cm", 15)
            
            # Find target paragraph
            for idx, para in enumerate(merged_doc.paragraphs):
                if target_text in para.text:
                    insert_image_after_paragraph(
                        merged_doc, idx, img_data["bytes"], caption, width
                    )
                    break
    
    # Save to buffer
    buffer = io.BytesIO()
    merged_doc.save(buffer)
    buffer.seek(0)
    return buffer


# ═══════════════════════════════════════════════════════
# SIDEBAR - Settings
# ═══════════════════════════════════════════════════════
with st.sidebar:
    st.header("⚙️ ตั้งค่าการรวมรายงาน")
    
    st.subheader("📌 หมายเลขบท")
    chapter_num = st.number_input(
        "หมายเลขบท (เช่น 4 = บทที่ 4)",
        min_value=1, max_value=99,
        value=st.session_state.chapter_num,
        key="chapter_input"
    )
    st.session_state.chapter_num = chapter_num
    
    st.subheader("🔢 ตัวเลือกการเรียงลำดับ")
    renumber_headings = st.checkbox("แก้ไขหมายเลขหัวข้อ", value=True)
    renumber_figures = st.checkbox("แก้ไขหมายเลขรูปภาพ/ตาราง", value=True)
    
    st.divider()
    
    st.subheader("📏 ตั้งค่าเอกสาร")
    page_size = st.selectbox("ขนาดกระดาษ", ["A4", "Letter"])
    
    st.divider()
    st.caption("พัฒนาโดย: ระบบรวมรายงาน v1.0")
    st.caption("สำหรับงานออกแบบโครงสร้างชั้นทาง")


# ═══════════════════════════════════════════════════════
# MAIN AREA
# ═══════════════════════════════════════════════════════
tab1, tab2, tab3, tab4 = st.tabs([
    "📁 Upload ไฟล์ Word", 
    "🖼️ Upload รูปภาพ", 
    "✏️ แก้ไขหมายเลขหัวข้อ",
    "📥 รวมและดาวน์โหลด"
])

# ─── TAB 1: Upload Word Files ───
with tab1:
    st.subheader("📁 Upload ไฟล์ Word แยกหัวข้อ")
    st.info("Upload ไฟล์ .docx แยกตามหัวข้อ ลำดับการ upload คือลำดับในรายงาน สามารถลากเปลี่ยนลำดับได้")
    
    uploaded_files = st.file_uploader(
        "เลือกไฟล์ Word (.docx)",
        type=["docx"],
        accept_multiple_files=True,
        key="word_uploader"
    )
    
    if uploaded_files:
        # Store uploaded files
        new_sections = []
        for uf in uploaded_files:
            file_bytes = uf.read()
            uf.seek(0)
            
            # Check if already in session
            existing = [s for s in st.session_state.sections if s["name"] == uf.name]
            if existing:
                existing[0]["file_bytes"] = file_bytes
                new_sections.append(existing[0])
            else:
                info = extract_docx_content(uf)
                uf.seek(0)
                new_sections.append({
                    "name": uf.name,
                    "file_bytes": file_bytes,
                    "info": info,
                    "order": len(new_sections),
                    "enabled": True,
                    "custom_section_num": None
                })
        
        st.session_state.sections = new_sections
    
    if st.session_state.sections:
        st.subheader(f"📋 ไฟล์ที่ Upload แล้ว ({len(st.session_state.sections)} ไฟล์)")
        
        for i, section in enumerate(st.session_state.sections):
            with st.expander(f"{'✅' if section['enabled'] else '⬜'} {i+1}. {section['name']}", expanded=False):
                col1, col2, col3 = st.columns([3, 1, 1])
                
                with col1:
                    info = section.get("info", {})
                    st.write(f"**ย่อหน้า:** {info.get('para_count', 'N/A')} | "
                            f"**ตาราง:** {info.get('table_count', 'N/A')} | "
                            f"**รูปภาพ:** {info.get('image_count', 'N/A')}")
                    
                    if info.get("headings"):
                        st.write("**หัวข้อที่พบ:**")
                        for h in info["headings"][:10]:
                            level = h["level"].replace("Heading ", "H")
                            st.write(f"  - [{level}] {h['text']}")
                
                with col2:
                    section["enabled"] = st.checkbox(
                        "รวมไฟล์นี้",
                        value=section.get("enabled", True),
                        key=f"enable_{i}"
                    )
                
                with col3:
                    # Move up/down buttons
                    if i > 0:
                        if st.button("⬆️ ขึ้น", key=f"up_{i}"):
                            sections = st.session_state.sections
                            sections[i], sections[i-1] = sections[i-1], sections[i]
                            st.rerun()
                    if i < len(st.session_state.sections) - 1:
                        if st.button("⬇️ ลง", key=f"down_{i}"):
                            sections = st.session_state.sections
                            sections[i], sections[i+1] = sections[i+1], sections[i]
                            st.rerun()
    else:
        st.warning("ยังไม่ได้ Upload ไฟล์ Word")


# ─── TAB 2: Upload Images ───
with tab2:
    st.subheader("🖼️ Upload รูปภาพสำหรับแทรกในรายงาน")
    st.info("Upload รูปภาพ เช่น แผนที่ตำแหน่งสำรวจปริมาณจราจร กราฟ CBR ฯลฯ แล้วกำหนดตำแหน่งที่จะแทรก")
    
    num_images = st.number_input("จำนวนรูปภาพที่ต้องการแทรก", min_value=0, max_value=20, value=len(st.session_state.images), key="num_images")
    
    for img_idx in range(num_images):
        st.divider()
        st.write(f"**รูปที่ {img_idx + 1}**")
        
        col1, col2 = st.columns([1, 1])
        
        with col1:
            img_file = st.file_uploader(
                f"เลือกรูปภาพ #{img_idx + 1}",
                type=["png", "jpg", "jpeg", "gif", "bmp"],
                key=f"img_upload_{img_idx}"
            )
            
            if img_file:
                img_bytes = img_file.read()
                img_file.seek(0)
                
                # Show preview
                st.image(img_file, caption=f"Preview: {img_file.name}", width=300)
                img_file.seek(0)
                
                if img_idx not in st.session_state.images:
                    st.session_state.images[img_idx] = {}
                st.session_state.images[img_idx]["bytes"] = img_bytes
                st.session_state.images[img_idx]["filename"] = img_file.name
        
        with col2:
            caption = st.text_input(
                f"คำบรรยายรูป (Caption) #{img_idx + 1}",
                value=f"รูปที่ {chapter_num}-{img_idx + 1} ",
                key=f"img_caption_{img_idx}"
            )
            
            insert_text = st.text_input(
                f"แทรกหลังข้อความ #{img_idx + 1}",
                placeholder="พิมพ์ข้อความในรายงานที่ต้องการแทรกรูปหลังจากนั้น",
                key=f"img_insert_{img_idx}"
            )
            
            width_cm = st.slider(
                f"ความกว้าง (ซม.) #{img_idx + 1}",
                min_value=5, max_value=20, value=15,
                key=f"img_width_{img_idx}"
            )
            
            if img_idx not in st.session_state.images:
                st.session_state.images[img_idx] = {}
            st.session_state.images[img_idx]["caption"] = caption
            st.session_state.images[img_idx]["insert_after_text"] = insert_text
            st.session_state.images[img_idx]["width_cm"] = width_cm


# ─── TAB 3: Edit Section Numbers ───
with tab3:
    st.subheader("✏️ แก้ไขหมายเลขบทและหัวข้อ")
    
    st.write(f"**บทปัจจุบัน: บทที่ {chapter_num}**")
    
    if st.session_state.sections:
        st.write("---")
        st.write("**แก้ไขหมายเลขหัวข้อ (Section Mapping)**")
        st.caption("กำหนดหมายเลขหัวข้อใหม่สำหรับแต่ละหัวข้อที่พบในไฟล์ Word")
        
        # Collect all headings from all sections
        all_headings = []
        for section in st.session_state.sections:
            if section.get("enabled", True) and section.get("info", {}).get("headings"):
                for h in section["info"]["headings"]:
                    match = re.match(r'^(\d+(?:\.\d+)*)\s*(.*)', h["text"])
                    if match:
                        all_headings.append({
                            "original_num": match.group(1),
                            "text": match.group(2),
                            "level": h["level"],
                            "source": section["name"]
                        })
        
        if all_headings:
            # Create mapping table
            if "section_mapping" not in st.session_state:
                st.session_state.section_mapping = {}
            
            # Header
            col_h1, col_h2, col_h3, col_h4 = st.columns([2, 3, 2, 2])
            col_h1.write("**เลขเดิม**")
            col_h2.write("**ข้อความ**")
            col_h3.write("**เลขใหม่**")
            col_h4.write("**ไฟล์ต้นทาง**")
            
            for idx, h in enumerate(all_headings):
                col1, col2, col3, col4 = st.columns([2, 3, 2, 2])
                
                col1.write(f"`{h['original_num']}`")
                col2.write(h["text"][:50] + ("..." if len(h["text"]) > 50 else ""))
                
                # Default: replace first number with chapter_num
                parts = h["original_num"].split(".")
                parts[0] = str(chapter_num)
                default_new = ".".join(parts)
                
                new_num = col3.text_input(
                    "เลขใหม่",
                    value=st.session_state.section_mapping.get(h["original_num"], default_new),
                    key=f"map_{idx}_{h['original_num']}",
                    label_visibility="collapsed"
                )
                st.session_state.section_mapping[h["original_num"]] = new_num
                
                col4.caption(h["source"][:20])
        else:
            st.info("ไม่พบหัวข้อที่มีหมายเลขในไฟล์ที่ Upload")
        
        st.divider()
        
        # Quick renumber tool
        st.write("**🔧 เครื่องมือเปลี่ยนเลขบทอัตโนมัติ**")
        
        col_a, col_b = st.columns(2)
        with col_a:
            new_chapter = st.number_input(
                "เปลี่ยนเลขบทเป็น",
                min_value=1, max_value=99,
                value=chapter_num,
                key="quick_chapter"
            )
        with col_b:
            if st.button("🔄 เปลี่ยนเลขบททั้งหมด", use_container_width=True):
                st.session_state.chapter_num = new_chapter
                # Update all mappings
                for h in all_headings:
                    parts = h["original_num"].split(".")
                    parts[0] = str(new_chapter)
                    st.session_state.section_mapping[h["original_num"]] = ".".join(parts)
                st.success(f"เปลี่ยนเลขบทเป็น {new_chapter} เรียบร้อย!")
                st.rerun()
        
        st.divider()
        
        # Preview renumbering
        if st.button("👁️ ดูตัวอย่างผลการเปลี่ยนเลข"):
            st.write("**ตัวอย่างผลลัพธ์:**")
            for h in all_headings:
                old = h["original_num"]
                new = st.session_state.section_mapping.get(old, old)
                if old != new:
                    st.write(f"  `{old}` → `{new}` {h['text'][:40]}")
                else:
                    st.write(f"  `{old}` (ไม่เปลี่ยน) {h['text'][:40]}")
    else:
        st.warning("กรุณา Upload ไฟล์ Word ก่อน")


# ─── TAB 4: Merge & Download ───
with tab4:
    st.subheader("📥 รวมรายงานและดาวน์โหลด")
    
    enabled_sections = [s for s in st.session_state.sections if s.get("enabled", True)]
    
    if enabled_sections:
        st.write(f"**ไฟล์ที่จะรวม ({len(enabled_sections)} ไฟล์):**")
        for i, s in enumerate(enabled_sections):
            st.write(f"  {i+1}. {s['name']}")
        
        # Images summary
        active_images = {k: v for k, v in st.session_state.images.items() if v.get("bytes")}
        if active_images:
            st.write(f"**รูปภาพที่จะแทรก ({len(active_images)} รูป):**")
            for k, v in active_images.items():
                st.write(f"  - {v.get('caption', 'N/A')} → แทรกหลัง: {v.get('insert_after_text', 'N/A')[:50]}")
        
        st.divider()
        
        col1, col2 = st.columns(2)
        with col1:
            output_filename = st.text_input(
                "ชื่อไฟล์ Output",
                value=f"บทที่_{chapter_num}_รายงานโครงสร้างชั้นทาง.docx"
            )
        
        with col2:
            st.write("")  # spacer
            st.write("")
            merge_button = st.button("🔀 รวมรายงาน", type="primary", use_container_width=True)
        
        if merge_button:
            with st.spinner("กำลังรวมรายงาน..."):
                try:
                    renumber_options = {
                        "renumber_headings": renumber_headings,
                        "renumber_figures": renumber_figures,
                        "section_mapping": st.session_state.get("section_mapping", {})
                    }
                    
                    result = merge_documents(
                        enabled_sections,
                        chapter_num,
                        st.session_state.images,
                        renumber_options
                    )
                    
                    if result:
                        st.success("✅ รวมรายงานสำเร็จ!")
                        
                        st.download_button(
                            label="📥 ดาวน์โหลดรายงาน",
                            data=result.getvalue(),
                            file_name=output_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary",
                            use_container_width=True
                        )
                    else:
                        st.error("เกิดข้อผิดพลาดในการรวมรายงาน")
                
                except Exception as e:
                    st.error(f"เกิดข้อผิดพลาด: {str(e)}")
                    st.exception(e)
    else:
        st.warning("กรุณา Upload และเลือกไฟล์ Word ที่ต้องการรวม")
    
    st.divider()
    
    # ─── Instructions ───
    with st.expander("📖 คู่มือการใช้งาน", expanded=False):
        st.markdown("""
        ### วิธีการใช้งาน
        
        **ขั้นตอนที่ 1: Upload ไฟล์ Word**
        - ไปที่แท็บ "Upload ไฟล์ Word"
        - Upload ไฟล์ .docx แยกตามหัวข้อ (เช่น 4.1, 4.2, 4.3, ...)
        - เรียงลำดับไฟล์ด้วยปุ่ม ⬆️⬇️
        - เลือก/ยกเลิกไฟล์ที่ต้องการรวม
        
        **ขั้นตอนที่ 2: Upload รูปภาพ (ถ้ามี)**
        - ไปที่แท็บ "Upload รูปภาพ"
        - กำหนดจำนวนรูปภาพ
        - Upload รูปและกำหนดตำแหน่งแทรก
        - ระบุ Caption และข้อความอ้างอิงตำแหน่ง
        
        **ขั้นตอนที่ 3: แก้ไขหมายเลข**
        - ไปที่แท็บ "แก้ไขหมายเลขหัวข้อ"
        - เปลี่ยนหมายเลขบท (เช่น บท 4 → บท 5)
        - แก้ไขหมายเลขหัวข้อย่อยทีละรายการ
        - ใช้เครื่องมือเปลี่ยนเลขบทอัตโนมัติ
        
        **ขั้นตอนที่ 4: รวมและดาวน์โหลด**
        - ไปที่แท็บ "รวมและดาวน์โหลด"
        - ตรวจสอบรายการไฟล์และรูปภาพ
        - กด "รวมรายงาน" แล้วดาวน์โหลด
        
        ### ตัวอย่างรูปแบบหัวข้อที่รองรับ
        - `4.1 แนวทางการวิเคราะห์ออกแบบโครงสร้างชั้นทาง`
        - `4.2.1 การสำรวจและวิเคราะห์ปริมาณการจราจร`
        - `4.4.4 ขั้นตอนการคำนวณความหนาชั้นทาง`
        
        ### หมายเหตุ
        - รองรับไฟล์ .docx เท่านั้น (ไม่รองรับ .doc)
        - รูปภาพรองรับ PNG, JPG, GIF, BMP
        - การแก้ไขหมายเลขจะเปลี่ยนเฉพาะตัวเลขนำหน้าหัวข้อ
        """)
