# -*- coding: utf-8 -*-
import streamlit as st
from docx import Document
from docxcompose.composer import Composer
import io
from pathlib import Path

# ---------------------------------------------------------
# CONFIG
# ---------------------------------------------------------
TEMPLATE_PATH = Path("template.docx")

SECTION_TITLES = {
    "truck_factor": "การคำนวณ Truck Factor",
    "esal_flex": "การคำนวณ ESALs ผิวทางลาดยาง (Flexible Pavement)",
    "esal_rigid": "การคำนวณ ESALs ผิวทางคอนกรีต (Rigid Pavement)",
    "cbr": "การวิเคราะห์ค่า CBR ที่เปอร์เซ็นต์ไทล์",
    "ac_design": "การออกแบบผิวทางลาดยาง (Flexible Pavement)",
    "jpcp_jrcp": "การออกแบบผิวทางคอนกรีต JPCP/JRCP",
    "k_jpcp": "การคำนวณ k-value สำหรับ JPCP/JRCP",
    "crcp": "การออกแบบผิวทางคอนกรีต CRCP",
    "k_crcp": "การคำนวณ k-value สำหรับ CRCP",
    "cost": "การประมาณราคาค่าก่อสร้าง"
}

# ---------------------------------------------------------
# UPDATED ORDER (ตามที่อาจารย์ต้องการ)
# ---------------------------------------------------------
UPLOAD_ORDER = [
    "truck_factor",
    "esal_flex",
    "esal_rigid",
    "cbr",
    "ac_design",
    "jpcp_jrcp",
    "k_jpcp",      # ← ย้ายมาต่อจาก jpcp_jrcp
    "crcp",
    "k_crcp",
    "cost"
]

# ---------------------------------------------------------
# MERGE FUNCTION
# ---------------------------------------------------------
def merge_documents(uploaded_files):
    """รวมไฟล์ Word ตามลำดับ พร้อมหัวข้อแบบ Heading 1"""

    master = Document(TEMPLATE_PATH)
    composer = Composer(master)

    for key in UPLOAD_ORDER:
        file = uploaded_files.get(key)
        if file is None:
            continue

        # เพิ่มหัวข้อ (Heading 1)
        header_doc = Document()
        h = header_doc.add_paragraph(style="Heading 1")
        h.add_run(SECTION_TITLES[key])

        temp_header = io.BytesIO()
        header_doc.save(temp_header)
        temp_header.seek(0)
        composer.append(Document(temp_header))

        # เพิ่มเนื้อหาไฟล์จริง
        file_bytes = file.read()
        file.seek(0)
        composer.append(Document(io.BytesIO(file_bytes)))

    return composer.doc


# ---------------------------------------------------------
# STREAMLIT UI
# ---------------------------------------------------------
st.set_page_config(page_title="Pavement Report Merger", layout="wide")

st.title("🛣️ โปรแกรมรวมรายงานออกแบบโครงสร้างชั้นทาง (Consultant Edition)")
st.write("อัปโหลดไฟล์รายงานแต่ละส่วน แล้วกดปุ่มเพื่อรวมเป็นไฟล์เดียว")

uploaded_files = {}

# UI อัปโหลดไฟล์
for key in UPLOAD_ORDER:
    uploaded_files[key] = st.file_uploader(
        f"{SECTION_TITLES[key]}",
        type=["docx"],
        key=key
    )

# นับจำนวนไฟล์
count_uploaded = sum(1 for f in uploaded_files.values() if f is not None)
st.info(f"อัปโหลดแล้ว {count_uploaded} จาก {len(UPLOAD_ORDER)} ไฟล์")

# ปุ่มรวมไฟล์
if st.button("🔄 รวมไฟล์และสร้างรายงาน"):
    if count_uploaded == 0:
        st.error("กรุณาอัปโหลดอย่างน้อย 1 ไฟล์")
    else:
        with st.spinner("กำลังรวมไฟล์..."):
            merged_doc = merge_documents(uploaded_files)

            output = io.BytesIO()
            merged_doc.save(output)
            output.seek(0)

            st.success("รวมไฟล์สำเร็จ!")

            st.download_button(
                label="📄 ดาวน์โหลดรายงานรวม (.docx)",
                data=output,
                file_name="Pavement_Report_Consultant.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
