"""
================================================================================
AASHTO 1993 Flexible Pavement Design - Streamlit Application (Version 5)
================================================================================
แอปพลิเคชันสำหรับออกแบบ Flexible Pavement ตามวิธี AASHTO 1993
ปรับปรุงตามมาตรฐานกรมทางหลวง (DOH Thailand)

[V5 Improvements - Friendly User Edition]
- Preset โครงสร้างถนนมาตรฐาน ทล. (Auto-fill)
- Tab Layout แทน 2 Columns (กว้างเต็มจอ)
- Quick Summary Card ด้านบน
- AC Sublayer แบบตาราง compact
- Input Validation & Warning messages
- Drainage Coefficient Reference Table
- Sensitivity Analysis Chart
- Session State Management ปรับปรุง
- Word Report + AC sublayer breakdown
- Bug fixes (fontsize, memory leak)

Author: รศ.ดร.อิทธิพล มีผล // ภาควิชาครุศาสตร์โยธา // มจพ.
Version: 5.0
================================================================================
"""

import streamlit as st
import numpy as np
import json
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import matplotlib.font_manager as fm
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ================================================================================
# CUSTOM ROOT-FINDING (แทน scipy.optimize.brentq)
# ================================================================================

def brentq(f, a, b, xtol=1e-12, maxiter=200):
    """
    Brent's method for root-finding — ไม่ต้องพึ่ง scipy
    หาค่า x ใน [a, b] ที่ f(x) = 0
    """
    fa, fb = f(a), f(b)
    if fa * fb > 0:
        raise ValueError(f"f(a) and f(b) must have different signs: f({a})={fa:.4f}, f({b})={fb:.4f}")
    if abs(fa) < xtol:
        return a
    if abs(fb) < xtol:
        return b

    c, fc = a, fa
    d = e = b - a

    for _ in range(maxiter):
        if fb * fc > 0:
            c, fc = a, fa
            d = e = b - a

        if abs(fc) < abs(fb):
            a, b, c = b, c, b
            fa, fb, fc = fb, fc, fb

        tol1 = 2.0 * 2.2e-16 * abs(b) + 0.5 * xtol
        m = 0.5 * (c - b)

        if abs(m) <= tol1 or fb == 0.0:
            return b

        if abs(e) >= tol1 and abs(fa) > abs(fb):
            s = fb / fa
            if a == c:
                p = 2.0 * m * s
                q = 1.0 - s
            else:
                q = fa / fc
                r = fb / fc
                p = s * (2.0 * m * q * (q - r) - (b - a) * (r - 1.0))
                q = (q - 1.0) * (r - 1.0) * (s - 1.0)
            if p > 0:
                q = -q
            else:
                p = -p
            if 2.0 * p < min(3.0 * m * q - abs(tol1 * q), abs(e * q)):
                e = d
                d = p / q
            else:
                d = m
                e = m
        else:
            d = m
            e = m

        a, fa = b, fb
        if abs(d) > tol1:
            b += d
        else:
            b += tol1 if m > 0 else -tol1
        fb = f(b)

    return b

# ================================================================================
# PAGE CONFIGURATION
# ================================================================================

st.set_page_config(
    page_title="Flexible Pavement Design (AASHTO 1993) v5",
    page_icon="🛣️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ================================================================================
# MATERIAL DATABASE - ตามมาตรฐานกรมทางหลวง (DOH Thailand)
# ================================================================================

MATERIALS = {
    # ============ ชั้นผิวทาง (Surface) ============
    "ผิวทางลาดยาง AC": {
        "layer_coeff": 0.40,
        "drainage_coeff": 1.0,
        "mr_psi": 362500,
        "mr_mpa": 2500,
        "layer_type": "surface",
        "color": "#1C1C1C",
        "short_name": "AC",
        "english_name": "Asphalt Concrete"
    },
    "ผิวทางลาดยาง PMA": {
        "layer_coeff": 0.40,
        "drainage_coeff": 1.0,
        "mr_psi": 536500,
        "mr_mpa": 3700,
        "layer_type": "surface",
        "color": "#2C2C2C",
        "short_name": "PMA",
        "english_name": "Polymer Modified Asphalt"
    },

    # ============ ชั้นพื้นทาง (Base) ============
    "พื้นทางซีเมนต์ CTB": {
        "layer_coeff": 0.18,
        "drainage_coeff": 1.0,
        "mr_psi": 174000,
        "mr_mpa": 1200,
        "layer_type": "base",
        "color": "#78909C",
        "short_name": "CTB",
        "english_name": "Cement Treated Base"
    },
    "พื้นทางหินคลุกผสมซีเมนต์ UCS 24.5 ksc.": {
        "layer_coeff": 0.15,
        "drainage_coeff": 1.0,
        "mr_psi": 123250,
        "mr_mpa": 850,
        "layer_type": "base",
        "color": "#607D8B",
        "short_name": "MOD.CRB",
        "english_name": "Mod.Crushed Rock Base"
    },
    "พื้นทางหินคลุก CBR 80%": {
        "layer_coeff": 0.13,
        "drainage_coeff": 1.0,
        "mr_psi": 50750,
        "mr_mpa": 350,
        "layer_type": "base",
        "color": "#795548",
        "short_name": "CAB",
        "english_name": "Crushed Rock Base"
    },
    "พื้นทางดินซีเมนต์ UCS 17.5 ksc.": {
        "layer_coeff": 0.13,
        "drainage_coeff": 1.0,
        "mr_psi": 50750,
        "mr_mpa": 350,
        "layer_type": "base",
        "color": "#8D6E63",
        "short_name": "SCB",
        "english_name": "Soil Cement Base"
    },
    "พื้นทางวัสดุหมุนเวียน (Recycling)": {
        "layer_coeff": 0.15,
        "drainage_coeff": 1.0,
        "mr_psi": 123250,
        "mr_mpa": 850,
        "layer_type": "base",
        "color": "#5D4037",
        "short_name": "RAP",
        "english_name": "Recycled Asphalt Pavement"
    },

    # ============ ชั้นรองพื้นทาง (Subbase) - วัสดุมวลรวม ============
    "รองพื้นทางวัสดุมวลรวม CBR 25%": {
        "layer_coeff": 0.10,
        "drainage_coeff": 1.0,
        "mr_psi": 21750,
        "mr_mpa": 150,
        "layer_type": "subbase",
        "color": "#FFB74D",
        "short_name": "GSB",
        "english_name": "Aggregate Subbase"
    },

    # ============ วัสดุคัดเลือก (Selected Material) ============
    "วัสดุคัดเลือก ก": {
        "layer_coeff": 0.08,
        "drainage_coeff": 1.0,
        "mr_psi": 14504,
        "mr_mpa": 100,
        "layer_type": "selected",
        "color": "#FFF176",
        "short_name": "SM-A",
        "english_name": "Selected Material"
    },

    # ============ ไม่ใช้วัสดุ (Skip layer) ============
    "ไม่ใช้วัสดุคัดเลือก (ใช้ดินทางทรพ)": {
        "layer_coeff": 0.00,
        "drainage_coeff": 1.0,
        "mr_psi": 0,
        "mr_mpa": 0,
        "layer_type": "none",
        "color": "#D7CCC8",
        "short_name": "NONE",
        "english_name": "None"
    }
}

# ================================================================================
# PRESET STRUCTURES - โครงสร้างมาตรฐาน ทล.
# ================================================================================

PRESET_STRUCTURES = {
    "--- เลือกโครงสร้างมาตรฐาน ---": None,
    "AC + CTB + GSB + SM (มาตรฐานหลัก)": {
        "description": "ผิวทาง AC / พื้นทางซีเมนต์ CTB / รองพื้นทาง GSB / วัสดุคัดเลือก",
        "num_layers": 4,
        "layers": [
            {"material": "ผิวทางลาดยาง AC", "thickness_cm": 15.0},
            {"material": "พื้นทางซีเมนต์ CTB", "thickness_cm": 15.0},
            {"material": "รองพื้นทางวัสดุมวลรวม CBR 25%", "thickness_cm": 15.0},
            {"material": "วัสดุคัดเลือก ก", "thickness_cm": 30.0},
        ]
    },
    "AC + MOD.CRB + GSB + SM": {
        "description": "ผิวทาง AC / หินคลุกผสมซีเมนต์ / รองพื้นทาง GSB / วัสดุคัดเลือก",
        "num_layers": 4,
        "layers": [
            {"material": "ผิวทางลาดยาง AC", "thickness_cm": 15.0},
            {"material": "พื้นทางหินคลุกผสมซีเมนต์ UCS 24.5 ksc.", "thickness_cm": 20.0},
            {"material": "รองพื้นทางวัสดุมวลรวม CBR 25%", "thickness_cm": 15.0},
            {"material": "วัสดุคัดเลือก ก", "thickness_cm": 30.0},
        ]
    },
    "AC + CAB + GSB + SM": {
        "description": "ผิวทาง AC / หินคลุก CBR 80% / รองพื้นทาง GSB / วัสดุคัดเลือก",
        "num_layers": 4,
        "layers": [
            {"material": "ผิวทางลาดยาง AC", "thickness_cm": 15.0},
            {"material": "พื้นทางหินคลุก CBR 80%", "thickness_cm": 20.0},
            {"material": "รองพื้นทางวัสดุมวลรวม CBR 25%", "thickness_cm": 15.0},
            {"material": "วัสดุคัดเลือก ก", "thickness_cm": 30.0},
        ]
    },
    "AC + SCB + GSB + SM": {
        "description": "ผิวทาง AC / ดินซีเมนต์ / รองพื้นทาง GSB / วัสดุคัดเลือก",
        "num_layers": 4,
        "layers": [
            {"material": "ผิวทางลาดยาง AC", "thickness_cm": 15.0},
            {"material": "พื้นทางดินซีเมนต์ UCS 17.5 ksc.", "thickness_cm": 20.0},
            {"material": "รองพื้นทางวัสดุมวลรวม CBR 25%", "thickness_cm": 15.0},
            {"material": "วัสดุคัดเลือก ก", "thickness_cm": 30.0},
        ]
    },
    "AC + CTB + GSB (ไม่ใช้ SM)": {
        "description": "ผิวทาง AC / พื้นทาง CTB / รองพื้นทาง GSB (ไม่มีวัสดุคัดเลือก)",
        "num_layers": 3,
        "layers": [
            {"material": "ผิวทางลาดยาง AC", "thickness_cm": 15.0},
            {"material": "พื้นทางซีเมนต์ CTB", "thickness_cm": 20.0},
            {"material": "รองพื้นทางวัสดุมวลรวม CBR 25%", "thickness_cm": 20.0},
        ]
    },
    "PMA + CTB + GSB + SM": {
        "description": "ผิวทาง PMA / พื้นทาง CTB / รองพื้นทาง GSB / วัสดุคัดเลือก",
        "num_layers": 4,
        "layers": [
            {"material": "ผิวทางลาดยาง PMA", "thickness_cm": 15.0},
            {"material": "พื้นทางซีเมนต์ CTB", "thickness_cm": 15.0},
            {"material": "รองพื้นทางวัสดุมวลรวม CBR 25%", "thickness_cm": 15.0},
            {"material": "วัสดุคัดเลือก ก", "thickness_cm": 30.0},
        ]
    },
    "AC + RAP + GSB + SM": {
        "description": "ผิวทาง AC / พื้นทาง Recycling / รองพื้นทาง GSB / วัสดุคัดเลือก",
        "num_layers": 4,
        "layers": [
            {"material": "ผิวทางลาดยาง AC", "thickness_cm": 15.0},
            {"material": "พื้นทางวัสดุหมุนเวียน (Recycling)", "thickness_cm": 20.0},
            {"material": "รองพื้นทางวัสดุมวลรวม CBR 25%", "thickness_cm": 15.0},
            {"material": "วัสดุคัดเลือก ก", "thickness_cm": 30.0},
        ]
    },
}

# ================================================================================
# RELIABILITY TABLE: Zr VALUES
# ================================================================================

RELIABILITY_ZR = {
    50: -0.000,
    60: -0.253,
    70: -0.524,
    75: -0.674,
    80: -0.841,
    85: -1.037,
    90: -1.282,
    91: -1.340,
    92: -1.405,
    93: -1.476,
    94: -1.555,
    95: -1.645,
    96: -1.751,
    97: -1.881,
    98: -2.054,
    99: -2.327,
    99.9: -3.090
}

# ================================================================================
# DRAINAGE COEFFICIENT TABLE (AASHTO Table 2.4)
# ================================================================================

DRAINAGE_TABLE = {
    "Excellent": {"description": "ระบายน้ำดีเยี่ยม (< 2 ชม.)", 
                  "values": {"<1%": 1.40, "1-5%": 1.35, "5-25%": 1.30, ">25%": 1.20}},
    "Good":      {"description": "ระบายน้ำดี (1 วัน)", 
                  "values": {"<1%": 1.35, "1-5%": 1.25, "5-25%": 1.15, ">25%": 1.00}},
    "Fair":      {"description": "ระบายน้ำพอใช้ (1 สัปดาห์)", 
                  "values": {"<1%": 1.25, "1-5%": 1.15, "5-25%": 1.05, ">25%": 0.80}},
    "Poor":      {"description": "ระบายน้ำไม่ดี (1 เดือน)", 
                  "values": {"<1%": 1.15, "1-5%": 1.05, "5-25%": 0.80, ">25%": 0.60}},
    "Very Poor": {"description": "ระบายน้ำไม่ดีมาก (ไม่ระบาย)", 
                  "values": {"<1%": 1.05, "1-5%": 0.80, "5-25%": 0.60, ">25%": 0.40}},
}

# DOH AC Sublayer Thickness Standards (mm)
DOH_THICKNESS_STANDARDS = {
    "Wearing Course": [40, 45, 50, 55, 60, 65, 70],
    "Binder Course": [40, 45, 50, 55, 60, 65, 70, 75, 80],
    "Base Course": [0, 70, 75, 80, 85, 90, 95, 100]
}

# ================================================================================
# CORE CALCULATION FUNCTIONS
# ================================================================================

def aashto_1993_equation(SN, W18, Zr, So, delta_psi, Mr):
    """AASHTO 1993 Main Design Equation for Flexible Pavement"""
    log_W18 = np.log10(W18)
    term1 = Zr * So
    term2 = 9.36 * np.log10(SN + 1) - 0.20
    numerator = np.log10(delta_psi / (4.2 - 1.5))
    denominator = 0.4 + (1094 / ((SN + 1) ** 5.19))
    term3 = numerator / denominator
    term4 = 2.32 * np.log10(Mr) - 8.07
    right_side = term1 + term2 + term3 + term4
    return right_side - log_W18


def calculate_sn_for_layer(W18, Zr, So, delta_psi, Mr):
    """Calculate required SN for a given subgrade/layer modulus"""
    def f(SN):
        return aashto_1993_equation(SN, W18, Zr, So, delta_psi, Mr)
    try:
        SN_required = brentq(f, 0.01, 25.0, xtol=1e-6, maxiter=100)
        return round(SN_required, 2)
    except ValueError:
        return None


def calculate_w18_supported(SN, Zr, So, delta_psi, Mr):
    """Calculate W18 that can be supported by a given SN"""
    term1 = Zr * So
    term2 = 9.36 * np.log10(SN + 1) - 0.20
    numerator = np.log10(delta_psi / (4.2 - 1.5))
    denominator = 0.4 + (1094 / ((SN + 1) ** 5.19))
    term3 = numerator / denominator
    term4 = 2.32 * np.log10(Mr) - 8.07
    log_W18 = term1 + term2 + term3 + term4
    return 10 ** log_W18


def calculate_layer_thicknesses(W18, Zr, So, delta_psi, subgrade_mr, layers, ac_sublayers=None):
    """Calculate minimum thickness for each layer using AASHTO 1993 method"""
    results = {
        'layers': [],
        'sn_values': [],
        'subgrade_mr': subgrade_mr,
        'total_sn_required': None,
        'total_sn_provided': 0,
        'ac_sublayers': ac_sublayers,
        'warnings': []  # เพิ่มระบบ warning
    }

    # Get active layers
    active_layers = [l for l in layers if l['material'] != "ไม่ใช้วัสดุคัดเลือก (ใช้ดินทางทรพ)"]
    if not active_layers:
        results['warnings'].append("⚠️ ไม่มีชั้นทางที่ active")
        return results

    num_layers = len(active_layers)
    sn_values = []

    # ===== Validation: ตรวจสอบลำดับ Mr =====
    for i in range(num_layers - 1):
        mr_current = MATERIALS[active_layers[i]['material']]['mr_psi']
        mr_next = MATERIALS[active_layers[i + 1]['material']]['mr_psi']
        if mr_current < mr_next:
            results['warnings'].append(
                f"⚠️ ชั้นที่ {i+1} ({active_layers[i]['material']}) มีค่า Mr = {mr_current:,} psi "
                f"ต่ำกว่าชั้นที่ {i+2} ({active_layers[i+1]['material']}) ที่มี Mr = {mr_next:,} psi "
                f"— ปกติชั้นบนควรมีค่า Mr สูงกว่าชั้นล่าง"
            )

    for i in range(num_layers):
        if i == num_layers - 1:
            mr_below = subgrade_mr
        else:
            mat_below = MATERIALS[active_layers[i + 1]['material']]
            mr_below = mat_below['mr_psi']

        sn_i = calculate_sn_for_layer(W18, Zr, So, delta_psi, mr_below)
        if sn_i is None:
            results['warnings'].append(
                f"⚠️ ไม่สามารถคำนวณ SN สำหรับชั้นที่ {i+1} ได้ "
                f"(Mr_below = {mr_below:,} psi) — ค่า W18 อาจสูงเกินไป หรือพารามิเตอร์ไม่เหมาะสม"
            )
        sn_values.append({
            'layer_index': i + 1,
            'mr_below': mr_below,
            'sn_required': sn_i
        })

    results['sn_values'] = sn_values
    results['total_sn_required'] = calculate_sn_for_layer(W18, Zr, So, delta_psi, subgrade_mr)

    if results['total_sn_required'] is None:
        results['warnings'].append(
            f"⚠️ ไม่สามารถคำนวณ SN_required จาก Subgrade Mr = {subgrade_mr:,} psi ได้ "
            f"— ลองปรับค่า W18, Reliability หรือ CBR"
        )

    cumulative_sn = 0

    for i, layer in enumerate(active_layers):
        mat = MATERIALS[layer['material']]
        a_i = layer.get('layer_coeff', mat['layer_coeff'])
        m_i = layer.get('drainage_coeff', 1.0)

        sn_required_at_layer = sn_values[i]['sn_required'] if sn_values[i]['sn_required'] else 0

        if a_i > 0 and m_i > 0:
            remaining_sn = max(0, sn_required_at_layer - cumulative_sn)
            min_thickness_inch = remaining_sn / (a_i * m_i)
            min_thickness_cm = min_thickness_inch * 2.54
        else:
            min_thickness_inch = 0
            min_thickness_cm = 0

        design_thickness_cm = layer['thickness_cm']
        design_thickness_inch = design_thickness_cm / 2.54

        sn_contribution = a_i * design_thickness_inch * m_i
        cumulative_sn += sn_contribution

        is_ok = design_thickness_cm >= min_thickness_cm

        layer_ac_sublayers = None
        if i == 0 and ac_sublayers is not None:
            layer_ac_sublayers = ac_sublayers

        results['layers'].append({
            'layer_no': i + 1,
            'material': layer['material'],
            'short_name': mat['short_name'],
            'english_name': mat.get('english_name', mat['short_name']),
            'mr_psi': mat['mr_psi'],
            'mr_mpa': mat['mr_mpa'],
            'a_i': a_i,
            'm_i': m_i,
            'sn_required_at_layer': sn_required_at_layer,
            'min_thickness_inch': round(min_thickness_inch, 2),
            'min_thickness_cm': round(min_thickness_cm, 1),
            'design_thickness_cm': design_thickness_cm,
            'design_thickness_inch': round(design_thickness_inch, 2),
            'sn_contribution': round(sn_contribution, 4),
            'cumulative_sn': round(cumulative_sn, 2),
            'is_ok': is_ok,
            'color': mat['color'],
            'ac_sublayers': layer_ac_sublayers
        })

    results['total_sn_provided'] = round(cumulative_sn, 2)
    return results


def check_design(sn_required, sn_provided):
    """Check if design is adequate"""
    if sn_required is None:
        return {
            'status': 'ERROR',
            'passed': False,
            'message': 'ไม่สามารถคำนวณ SN_required ได้',
            'safety_margin': 0
        }
    safety_margin = sn_provided - sn_required
    passed = sn_provided >= sn_required
    return {
        'status': 'OK' if passed else 'NG',
        'passed': passed,
        'safety_margin': round(safety_margin, 2),
        'message': f"SN_provided ({sn_provided:.2f}) {'≥' if passed else '<'} SN_required ({sn_required:.2f})"
    }


# ================================================================================
# SENSITIVITY ANALYSIS
# ================================================================================

def plot_sensitivity_cbr(W18, Zr, So, delta_psi, current_cbr):
    """Plot SN_required vs CBR"""
    cbr_range = np.linspace(2, 20, 50)
    sn_values = []
    for cbr in cbr_range:
        mr = 1500 * cbr
        sn = calculate_sn_for_layer(W18, Zr, So, delta_psi, mr)
        sn_values.append(sn if sn else np.nan)

    fig, ax = plt.subplots(figsize=(8, 4))
    ax.plot(cbr_range, sn_values, 'b-', linewidth=2, label='SN required')
    
    # Mark current CBR
    current_mr = 1500 * current_cbr
    current_sn = calculate_sn_for_layer(W18, Zr, So, delta_psi, current_mr)
    if current_sn:
        ax.plot(current_cbr, current_sn, 'ro', markersize=12, label=f'Current: CBR={current_cbr}%, SN={current_sn:.2f}')
    
    ax.set_xlabel('CBR (%)', fontsize=12)
    ax.set_ylabel('SN Required', fontsize=12)
    ax.set_title('Sensitivity: SN Required vs CBR', fontsize=14, fontweight='bold')
    ax.legend(fontsize=11)
    ax.grid(True, alpha=0.3)
    try:
        plt.tight_layout()
    except Exception:
        pass
    return fig


def plot_sensitivity_w18(Zr, So, delta_psi, Mr, current_w18):
    """Plot SN_required vs W18"""
    w18_range = np.logspace(5, 8.5, 50)  # 100,000 to ~300M
    sn_values = []
    for w18 in w18_range:
        sn = calculate_sn_for_layer(w18, Zr, So, delta_psi, Mr)
        sn_values.append(sn if sn else np.nan)

    fig, ax = plt.subplots(figsize=(8, 4))
    ax.semilogx(w18_range, sn_values, 'g-', linewidth=2, label='SN required')
    
    # Mark current W18
    current_sn = calculate_sn_for_layer(current_w18, Zr, So, delta_psi, Mr)
    if current_sn:
        ax.semilogx(current_w18, current_sn, 'ro', markersize=12, 
                     label=f'Current: W18={current_w18/1e6:.2f}M, SN={current_sn:.2f}')
    
    ax.set_xlabel('W₁₈ (ESALs)', fontsize=12)
    ax.set_ylabel('SN Required', fontsize=12)
    ax.set_title('Sensitivity: SN Required vs W₁₈', fontsize=14, fontweight='bold')
    ax.legend(fontsize=11)
    ax.grid(True, alpha=0.3)
    try:
        plt.tight_layout()
    except Exception:
        pass
    return fig


# ================================================================================
# VISUALIZATION FUNCTIONS
# ================================================================================

def _get_thai_fonts():
    """
    หา Thai font สำหรับ matplotlib
    - ค้นหาในระบบ (Garuda จาก packages.txt / Loma / Noto)
    - Fallback = DejaVu Sans (ไม่แสดงไทย แต่ไม่ crash)
    Return: (font_regular, font_bold, has_thai)
    """
    import os
    
    # ค้นหา font ในระบบ (Garuda จาก packages.txt เป็นอันดับแรก)
    sys_candidates = [
        ('/usr/share/fonts/truetype/tlwg/Garuda.ttf', '/usr/share/fonts/truetype/tlwg/Garuda-Bold.ttf'),
        ('/usr/share/fonts/opentype/tlwg/Garuda.otf', '/usr/share/fonts/opentype/tlwg/Garuda-Bold.otf'),
        ('/usr/share/fonts/truetype/tlwg/Loma.ttf', '/usr/share/fonts/truetype/tlwg/Loma-Bold.ttf'),
        ('/usr/share/fonts/opentype/tlwg/Loma.otf', '/usr/share/fonts/opentype/tlwg/Loma-Bold.otf'),
        ('/usr/share/fonts/truetype/noto/NotoSansThai-Regular.ttf', '/usr/share/fonts/truetype/noto/NotoSansThai-Bold.ttf'),
    ]
    for reg, bold in sys_candidates:
        if os.path.exists(reg):
            fp_r = fm.FontProperties(fname=reg)
            fp_b = fm.FontProperties(fname=bold) if os.path.exists(bold) else fm.FontProperties(fname=reg)
            return fp_r, fp_b, True
    
    # Fallback — DejaVu Sans (English only)
    return (fm.FontProperties(family='DejaVu Sans'),
            fm.FontProperties(family='DejaVu Sans', weight='bold'),
            False)


# Cache font ไว้ใน session (เรียกครั้งเดียว)
@st.cache_resource
def get_cached_thai_fonts():
    """Cache Thai font resource เพื่อไม่ต้องค้นหาซ้ำทุก rerun"""
    return _get_thai_fonts()


def plot_pavement_section(layers_result, subgrade_mr=None, subgrade_cbr=None, lang='en'):
    """Draw vertical pavement section diagram — auto fallback to English if no Thai font"""

    plt.rcParams['font.family'] = 'DejaVu Sans'

    # ตรวจสอบ Thai font
    thai_font = thai_font_bold = None
    has_thai = False
    if lang == 'th':
        thai_font, thai_font_bold, has_thai = get_cached_thai_fonts()
        if not has_thai:
            lang = 'en'  # fallback

    def _fp(bold=False):
        """Return fontproperties kwarg dict"""
        if has_thai:
            return {'fontproperties': thai_font_bold if bold else thai_font}
        return {}

    if not layers_result:
        fig, ax = plt.subplots(figsize=(12, 8))
        ax.text(0.5, 0.5, 'No layers defined', ha='center', va='center', fontsize=14)
        ax.axis('off')
        return fig

    valid_layers = [l for l in layers_result if l.get('design_thickness_cm', 0) > 0]
    if not valid_layers:
        fig, ax = plt.subplots(figsize=(12, 8))
        ax.text(0.5, 0.5, 'No valid layers', ha='center', va='center', fontsize=14)
        ax.axis('off')
        return fig

    # Expand AC sublayers
    expanded_layers = []
    for layer in valid_layers:
        ac_sub = layer.get('ac_sublayers', None)
        if ac_sub is not None and layer['layer_no'] == 1:
            sub_info = [
                ('wearing', '#1C1C1C', 'ผิวทาง (Wearing Course)',   'Wearing Course'),
                ('binder',  '#333333', 'ยึดเกาะ (Binder Course)',   'Binder Course'),
                ('base',    '#4A4A4A', 'รองผิวทาง (Base Course)',   'Base Course'),
            ]
            for key, color, th_name, en_name in sub_info:
                if ac_sub[key] > 0:
                    expanded_layers.append({
                        'design_thickness_cm': ac_sub[key],
                        'material': th_name if lang == 'th' else en_name,
                        'english_name': en_name,
                        'short_name': key[:2].upper() + 'C',
                        'color': color, 'mr_mpa': layer['mr_mpa'],
                        'is_sublayer': True
                    })
        else:
            expanded_layers.append(layer)
    draw_layers = expanded_layers

    total_thickness = sum(l['design_thickness_cm'] for l in draw_layers)

    fig, ax = plt.subplots(figsize=(12, 9))
    width = 3
    x_center = 7
    x_start = x_center - width / 2

    min_display_height = 6
    display_heights = [max(l['design_thickness_cm'], min_display_height) for l in draw_layers]
    total_display = sum(display_heights)

    dark_colors = ['#1C1C1C', '#2C2C2C', '#333333', '#4A4A4A', '#78909C', '#607D8B',
                   '#795548', '#8D6E63', '#5D4037', '#6D4C41', '#455A64']

    y_current = total_display

    for i, layer in enumerate(draw_layers):
        thickness = layer['design_thickness_cm']
        display_h = display_heights[i]
        color = layer.get('color', '#CCCCCC')
        e_mpa = layer.get('mr_mpa', 0)
        is_sublayer = layer.get('is_sublayer', False)

        if lang == 'th':
            name = layer.get('material', layer.get('short_name', f'Layer {i+1}'))
        else:
            name = layer.get('english_name', layer.get('short_name', f'Layer {i+1}'))

        y_bottom = y_current - display_h
        ls = '--' if is_sublayer else '-'
        lw = 1 if is_sublayer else 2

        rect = mpatches.Rectangle((x_start, y_bottom), width, display_h,
                                  linewidth=lw, linestyle=ls, edgecolor='black', facecolor=color)
        ax.add_patch(rect)

        yc = y_bottom + display_h / 2
        tc = 'white' if color in dark_colors else 'black'

        fs_center = 14 if is_sublayer else 16
        ax.text(x_center, yc, f'{thickness:.0f} cm',
                ha='center', va='center', fontsize=fs_center, fontweight='bold', color=tc)

        fs_name = 12 if is_sublayer else 14
        ax.text(x_start - 0.5, yc, name,
                ha='right', va='center', fontsize=fs_name, fontweight='bold', color='black', **_fp(True))

        if e_mpa and e_mpa > 0 and not is_sublayer:
            ax.text(x_start + width + 0.5, yc, f'E = {e_mpa:,.0f} MPa',
                    ha='left', va='center', fontsize=12, color='#0066CC')

        y_current = y_bottom

    # ===== Subgrade — hatch ก่อน แล้วทับด้วย box สีทึบสำหรับข้อความ =====
    sg_h = 6
    sg_yb = -sg_h
    # วาด hatch background
    ax.add_patch(mpatches.Rectangle(
        (x_start, sg_yb), width, sg_h,
        linewidth=2, edgecolor='black', facecolor='#D7CCC8', hatch='///'))
    
    # วาด box สีทึบตรงกลางสำหรับข้อความ (ไม่ให้ hatch ทับ)
    text_box_h = 3.5
    text_box_w = width * 0.85
    ax.add_patch(mpatches.FancyBboxPatch(
        (x_center - text_box_w / 2, sg_yb + (sg_h - text_box_h) / 2),
        text_box_w, text_box_h,
        boxstyle="round,pad=0.2",
        facecolor='#EFEBE9', edgecolor='#8D6E63', linewidth=1.5, alpha=0.95))

    if lang == 'th':
        sg_label = 'ดินเดิม (Subgrade)'
    else:
        sg_label = 'Subgrade'
    if subgrade_cbr:
        sg_label += f'\nCBR = {subgrade_cbr:.1f}%'
    ax.text(x_center, sg_yb + sg_h / 2, sg_label,
            ha='center', va='center', fontsize=12, fontweight='bold', color='#5D4037', **_fp(True))

    if subgrade_mr:
        ax.text(x_start + width + 0.5, sg_yb + sg_h / 2, f'Mr = {subgrade_mr:,} psi',
                ha='left', va='center', fontsize=12, color='#0066CC')

    # ===== Total thickness arrow =====
    ax.annotate('', xy=(x_start + width + 3.5, total_display),
                xytext=(x_start + width + 3.5, 0),
                arrowprops=dict(arrowstyle='<->', color='red', lw=2))

    if lang == 'th':
        total_label = f'รวม\n{total_thickness:.0f} cm'
    else:
        total_label = f'Total\n{total_thickness:.0f} cm'
    ax.text(x_start + width + 4, total_display / 2, total_label,
            ha='left', va='center', fontsize=14, color='red', fontweight='bold', **_fp(True))

    margin = 10
    ax.set_xlim(0, 15)
    ax.set_ylim(-sg_h - 4, total_display + margin)
    ax.axis('off')

    # Title
    if lang == 'th':
        title_text = 'รูปตัดโครงสร้างชั้นทาง'
    else:
        title_text = 'Pavement Structure'
    ax.set_title(title_text, fontsize=20, fontweight='bold', pad=20, **_fp(True))

    # Bottom box
    if lang == 'th':
        box_text = f'ความหนารวมโครงสร้างชั้นทาง: {total_thickness:.0f} cm'
    else:
        box_text = f'Total Pavement Thickness: {total_thickness:.0f} cm'
    ax.text(x_center, -sg_h - 2, box_text,
            ha='center', va='center', fontsize=15, fontweight='bold', **_fp(True),
            bbox=dict(boxstyle='round', facecolor='lightyellow', alpha=0.9, edgecolor='orange'))

    try:
        plt.tight_layout()
    except Exception:
        pass
    return fig


def get_figure_as_bytes(fig):
    """Convert matplotlib figure to bytes"""
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    buf.seek(0)
    return buf


# ================================================================================
# WORD EXPORT FUNCTION
# ================================================================================

def set_thai_font(run, size_pt=15, bold=False):
    """Set TH Sarabun New font for Thai text"""
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(size_pt)
    run.bold = bold
    run._element.rPr.rFonts.set(
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cs', 'TH Sarabun New')

def set_equation_font(run, size_pt=11, bold=False, italic=True):
    """Set Times New Roman font for equations"""
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic

def add_thai_paragraph(doc, text, size_pt=15, bold=False, alignment=None):
    """Add paragraph with Thai font"""
    para = doc.add_paragraph()
    if alignment:
        para.alignment = alignment
    run = para.add_run(text)
    set_thai_font(run, size_pt, bold)
    return para

def add_equation_paragraph(doc, text, size_pt=11, bold=False, italic=True):
    """Add paragraph with equation font"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_equation_font(run, size_pt, bold, italic)
    return para


def create_word_report(project_title, inputs, calc_results, design_check, fig):
    """Create Word document report with step-by-step calculations"""
    doc = Document()

    # ========================================
    # TITLE
    # ========================================
    title = doc.add_heading('รายงานการออกแบบ Flexible Pavement', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_thai_font(run, size_pt=24, bold=True)

    heading1 = doc.add_heading(f'โครงการ: {project_title}', level=1)
    for run in heading1.runs:
        set_thai_font(run, size_pt=18, bold=True)

    add_thai_paragraph(doc, f'วันที่ออกแบบ: {datetime.now().strftime("%d/%m/%Y %H:%M")}', size_pt=15)

    # ========================================
    # SECTION 1: Design Method
    # ========================================
    heading2 = doc.add_heading('1. วิธีการออกแบบ', level=2)
    for run in heading2.runs:
        set_thai_font(run, size_pt=16, bold=True)

    add_thai_paragraph(doc,
        'การออกแบบโครงสร้างถนนใช้วิธี AASHTO 1993 Guide for Design of Pavement Structures '
        'ตามมาตรฐานกรมทางหลวง โดยใช้สมการหลักดังนี้:', size_pt=15)

    add_equation_paragraph(doc,
        'log₁₀(W₁₈) = Zᵣ·Sₒ + 9.36·log₁₀(SN+1) - 0.20 + '
        'log₁₀(ΔPSI/2.7) / [0.4 + 1094/(SN+1)⁵·¹⁹] + 2.32·log₁₀(Mᵣ) - 8.07',
        size_pt=11, italic=True)

    # ========================================
    # SECTION 2: Input Parameters
    # ========================================
    heading2_2 = doc.add_heading('2. ข้อมูลนำเข้า (Design Inputs)', level=2)
    for run in heading2_2.runs:
        set_thai_font(run, size_pt=16, bold=True)

    input_table = doc.add_table(rows=1, cols=3)
    input_table.style = 'Table Grid'

    headers = ['พารามิเตอร์', 'ค่า', 'หน่วย']
    for i, header in enumerate(headers):
        cell = input_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_thai_font(run, size_pt=15, bold=True)

    input_data = [
        ('Design ESALs (W₁₈)', f'{inputs["W18"]:,.0f}', '18-kip ESAL'),
        ('Reliability (R)', f'{inputs["reliability"]}', '%'),
        ('Standard Normal Deviate (Zᵣ)', f'{inputs["Zr"]:.3f}', '-'),
        ('Overall Standard Deviation (Sₒ)', f'{inputs["So"]:.2f}', '-'),
        ('Initial Serviceability (P₀)', f'{inputs["P0"]:.1f}', '-'),
        ('Terminal Serviceability (Pₜ)', f'{inputs["Pt"]:.1f}', '-'),
        ('ΔPSI = P₀ - Pₜ', f'{inputs["delta_psi"]:.1f}', '-'),
        ('Subgrade CBR', f'{inputs.get("CBR", "-")}', '%'),
        ('Subgrade Mᵣ = 1500 × CBR', f'{inputs["Mr"]:,.0f}', 'psi'),
    ]

    for param, value, unit in input_data:
        row = input_table.add_row()
        row.cells[0].text = param
        row.cells[1].text = value
        row.cells[2].text = unit
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_thai_font(run, size_pt=15)

    # ========================================
    # SECTION 3: Material Properties (+ AC Sublayers)
    # ========================================
    heading2_3 = doc.add_heading('3. คุณสมบัติวัสดุชั้นทาง', level=2)
    for run in heading2_3.runs:
        set_thai_font(run, size_pt=16, bold=True)

    mat_table = doc.add_table(rows=1, cols=6)
    mat_table.style = 'Table Grid'

    mat_headers = ['ชั้น', 'วัสดุ', 'aᵢ', 'mᵢ', 'Mᵣ (psi)', 'E (MPa)']
    for i, header in enumerate(mat_headers):
        cell = mat_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_thai_font(run, size_pt=15, bold=True)

    for layer in calc_results['layers']:
        row = mat_table.add_row()
        row.cells[0].text = str(layer['layer_no'])
        row.cells[1].text = layer['material']
        row.cells[2].text = f'{layer["a_i"]:.2f}'
        row.cells[3].text = f'{layer["m_i"]:.2f}'
        row.cells[4].text = f'{layer["mr_psi"]:,}'
        row.cells[5].text = f'{layer["mr_mpa"]:,}'
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_thai_font(run, size_pt=15)

    # AC Sublayer Breakdown (ถ้ามี)
    ac_sub = calc_results.get('ac_sublayers', None)
    if ac_sub is not None:
        doc.add_paragraph()
        add_thai_paragraph(doc, 'รายละเอียดชั้นย่อยผิวทาง AC:', size_pt=15, bold=True)
        
        sub_table = doc.add_table(rows=1, cols=3)
        sub_table.style = 'Table Grid'
        sub_headers = ['ชั้นย่อย', 'ความหนา (cm)', 'ความหนา (mm)']
        for i, header in enumerate(sub_headers):
            cell = sub_table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_thai_font(run, size_pt=15, bold=True)
        
        sub_data = [
            ('Wearing Course', ac_sub['wearing']),
            ('Binder Course', ac_sub['binder']),
            ('Base Course', ac_sub['base']),
            ('รวม', ac_sub['total']),
        ]
        for name, thick_cm in sub_data:
            row = sub_table.add_row()
            row.cells[0].text = name
            row.cells[1].text = f'{thick_cm:.1f}'
            row.cells[2].text = f'{thick_cm * 10:.0f}'
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_thai_font(run, size_pt=15)

    # ========================================
    # SECTION 4: Step-by-Step Calculation
    # ========================================
    heading2_4 = doc.add_heading('4. ขั้นตอนการคำนวณความหนาชั้นทาง', level=2)
    for run in heading2_4.runs:
        set_thai_font(run, size_pt=16, bold=True)

    add_thai_paragraph(doc,
        'การคำนวณความหนาขั้นต่ำของแต่ละชั้น ใช้หลักการว่า Structural Number (SN) '
        'ที่จุดใดๆ ต้องมากกว่าหรือเท่ากับ SN ที่ต้องการ โดยคำนวณจากค่า Mᵣ ของชั้นถัดไป',
        size_pt=15)

    for layer in calc_results['layers']:
        layer_heading = doc.add_heading(f'ชั้นที่ {layer["layer_no"]}: {layer["material"]}', level=3)
        for run in layer_heading.runs:
            set_thai_font(run, size_pt=15, bold=True)

        # Material properties
        add_thai_paragraph(doc, 'ข้อมูลวัสดุ:', size_pt=15, bold=True)
        props_para = doc.add_paragraph()
        run1 = props_para.add_run(f'    • Mᵣ = {layer["mr_psi"]:,} psi = {layer["mr_mpa"]:,} MPa\n')
        set_thai_font(run1, size_pt=15)
        run2 = props_para.add_run(f'    • Layer Coefficient (a{layer["layer_no"]}) = {layer["a_i"]:.2f}\n')
        set_thai_font(run2, size_pt=15)
        run3 = props_para.add_run(f'    • Drainage Coefficient (m{layer["layer_no"]}) = {layer["m_i"]:.2f}')
        set_thai_font(run3, size_pt=15)

        sn_at_layer = layer['sn_required_at_layer']
        
        # SN calculation
        add_thai_paragraph(doc, 'การคำนวณ SN:', size_pt=15, bold=True)
        sn_para = doc.add_paragraph()
        sn_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sn_run = sn_para.add_run(f'จากสมการ AASHTO 1993:  SN{layer["layer_no"]} = {sn_at_layer:.2f}')
        set_equation_font(sn_run, size_pt=11, bold=True, italic=False)

        # Thickness calculation
        add_thai_paragraph(doc, 'การคำนวณความหนาขั้นต่ำ:', size_pt=15, bold=True)

        if layer['layer_no'] == 1:
            formula_text = (f'D₁ ≥ SN₁ / (a₁ × m₁) = {sn_at_layer:.2f} / '
                          f'({layer["a_i"]:.2f} × {layer["m_i"]:.2f})')
            add_equation_paragraph(doc, formula_text, size_pt=11, italic=True)
        else:
            prev_sn = calc_results['layers'][layer['layer_no']-2]['cumulative_sn']
            formula_text = (f'D{layer["layer_no"]} ≥ (SN{layer["layer_no"]} - SNₚᵣₑᵥ) / '
                          f'(a{layer["layer_no"]} × m{layer["layer_no"]}) = '
                          f'({sn_at_layer:.2f} - {prev_sn:.2f}) / '
                          f'({layer["a_i"]:.2f} × {layer["m_i"]:.2f})')
            add_equation_paragraph(doc, formula_text, size_pt=11, italic=True)

        # Results
        result_para = doc.add_paragraph()
        result_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        result_run = result_para.add_run(
            f'D{layer["layer_no"]}(min) = {layer["min_thickness_inch"]:.2f} นิ้ว = {layer["min_thickness_cm"]:.1f} ซม.')
        set_equation_font(result_run, size_pt=11, bold=True, italic=False)

        # Design thickness
        add_thai_paragraph(doc, 'เลือกใช้ความหนา:', size_pt=15, bold=True)
        design_para = doc.add_paragraph()
        design_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        design_run = design_para.add_run(
            f'D{layer["layer_no"]}(design) = {layer["design_thickness_cm"]:.0f} ซม. '
            f'({layer["design_thickness_inch"]:.2f} นิ้ว)')
        set_equation_font(design_run, size_pt=11, bold=True, italic=False)

        # SN contribution
        add_thai_paragraph(doc, 'SN contribution:', size_pt=15, bold=True)
        contrib_text = (f'ΔSN{layer["layer_no"]} = a{layer["layer_no"]} × D{layer["layer_no"]} × '
                       f'm{layer["layer_no"]} = {layer["a_i"]:.2f} × {layer["design_thickness_inch"]:.2f} × '
                       f'{layer["m_i"]:.2f} = {layer["sn_contribution"]:.3f}')
        add_equation_paragraph(doc, contrib_text, size_pt=11, italic=False)

        # Cumulative SN
        cum_para = doc.add_paragraph()
        cum_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cum_run = cum_para.add_run(f'ΣSN = {layer["cumulative_sn"]:.2f}')
        set_equation_font(cum_run, size_pt=11, bold=True, italic=False)

        # Status
        status_text = '✓ OK' if layer['is_ok'] else '✗ NG - ต้องเพิ่มความหนา'
        status_para = doc.add_paragraph()
        status_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        status_run = status_para.add_run(f'สถานะ: {status_text}')
        set_thai_font(status_run, size_pt=15, bold=True)
        doc.add_paragraph()

    # ========================================
    # SECTION 5: SN Summary Table
    # ========================================
    heading2_5 = doc.add_heading('5. ตารางสรุปการคำนวณ Structural Number', level=2)
    for run in heading2_5.runs:
        set_thai_font(run, size_pt=16, bold=True)

    sn_table = doc.add_table(rows=1, cols=8)
    sn_table.style = 'Table Grid'

    sn_headers = ['ชั้น', 'วัสดุ', 'aᵢ', 'mᵢ', 'Dᵢ (นิ้ว)', 'Dᵢ (ซม.)', 'ΔSNᵢ', 'ΣSN']
    for i, header in enumerate(sn_headers):
        cell = sn_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_thai_font(run, size_pt=15, bold=True)

    for layer in calc_results['layers']:
        row = sn_table.add_row()
        row.cells[0].text = str(layer['layer_no'])
        row.cells[1].text = layer['material']
        row.cells[2].text = f'{layer["a_i"]:.2f}'
        row.cells[3].text = f'{layer["m_i"]:.2f}'
        row.cells[4].text = f'{layer["design_thickness_inch"]:.2f}'
        row.cells[5].text = f'{layer["design_thickness_cm"]:.0f}'
        row.cells[6].text = f'{layer["sn_contribution"]:.3f}'
        row.cells[7].text = f'{layer["cumulative_sn"]:.2f}'
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_thai_font(run, size_pt=15)

    doc.add_paragraph()
    add_equation_paragraph(doc, 'สูตร: SN = Σ(aᵢ × Dᵢ × mᵢ)', size_pt=11, italic=True)

    # ========================================
    # SECTION 6: Design Verification
    # ========================================
    heading2_6 = doc.add_heading('6. ผลการตรวจสอบการออกแบบ', level=2)
    for run in heading2_6.runs:
        set_thai_font(run, size_pt=16, bold=True)

    result_table = doc.add_table(rows=4, cols=2)
    result_table.style = 'Table Grid'

    result_data = [
        ('SN Required (จากสมการ AASHTO)', f'{calc_results["total_sn_required"]:.2f}'),
        ('SN Provided (จากชั้นทาง)', f'{calc_results["total_sn_provided"]:.2f}'),
        ('Safety Margin', f'{design_check["safety_margin"]:.2f}'),
        ('ผลการตรวจสอบ', 'ผ่าน (OK)' if design_check['passed'] else 'ไม่ผ่าน (NG)'),
    ]

    for i, (param, value) in enumerate(result_data):
        result_table.rows[i].cells[0].text = param
        result_table.rows[i].cells[1].text = value
        for cell in result_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_thai_font(run, size_pt=15)

    # W18 Supported
    doc.add_paragraph()
    w18_supported = calculate_w18_supported(
        calc_results['total_sn_provided'], 
        inputs['Zr'], inputs['So'], inputs['delta_psi'], inputs['Mr']
    )
    add_thai_paragraph(doc, f'W₁₈ ที่โครงสร้างรองรับได้ = {w18_supported/1e6:,.2f} ล้าน ESALs', 
                       size_pt=15, bold=True)
    add_thai_paragraph(doc, f'W₁₈ ที่ออกแบบ = {inputs["W18"]/1e6:,.2f} ล้าน ESALs', size_pt=15)

    # Conclusion
    doc.add_paragraph()
    if design_check['passed']:
        conclusion_text = (f'สรุป: การออกแบบผ่านเกณฑ์ เนื่องจาก SN_provided ({calc_results["total_sn_provided"]:.2f}) ≥ '
            f'SN_required ({calc_results["total_sn_required"]:.2f})')
        add_thai_paragraph(doc, conclusion_text, size_pt=15, bold=True)
    else:
        add_thai_paragraph(doc, 'สรุป: การออกแบบไม่ผ่านเกณฑ์ กรุณาปรับเพิ่มความหนาชั้นทาง', 
                          size_pt=15, bold=True)

    # ========================================
    # SECTION 7: Figure
    # ========================================
    heading2_7 = doc.add_heading('7. ภาพตัดขวางโครงสร้างถนน', level=2)
    for run in heading2_7.runs:
        set_thai_font(run, size_pt=16, bold=True)

    fig_bytes = get_figure_as_bytes(fig)
    doc.add_picture(fig_bytes, width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ========================================
    # SECTION 8: สรุปโครงสร้างชั้นทางที่ออกแบบ
    # ========================================
    heading2_8 = doc.add_heading('8. สรุปโครงสร้างชั้นทางที่ออกแบบด้วยวิธี AASHTO 1993', level=2)
    for run in heading2_8.runs:
        set_thai_font(run, size_pt=16, bold=True)

    # --- สร้างรายการชั้นทาง ---
    structure_rows = []  # [(ลำดับ, ชนิดวัสดุ, ความหนา ซม.)]
    row_num = 1

    # ชั้นที่ 1: AC — แยกชั้นย่อย (ถ้ามี)
    ac_sub = calc_results.get('ac_sublayers', None)
    first_layer = calc_results['layers'][0] if calc_results['layers'] else None

    if ac_sub is not None and first_layer:
        if ac_sub.get('wearing', 0) > 0:
            structure_rows.append((row_num, 'Wearing Course', f"{ac_sub['wearing']:.0f}"))
            row_num += 1
        if ac_sub.get('binder', 0) > 0:
            structure_rows.append((row_num, 'Binder Course', f"{ac_sub['binder']:.0f}"))
            row_num += 1
        if ac_sub.get('base', 0) > 0:
            structure_rows.append((row_num, 'Base Course', f"{ac_sub['base']:.0f}"))
            row_num += 1
        # ชั้นที่ 2 เป็นต้นไป
        for layer in calc_results['layers'][1:]:
            mat_name = layer['material']
            # แปลงชื่อวัสดุให้สั้นลงตามรูปแบบ
            short = mat_name
            short = short.replace('พื้นทางหินคลุกผสมซีเมนต์ UCS 24.5 ksc.', 'หินคลุกผสมซีเมนต์ UCS ≥ 24.5 ksc')
            short = short.replace('พื้นทางหินคลุก CBR 80%', 'หินคลุก CBR ≥ 80%')
            short = short.replace('พื้นทางซีเมนต์ CTB', 'ซีเมนต์ CTB')
            short = short.replace('พื้นทางดินซีเมนต์ UCS 17.5 ksc.', 'ดินซีเมนต์ UCS ≥ 17.5 ksc')
            short = short.replace('พื้นทางวัสดุหมุนเวียน (Recycling)', 'วัสดุหมุนเวียน (Recycling)')
            short = short.replace('รองพื้นทางวัสดุมวลรวม CBR 25%', 'รองพื้นทางวัสดุมวลรวม CBR ≥ 25%')
            short = short.replace('วัสดุคัดเลือก ก', 'วัสดุคัดเลือก ก')
            structure_rows.append((row_num, short, f"{layer['design_thickness_cm']:.0f}"))
            row_num += 1
    else:
        # ไม่มี sublayer — แสดงทุกชั้นปกติ
        for layer in calc_results['layers']:
            mat_name = layer['material']
            short = mat_name
            short = short.replace('ผิวทางลาดยาง AC', 'ผิวทางลาดยาง AC')
            short = short.replace('ผิวทางลาดยาง PMA', 'ผิวทางลาดยาง PMA')
            short = short.replace('พื้นทางหินคลุกผสมซีเมนต์ UCS 24.5 ksc.', 'หินคลุกผสมซีเมนต์ UCS ≥ 24.5 ksc')
            short = short.replace('พื้นทางหินคลุก CBR 80%', 'หินคลุก CBR ≥ 80%')
            short = short.replace('พื้นทางซีเมนต์ CTB', 'ซีเมนต์ CTB')
            short = short.replace('พื้นทางดินซีเมนต์ UCS 17.5 ksc.', 'ดินซีเมนต์ UCS ≥ 17.5 ksc')
            short = short.replace('พื้นทางวัสดุหมุนเวียน (Recycling)', 'วัสดุหมุนเวียน (Recycling)')
            short = short.replace('รองพื้นทางวัสดุมวลรวม CBR 25%', 'รองพื้นทางวัสดุมวลรวม CBR ≥ 25%')
            short = short.replace('วัสดุคัดเลือก ก', 'วัสดุคัดเลือก ก')
            structure_rows.append((row_num, short, f"{layer['design_thickness_cm']:.0f}"))
            row_num += 1

    # เพิ่มแถวดินคันทาง
    cbr_val = inputs.get('CBR', 3.0)
    structure_rows.append((row_num, 'ดินคันทาง', f'CBR ≥ {cbr_val:.1f} %'))

    # --- หัวข้อย่อย: ชื่อชั้นผิวทาง ---
    surface_mat_name = calc_results['layers'][0]['material'] if calc_results['layers'] else 'ผิวทางลาดยาง'
    sub_heading = doc.add_heading(f'รูปแบบที่: {surface_mat_name}', level=3)
    for run in sub_heading.runs:
        set_thai_font(run, size_pt=15, bold=True)

    # --- สร้างตาราง ---
    num_rows = 1 + len(structure_rows)  # header + data rows
    summary_table = doc.add_table(rows=num_rows, cols=3)
    summary_table.style = 'Table Grid'
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    for row in summary_table.rows:
        row.cells[0].width = Cm(2.0)
        row.cells[1].width = Cm(10.0)
        row.cells[2].width = Cm(4.0)

    # Header row
    header_texts = ['ลำดับ', 'ชนิดวัสดุ', 'ความหนา (ซม.)']
    for j, text in enumerate(header_texts):
        cell = summary_table.rows[0].cells[j]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        set_thai_font(run, size_pt=15, bold=True)
        # Header shading (สีฟ้าอ่อน)
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), 'BDD7EE')  # สีฟ้าอ่อน ตามมาตรฐาน
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)  # สีดำ
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_pr.append(shading)

    # Data rows
    for i, (num, mat_name, thickness) in enumerate(structure_rows):
        row_idx = i + 1
        row = summary_table.rows[row_idx]

        # ลำดับ
        cell0 = row.cells[0]
        cell0.text = ''
        p0 = cell0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(str(num))
        set_thai_font(r0, size_pt=15)

        # ชนิดวัสดุ
        cell1 = row.cells[1]
        cell1.text = ''
        p1 = cell1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run(mat_name)
        set_thai_font(r1, size_pt=15)

        # ความหนา
        cell2 = row.cells[2]
        cell2.text = ''
        p2 = cell2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(thickness)
        set_thai_font(r2, size_pt=15)

    # --- รูปตัดโครงสร้างชั้นทาง ---
    doc.add_paragraph()
    add_thai_paragraph(doc, 'รูปตัดโครงสร้างชั้นทาง', size_pt=15, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)

    fig_bytes_section8 = get_figure_as_bytes(fig)
    doc.add_picture(fig_bytes_section8, width=Inches(5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ========================================
    # Footer
    # ========================================
    doc.add_paragraph()
    footer_para = add_thai_paragraph(doc, 
        'รายงานนี้สร้างโดยแอปพลิเคชัน AASHTO 1993 Flexible Pavement Design v5.0\n'
        'พัฒนาโดย รศ.ดร.อิทธิพล มีผล // ภาควิชาครุศาสตร์โยธา // มจพ.',
        size_pt=12)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes


# ================================================================================
# WORD REPORT WITH INTRO SECTION (สำหรับรวมกับรายงานอื่น)
# ================================================================================

def set_thai_distribute(para):
    """ตั้ง Thai Distributed alignment ผ่าน XML"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = para._element.get_or_add_pPr()
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'thaiDistribute')
    pPr.append(jc)


def add_table_header_shading(cell, fill_hex='D9E2F3'):
    """เพิ่มพื้นหลังสีให้ cell header"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), fill_hex)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(shading)


def create_word_report_intro(project_title, inputs, calc_results, design_check, fig, report_settings):
    """
    สร้างรายงาน Word รูปแบบสำหรับรวมกับรายงานอื่น
    โครงสร้างครบ:
      {sec_no}      หัวข้อ + เกริ่นนำ
      {sec_no}.1    วิธีการออกแบบ
      {sec_no}.2    ข้อมูลนำเข้า (Design Inputs) + ตาราง
      {sec_no}.3    คุณสมบัติวัสดุชั้นทาง + ตาราง
      {sec_no}.4    ขั้นตอนการคำนวณความหนาชั้นทาง + รูปตัดขวาง
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()

    # ตั้ง Normal style
    style = doc.styles['Normal']
    style.font.name = 'TH SarabunPSK'
    style.font.size = Pt(15)
    try:
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')
    except Exception:
        pass

    # ------------------------------------------------------------------
    # Helper functions (inline)
    # ------------------------------------------------------------------
    def _run(para, text, size=15, bold=False, italic=False, color=None, underline=False):
        r = para.add_run(text)
        r.font.name = 'TH SarabunPSK'
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        r.underline = underline
        if color:
            r.font.color.rgb = color
        try:
            r._element.rPr.rFonts.set(qn('w:cs'), 'TH SarabunPSK')
        except Exception:
            pass
        return r

    def _heading_para(text, size=15, bold=True, underline=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(4)
        _run(p, text, size=size, bold=bold, underline=underline)
        return p

    def _body_para(parts, indent_cm=1.25):
        """parts = list of (text, bold)"""
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(indent_cm)
        p.paragraph_format.space_after = Pt(4)
        set_thai_distribute(p)
        for text, bold in parts:
            _run(p, text, bold=bold)
        return p

    def _table_caption(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        _run(p, text, bold=True)

    def _fig_caption(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        _run(p, text, size=14, bold=True)

    def _cell_run(cell, text, size=14, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.CENTER):
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = align
        _run(p, text, size=size, bold=bold, color=color)

    def _make_table_header(tbl, headers, widths_cm=None, fill='D9E2F3'):
        row = tbl.rows[0]
        for i, hdr in enumerate(headers):
            cell = row.cells[i]
            if widths_cm:
                cell.width = Cm(widths_cm[i])
            _cell_run(cell, hdr, bold=True)
            add_table_header_shading(cell, fill)

    # ------------------------------------------------------------------
    # รับค่า report settings
    # ------------------------------------------------------------------
    sec_no  = report_settings.get('section_number', '4.4')
    tbl_no1 = report_settings.get('table_number_inputs', '4-8')
    tbl_no2 = report_settings.get('table_number_materials', '4-9')
    fig_no  = report_settings.get('figure_number', '4-8')
    sec_title   = report_settings.get('section_title', 'การออกแบบผิวทางลาดยาง (Flexible Pavement)')
    tbl_cap1    = report_settings.get('table_caption_inputs',    'ค่าพารามิเตอร์ที่ใช้ในการออกแบบผิวทางยืดหยุ่น')
    tbl_cap2    = report_settings.get('table_caption_materials', 'ค่าสัมประสิทธิ์และค่าโมดูลัสของวัสดุโครงสร้างชั้นทาง')
    fig_cap     = report_settings.get('figure_caption', 'รูปตัดโครงสร้างชั้นทางที่ออกแบบ')

    # ค่าคำนวณ
    W18_val     = inputs.get('W18', 0)
    reliability = inputs.get('reliability', 90)
    CBR_val     = inputs.get('CBR', 3.0)
    Mr_val      = inputs.get('Mr', 4500)
    Zr_val      = inputs.get('Zr', -1.282)
    So_val      = inputs.get('So', 0.45)
    P0_val      = inputs.get('P0', 4.2)
    Pt_val      = inputs.get('Pt', 2.5)
    dpsi_val    = inputs.get('delta_psi', 1.7)
    sn_req      = calc_results.get('total_sn_required', 0)
    sn_prov     = calc_results.get('total_sn_provided', 0)
    total_thick = sum(l['design_thickness_cm'] for l in calc_results.get('layers', []))
    num_layers  = len(calc_results.get('layers', []))
    passed_txt  = 'ผ่านเกณฑ์' if design_check.get('passed') else 'ไม่ผ่านเกณฑ์'
    RED = RGBColor(255, 0, 0)

    # ==================================================================
    # 4.4  หัวข้อหลัก + บทเกริ่นนำ
    # ==================================================================
    _heading_para(f'{sec_no}\t{sec_title}', size=16, bold=True)

    _body_para([
        ('        ถนนลาดยางซึ่งประกอบด้วยวัสดุงานทางหลายชนิด เนื่องจาก หน่วยแรงจากน้ำหนักจราจร'
         'จะมีความเข้มข้นสูงสุดบนผิวทาง แอสฟัลต์คอนกรีตจึงนำมาใช้เป็นวัสดุ ผิวทาง '
         'และใช้วัสดุที่มีคุณภาพด้อยลงมา ได้แก่ วัสดุท้องถิ่น (Local Materials) '
         'หรือวัสดุที่มีราคาถูก ในระดับลึกลงไป โดยวางซ้อนกันเป็นชั้น ๆ อย่างเป็นระบบ '
         '(Multi-layer System) เหนือดินฐานราก (Subgrade)', False),
    ], indent_cm=1.25)

    # ==================================================================
    # 4.4.1  วิธีการออกแบบ
    # ==================================================================
    _heading_para(f'{sec_no}.1\tวิธีการออกแบบ', size=15, bold=True)

    _body_para([
        ('        การออกแบบโครงสร้างถนนแบบยืดหยุ่น (Flexible Pavement) ใช้วิธี ', False),
        ('AASHTO 1993 Guide for Design of Pavement Structures', True),
        (' โดยใช้สมการหลักดังนี้', False),
    ], indent_cm=1.25)

    # สมการ AASHTO — ใช้ Times New Roman
    eq_para = doc.add_paragraph()
    eq_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq_para.paragraph_format.space_before = Pt(4)
    eq_para.paragraph_format.space_after  = Pt(4)
    eq_run = eq_para.add_run(
        'log\u2081\u2080(W\u2081\u2088) = Z\u1d63\u00b7S\u2092 + 9.36\u00b7log\u2081\u2080(SN+1) \u2212 0.20 + '
        'log\u2081\u2080(\u0394PSI/2.7) / [0.4 + 1094/(SN+1)\u2075\u00b7\u00b9\u2079] + '
        '2.32\u00b7log\u2081\u2080(M\u1d63) \u2212 8.07'
    )
    eq_run.font.name = 'Times New Roman'
    eq_run.font.size = Pt(12)
    eq_run.italic = True

    # ==================================================================
    # 4.4.2  ข้อมูลนำเข้า (Design Inputs)
    # ==================================================================
    _heading_para(f'{sec_no}.2\tข้อมูลนำเข้า (Design Inputs)', size=15, bold=True)

    _body_para([
        ('        ในการออกแบบโครงสร้างถนนยืดหยุ่น การกำหนดค่าพารามิเตอร์นำเข้า (Design Inputs) '
         'ถือเป็นขั้นตอนสำคัญที่มีผลโดยตรงต่อความถูกต้องและความน่าเชื่อถือของแบบโครงสร้างถนนที่ต้องการ '
         'เนื่องจากค่าพารามิเตอร์แต่ละตัวสะท้อนให้เห็นสภาพการใช้งานจริงของโครงสร้างถนน '
         'ปริมาณการจราจรตลอดอายุการใช้งาน ระดับความน่าเชื่อถือที่ยอมรับได้ '
         'รวมถึงคุณสมบัติของวัสดุและชั้นดินรองรับในพื้นที่โครงการ '
         'สำหรับโครงการนี้ ที่ปรึกษาได้กำหนดค่าพารามิเตอร์หลักที่ใช้ในการออกแบบตามแนวทางของ AASHTO '
         'ซึ่งประกอบด้วยข้อมูลด้านความสามารถในการรับน้ำหนักของโครงสร้างชั้นทาง ปริมาณจราจรที่โครงสร้าง'
         'ถนนต้องรองรับตลอดอายุการใช้งาน รวมถึงคุณสมบัติของชั้นดินที่ต้องซ่อมบำรุงหรือปรับปรุงใหม่ '
         'รวมถึงคุณสมบัติของดินชั้นรองรับ รายละเอียดของค่าพารามิเตอร์ทั้งหมดแสดงในตารางที่ ', False),
        (f'{tbl_no1}', True),
    ], indent_cm=1.25)

    _table_caption(f'ตารางที่ {tbl_no1}  {tbl_cap1}')

    t1 = doc.add_table(rows=1, cols=3)
    t1.style = 'Table Grid'
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    _make_table_header(t1, ['พารามิเตอร์', 'ค่า', 'หน่วย'], widths_cm=[9, 4, 3])

    input_data = [
        ('Design ESALs (W\u2081\u2088)',               f'{W18_val:,.0f}',        '18-kip ESAL'),
        ('Reliability (R)',                             f'{reliability}',         '%'),
        ('Standard Normal Deviate (Z\u1d63)',           f'{Zr_val:.3f}',          '-'),
        ('Overall Standard Deviation (S\u2092)',        f'{So_val:.2f}',          '-'),
        ('Initial Serviceability (P\u2080)',            f'{P0_val:.1f}',          '-'),
        ('Terminal Serviceability (P\u209c)',           f'{Pt_val:.1f}',          '-'),
        ('\u0394PSI = P\u2080 \u2212 P\u209c',         f'{dpsi_val:.1f}',         '-'),
        ('Subgrade CBR',                                f'{CBR_val}',             '%'),
        ('Subgrade M\u1d63 = 1500 \u00d7 CBR',         f'{Mr_val:,.0f}',         'psi'),
    ]

    for param, value, unit in input_data:
        row = t1.add_row()
        _cell_run(row.cells[0], param, align=WD_ALIGN_PARAGRAPH.LEFT)
        _cell_run(row.cells[1], value, bold=True, color=RED)
        _cell_run(row.cells[2], unit)

    # ==================================================================
    # 4.4.3  คุณสมบัติวัสดุชั้นทาง
    # ==================================================================
    doc.add_paragraph()
    _heading_para(f'{sec_no}.3\tคุณสมบัติวัสดุชั้นทาง', size=15, bold=True)

    _body_para([
        ('        วัสดุโครงสร้างชั้นทางแต่ละชนิดมีค่าสัมประสิทธิ์ชั้นทาง (Layer Coefficient) '
         'และค่าสัมประสิทธิ์การระบายน้ำ (Drained Coefficient) โดยที่ปรึกษาเลือกใช้วัสดุ'
         'และแสดงค่าสัมประสิทธิ์รวมถึงค่าโมดูลัสของวัสดุต่างๆ ดังแสดงในตารางที่ ', False),
        (f'{tbl_no2}', True),
    ], indent_cm=1.25)

    _table_caption(f'ตารางที่ {tbl_no2}  {tbl_cap2}')

    # ตาราง 4 คอลัมน์ตามภาพตัวอย่าง: ชั้น | วัสดุ | ai | mi | Mr(psi) | E(MPa)
    t2 = doc.add_table(rows=1, cols=6)
    t2.style = 'Table Grid'
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    _make_table_header(t2, ['ชั้น', 'วัสดุ', 'a\u1d62', 'm\u1d62', 'M\u1d63 (psi)', 'E (MPa)'],
                       widths_cm=[1.5, 7, 1.5, 1.5, 2.5, 2])

    for layer in calc_results.get('layers', []):
        row = t2.add_row()
        _cell_run(row.cells[0], str(layer['layer_no']))
        _cell_run(row.cells[1], layer['material'], align=WD_ALIGN_PARAGRAPH.LEFT)
        _cell_run(row.cells[2], f'{layer["a_i"]:.2f}')
        _cell_run(row.cells[3], f'{layer["m_i"]:.2f}')
        _cell_run(row.cells[4], f'{layer["mr_psi"]:,}')
        _cell_run(row.cells[5], f'{layer["mr_mpa"]:,}')

    # ==================================================================
    # 4.4.4  ขั้นตอนการคำนวณความหนาชั้นทาง
    # ==================================================================
    doc.add_paragraph()
    _heading_para(f'{sec_no}.4\tขั้นตอนการคำนวณความหนาชั้นทาง', size=15, bold=True)

    _body_para([
        ('        การคำนวณความหนาขั้นต่ำของแต่ละชั้น ใช้หลักการว่า Structural Number (SN) '
         'ที่จุดใดๆ ต้องมากกว่าหรือเท่ากับ SN ที่ต้องการ โดยคำนวณจากค่า M\u1d63 ของชั้นถัดไป', False),
    ], indent_cm=1.25)

    # --- ชั้นทีละชั้น ---
    for layer in calc_results.get('layers', []):
        sn_at = layer['sn_required_at_layer']
        layer_no = layer['layer_no']

        # หัวข้อชั้น
        sub_p = doc.add_paragraph()
        sub_p.paragraph_format.space_before = Pt(6)
        sub_p.paragraph_format.space_after  = Pt(2)
        sub_p.paragraph_format.left_indent  = Cm(1.0)
        _run(sub_p, f'ชั้นที่ {layer_no}: {layer["material"]}', bold=True, underline=True)

        # ข้อมูลวัสดุ
        mat_p = doc.add_paragraph()
        mat_p.paragraph_format.left_indent = Cm(2.0)
        mat_p.paragraph_format.space_after = Pt(2)
        _run(mat_p,
             f'M\u1d63 = {layer["mr_psi"]:,} psi = {layer["mr_mpa"]:,} MPa   '
             f'| a{layer_no} = {layer["a_i"]:.2f}   '
             f'| m{layer_no} = {layer["m_i"]:.2f}')

        # SN required
        sn_p = doc.add_paragraph()
        sn_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sn_run = sn_p.add_run(f'SN{layer_no} = {sn_at:.2f}   (จากสมการ AASHTO 1993)')
        sn_run.font.name = 'Times New Roman'
        sn_run.font.size = Pt(11)
        sn_run.bold = True

        # สูตรความหนา
        formula_p = doc.add_paragraph()
        formula_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if layer_no == 1:
            formula_txt = (f'D\u2081 \u2265 SN\u2081 / (a\u2081 \u00d7 m\u2081) = '
                           f'{sn_at:.2f} / ({layer["a_i"]:.2f} \u00d7 {layer["m_i"]:.2f}) = '
                           f'{layer["min_thickness_inch"]:.2f} \u0e19\u0e34\u0e49\u0e27 = '
                           f'{layer["min_thickness_cm"]:.1f} \u0e0b\u0e21.')
        else:
            prev_sn = calc_results['layers'][layer_no - 2]['cumulative_sn']
            formula_txt = (f'D{layer_no} \u2265 (SN{layer_no} \u2212 SN{layer_no-1}) / '
                           f'(a{layer_no} \u00d7 m{layer_no}) = '
                           f'({sn_at:.2f} \u2212 {prev_sn:.2f}) / ({layer["a_i"]:.2f} \u00d7 {layer["m_i"]:.2f}) = '
                           f'{layer["min_thickness_inch"]:.2f} \u0e19\u0e34\u0e49\u0e27 = '
                           f'{layer["min_thickness_cm"]:.1f} \u0e0b\u0e21.')
        f_run = formula_p.add_run(formula_txt)
        f_run.font.name = 'Times New Roman'
        f_run.font.size = Pt(11)
        f_run.italic = True

        # ความหนาที่เลือก + SN contribution
        res_p = doc.add_paragraph()
        res_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        res_p.paragraph_format.space_after = Pt(2)
        status_sym = '\u2713 OK' if layer['is_ok'] else '\u2717 NG'
        res_run = res_p.add_run(
            f'D{layer_no}(design) = {layer["design_thickness_cm"]:.0f} \u0e0b\u0e21.   '
            f'| \u0394SN{layer_no} = {layer["sn_contribution"]:.3f}   '
            f'| \u03a3SN = {layer["cumulative_sn"]:.2f}   '
            f'| {status_sym}'
        )
        res_run.font.name = 'Times New Roman'
        res_run.font.size = Pt(11)
        res_run.bold = True
        res_run.font.color.rgb = RGBColor(0, 112, 0) if layer['is_ok'] else RED

    # ------------------------------------------------------------------
    # สรุป SN รวม
    # ------------------------------------------------------------------
    doc.add_paragraph()
    sum_p = doc.add_paragraph()
    sum_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sum_p.paragraph_format.space_before = Pt(6)
    _run(sum_p,
         f'SN_required = {sn_req:.2f}   |   SN_provided = {sn_prov:.2f}   |   '
         f'Safety Margin = {design_check.get("safety_margin", sn_prov - sn_req):.2f}   |   '
         f'ผลการออกแบบ: {passed_txt}',
         bold=True, size=15,
         color=RGBColor(0, 112, 0) if design_check.get('passed') else RED)

    # ------------------------------------------------------------------
    # ตารางสรุปโครงสร้างชั้นทาง (รูปแบบ Section 8)
    # ------------------------------------------------------------------
    doc.add_paragraph()

    # หัวข้อสรุป
    surf_name = calc_results['layers'][0]['material'] if calc_results.get('layers') else 'ผิวทางลาดยาง'
    surf_p = doc.add_paragraph()
    surf_p.paragraph_format.space_before = Pt(6)
    surf_p.paragraph_format.space_after  = Pt(4)
    _run(surf_p, f'รูปแบบที่: {surf_name}', bold=True)

    # สร้างรายการชั้นทาง
    structure_rows = []
    row_num = 1
    ac_sub = calc_results.get('ac_sublayers', None)
    first_layer = calc_results['layers'][0] if calc_results.get('layers') else None

    def _short_name(mat_name):
        return (mat_name
            .replace('พื้นทางหินคลุกผสมซีเมนต์ UCS 24.5 ksc.', 'หินคลุกผสมซีเมนต์ UCS \u2265 24.5 ksc')
            .replace('พื้นทางหินคลุก CBR 80%',                  'หินคลุก CBR \u2265 80%')
            .replace('พื้นทางซีเมนต์ CTB',                      'ซีเมนต์ CTB')
            .replace('พื้นทางดินซีเมนต์ UCS 17.5 ksc.',         'ดินซีเมนต์ UCS \u2265 17.5 ksc')
            .replace('พื้นทางวัสดุหมุนเวียน (Recycling)',       'วัสดุหมุนเวียน (Recycling)')
            .replace('รองพื้นทางวัสดุมวลรวม CBR 25%',           'รองพื้นทางวัสดุมวลรวม CBR \u2265 25%')
        )

    if ac_sub is not None and first_layer:
        for key, label in [('wearing', 'Wearing Course'), ('binder', 'Binder Course'), ('base', 'Base Course')]:
            if ac_sub.get(key, 0) > 0:
                structure_rows.append((row_num, label, f"{ac_sub[key]:.0f}"))
                row_num += 1
        for layer in calc_results['layers'][1:]:
            structure_rows.append((row_num, _short_name(layer['material']),
                                   f"{layer['design_thickness_cm']:.0f}"))
            row_num += 1
    else:
        for layer in calc_results.get('layers', []):
            structure_rows.append((row_num, _short_name(layer['material']),
                                   f"{layer['design_thickness_cm']:.0f}"))
            row_num += 1

    # แถวดินคันทาง
    cbr_val = inputs.get('CBR', 3.0)
    structure_rows.append((row_num, 'ดินคันทาง', f'CBR \u2265 {cbr_val:.1f} %'))

    # สร้างตาราง 3 คอลัมน์
    from docx.oxml import OxmlElement as _OxmlElement
    from docx.oxml.ns import qn as _qn

    num_rows_tbl = 1 + len(structure_rows)
    sum_tbl = doc.add_table(rows=num_rows_tbl, cols=3)
    sum_tbl.style = 'Table Grid'
    sum_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r in sum_tbl.rows:
        r.cells[0].width = Cm(2.0)
        r.cells[1].width = Cm(10.0)
        r.cells[2].width = Cm(4.0)

    # Header row (สีฟ้าอ่อน BDD7EE)
    for j, hdr_txt in enumerate(['ลำดับ', 'ชนิดวัสดุ', 'ความหนา (ซม.)']):
        cell = sum_tbl.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, hdr_txt, bold=True)
        shd = _OxmlElement('w:shd')
        shd.set(_qn('w:val'), 'clear')
        shd.set(_qn('w:color'), 'auto')
        shd.set(_qn('w:fill'), 'BDD7EE')
        cell._tc.get_or_add_tcPr().append(shd)

    # Data rows
    for i, (num, mat_name, thickness) in enumerate(structure_rows):
        row = sum_tbl.rows[i + 1]
        # ลำดับ
        row.cells[0].text = ''
        p0 = row.cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p0, str(num))
        # ชนิดวัสดุ
        row.cells[1].text = ''
        p1 = row.cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _run(p1, mat_name)
        # ความหนา
        row.cells[2].text = ''
        p2 = row.cells[2].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p2, thickness)

    # ------------------------------------------------------------------
    # รูปตัดขวาง + caption ใต้รูป
    # ------------------------------------------------------------------
    doc.add_paragraph()
    fig_bytes_intro = get_figure_as_bytes(fig)
    doc.add_picture(fig_bytes_intro, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _fig_caption(f'รูปที่ {fig_no}  {fig_cap}')

    # ==================================================================
    # Footer
    # ==================================================================
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(footer_p,
         'พัฒนาโดย รศ.ดร.อิทธิพล มีผล // ภาควิชาครุศาสตร์โยธา // มจพ.',
         size=12, italic=True)

    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes


# ================================================================================
# STREAMLIT USER INTERFACE - Tab Layout (V5)
# ================================================================================

def main():
    """Main Streamlit application"""

    # ========================================
    # HEADER
    # ========================================
    st.title("🛣️  Flexible Pavement Design (AASHTO 1993) v5")
    st.markdown("**แอปพลิเคชันออกแบบโครงสร้างทางแบบยืดหยุ่น ตามวิธีการ AASHTO (1993) — มาตรฐานกรมทางหลวง**")

    # ========================================
    # SIDEBAR: Project Info, Preset, JSON, Material DB
    # ========================================
    with st.sidebar:
        st.header("📋 ข้อมูลโครงการ")
        
        project_title = st.text_input(
            "ชื่อโครงการ",
            value=st.session_state.get('input_project_title', "โครงการออกแบบถนน"),
            key="project_title_input"
        )

        st.markdown("---")

        # ===== PRESET STRUCTURES =====
        st.header("🏗️ Preset โครงสร้าง ทล.")
        
        preset_names = list(PRESET_STRUCTURES.keys())
        selected_preset = st.selectbox(
            "เลือกโครงสร้างมาตรฐาน",
            options=preset_names,
            index=0,
            help="เลือกเพื่อ Auto-fill วัสดุและความหนาเริ่มต้น (แก้ไขเพิ่มเติมได้)"
        )

        if selected_preset != "--- เลือกโครงสร้างมาตรฐาน ---":
            preset = PRESET_STRUCTURES[selected_preset]
            if preset:
                st.info(f"📋 {preset['description']}")
                if st.button("✅ ใช้โครงสร้างนี้", type="primary"):
                    st.session_state['input_num_layers'] = preset['num_layers']
                    for i, layer in enumerate(preset['layers']):
                        st.session_state[f'layer{i+1}_mat'] = layer['material']
                        st.session_state[f'layer{i+1}_thick'] = layer['thickness_cm']
                        mat = MATERIALS[layer['material']]
                        st.session_state[f'layer{i+1}_a'] = mat['layer_coeff']
                        st.session_state[f'layer{i+1}_m'] = mat['drainage_coeff']
                    st.session_state['use_ac_sublayers'] = False
                    st.session_state['ac_sublayers'] = None
                    st.rerun()

        st.markdown("---")

        # ===== JSON UPLOAD/DOWNLOAD =====
        st.header("💾 บันทึก/โหลดข้อมูล")
        uploaded_json = st.file_uploader(
            "📂 โหลดข้อมูลจากไฟล์ JSON", type=['json'],
            help="อัปโหลดไฟล์ JSON ที่บันทึกไว้ก่อนหน้า"
        )

        if uploaded_json is not None:
            try:
                loaded_data = json.load(uploaded_json)
                file_id = f"{uploaded_json.name}_{uploaded_json.size}"
                if st.session_state.get('last_uploaded_file') != file_id:
                    st.session_state['last_uploaded_file'] = file_id
                    st.session_state['loaded_json'] = loaded_data
                    st.session_state['input_W18'] = loaded_data.get('W18', 5000000)
                    st.session_state['input_reliability'] = loaded_data.get('reliability', 90)
                    st.session_state['input_So'] = loaded_data.get('So', 0.45)
                    st.session_state['input_P0'] = loaded_data.get('P0', 4.2)
                    st.session_state['input_Pt'] = loaded_data.get('Pt', 2.5)
                    st.session_state['input_CBR'] = loaded_data.get('CBR', 5.0)
                    st.session_state['input_num_layers'] = loaded_data.get('num_layers', 4)
                    st.session_state['input_project_title'] = loaded_data.get('project_title', 'โครงการออกแบบถนน')
                    # Load report settings
                    rs = loaded_data.get('report_settings', {})
                    for key, default in [
                        ('section_number', '4.4'),
                        ('table_number_inputs', '4-8'),
                        ('table_number_materials', '4-9'),
                        ('figure_number', '4-8'),
                        ('section_title', 'การออกแบบผิวทางลาดยาง (Flexible Pavement)'),
                        ('table_caption_inputs', 'ค่าพารามิเตอร์ที่ใช้ในการออกแบบผิวทางยืดหยุ่น'),
                        ('table_caption_materials', 'ค่าสัมประสิทธิ์และค่าโมดูลัสของวัสดุโครงสร้างชั้นทาง'),
                        ('figure_caption', 'รูปตัดโครงสร้างชั้นทางที่ออกแบบ'),
                    ]:
                        if key in rs:
                            st.session_state[f'rs_{key}'] = rs[key]
                    layers = loaded_data.get('layers', [])
                    for i, layer in enumerate(layers):
                        st.session_state[f'layer{i+1}_mat'] = layer.get('material', '')
                        st.session_state[f'layer{i+1}_thick'] = layer.get('thickness_cm', 15.0)
                        st.session_state[f'layer{i+1}_m'] = layer.get('drainage_coeff', 1.0)
                    st.success("✅ โหลดข้อมูลสำเร็จ!")
                    st.rerun()
            except Exception as e:
                st.error(f"❌ ไม่สามารถอ่านไฟล์ได้: {e}")

        st.markdown("---")

        # ===== Figure Language =====
        st.header("🖼️ ตั้งค่ารูปภาพ")
        figure_language = st.radio(
            "ภาษาในรูปภาพ",
            options=["English", "ภาษาไทย"],
            index=0,
            help="เลือกภาษาสำหรับแสดงในรูปตัดขวาง"
        )

        st.markdown("---")

        # ===== Material Database =====
        st.header("📚 ฐานข้อมูลวัสดุ (ทล.)")
        with st.expander("ดูค่า สปส. วัสดุทั้งหมด"):
            for mat_name, props in MATERIALS.items():
                if props['layer_coeff'] > 0:
                    st.markdown(f"**{mat_name}**")
                    st.markdown(f"- a = {props['layer_coeff']}, m = {props['drainage_coeff']}")
                    st.markdown(f"- MR = {props['mr_psi']:,} psi ({props['mr_mpa']:,} MPa)")
                    st.markdown("---")

    # ========================================
    # MAIN CONTENT — TABS
    # ========================================
    tab_input, tab_layers, tab_results, tab_report = st.tabs([
        "📝 ข้อมูลนำเข้า", "🏗️ ชั้นทาง", "📊 ผลลัพธ์", "📄 รายงาน"
    ])

    # ========================================
    # TAB 1: DESIGN INPUTS
    # ========================================
    with tab_input:
        st.header("📝 Design Inputs")

        col_t1, col_t2 = st.columns(2)

        with col_t1:
            st.subheader("1️⃣ Traffic & Reliability")

            W18 = st.number_input(
                "Design ESALs (W₁₈)",
                min_value=100000, max_value=250000000,
                value=st.session_state.get('input_W18', 5000000),
                step=100000, format="%d",
                help="จำนวน 18-kip ESAL ตลอดอายุการใช้งาน (สูงสุด 250 ล้าน)",
                key="input_W18"
            )
            esal_million = W18 / 1000000
            st.markdown(
                f'<p style="color: #1E90FF; font-size: 20px; font-weight: bold;">'
                f'💡 W₁₈ = {esal_million:,.2f} ล้าน ESALs</p>',
                unsafe_allow_html=True)

            reliability_options = list(RELIABILITY_ZR.keys())
            current_reliability = st.session_state.get('input_reliability', 90)
            default_reliability_idx = (reliability_options.index(current_reliability) 
                                       if current_reliability in reliability_options 
                                       else reliability_options.index(90))

            reliability = st.selectbox(
                "Reliability Level (R)", options=reliability_options,
                index=default_reliability_idx, key="input_reliability"
            )
            Zr = RELIABILITY_ZR[reliability]
            st.info(f"Zᵣ = {Zr:.3f}")

            So = st.number_input(
                "Overall Standard Deviation (Sₒ)",
                min_value=0.30, max_value=0.60,
                value=st.session_state.get('input_So', 0.45),
                step=0.01, format="%.2f", key="input_So"
            )

        with col_t2:
            st.subheader("2️⃣ Serviceability")
            col_p1, col_p2 = st.columns(2)
            with col_p1:
                P0 = st.number_input("P₀ (Initial)", min_value=3.0, max_value=5.0,
                    value=st.session_state.get('input_P0', 4.2), step=0.1, key="input_P0")
            with col_p2:
                Pt = st.number_input("Pₜ (Terminal)", min_value=1.5, max_value=3.5,
                    value=st.session_state.get('input_Pt', 2.5), step=0.1, key="input_Pt")

            delta_psi = P0 - Pt
            st.success(f"**ΔPSI = {delta_psi:.1f}**")

            st.subheader("3️⃣ Subgrade (ดินเดิม/ดินถม)")
            CBR = st.number_input("CBR (%)", min_value=1.0, max_value=30.0,
                value=st.session_state.get('input_CBR', 5.0), step=0.5,
                help="ค่า CBR ของดินเดิมหรือดินถมคันทาง", key="input_CBR")
            Mr = int(1500 * CBR)
            st.info(f"**Mᵣ = 1,500 × CBR = 1,500 × {CBR:.1f} = {Mr:,} psi**")

        # ===== Drainage Coefficient Reference =====
        with st.expander("📖 ตาราง Drainage Coefficient (AASHTO Table 2.4)"):
            st.markdown("**ค่าสัมประสิทธิ์การระบายน้ำ (mᵢ) — AASHTO 1993 Table 2.4**")
            st.markdown("ค่า default กรมทางหลวง = **1.0** (สภาพการระบายน้ำดี)")
            
            drain_data = []
            for quality, info in DRAINAGE_TABLE.items():
                row = {"คุณภาพการระบายน้ำ": f"{quality} — {info['description']}"}
                for pct, val in info['values'].items():
                    row[f"เวลาอิ่มตัว {pct}"] = f"{val:.2f}"
                drain_data.append(row)
            st.table(drain_data)

    # ========================================
    # TAB 2: LAYER CONFIGURATION
    # ========================================
    with tab_layers:
        st.header("🏗️ Layer Configuration")

        num_layers = st.slider(
            "จำนวนชั้นทาง", min_value=2, max_value=6,
            value=st.session_state.get('input_num_layers', 4),
            help="เลือกจำนวนชั้นทาง (2-6 ชั้น)", key="input_num_layers"
        )

        all_materials = [m for m, p in MATERIALS.items() if p['layer_type'] != 'none']
        surface_materials = [m for m, p in MATERIALS.items() if p['layer_type'] == 'surface']
        
        layer_data = []
        status_placeholders = {}

        # ===== ชั้นที่ 1: ผิวทาง =====
        st.subheader("🔶 ชั้นที่ 1: ผิวทาง (Surface)")
        
        layer1_mat_default = st.session_state.get('layer1_mat', surface_materials[0])
        layer1_mat_idx = (surface_materials.index(layer1_mat_default) 
                         if layer1_mat_default in surface_materials else 0)

        layer1_mat = st.selectbox(
            "เลือกวัสดุ", options=surface_materials,
            index=layer1_mat_idx, key="layer1_mat"
        )

        # ===== AC Sublayer (Compact Table) =====
        use_sublayers = st.checkbox(
            "📐 แบ่งชั้นย่อยผิวทาง AC (Wearing, Binder, Base Course)",
            value=st.session_state.get('use_ac_sublayers', False),
            help="แบ่งชั้น AC ออกเป็น 3 ชั้นย่อย ตามมาตรฐานกรมทางหลวง",
            key="use_ac_sublayers"
        )

        mat_props_1 = MATERIALS[layer1_mat]
        default_a1 = mat_props_1['layer_coeff']
        default_m1 = mat_props_1['drainage_coeff']

        if use_sublayers:
            st.info("📋 ความหนามาตรฐาน ทล.: Wearing 40-70 มม. / Binder 40-80 มม. / Base 70-100 มม.")
            
            # Compact 3-column table
            col_w, col_b, col_bc = st.columns(3)
            with col_w:
                st.markdown("**Wearing Course**")
                wearing_options = ["กำหนดเอง"] + [f"{t} มม." for t in DOH_THICKNESS_STANDARDS["Wearing Course"]]
                wearing_std = st.selectbox("มาตรฐาน ทล.", wearing_options, index=0, key="wearing_std_select")
                if wearing_std != "กำหนดเอง":
                    wearing_thick = int(wearing_std.replace(" มม.", "")) / 10
                    st.metric("ความหนา", f"{wearing_thick:.1f} cm")
                else:
                    wearing_thick = st.number_input("ความหนา (cm)", 1.0, 15.0,
                        value=st.session_state.get('wearing_thick_val', 5.0), step=0.5, key="wearing_thick")

            with col_b:
                st.markdown("**Binder Course**")
                binder_options = ["กำหนดเอง"] + [f"{t} มม." for t in DOH_THICKNESS_STANDARDS["Binder Course"]]
                binder_std = st.selectbox("มาตรฐาน ทล.", binder_options, index=0, key="binder_std_select")
                if binder_std != "กำหนดเอง":
                    binder_thick = int(binder_std.replace(" มม.", "")) / 10
                    st.metric("ความหนา", f"{binder_thick:.1f} cm")
                else:
                    binder_thick = st.number_input("ความหนา (cm)", 1.0, 15.0,
                        value=st.session_state.get('binder_thick_val', 7.0), step=0.5, key="binder_thick")

            with col_bc:
                st.markdown("**Base Course**")
                base_options = ["กำหนดเอง", "ไม่ใช้"] + [f"{t} มม." for t in DOH_THICKNESS_STANDARDS["Base Course"] if t > 0]
                base_std = st.selectbox("มาตรฐาน ทล.", base_options, index=0, key="base_std_select")
                if base_std == "ไม่ใช้":
                    base_course_thick = 0.0
                    st.metric("ความหนา", "0.0 cm")
                elif base_std != "กำหนดเอง":
                    base_course_thick = int(base_std.replace(" มม.", "")) / 10
                    st.metric("ความหนา", f"{base_course_thick:.1f} cm")
                else:
                    base_course_thick = st.number_input("ความหนา (cm)", 0.0, 15.0,
                        value=st.session_state.get('base_thick_val', 10.0), step=0.5, key="base_course_thick")

            layer1_thick = wearing_thick + binder_thick + base_course_thick
            st.markdown(
                f'<p style="color: #1E90FF; font-size: 18px; font-weight: bold;">'
                f'📏 ความหนารวม AC = {wearing_thick:.1f} + {binder_thick:.1f} + {base_course_thick:.1f} = {layer1_thick:.1f} cm</p>',
                unsafe_allow_html=True)

            st.session_state['ac_sublayers'] = {
                'wearing': wearing_thick, 'binder': binder_thick,
                'base': base_course_thick, 'total': layer1_thick
            }

            # a and m for AC sublayer
            col_am1, col_am2 = st.columns(2)
            with col_am1:
                st.markdown(f"a₁ <span style='color:#1E90FF;font-size:12px;'>(default={default_a1:.2f})</span>", unsafe_allow_html=True)
                layer1_a = st.number_input("a1", 0.10, 0.50,
                    value=st.session_state.get('layer1_a', default_a1), step=0.01,
                    key="layer1_a", label_visibility="collapsed")
            with col_am2:
                st.markdown(f"m₁ <span style='color:#1E90FF;font-size:12px;'>(default={default_m1:.2f})</span>", unsafe_allow_html=True)
                layer1_m = st.number_input("m1", 0.5, 1.5,
                    value=st.session_state.get('layer1_m', default_m1), step=0.05,
                    key="layer1_m", label_visibility="collapsed")
        else:
            # No sublayers
            st.session_state['ac_sublayers'] = None
            col_a, col_b, col_c = st.columns(3)
            with col_a:
                layer1_thick = st.number_input("ความหนา (cm)", 1.0, 30.0,
                    value=st.session_state.get('layer1_thick', 5.0), step=1.0, key="layer1_thick")
            with col_b:
                st.markdown(f"a₁ <span style='color:#1E90FF;font-size:12px;'>(default={default_a1:.2f})</span>", unsafe_allow_html=True)
                layer1_a = st.number_input("a1", 0.10, 0.50,
                    value=st.session_state.get('layer1_a', default_a1), step=0.01,
                    key="layer1_a", label_visibility="collapsed")
            with col_c:
                st.markdown(f"m₁ <span style='color:#1E90FF;font-size:12px;'>(default={default_m1:.2f})</span>", unsafe_allow_html=True)
                layer1_m = st.number_input("m1", 0.5, 1.5,
                    value=st.session_state.get('layer1_m', default_m1), step=0.05,
                    key="layer1_m", label_visibility="collapsed")

        st.markdown(f'<p style="color: #1E90FF; font-size: 14px;">E = {mat_props_1["mr_mpa"]:,} MPa</p>', unsafe_allow_html=True)
        status_placeholders[1] = st.empty()

        layer_data.append({
            'material': layer1_mat,
            'thickness_cm': layer1_thick,
            'layer_coeff': layer1_a,
            'drainage_coeff': layer1_m
        })

        # ===== ชั้นที่ 2-6 =====
        default_materials = [
            "พื้นทางซีเมนต์ CTB",
            "รองพื้นทางวัสดุมวลรวม CBR 25%",
            "วัสดุคัดเลือก ก",
            "วัสดุคัดเลือก ก",
            "วัสดุคัดเลือก ก"
        ]
        default_thickness = [15.0, 15.0, 30.0, 30.0, 30.0]
        layer_icons = ['🔷', '🔶', '🟢', '🟡', '🔴']

        for i in range(2, num_layers + 1):
            st.markdown("---")
            st.subheader(f"{layer_icons[i-2]} ชั้นที่ {i}")

            layer_i_mat_default = st.session_state.get(f'layer{i}_mat', default_materials[i-2])
            if layer_i_mat_default in all_materials:
                default_idx = all_materials.index(layer_i_mat_default)
            else:
                default_idx = (all_materials.index(default_materials[i-2]) 
                              if default_materials[i-2] in all_materials else 0)

            layer_mat = st.selectbox(
                f"เลือกวัสดุชั้นที่ {i}", options=all_materials,
                index=min(default_idx, len(all_materials)-1), key=f"layer{i}_mat"
            )

            mat_props = MATERIALS[layer_mat]
            default_a = mat_props['layer_coeff']
            default_m = mat_props['drainage_coeff']

            # Auto-reset a, m when material changes
            prev_mat_key = f'layer{i}_prev_mat'
            if prev_mat_key not in st.session_state:
                st.session_state[prev_mat_key] = layer_mat
            if st.session_state[prev_mat_key] != layer_mat:
                st.session_state[f'layer{i}_a'] = default_a
                st.session_state[f'layer{i}_m'] = default_m
                st.session_state[prev_mat_key] = layer_mat

            col_c, col_d, col_e = st.columns(3)
            with col_c:
                layer_thick = st.number_input("ความหนา (cm)", 1.0, 150.0,
                    value=st.session_state.get(f'layer{i}_thick', default_thickness[i-2]),
                    step=5.0, key=f"layer{i}_thick")
            with col_d:
                st.markdown(f"a{i} <span style='color:#1E90FF;font-size:12px;'>(default={default_a:.2f})</span>", unsafe_allow_html=True)
                layer_a = st.number_input(f"a{i}", 0.01, 0.50,
                    value=st.session_state.get(f'layer{i}_a', default_a), step=0.01,
                    key=f"layer{i}_a", label_visibility="collapsed")
            with col_e:
                st.markdown(f"m{i} <span style='color:#1E90FF;font-size:12px;'>(default={default_m:.2f})</span>", unsafe_allow_html=True)
                layer_m = st.number_input(f"m{i}", 0.5, 1.5,
                    value=st.session_state.get(f'layer{i}_m', default_m), step=0.05,
                    key=f"layer{i}_m", label_visibility="collapsed")

            st.markdown(f'<p style="color: #1E90FF; font-size: 14px;">E = {mat_props["mr_mpa"]:,} MPa</p>', unsafe_allow_html=True)
            status_placeholders[i] = st.empty()

            layer_data.append({
                'material': layer_mat,
                'thickness_cm': layer_thick,
                'layer_coeff': layer_a,
                'drainage_coeff': layer_m
            })

    # ========================================
    # CALCULATION
    # ========================================
    inputs = {
        'W18': W18, 'reliability': reliability, 'Zr': Zr, 'So': So,
        'P0': P0, 'Pt': Pt, 'delta_psi': delta_psi, 'CBR': CBR, 'Mr': Mr
    }
    ac_sublayers = st.session_state.get('ac_sublayers', None)
    calc_results = calculate_layer_thicknesses(W18, Zr, So, delta_psi, Mr, layer_data, ac_sublayers)
    design_check = check_design(calc_results['total_sn_required'], calc_results['total_sn_provided'])

    # Fill status placeholders in Layer tab
    for layer in calc_results['layers']:
        layer_no = layer['layer_no']
        if layer_no in status_placeholders:
            with status_placeholders[layer_no]:
                if layer['is_ok']:
                    st.success(f"✅ ผ่าน (ต้องการ ≥ {layer['min_thickness_cm']:.1f} cm)")
                else:
                    shortage = layer['min_thickness_cm'] - layer['design_thickness_cm']
                    st.error(f"❌ ไม่ผ่าน (ต้องเพิ่มอีก {shortage:.1f} cm)")

    # ========================================
    # TAB 3: RESULTS
    # ========================================
    with tab_results:

        # ===== QUICK SUMMARY CARD =====
        st.markdown("### 🎯 สรุปผลการออกแบบ (Quick Summary)")
        
        if design_check['passed']:
            st.markdown(
                f"""<div style="background-color: #d4edda; border: 2px solid #28a745; border-radius: 10px; 
                padding: 20px; text-align: center; margin-bottom: 20px;">
                <h2 style="color: #28a745; margin: 0;">✅ PASS — การออกแบบผ่านเกณฑ์</h2>
                <p style="font-size: 18px; margin: 10px 0;">
                SN<sub>provided</sub> = <b>{calc_results['total_sn_provided']:.2f}</b> &nbsp;≥&nbsp; 
                SN<sub>required</sub> = <b>{calc_results['total_sn_required']:.2f}</b>
                &nbsp;&nbsp;|&nbsp;&nbsp; Safety Margin = <b>{design_check['safety_margin']:.2f}</b>
                </p></div>""", unsafe_allow_html=True)
        else:
            st.markdown(
                f"""<div style="background-color: #f8d7da; border: 2px solid #dc3545; border-radius: 10px; 
                padding: 20px; text-align: center; margin-bottom: 20px;">
                <h2 style="color: #dc3545; margin: 0;">❌ FAIL — การออกแบบไม่ผ่าน</h2>
                <p style="font-size: 18px; margin: 10px 0;">
                SN<sub>provided</sub> = <b>{calc_results['total_sn_provided']:.2f}</b> &nbsp;&lt;&nbsp; 
                SN<sub>required</sub> = <b>{calc_results['total_sn_required']:.2f}</b>
                &nbsp;&nbsp;|&nbsp;&nbsp; ขาดอีก = <b>{abs(design_check['safety_margin']):.2f}</b>
                </p></div>""", unsafe_allow_html=True)

        # ===== WARNINGS =====
        warnings = calc_results.get('warnings', [])
        if warnings:
            for w in warnings:
                st.warning(w)

        # ===== W18 Supported =====
        w18_supported = calculate_w18_supported(
            calc_results['total_sn_provided'], Zr, So, delta_psi, Mr
        )
        w18_supported_million = w18_supported / 1_000_000
        w18_diff_percent = ((w18_supported - W18) / W18) * 100

        w18_col1, w18_col2 = st.columns(2)
        with w18_col1:
            st.metric("W₁₈ ออกแบบ", f"{W18/1e6:,.2f} ล้าน")
        with w18_col2:
            delta_str = f"{w18_diff_percent:+.1f}%"
            st.metric("W₁₈ รองรับได้", f"{w18_supported_million:,.2f} ล้าน",
                      delta=delta_str, delta_color="normal" if w18_diff_percent >= 0 else "inverse")

        st.markdown("---")

        # ===== STEP-BY-STEP CALCULATION =====
        st.subheader("🔢 ขั้นตอนการคำนวณความหนาแต่ละชั้น")

        for layer in calc_results['layers']:
            with st.container():
                layer_status = "✅" if layer['is_ok'] else "❌"
                st.markdown(f"### {layer_status} ชั้นที่ {layer['layer_no']}: {layer['material']}")

                # AC sublayer info
                layer_ac_sub = layer.get('ac_sublayers', None)
                if layer_ac_sub is not None and layer['layer_no'] == 1:
                    st.info(f"**📐 แบ่งชั้นย่อย AC:** "
                           f"Wearing = {layer_ac_sub['wearing']:.1f} cm | "
                           f"Binder = {layer_ac_sub['binder']:.1f} cm | "
                           f"Base = {layer_ac_sub['base']:.1f} cm | "
                           f"**รวม = {layer_ac_sub['total']:.1f} cm**")

                col_a, col_b = st.columns([1, 1])
                with col_a:
                    st.markdown("**ข้อมูลวัสดุ:**")
                    st.markdown(f"- E (MPa) = **{layer['mr_mpa']:,}**")
                    st.markdown(f"- Mᵣ (psi) = **{layer['mr_psi']:,}**")
                    st.markdown(f"- Layer Coefficient (a{layer['layer_no']}) = **{layer['a_i']:.2f}**")
                    st.markdown(f"- Drain Coefficient (m{layer['layer_no']}) = **{layer['m_i']:.2f}**")

                with col_b:
                    st.markdown("**จากสมการ AASHTO:**")
                    sn_at_layer = layer['sn_required_at_layer']
                    st.latex(f"SN_{{{layer['layer_no']}}} = {sn_at_layer:.2f}")

                # Thickness formula
                st.markdown("**คำนวณความหนาผิวทาง:**")
                if layer['layer_no'] == 1:
                    st.latex(f"D_{{1}} \\geq \\frac{{SN_{{1}}}}{{a_{{1}} \\times m_{{1}}}} = "
                            f"\\frac{{{sn_at_layer:.2f}}}{{{layer['a_i']:.2f} \\times {layer['m_i']:.2f}}} = "
                            f"{layer['min_thickness_inch']:.2f} \\text{{ นิ้ว}}")
                else:
                    prev_sn = calc_results['layers'][layer['layer_no']-2]['cumulative_sn']
                    st.latex(f"D_{{{layer['layer_no']}}} \\geq "
                            f"\\frac{{SN_{{{layer['layer_no']}}} - SN_{{prev}}}}"
                            f"{{a_{{{layer['layer_no']}}} \\times m_{{{layer['layer_no']}}}}} = "
                            f"\\frac{{{sn_at_layer:.2f} - {prev_sn:.2f}}}"
                            f"{{{layer['a_i']:.2f} \\times {layer['m_i']:.2f}}} = "
                            f"{layer['min_thickness_inch']:.2f} \\text{{ นิ้ว}}")

                result_cols = st.columns(4)
                with result_cols[0]:
                    st.metric("ความหนาขั้นต่ำ", f"{layer['min_thickness_cm']:.1f} cm")
                with result_cols[1]:
                    st.metric("ความหนาที่เลือก", f"{layer['design_thickness_cm']:.0f} cm",
                             delta=f"{layer['design_thickness_cm'] - layer['min_thickness_cm']:.1f} cm")
                with result_cols[2]:
                    st.metric("SN contribution", f"{layer['sn_contribution']:.3f}")
                with result_cols[3]:
                    st.metric("Cumulative SN", f"{layer['cumulative_sn']:.2f}")

                if layer['is_ok']:
                    st.success(f"✅ **OK** — ความหนาเพียงพอ ({layer['design_thickness_cm']:.0f} ≥ {layer['min_thickness_cm']:.1f} cm)")
                else:
                    st.error(f"❌ **NG** — ต้องเพิ่มความหนาอีก {layer['min_thickness_cm'] - layer['design_thickness_cm']:.1f} cm")
                st.markdown("---")

        # ===== SN TABLE =====
        with st.expander("📋 ตารางสรุปการคำนวณ SN"):
            table_data = []
            for layer in calc_results['layers']:
                table_data.append({
                    'ชั้น': layer['layer_no'],
                    'วัสดุ': layer['short_name'],
                    'aᵢ': layer['a_i'],
                    'Dᵢ (cm)': layer['design_thickness_cm'],
                    'Dᵢ (in)': layer['design_thickness_inch'],
                    'mᵢ': layer['m_i'],
                    'E (MPa)': layer['mr_mpa'],
                    'SN contrib.': layer['sn_contribution'],
                    'SN cumul.': layer['cumulative_sn']
                })
            st.table(table_data)
            st.markdown(f"""
            **สูตรการคำนวณ:** $SN = \\sum_{{i=1}}^{{n}} a_i \\times D_i \\times m_i$
            
            **ผลลัพธ์:** SN_provided = {calc_results['total_sn_provided']:.2f} | SN_required = {calc_results['total_sn_required']:.2f}
            """)

        # ===== PAVEMENT SECTION FIGURE =====
        st.subheader("📐 ภาพตัดขวางโครงสร้างถนน")
        fig_lang = 'th' if figure_language == "ภาษาไทย" else 'en'
        fig = plot_pavement_section(calc_results['layers'], Mr, CBR, lang=fig_lang)
        st.pyplot(fig)
        plt.close(fig)  # Fix memory leak

        # ===== SENSITIVITY ANALYSIS =====
        st.subheader("📈 Sensitivity Analysis")
        
        sens_col1, sens_col2 = st.columns(2)
        with sens_col1:
            fig_cbr = plot_sensitivity_cbr(W18, Zr, So, delta_psi, CBR)
            st.pyplot(fig_cbr)
            plt.close(fig_cbr)
        with sens_col2:
            fig_w18 = plot_sensitivity_w18(Zr, So, delta_psi, Mr, W18)
            st.pyplot(fig_w18)
            plt.close(fig_w18)

    # ========================================
    # TAB 4: REPORT & EXPORT
    # ========================================
    with tab_report:
        st.header("📄 ส่งออกรายงาน")

        # ============================================================
        # REPORT SETTINGS: เลขหัวข้อ / ตาราง / รูป / คำบรรยาย
        # ============================================================
        st.markdown("### 📝 ตั้งค่าหมายเลขหัวข้อและตารางสำหรับรายงาน Word")

        col_num1, col_num2, col_num3 = st.columns(3)
        with col_num1:
            rs_section_number = st.text_input(
                "เลขหัวข้อ",
                value=st.session_state.get('rs_section_number', '4.4'),
                key='rs_section_number'
            )
        with col_num2:
            rs_table_number_inputs = st.text_input(
                "เลขตารางพารามิเตอร์",
                value=st.session_state.get('rs_table_number_inputs', '4-8'),
                key='rs_table_number_inputs'
            )
        with col_num3:
            rs_table_number_materials = st.text_input(
                "เลขตารางวัสดุ",
                value=st.session_state.get('rs_table_number_materials', '4-9'),
                key='rs_table_number_materials'
            )

        rs_figure_number = st.text_input(
            "เลขรูป",
            value=st.session_state.get('rs_figure_number', '4-8'),
            key='rs_figure_number'
        )

        rs_section_title = st.text_input(
            "ชื่อหัวข้อ",
            value=st.session_state.get('rs_section_title', 'การออกแบบผิวทางลาดยาง (Flexible Pavement)'),
            key='rs_section_title'
        )

        col_cap1, col_cap2 = st.columns(2)
        with col_cap1:
            rs_table_caption_inputs = st.text_input(
                "คำบรรยายตารางพารามิเตอร์",
                value=st.session_state.get('rs_table_caption_inputs', 'ค่าพารามิเตอร์ที่ใช้ในการออกแบบผิวทางยืดหยุ่น'),
                key='rs_table_caption_inputs'
            )
        with col_cap2:
            rs_table_caption_materials = st.text_input(
                "คำบรรยายตารางวัสดุ",
                value=st.session_state.get('rs_table_caption_materials', 'ค่าสัมประสิทธิ์และค่าโมดูลัสของวัสดุโครงสร้างชั้นทาง'),
                key='rs_table_caption_materials'
            )

        rs_figure_caption = st.text_input(
            "คำบรรยายรูป",
            value=st.session_state.get('rs_figure_caption', 'รูปตัดโครงสร้างชั้นทางที่ออกแบบ'),
            key='rs_figure_caption'
        )

        # รวบรวม report_settings
        report_settings = {
            'section_number':          rs_section_number,
            'table_number_inputs':     rs_table_number_inputs,
            'table_number_materials':  rs_table_number_materials,
            'figure_number':           rs_figure_number,
            'section_title':           rs_section_title,
            'table_caption_inputs':    rs_table_caption_inputs,
            'table_caption_materials': rs_table_caption_materials,
            'figure_caption':          rs_figure_caption,
        }

        st.markdown("---")

        # ============================================================
        # PREVIEW บทเกริ่นนำ (HTML)
        # ============================================================
        st.markdown("### 👁️ Preview บทเกริ่นนำ")

        total_thick_prev = sum(l['design_thickness_cm'] for l in calc_results['layers'])
        num_layers_prev  = len(calc_results['layers'])
        passed_prev      = 'ผ่านเกณฑ์' if design_check['passed'] else 'ไม่ผ่านเกณฑ์'

        def hl_purple(val):
            return f'<span style="background-color:#D8B4FE;padding:1px 4px;border-radius:3px;font-weight:bold;">{val}</span>'

        def hl_yellow(val):
            return f'<span style="background-color:#FDE68A;padding:1px 4px;border-radius:3px;font-weight:bold;">{val}</span>'

        intro_html = f"""
        <div style="background:#f9f9f9;padding:15px 20px;border-radius:8px;border:1px solid #ddd;
                    font-family:'TH SarabunPSK',Sarabun,sans-serif;font-size:16px;line-height:1.9;">
            <p style="font-weight:bold;margin-bottom:5px;">
                {hl_yellow(rs_section_number)}&nbsp;&nbsp;{hl_yellow(rs_section_title)}
            </p>
            <p style="text-indent:40px;text-align:justify;text-justify:inter-character;margin-top:8px;">
                ถนนลาดยางซึ่งประกอบด้วยวัสดุงานทางหลายชนิด การออกแบบโครงสร้างถนนแบบยืดหยุ่น (Flexible Pavement)
                ใช้วิธี AASHTO 1993 Guide for Design of Pavement Structures โดยพิจารณาปัจจัยด้านปริมาณจราจรสะสม ESALs
                ความน่าเชื่อถือ และคุณสมบัติของดินรองรับ
                สำหรับโครงการนี้ที่ปรึกษาได้กำหนดค่าพารามิเตอร์หลักในการออกแบบ ได้แก่
                ปริมาณ W&#8321;&#8328; = {hl_purple(f"{W18:,.0f}")} 18-kip ESALs
                ที่ระดับความน่าเชื่อถือ (Reliability) = {hl_purple(reliability)} %
                โดยมีดินเดิมค่า CBR = {hl_purple(f"{CBR:.1f}")} % (M&#7523; = {hl_purple(f"{Mr:,.0f}")} psi)
                ผลการออกแบบได้โครงสร้างชั้นทาง {hl_purple(num_layers_prev)} ชั้น
                ที่ SN&#8203;_required = {hl_purple(f"{calc_results['total_sn_required']:.2f}")}
                และ SN&#8203;_provided = {hl_purple(f"{calc_results['total_sn_provided']:.2f}")}
                ความหนารวม {hl_purple(f"{total_thick_prev:.0f}")} ซม.
                การออกแบบ{hl_purple(passed_prev)}
                ดังแสดงผลการวิเคราะห์ใน<b>ตารางที่ {hl_yellow(rs_table_number_inputs)}</b>
                และ<b>ตารางที่ {hl_yellow(rs_table_number_materials)}</b>
                และ<b>รูปที่ {hl_yellow(rs_figure_number)}</b>
            </p>
        </div>
        """
        st.markdown(intro_html, unsafe_allow_html=True)
        st.caption("🟣 สีม่วง = ดึงจากผลคำนวณอัตโนมัติ | 🟡 สีเหลือง = ผู้ใช้กรอกเอง")

        st.markdown("---")

        # ============================================================
        # EXPORT BUTTONS
        # ============================================================
        col_exp0, col_exp1, col_exp2, col_exp3 = st.columns(4)

        with col_exp0:
            if st.button("📋 สร้างรายงานแบบมีเกริ่นนำ", type="primary",
                         help="รายงานรูปแบบสำหรับรวมกับบทรายงานอื่น — มีหัวข้อ, เกริ่นนำ, ตาราง, รูป"):
                with st.spinner("กำลังสร้างรายงาน..."):
                    fig_intro = plot_pavement_section(calc_results['layers'], Mr, CBR, lang='th')
                    doc_intro_bytes = create_word_report_intro(
                        project_title, inputs, calc_results, design_check, fig_intro, report_settings
                    )
                    plt.close(fig_intro)
                    st.download_button(
                        label="⬇️ ดาวน์โหลดรายงานแบบเกริ่นนำ",
                        data=doc_intro_bytes,
                        file_name=f"Flexible_Intro_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

        with col_exp1:
            if st.button("📝 สร้างรายงาน Word (เต็ม)"):
                with st.spinner("กำลังสร้างรายงาน..."):
                    fig_thai = plot_pavement_section(calc_results['layers'], Mr, CBR, lang='th')
                    doc_bytes = create_word_report(project_title, inputs, calc_results, design_check, fig_thai)
                    plt.close(fig_thai)
                    st.download_button(
                        label="⬇️ ดาวน์โหลดรายงาน Word (เต็ม)",
                        data=doc_bytes,
                        file_name=f"AASHTO_Flexible_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

        with col_exp2:
            fig_export = plot_pavement_section(calc_results['layers'], Mr, CBR, lang=fig_lang)
            fig_bytes = get_figure_as_bytes(fig_export)
            plt.close(fig_export)
            st.download_button(
                label="📸 ดาวน์โหลดรูปตัดขวาง (PNG)",
                data=fig_bytes,
                file_name=f"Pavement_Section_{datetime.now().strftime('%Y%m%d_%H%M')}.png",
                mime="image/png"
            )

        with col_exp3:
            export_data = {
                'project_title': project_title,
                'W18': W18,
                'reliability': reliability,
                'So': So,
                'P0': P0,
                'Pt': Pt,
                'CBR': CBR,
                'num_layers': num_layers,
                'layers': layer_data,
                'ac_sublayers': st.session_state.get('ac_sublayers', None),
                'report_settings': report_settings,
            }
            json_str = json.dumps(export_data, ensure_ascii=False, indent=2)
            st.download_button(
                label="💾 ดาวน์โหลดข้อมูล (JSON)",
                data=json_str,
                file_name=f"Flexible_Input_{datetime.now().strftime('%Y%m%d_%H%M')}.json",
                mime="application/json"
            )

        st.markdown("---")

        # ===== Summary in report tab =====
        st.subheader("📊 สรุปผลการออกแบบ")
        
        summary_data = [
            ("ชื่อโครงการ", project_title),
            ("W₁₈ (Design ESALs)", f"{W18:,.0f} ({W18/1e6:,.2f} ล้าน)"),
            ("Reliability", f"{reliability}%"),
            ("CBR", f"{CBR:.1f}%"),
            ("Mᵣ (Subgrade)", f"{Mr:,} psi"),
            ("SN Required", f"{calc_results['total_sn_required']:.2f}"),
            ("SN Provided", f"{calc_results['total_sn_provided']:.2f}"),
            ("Safety Margin", f"{design_check['safety_margin']:.2f}"),
            ("ผลการตรวจสอบ", "✅ PASS" if design_check['passed'] else "❌ FAIL"),
        ]
        st.table(summary_data)

    # ===== FOOTER =====
    st.markdown("---")
    st.markdown("""
    <div style='text-align: center; color: gray;'>
    <p>AASHTO 1993 Flexible Pavement Design Application v5.0</p>
    <p>พัฒนาโดย รศ.ดร.อิทธิพล มีผล // ภาควิชาครุศาสตร์โยธา // มจพ.</p>
    </div>
    """, unsafe_allow_html=True)


# ================================================================================
# ENTRY POINT
# ================================================================================

if __name__ == "__main__":
    main()
