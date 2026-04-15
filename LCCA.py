#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
================================================================================
โปรแกรมวิเคราะห์ต้นทุนตลอดอายุการใช้งานผิวทาง (LCCA) - Integrated v2.0
Life-Cycle Cost Analysis for Pavement Alternatives
================================================================================
พัฒนาโดย: รศ.ดร.อิทธิพล มีผล
ภาควิชาครุศาสตร์โยธา มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ (KMUTNB)

การปรับปรุง v2.1.1:
  - ยกเลิก field สายทาง/ตอนควบคุม + ปีงบประมาณ
  - เปลี่ยน กม.ที่ → ระยะทางรวม + Section ขนาดถนน (คำนวณพื้นที่อัตโนมัติ)
  - CBR ย้ายไปตัวแปรร่วม TAB 2
  - Progress Indicator + Warning + Badge ✅/⚠️
  - แสดงผลทั้ง บาท/ตร.ม./ปี และ บาท/กม./ปี
  - Breakeven Year Analysis
  - Cumulative Cost Timeline Graph
  - CSS/UI — Metric cards + color band + result highlight ... แก้ไขรายงาน word
================================================================================
"""

import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from dataclasses import dataclass
from typing import List, Dict
import json, io
from datetime import datetime
from itertools import combinations

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    from docx import Document as WordDocument
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ─────────────────────────────────────────────────────────────────────────────
st.set_page_config(page_title="LCCA Pavement v2.0", page_icon="🛣️", layout="wide")

# =============================================================================
# GLOBAL CSS
# =============================================================================
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Sarabun:wght@300;400;600;700&display=swap');

/* ── Global font override ── */
html, body, [class*="css"], .stApp, .stMarkdown,
button, input, select, textarea, label,
.stDataFrame, .stSelectbox, .stNumberInput,
.stTextInput, .stButton, .stTabs {
    font-family: 'Sarabun', sans-serif !important;
}

/* ── Card Metric ── */
.metric-card {
    border-radius: 12px;
    padding: 16px 20px;
    margin: 6px 0;
    color: #fff;
    box-shadow: 0 4px 12px rgba(0,0,0,0.15);
}
.metric-card .label {
    font-size: 14px; font-weight: 600;
    opacity: 0.88; margin-bottom: 4px;
    font-family: 'Sarabun', sans-serif;
}
.metric-card .value {
    font-size: 28px; font-weight: 700;
    font-family: 'Sarabun', sans-serif;
}
.metric-card .sub {
    font-size: 13px; opacity: 0.80; margin-top: 2px;
    font-family: 'Sarabun', sans-serif;
}
.card-blue   { background: linear-gradient(135deg,#1565C0,#42A5F5); }
.card-orange { background: linear-gradient(135deg,#E65100,#FFA726); }
.card-green  { background: linear-gradient(135deg,#1B5E20,#66BB6A); }
.card-gold   { background: linear-gradient(135deg,#827717,#FFEE58); color:#333 !important; }
.card-purple { background: linear-gradient(135deg,#4A148C,#AB47BC); }
.card-teal   { background: linear-gradient(135deg,#004D40,#26A69A); }

/* ── Section header band ── */
.sec-header {
    background: linear-gradient(90deg,#1565C0 0%,#42A5F5 100%);
    color: #fff; padding: 8px 16px; border-radius: 8px;
    font-size: 17px; font-weight: 700; margin: 12px 0 8px 0;
    font-family: 'Sarabun', sans-serif;
}
.sec-header-orange {
    background: linear-gradient(90deg,#E65100 0%,#FFA726 100%);
    color:#fff; padding:8px 16px; border-radius:8px;
    font-size:17px; font-weight:700; margin:12px 0 8px 0;
    font-family:'Sarabun',sans-serif;
}
.sec-header-green {
    background: linear-gradient(90deg,#1B5E20 0%,#66BB6A 100%);
    color:#fff; padding:8px 16px; border-radius:8px;
    font-size:17px; font-weight:700; margin:12px 0 8px 0;
    font-family:'Sarabun',sans-serif;
}

/* ── Progress bar ── */
.progress-wrap {
    display:flex; gap:8px; margin:8px 0 16px 0; align-items:center;
}
.prog-step {
    flex:1; padding:8px 4px; border-radius:8px; text-align:center;
    font-size:14px; font-weight:600; border:2px solid #ccc;
    font-family:'Sarabun',sans-serif;
}
.prog-done  { background:#E8F5E9; border-color:#43A047; color:#1B5E20; }
.prog-warn  { background:#FFF8E1; border-color:#FFA000; color:#E65100; }
.prog-idle  { background:#F5F5F5; border-color:#BDBDBD; color:#757575; }
.prog-arrow { font-size:18px; color:#BDBDBD; flex:0; }

/* ── Badge ── */
.badge-ok   { background:#E8F5E9; color:#1B5E20; border:1px solid #43A047;
              border-radius:20px; padding:3px 12px; font-size:14px;
              font-weight:600; font-family:'Sarabun',sans-serif; }
.badge-warn { background:#FFF3E0; color:#E65100; border:1px solid #FFA000;
              border-radius:20px; padding:3px 12px; font-size:14px;
              font-weight:600; font-family:'Sarabun',sans-serif; }

/* ── Best result highlight ── */
.best-row {
    background: linear-gradient(90deg,#E8F5E9,#F1F8E9);
    border-left: 5px solid #43A047;
    border-radius: 6px; padding: 10px 14px; margin: 4px 0;
    font-family:'Sarabun',sans-serif; font-size:15px;
}

/* ── Road preview box ── */
.road-box {
    background: linear-gradient(135deg,#263238,#455A64);
    color:#fff; border-radius:10px; padding:14px 18px;
    font-family:'Sarabun',sans-serif; font-size:16px;
    margin-top:8px;
}
.road-box .road-val { font-size:24px; font-weight:700; color:#FFD54F; }

/* ── Info band ── */
.info-band {
    background:#E3F2FD; border-left:4px solid #1565C0;
    border-radius:6px; padding:10px 14px; margin:8px 0;
    font-family:'Sarabun',sans-serif; font-size:15px; color:#0D47A1;
}
.warn-band {
    background:#FFF8E1; border-left:4px solid #FFA000;
    border-radius:6px; padding:10px 14px; margin:8px 0;
    font-family:'Sarabun',sans-serif; font-size:15px; color:#E65100;
}
</style>
""", unsafe_allow_html=True)

# =============================================================================
# SECTION A: LOOKUP TABLES
# =============================================================================
X1_MAP = {
    "High Type (AC/PM บนหินคลุก)": 0.00,
    "Intermediate Type (AC/PM บน Stabilized)": 0.50,
    "Low Type (ST บน Soil-Aggregate)": 1.00,
}
X2_BREAKS = [(0,2,1.00),(2.01,3,0.75),(3.01,4,0.50),(4.01,5,0.25),(5.01,999,0.00)]
X3_OPTIONS = {
    "0 – 500       (X3=0.00)":0.00, "501 – 600     (X3=0.04)":0.04,
    "601 – 700     (X3=0.08)":0.08, "701 – 800     (X3=0.12)":0.12,
    "801 – 900     (X3=0.16)":0.16, "901 – 1,000   (X3=0.20)":0.20,
    "1,001 – 1,100 (X3=0.24)":0.24, "1,101 – 1,200 (X3=0.29)":0.29,
    "1,201 – 1,300 (X3=0.33)":0.33, "1,301 – 1,400 (X3=0.37)":0.37,
    "1,401 – 1,500 (X3=0.41)":0.41, "1,501 – 1,600 (X3=0.45)":0.45,
    "1,601 – 1,700 (X3=0.49)":0.49, "1,701 – 1,800 (X3=0.53)":0.53,
    "1,801 – 1,900 (X3=0.57)":0.57, "1,901 – 2,000 (X3=0.61)":0.61,
    "2,001 – 2,200 (X3=0.69)":0.69, "2,201 – 2,400 (X3=0.78)":0.78,
    "2,401 – 2,600 (X3=0.86)":0.86, "2,601 – 2,800 (X3=0.94)":0.94,
    "2,801 – 3,000 (X3=1.02)":1.02, "3,001 – 3,300 (X3=1.14)":1.14,
    "3,301 – 3,600 (X3=1.27)":1.27, "3,601 – 3,900 (X3=1.37)":1.37,
    "3,901 – 4,200 (X3=1.51)":1.51, "4,201 – 4,500 (X3=1.64)":1.64,
    "4,501 – 4,800 (X3=1.76)":1.76, "4,801 – 5,100 (X3=1.88)":1.88,
    "5,101 – 5,400 (X3=2.00)":2.00, "5,401 – 5,700 (X3=2.13)":2.13,
    "5,701+         (X3=2.25)":2.25,
}
X4_BREAKS = [(0,3,0.00),(4,4,0.20),(5,5,0.40),(6,6,0.60),(7,7,0.80),
             (8,8,1.00),(9,9,1.20),(10,10,1.40),(11,11,1.60),(12,99999,1.80)]
X5_BREAKS = [(0,5.49,0.00),(5.50,5.99,0.02),(6.00,6.49,0.05),(6.50,6.99,0.10),(7.00,9999,0.19)]
TERRAIN_MAP  = {"ที่ราบ (0-3%)":"P","ลูกเนิน (3-5%)":"R",
                "ลูกเนินสลับเขา (5-7%)":"RM","เขา (>7%)":"S"}
TERRAIN_KEYS = list(TERRAIN_MAP.keys())
X6_MAP = {"P":0.00,"R":0.02,"RM":0.04,"S":0.07}
Y3_MAP = {"P":0.00,"R":0.24,"RM":0.36,"S":0.48}
Y4_MAP = {"P":0.00,"R":0.24,"RM":0.36,"S":0.48}
Y6_MAP = {"P":0.00,"R":0.04,"RM":0.08,"S":0.12}
Y1_BREAKS = [(0,40,0.00),(40.01,60,0.10),(60.01,80,0.20),(80.01,9999,0.30)]
Y2_BREAKS = [(0,1.75,0.00),(1.76,2.00,0.10),(2.01,2.25,0.15),(2.26,9999,0.20)]
Y5_BREAKS = [(0,20.99,0.00),(21,25,0.02),(25.01,30,0.04),(30.01,9999,0.06)]
Z1_MAP    = {1:0.00,2:0.25,3:0.50,4:0.75,5:1.00,6:1.30,7:1.60,8:2.00}
Z2_BREAKS = [(0,2,1.00),(2.01,3,0.75),(3.01,4,0.50),(4.01,5,0.25),(5.01,999,0.00)]
Z3_OPTIONS = {
    "0 – 1,000        (Z3=0.00)":0.00, "1,001 – 2,000    (Z3=0.20)":0.20,
    "2,001 – 3,000    (Z3=0.30)":0.30, "3,001 – 4,000    (Z3=0.50)":0.50,
    "4,001 – 5,000    (Z3=0.75)":0.75, "5,001 – 6,000    (Z3=1.00)":1.00,
    "6,001 – 7,000    (Z3=1.25)":1.25, "7,001 – 8,000    (Z3=1.50)":1.50,
    "8,001 – 9,000    (Z3=1.75)":1.75, "9,001 – 10,000   (Z3=2.00)":2.00,
    "10,001 – 15,000  (Z3=2.50)":2.50, "15,001+           (Z3=3.00)":3.00,
}
Z4_BREAKS = [(0,6.49,0.00),(6.50,6.99,0.08),(7.00,9999,0.17)]

# =============================================================================
# SECTION B: CALCULATION FUNCTIONS
# =============================================================================
def lookup_range(value, breaks):
    for lo, hi, v in breaks:
        if lo <= value <= hi:
            return v
    return breaks[-1][2]

def calc_Ka_average(x1, cbr, x3, x4_start, x5_width, x6_code,
                    y1_row, y2_shoulder, terrain_code, y5_bridge, n_years):
    X1=x1; X2=lookup_range(cbr,X2_BREAKS); X3=x3
    X5=lookup_range(x5_width,X5_BREAKS);   X6=X6_MAP[x6_code]
    Y1=lookup_range(y1_row,Y1_BREAKS);     Y2=lookup_range(y2_shoulder,Y2_BREAKS)
    Y3=Y3_MAP[terrain_code]; Y4=Y4_MAP[terrain_code]
    Y5=lookup_range(y5_bridge,Y5_BREAKS);  Y6=Y6_MAP[terrain_code]
    ka_list, rows = [], []
    for yr in range(1, n_years+1):
        age = x4_start + (yr-1)
        X4  = lookup_range(age, X4_BREAKS)
        Ka  = 1 + 0.50*(X1+X2+X3+X4+X5+X6+Y1+Y2+Y3+Y4+Y5+Y6)
        ka_list.append(Ka)
        rows.append({"ปี":yr,"อายุ (ปี)":age,"X4":X4,"Ka":round(Ka,4)})
    fixed = {"X1":X1,"X2":X2,"X3":X3,"X5":X5,"X6":X6,
             "Y1":Y1,"Y2":Y2,"Y3":Y3,"Y4":Y4,"Y5":Y5,"Y6":Y6}
    return round(np.mean(ka_list),4), pd.DataFrame(rows), fixed

def calc_Kc(z1_idx, cbr, z3, z4_width,
            y1_row, y2_shoulder, terrain_code, y5_bridge):
    Z1=Z1_MAP.get(z1_idx,0); Z2=lookup_range(cbr,Z2_BREAKS)
    Z3=z3; Z4=lookup_range(z4_width,Z4_BREAKS)
    Y1=lookup_range(y1_row,Y1_BREAKS); Y2=lookup_range(y2_shoulder,Y2_BREAKS)
    Y3=Y3_MAP[terrain_code]; Y4=Y4_MAP[terrain_code]
    Y5=lookup_range(y5_bridge,Y5_BREAKS); Y6=Y6_MAP[terrain_code]
    Kc = 1 + 0.50*(Z1+Z2+Z3+Z4+Y1+Y2+Y3+Y4+Y5+Y6)
    factors = {"Z1":Z1,"Z2":Z2,"Z3":Z3,"Z4":Z4,
               "Y1":Y1,"Y2":Y2,"Y3":Y3,"Y4":Y4,"Y5":Y5,"Y6":Y6}
    return round(Kc,4), factors

# =============================================================================
# SECTION C: LCCA DATA STRUCTURES
# =============================================================================
@dataclass
class MaintAct:
    name: str; unit_cost: float; start_year: int; frequency: int = 0

@dataclass
class RehabAct:
    name: str; unit_cost: float; year: int

@dataclass
class PavAlt:
    name: str; pave_type: str; construction_cost: float; area: float
    maintenance: List[MaintAct]; rehab: List[RehabAct]
    salvage_pct: float = 20.0; enabled: bool = True

def calc_pv(cost, yr, dr):
    return cost * (1+dr)**(-yr) if yr >= 0 else 0.0

def calc_eac(pw, dr, n):
    if n <= 0 or dr <= 0: return 0.0
    return pw * dr*(1+dr)**n / ((1+dr)**n - 1)

def build_cashflow(alt: PavAlt, n: int, dr: float, inc_salvage: bool) -> pd.DataFrame:
    rows, area = [], alt.area
    rehab_yrs = sorted([r.year for r in alt.rehab if r.year <= n])
    rehab_set = set(rehab_yrs)
    # ปีที่ 0
    c0 = alt.construction_cost * area
    rows.append({"ปี":0,"กิจกรรม":"ก่อสร้างเริ่มต้น","ประเภท":"ก่อสร้าง",
                 "ต้นทุน/หน่วย":alt.construction_cost,"ต้นทุนตามปี":c0,
                 "PW_factor":1.0,"มูลค่าปัจจุบัน":c0})
    # บำรุงรักษา
    for m in alt.maintenance:
        if m.frequency > 0:
            cps = [0] + rehab_yrs
            for idx, cp in enumerate(cps):
                end = cps[idx+1] if idx+1 < len(cps) else n+1
                yr  = cp + m.frequency
                while yr < end and yr <= n:
                    if yr not in rehab_set:
                        c = m.unit_cost*area; pwf=(1+dr)**(-yr)
                        rows.append({"ปี":yr,"กิจกรรม":m.name,"ประเภท":"บำรุงรักษา",
                                     "ต้นทุน/หน่วย":m.unit_cost,"ต้นทุนตามปี":c,
                                     "PW_factor":pwf,"มูลค่าปัจจุบัน":c*pwf})
                    yr += m.frequency
        else:
            if m.start_year <= n and m.start_year not in rehab_set:
                c=m.unit_cost*area; pwf=(1+dr)**(-m.start_year)
                rows.append({"ปี":m.start_year,"กิจกรรม":m.name,"ประเภท":"บำรุงรักษา",
                             "ต้นทุน/หน่วย":m.unit_cost,"ต้นทุนตามปี":c,
                             "PW_factor":pwf,"มูลค่าปัจจุบัน":c*pwf})
    # ฟื้นฟู
    last_cost, last_yr = alt.construction_cost*area, 0
    for r in alt.rehab:
        if r.year <= n:
            c=r.unit_cost*area; pwf=(1+dr)**(-r.year)
            rows.append({"ปี":r.year,"กิจกรรม":r.name,"ประเภท":"ฟื้นฟูสภาพ",
                         "ต้นทุน/หน่วย":r.unit_cost,"ต้นทุนตามปี":c,
                         "PW_factor":pwf,"มูลค่าปัจจุบัน":c*pwf})
            last_cost, last_yr = c, r.year
    # มูลค่าซาก
    if inc_salvage:
        life = {"Flexible":15,"AC":15,"JPCP":20,"JRCP":20,"CRCP":25}
        exp  = next((v for k,v in life.items() if k in alt.pave_type), 20)
        dep  = last_cost*(1-alt.salvage_pct/100)/exp
        sv   = max(last_cost - dep*(n-last_yr), last_cost*alt.salvage_pct/100)
        pwf  = (1+dr)**(-n)
        rows.append({"ปี":n,"กิจกรรม":"มูลค่าซาก","ประเภท":"มูลค่าซาก",
                     "ต้นทุน/หน่วย":-sv/area,"ต้นทุนตามปี":-sv,
                     "PW_factor":pwf,"มูลค่าปัจจุบัน":-sv*pwf})
    return pd.DataFrame(rows).sort_values(["ปี","กิจกรรม"]).reset_index(drop=True)

def analyze_lcca(alts, n, dr, inc_salvage):
    rows, cf_dict = [], {}
    for alt in [a for a in alts if a.enabled]:
        cf   = build_cashflow(alt, n, dr, inc_salvage)
        cf_dict[alt.name] = cf
        pw   = cf["มูลค่าปัจจุบัน"].sum()
        eac  = calc_eac(pw, dr, n)
        area = alt.area
        rows.append({
            "ทางเลือก": alt.name,
            "ประเภทผิวทาง": alt.pave_type,
            "พื้นที่ (ตร.ม./กม.)": area,
            "ต้นทุนก่อสร้าง (บาท/ตร.ม.)": alt.construction_cost,
            "ต้นทุนก่อสร้าง (ล้านบาท/กม.)": round(alt.construction_cost * area / 1e6, 4),
            "PW_ก่อสร้าง": cf[cf["ประเภท"]=="ก่อสร้าง"]["มูลค่าปัจจุบัน"].sum(),
            "PW_บำรุงรักษา": cf[cf["ประเภท"]=="บำรุงรักษา"]["มูลค่าปัจจุบัน"].sum(),
            "PW_ฟื้นฟูสภาพ": cf[cf["ประเภท"]=="ฟื้นฟูสภาพ"]["มูลค่าปัจจุบัน"].sum(),
            "PW_มูลค่าซาก": cf[cf["ประเภท"]=="มูลค่าซาก"]["มูลค่าปัจจุบัน"].sum(),
            "NPV (บาท/กม.)": pw,
            "NPV (ล้านบาท/กม.)": round(pw / 1e6, 4),
            "EAC (บาท/กม./ปี)": eac,
            "EAC (ล้านบาท/กม./ปี)": round(eac / 1e6, 4),
            "EAC (บาท/ตร.ม./ปี)": eac / area if area > 0 else 0,
        })
    df = pd.DataFrame(rows)
    if len(df) > 0:
        df = df.sort_values("NPV (บาท/กม.)").reset_index(drop=True)
        df.insert(0,"อันดับ", range(1, len(df)+1))
    return df, cf_dict

def calc_breakeven(cf_dict, n, dr):
    """คำนวณ Breakeven Year ระหว่างทุกคู่ทางเลือก"""
    names = list(cf_dict.keys())
    results = []
    for a, b in combinations(names, 2):
        cf_a = cf_dict[a]; cf_b = cf_dict[b]
        # สร้าง cumulative PW ปีต่อปี (ปี 0..n)
        cum_a, cum_b = 0.0, 0.0
        be_yr = None
        prev_diff = None
        for yr in range(0, n+1):
            rows_a = cf_a[cf_a["ปี"]==yr]["มูลค่าปัจจุบัน"].sum()
            rows_b = cf_b[cf_b["ปี"]==yr]["มูลค่าปัจจุบัน"].sum()
            cum_a += rows_a; cum_b += rows_b
            diff = cum_a - cum_b
            if prev_diff is not None and prev_diff * diff < 0:
                be_yr = yr
                break
            prev_diff = diff
        results.append({
            "คู่เปรียบเทียบ": f"{a} vs {b}",
            "Breakeven Year": be_yr if be_yr else f">{n}",
            "หมายเหตุ": f"{b} คุ้มกว่า {a} หลังปีที่ {be_yr}" if be_yr else f"ไม่มี crossover ใน {n} ปี"
        })
    return pd.DataFrame(results)

def build_cumulative(cf_dict, n, dr):
    """สร้าง DataFrame cumulative PW รายปีสำหรับทุกทางเลือก"""
    rows = []
    for name, cf in cf_dict.items():
        cum = 0.0
        for yr in range(0, n+1):
            cum += cf[cf["ปี"]==yr]["มูลค่าปัจจุบัน"].sum()
            rows.append({"ปี":yr, "ทางเลือก":name, "Cumulative NPV (บาท)":cum})
    return pd.DataFrame(rows)

# =============================================================================
# SECTION D: SESSION STATE
# =============================================================================
def init_state():
    d = {
        # TAB1
        "project_name": "โครงการก่อสร้างทางหลวง",
        "project_km_total": 1.0,
        "lane_width": 3.50,
        "lanes_per_dir": 2,
        "shoulder_out": 2.50,
        "shoulder_in": 1.50,
        "road_total_width": 22.0,
        "road_area_sqm": 22000.0,
        "cost_ac":0.0,"cost_jpcp":0.0,"cost_jrcp":0.0,"cost_crcp":0.0,
        "thick_ac":18.0,"thick_jpcp":28.0,"thick_jrcp":25.0,"thick_crcp":25.0,
        "base_sec_routine":"3.8","base_sec_lcca":"3.9",
        # TAB2 shared
        "cbr_shared": 3.0,
        "y_row": 40.0,"y_shoulder":1.75,
        "y_terrain": TERRAIN_KEYS[0],"y_bridge":0.0,
        # TAB2 AC
        "ac_x1_key": list(X1_MAP.keys())[0],
        "ac_x3_key": list(X3_OPTIONS.keys())[0],
        "ac_x4_age": 0,"ac_x5_width":7.0,
        "ac_x6_terrain": TERRAIN_KEYS[0],
        "ac_Na":35000.0,"ac_Km":1.0,
        # TAB2 Concrete
        "cc_z1_idx":1,
        "cc_z3_key": list(Z3_OPTIONS.keys())[0],
        "cc_z4_width":7.0,"cc_Nc":35000.0,"cc_Km":1.0,
        # Results TAB2→3
        "ka_avg":None,"kc_val":None,
        "routine_ac_sqm":None,"routine_cc_sqm":None,
        "routine_ac_km":None,"routine_cc_km":None,
        # TAB3
        "lcca_n":20,"lcca_dr":0.06,"lcca_salvage":True,
        "lcca_alternatives":None,
        # Flags
        "tab1_done":False,"tab2_done":False,"tab3_done":False,
        "tab2_dirty":False,
        "show_gravel":False,
        "json_version":0,
    }
    for k,v in d.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_state()
ss = st.session_state

# =============================================================================
# SECTION E: PROGRESS INDICATOR
# =============================================================================
def render_progress():
    t1 = "prog-done" if ss["tab1_done"] else "prog-idle"
    t2 = "prog-done" if ss["tab2_done"] else ("prog-warn" if ss["tab2_dirty"] else "prog-idle")
    t3 = "prog-done" if ss["tab3_done"] else "prog-idle"
    t4 = "prog-done" if ss["tab3_done"] else "prog-idle"
    i1 = "✅" if ss["tab1_done"] else "⭕"
    i2 = "✅" if ss["tab2_done"] else ("⚠️" if ss["tab2_dirty"] else "⭕")
    i3 = "✅" if ss["tab3_done"] else "⭕"
    i4 = "📄" if ss["tab3_done"] else "🔒"
    st.markdown(f"""
    <div class="progress-wrap">
      <div class="prog-step {t1}">{i1} TAB 1<br><small>ข้อมูลโครงการ</small></div>
      <div class="prog-arrow">›</div>
      <div class="prog-step {t2}">{i2} TAB 2<br><small>Routine Cost</small></div>
      <div class="prog-arrow">›</div>
      <div class="prog-step {t3}">{i3} TAB 3<br><small>LCCA</small></div>
      <div class="prog-arrow">›</div>
      <div class="prog-step {t4}">{i4} TAB 4<br><small>Word Report</small></div>
    </div>""", unsafe_allow_html=True)

# =============================================================================
# SECTION F: WORD REPORT HELPERS
# =============================================================================
def set_run_font(run, font="TH SarabunPSK", size=None, bold=False):
    if size is None: size=Pt(14)
    run.font.name=font; run.font.size=size; run.font.bold=bold
    rPr=run._r.get_or_add_rPr()
    rF=rPr.find(qn("w:rFonts"))
    if rF is None:
        rF=OxmlElement("w:rFonts"); rPr.insert(0,rF)
    rF.set(qn("w:ascii"),font); rF.set(qn("w:hAnsi"),font); rF.set(qn("w:cs"),font)

def add_thai_para(doc, text="", bold=False, first_indent=True):
    p=doc.add_paragraph()
    pPr=p._p.get_or_add_pPr()
    jc=OxmlElement("w:jc"); jc.set(qn("w:val"),"thaiDistribute"); pPr.append(jc)
    if first_indent:
        ind=OxmlElement("w:ind"); ind.set(qn("w:firstLine"),"720"); pPr.append(ind)
    if text:
        run=p.add_run(text); set_run_font(run, bold=bold)
    return p

def add_heading_w(doc, text, level=1):
    p=doc.add_heading(text, level=level)
    for r in p.runs: set_run_font(r, size=Pt(15 if level==1 else 14), bold=True)
    return p

def add_table_w(doc, headers, rows, col_widths=None):
    t=doc.add_table(rows=1, cols=len(headers))
    t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for i,w in enumerate(col_widths): t.columns[i].width=Cm(w)
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.paragraphs[0].clear()
        set_run_font(c.paragraphs[0].add_run(str(h)), bold=True)
    for rd in rows:
        row=t.add_row()
        for i,v in enumerate(rd):
            c=row.cells[i]; c.paragraphs[0].clear()
            set_run_font(c.paragraphs[0].add_run(str(v)))
    return t

def generate_word(summary_df, cf_dict, n, dr, alts,
                  base_sec_routine="3.8", base_sec_lcca="3.9"):
    """
    Word Report รูปแบบที่ปรึกษา
    โครงสร้าง:
      X.X   การคำนวณงบประมาณงานบำรุงปกติ  (base_sec_routine)
      X.X+1 วิเคราะห์ LCCA               (base_sec_lcca)
    """
    doc = WordDocument()
    doc.styles["Normal"].font.name = "TH SarabunPSK"
    doc.styles["Normal"].font.size = Pt(14)
    for sec in doc.sections:
        sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5)
        sec.left_margin=Cm(3.0); sec.right_margin=Cm(2.5)

    # ── Section counter helpers ──────────────────────────────────────────────
    def make_sc(base):
        parts = base.strip().split(".")
        class SC:
            prefix = base.strip()
            h1 = 0; h2 = 0
        return SC

    def next_h1(SC, title, doc=doc):
        SC.h1 += 1; SC.h2 = 0
        add_heading_w(doc, f"{SC.prefix}.{SC.h1}  {title}", 1)

    def next_h2(SC, title, doc=doc):
        SC.h2 += 1
        add_heading_w(doc, f"{SC.prefix}.{SC.h1}.{SC.h2}  {title}", 2)

    def section_title(base, title, doc=doc):
        """หัวข้อระดับบน เช่น 3.8  การคำนวณงบประมาณงานบำรุงปกติ"""
        add_heading_w(doc, f"{base}  {title}", 1)

    SCr = make_sc(base_sec_routine)
    SCl = make_sc(base_sec_lcca)

    # ── ปก ──────────────────────────────────────────────────────────────────
    for txt, sz, bold in [
        ("รายการคำนวณการวิเคราะห์ต้นทุนตลอดอายุการใช้งานผิวทาง", 18, True),
        ("Life-Cycle Cost Analysis for Pavement Alternatives", 16, True),
        (ss["project_name"], 15, True),
        (f"วันที่จัดทำ: {datetime.now().strftime('%d/%m/%Y %H:%M')}", 13, False),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(txt), size=Pt(sz), bold=bold)
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # PART 1: การคำนวณงบประมาณงานบำรุงปกติ
    # ════════════════════════════════════════════════════════════════════════
    section_title(base_sec_routine, "การคำนวณงบประมาณงานบำรุงปกติ")
    add_thai_para(doc,
        "การคำนวณงบประมาณงานบำรุงปกติใช้วิธีการกำหนดค่าสัมประสิทธิ์ปรับแก้ (K) "
        "เพื่อสะท้อนลักษณะเฉพาะของแต่ละสายทาง ได้แก่ ประเภทและสภาพผิวทาง "
        "ความสามารถรับแรงของดินคันทาง ปริมาณจราจร อายุการใช้งาน ขนาดเรขาคณิตของถนน "
        "สภาพภูมิประเทศ รวมถึงองค์ประกอบของงานบำรุงรักษาที่เกี่ยวเนื่อง "
        "โดยแบ่งการคำนวณออกเป็น 2 ประเภทตามลักษณะผิวทาง ได้แก่ "
        "ผิวแอสฟัลท์คอนกรีต (Ka) และผิวคอนกรีตซีเมนต์ (Kc)")
    doc.add_paragraph()

    # X.X.1 อัตราค่าบำรุงมาตรฐาน
    next_h1(SCr, "อัตราค่าบำรุงมาตรฐาน (N) และค่า Factor วัสดุ (Km)")
    add_thai_para(doc,
        "อัตราค่าบำรุงมาตรฐาน (N) คือค่าใช้จ่ายพื้นฐานต่อกิโลเมตรต่อปี "
        "สำหรับผิวทางแต่ละประเภทก่อนการปรับแก้ด้วยค่าสัมประสิทธิ์ K "
        "โดยมีค่า Factor วัสดุ (Km) ใช้ปรับตามราคาวัสดุปัจจุบัน ดังแสดงในตาราง")
    add_table_w(doc, ["ประเภทผิวทาง","N มาตรฐาน (บาท/กม./ปี)","Km วัสดุ"], [
        ["ผิวแอสฟัลท์ (Ka)", f"{ss['ac_Na']:,.0f}", f"{ss['ac_Km']:.3f}"],
        ["ผิวคอนกรีต (Kc)",  f"{ss['cc_Nc']:,.0f}", f"{ss['cc_Km']:.3f}"],
    ], col_widths=[6, 5, 5])
    doc.add_paragraph()

    # X.X.2 สูตรและค่า Factor
    next_h1(SCr, "สูตรการคำนวณและค่า Factor")
    add_thai_para(doc,
        "ค่าสัมประสิทธิ์ปรับแก้ (K) คำนวณจากผลรวมของค่า Factor ที่สะท้อนลักษณะของสายทาง "
        "โดยแบ่งเป็น Factor X (เกี่ยวกับผิวทางและปริมาณจราจร) "
        "และ Factor Y (เกี่ยวกับเขตทาง ภูมิประเทศ และงานบำรุงรักษาที่เกี่ยวเนื่อง)")

    # X.X.2.1 Ka
    next_h2(SCr, "ผิวแอสฟัลท์ (Ka)")
    add_thai_para(doc,
        "ค่าสัมประสิทธิ์ปรับแก้สำหรับผิวแอสฟัลท์คอนกรีต (Ka) คำนวณตามสมการ",
        first_indent=False)
    p_eq = doc.add_paragraph(); p_eq.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(p_eq.add_run(
        "Ka = 1 + 0.50 × (X1 + X2 + X3 + X4 + X5 + X6 + Y1 + Y2 + Y3 + Y4 + Y5 + Y6)"),
        bold=True, size=Pt(13))
    add_thai_para(doc, "ผลการคำนวณค่า Factor สำหรับผิวแอสฟัลท์แสดงดังนี้", first_indent=False)

    ka_fixed = ss.get("_ka_fixed", {})
    ka_df    = ss.get("_ka_df")
    ka_avg   = ss.get("ka_avg", 0)
    if ka_fixed:
        fac_rows = [
            ["X1","ลักษณะผิวทางและพื้นทาง",        f"{ka_fixed.get('X1',0):.4f}"],
            ["X2","CBR ดินเดิม",                    f"{ka_fixed.get('X2',0):.4f}"],
            ["X3","ปริมาณจราจร AADT",               f"{ka_fixed.get('X3',0):.4f}"],
            ["X4","อายุบริการ (เฉลี่ย)",             f"(รายปี — ดูตารางด้านล่าง)"],
            ["X5","ความกว้างผิวทาง",                 f"{ka_fixed.get('X5',0):.4f}"],
            ["X6","ภูมิประเทศ",                      f"{ka_fixed.get('X6',0):.4f}"],
            ["Y1","ความกว้างเขตทาง",                 f"{ka_fixed.get('Y1',0):.4f}"],
            ["Y2","ไหล่ทางกว้างสุด 1 ข้าง",         f"{ka_fixed.get('Y2',0):.4f}"],
            ["Y3","จราจรสงเคราะห์",                  f"{ka_fixed.get('Y3',0):.4f}"],
            ["Y4","ท่อระบายน้ำ",                     f"{ka_fixed.get('Y4',0):.4f}"],
            ["Y5","สะพาน",                           f"{ka_fixed.get('Y5',0):.4f}"],
            ["Y6","ทำความสะอาดระบาย",                f"{ka_fixed.get('Y6',0):.4f}"],
            ["Ka เฉลี่ย","ค่าสัมประสิทธิ์เฉลี่ยตลอด analysis period", f"{ka_avg:.4f}"],
        ]
        add_table_w(doc, ["Factor","คำอธิบาย","ค่าที่ใช้"], fac_rows,
                    col_widths=[2, 8, 3])
        doc.add_paragraph()

    # ตาราง Ka รายปี
    if ka_df is not None and len(ka_df) > 0:
        add_thai_para(doc, "ตาราง Ka รายปี (X4 เปลี่ยนตามอายุ):", bold=True, first_indent=False)
        ka_rows = [[str(int(r["ปี"])), str(int(r["อายุ (ปี)"])),
                    f"{r['X4']:.2f}", f"{r['Ka']:.4f}"]
                   for _, r in ka_df.iterrows()]
        add_table_w(doc, ["ปีที่","อายุ (ปี)","X4","Ka"], ka_rows,
                    col_widths=[2.5, 3, 3, 3])
        doc.add_paragraph()

    # X.X.2.2 Kc
    next_h2(SCr, "ผิวคอนกรีต (Kc)")
    add_thai_para(doc,
        "ค่าสัมประสิทธิ์ปรับแก้สำหรับผิวคอนกรีต (Kc) คำนวณตามสมการ",
        first_indent=False)
    p_eq2 = doc.add_paragraph()
    set_run_font(p_eq2.add_run(
        "Kc = 1 + 0.50 × (Z1 + Z2 + Z3 + Z4 + Y1 + Y2 + Y3 + Y4 + Y5 + Y6)"),
        bold=True, size=Pt(13))
    add_thai_para(doc, "ผลการคำนวณค่า Factor สำหรับผิวคอนกรีตแสดงดังนี้", first_indent=False)
    kc_fac = ss.get("_kc_fac", {})
    kc_val = ss.get("kc_val", 0)
    if kc_fac:
        fac_rows_c = [
            ["Z1","ดัชนีสภาพผิวทาง",           f"{kc_fac.get('Z1',0):.4f}"],
            ["Z2","CBR ดินคันทาง",              f"{kc_fac.get('Z2',0):.4f}"],
            ["Z3","ปริมาณจราจร AADT",           f"{kc_fac.get('Z3',0):.4f}"],
            ["Z4","ความกว้างผิวทาง",             f"{kc_fac.get('Z4',0):.4f}"],
            ["Y1","ความกว้างเขตทาง",             f"{kc_fac.get('Y1',0):.4f}"],
            ["Y2","ไหล่ทางกว้างสุด 1 ข้าง",     f"{kc_fac.get('Y2',0):.4f}"],
            ["Y3","จราจรสงเคราะห์",              f"{kc_fac.get('Y3',0):.4f}"],
            ["Y4","ท่อระบายน้ำ",                 f"{kc_fac.get('Y4',0):.4f}"],
            ["Y5","สะพาน",                       f"{kc_fac.get('Y5',0):.4f}"],
            ["Y6","ทำความสะอาดระบาย",            f"{kc_fac.get('Y6',0):.4f}"],
            ["Kc","ค่าสัมประสิทธิ์คอนกรีต",     f"{kc_val:.4f}"],
        ]
        add_table_w(doc, ["Factor","คำอธิบาย","ค่าที่ใช้"], fac_rows_c,
                    col_widths=[2, 8, 3])
    doc.add_paragraph()

    # X.X.3 วิธีการคำนวณงบประมาณ
    next_h1(SCr, "วิธีการคำนวณงบประมาณงานบำรุงปกติรายปี")
    add_thai_para(doc,
        "งบประมาณงานบำรุงปกติรายปีของแต่ละสายทางคำนวณจากผลคูณของระยะทาง "
        "ค่าสัมประสิทธิ์ปรับแก้ (K) ค่า Factor วัสดุ (Km) และอัตราค่าบำรุงมาตรฐาน (N) "
        "โดยปัดค่าที่ได้เป็นหลักร้อยบาทตามแนวทางของกรมทางหลวง", first_indent=False)
    for formula in [
        "งบประมาณ (บาท/ปี) = ระยะทาง (กม.) × K × Km × N  (ปัดเป็นหลักร้อย)",
        "ระยะเทียบเท่า (กม.) = ระยะจริง (กม.) × (จำนวนช่องจราจร / 2)",
        "Workload (หน่วย) = ระยะเทียบเท่า (กม.) × K'",
    ]:
        p_f = doc.add_paragraph(); p_f.paragraph_format.left_indent = Cm(1)
        set_run_font(p_f.add_run(formula), bold=True, size=Pt(13))
    doc.add_paragraph()

    # X.X.4 ผลการคำนวณงบประมาณ (ตาราง Ka/Kc)
    next_h1(SCr, "ผลการคำนวณค่าบำรุงรักษาประจำปี")
    add_thai_para(doc,
        "ตารางต่อไปนี้แสดงผลการคำนวณค่าบำรุงรักษาประจำปีสำหรับผิวทางแต่ละประเภท "
        "โดยแสดงทั้งหน่วย บาท/ตร.ม./ปี และ บาท/กม./ปี เพื่อประกอบการวิเคราะห์ LCCA")
    mrows = []
    r_ac_km  = ss.get("routine_ac_km", 0)
    r_cc_km  = ss.get("routine_cc_km", 0)
    r_ac_sqm = ss.get("routine_ac_sqm", 0)
    r_cc_sqm = ss.get("routine_cc_sqm", 0)
    if ss.get("ka_avg"):
        mrows.append(["AC (ลาดยาง)", f"Ka = {ss['ka_avg']:.4f}",
                      f"{r_ac_sqm:.4f}", f"{r_ac_km:,.2f}", f"{r_ac_km/1e6:.4f}"])
    if ss.get("kc_val"):
        mrows.append(["Concrete", f"Kc = {ss['kc_val']:.4f}",
                      f"{r_cc_sqm:.4f}", f"{r_cc_km:,.2f}", f"{r_cc_km/1e6:.4f}"])
    if mrows:
        add_table_w(doc,
            ["ผิวทาง","K เฉลี่ย","บาท/ตร.ม./ปี","บาท/กม./ปี","ล้านบาท/กม./ปี"],
            mrows, col_widths=[3, 3, 3.5, 4, 4])
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # PART 2: วิเคราะห์ LCCA
    # ════════════════════════════════════════════════════════════════════════
    section_title(base_sec_lcca,
        "วิเคราะห์ต้นทุนตลอดอายุการใช้งานผิวทาง (Life-Cycle Cost Analysis)")
    add_thai_para(doc,
        "การวิเคราะห์ต้นทุนตลอดอายุการใช้งาน (LCCA) "
        "เป็นเครื่องมือทางเศรษฐศาสตร์วิศวกรรมที่ใช้เปรียบเทียบทางเลือกการลงทุนต่างๆ "
        "โดยพิจารณาต้นทุนทั้งหมดตลอดอายุการใช้งานของโครงการ ประกอบด้วย "
        "ต้นทุนก่อสร้างเริ่มต้น ต้นทุนบำรุงรักษา ต้นทุนฟื้นฟูสภาพ และมูลค่าซากปลายทาง "
        "โดยแปลงต้นทุนทั้งหมดมาเป็นมูลค่าปัจจุบัน (Present Worth) เพื่อเปรียบเทียบในฐานเดียวกัน")
    doc.add_paragraph()

    # X.X.1 ทฤษฎี
    next_h1(SCl, "ทฤษฎี Life-Cycle Cost Analysis (LCCA)")
    add_thai_para(doc,
        "LCCA เป็นวิธีการประเมินและเปรียบเทียบต้นทุนตลอดวงจรชีวิตของทางเลือกต่างๆ "
        "ตามมาตรฐาน FHWA-SA-98-079 และ AASHTO Guide for Design of Pavement Structures "
        "โดยคำนึงถึงมูลค่าของเงินตามเวลา (Time Value of Money) "
        "ผ่านการใช้อัตราคิดลด (Discount Rate) ในการแปลงต้นทุนในอนาคตมาเป็นมูลค่าปัจจุบัน")
    doc.add_paragraph()

    # X.X.2 สูตร PW
    next_h1(SCl, "สูตรมูลค่าปัจจุบัน (Present Worth)")
    add_thai_para(doc, "สูตรแปลงต้นทุนในอนาคตมาเป็นมูลค่าปัจจุบัน:", first_indent=False)
    p_pw = doc.add_paragraph(); p_pw.paragraph_format.left_indent = Cm(1)
    set_run_font(p_pw.add_run("PW = FV × (1 + i)^(-n)"), bold=True, size=Pt(13))
    for line in [
        "PW  = มูลค่าปัจจุบัน (Present Worth)",
        "FV  = มูลค่าอนาคต (Future Value)",
        f"i   = อัตราคิดลด (Discount Rate) = {dr*100:.1f}%",
        "n   = จำนวนปีนับจากปัจจุบัน",
    ]:
        p_l = doc.add_paragraph(); p_l.paragraph_format.left_indent = Cm(1.5)
        set_run_font(p_l.add_run(line), size=Pt(13))
    doc.add_paragraph()

    # X.X.3 สูตร EAC
    next_h1(SCl, "สูตรต้นทุนเฉลี่ยรายปี (EAC)")
    add_thai_para(doc, "สูตรแปลงมูลค่าปัจจุบันรวมเป็นต้นทุนเฉลี่ยต่อปี:", first_indent=False)
    p_eac = doc.add_paragraph(); p_eac.paragraph_format.left_indent = Cm(1)
    set_run_font(p_eac.add_run("EAC = PW × [i × (1 + i)^n] / [(1 + i)^n - 1]"),
                 bold=True, size=Pt(13))
    for line in [
        "EAC = ต้นทุนเฉลี่ยรายปี (Equivalent Annual Cost)",
        "PW  = มูลค่าปัจจุบันรวม",
        f"i   = อัตราคิดลด = {dr*100:.1f}%",
        f"n   = ระยะเวลาวิเคราะห์ = {n} ปี",
    ]:
        p_l = doc.add_paragraph(); p_l.paragraph_format.left_indent = Cm(1.5)
        set_run_font(p_l.add_run(line), size=Pt(13))
    doc.add_paragraph()

    # X.X.4 ทางเลือกที่วิเคราะห์
    next_h1(SCl, "ทางเลือกผิวทางที่วิเคราะห์")
    add_thai_para(doc,
        f"การวิเคราะห์ครอบคลุมผิวทาง {len([a for a in alts if a.enabled])} ทางเลือก "
        f"สำหรับถนนความกว้าง {ss.get('road_total_width',22):.2f} ม. "
        f"ระยะวิเคราะห์ {n} ปี ที่อัตราคิดลด {dr*100:.1f}% ต่อปี "
        f"พื้นที่คำนวณ {ss.get('road_total_width',22)*1000:,.0f} ตร.ม./กม.")
    thick_map = {"Flexible":ss.get("thick_ac",18), "AC":ss.get("thick_ac",18),
                 "JPCP":ss.get("thick_jpcp",28), "JRCP":ss.get("thick_jrcp",25),
                 "CRCP":ss.get("thick_crcp",25)}
    alt_rows = []
    for alt in [a for a in alts if a.enabled]:
        th = next((v for k,v in thick_map.items() if k in alt.pave_type), 0)
        alt_rows.append([alt.name, alt.pave_type, f"{th:.1f}",
                         f"{alt.area:,.0f}", f"{alt.construction_cost:,.2f}",
                         f"{alt.construction_cost*alt.area/1e6:.4f}"])
    add_table_w(doc,
        ["ทางเลือก","ประเภท","ความหนา (ซม.)","พื้นที่ (ตร.ม./กม.)",
         "ต้นทุนก่อสร้าง (บาท/ตร.ม.)","ต้นทุนก่อสร้าง (ล้านบาท/กม.)"],
        alt_rows, col_widths=[3.5,2.5,2.5,3,3.5,3.5])
    doc.add_paragraph()

    # X.X.5 ผลการวิเคราะห์
    next_h1(SCl, "ผลการวิเคราะห์ LCCA")
    add_thai_para(doc,
        f"ตารางต่อไปนี้แสดงผลการวิเคราะห์ต้นทุนตลอดอายุการใช้งาน "
        f"สำหรับระยะเวลาวิเคราะห์ {n} ปี ที่อัตราคิดลด {dr*100:.1f}% ต่อปี")
    if len(summary_df) > 0:
        trows = []
        for _, r in summary_df.iterrows():
            trows.append([
                str(int(r["อันดับ"])), r["ทางเลือก"], r["ประเภทผิวทาง"],
                f"{r['NPV (บาท/กม.)']:,.0f}",
                f"{r['NPV (ล้านบาท/กม.)']:,.4f}",
                f"{r['EAC (บาท/กม./ปี)']:,.0f}",
                f"{r['EAC (บาท/ตร.ม./ปี)']:,.2f}",
            ])
        add_table_w(doc,
            ["อันดับ","ทางเลือก","ประเภท",
             "NPV (บาท/กม.)","NPV (ล้านบ./กม.)",
             "EAC (บาท/กม./ปี)","EAC (บ./ตร.ม./ปี)"],
            trows, col_widths=[1.2,3,2,4,3.5,4,3])
        doc.add_paragraph()

    # X.X.6 องค์ประกอบต้นทุน
    next_h1(SCl, "องค์ประกอบต้นทุน (มูลค่าปัจจุบัน)")
    add_thai_para(doc,
        "ตารางต่อไปนี้แสดงองค์ประกอบมูลค่าปัจจุบันของต้นทุนแต่ละประเภท "
        "แยกตามก่อสร้าง บำรุงรักษา ฟื้นฟูสภาพ และมูลค่าซาก")
    if len(summary_df) > 0:
        comp_rows = []
        for _, r in summary_df.iterrows():
            comp_rows.append([
                r["ทางเลือก"],
                f"{r['PW_ก่อสร้าง']:,.0f}",
                f"{r['PW_บำรุงรักษา']:,.0f}",
                f"{r['PW_ฟื้นฟูสภาพ']:,.0f}",
                f"{r['PW_มูลค่าซาก']:,.0f}",
                f"{r['NPV (บาท/กม.)']:,.0f}",
            ])
        add_table_w(doc,
            ["ทางเลือก","ก่อสร้าง (บาท)","บำรุงรักษา (บาท)",
             "ฟื้นฟูสภาพ (บาท)","มูลค่าซาก (บาท)","รวม NPV (บาท/กม.)"],
            comp_rows, col_widths=[3,3.5,3.5,3,3,4])
        doc.add_paragraph()

    # X.X.7 กระแสเงินสดรายทางเลือก
    next_h1(SCl, "รายละเอียดกระแสเงินสดแต่ละทางเลือก")
    for alt_name, cf in cf_dict.items():
        next_h2(SCl, alt_name)
        pw  = cf["มูลค่าปัจจุบัน"].sum()
        eac = calc_eac(pw, dr, n)
        add_thai_para(doc,
            f"NPV รวม = {pw:,.0f} บาท/กม.  |  EAC = {eac:,.0f} บาท/กม./ปี",
            first_indent=False)
        cf_rows = []
        for _, r in cf.iterrows():
            cf_rows.append([
                str(int(r["ปี"])), r["กิจกรรม"], r["ประเภท"],
                f"{r['ต้นทุน/หน่วย']:,.2f}",
                f"{r['ต้นทุนตามปี']:,.0f}",
                f"{r['PW_factor']:.4f}",
                f"{r['มูลค่าปัจจุบัน']:,.0f}",
            ])
        add_table_w(doc,
            ["ปี","กิจกรรม","ประเภท","บาท/หน่วย",
             "ต้นทุนตามปี (บาท)","PW Factor","มูลค่าปัจจุบัน (บาท)"],
            cf_rows, col_widths=[1,4.5,2.5,2.5,3,2,3.5])
        doc.add_paragraph()

    # X.X.8 สรุปผลการวิเคราะห์ (auto-generate — ย่อหน้าเดียว หน่วย บาท/บาทต่อปี)
    next_h1(SCl, "สรุปผลการวิเคราะห์")
    if len(summary_df) > 0:
        best = summary_df.iloc[0]
        npv_best = best["NPV (บาท/กม.)"]
        eac_best = best["EAC (บาท/กม./ปี)"]

        # สร้างข้อความเปรียบเทียบต่อเนื่องในประโยคเดียว
        compare_parts = []
        for _, r in summary_df.iloc[1:].iterrows():
            pct = (r["NPV (บาท/กม.)"] - npv_best) / r["NPV (บาท/กม.)"] * 100
            compare_parts.append(
                f"ประหยัดกว่า{r['ทางเลือก']} คิดเป็นร้อยละ {pct:.1f}")
        compare_str = " ".join(compare_parts)

        summary_text = (
            f"จากการวิเคราะห์ต้นทุนตลอดวงจรชีวิตของทางเลือกผิวทาง "
            f"{len(summary_df)} ประเภท พบว่า {best['ทางเลือก']} "
            f"เป็นทางเลือกที่มีความคุ้มค่าทางเศรษฐศาสตร์สูงสุด "
            f"โดยมีมูลค่าปัจจุบันรวมเท่ากับ {npv_best:,.0f} บาท "
            f"และต้นทุนเฉลี่ยรายปี (EAC) ที่ {eac_best:,.0f} บาท/ปี "
            f"เมื่อเปรียบเทียบกับทางเลือกอื่น {compare_str} "
            f"ดังนั้น จึงมีข้อเสนอแนะให้เลือกใช้ {best['ทางเลือก']} "
            f"เป็นทางเลือกหลักในการออกแบบ "
            f"เนื่องจากให้ต้นทุนรวมต่ำที่สุดตลอดอายุการใช้งาน {n} ปี "
            f"ของโครงการ{ss.get('project_name','')}"
        )
        add_thai_para(doc, summary_text)

    # ── อ้างอิง ──────────────────────────────────────────────────────────────
    doc.add_paragraph()
    next_h1(SCl, "เอกสารอ้างอิง")
    for ref in [
        "FHWA-SA-98-079: Life-Cycle Cost Analysis in Pavement Design",
        "AASHTO Guide for Design of Pavement Structures (1993)",
        "NCHRP Report 703: Guide for Pavement-Type Selection",
        "คู่มือการคิดค่าปริมาณงานและงานบำรุงปกติ กองบำรุง กรมทางหลวง มกราคม พ.ศ. 2538",
    ]:
        p_r = doc.add_paragraph()
        p_r.paragraph_format.left_indent = Cm(1)
        set_run_font(p_r.add_run(f"- {ref}"), size=Pt(13))

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf


# =============================================================================
# SECTION G: JSON HELPERS
# =============================================================================
def build_json():
    alts=ss.get("lcca_alternatives") or []
    return {
        "app":"LCCA_v2","saved_at":datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "project_name":ss["project_name"],
        "project_km_total":ss["project_km_total"],
        "lane_width":ss["lane_width"],"lanes_per_dir":ss["lanes_per_dir"],
        "shoulder_out":ss["shoulder_out"],"shoulder_in":ss["shoulder_in"],
        "cost_ac":ss["cost_ac"],"cost_jpcp":ss["cost_jpcp"],
        "cost_jrcp":ss["cost_jrcp"],"cost_crcp":ss["cost_crcp"],
        "cbr_shared":ss["cbr_shared"],
        "lcca_n":ss["lcca_n"],"lcca_dr":ss["lcca_dr"],"lcca_salvage":ss["lcca_salvage"],
        "ka_avg":ss.get("ka_avg"),"kc_val":ss.get("kc_val"),
        "routine_ac_sqm":ss.get("routine_ac_sqm"),"routine_cc_sqm":ss.get("routine_cc_sqm"),
        "routine_ac_km":ss.get("routine_ac_km"),"routine_cc_km":ss.get("routine_cc_km"),
        "alternatives":[{
            "name":a.name,"pave_type":a.pave_type,
            "construction_cost":a.construction_cost,"area":a.area,
            "salvage_pct":a.salvage_pct,"enabled":a.enabled,
            "maintenance":[{"name":m.name,"unit_cost":m.unit_cost,
                            "start_year":m.start_year,"frequency":m.frequency}
                           for m in a.maintenance],
            "rehab":[{"name":r.name,"unit_cost":r.unit_cost,"year":r.year}
                     for r in a.rehab],
        } for a in alts],
    }

def load_json(data):
    for k in ["project_name","project_km_total","lane_width","lanes_per_dir",
              "shoulder_out","shoulder_in","cost_ac","cost_jpcp","cost_jrcp","cost_crcp",
              "cbr_shared","lcca_n","lcca_dr","lcca_salvage",
              "ka_avg","kc_val","routine_ac_sqm","routine_cc_sqm",
              "routine_ac_km","routine_cc_km"]:
        if k in data: ss[k]=data[k]
    alts=[]
    for a in data.get("alternatives",[]):
        alts.append(PavAlt(
            name=a["name"],pave_type=a["pave_type"],
            construction_cost=a["construction_cost"],area=a["area"],
            salvage_pct=a.get("salvage_pct",20.0),enabled=a.get("enabled",True),
            maintenance=[MaintAct(m["name"],m["unit_cost"],m["start_year"],m["frequency"])
                         for m in a.get("maintenance",[])],
            rehab=[RehabAct(r["name"],r["unit_cost"],r["year"])
                   for r in a.get("rehab",[])],
        ))
    if alts: ss["lcca_alternatives"]=alts
    ss["tab1_done"]=True; ss["tab2_done"]=bool(ss.get("ka_avg"))

# =============================================================================
# SECTION H: MAIN UI
# =============================================================================
st.markdown('<h2 style="margin-bottom:4px">🛣️ LCCA Pavement Integrated v2.0</h2>', unsafe_allow_html=True)
st.caption("Life-Cycle Cost Analysis for Pavement Alternatives | รศ.ดร.อิทธิพล มีผล | KMUTNB")
render_progress()
st.divider()

tab1, tab2, tab3, tab4 = st.tabs([
    "📋 TAB 1  ข้อมูลโครงการ",
    "🔧 TAB 2  Routine Cost",
    "📊 TAB 3  LCCA Analysis",
    "📄 TAB 4  Word Report",
])

# ─────────────────────────────────────────────────────────────────────────────
# TAB 1
# ─────────────────────────────────────────────────────────────────────────────
with tab1:
    st.markdown('<div class="sec-header">📋 ข้อมูลโครงการและราคาก่อสร้าง</div>', unsafe_allow_html=True)

    # JSON I/O
    cj1, cj2 = st.columns(2)
    with cj1:
        uj = st.file_uploader("📂 โหลดโครงการ (JSON)", type="json", key="jload_t1")
        if uj:
            try:
                load_json(json.load(uj)); st.success("✅ โหลดสำเร็จ"); st.rerun()
            except Exception as e: st.error(f"โหลดไม่ได้: {e}")
    with cj2:
        st.download_button("💾 บันทึก JSON",
            data=json.dumps(build_json(),ensure_ascii=False,indent=2).encode("utf-8"),
            file_name=f"LCCA_{ss['project_name'].replace(' ','_')}.json",
            mime="application/json", key="jdl_t1")

    st.divider()

    # ── ข้อมูลโครงการ ─────────────────────────────────────────────────────
    st.markdown('<div class="sec-header">🏗️ ข้อมูลโครงการ</div>', unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    with c1:
        ss["project_name"] = st.text_input("ชื่อโครงการ", value=ss["project_name"], key="pn_t1")
    with c2:
        ss["project_km_total"] = st.number_input(
            "ระยะทางรวม (กม.)", min_value=0.01, value=float(ss["project_km_total"]),
            step=0.1, format="%.3f", key="km_t1")

    st.divider()

    # ── ขนาดถนน ───────────────────────────────────────────────────────────
    st.markdown('<div class="sec-header-orange">🛤️ ขนาดถนน (คำนวณพื้นที่อัตโนมัติ)</div>', unsafe_allow_html=True)
    rc1, rc2, rc3, rc4 = st.columns(4)
    with rc1:
        ss["lane_width"] = st.number_input("ความกว้างช่องจราจร (ม.)",
            min_value=2.5, max_value=5.0, value=float(ss["lane_width"]),
            step=0.25, format="%.2f", key="lw_t1")
    with rc2:
        ss["lanes_per_dir"] = st.selectbox("ช่องจราจร/ทิศทาง",
            options=[1,2,3,4], index=ss["lanes_per_dir"]-1, key="ld_t1")
    with rc3:
        ss["shoulder_out"] = st.number_input("ไหล่ทางนอก (ม.)",
            min_value=0.0, value=float(ss["shoulder_out"]),
            step=0.25, format="%.2f", key="so_t1")
    with rc4:
        ss["shoulder_in"] = st.number_input("ไหล่ทางใน (ม.)",
            min_value=0.0, value=float(ss["shoulder_in"]),
            step=0.25, format="%.2f", key="si_t1")

    # คำนวณพื้นที่
    lanes_total   = ss["lanes_per_dir"] * 2
    pavement_w    = lanes_total * ss["lane_width"]
    shoulder_w    = 2 * ss["shoulder_out"] + 2 * ss["shoulder_in"]
    total_w       = pavement_w + shoulder_w
    area_per_km   = total_w * 1000
    total_area    = area_per_km * ss["project_km_total"]
    ss["road_total_width"] = total_w
    ss["road_area_sqm"]    = total_area

    st.markdown(f"""
    <div class="road-box">
      🛣️ &nbsp;
      <b>ช่องรวม:</b> {lanes_total} ช่อง &nbsp;|&nbsp;
      <b>ผิวจราจร:</b> {pavement_w:.2f} ม. &nbsp;|&nbsp;
      <b>ไหล่ทาง:</b> {shoulder_w:.2f} ม. &nbsp;|&nbsp;
      <b>กว้างรวม:</b> <span class="road-val">{total_w:.2f} ม.</span>
      <br>📐 &nbsp;
      <b>พื้นที่/กม.:</b> {area_per_km:,.0f} ตร.ม./กม. &nbsp;|&nbsp;
      <b>ระยะทาง:</b> {ss['project_km_total']:.3f} กม. &nbsp;|&nbsp;
      <b>พื้นที่รวม:</b> <span class="road-val">{total_area:,.0f} ตร.ม.</span>
    </div>""", unsafe_allow_html=True)

    st.divider()

    # ── ราคาก่อสร้าง ──────────────────────────────────────────────────────
    st.markdown('<div class="sec-header-green">💰 ราคาโครงสร้างชั้นทาง (บาท/ตร.ม.)</div>', unsafe_allow_html=True)

    # Upload Excel
    ef = st.file_uploader("📤 Upload Excel ราคาก่อสร้าง", type=["xlsx","xls"], key="exc_t1")
    if ef and OPENPYXL_AVAILABLE:
        # ใช้ hash ของไฟล์เพื่อตรวจว่าโหลดแล้วหรือยัง (ป้องกัน rerun ซ้ำ)
        import hashlib
        file_hash = hashlib.md5(ef.read()).hexdigest(); ef.seek(0)
        if ss.get("_excel_hash") != file_hash:
            try:
                df_c = pd.read_excel(ef, header=2)
                cm   = {"AC":"cost_ac","JPCP":"cost_jpcp","JRCP":"cost_jrcp","CRCP":"cost_crcp"}
                loaded = []
                for _,row in df_c.iterrows():
                    k=str(row.iloc[0]).strip().upper()
                    if k in cm and pd.notna(row.iloc[2]) and float(row.iloc[2])>0:
                        ss[cm[k]]=float(row.iloc[2]); loaded.append(k)
                if loaded:
                    ss["_excel_hash"] = file_hash
                    ss["_cost_ver"]   = ss.get("_cost_ver",0) + 1  # บังคับ widget reset
                    st.success(f"✅ โหลดราคาจาก Excel สำเร็จ: {', '.join(loaded)}")
                    st.rerun()
                else:
                    st.warning("⚠️ ไม่พบราคาในไฟล์ — ตรวจสอบ format ว่าตรงกับ Template")
            except Exception as e: st.error(f"อ่าน Excel ไม่ได้: {e}")
        else:
            st.success(f"✅ โหลดราคาจาก Excel แล้ว (AC={ss['cost_ac']:,.2f}, JPCP={ss['cost_jpcp']:,.2f}, JRCP={ss['cost_jrcp']:,.2f}, CRCP={ss['cost_crcp']:,.2f})")

    # Template download
    if OPENPYXL_AVAILABLE:
        buf_t=io.BytesIO(); wb=openpyxl.Workbook(); ws=wb.active
        ws.merge_cells("A1:C1")
        ws["A1"]="ข้อมูลสำหรับวิเคราะห์ LCCA (Life-Cycle Cost Analysis)"
        ws["A1"].font=Font(name="TH SarabunPSK",bold=True,size=14,color="FFFFFF")
        ws["A1"].fill=PatternFill("solid",fgColor="1F4E79")
        ws["A1"].alignment=Alignment(horizontal="center",vertical="center")
        ws["A2"]="💡 กรอกข้อมูลในแถวที่ 4-7 → บันทึกไฟล์ → อัปโหลดในโปรแกรม"
        ws["A2"].font=Font(name="TH SarabunPSK",size=12,italic=True)
        for ci,h in enumerate(["ผิวทาง","ประเภทผิวทาง","ต้นทุนก่อสร้าง (บาท/ตร.ม.)"],1):
            c=ws.cell(row=3,column=ci,value=h)
            c.font=Font(name="TH SarabunPSK",bold=True,size=13,color="FFFFFF")
            c.fill=PatternFill("solid",fgColor="2E75B6")
            c.alignment=Alignment(horizontal="center",vertical="center")
        for ri,(a,b) in enumerate([("AC","ลาดยาง"),("JPCP","คอนกรีต"),
                                   ("JRCP","คอนกรีต"),("CRCP","คอนกรีต")],4):
            ws.cell(row=ri,column=1,value=a).font=Font(name="TH SarabunPSK",bold=True,size=13)
            ws.cell(row=ri,column=2,value=b).font=Font(name="TH SarabunPSK",size=13)
            ws.cell(row=ri,column=3,value="").font=Font(name="TH SarabunPSK",size=13)
        ws.column_dimensions["A"].width=10; ws.column_dimensions["B"].width=18
        ws.column_dimensions["C"].width=30
        wb.save(buf_t); buf_t.seek(0)
        st.download_button("📥 ดาวน์โหลด Excel Template",data=buf_t,
            file_name="LCCA_cost_template.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="tmpl_t1")

    # Manual input — ใช้ _cost_ver เป็น key suffix เพื่อบังคับ Streamlit reset value เมื่อโหลด Excel ใหม่
    _cv = ss.get("_cost_ver", 0)
    st.markdown("**กรอกราคาด้วยตนเอง:**")
    ca, cj, cjr, cc = st.columns(4)
    with ca:
        ss["cost_ac"]   = st.number_input("AC (ลาดยาง)",   min_value=0.0, value=float(ss["cost_ac"]),   step=10.0, format="%.2f", key=f"cac_{_cv}")
    with cj:
        ss["cost_jpcp"] = st.number_input("JPCP (คอนกรีต)",min_value=0.0, value=float(ss["cost_jpcp"]), step=10.0, format="%.2f", key=f"cjp_{_cv}")
    with cjr:
        ss["cost_jrcp"] = st.number_input("JRCP (คอนกรีต)",min_value=0.0, value=float(ss["cost_jrcp"]), step=10.0, format="%.2f", key=f"cjr_{_cv}")
    with cc:
        ss["cost_crcp"] = st.number_input("CRCP (คอนกรีต)",min_value=0.0, value=float(ss["cost_crcp"]), step=10.0, format="%.2f", key=f"ccr_{_cv}")

    # แสดง metric cards
    costs = {"AC":ss["cost_ac"],"JPCP":ss["cost_jpcp"],"JRCP":ss["cost_jrcp"],"CRCP":ss["cost_crcp"]}
    card_colors = ["card-blue","card-orange","card-purple","card-teal"]
    if any(v>0 for v in costs.values()):
        st.markdown("**ราคาก่อสร้างที่กำหนด:**")
        cols = st.columns(4)
        for i,(name,val) in enumerate(costs.items()):
            with cols[i]:
                color = card_colors[i]
                st.markdown(f"""
                <div class="metric-card {color}">
                  <div class="label">💰 {name}</div>
                  <div class="value">{val:,.0f}</div>
                  <div class="sub">บาท/ตร.ม.</div>
                </div>""", unsafe_allow_html=True)

    # อัปเดต progress TAB1
    tab1_ok = (ss["project_name"] and ss["project_km_total"]>0 and
               any(v>0 for v in [ss["cost_ac"],ss["cost_jpcp"],ss["cost_jrcp"],ss["cost_crcp"]]))
    if tab1_ok and not ss["tab1_done"]:
        ss["tab1_done"]=True; st.rerun()
    elif not tab1_ok:
        ss["tab1_done"]=False

    if tab1_ok:
        st.markdown('<span class="badge-ok">✅ TAB 1 ครบ — พร้อมไป TAB 2</span>', unsafe_allow_html=True)
    else:
        st.markdown('<span class="badge-warn">⚠️ กรอกชื่อโครงการ + ระยะทาง + ราคาก่อสร้างอย่างน้อย 1 รายการ</span>', unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# TAB 2
# ─────────────────────────────────────────────────────────────────────────────
with tab2:
    st.markdown('<div class="sec-header">🔧 Routine Cost Calculator</div>', unsafe_allow_html=True)

    if not ss["tab1_done"]:
        st.markdown('<div class="warn-band">⚠️ กรุณากรอกข้อมูลใน TAB 1 ให้ครบก่อน</div>', unsafe_allow_html=True)
    else:
        # ── ตัวแปรร่วม ──────────────────────────────────────────────────────
        st.markdown('<div class="sec-header">📌 ตัวแปรร่วม (CBR + Y1–Y6)</div>', unsafe_allow_html=True)
        yc1, yc2, yc3, yc4, yc5 = st.columns(5)
        with yc1:
            ss["cbr_shared"] = st.number_input("CBR ดินเดิม (%) — ใช้ร่วม AC & Concrete",
                min_value=0.0, max_value=20.0, value=float(ss["cbr_shared"]), step=0.5, key="cbr_t2")
        with yc2:
            ss["y_row"]      = st.number_input("Y1: กว้างเขตทาง (ม.)",
                min_value=0.0, value=float(ss["y_row"]), step=1.0, key="yr_t2")
        with yc3:
            ss["y_shoulder"] = st.number_input("Y2: ไหล่ทางกว้างสุด 1 ข้าง (ม.)",
                min_value=0.0, value=float(ss["y_shoulder"]), step=0.25, key="ys_t2")
        with yc4:
            ss["y_terrain"]  = st.selectbox("Y3/Y4/Y6: ภูมิประเทศ", TERRAIN_KEYS,
                index=TERRAIN_KEYS.index(ss["y_terrain"]), key="yt_t2")
        with yc5:
            ss["y_bridge"]   = st.number_input("Y5: สะพาน (ม./กม.)",
                min_value=0.0, value=float(ss["y_bridge"]), step=1.0, key="yb_t2")

        terrain_code = TERRAIN_MAP[ss["y_terrain"]]

        with st.expander("ดูค่า Y factors"):
            Y1=lookup_range(ss["y_row"],Y1_BREAKS); Y2=lookup_range(ss["y_shoulder"],Y2_BREAKS)
            Y3=Y3_MAP[terrain_code]; Y4=Y4_MAP[terrain_code]
            Y5=lookup_range(ss["y_bridge"],Y5_BREAKS); Y6=Y6_MAP[terrain_code]
            st.dataframe(pd.DataFrame({"Factor":["Y1","Y2","Y3","Y4","Y5","Y6"],
                "คำอธิบาย":["เขตทาง","ไหล่ทาง","จราจรสงเคราะห์","ท่อระบายน้ำ","สะพาน","ระบายน้ำ"],
                "ค่า":[Y1,Y2,Y3,Y4,Y5,Y6]}), hide_index=True)

        st.divider()

        # ── Section A & B ────────────────────────────────────────────────────
        col_a, col_b = st.columns(2)
        with col_a:
            st.markdown('<div class="sec-header">🔵 Section A: ผิวแอสฟัลท์ (Ka)</div>', unsafe_allow_html=True)
            ss["ac_x1_key"]    = st.selectbox("X1: ลักษณะผิวทาง", list(X1_MAP.keys()),
                index=list(X1_MAP.keys()).index(ss["ac_x1_key"]), key="ax1_t2")
            ss["ac_x3_key"]    = st.selectbox("X3: AADT", list(X3_OPTIONS.keys()),
                index=list(X3_OPTIONS.keys()).index(ss["ac_x3_key"]), key="ax3_t2")
            ss["ac_x4_age"]    = st.number_input("X4: อายุปัจจุบัน (ปี) — ปีเริ่มต้น",
                min_value=0, value=int(ss["ac_x4_age"]), step=1, key="ax4_t2")
            ss["ac_x5_width"]  = st.number_input("X5: กว้างผิวทาง (ม.)",
                min_value=4.0, value=float(ss["ac_x5_width"]), step=0.5, key="ax5_t2")
            ss["ac_x6_terrain"]= st.selectbox("X6: ภูมิประเทศ (AC)", TERRAIN_KEYS,
                index=TERRAIN_KEYS.index(ss["ac_x6_terrain"]), key="ax6_t2")
            ss["ac_Na"] = st.number_input("Na (บาท/กม./ปี)",
                min_value=1000.0, value=float(ss["ac_Na"]), step=500.0, key="ana_t2")
            ss["ac_Km"] = st.number_input("Km วัสดุ AC", min_value=0.1,
                value=float(ss["ac_Km"]), step=0.05, format="%.3f", key="akm_t2")

        with col_b:
            st.markdown('<div class="sec-header-orange">🟠 Section B: ผิวคอนกรีต (Kc)</div>', unsafe_allow_html=True)
            ss["cc_z1_idx"]   = st.selectbox("Z1: ดัชนีสภาพผิวทาง (1=ดีมาก…8=แย่มาก)",
                list(range(1,9)), index=ss["cc_z1_idx"]-1, key="cz1_t2")
            ss["cc_z3_key"]   = st.selectbox("Z3: AADT", list(Z3_OPTIONS.keys()),
                index=list(Z3_OPTIONS.keys()).index(ss["cc_z3_key"]), key="cz3_t2")
            ss["cc_z4_width"] = st.number_input("Z4: กว้างผิวทาง (ม.)",
                min_value=4.0, value=float(ss["cc_z4_width"]), step=0.5, key="cz4_t2")
            ss["cc_Nc"] = st.number_input("Nc (บาท/กม./ปี)",
                min_value=1000.0, value=float(ss["cc_Nc"]), step=500.0, key="cnc_t2")
            ss["cc_Km"] = st.number_input("Km วัสดุ Concrete", min_value=0.1,
                value=float(ss["cc_Km"]), step=0.05, format="%.3f", key="ckm_t2")

        st.divider()
        n_lcca = ss.get("lcca_n",20)
        st.markdown(f'<div class="info-band">ℹ️ Ka คำนวณรายปีตลอด <b>{n_lcca} ปี</b> (ตาม Analysis Period TAB 3) — X4 เปลี่ยนตามอายุจริงแต่ละปี แล้วหาค่าเฉลี่ย</div>', unsafe_allow_html=True)

        if st.button("🔄 คำนวณ Ka และ Kc", type="primary", key="calc_t2"):
            X1v  = X1_MAP[ss["ac_x1_key"]]
            X3v  = X3_OPTIONS[ss["ac_x3_key"]]
            X6c  = TERRAIN_MAP[ss["ac_x6_terrain"]]
            Z3v  = Z3_OPTIONS[ss["cc_z3_key"]]

            ka_avg, ka_df, ka_fixed = calc_Ka_average(
                X1v, ss["cbr_shared"], X3v, ss["ac_x4_age"],
                ss["ac_x5_width"], X6c,
                ss["y_row"], ss["y_shoulder"], terrain_code, ss["y_bridge"], n_lcca)

            Kc, kc_fac = calc_Kc(
                ss["cc_z1_idx"], ss["cbr_shared"], Z3v, ss["cc_z4_width"],
                ss["y_row"], ss["y_shoulder"], terrain_code, ss["y_bridge"])

            # บาท/ตร.ม./ปี และ บาท/กม./ปี
            # Na × Ka × Km = บาท/กม./ปี
            r_ac_km = ss["ac_Na"] * ka_avg * ss["ac_Km"]
            r_cc_km = ss["cc_Nc"] * Kc     * ss["cc_Km"]
            # หารด้วยพื้นที่จริงต่อ กม. = ความกว้างถนนรวม (TAB1) × 1,000
            # ไม่ใช่ X5/Z4 ซึ่งเป็นเฉพาะผิวจราจร ไม่รวมไหล่ทาง
            total_w  = ss.get("road_total_width", 22.0)
            r_ac_sqm = r_ac_km / (total_w * 1000)
            r_cc_sqm = r_cc_km / (total_w * 1000)

            # แสดงที่มาของสูตรใน info band
            st.markdown(f'''<div class="info-band">
            📐 <b>สูตร Routine Cost:</b><br>
            AC: Na({ss["ac_Na"]:,.0f}) × Ka({ka_avg:.4f}) × Km({ss["ac_Km"]:.3f}) 
            = <b>{r_ac_km:,.2f} บาท/กม./ปี</b> ÷ {total_w:.1f}ม.×1,000 
            = <b>{r_ac_sqm:.4f} บาท/ตร.ม./ปี</b><br>
            Concrete: Nc({ss["cc_Nc"]:,.0f}) × Kc({Kc:.4f}) × Km({ss["cc_Km"]:.3f}) 
            = <b>{r_cc_km:,.2f} บาท/กม./ปี</b> ÷ {total_w:.1f}ม.×1,000 
            = <b>{r_cc_sqm:.4f} บาท/ตร.ม./ปี</b>
            </div>''', unsafe_allow_html=True)
            ss["ka_avg"]=ka_avg; ss["kc_val"]=Kc
            ss["routine_ac_sqm"]=round(r_ac_sqm,4); ss["routine_cc_sqm"]=round(r_cc_sqm,4)
            ss["routine_ac_km"]=round(r_ac_km,2);   ss["routine_cc_km"]=round(r_cc_km,2)
            ss["_ka_df"]=ka_df; ss["_kc_fac"]=kc_fac; ss["_ka_fixed"]=ka_fixed
            ss["tab2_done"]=True; ss["tab2_dirty"]=False
            st.success("✅ คำนวณสำเร็จ — ส่งค่าไป TAB 3 อัตโนมัติแล้ว")

        # Badge สถานะ
        if ss["tab2_dirty"] and ss["tab2_done"]:
            st.markdown('<span class="badge-warn">⚠️ ข้อมูลเปลี่ยน กรุณากด คำนวณ Ka และ Kc อีกครั้ง</span>', unsafe_allow_html=True)
        elif ss["tab2_done"]:
            st.markdown('<span class="badge-ok">✅ Ka / Kc ส่งไป TAB 3 แล้ว</span>', unsafe_allow_html=True)

        # แสดงผล
        if ss.get("ka_avg") is not None:
            mc1,mc2,mc3,mc4,mc5,mc6 = st.columns(6)
            cards = [
                (mc1,"card-blue","Ka เฉลี่ย",f"{ss['ka_avg']:.4f}",""),
                (mc2,"card-orange","Kc",f"{ss['kc_val']:.4f}",""),
                (mc3,"card-teal","บำรุง AC",f"{ss['routine_ac_sqm']:.4f}","บาท/ตร.ม./ปี"),
                (mc4,"card-teal","บำรุง AC",f"{ss['routine_ac_km']:,.2f}","บาท/กม./ปี"),
                (mc5,"card-purple","บำรุง Concrete",f"{ss['routine_cc_sqm']:.4f}","บาท/ตร.ม./ปี"),
                (mc6,"card-purple","บำรุง Concrete",f"{ss['routine_cc_km']:,.2f}","บาท/กม./ปี"),
            ]
            for col,color,label,val,sub in cards:
                with col:
                    st.markdown(f"""
                    <div class="metric-card {color}">
                      <div class="label">{label}</div>
                      <div class="value">{val}</div>
                      <div class="sub">{sub}</div>
                    </div>""", unsafe_allow_html=True)

            cd1,cd2 = st.columns(2)
            with cd1:
                st.markdown("**Ka รายปี (X4 เปลี่ยนตามอายุ)**")
                st.dataframe(ss["_ka_df"], hide_index=True, use_container_width=True, height=280)
            with cd2:
                st.markdown("**Kc Factors**")
                kf=ss.get("_kc_fac",{})
                st.dataframe(pd.DataFrame({"Factor":list(kf.keys()),
                    "ค่า":[round(v,4) for v in kf.values()]}),
                    hide_index=True, use_container_width=True)

        # Legacy gravel
        with st.expander("⚙️ ผิวลูกรัง (Legacy)"):
            ss["show_gravel"] = st.checkbox("แสดงผิวลูกรัง",value=ss["show_gravel"],key="sg_t2")
            if ss["show_gravel"]:
                st.warning("⚠️ DOH เลิกใช้แล้ว — สำรองไว้สำหรับงานวิจัย/อ้างอิงเท่านั้น")

# ─────────────────────────────────────────────────────────────────────────────
# TAB 3
# ─────────────────────────────────────────────────────────────────────────────
with tab3:
    st.markdown('<div class="sec-header">📊 LCCA Analysis</div>', unsafe_allow_html=True)

    # Warning checks
    warns = []
    if not ss["tab1_done"]: warns.append("TAB 1 ยังไม่ครบ (ชื่อโครงการ / ระยะทาง / ราคา)")
    if not ss["tab2_done"]: warns.append("TAB 2 ยังไม่ได้คำนวณ Ka / Kc")
    if ss["tab2_dirty"]:    warns.append("ข้อมูล TAB 2 เปลี่ยน กรุณาคำนวณ Ka/Kc ใหม่")
    if warns:
        for w in warns:
            st.markdown(f'<div class="warn-band">⚠️ {w}</div>', unsafe_allow_html=True)

    # พารามิเตอร์
    st.markdown('<div class="sec-header">⚙️ พารามิเตอร์การวิเคราะห์</div>', unsafe_allow_html=True)
    pc1,pc2,pc3 = st.columns(3)
    with pc1:
        ss["lcca_n"]  = st.number_input("ระยะเวลาวิเคราะห์ (ปี)",
            min_value=5,max_value=50,value=int(ss["lcca_n"]),step=1,key="ln_t3")
    with pc2:
        ss["lcca_dr"] = st.number_input("อัตราคิดลด (%/ปี)",
            min_value=1.0,max_value=20.0,value=float(ss["lcca_dr"])*100,
            step=0.5,key="ld_t3") / 100.0
    with pc3:
        ss["lcca_salvage"] = st.checkbox("รวมมูลค่าซาก",value=ss["lcca_salvage"],key="ls_t3")

    n=ss["lcca_n"]; dr=ss["lcca_dr"]

    st.divider()

    # Alternatives
    st.markdown('<div class="sec-header-orange">🏗️ ทางเลือกผิวทาง</div>', unsafe_allow_html=True)
    costs = {"AC":ss["cost_ac"],"JPCP":ss["cost_jpcp"],"JRCP":ss["cost_jrcp"],"CRCP":ss["cost_crcp"]}
    r_ac  = ss.get("routine_ac_sqm") or 0.0
    r_cc  = ss.get("routine_cc_sqm") or 0.0
    # ใช้พื้นที่ต่อ กม. = ความกว้างถนนรวม × 1,000 ม.
    # ผลลัพธ์ NPV/EAC จะเป็น บาท/กม.
    area = ss.get("road_total_width", 22.0) * 1000.0

    area_per_km = ss.get("road_total_width", 22.0) * 1000.0
    st.markdown(f'<div class="info-band">📐 พื้นที่คำนวณ = <b>{area_per_km:,.0f} ตร.ม./กม.</b> (กว้าง {ss.get("road_total_width",22.0):.2f} ม. × 1,000 ม.) — ผลลัพธ์ NPV/EAC เป็น <b>บาท/กม.</b> และ <b>ล้านบาท/กม.</b></div>', unsafe_allow_html=True)

    if st.button("🔄 สร้าง/รีเซ็ต Alternatives จาก TAB 1 & 2", key="gen_t3"):
        if not ss["tab1_done"] or not ss["tab2_done"]:
            st.error("กรุณากรอก TAB 1 และคำนวณ TAB 2 ก่อน")
        else:
            alts=[]
            for nm,pt,cost,mc in [
                ("ผิวทางยืดหยุ่น (AC)","Flexible",costs["AC"],r_ac),
                ("JPCP","JPCP",costs["JPCP"],r_cc),
                ("JRCP","JRCP",costs["JRCP"],r_cc),
                ("CRCP","CRCP",costs["CRCP"],r_cc),
            ]:
                if cost<=0: continue
                rehab_yr=max(10,n//2) if "AC" in pt or "Flex" in pt else max(15,int(n*0.75))
                if "AC" in pt or "Flex" in pt:
                    mlist=[MaintAct("บำรุงรักษาประจำปี ",mc,1,1),
                           MaintAct("ฉาบผิว(Seal Coating)",mc*0.8,3,3)]
                    rlist=[RehabAct("Overlay AC 50 มม.",390,rehab_yr)]
                    sv=20.0
                else:
                    mlist=[MaintAct("บำรุงรักษาประจำปี ",mc,1,1),
                           MaintAct("Joint Maintenance",mc*0.5,5,5)]
                    rlist=[]; sv=30.0
                alts.append(PavAlt(nm,pt,cost,area,mlist,rlist,sv))
            ss["lcca_alternatives"]=alts
            st.success(f"✅ สร้าง {len(alts)} ทางเลือก")

    alts=ss.get("lcca_alternatives") or []
    if alts:
        for ai,alt in enumerate(alts):
            with st.expander(f"✏️ {alt.name} | {alt.construction_cost:,.0f} บาท/ตร.ม."):
                e1,e2=st.columns(2)
                with e1:
                    alts[ai].construction_cost=st.number_input("ต้นทุนก่อสร้าง (บาท/ตร.ม.)",
                        min_value=0.0,value=float(alt.construction_cost),step=10.0,key=f"ec_{ai}")
                    alts[ai].salvage_pct=st.number_input("มูลค่าซาก (%)",
                        min_value=0.0,max_value=100.0,value=float(alt.salvage_pct),step=1.0,key=f"es_{ai}")
                    alts[ai].enabled=st.checkbox("เปิดใช้งาน",value=alt.enabled,key=f"ee_{ai}")
                with e2:
                    st.markdown("**บำรุงรักษา:**")
                    for mi,m in enumerate(alt.maintenance):
                        mc1,mc2=st.columns([2,1])
                        with mc1: alts[ai].maintenance[mi].unit_cost=st.number_input(
                            f"{m.name} (บาท/ตร.ม./ปี)",min_value=0.0,value=float(m.unit_cost),step=1.0,key=f"mc_{ai}_{mi}")
                        with mc2: alts[ai].maintenance[mi].frequency=st.number_input(
                            "ความถี่ (ปี/ครั้ง)",min_value=0,value=int(m.frequency),step=1,key=f"mf_{ai}_{mi}")
                    st.markdown("**ฟื้นฟูสภาพ:**")
                    for ri2,r in enumerate(alt.rehab):
                        rc1,rc2=st.columns([2,1])
                        with rc1: alts[ai].rehab[ri2].unit_cost=st.number_input(
                            f"{r.name} (บาท/ตร.ม.)",min_value=0.0,value=float(r.unit_cost),step=10.0,key=f"rc_{ai}_{ri2}")
                        with rc2: alts[ai].rehab[ri2].year=st.number_input(
                            "ดำเนินการปีที่",min_value=1,max_value=n,value=int(r.year),step=1,key=f"ry_{ai}_{ri2}")
        ss["lcca_alternatives"]=alts

        st.divider()
        if st.button("🚀 คำนวณ LCCA", type="primary", key="run_t3"):
            if warns:
                st.error("กรุณาแก้ไข warning ด้านบนก่อน")
            else:
                with st.spinner("กำลังคำนวณ..."):
                    sdf,cfd=analyze_lcca(alts,n,dr,ss["lcca_salvage"])
                    ss["_lcca_sum"]=sdf; ss["_lcca_cf"]=cfd
                    ss["tab3_done"]=True
                    st.success("✅ คำนวณ LCCA สำเร็จ — ดูผลด้านล่าง")

    # ── แสดงผล ──────────────────────────────────────────────────────────────
    sdf=ss.get("_lcca_sum"); cfd=ss.get("_lcca_cf",{})
    if sdf is not None and len(sdf)>0:
        st.markdown('<div class="sec-header-green">🏆 สรุปผล LCCA</div>', unsafe_allow_html=True)
        best=sdf.iloc[0]

        # Best highlight
        st.markdown(f"""
        <div class="best-row">
          🥇 <b>ทางเลือกที่ดีที่สุด: {best['ทางเลือก']}</b> ({best['ประเภทผิวทาง']})<br>
          NPV = <b>{best['NPV (ล้านบาท/กม.)']:,.4f} ล้านบาท/กม.</b> &nbsp;|&nbsp;
          EAC = <b>{best['EAC (ล้านบาท/กม./ปี)']:,.4f} ล้านบาท/กม./ปี</b> &nbsp;|&nbsp;
          EAC = <b>{best['EAC (บาท/ตร.ม./ปี)']:,.2f} บาท/ตร.ม./ปี</b>
        </div>""", unsafe_allow_html=True)

        # ตาราง metric cards ทุกทางเลือก
        cols_res=st.columns(len(sdf))
        card_c=["card-green","card-blue","card-orange","card-purple"]
        for i,(_,row) in enumerate(sdf.iterrows()):
            with cols_res[i]:
                badge="🥇" if i==0 else ("🥈" if i==1 else "🥉")
                st.markdown(f"""
                <div class="metric-card {card_c[i%4]}">
                  <div class="label">{badge} อันดับ {int(row['อันดับ'])} — {row['ทางเลือก']}</div>
                  <div class="value">{row['NPV (ล้านบาท/กม.)']:,.4f}</div>
                  <div class="sub">ล้านบาท/กม. (NPV) | EAC {row['EAC (ล้านบาท/กม./ปี)']:,.4f} ล้านบ./กม./ปี</div>
                </div>""", unsafe_allow_html=True)

        # ตารางสรุปเต็ม
        st.markdown("**ตารางสรุปเปรียบเทียบ:**")
        fmt_cols={
                  "ต้นทุนก่อสร้าง (บาท/ตร.ม.)":"{:,.2f}",
                  "ต้นทุนก่อสร้าง (ล้านบาท/กม.)":"{:,.4f}",
                  "NPV (บาท/กม.)":"{:,.0f}",
                  "NPV (ล้านบาท/กม.)":"{:,.4f}",
                  "EAC (บาท/กม./ปี)":"{:,.0f}",  # legacy
                  "EAC (ล้านบาท/กม./ปี)":"{:,.4f}",
                  "EAC (บาท/ตร.ม./ปี)":"{:,.2f}"}
        show_cols=["อันดับ","ทางเลือก","ประเภทผิวทาง",
                   "ต้นทุนก่อสร้าง (บาท/ตร.ม.)",
                   "ต้นทุนก่อสร้าง (ล้านบาท/กม.)",
                   "NPV (ล้านบาท/กม.)",
                   "EAC (บาท/กม./ปี)","EAC (ล้านบาท/กม./ปี)",
                   "EAC (บาท/ตร.ม./ปี)"]
        st.dataframe(sdf[show_cols].style.format(fmt_cols), hide_index=True, use_container_width=True)

        # กราฟ Stacked Bar NPV
        st.markdown('<div class="sec-header">📊 เปรียบเทียบ NPV แยกประเภทต้นทุน</div>', unsafe_allow_html=True)
        fig_bar=go.Figure()
        color_map={"ก่อสร้าง":"#1565C0","บำรุงรักษา":"#F57C00",
                   "ฟื้นฟูสภาพ":"#C62828","มูลค่าซาก":"#2E7D32"}
        for pt,col_k in [("ก่อสร้าง","PW_ก่อสร้าง"),("บำรุงรักษา","PW_บำรุงรักษา"),
                         ("ฟื้นฟูสภาพ","PW_ฟื้นฟูสภาพ"),("มูลค่าซาก","PW_มูลค่าซาก")]:
            fig_bar.add_trace(go.Bar(name=pt,x=sdf["ทางเลือก"],y=sdf[col_k],
                marker_color=color_map[pt]))
        fig_bar.update_layout(barmode="relative",height=420,
            title="มูลค่าปัจจุบัน (NPV) แยกตามประเภทต้นทุน",
            yaxis_title="บาท",xaxis_title="ทางเลือก",
            paper_bgcolor="rgba(0,0,0,0)",plot_bgcolor="rgba(0,0,0,0)")
        st.plotly_chart(fig_bar, use_container_width=True)

        # กราฟ Cumulative Cost Timeline (2.3)
        st.markdown('<div class="sec-header">📈 Cumulative Cost Timeline</div>', unsafe_allow_html=True)
        cum_df=build_cumulative(cfd,n,dr)
        fig_cum=px.line(cum_df,x="ปี",y="Cumulative NPV (บาท)",color="ทางเลือก",
            markers=True,height=450,
            title=f"ต้นทุนสะสม (Cumulative NPV) ตลอด {n} ปี",
            labels={"Cumulative NPV (บาท)":"Cumulative NPV (บาท)","ปี":"ปีที่"})
        fig_cum.update_layout(paper_bgcolor="rgba(0,0,0,0)",plot_bgcolor="rgba(248,249,250,1)",
            legend=dict(yanchor="top",y=0.99,xanchor="left",x=0.01))
        st.plotly_chart(fig_cum, use_container_width=True)

        # Breakeven Year (2.2)
        st.markdown('<div class="sec-header">⚖️ Breakeven Year Analysis</div>', unsafe_allow_html=True)
        be_df=calc_breakeven(cfd,n,dr)
        if len(be_df)>0:
            st.dataframe(be_df, hide_index=True, use_container_width=True)
            for _,r in be_df.iterrows():
                if isinstance(r["Breakeven Year"], int):
                    st.markdown(f'<div class="info-band">📌 {r["หมายเหตุ"]}</div>', unsafe_allow_html=True)
                else:
                    st.markdown(f'<div class="warn-band">📌 {r["หมายเหตุ"]}</div>', unsafe_allow_html=True)

        # Sensitivity (อัตราคิดลด)
        st.markdown('<div class="sec-header">📉 Sensitivity Analysis — อัตราคิดลด</div>', unsafe_allow_html=True)
        sens_rows=[]
        for dr_i in np.linspace(max(dr-0.03,0.01),dr+0.03,7):
            for alt in [a for a in alts if a.enabled]:
                cf_i=build_cashflow(alt,n,dr_i,ss["lcca_salvage"])
                sens_rows.append({"อัตราคิดลด (%)":round(dr_i*100,1),
                                  "ทางเลือก":alt.name,
                                  "NPV (บาท)":cf_i["มูลค่าปัจจุบัน"].sum()})
        fig_sens=px.line(pd.DataFrame(sens_rows),x="อัตราคิดลด (%)",y="NPV (บาท)",
            color="ทางเลือก",markers=True,height=420,
            title="Sensitivity Analysis — ผลกระทบของอัตราคิดลดต่อ NPV")
        fig_sens.update_layout(paper_bgcolor="rgba(0,0,0,0)",
            plot_bgcolor="rgba(248,249,250,1)")
        st.plotly_chart(fig_sens, use_container_width=True)

        # ตารางกระแสเงินสด
        st.markdown('<div class="sec-header">💰 กระแสเงินสดรายทางเลือก</div>', unsafe_allow_html=True)
        sel=st.selectbox("เลือกทางเลือก",list(cfd.keys()),key="cfsel_t3")
        if sel in cfd:
            cf_s=cfd[sel].copy()
            cf_s["ต้นทุน/หน่วย"]=cf_s["ต้นทุน/หน่วย"].map(lambda x:f"{x:,.2f}")
            cf_s["ต้นทุนตามปี"]=cf_s["ต้นทุนตามปี"].map(lambda x:f"{x:,.0f}")
            cf_s["PW_factor"]=cf_s["PW_factor"].map(lambda x:f"{x:.4f}")
            cf_s["มูลค่าปัจจุบัน"]=cf_s["มูลค่าปัจจุบัน"].map(lambda x:f"{x:,.0f}")
            st.dataframe(cf_s, hide_index=True, use_container_width=True, height=400)

# ─────────────────────────────────────────────────────────────────────────────
# TAB 4
# ─────────────────────────────────────────────────────────────────────────────
with tab4:
    st.markdown('<div class="sec-header">📄 Word Report — รูปแบบที่ปรึกษา</div>', unsafe_allow_html=True)
    if not DOCX_AVAILABLE:
        st.error("❌ pip install python-docx")
    elif not ss.get("tab3_done"):
        st.markdown('<div class="warn-band">⚠️ กรุณาคำนวณ LCCA ใน TAB 3 ก่อน</div>', unsafe_allow_html=True)
    else:
        sdf=ss.get("_lcca_sum"); cfd=ss.get("_lcca_cf",{})
        if sdf is not None and len(sdf)>0:
            st.markdown('<div class="best-row">✅ พร้อมสร้างรายงาน — ข้อมูลครบทุก TAB</div>', unsafe_allow_html=True)

            cols_r=st.columns(4)
            with cols_r[0]:
                ss["base_sec_routine"] = st.text_input(
                    "Base Section — Routine Cost (เช่น 3.8)",
                    value=ss.get("base_sec_routine","3.8"), key="bsr_t4")
            with cols_r[1]:
                # smart default: LCCA = Routine + 1
                try:
                    parts = ss["base_sec_routine"].split(".")
                    auto_lcca = f"{parts[0]}.{int(parts[1])+1}" if len(parts)>1 else str(int(parts[0])+1)
                except: auto_lcca = "3.9"
                ss["base_sec_lcca"] = st.text_input(
                    "Base Section — LCCA (เช่น 3.9)",
                    value=ss.get("base_sec_lcca", auto_lcca), key="bsl_t4")
            with cols_r[2]: st.write("")

            st.dataframe(sdf[["อันดับ","ทางเลือก","NPV (ล้านบาท/กม.)","EAC (ล้านบาท/กม./ปี)","EAC (บาท/ตร.ม./ปี)"]]\
                .style.format({"NPV (ล้านบาท/กม.)":"{:,.4f}","EAC (ล้านบาท/กม./ปี)":"{:,.4f}","EAC (บาท/ตร.ม./ปี)":"{:,.2f}"}),
                hide_index=True,use_container_width=True)

            # ── ความหนา (optional — สำหรับแสดงในรายงานเท่านั้น) ──────────
            with st.expander("📏 ความหนาโครงสร้างชั้นทาง (optional — แสดงในรายงานเท่านั้น ไม่มีผลต่อการคำนวณ)"):
                st.caption("กรอกเพื่อแสดงในตาราง 3.9.4 ของรายงาน ถ้าไม่กรอกจะแสดงเป็น 0")
                th1, th2, th3, th4 = st.columns(4)
                with th1: ss["thick_ac"]   = st.number_input("AC (ซม.)",   min_value=0.0, value=float(ss.get("thick_ac",18.0)),   step=1.0, format="%.1f", key="tac_t4")
                with th2: ss["thick_jpcp"] = st.number_input("JPCP (ซม.)", min_value=0.0, value=float(ss.get("thick_jpcp",28.0)), step=1.0, format="%.1f", key="tjp_t4")
                with th3: ss["thick_jrcp"] = st.number_input("JRCP (ซม.)", min_value=0.0, value=float(ss.get("thick_jrcp",25.0)), step=1.0, format="%.1f", key="tjr_t4")
                with th4: ss["thick_crcp"] = st.number_input("CRCP (ซม.)", min_value=0.0, value=float(ss.get("thick_crcp",25.0)), step=1.0, format="%.1f", key="tcr_t4")

            if st.button("📋 สร้างรายงาน Word", type="primary", key="gen_w_t4"):
                with st.spinner("กำลังสร้างรายงาน..."):
                    try:
                        buf=generate_word(sdf,cfd,ss["lcca_n"],ss["lcca_dr"],
                                         ss.get("lcca_alternatives",[]),
                                         base_sec_routine=ss.get("base_sec_routine","3.8"),
                                         base_sec_lcca=ss.get("base_sec_lcca","3.9"))
                        proj=ss["project_name"].replace(" ","_")
                        st.download_button("⬇️ ดาวน์โหลด Word Report", data=buf,
                            file_name=f"LCCA_{proj}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="dl_w_t4")
                        st.success("✅ สำเร็จ")
                    except Exception as e:
                        st.error(f"สร้างรายงานไม่ได้: {e}"); st.exception(e)
