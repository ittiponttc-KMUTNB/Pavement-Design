 # -*- coding: utf-8 -*-
"""
โปรแกรมรวมรายงานออกแบบโครงสร้างชั้นทาง v3.0
Pavement Design Report Merger - Refactored Version
"""

import streamlit as st
from dataclasses import dataclass
from typing import List, Optional
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import io

# ═══════════════════════════════════════════════════════════════
# CONFIGURATION
# ═══════════════════════════════════════════════════════════════

@dataclass
class ReportSection:
    id: str
    title: str
    category: str
    required: bool = False
    description: str = ""

DEFAULT_SECTIONS = [
    ReportSection("truck_factor", "การคำนวณ Truck Factor", "1. พื้นฐานข้อมูล", False, "ถ้ามี"),
    ReportSection("esals_flex", "ESALs สำหรับผิวทางลาดยาง", "2. การคำนวณ ESALs", True, "Flexible Pavement"),
    ReportSection("esals_rigid", "ESALs สำหรับผิวทางคอนกรีต", "2. การคำนวณ ESALs", True, "Rigid Pavement"),
    ReportSection("cbr", "การวิเคราะห์ค่า CBR ที่เปอร์เซ็นต์ไทล์", "3. การวิเคราะห์", True),
    ReportSection("flex_design", "การออกแบบผิวทางลาดยาง", "4. การออกแบบ", True, "AASHTO 1993"),
    ReportSection("jpcp_design", "การออกแบบ JPCP/JRCP", "5. การออกแบบคอนกรีต", True, "Jointed Plain/Reinforced"),
    ReportSection("crcp_design", "การออกแบบ CRCP", "5. การออกแบบคอนกรีต", False, "Continuously Reinforced"),
    ReportSection("k_jpcp", "Corrected k-value สำหรับ JPCP/JRCP", "6. การคำนวณ k-value", True),
    ReportSection("k_crcp", "Corrected k-value สำหรับ CRCP", "6. การคำนวณ k-value", False),
    ReportSection("cost", "การประมาณราคาค่าก่อสร้าง", "7. การประมาณราคา", False, "ถ้ามี"),
]

# ═══════════════════════════════════════════════════════════════
# UI SETUP
# ═══════════════════════════════════════════════════════════════

st.set_page_config(page_title="รวมรายงานโครงสร้างชั้นทาง v3.0", page_icon="🛣️", layout="wide")

st.markdown("""
<style>
    .main-header { font-size: 32px; font-weight: bold; text-align: center; 
                   background: linear-gradient(90deg, #667eea, #764ba2); 
                   -webkit-background-clip: text; -webkit-text-fill-color: transparent; }
    .section-card { background: #f8fafc; border-radius: 12px; padding: 16px; 
                    border-left: 4px solid #667eea; margin: 8px 0; }
    .file-item { background: white; border-radius: 8px; padding: 12px; 
                 box-shadow: 0 2px 4px rgba(0,0,0,0.1); margin: 4px 0;
                 display: flex; align-items: center; gap: 12px; }
    .drag-handle { cursor: grab; color: #94a3b8; font-size: 20px; }
    .status-badge { padding: 4px 12px; border-radius: 20px; font-size: 12px; font-weight: bold; }
    .status-ready { background: #dcfce7; color: #166534; }
    .status-pending { background: #fee2e2; color: #991b1b; }
    .merge-btn { background: linear-gradient(90deg, #667eea, #764ba2) !important; 
                 color: white !important; font-weight: bold !important; 
                 border-radius: 25px !important; padding: 12px 32px !important; }
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════
# UTILITY FUNCTIONS
# ═══════════════════════════════════════════════════════════════

def set_thai_font(run, name="TH Sarabun New", size=16, bold=False):
    """ตั้งค่าฟอนต์ไทย"""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    for attr in ['w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia']:
        rFonts.set(qn(attr), name)

def set_a4_margins(section):
    """ตั้งค่าหน้ากระดาษ A4"""
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    section.orientation = WD_ORIENT.PORTRAIT
    for margin in ['left', 'right', 'top', 'bottom']:
        setattr(section, f'{margin}_margin', Cm(2.5))

def copy_element(source, target_doc):
    """คัดลอก element จากเอกสารต้นทางไปยังเอกสารปลายทาง"""
    if source.tag.endswith('p'):  # Paragraph
        new_para = target_doc.add_paragraph()
        for run in source.runs:
            new_run = new_para.add_run(run.text)
            set_thai_font(new_run, 
                         name=run.font.name or "TH Sarabun New",
                         size=run.font.size.pt if run.font.size else 16,
                         bold=run.font.bold or False)
        new_para.alignment = source.alignment
        
    elif source.tag.endswith('tbl'):  # Table
        rows, cols = len(source.rows), len(source.columns)
        new_table = target_doc.add_table(rows=rows, cols=cols)
        new_table.style = source.style
        
        for i, row in enumerate(source.rows):
            for j, cell in enumerate(row.cells):
                new_cell = new_table.rows[i].cells[j]
                new_cell.text = cell.text
                # Copy cell formatting
                if cell.paragraphs:
                    new_cell.paragraphs[0].alignment = cell.paragraphs[0].alignment

def create_cover_page(doc, project, date_str):
    """สร้างหน้าปก"""
    doc.add_paragraph("\n" * 6)
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("รายงานการออกแบบโครงสร้างชั้นทาง")
    set_thai_font(run, size=28, bold=True)
    
    if project:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"\n{project}")
        set_thai_font(run, size=22, bold=True)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n\n\n{date_str}")
    set_thai_font(run, size=18)
    
    doc.add_page_break()

def create_toc(doc, sections_with_files):
    """สร้างสารบัญ"""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("สารบัญ")
    set_thai_font(run, size=20, bold=True)
    doc.add_paragraph()
    
    for i, (section, _) in enumerate(sections_with_files, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {section.title}")
        set_thai_font(run, size=16)
    
    doc.add_page_break()

def merge_documents(sections_with_files, project, date_str):
    """รวมเอกสารทั้งหมด"""
    merged = Document()
    set_a4_margins(merged.sections[0])
    
    # หน้าปก
    create_cover_page(merged, project, date_str)
    
    # สารบัญ
    create_toc(merged, sections_with_files)
    
    # เนื้อหาแต่ละส่วน
    for i, (section, file_bytes) in enumerate(sections_with_files, 1):
        # หัวข้อส่วน
        header = merged.add_paragraph()
        run = header.add_run(f"{i}. {section.title}")
        set_thai_font(run, size=20, bold=True)
        merged.add_paragraph()
        
        # เนื้อหาจากไฟล์
        try:
            source = Document(io.BytesIO(file_bytes))
            for element in source.element.body:
                copy_element(element, merged)
            merged.add_page_break()
        except Exception as e:
            p = merged.add_paragraph()
            run = p.add_run(f"[Error loading file: {str(e)}]")
            set_thai_font(run, size=12)
    
    return merged

# ═══════════════════════════════════════════════════════════════
# MAIN UI
# ═══════════════════════════════════════════════════════════════

st.markdown('<h1 class="main-header">🛣️ โปรแกรมรวมรายงานโครงสร้างชั้นทาง v3.0</h1>', unsafe_allow_html=True)

# Sidebar: ตั้งค่าโครงการ
with st.sidebar:
    st.header("⚙️ ตั้งค่าโครงการ")
    project_name = st.text_input("ชื่อโครงการ", placeholder="โครงการก่อสร้างถนน...")
    report_date = st.date_input("วันที่รายงาน", datetime.now())
    
    st.markdown("---")
    st.header("📋 รูปแบบเลขหัวข้อ")
    numbering = st.radio("รูปแบบ", ["1, 2, 3...", "1.1, 1.2...", "A, B, C..."], horizontal=True)
    
    st.markdown("---")
    st.caption("พัฒนาโดย ภาควิชาครุศาสตร์โยธา มจพ. © 2025")

# Main: อัปโหลดและจัดการไฟล์
st.subheader("📤 อัปโหลดไฟล์รายงาน")

# จัดกลุ่มตาม category
categories = {}
for s in DEFAULT_SECTIONS:
    categories.setdefault(s.category, []).append(s)

uploaded_files = {}
files_order = []

# แสดงเป็นกลุ่ม
for cat_name, sections in categories.items():
    with st.expander(f"**{cat_name}**", expanded=True):
        for section in sections:
            col1, col2 = st.columns([3, 1])
            
            with col1:
                file = st.file_uploader(
                    f"**{section.title}**" + (f" *" if section.required else ""),
                    type=["docx"],
                    key=f"up_{section.id}",
                    help=section.description
                )
            
            with col2:
                if file:
                    st.markdown('<span class="status-badge status-ready">✅ พร้อม</span>', 
                               unsafe_allow_html=True)
                    uploaded_files[section] = file.getvalue()
                    files_order.append(section)
                else:
                    st.markdown('<span class="status-badge status-pending">⏳ รอไฟล์</span>' 
                               + (' *' if section.required else ''), 
                               unsafe_allow_html=True)

# แสดงสรุปและปุ่มดำเนินการ
st.markdown("---")

ready_count = len(uploaded_files)
required_count = sum(1 for s in DEFAULT_SECTIONS if s.required)
required_ready = sum(1 for s in uploaded_files if s.required)

col1, col2, col3 = st.columns([2, 2, 2])

with col1:
    st.metric("ไฟล์ที่พร้อม", f"{ready_count}/{len(DEFAULT_SECTIONS)}")

with col2:
    if required_ready >= required_count:
        st.success(f"✅ ครบตามบังคับ ({required_ready}/{required_count})")
    else:
        st.warning(f"⚠️ ขาดบังคับอีก {required_count - required_ready} ไฟล์")

with col3:
    if ready_count > 0 and required_ready >= required_count:
        if st.button("🔄 รวมไฟล์และดาวน์โหลด", type="primary", use_container_width=True):
            with st.spinner("กำลังรวมเอกสาร..."):
                try:
                    # เรียงตามลำดับใน DEFAULT_SECTIONS
                    ordered = [(s, uploaded_files[s]) for s in DEFAULT_SECTIONS if s in uploaded_files]
                    
                    merged = merge_documents(
                        ordered, 
                        project_name, 
                        report_date.strftime("%d/%m/%Y")
                    )
                    
                    # Export
                    output = io.BytesIO()
                    merged.save(output)
                    output.seek(0)
                    
                    filename = f"รายงานออกแบบ_{project_name or 'โครงสร้างชั้นทาง'}_{report_date:%Y%m%d}.docx"
                    
                    st.download_button(
                        "📥 ดาวน์โหลดรายงาน (.docx)",
                        data=output.getvalue(),
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                    
                    st.balloons()
                    st.success(f"✅ รวมสำเร็จ! ({ready_count} ไฟล์)")
                    
                except Exception as e:
                    st.error(f"❌ เกิดข้อผิดพลาด: {str(e)}")
                    st.exception(e)
    else:
        st.button("🔄 รวมไฟล์และดาวน์โหลด", disabled=True, use_container_width=True)
        if ready_count == 0:
            st.error("กรุณาอัปโหลดไฟล์อย่างน้อย 1 ไฟล์")
        elif required_ready < required_count:
            st.error("กรุณาอัปโหลดไฟล์ที่บังคับให้ครบ")

# Footer
st.markdown("---")
st.caption("หมายเหตุ: ไฟล์ที่มีเครื่องหมาย * จำเป็นต้องอัปโหลด")
