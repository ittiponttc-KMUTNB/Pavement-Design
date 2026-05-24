"""
================================================================================
ระบบวิเคราะห์ค่าก่อสร้างโครงสร้างชั้นทาง + วิเคราะห์ต้นทุนตลอดอายุการใช้งาน
Pavement Structure Cost + Life-Cycle Cost Analysis (Integrated v1.0)
================================================================================
พัฒนาโดย: รศ.ดร.อิทธิพล มีผล
ภาควิชาครุศาสตร์โยธา คณะครุศาสตร์อุตสาหกรรม
มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ (KMUTNB)

โครงสร้างโปรแกรม:
  Tab A — 🏗️ โครงสร้างชั้นทาง  (Cost-Pavement-Structure)
    ├── Sub: Layer Editor (AC / JPCP / JRCP / CRCP)
    ├── Sub: ราคาวัสดุ (Price Library)
    └── Sub: สรุปต้นทุน & Word Report
  Tab B — 📊 LCCA  (Life-Cycle Cost Analysis)
    ├── Sub: Routine Cost (Ka / Kc)
    ├── Sub: กำหนดทางเลือก  ← Auto-fill cost_sqm จาก Tab A
    ├── Sub: ผลการวิเคราะห์
    └── Sub: Word Report

Session-state namespaces:
  cs_*   → Cost-Structure module
  lc_*   → LCCA module
  sb_*   → Sidebar / shared project info
================================================================================
"""

import ast
import io
import json
import hashlib
from dataclasses import dataclass
from datetime import datetime
from itertools import combinations
from typing import List

import numpy as np
import pandas as pd
import streamlit as st

from cross_section_tab import render_cross_section_tab

# ── optional imports ──────────────────────────────────────────────────────────
try:
    import plotly.express as px
    import plotly.graph_objects as go
    PLOTLY_OK = True
except ImportError:
    PLOTLY_OK = False

try:
    from docx import Document as WordDoc
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    import openpyxl
    from openpyxl.styles import Alignment as XlAlign, Font as XlFont, PatternFill
    OPENPYXL_OK = True
except ImportError:
    OPENPYXL_OK = False

# ── syntax check ──────────────────────────────────────────────────────────────
_src = open(__file__).read()
try:
    ast.parse(_src)
except SyntaxError as _e:
    raise SyntaxError(f"[ast.parse] {__file__}: {_e}") from _e

# ══════════════════════════════════════════════════════════════════════════════
# PAGE CONFIG
# ══════════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="PaveCost + LCCA",
    page_icon="🛣️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ══════════════════════════════════════════════════════════════════════════════
# GLOBAL CSS
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Sans+Thai:wght@300;400;500;600&family=Space+Grotesk:wght@400;500;600;700&display=swap');

html, body, [class*="css"] { font-family: 'IBM Plex Sans Thai', sans-serif; }

/* ── Main header ── */
.main-header {
    background: linear-gradient(135deg, #0f2942 0%, #1a4a7a 60%, #0d7377 100%);
    border-radius: 16px; padding: 1.6rem 2rem; margin-bottom: 1.2rem;
}
.main-header h1 { font-family:'Space Grotesk',sans-serif; color:#fff;
    font-size:1.6rem; font-weight:700; margin:0; }
.main-header p  { color:rgba(255,255,255,.65); font-size:.88rem; margin:.3rem 0 0 0; }

/* ── Metric cards ── */
.mc { background:#fff; border:1px solid #e8edf2; border-radius:12px;
    padding:1rem 1.2rem; box-shadow:0 2px 8px rgba(0,0,0,.05); }
.mc .lbl { font-size:.76rem; color:#6b7a8d; font-weight:500;
    text-transform:uppercase; letter-spacing:.5px; margin-bottom:.25rem; }
.mc .val { font-family:'Space Grotesk',sans-serif; font-size:1.45rem;
    font-weight:700; color:#0f2942; }
.mc .sub { font-size:.78rem; color:#8a95a3; margin-top:.2rem; }

/* colored metric cards (LCCA) */
.mcc { border-radius:12px; padding:14px 18px; margin:4px 0;
    color:#fff; box-shadow:0 4px 12px rgba(0,0,0,.15); }
.mcc .lbl { font-size:13px; font-weight:600; opacity:.88; margin-bottom:3px; }
.mcc .val { font-size:26px; font-weight:700; }
.mcc .sub { font-size:12px; opacity:.80; margin-top:2px; }
.c-blue   { background:linear-gradient(135deg,#1565C0,#42A5F5); }
.c-orange { background:linear-gradient(135deg,#E65100,#FFA726); }
.c-green  { background:linear-gradient(135deg,#1B5E20,#66BB6A); }
.c-purple { background:linear-gradient(135deg,#4A148C,#AB47BC); }
.c-teal   { background:linear-gradient(135deg,#004D40,#26A69A); }
.c-gold   { background:linear-gradient(135deg,#827717,#FFEE58); color:#333!important; }

/* ── Section card ── */
.sc { background:#f8fafc; border:1px solid #e2e8f0;
    border-left:4px solid #1a4a7a; border-radius:10px;
    padding:.9rem 1.1rem; margin-bottom:.9rem; }

/* ── Section header band ── */
.sh      { background:linear-gradient(90deg,#1565C0,#42A5F5);
    color:#fff; padding:7px 14px; border-radius:7px;
    font-size:16px; font-weight:700; margin:10px 0 7px 0; }
.sh-org  { background:linear-gradient(90deg,#E65100,#FFA726);
    color:#fff; padding:7px 14px; border-radius:7px;
    font-size:16px; font-weight:700; margin:10px 0 7px 0; }
.sh-grn  { background:linear-gradient(90deg,#1B5E20,#66BB6A);
    color:#fff; padding:7px 14px; border-radius:7px;
    font-size:16px; font-weight:700; margin:10px 0 7px 0; }

/* ── Progress ── */
.prog-wrap { display:flex; gap:7px; margin:6px 0 14px 0; align-items:center; }
.prog-step { flex:1; padding:7px 4px; border-radius:8px; text-align:center;
    font-size:13px; font-weight:600; border:2px solid #ccc; }
.prog-done { background:#E8F5E9; border-color:#43A047; color:#1B5E20; }
.prog-warn { background:#FFF8E1; border-color:#FFA000; color:#E65100; }
.prog-idle { background:#F5F5F5; border-color:#BDBDBD; color:#757575; }
.prog-arr  { font-size:17px; color:#BDBDBD; flex:0; }

/* ── Badges ── */
.badge-ok   { background:#E8F5E9; color:#1B5E20; border:1px solid #43A047;
    border-radius:20px; padding:3px 11px; font-size:13px; font-weight:600; }
.badge-warn { background:#FFF3E0; color:#E65100; border:1px solid #FFA000;
    border-radius:20px; padding:3px 11px; font-size:13px; font-weight:600; }

/* ── Info / warn bands ── */
.info-band { background:#E3F2FD; border-left:4px solid #1565C0;
    border-radius:6px; padding:9px 13px; margin:7px 0;
    font-size:14px; color:#0D47A1; }
.warn-band { background:#FFF8E1; border-left:4px solid #FFA000;
    border-radius:6px; padding:9px 13px; margin:7px 0;
    font-size:14px; color:#E65100; }

/* ── Road preview box ── */
.road-box { background:linear-gradient(135deg,#263238,#455A64);
    color:#fff; border-radius:10px; padding:12px 16px; font-size:15px; margin-top:7px; }
.road-box .rv { font-size:22px; font-weight:700; color:#FFD54F; }

/* ── Best result ── */
.best-row { background:linear-gradient(90deg,#E8F5E9,#F1F8E9);
    border-left:5px solid #43A047; border-radius:6px;
    padding:9px 13px; margin:4px 0; font-size:14px; }

/* ── Selectbox green ── */
div[data-baseweb="select"]>div { background-color:#f0faf4!important;
    border-color:#52b788!important; border-radius:8px!important; }
div[data-baseweb="select"] span { color:#1b5e20!important; font-weight:500!important; }
div[data-baseweb="menu"] { background-color:#f1f8e9!important; }
div[data-baseweb="menu"] li { background-color:#f1f8e9!important; color:#1b5e20!important; }
div[data-baseweb="menu"] li:hover { background-color:#c8e6c9!important; }

/* ── Tab styling ── */
button[data-baseweb="tab"] {
    font-family:'Space Grotesk',sans-serif!important;
    font-weight:600!important; font-size:.88rem!important; }

/* ── Footer ── */
.footer { text-align:center; color:#94a3b8; font-size:.78rem;
    padding:1.2rem 0 .5rem; border-top:1px solid #e2e8f0; margin-top:1.5rem; }

/* ── auto-fill badge ── */
.autofill { background:#e0f2f1; color:#00695c; border:1px solid #26a69a;
    border-radius:12px; padding:2px 10px; font-size:12px; font-weight:600; }

/* ── Sync panel ── */
.sync-panel {
    background:linear-gradient(135deg,#f0fdf4,#ecfdf5);
    border:1.5px solid #86efac; border-radius:12px;
    padding:14px 18px; margin:10px 0;
}
.sync-panel .sync-title { font-weight:700; color:#14532d; font-size:14px; margin-bottom:8px; }
.sync-panel .sync-ts    { font-size:12px; color:#6b7a8d; margin-top:6px; }
.sync-pill {
    display:inline-block; background:#dcfce7; color:#15803d;
    border:1px solid #86efac; border-radius:20px;
    padding:3px 12px; font-size:13px; font-weight:700; margin:2px 4px 2px 0;
}
.sync-pill-zero {
    display:inline-block; background:#f3f4f6; color:#9ca3af;
    border:1px solid #d1d5db; border-radius:20px;
    padding:3px 12px; font-size:13px; font-weight:600; margin:2px 4px 2px 0;
}
.sync-warn-banner {
    background:#fff7ed; border:1.5px solid #fb923c;
    border-radius:10px; padding:10px 16px; margin:8px 0;
    font-size:13px; color:#9a3412; font-weight:600;
}
.locked-badge {
    background:#f0f9ff; color:#0369a1; border:1px solid #7dd3fc;
    border-radius:12px; padding:2px 10px; font-size:12px; font-weight:600;
}
.manual-badge {
    background:#fef3c7; color:#92400e; border:1px solid #fcd34d;
    border-radius:12px; padding:2px 10px; font-size:12px; font-weight:600;
}
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# ─── MODULE A: COST-STRUCTURE — CONSTANTS & PRICE TABLES ─────────────────────
# ══════════════════════════════════════════════════════════════════════════════

DEFAULT_AC_TON_PRICES: dict = {
    'PMA Wearing Course': 3100,
    'AC Wearing Course':  2973,
    'AC Binder Course':   2929,
    'AC Base Course':     1795,
}
DEFAULT_AC_DENSITY: float = 2.4

DEFAULT_AC_PRICES: dict = {
    'PMA Wearing Course': {2.5:186,3:223,4:298,5:372,6:446,7:521,8:595,9:670,10:744},
    'AC Wearing Course':  {2.5:178,3:214,4:285,5:357,6:428,7:499,8:571,9:642,10:714},
    'AC Binder Course':   {2.5:176,3:211,4:281,5:351,6:422,7:492,8:562,9:633,10:703},
    'AC Base Course':     {2.5:108,3:129,4:172,5:215,6:258,7:302,8:345,9:388,10:431},
}

DEFAULT_CONCRETE_CUM_PRICES: dict = {'JPCP':2732,'JRCP':3077,'CRCP':3663}

def _calc_concrete_prices(cum_prices: dict) -> dict:
    thicknesses = [20,25,28,30,32,35]
    return {pt:{t:round(cum*t/100,0) for t in thicknesses}
            for pt,cum in cum_prices.items()}

DEFAULT_CONCRETE_PRICES: dict = _calc_concrete_prices(DEFAULT_CONCRETE_CUM_PRICES)

DEFAULT_BASE_PRICES: dict = {
    'Cement Treated Base (UCS 40 ksc)':                    1096,
    'Cement Modified Crushed Rock Base (UCS 24.5 ksc)':     919,
    'Crushed Rock Base Course':                              583,
    'Soil Cement Subbase (UCS 7 ksc)':                      854,
    'Soil Aggregate Subbase':                                375,
    'Selected Material A':                                   375,
    'Embankment':                                            352,
    'Sand Embankment':                                       220,
    'Prime Coat':                                           37.47,
    'Non Woven Geotextile':                                  78,
    'Wire Mesh':                                            100,
    'Tack Coat':                                             20,
}

BASE_MATERIAL_LIST = [
    'Cement Treated Base (UCS 40 ksc)',
    'Cement Modified Crushed Rock Base (UCS 24.5 ksc)',
    'Crushed Rock Base Course',
    'Soil Cement Subbase (UCS 7 ksc)',
    'Soil Aggregate Subbase',
    'Selected Material A',
    'Embankment',
    'Sand Embankment',
]
SQMKEYS = {'Prime Coat','Non Woven Geotextile','Wire Mesh','Tack Coat'}
BASE_KEYWORDS = ['crushed rock','soil aggregate','soil cement','cement modified',
                 'cement treated','selected material','embankment','sand embankment']

# ══════════════════════════════════════════════════════════════════════════════
# ─── MODULE B: LCCA — LOOKUP TABLES ──────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

X1_MAP = {
    "High Type (AC/PM บนหินคลุก)":                0.00,
    "Intermediate Type (AC/PM บน Stabilized)":    0.50,
    "Low Type (ST บน Soil-Aggregate)":             1.00,
}
X2_BREAKS = [(0,2,1.00),(2.01,3,.75),(3.01,4,.50),(4.01,5,.25),(5.01,999,0.00)]
X3_OPTIONS = {
    "0 – 500       (X3=0.00)":0.00,"501 – 600     (X3=0.04)":0.04,
    "601 – 700     (X3=0.08)":0.08,"701 – 800     (X3=0.12)":0.12,
    "801 – 900     (X3=0.16)":0.16,"901 – 1,000   (X3=0.20)":0.20,
    "1,001 – 1,100 (X3=0.24)":0.24,"1,101 – 1,200 (X3=0.29)":0.29,
    "1,201 – 1,300 (X3=0.33)":0.33,"1,301 – 1,400 (X3=0.37)":0.37,
    "1,401 – 1,500 (X3=0.41)":0.41,"1,501 – 1,600 (X3=0.45)":0.45,
    "1,601 – 1,700 (X3=0.49)":0.49,"1,701 – 1,800 (X3=0.53)":0.53,
    "1,801 – 1,900 (X3=0.57)":0.57,"1,901 – 2,000 (X3=0.61)":0.61,
    "2,001 – 2,200 (X3=0.69)":0.69,"2,201 – 2,400 (X3=0.78)":0.78,
    "2,401 – 2,600 (X3=0.86)":0.86,"2,601 – 2,800 (X3=0.94)":0.94,
    "2,801 – 3,000 (X3=1.02)":1.02,"3,001 – 3,300 (X3=1.14)":1.14,
    "3,301 – 3,600 (X3=1.27)":1.27,"3,601 – 3,900 (X3=1.37)":1.37,
    "3,901 – 4,200 (X3=1.51)":1.51,"4,201 – 4,500 (X3=1.64)":1.64,
    "4,501 – 4,800 (X3=1.76)":1.76,"4,801 – 5,100 (X3=1.88)":1.88,
    "5,101 – 5,400 (X3=2.00)":2.00,"5,401 – 5,700 (X3=2.13)":2.13,
    "5,701+         (X3=2.25)":2.25,
}
X4_BREAKS = [(0,3,.00),(4,4,.20),(5,5,.40),(6,6,.60),(7,7,.80),
             (8,8,1.00),(9,9,1.20),(10,10,1.40),(11,11,1.60),(12,99999,1.80)]
X5_BREAKS = [(0,5.49,.00),(5.50,5.99,.02),(6.00,6.49,.05),(6.50,6.99,.10),(7.00,9999,.19)]
TERRAIN_MAP  = {"ที่ราบ (0-3%)":"P","ลูกเนิน (3-5%)":"R",
                "ลูกเนินสลับเขา (5-7%)":"RM","เขา (>7%)":"S"}
TERRAIN_KEYS = list(TERRAIN_MAP.keys())
X6_MAP = {"P":0.00,"R":0.02,"RM":0.04,"S":0.07}
Y3_MAP = {"P":0.00,"R":0.24,"RM":0.36,"S":0.48}
Y4_MAP = {"P":0.00,"R":0.24,"RM":0.36,"S":0.48}
Y6_MAP = {"P":0.00,"R":0.04,"RM":0.08,"S":0.12}
Y1_BREAKS = [(0,40,.00),(40.01,60,.10),(60.01,80,.20),(80.01,9999,.30)]
Y2_BREAKS = [(0,1.75,.00),(1.76,2.00,.10),(2.01,2.25,.15),(2.26,9999,.20)]
Y5_BREAKS = [(0,20.99,.00),(21,25,.02),(25.01,30,.04),(30.01,9999,.06)]
Z1_MAP    = {1:0.00,2:0.25,3:0.50,4:0.75,5:1.00,6:1.30,7:1.60,8:2.00}
Z2_BREAKS = [(0,2,1.00),(2.01,3,.75),(3.01,4,.50),(4.01,5,.25),(5.01,999,0.00)]
Z3_OPTIONS = {
    "0 – 1,000        (Z3=0.00)":0.00,"1,001 – 2,000    (Z3=0.20)":0.20,
    "2,001 – 3,000    (Z3=0.30)":0.30,"3,001 – 4,000    (Z3=0.50)":0.50,
    "4,001 – 5,000    (Z3=0.75)":0.75,"5,001 – 6,000    (Z3=1.00)":1.00,
    "6,001 – 7,000    (Z3=1.25)":1.25,"7,001 – 8,000    (Z3=1.50)":1.50,
    "8,001 – 9,000    (Z3=1.75)":1.75,"9,001 – 10,000   (Z3=2.00)":2.00,
    "10,001 – 15,000  (Z3=2.50)":2.50,"15,001+           (Z3=3.00)":3.00,
}
Z4_BREAKS = [(0,6.49,.00),(6.50,6.99,.08),(7.00,9999,.17)]

# ══════════════════════════════════════════════════════════════════════════════
# ─── PRICE LIBRARY HELPERS ───────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

def get_price_library() -> dict:
    if 'price_library' in st.session_state:
        return st.session_state['price_library']
    return {'ac_prices':DEFAULT_AC_PRICES,
            'concrete_prices':DEFAULT_CONCRETE_PRICES,
            'base_prices':DEFAULT_BASE_PRICES}


def lookup_price(name: str, thickness: float, ptype: str = 'AC') -> float:
    lib = get_price_library()
    n = name.lower()

    def _nearest(d: dict, t: float) -> float:
        if not d: return 0.0
        if t in d: return float(d[t])
        return float(d[min(d.keys(), key=lambda x: abs(x-t))])

    if 'pma' in n:
        return _nearest(lib['ac_prices'].get('PMA Wearing Course',{}), thickness)
    if 'wearing' in n:
        return _nearest(lib['ac_prices'].get('AC Wearing Course',{}), thickness)
    if 'binder' in n:
        return _nearest(lib['ac_prices'].get('AC Binder Course',{}), thickness)
    if ('asphalt' in n and 'base' in n) or 'ac base' in n or 'interlayer' in n:
        return _nearest(lib['ac_prices'].get('AC Base Course',{}), thickness)

    _conc_cum = st.session_state.get('concrete_cum_prices', DEFAULT_CONCRETE_CUM_PRICES)
    for ct in ('jpcp','jrcp','crcp'):
        if ct in n:
            cum = float(_conc_cum.get(ct.upper(), DEFAULT_CONCRETE_CUM_PRICES.get(ct.upper(),0)))
            return round(cum*thickness/100,2) if cum>0 else _nearest(lib['concrete_prices'].get(ct.upper(),{}), thickness)
    if ptype in ('JPCP','JRCP','CRCP') and ('concrete' in n or 'ksc' in n or '350' in n or 'slab' in n):
        cum = float(_conc_cum.get(ptype, DEFAULT_CONCRETE_CUM_PRICES.get(ptype,0)))
        return round(cum*thickness/100,2) if cum>0 else _nearest(lib['concrete_prices'].get(ptype,{}), thickness)

    if 'tack' in n:   return float(lib['base_prices'].get('Tack Coat',20))
    if 'prime' in n:  return float(lib['base_prices'].get('Prime Coat',37.47))
    if 'geotextile' in n: return float(lib['base_prices'].get('Non Woven Geotextile',78))
    if 'wire' in n:   return float(lib['base_prices'].get('Wire Mesh',100))

    for key in BASE_MATERIAL_LIST:
        if key.lower() == n: return float(lib['base_prices'].get(key,0))
    for key in BASE_MATERIAL_LIST:
        if key.lower() in n and len(key)>5: return float(lib['base_prices'].get(key,0))
    for key,val in lib['base_prices'].items():
        if key.lower() == n: return float(val)
    return 0.0


def _calc_ac_prices_from_ton(ton_prices: dict, density: float = DEFAULT_AC_DENSITY) -> dict:
    thicknesses = [2.5,3,4,5,6,7,8,9,10]
    result = {}
    for mat,ton_p in ton_prices.items():
        if ton_p>0:
            result[mat] = {t:round(ton_p*density*t/100,2) for t in thicknesses}
        else:
            result[mat] = dict(DEFAULT_AC_PRICES.get(mat,{}))
    return result

# ══════════════════════════════════════════════════════════════════════════════
# ─── DEFAULT LAYERS / JOINTS ─────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

def get_default_layers(ptype: str, area_per_km: float = 22000) -> list:
    lib = get_price_library()
    ac = lib['ac_prices']; cp = lib['concrete_prices']; bp = lib['base_prices']
    def ap(mat,t):
        return ac.get(mat,{}).get(t, list(ac.get(mat,{0:0}).values())[0] if ac.get(mat) else 0)
    if ptype=='AC':
        return [
            {'name':'AC Wearing Course','thickness':7,'unit':'cm','quantity':area_per_km,'qty_unit':'sq.m','unit_cost':ap('AC Wearing Course',7)},
            {'name':'AC Binder Course','thickness':7,'unit':'cm','quantity':area_per_km,'qty_unit':'sq.m','unit_cost':ap('AC Binder Course',7)},
            {'name':'AC Base Course','thickness':10,'unit':'cm','quantity':area_per_km,'qty_unit':'sq.m','unit_cost':ap('AC Base Course',10)},
            {'name':'Tack Coat','thickness':1,'unit':'Layer','quantity':area_per_km*2,'qty_unit':'sq.m','unit_cost':float(bp.get('Tack Coat',20))},
            {'name':'Prime Coat','thickness':1,'unit':'Layer','quantity':area_per_km,'qty_unit':'sq.m','unit_cost':float(bp.get('Prime Coat',37.47))},
        ]
    if ptype=='JPCP':
        return [{'name':'Concrete Slab (JPCP)','thickness':28,'unit':'cm','quantity':area_per_km,'qty_unit':'sq.m','unit_cost':cp.get('JPCP',{}).get(28,1000)}]
    if ptype=='JRCP':
        return [{'name':'Concrete Slab (JRCP)','thickness':28,'unit':'cm','quantity':area_per_km,'qty_unit':'sq.m','unit_cost':cp.get('JRCP',{}).get(28,1002)}]
    if ptype=='CRCP':
        return [{'name':'Concrete Slab (CRCP)','thickness':25,'unit':'cm','quantity':area_per_km,'qty_unit':'sq.m','unit_cost':cp.get('CRCP',{}).get(25,1245)}]
    return []


def get_default_base_layers(ptype: str, area_per_km: float = 22000) -> list:
    lib = get_price_library(); bp = lib['base_prices']
    def b(mat,t): return float(bp.get(mat,0))*t/100
    if ptype=='AC':
        return [
            {'name':'Crushed Rock Base Course','thickness':20,'unit':'cm','quantity':area_per_km,'qty_unit':'sq.m','unit_cost':b('Crushed Rock Base Course',20),'cost_cum':float(bp.get('Crushed Rock Base Course',583))},
            {'name':'Soil Aggregate Subbase','thickness':30,'unit':'cm','quantity':area_per_km,'qty_unit':'sq.m','unit_cost':b('Soil Aggregate Subbase',30),'cost_cum':float(bp.get('Soil Aggregate Subbase',375))},
            {'name':'Sand Embankment','thickness':40,'unit':'cm','quantity':area_per_km,'qty_unit':'sq.m','unit_cost':b('Sand Embankment',40),'cost_cum':float(bp.get('Sand Embankment',220))},
        ]
    return [
        {'name':'Cement Modified Crushed Rock Base (UCS 24.5 ksc)','thickness':20,'unit':'cm','quantity':area_per_km,'qty_unit':'sq.m','unit_cost':b('Cement Modified Crushed Rock Base (UCS 24.5 ksc)',20),'cost_cum':float(bp.get('Cement Modified Crushed Rock Base (UCS 24.5 ksc)',864))},
        {'name':'Soil Aggregate Subbase','thickness':20,'unit':'cm','quantity':area_per_km,'qty_unit':'sq.m','unit_cost':b('Soil Aggregate Subbase',20),'cost_cum':float(bp.get('Soil Aggregate Subbase',375))},
        {'name':'Sand Embankment','thickness':50,'unit':'cm','quantity':area_per_km,'qty_unit':'sq.m','unit_cost':b('Sand Embankment',50),'cost_cum':float(bp.get('Sand Embankment',220))},
    ]


def get_default_joints(ptype: str, area_per_km: float = 22000, road_length: float = 1.0) -> list:
    width_m = area_per_km/1000
    if ptype=='JPCP':
        return [
            {'name':'Transverse Joint @4m','quantity':(road_length*1000/4)*width_m,'qty_unit':'m','unit_cost':430},
            {'name':'Longitudinal Joint','quantity':road_length*1000,'qty_unit':'m','unit_cost':120},
        ]
    if ptype=='JRCP':
        return [
            {'name':'Transverse Joint @10m','quantity':(road_length*1000/10)*width_m,'qty_unit':'m','unit_cost':430},
            {'name':'Longitudinal Joint','quantity':road_length*1000,'qty_unit':'m','unit_cost':120},
        ]
    if ptype=='CRCP':
        return [
            {'name':'Longitudinal Steel (CRCP)','quantity':road_length*1000,'qty_unit':'m','unit_cost':200},
            {'name':'Transverse Joint (End)','quantity':0,'qty_unit':'m','unit_cost':500},
        ]
    return []

# ══════════════════════════════════════════════════════════════════════════════
# ─── CALCULATE LAYER / JOINT COST ────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

def calculate_layer_cost(layers: list, road_length_km: float = 1.0) -> tuple:
    total = 0.0; details = []
    for layer in layers:
        qty_raw   = float(layer['quantity'])
        unit_cost = float(layer['unit_cost'])
        cost      = qty_raw*unit_cost; total += cost
        name_lower = layer['name'].lower()
        is_base = any(kw in name_lower for kw in BASE_KEYWORDS)
        if is_base:
            thick_cm = float(layer.get('thickness',1))
            u = layer.get('unit','cm').lower()
            price_cum = (float(layer['cost_cum']) if layer.get('cost_cum')
                         else (unit_cost/(thick_cm/100) if thick_cm>0 and u in ('cm','ซม.','ซ.ม.') else unit_cost))
            qty_display = qty_raw*thick_cm/100 if (thick_cm>0 and u in ('cm','ซม.','ซ.ม.')) else qty_raw
            details.append({
                'รายการ':layer['name'],'ความหนา':f"{layer['thickness']} {layer['unit']}",
                'ปริมาณ':qty_display,'หน่วย':'ลบ.ม.',
                'ราคา/หน่วย':unit_cost,'ราคา/หน่วย (แสดง)':f"{price_cum:,.0f}",
                'หน่วยราคา':'บาท/ลบ.ม.','มูลค่า (บาท)':cost,
            })
        else:
            details.append({
                'รายการ':layer['name'],'ความหนา':f"{layer['thickness']} {layer['unit']}",
                'ปริมาณ':qty_raw,'หน่วย':'ตร.ม.',
                'ราคา/หน่วย':unit_cost,'ราคา/หน่วย (แสดง)':f"{unit_cost:,.0f}",
                'หน่วยราคา':'บาท/ตร.ม.','มูลค่า (บาท)':cost,
            })
    return total, details


def calculate_joint_cost(joints: list, road_length_km: float = 1.0, include_joints: bool = True) -> tuple:
    total = 0.0; details = []
    for joint in joints:
        qty  = float(joint['quantity'])
        cost = qty*float(joint['unit_cost']) if include_joints else 0.0
        total += cost
        unit_th = 'ม.' if joint.get('qty_unit','m')=='m' else joint.get('qty_unit','m')
        details.append({
            'รายการ':joint['name'],'ความหนา':'-','ปริมาณ':qty,'หน่วย':unit_th,
            'ราคา/หน่วย':float(joint['unit_cost']),'ราคา/หน่วย (แสดง)':f"{float(joint['unit_cost']):,.0f}",
            'หน่วยราคา':'บาท/ม.','มูลค่า (บาท)':cost,
        })
    return total, details

# ══════════════════════════════════════════════════════════════════════════════
# ─── RENDER LAYER EDITOR ─────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

def render_layer_editor(ptype: str, key_prefix: str,
                        total_width: float, road_length: float, v: int = 0) -> list:
    area_per_km = total_width*1000
    proj_area   = area_per_km*road_length
    lib         = get_price_library()
    is_concrete = ptype in ('JPCP','JRCP','CRCP')
    updated_layers: list = []

    def _price(name, thick):
        p = lookup_price(name, thick, ptype)
        return float(p) if p>0 else 0.0

    # ── SECTION A: ผิวทาง ────────────────────────────────────────────────────
    st.markdown('<div class="sc"><b>🏗️ ผิวทาง</b> <span style="color:#6b7a8d;font-size:.83rem"> — ราคาดึงจาก Library อัตโนมัติ แก้ได้ที่ Sub-tab 💰 ราคาวัสดุ</span></div>', unsafe_allow_html=True)

    ac_layer_count = 0

    if not is_concrete:
        hc = st.columns([3.5,1.2,1.8])
        hc[0].markdown("<span style='color:#6b7a8d;font-size:.82rem;font-weight:600'>รายการ</span>",unsafe_allow_html=True)
        hc[1].markdown("<span style='color:#6b7a8d;font-size:.82rem;font-weight:600'>หนา (cm)</span>",unsafe_allow_html=True)
        hc[2].markdown("<div style='color:#6b7a8d;font-size:.82rem;font-weight:600;text-align:right'>ราคา (บาท/ตร.ม.)</div>",unsafe_allow_html=True)

        # Row 1: Wearing Course
        _wr_key = f"{key_prefix}_wearing_type_v{v}"
        if _wr_key not in st.session_state: st.session_state[_wr_key]='AC Wearing Course'
        r1 = st.columns([3.5,1.2,1.8])
        with r1[0]:
            wearing_type = st.radio("Wearing",['AC Wearing Course','PMA Wearing Course'],
                index=0 if st.session_state[_wr_key]=='AC Wearing Course' else 1,
                horizontal=True, key=f"{key_prefix}_wearing_radio_v{v}", label_visibility="collapsed")
            st.session_state[_wr_key]=wearing_type
        with r1[1]:
            _wt_key=f"{key_prefix}_sthick_w_v{v}"
            if _wt_key not in st.session_state: st.session_state[_wt_key]=5.0
            st.number_input("หนา W",min_value=0.5,max_value=20.0,step=0.5,format="%.1f",
                key=_wt_key, label_visibility="collapsed")
        wearing_thick=float(st.session_state[_wt_key])
        with r1[2]:
            wearing_price=_price(wearing_type,wearing_thick)
            st.markdown(f"<div style='padding:8px 0;font-weight:700;color:#0f2942;text-align:right'>{wearing_price:,.2f}</div>",unsafe_allow_html=True)
        ac_layer_count+=1
        updated_layers.append({'name':wearing_type,'thickness':wearing_thick,'unit':'cm',
            'quantity':proj_area,'qty_unit':'sq.m','unit_cost':wearing_price,'cost_per_sqm':wearing_price})

        # Row 2: Binder Course
        r2=st.columns([3.5,1.2,1.8])
        with r2[0]: st.markdown("<div style='padding:8px 0;font-weight:600'>AC Binder Course</div>",unsafe_allow_html=True)
        with r2[1]:
            _bt_key=f"{key_prefix}_sthick_b_v{v}"
            if _bt_key not in st.session_state: st.session_state[_bt_key]=5.0
            st.number_input("หนา B",min_value=0.5,max_value=20.0,step=0.5,format="%.1f",
                key=_bt_key, label_visibility="collapsed")
        binder_thick=float(st.session_state[_bt_key])
        with r2[2]:
            binder_price=_price('AC Binder Course',binder_thick)
            st.markdown(f"<div style='padding:8px 0;font-weight:700;color:#0f2942;text-align:right'>{binder_price:,.2f}</div>",unsafe_allow_html=True)
        ac_layer_count+=1
        updated_layers.append({'name':'AC Binder Course','thickness':binder_thick,'unit':'cm',
            'quantity':proj_area,'qty_unit':'sq.m','unit_cost':binder_price,'cost_per_sqm':binder_price})

        # Row 3: Base Course (checkbox)
        r3=st.columns([3.5,1.2,1.8])
        with r3[0]: use_base=st.checkbox("AC Base Course",value=True,key=f"{key_prefix}_use_base_v{v}")
        if use_base:
            with r3[1]:
                _basethick_key=f"{key_prefix}_sthick_base_v{v}"
                if _basethick_key not in st.session_state: st.session_state[_basethick_key]=8.0
                st.number_input("หนา Base",min_value=0.5,max_value=30.0,step=0.5,format="%.1f",
                    key=_basethick_key, label_visibility="collapsed")
            base_thick=float(st.session_state[_basethick_key])
            with r3[2]:
                base_price=_price('AC Base Course',base_thick)
                st.markdown(f"<div style='padding:8px 0;font-weight:700;color:#0f2942;text-align:right'>{base_price:,.2f}</div>",unsafe_allow_html=True)
            ac_layer_count+=1
            updated_layers.append({'name':'AC Base Course','thickness':base_thick,'unit':'cm',
                'quantity':proj_area,'qty_unit':'sq.m','unit_cost':base_price,'cost_per_sqm':base_price})

        # Row 4: Tack Coat
        tack_times=max(ac_layer_count-1,1); tack_qty=proj_area*tack_times
        tack_price=float(get_price_library()['base_prices'].get('Tack Coat',20))
        r4=st.columns([3.5,1.2,1.8])
        with r4[0]: st.markdown(f"<div style='padding:8px 0;color:#0f2942'>Tack Coat <span style='color:#6b7a8d;font-size:.82rem'>({tack_times} ครั้ง × {proj_area:,.0f} = {tack_qty:,.0f} ตร.ม.)</span></div>",unsafe_allow_html=True)
        with r4[1]: st.markdown("<div style='padding:8px 0;color:#94a3b8;font-size:.85rem'>auto</div>",unsafe_allow_html=True)
        with r4[2]: st.markdown(f"<div style='padding:8px 0;font-weight:700;color:#0f2942;text-align:right'>{tack_price:,.2f}</div>",unsafe_allow_html=True)
        updated_layers.append({'name':'Tack Coat','thickness':1,'unit':'Layer',
            'quantity':tack_qty,'qty_unit':'sq.m','unit_cost':tack_price,'cost_per_sqm':tack_price})

        # Row 5: Prime Coat
        prime_price=float(get_price_library()['base_prices'].get('Prime Coat',37.47))
        r5=st.columns([3.5,1.2,1.8])
        with r5[0]: st.markdown(f"<div style='padding:8px 0;color:#0f2942'>Prime Coat <span style='color:#6b7a8d;font-size:.82rem'>({proj_area:,.0f} ตร.ม.)</span></div>",unsafe_allow_html=True)
        with r5[1]: st.markdown("<div style='padding:8px 0;color:#94a3b8;font-size:.85rem'>auto</div>",unsafe_allow_html=True)
        with r5[2]: st.markdown(f"<div style='padding:8px 0;font-weight:700;color:#0f2942;text-align:right'>{prime_price:,.2f}</div>",unsafe_allow_html=True)
        updated_layers.append({'name':'Prime Coat','thickness':1,'unit':'Layer',
            'quantity':proj_area,'qty_unit':'sq.m','unit_cost':prime_price,'cost_per_sqm':prime_price})

    else:
        # Concrete slab
        slab_name=f'Concrete Slab ({ptype})'
        _sh=st.columns([2,1,1.5])
        _sh[0].markdown("<span style='color:#6b7a8d;font-size:.82rem;font-weight:600'>รายการ</span>",unsafe_allow_html=True)
        _sh[1].markdown("<span style='color:#6b7a8d;font-size:.82rem;font-weight:600'>หนา (cm)</span>",unsafe_allow_html=True)
        _sh[2].markdown("<div style='color:#6b7a8d;font-size:.82rem;font-weight:600;text-align:right'>ราคา (บาท/ตร.ม.)</div>",unsafe_allow_html=True)
        c1,c2,c3=st.columns([2,1,1.5])
        with c1: st.markdown(f"**{slab_name}**")
        with c2:
            _slab_key=f"{key_prefix}_sthick_slab_v{v}"
            _def_t=28.0 if ptype in ('JPCP','JRCP') else 25.0
            if _slab_key not in st.session_state: st.session_state[_slab_key]=_def_t
            st.number_input("ความหนา (cm)",min_value=15.0,max_value=50.0,step=1.0,format="%.0f",
                key=_slab_key, label_visibility="collapsed")
        slab_thick=float(st.session_state[_slab_key])
        with c3:
            slab_price=_price(slab_name,slab_thick)
            st.markdown(f"<div style='padding:8px 0;font-weight:700;color:#0f2942;text-align:right'>{slab_price:,.2f}</div>",unsafe_allow_html=True)
        updated_layers.append({'name':slab_name,'thickness':slab_thick,'unit':'cm',
            'quantity':proj_area,'qty_unit':'sq.m','unit_cost':slab_price,'cost_per_sqm':slab_price})

    # ── SECTION B: วัสดุประกอบ (Concrete เท่านั้น) ───────────────────────────
    if is_concrete:
        st.markdown("---")
        st.markdown("**🔧 วัสดุประกอบ** — ติ๊กเลือกและปรับแก้ได้")
        col_cb1,col_cb2=st.columns(2)
        with col_cb1:
            use_acil=st.checkbox("AC Interlayer รองใต้แผ่นคอนกรีต",value=True,
                key=f"{key_prefix}_use_acil_v{v}")
        with col_cb2:
            use_pc=st.checkbox("Prime Coat",value=True,
                key=f"{key_prefix}_use_pc_v{v}")
        col_cb3,col_cb4=st.columns(2)
        with col_cb3:
            use_geo=st.checkbox("Non Woven Geotextile",value=True,
                key=f"{key_prefix}_use_geo_v{v}")
        with col_cb4:
            use_wire=st.checkbox(
                "Wire Mesh" if ptype!='JPCP' else "Wire Mesh (ไม่ใช้กับ JPCP)",
                value=(ptype!='JPCP'), disabled=(ptype=='JPCP'),
                key=f"{key_prefix}_use_wire_v{v}")

        if use_acil:
            c1,c2=st.columns([2,2])
            _acil_thick_key=f"{key_prefix}_acil_thick_v{v}"
            if _acil_thick_key not in st.session_state: st.session_state[_acil_thick_key]=5.0
            with c1: st.number_input("ความหนา AC Interlayer (cm)",min_value=1.0,max_value=15.0,step=1.0,key=_acil_thick_key)
            acil_thick=float(st.session_state[_acil_thick_key])
            with c2:
                acil_price=_price('AC Binder Course',acil_thick)
                if acil_price==0: acil_price=251.0
                st.markdown(f"AC Interlayer {acil_thick:.0f} cm → **{acil_price:,.2f}** บาท/ตร.ม.")
            updated_layers.append({'name':f'AC Interlayer ({acil_thick:.0f} cm)',
                'thickness':acil_thick,'unit':'cm','quantity':proj_area,'qty_unit':'sq.m',
                'unit_cost':acil_price,'cost_per_sqm':acil_price})

        if use_pc:
            pc_price=float(lib['base_prices'].get('Prime Coat',37.47))
            c1,c2=st.columns([3,1])
            with c1: st.caption("Prime Coat — ราดบน Base Course ก่อนปู AC Interlayer")
            with c2: st.markdown(f"<div style='padding:6px 0;font-weight:700;color:#0f2942;text-align:right'>{pc_price:,.2f}</div>",unsafe_allow_html=True)
            updated_layers.append({'name':'Prime Coat','thickness':1,'unit':'Layer',
                'quantity':proj_area,'qty_unit':'sq.m','unit_cost':pc_price,'cost_per_sqm':pc_price})

        if use_geo:
            geo_price=float(lib['base_prices'].get('Non Woven Geotextile',78))
            c1,c2=st.columns([3,1])
            with c1: st.caption("Non Woven Geotextile — รองใต้แผ่นคอนกรีต")
            with c2: st.markdown(f"<div style='padding:6px 0;font-weight:700;color:#0f2942;text-align:right'>{geo_price:,.2f}</div>",unsafe_allow_html=True)
            updated_layers.append({'name':'Non Woven Geotextile','thickness':1,'unit':'ชั้น',
                'quantity':proj_area,'qty_unit':'sq.m','unit_cost':geo_price,'cost_per_sqm':geo_price})

        if use_wire and ptype!='JPCP':
            wire_price=float(lib['base_prices'].get('Wire Mesh',100))
            c1,c2=st.columns([3,1])
            with c1: st.caption("Wire Mesh — ตะแกรงเหล็กในแผ่นคอนกรีต")
            with c2: st.markdown(f"<div style='padding:6px 0;font-weight:700;color:#0f2942;text-align:right'>{wire_price:,.2f}</div>",unsafe_allow_html=True)
            updated_layers.append({'name':'Wire Mesh','thickness':1,'unit':'ชั้น',
                'quantity':proj_area,'qty_unit':'sq.m','unit_cost':wire_price,'cost_per_sqm':wire_price})

    # ── SECTION C: พื้นทาง / รองพื้นทาง ─────────────────────────────────────
    st.markdown("---")
    st.caption("💡 ราคาดึงจาก Library อัตโนมัติ — แก้ราคาได้ที่ Sub-tab 💰 ราคาวัสดุ")

    sk_base_rows = f"{key_prefix}_base_rows_v{v}"
    sk_price_ver = f"{key_prefix}_base_price_ver_v{v}"
    sk_copy_flag = f"{key_prefix}_do_copy_base_v{v}"
    cur_price_ver = id(st.session_state.get('price_library',{}))

    hcol1,hcol2=st.columns([3,1])
    with hcol1:
        st.markdown('<div class="sc"><b>🪨 พื้นทาง / รองพื้นทาง</b> <span style="color:#6b7a8d;font-size:.83rem">บาท/ลบ.ม. × ความหนา = บาท/ตร.ม.</span></div>',unsafe_allow_html=True)
    with hcol2:
        if is_concrete and ptype!='JPCP':
            if st.button("📋 คัดลอก Base จาก JPCP",key=f"{key_prefix}_copy_base_v{v}",
                         use_container_width=True,type="secondary"):
                st.session_state[sk_copy_flag]=True

    if st.session_state.get(sk_copy_flag):
        jpcp_sk=f"jpcp_base_rows_v{v}"
        if jpcp_sk in st.session_state:
            st.session_state[sk_base_rows]=[dict(r) for r in st.session_state[jpcp_sk]]
            _cv_key=f"{key_prefix}_base_cv"
            st.session_state[_cv_key]=st.session_state.get(_cv_key,0)+1
            st.session_state[sk_copy_flag]=False
            st.success("✅ คัดลอก Base จาก JPCP สำเร็จ")
        else:
            st.warning("⚠️ ยังไม่มีข้อมูล Base ของ JPCP — กรุณาตั้งค่า JPCP ก่อน")
            st.session_state[sk_copy_flag]=False

    _lib_cum={m:lookup_price(m,20) for m in BASE_MATERIAL_LIST}

    if sk_base_rows not in st.session_state:
        _def_base=get_default_base_layers(ptype,area_per_km)
        st.session_state[sk_base_rows]=[{'name':r['name'],'thickness':r['thickness'],'cost_cum':r['cost_cum']} for r in _def_base]
        st.session_state[sk_price_ver]=cur_price_ver
    elif st.session_state.get(sk_price_ver)!=cur_price_ver:
        st.session_state[sk_base_rows]=[
            {'name':r['name'],'thickness':r['thickness'],
             'cost_cum':_lib_cum.get(r['name'],lookup_price(r['name'],20))}
            for r in st.session_state[sk_base_rows]
        ]
        st.session_state[sk_price_ver]=cur_price_ver

    _cv=st.session_state.get(f"{key_prefix}_base_cv",0)
    _cv_sf=f"cv{_cv}"
    _cur_rows=st.session_state[sk_base_rows]
    _nb_key=f"{key_prefix}_num_base_{_cv_sf}_v{v}"
    if _nb_key not in st.session_state: st.session_state[_nb_key]=len(_cur_rows)
    st.number_input("จำนวนชั้นพื้นทาง/รองพื้นทาง",min_value=0,max_value=8,step=1,key=_nb_key)
    num_base=int(st.session_state[_nb_key])

    _base_cols=[3,1.2,1.5,1.5]
    hdr=st.columns(_base_cols)
    _hs="color:#6b7a8d;font-size:.82rem;font-weight:600"
    hdr[0].markdown(f"<span style='{_hs}'>วัสดุ</span>",unsafe_allow_html=True)
    hdr[1].markdown(f"<span style='{_hs}'>หนา (cm)</span>",unsafe_allow_html=True)
    hdr[2].markdown(f"<span style='{_hs};display:block;text-align:right'>ราคา (บาท/ลบ.ม.)</span>",unsafe_allow_html=True)
    hdr[3].markdown(f"<span style='{_hs};display:block;text-align:right'>ราคา (บาท/ตร.ม.)</span>",unsafe_allow_html=True)

    sk_prev_names=f"{key_prefix}_prev_names_v{v}"
    prev_names=st.session_state.get(sk_prev_names,{})
    new_rows=[]
    for i in range(num_base):
        prev=_cur_rows[i] if i<len(_cur_rows) else {'name':BASE_MATERIAL_LIST[0],'thickness':20.0,'cost_cum':_lib_cum.get(BASE_MATERIAL_LIST[0],0)}
        prev_name=str(prev.get('name',BASE_MATERIAL_LIST[0]))
        prev_thick=float(prev.get('thickness',20.0) or 20.0)
        prev_cum=float(prev.get('cost_cum',0) or 0)
        if prev_cum==0: prev_cum=_lib_cum.get(prev_name,0)
        cols=st.columns(_base_cols)
        with cols[0]:
            try: name_idx=BASE_MATERIAL_LIST.index(prev_name)
            except ValueError: name_idx=0
            sel_name=st.selectbox("วัสดุ",BASE_MATERIAL_LIST,index=name_idx,
                key=f"{key_prefix}_bname_{i}_{_cv_sf}_v{v}",label_visibility="collapsed")
        with cols[1]:
            bthick_key=f"{key_prefix}_bthick_{i}_{_cv_sf}_v{v}"
            if bthick_key not in st.session_state: st.session_state[bthick_key]=float(prev_thick)
            st.number_input("หนา",min_value=0.0,step=5.0,format="%.0f",
                key=bthick_key,label_visibility="collapsed")
            sel_thick=float(st.session_state.get(bthick_key,prev_thick))

        last_rendered=prev_names.get(i,prev_name)
        if sel_name!=last_rendered:
            wkey=f"{key_prefix}_bcum_{i}_{_cv_sf}_v{v}"
            if wkey in st.session_state: del st.session_state[wkey]
            prev_cum=_lib_cum.get(sel_name,0)

        bcum_key=f"{key_prefix}_bcum_{i}_{_cv_sf}_v{v}"
        if bcum_key not in st.session_state: st.session_state[bcum_key]=float(prev_cum)

        sel_cum=float(_lib_cum.get(sel_name,0))
        with cols[2]:
            st.markdown(f'<div style="padding:8px 0;font-weight:700;color:#0f2942;text-align:right">{sel_cum:,.0f}</div>',unsafe_allow_html=True)
        cost_sqm_base=sel_cum*sel_thick/100 if sel_thick>0 else 0.0
        with cols[3]:
            st.markdown(f'<div style="padding:8px 0;font-weight:700;color:#0f2942;text-align:right">{cost_sqm_base:,.2f}</div>',unsafe_allow_html=True)

        new_rows.append({'name':sel_name,'thickness':sel_thick,'cost_cum':sel_cum})
        if sel_thick>0 and sel_name:
            updated_layers.append({'name':sel_name,'thickness':sel_thick,'unit':'cm',
                'quantity':proj_area,'qty_unit':'sq.m',
                'unit_cost':cost_sqm_base,'cost_per_sqm':cost_sqm_base,'cost_cum':sel_cum})

    st.session_state[sk_base_rows]=new_rows
    st.session_state[sk_prev_names]={i:r['name'] for i,r in enumerate(new_rows)}
    if ptype=='JPCP': st.session_state[f"jpcp_base_rows_v{v}"]=[dict(r) for r in new_rows]
    return updated_layers

# ══════════════════════════════════════════════════════════════════════════════
# ─── RENDER JOINT EDITOR ─────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

def render_joint_editor(ptype: str, key_prefix: str,
                        area_per_km: float, road_length: float, v: int = 0) -> tuple:
    total_area=area_per_km*road_length
    width_m=area_per_km/1000
    lane_w=st.session_state.get('sb_lane_width',3.5)
    spacing={'JPCP':4,'JRCP':10,'CRCP':0}.get(ptype,4)
    if ptype in ('JPCP','JRCP'):
        auto_trans=(road_length*1000/spacing)*width_m
        auto_long=max(1,round(width_m/lane_w)-1)*road_length*1000
    else:
        auto_trans=0; auto_long=road_length*1000

    sk_joint_init=f"{key_prefix}_joint_init_v{v}"
    if sk_joint_init not in st.session_state:
        defaults_j=get_default_joints(ptype,area_per_km,road_length)
        rows_init=[]
        for j in defaults_j:
            name=j['name']; qty=j['quantity']
            if 'transverse' in name.lower(): qty=auto_trans if ptype in ('JPCP','JRCP') else qty
            elif 'longitudinal' in name.lower() or 'steel' in name.lower(): qty=auto_long
            rows_init.append({'name':name,'quantity':qty,'unit_cost':float(j['unit_cost'])})
        st.session_state[sk_joint_init]=pd.DataFrame(rows_init)

    st.markdown("---")
    col_h1,col_h2=st.columns([3,1])
    with col_h1:
        if ptype=='CRCP':
            st.markdown('<div class="sc"><b>⛓️ Longitudinal Steel & Transverse Joint (CRCP)</b></div>',unsafe_allow_html=True)
        else:
            st.markdown(f'<div class="sc"><b>🔗 รอยต่อ (Joints) — {ptype} ระยะ {spacing} ม.</b></div>',unsafe_allow_html=True)
    with col_h2:
        _cb_label="รวมราคา Steel & Joints" if ptype=='CRCP' else "รวมราคา Joints"
        include_joints=st.checkbox(_cb_label,value=True,key=f"{key_prefix}_include_joints_v{v}")

    ek_joint=f"{key_prefix}_joint_editor_v{v}"
    st.data_editor(st.session_state[sk_joint_init],
        column_config={
            'name':st.column_config.TextColumn('รายการ',width='large'),
            'quantity':st.column_config.NumberColumn('ปริมาณ (ม.)',min_value=0.0,step=100.0,format='%.0f'),
            'unit_cost':st.column_config.NumberColumn('ราคา/ม. (บาท)',min_value=0.0,step=10.0,format='%.0f'),
        }, num_rows='dynamic', use_container_width=True,
        key=ek_joint, hide_index=True)

    _jstate=st.session_state.get(ek_joint,{})
    edited_joint=st.session_state[sk_joint_init].copy()
    if isinstance(_jstate,dict):
        for idx_str,changes in _jstate.get("edited_rows",{}).items():
            idx=int(idx_str)
            if idx<len(edited_joint):
                for col,val in changes.items(): edited_joint.at[idx,col]=val
        for new_row in _jstate.get("added_rows",[]):
            edited_joint=pd.concat([edited_joint,pd.DataFrame([new_row])],ignore_index=True)
        del_idxs=_jstate.get("deleted_rows",[])
        if del_idxs: edited_joint=edited_joint.drop(index=del_idxs).reset_index(drop=True)

    joint_total_cost=0.0; updated_joints=[]
    for _,row in edited_joint.iterrows():
        qty=float(row.get('quantity',0) or 0)
        uc=float(row.get('unit_cost',0) or 0)
        cpsqm=(qty*uc/total_area) if total_area>0 else 0.0
        joint_total_cost+=qty*uc
        updated_joints.append({'name':str(row.get('name','') or ''),'quantity':qty,
            'qty_unit':'m','unit_cost':uc,'cost_per_sqm':cpsqm})
    if total_area>0:
        st.caption(f"รวม Joints = **{joint_total_cost/total_area:,.2f}** บาท/ตร.ม. | **{joint_total_cost/1e6:,.3f}** ล้านบาท/โครงการ")
    return updated_joints, include_joints

# ══════════════════════════════════════════════════════════════════════════════
# ─── EXCEL PRICE LIBRARY ─────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

def generate_excel_template() -> bytes:
    lib=get_price_library()
    ac_ton=st.session_state.get('ac_ton_prices',DEFAULT_AC_TON_PRICES)
    ac_rows=[{'Material':mat,'Price (Baht/ton)':float(ac_ton.get(mat,0))} for mat in DEFAULT_AC_TON_PRICES]
    _conc_cum=st.session_state.get('concrete_cum_prices',DEFAULT_CONCRETE_CUM_PRICES)
    conc_rows=[{'ประเภท':ct,'บาท/ลบ.ม.':float(_conc_cum.get(ct,DEFAULT_CONCRETE_CUM_PRICES.get(ct,0)))} for ct in ['JPCP','JRCP','CRCP']]
    base_rows=[{'Material':k,'Price':v,'Unit':'Baht/sq.m' if k in SQMKEYS else 'Baht/cu.m'} for k,v in lib['base_prices'].items()]
    output=io.BytesIO()
    with pd.ExcelWriter(output,engine='openpyxl') as writer:
        pd.DataFrame(ac_rows).to_excel(writer,sheet_name='AC_Prices',index=False)
        pd.DataFrame(conc_rows).to_excel(writer,sheet_name='Concrete_Prices',index=False)
        pd.DataFrame(base_rows).to_excel(writer,sheet_name='Base_Materials',index=False)
    output.seek(0); return output.getvalue()


def load_excel_price_library(uploaded_file) -> dict:
    ac_df=pd.read_excel(uploaded_file,sheet_name='AC_Prices')
    conc_df=pd.read_excel(uploaded_file,sheet_name='Concrete_Prices')
    base_df=pd.read_excel(uploaded_file,sheet_name='Base_Materials')
    ac_ton:dict=dict(DEFAULT_AC_TON_PRICES)
    if 'Price (Baht/ton)' in ac_df.columns:
        for _,row in ac_df.iterrows():
            try:
                mat=str(row['Material']); val=row['Price (Baht/ton)']
                if pd.notna(mat) and pd.notna(val) and float(val)>0: ac_ton[mat]=float(val)
            except: pass
        st.session_state['ac_ton_prices']=ac_ton
    density=st.session_state.get('tab2_density',DEFAULT_AC_DENSITY)
    ac_prices=_calc_ac_prices_from_ton(ac_ton,density)
    for mat,dp in DEFAULT_AC_PRICES.items():
        if mat not in ac_prices: ac_prices[mat]=dict(dp)
    conc_cum:dict=dict(DEFAULT_CONCRETE_CUM_PRICES)
    _type_col='ประเภท' if 'ประเภท' in conc_df.columns else 'Type'
    _price_col=next((c for c in ['บาท/ลบ.ม.','Price (Baht/cu.m)','Price'] if c in conc_df.columns),None)
    if _type_col in conc_df.columns and _price_col:
        for _,row in conc_df.iterrows():
            try:
                ct=str(row[_type_col]).strip(); val=row[_price_col]
                if pd.notna(ct) and pd.notna(val) and float(val)>0: conc_cum[ct]=float(val)
            except: pass
    st.session_state['concrete_cum_prices']=conc_cum
    conc_prices=_calc_concrete_prices(conc_cum)
    base_prices:dict=dict(DEFAULT_BASE_PRICES)
    _bmat_col='Material' if 'Material' in base_df.columns else 'วัสดุ'
    _bprice_col=next((c for c in ['Price','Price (Baht/cu.m)','ราคา (บาท/ลบ.ม.)','ราคา (บาท/ตร.ม.)'] if c in base_df.columns),None)
    if _bmat_col in base_df.columns and _bprice_col:
        for _,row in base_df.iterrows():
            try:
                mat=str(row[_bmat_col]).strip(); val=row[_bprice_col]
                if pd.notna(mat) and pd.notna(val) and float(val)>0: base_prices[mat]=float(val)
            except: pass
    return {'ac_prices':ac_prices,'concrete_prices':conc_prices,'base_prices':base_prices}

# ══════════════════════════════════════════════════════════════════════════════
# ─── LCCA DATA STRUCTURES & CALCULATION ──────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

@dataclass
class MaintAct:
    name: str; unit_cost: float; start_year: int; frequency: int = 0

@dataclass
class RehabAct:
    name: str; unit_cost: float; year: int

@dataclass
class PavAlt:
    name: str; pave_type: str; construction_cost: float; area: float
    maintenance: List[MaintAct]; rehab: List[RehabAct]
    salvage_pct: float = 20.0; enabled: bool = True


def lookup_range(value, breaks):
    for lo,hi,v in breaks:
        if lo<=value<=hi: return v
    return breaks[-1][2]


def calc_Ka_average(x1,cbr,x3,x4_start,x5_width,x6_code,y1_row,y2_shoulder,terrain_code,y5_bridge,n_years):
    X1=x1; X2=lookup_range(cbr,X2_BREAKS); X3=x3
    X5=lookup_range(x5_width,X5_BREAKS); X6=X6_MAP[x6_code]
    Y1=lookup_range(y1_row,Y1_BREAKS); Y2=lookup_range(y2_shoulder,Y2_BREAKS)
    Y3=Y3_MAP[terrain_code]; Y4=Y4_MAP[terrain_code]
    Y5=lookup_range(y5_bridge,Y5_BREAKS); Y6=Y6_MAP[terrain_code]
    ka_list,rows=[],[]
    for yr in range(1,n_years+1):
        age=x4_start+(yr-1); X4=lookup_range(age,X4_BREAKS)
        Ka=1+0.50*(X1+X2+X3+X4+X5+X6+Y1+Y2+Y3+Y4+Y5+Y6)
        ka_list.append(Ka); rows.append({"ปี":yr,"อายุ (ปี)":age,"X4":X4,"Ka":round(Ka,4)})
    fixed={"X1":X1,"X2":X2,"X3":X3,"X5":X5,"X6":X6,"Y1":Y1,"Y2":Y2,"Y3":Y3,"Y4":Y4,"Y5":Y5,"Y6":Y6}
    return round(np.mean(ka_list),4), pd.DataFrame(rows), fixed


def calc_Kc(z1_idx,cbr,z3,z4_width,y1_row,y2_shoulder,terrain_code,y5_bridge):
    Z1=Z1_MAP.get(z1_idx,0); Z2=lookup_range(cbr,Z2_BREAKS)
    Z3=z3; Z4=lookup_range(z4_width,Z4_BREAKS)
    Y1=lookup_range(y1_row,Y1_BREAKS); Y2=lookup_range(y2_shoulder,Y2_BREAKS)
    Y3=Y3_MAP[terrain_code]; Y4=Y4_MAP[terrain_code]
    Y5=lookup_range(y5_bridge,Y5_BREAKS); Y6=Y6_MAP[terrain_code]
    Kc=1+0.50*(Z1+Z2+Z3+Z4+Y1+Y2+Y3+Y4+Y5+Y6)
    factors={"Z1":Z1,"Z2":Z2,"Z3":Z3,"Z4":Z4,"Y1":Y1,"Y2":Y2,"Y3":Y3,"Y4":Y4,"Y5":Y5,"Y6":Y6}
    return round(Kc,4), factors


def calc_pv(cost,yr,dr): return cost*(1+dr)**(-yr) if yr>=0 else 0.0
def calc_eac(pw,dr,n):
    if n<=0 or dr<=0: return 0.0
    return pw*dr*(1+dr)**n/((1+dr)**n-1)


def build_cashflow(alt: PavAlt, n:int, dr:float, inc_salvage:bool) -> pd.DataFrame:
    rows=[]; area=alt.area
    rehab_yrs=sorted([r.year for r in alt.rehab if r.year<=n])
    rehab_set=set(rehab_yrs)
    c0=alt.construction_cost*area
    rows.append({"ปี":0,"กิจกรรม":"ก่อสร้างเริ่มต้น","ประเภท":"ก่อสร้าง",
        "ต้นทุน/หน่วย":alt.construction_cost,"ต้นทุนตามปี":c0,"PW_factor":1.0,"มูลค่าปัจจุบัน":c0})
    for m in alt.maintenance:
        if m.frequency>0:
            cps=[0]+rehab_yrs
            for idx,cp in enumerate(cps):
                end=cps[idx+1] if idx+1<len(cps) else n+1
                yr=cp+m.frequency
                while yr<end and yr<=n:
                    if yr not in rehab_set:
                        c=m.unit_cost*area; pwf=(1+dr)**(-yr)
                        rows.append({"ปี":yr,"กิจกรรม":m.name,"ประเภท":"บำรุงรักษา",
                            "ต้นทุน/หน่วย":m.unit_cost,"ต้นทุนตามปี":c,"PW_factor":pwf,"มูลค่าปัจจุบัน":c*pwf})
                    yr+=m.frequency
        else:
            if m.start_year<=n and m.start_year not in rehab_set:
                c=m.unit_cost*area; pwf=(1+dr)**(-m.start_year)
                rows.append({"ปี":m.start_year,"กิจกรรม":m.name,"ประเภท":"บำรุงรักษา",
                    "ต้นทุน/หน่วย":m.unit_cost,"ต้นทุนตามปี":c,"PW_factor":pwf,"มูลค่าปัจจุบัน":c*pwf})
    last_cost,last_yr=alt.construction_cost*area,0
    for r in alt.rehab:
        if r.year<=n:
            c=r.unit_cost*area; pwf=(1+dr)**(-r.year)
            rows.append({"ปี":r.year,"กิจกรรม":r.name,"ประเภท":"ฟื้นฟูสภาพ",
                "ต้นทุน/หน่วย":r.unit_cost,"ต้นทุนตามปี":c,"PW_factor":pwf,"มูลค่าปัจจุบัน":c*pwf})
            last_cost,last_yr=c,r.year
    if inc_salvage:
        life={"Flexible":15,"AC":15,"JPCP":20,"JRCP":20,"CRCP":25}
        exp=next((v for k,v in life.items() if k in alt.pave_type),20)
        dep=last_cost*(1-alt.salvage_pct/100)/exp
        sv=max(last_cost-dep*(n-last_yr),last_cost*alt.salvage_pct/100)
        pwf=(1+dr)**(-n)
        rows.append({"ปี":n,"กิจกรรม":"มูลค่าซาก","ประเภท":"มูลค่าซาก",
            "ต้นทุน/หน่วย":-sv/area,"ต้นทุนตามปี":-sv,"PW_factor":pwf,"มูลค่าปัจจุบัน":-sv*pwf})
    return pd.DataFrame(rows).sort_values(["ปี","กิจกรรม"]).reset_index(drop=True)


def analyze_lcca(alts, n, dr, inc_salvage):
    rows,cf_dict=[],{}
    for alt in [a for a in alts if a.enabled]:
        cf=build_cashflow(alt,n,dr,inc_salvage); cf_dict[alt.name]=cf
        pw=cf["มูลค่าปัจจุบัน"].sum(); eac=calc_eac(pw,dr,n); area=alt.area
        rows.append({"ทางเลือก":alt.name,"ประเภทผิวทาง":alt.pave_type,"พื้นที่ (ตร.ม./กม.)":area,
            "ต้นทุนก่อสร้าง (บาท/ตร.ม.)":alt.construction_cost,
            "ต้นทุนก่อสร้าง (ล้านบาท/กม.)":round(alt.construction_cost*area/1e6,4),
            "PW_ก่อสร้าง":cf[cf["ประเภท"]=="ก่อสร้าง"]["มูลค่าปัจจุบัน"].sum(),
            "PW_บำรุงรักษา":cf[cf["ประเภท"]=="บำรุงรักษา"]["มูลค่าปัจจุบัน"].sum(),
            "PW_ฟื้นฟูสภาพ":cf[cf["ประเภท"]=="ฟื้นฟูสภาพ"]["มูลค่าปัจจุบัน"].sum(),
            "PW_มูลค่าซาก":cf[cf["ประเภท"]=="มูลค่าซาก"]["มูลค่าปัจจุบัน"].sum(),
            "NPV (บาท/กม.)":pw,"NPV (ล้านบาท/กม.)":round(pw/1e6,4),
            "EAC (บาท/กม./ปี)":eac,"EAC (ล้านบาท/กม./ปี)":round(eac/1e6,4),
            "EAC (บาท/ตร.ม./ปี)":eac/area if area>0 else 0})
    df=pd.DataFrame(rows)
    if len(df)>0:
        df=df.sort_values("NPV (บาท/กม.)").reset_index(drop=True)
        df.insert(0,"อันดับ",range(1,len(df)+1))
    return df, cf_dict


def calc_breakeven(cf_dict,n,dr):
    names=list(cf_dict.keys()); results=[]
    for a,b in combinations(names,2):
        cf_a=cf_dict[a]; cf_b=cf_dict[b]
        cum_a,cum_b=0.0,0.0; be_yr=None; prev_diff=None
        for yr in range(0,n+1):
            cum_a+=cf_a[cf_a["ปี"]==yr]["มูลค่าปัจจุบัน"].sum()
            cum_b+=cf_b[cf_b["ปี"]==yr]["มูลค่าปัจจุบัน"].sum()
            diff=cum_a-cum_b
            if prev_diff is not None and prev_diff*diff<0: be_yr=yr; break
            prev_diff=diff
        results.append({"คู่เปรียบเทียบ":f"{a} vs {b}",
            "Breakeven Year":be_yr if be_yr else f">{n}",
            "หมายเหตุ":f"{b} คุ้มกว่า {a} หลังปีที่ {be_yr}" if be_yr else f"ไม่มี crossover ใน {n} ปี"})
    return pd.DataFrame(results)


def build_cumulative(cf_dict,n,dr):
    rows=[]
    for name,cf in cf_dict.items():
        cum=0.0
        for yr in range(0,n+1):
            cum+=cf[cf["ปี"]==yr]["มูลค่าปัจจุบัน"].sum()
            rows.append({"ปี":yr,"ทางเลือก":name,"Cumulative NPV (บาท)":cum})
    return pd.DataFrame(rows)

# ══════════════════════════════════════════════════════════════════════════════
# ─── WORD REPORT — COMBINED (Cost Structure + Routine Cost + LCCA) ────────────
# ══════════════════════════════════════════════════════════════════════════════

def _set_rf(run, size:int=16, bold:bool=False, italic:bool=False):
    if not DOCX_OK: return
    run.font.name='TH SarabunPSK'; run.font.size=Pt(size)
    run.font.bold=bold; run.font.italic=italic
    rPr=run._r.get_or_add_rPr(); rFonts=rPr.get_or_add_rFonts()
    for attr in ('w:eastAsia','w:ascii','w:hAnsi'): rFonts.set(qn(attr),'TH SarabunPSK')


def _add_thai_para(doc, text="", bold=False, first_indent=True):
    if not DOCX_OK: return
    p=doc.add_paragraph()
    pPr=p._p.get_or_add_pPr()
    jc=OxmlElement("w:jc"); jc.set(qn("w:val"),"thaiDistribute"); pPr.append(jc)
    if first_indent:
        ind=OxmlElement("w:ind"); ind.set(qn("w:firstLine"),"720"); pPr.append(ind)
    if text:
        run=p.add_run(text); _set_rf(run,bold=bold)
    return p


def _add_hdg_w(doc, text, level=1, size=None):
    p=doc.add_heading(text, level=level)
    sz=size if size else (16 if level==1 else 15)
    for r in p.runs: _set_rf(r, size=sz, bold=True)
    return p


def _add_tbl_w(doc, headers, rows, col_widths=None):
    if not DOCX_OK: return
    t=doc.add_table(rows=1,cols=len(headers)); t.style="Table Grid"
    t.alignment=WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for i,w in enumerate(col_widths): t.columns[i].width=Cm(w)
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.paragraphs[0].clear()
        _set_rf(c.paragraphs[0].add_run(str(h)),bold=True)
    for rd in rows:
        row=t.add_row()
        for i,val in enumerate(rd):
            c=row.cells[i]; c.paragraphs[0].clear()
            _set_rf(c.paragraphs[0].add_run(str(val)))
    return t


def generate_word_combined(
    project_info: dict,
    all_details:  dict,
    summary_df,
    cf_dict:      dict,
    n:            int,
    dr:           float,
    alts:         list,
    ss:           dict,
    base_sec:     str = "4.7",
    intro_text:   str = "",
) -> io.BytesIO:
    """
    รายงาน Word รวม แบบที่ปรึกษา — 3 ส่วนในไฟล์เดียว
      {base_sec}   ราคาก่อสร้างโครงสร้างชั้นทาง
      {base+1}     งบประมาณบำรุงรักษาประจำปี (Routine Cost)
      {base+2}     วิเคราะห์ต้นทุนตลอดอายุการใช้งาน (LCCA)
    """
    if not DOCX_OK: raise ImportError("python-docx ไม่พร้อม")

    def _next_sec(s: str, offset: int) -> str:
        parts = s.strip().split(".")
        try:
            return ".".join(parts[:-1] + [str(int(parts[-1]) + offset)])
        except ValueError:
            return s

    sec_cost    = base_sec
    sec_routine = _next_sec(base_sec, 1)
    sec_lcca    = _next_sec(base_sec, 2)

    doc = WordDoc()
    doc.styles['Normal'].font.name = 'TH SarabunPSK'
    doc.styles['Normal'].font.size = Pt(16)
    for sec in doc.sections:
        sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5)
        sec.left_margin=Cm(3.0); sec.right_margin=Cm(2.5)

    # ── helpers ──────────────────────────────────────────────────────────────
    def hdg(text, size=16, bold=True, uline=False, sb=8, sa=3):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(sb)
        para.paragraph_format.space_after  = Pt(sa)
        run = para.add_run(text)
        _set_rf(run, size=size, bold=bold)
        run.underline = uline
        return para

    def body(text, indent=True):
        """ย่อหน้าเนื้อหา — Thai justify + เยื้องบรรทัดแรก"""
        _add_thai_para(doc, text, first_indent=indent)

    def make_sc(base):
        class SC:
            prefix = base.strip(); h1 = 0; h2 = 0
        return SC

    def next_h1(SC, title):
        SC.h1 += 1; SC.h2 = 0
        _add_hdg_w(doc, f"{SC.prefix}.{SC.h1}  {title}", 1)

    def next_h2(SC, title):
        SC.h2 += 1
        _add_hdg_w(doc, f"{SC.prefix}.{SC.h1}.{SC.h2}  {title}", 2)

    # ── ดึงค่าจาก session state ───────────────────────────────────────────────
    length   = float(project_info.get('length', 1))
    tw       = float(project_info.get('total_width', 22))
    nl       = int(project_info.get('num_lanes', 4))
    dl       = int(project_info.get('design_life', 20))
    proj_nm  = project_info.get('name', 'โครงการ')
    ac_Na    = ss.get('lc_ac_Na', 35000);  ac_Km = ss.get('lc_ac_Km', 1.0)
    cc_Nc    = ss.get('lc_cc_Nc', 35000);  cc_Km = ss.get('lc_cc_Km', 1.0)
    ka_avg   = ss.get('lc_ka_avg', 0) or 0
    kc_val   = ss.get('lc_kc_val', 0) or 0
    r_ac_sqm = ss.get('lc_routine_ac_sqm', 0) or 0
    r_ac_km  = ss.get('lc_routine_ac_km',  0) or 0
    r_cc_sqm = ss.get('lc_routine_cc_sqm', 0) or 0
    r_cc_km  = ss.get('lc_routine_cc_km',  0) or 0
    ka_fixed = ss.get('_lc_ka_fixed', {})
    ka_df    = ss.get('_lc_ka_df')
    kc_fac   = ss.get('_lc_kc_fac', {})
    enabled_alts = [a for a in alts if a.enabled]

    # ════════════════════════════════════════════════════════════════════════
    # ปก
    # ════════════════════════════════════════════════════════════════════════
    for txt, sz, bl in [
        ("รายงานวิเคราะห์โครงสร้างชั้นทางและต้นทุนตลอดอายุการใช้งาน", 20, True),
        ("Pavement Structure Cost Analysis & Life-Cycle Cost Analysis", 16, True),
        (proj_nm, 15, True),
        (f"วันที่จัดทำ: {datetime.now().strftime('%d/%m/%Y')}", 13, False),
        ("จัดทำโดย รศ.ดร.อิทธิพล มีผล", 13, True),
        ("ภาควิชาครุศาสตร์โยธา คณะครุศาสตร์อุตสาหกรรม", 13, False),
        ("มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ (มจพ.)", 13, False),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_rf(p.add_run(txt), size=sz, bold=bl)
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # ส่วนที่ 1 — ราคาก่อสร้างโครงสร้างชั้นทาง
    # ════════════════════════════════════════════════════════════════════════
    hdg(f"{sec_cost}  ราคาก่อสร้างโครงสร้างชั้นทาง", size=18, uline=True, sb=14)
    body(
        f"การวิเคราะห์ราคาก่อสร้างโครงสร้างชั้นทางฉบับนี้จัดทำขึ้นเพื่อเปรียบเทียบ"
        f"ทางเลือกโครงสร้างชั้นทางประเภทต่าง ๆ สำหรับ{proj_nm} "
        f"ถนน {nl} ช่องจราจร ความกว้างรวม {tw:.2f} เมตร ระยะทาง {length:.2f} กิโลเมตร "
        f"อายุออกแบบ {dl} ปี โดยครอบคลุมทั้งผิวทางแอสฟัลต์คอนกรีต (AC) "
        f"และผิวทางคอนกรีตซีเมนต์ (JPCP, JRCP, CRCP) "
        f"การวิเคราะห์อ้างอิงราคาวัสดุและค่าก่อสร้างตามมาตรฐานกรมทางหลวง"
    )
    if intro_text:
        body(intro_text)
    doc.add_paragraph()

    # ── 1.1 ข้อมูลโครงการ ────────────────────────────────────────────────
    SCc = make_sc(sec_cost)
    next_h1(SCc, "ข้อมูลโครงการและขนาดถนน")
    _add_tbl_w(doc, ["รายการ", "ค่า"], [
        ["ชื่อโครงการ",      proj_nm],
        ["ระยะทางรวม",       f"{length:.2f} กม."],
        ["ความกว้างรวม",     f"{tw:.2f} ม."],
        ["จำนวนช่องจราจร",  f"{nl} ช่อง"],
        ["อายุออกแบบ",      f"{dl} ปี"],
        ["พื้นที่ทาง/กม.",   f"{tw*1000:,.0f} ตร.ม./กม."],
        ["พื้นที่ทางรวม",    f"{tw*1000*length:,.0f} ตร.ม."],
    ], col_widths=[7, 9])
    doc.add_paragraph()

    # ── 1.2 รายละเอียดวัสดุและราคา ───────────────────────────────────────
    next_h1(SCc, "รายละเอียดวัสดุและราคาโครงสร้างชั้นทางแต่ละประเภท")
    body(
        "ตารางต่อไปนี้แสดงรายการวัสดุ ปริมาณ และราคาโครงสร้างชั้นทางแต่ละประเภท "
        "โดยราคาวัสดุอ้างอิงตามมาตรฐานกรมบัญชีกลางและราคาท้องตลาดปัจจุบัน "
        "ราคาวัสดุพื้นทางคิดในหน่วย บาท/ลูกบาศก์เมตร คูณความหนาเพื่อให้ได้ บาท/ตารางเมตร"
    )
    doc.add_paragraph()
    summary_cost = []
    for ptype, data in all_details.items():
        sname   = data.get('name', ptype)
        details = [d for d in data.get('details', []) if float(d.get('มูลค่า (บาท)', 0)) != 0]
        hdg(f"ผิวทางประเภท {sname}", size=16, bold=True, sb=6, sa=2)
        if details:
            table = doc.add_table(rows=len(details)+2, cols=5)
            table.style = 'Table Grid'
            cw = [Cm(6.5), Cm(2.5), Cm(1.8), Cm(3.5), Cm(3.5)]
            for row in table.rows:
                for idx, cell in enumerate(row.cells): cell.width = cw[idx]
            for j, h in enumerate(['รายการ','ปริมาณ','หน่วย','ราคา/หน่วย (บาท)','มูลค่า (บาท)']):
                ph = table.rows[0].cells[j].paragraphs[0]
                ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_rf(ph.add_run(h), size=15, bold=True)
            subtotal = 0.0
            for i, d in enumerate(details):
                rc = table.rows[i+1].cells
                _set_rf(rc[0].paragraphs[0].add_run(str(d['รายการ'])), size=15)
                rc[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _set_rf(rc[1].paragraphs[0].add_run(f"{d['ปริมาณ']:,.0f}"), size=15)
                rc[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_rf(rc[2].paragraphs[0].add_run(d.get('หน่วย','ตร.ม.')), size=15)
                rc[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _set_rf(rc[3].paragraphs[0].add_run(d.get('ราคา/หน่วย (แสดง)','')), size=15)
                _set_rf(rc[3].paragraphs[0].add_run(f" ({d.get('หน่วยราคา','บาท/ตร.ม.')})"), size=12)
                rc[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _set_rf(rc[4].paragraphs[0].add_run(f"{d['มูลค่า (บาท)']:,.0f}"), size=15)
                subtotal += d['มูลค่า (บาท)']
            lr = table.rows[len(details)+1]
            lr.cells[0].merge(lr.cells[3])
            lr.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_rf(lr.cells[0].paragraphs[0].add_run(f"รวม {sname}"), size=15, bold=True)
            lr.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_rf(lr.cells[4].paragraphs[0].add_run(f"{subtotal:,.0f}"), size=15, bold=True)
            doc.add_paragraph()
            summary_cost.append({
                'name': sname, 'total': subtotal,
                'cost_per_km': data.get('cost_per_km', 0),
                'cost_sqm':    data.get('cost_sqm', 0),
            })

    # ── 1.3 สรุปเปรียบเทียบราคา ──────────────────────────────────────────
    next_h1(SCc, "สรุปและเปรียบเทียบราคาก่อสร้างโครงสร้างชั้นทาง")
    body(
        "ตารางสรุปด้านล่างแสดงราคาก่อสร้างโครงสร้างชั้นทางแต่ละประเภทในหน่วย "
        "บาท/ตารางเมตร และ ล้านบาท/กิโลเมตร เพื่อใช้เปรียบเทียบและประกอบการ"
        "ตัดสินใจเลือกโครงสร้างที่เหมาะสม"
    )
    if summary_cost:
        st2 = doc.add_table(rows=len(summary_cost)+1, cols=4)
        st2.style = 'Table Grid'
        for j, h in enumerate(['ชนิดโครงสร้าง','มูลค่ารวม/กม. (บาท)','ราคา/กม. (ล้านบาท)','ราคา/ตร.ม. (บาท)']):
            ph = st2.rows[0].cells[j].paragraphs[0]
            ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_rf(ph.add_run(h), size=15, bold=True)
        for i, item in enumerate(summary_cost):
            tpk = item['total']/length if length > 0 else 0
            vals   = [item['name'], f"{tpk:,.0f}", f"{item['cost_per_km']:.3f}", f"{item['cost_sqm']:,.2f}"]
            aligns = [WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.RIGHT]*3
            for j, (val, align) in enumerate(zip(vals, aligns)):
                cell = st2.rows[i+1].cells[j]; cell.text = ''
                p = cell.paragraphs[0]; p.alignment = align
                _set_rf(p.add_run(val), size=15)
        doc.add_paragraph()
        # สรุปย่อหน้าท้ายส่วนที่ 1
        if len(summary_cost) >= 2:
            sorted_c = sorted(summary_cost, key=lambda x: x['cost_sqm'])
            cheapest = sorted_c[0]; priciest = sorted_c[-1]
            body(
                f"จากการวิเคราะห์ราคาก่อสร้าง พบว่าโครงสร้างชั้นทางประเภท {cheapest['name']} "
                f"มีราคาก่อสร้างต่ำที่สุดที่ {cheapest['cost_sqm']:,.2f} บาท/ตร.ม. "
                f"({cheapest['cost_per_km']:.3f} ล้านบาท/กม.) "
                f"ในขณะที่ {priciest['name']} มีราคาสูงที่สุดที่ {priciest['cost_sqm']:,.2f} บาท/ตร.ม. "
                f"อย่างไรก็ตาม ราคาก่อสร้างเพียงอย่างเดียวไม่เพียงพอต่อการตัดสินใจ "
                f"จำเป็นต้องพิจารณาต้นทุนตลอดอายุการใช้งาน (LCCA) ประกอบด้วย "
                f"ซึ่งจะแสดงในหัวข้อ {sec_lcca}"
            )
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # แทรกรูป Cross-Section (ถ้าผู้ใช้ Generate ไว้ใน Tab Cross-Section)
    # ════════════════════════════════════════════════════════════════════════
    img_bytes = ss.get('cs_last_img_bytes')
    if img_bytes:
        next_h1(SCc, "รูปแบบหน้าตัดโครงสร้างชั้นทาง (Cross-Section)")
        body("รูปต่อไปนี้แสดงหน้าตัดโครงสร้างชั้นทางที่ได้จากการออกแบบ "
             "เพื่อใช้ประกอบการประเมินราคาและรายงาน:")
        img_stream = io.BytesIO(img_bytes)
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(img_stream, width=Cm(15.0))
        f_no    = ss.get('cs_last_fig_no', '3.9-1')
        p_type  = ss.get('cs_last_ptype',  'โครงสร้างชั้นทาง')
        p_cap   = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_rf(p_cap.add_run(f"รูปที่ {f_no}  รูปแบบหน้าตัด {p_type}"),
                size=14, bold=False)
        doc.add_paragraph()
    # ════════════════════════════════════════════════════════════════════════

    # ════════════════════════════════════════════════════════════════════════
    # ส่วนที่ 2 — งบประมาณบำรุงรักษาประจำปี (Routine Cost)
    # ════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    hdg(f"{sec_routine}  งบประมาณบำรุงรักษาประจำปี", size=18, uline=True, sb=14)
    body(
        "งบประมาณบำรุงรักษาประจำปีคำนวณโดยวิธีกำหนดค่าสัมประสิทธิ์ปรับแก้ (K) "
        "ตามแนวทางคู่มือการคิดค่าปริมาณงานและงานบำรุงปกติ กองบำรุง กรมทางหลวง (2538) "
        "ค่าสัมประสิทธิ์ K สะท้อนลักษณะเฉพาะของแต่ละสายทาง ได้แก่ "
        "ประเภทผิวทาง ปริมาณจราจร สภาพดินคันทาง อายุบริการ เขตทาง ภูมิประเทศ "
        "และสิ่งปลูกสร้างประกอบทาง แบ่งออกเป็นสองประเภทตามลักษณะผิวทาง "
        "คือ Ka สำหรับผิวแอสฟัลต์คอนกรีต และ Kc สำหรับผิวคอนกรีตซีเมนต์"
    )
    doc.add_paragraph()

    SCr = make_sc(sec_routine)

    # ── 2.1 พารามิเตอร์ N และ Km ─────────────────────────────────────────
    next_h1(SCr, "อัตราค่าบำรุงมาตรฐาน (N) และค่า Factor วัสดุ (Km)")
    body(
        f"อัตราค่าบำรุงมาตรฐาน (N) คือค่าใช้จ่ายพื้นฐานต่อกิโลเมตรต่อปีก่อนปรับแก้ "
        f"ค่า Factor วัสดุ (Km) ใช้ปรับตามราคาวัสดุปัจจุบัน "
        f"ในการวิเคราะห์ครั้งนี้ใช้ N = {ac_Na:,.0f} บาท/กม./ปี (AC) "
        f"และ {cc_Nc:,.0f} บาท/กม./ปี (Concrete) พร้อม Km ดังตาราง"
    )
    _add_tbl_w(doc, ["ประเภทผิวทาง","N มาตรฐาน (บาท/กม./ปี)","Km วัสดุ"], [
        ["ผิวแอสฟัลต์ (Ka)", f"{ac_Na:,.0f}", f"{ac_Km:.3f}"],
        ["ผิวคอนกรีต (Kc)",  f"{cc_Nc:,.0f}", f"{cc_Km:.3f}"],
    ], col_widths=[6, 5, 5])
    doc.add_paragraph()

    # ── 2.2 สูตรและค่า Factor ────────────────────────────────────────────
    next_h1(SCr, "สูตรการคำนวณและค่า Factor")
    body(
        "ค่าสัมประสิทธิ์ K คำนวณจากผลรวมของค่า Factor ที่สะท้อนลักษณะของสายทาง "
        "โดยแบ่งเป็น Factor X และ Z (เกี่ยวกับผิวทางและปริมาณจราจร) "
        "และ Factor Y (เกี่ยวกับเขตทาง ภูมิประเทศ และสิ่งปลูกสร้างประกอบทาง)"
    )
    for formula in [
        "Ka = 1 + 0.50 × (X1 + X2 + X3 + X4 + X5 + X6 + Y1 + Y2 + Y3 + Y4 + Y5 + Y6)",
        "Kc = 1 + 0.50 × (Z1 + Z2 + Z3 + Z4 + Y1 + Y2 + Y3 + Y4 + Y5 + Y6)",
        "งบประมาณ (บาท/ปี) = ระยะทาง (กม.) × K × Km × N",
    ]:
        p_f = doc.add_paragraph(); p_f.paragraph_format.left_indent = Cm(1.5)
        _set_rf(p_f.add_run(formula), size=14, bold=True)
    doc.add_paragraph()

    # ── 2.3 ผลการคำนวณ Ka ────────────────────────────────────────────────
    next_h1(SCr, "ผลการคำนวณค่าสัมประสิทธิ์และงบประมาณ")
    if ka_fixed:
        next_h2(SCr, "ผิวแอสฟัลต์ — ค่า Factor และ Ka")
        body(
            f"ค่า Factor ที่ใช้ในการคำนวณ Ka สำหรับผิวแอสฟัลต์คอนกรีตแสดงในตารางต่อไปนี้ "
            f"ค่า Ka เฉลี่ยตลอดระยะเวลาวิเคราะห์ = {ka_avg:.4f}"
        )
        fac_rows = [
            ["X1","ลักษณะผิวทางและพื้นทาง",     f"{ka_fixed.get('X1',0):.4f}"],
            ["X2","CBR ดินเดิม",                 f"{ka_fixed.get('X2',0):.4f}"],
            ["X3","ปริมาณจราจร AADT",            f"{ka_fixed.get('X3',0):.4f}"],
            ["X4","อายุบริการ (เฉลี่ยตลอดช่วง)", "(ดูตารางรายปี)"],
            ["X5","ความกว้างผิวทาง",              f"{ka_fixed.get('X5',0):.4f}"],
            ["X6","ภูมิประเทศ",                   f"{ka_fixed.get('X6',0):.4f}"],
            ["Y1","ความกว้างเขตทาง",              f"{ka_fixed.get('Y1',0):.4f}"],
            ["Y2","ไหล่ทางกว้างสุด 1 ข้าง",      f"{ka_fixed.get('Y2',0):.4f}"],
            ["Y3","จราจรสงเคราะห์",               f"{ka_fixed.get('Y3',0):.4f}"],
            ["Y4","ท่อระบายน้ำ",                  f"{ka_fixed.get('Y4',0):.4f}"],
            ["Y5","สะพาน",                        f"{ka_fixed.get('Y5',0):.4f}"],
            ["Y6","ทำความสะอาดและระบายน้ำ",       f"{ka_fixed.get('Y6',0):.4f}"],
            ["Ka เฉลี่ย","ค่าสัมประสิทธิ์เฉลี่ยตลอดช่วงวิเคราะห์", f"{ka_avg:.4f}"],
        ]
        _add_tbl_w(doc, ["Factor","คำอธิบาย","ค่าที่ใช้"], fac_rows, col_widths=[2,8,3])
        doc.add_paragraph()
    if ka_df is not None and len(ka_df) > 0:
        next_h2(SCr, "ตาราง Ka รายปี (X4 เปลี่ยนตามอายุ)")
        body("ค่า X4 เปลี่ยนแปลงตามอายุบริการของผิวทางในแต่ละปี ส่งผลให้ Ka แตกต่างกันในแต่ละปี ดังตาราง")
        ka_rows = [[str(int(r["ปี"])), str(int(r["อายุ (ปี)"])), f"{r['X4']:.2f}", f"{r['Ka']:.4f}"]
                   for _, r in ka_df.iterrows()]
        _add_tbl_w(doc, ["ปีที่","อายุ (ปี)","X4","Ka"], ka_rows, col_widths=[2.5,3,3,3])
        doc.add_paragraph()
    if kc_fac:
        next_h2(SCr, "ผิวคอนกรีต — ค่า Factor และ Kc")
        body(
            f"ค่า Factor ที่ใช้ในการคำนวณ Kc สำหรับผิวคอนกรีตซีเมนต์แสดงในตารางต่อไปนี้ "
            f"ค่า Kc = {kc_val:.4f}"
        )
        fac_rows_c = [
            ["Z1","ดัชนีสภาพผิวทาง",      f"{kc_fac.get('Z1',0):.4f}"],
            ["Z2","CBR ดินคันทาง",         f"{kc_fac.get('Z2',0):.4f}"],
            ["Z3","ปริมาณจราจร AADT",     f"{kc_fac.get('Z3',0):.4f}"],
            ["Z4","ความกว้างผิวทาง",       f"{kc_fac.get('Z4',0):.4f}"],
            ["Y1","ความกว้างเขตทาง",       f"{kc_fac.get('Y1',0):.4f}"],
            ["Y2","ไหล่ทางกว้างสุด 1 ข้าง",f"{kc_fac.get('Y2',0):.4f}"],
            ["Y3","จราจรสงเคราะห์",        f"{kc_fac.get('Y3',0):.4f}"],
            ["Y4","ท่อระบายน้ำ",           f"{kc_fac.get('Y4',0):.4f}"],
            ["Y5","สะพาน",                 f"{kc_fac.get('Y5',0):.4f}"],
            ["Y6","ทำความสะอาดและระบายน้ำ",f"{kc_fac.get('Y6',0):.4f}"],
            ["Kc","ค่าสัมประสิทธิ์คอนกรีต",f"{kc_val:.4f}"],
        ]
        _add_tbl_w(doc, ["Factor","คำอธิบาย","ค่าที่ใช้"], fac_rows_c, col_widths=[2,8,3])
        doc.add_paragraph()

    # ── 2.4 สรุปงบประมาณบำรุงปกติ ────────────────────────────────────────
    next_h1(SCr, "สรุปงบประมาณบำรุงรักษาประจำปี")
    body(
        f"งบประมาณบำรุงรักษาประจำปีสำหรับโครงการระยะทาง {length:.2f} กม. "
        f"คำนวณจากสูตร: งบประมาณ = ระยะทาง × K × Km × N แสดงดังตาราง"
    )
    mrows = []
    if ka_avg: mrows.append(["AC (ลาดยาง)", f"Ka = {ka_avg:.4f}",
        f"{r_ac_sqm:.4f}", f"{r_ac_km:,.2f}", f"{r_ac_km/1e6:.4f}"])
    if kc_val: mrows.append(["Concrete",    f"Kc = {kc_val:.4f}",
        f"{r_cc_sqm:.4f}", f"{r_cc_km:,.2f}", f"{r_cc_km/1e6:.4f}"])
    if mrows:
        _add_tbl_w(doc, ["ผิวทาง","K เฉลี่ย","บาท/ตร.ม./ปี","บาท/กม./ปี","ล้านบาท/กม./ปี"],
            mrows, col_widths=[3,3,3.5,4,4])
    doc.add_paragraph()
    if ka_avg and kc_val:
        body(
            f"งบประมาณบำรุงปกติผิวแอสฟัลต์สูงกว่าผิวคอนกรีตประมาณ "
            f"{(r_ac_sqm/r_cc_sqm - 1)*100:.1f}% "
            f"({r_ac_sqm:.4f} เทียบกับ {r_cc_sqm:.4f} บาท/ตร.ม./ปี) "
            f"ซึ่งค่านี้จะนำไปใช้เป็นต้นทุนบำรุงรักษาประจำปีในการวิเคราะห์ LCCA ต่อไป"
        )
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # ส่วนที่ 3 — LCCA
    # ════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    hdg(f"{sec_lcca}  วิเคราะห์ต้นทุนตลอดอายุการใช้งานผิวทาง (LCCA)", size=18, uline=True, sb=14)
    body(
        "การวิเคราะห์ต้นทุนตลอดอายุการใช้งาน (Life-Cycle Cost Analysis: LCCA) "
        "เป็นเครื่องมือทางเศรษฐศาสตร์วิศวกรรมสำหรับเปรียบเทียบทางเลือกการลงทุน "
        "โดยพิจารณาต้นทุนทุกประเภทตลอดอายุโครงการ ได้แก่ ต้นทุนก่อสร้างเริ่มต้น "
        "ต้นทุนบำรุงรักษาประจำปี ต้นทุนฟื้นฟูสภาพ และมูลค่าซากปลายโครงการ "
        "ต้นทุนในอนาคตทุกรายการถูกแปลงเป็นมูลค่าปัจจุบัน (Present Worth) "
        "โดยใช้อัตราคิดลด (Discount Rate) เพื่อให้สามารถเปรียบเทียบในฐานเดียวกันได้ "
        "วิธีนี้อ้างอิงตามมาตรฐาน FHWA-SA-98-079 และ AASHTO Guide for Design of Pavement Structures (1993)"
    )
    doc.add_paragraph()

    SCl = make_sc(sec_lcca)

    # ── 3.1 ทฤษฎีและสูตร ────────────────────────────────────────────────
    next_h1(SCl, "ทฤษฎีและสูตรที่ใช้ในการวิเคราะห์")
    body("สูตรมูลค่าปัจจุบัน (Present Worth) และต้นทุนเฉลี่ยรายปี (EAC) มีดังนี้")
    for line in [
        f"PW  = FV × (1 + i)^(−n)                              [i = {dr*100:.1f}%/ปี]",
        f"EAC = PW × [i × (1+i)^n] / [(1+i)^n − 1]    [n = {n} ปี]",
        "เมื่อ  PW = มูลค่าปัจจุบัน,  FV = ต้นทุนในอนาคต,  EAC = ต้นทุนเฉลี่ยรายปี",
    ]:
        p_l = doc.add_paragraph(); p_l.paragraph_format.left_indent = Cm(1.5)
        _set_rf(p_l.add_run(line), size=14, bold=(line.startswith("PW") or line.startswith("EAC")))
    doc.add_paragraph()

    # ── 3.2 พารามิเตอร์ ──────────────────────────────────────────────────
    next_h1(SCl, "พารามิเตอร์การวิเคราะห์")
    _add_tbl_w(doc, ["พารามิเตอร์","ค่าที่ใช้","หมายเหตุ"], [
        ["ระยะเวลาวิเคราะห์",  f"{n} ปี",          "ตามมาตรฐาน DOH"],
        ["อัตราคิดลด",         f"{dr*100:.1f}%/ปี", "อัตราคิดลดทางสังคม"],
        ["มูลค่าซากปลายโครงการ","รวม" if ss.get("lc_salvage",True) else "ไม่รวม", ""],
        ["พื้นที่คำนวณ",        f"{tw*1000:,.0f} ตร.ม./กม.", f"กว้าง {tw:.2f} ม."],
    ], col_widths=[5,4,7])
    doc.add_paragraph()

    # ── 3.3 ทางเลือกที่วิเคราะห์ ─────────────────────────────────────────
    next_h1(SCl, "ทางเลือกผิวทางที่วิเคราะห์")
    body(
        f"การวิเคราะห์ครั้งนี้พิจารณาทางเลือกทั้งหมด {len(enabled_alts)} ทางเลือก "
        "ได้แก่ผิวทางแอสฟัลต์คอนกรีต (AC) และผิวทางคอนกรีตซีเมนต์ประเภทต่าง ๆ "
        "โดยต้นทุนก่อสร้างนำมาจากการคำนวณในหัวข้อที่ผ่านมา "
        "ส่วนต้นทุนบำรุงรักษาใช้ค่าจากการคำนวณ Routine Cost"
    )
    if enabled_alts:
        alt_rows = [[a.name, a.pave_type, f"{a.construction_cost:,.2f}",
                     f"{a.salvage_pct:.0f}%"] for a in enabled_alts]
        _add_tbl_w(doc, ["ทางเลือก","ประเภท","ต้นทุนก่อสร้าง (บาท/ตร.ม.)","มูลค่าซาก"],
            alt_rows, col_widths=[5,3,5,3])
    doc.add_paragraph()

    # ── 3.4 ผลการวิเคราะห์ ───────────────────────────────────────────────
    next_h1(SCl, "ผลการวิเคราะห์ LCCA")
    if summary_df is not None and len(summary_df) > 0:
        body(
            "ผลการวิเคราะห์แสดงมูลค่าปัจจุบันสุทธิ (NPV) และต้นทุนเฉลี่ยรายปี (EAC) "
            f"ของแต่ละทางเลือกตลอดระยะเวลาวิเคราะห์ {n} ปี ที่อัตราคิดลด {dr*100:.1f}% ต่อปี "
            "ทางเลือกที่มี NPV และ EAC ต่ำที่สุดถือว่าคุ้มค่าทางเศรษฐศาสตร์มากที่สุด"
        )
        rows_out = []
        for _, row in summary_df.iterrows():
            rows_out.append([
                str(int(row['อันดับ'])), row['ทางเลือก'], row['ประเภทผิวทาง'],
                f"{row['ต้นทุนก่อสร้าง (บาท/ตร.ม.)']:,.2f}",
                f"{row['NPV (ล้านบาท/กม.)']:,.4f}",
                f"{row['EAC (ล้านบาท/กม./ปี)']:,.4f}",
                f"{row['EAC (บาท/ตร.ม./ปี)']:,.2f}",
            ])
        _add_tbl_w(doc,
            ["อันดับ","ทางเลือก","ประเภท","ก่อสร้าง\n(บ./ตร.ม.)","NPV\n(ล้าน/กม.)","EAC\n(ล้าน/กม./ปี)","EAC\n(บ./ตร.ม./ปี)"],
            rows_out, col_widths=[1.5, 4.5, 2.5, 3, 3, 3.5, 3.5])
        doc.add_paragraph()

    # ── 3.5 กระแสเงินสดรายทางเลือก ──────────────────────────────────────
    next_h1(SCl, "กระแสเงินสด (Cash Flow) รายทางเลือก")
    body(
        "ตารางกระแสเงินสดต่อไปนี้แสดงต้นทุนแต่ละรายการในแต่ละปี "
        "พร้อมมูลค่าปัจจุบัน (PW) โดยคิดลดกลับมา ณ ปีที่ 0 "
        "ด้วย Discount Factor = (1 + i)^(−n)"
    )
    for alt_name, cf in cf_dict.items():
        _add_hdg_w(doc, f"ทางเลือก: {alt_name}", level=2)
        cf_rows = []
        for _, row in cf.iterrows():
            cf_rows.append([
                str(int(row['ปี'])), row['กิจกรรม'], row['ประเภท'],
                f"{row['ต้นทุน/หน่วย']:,.2f}",
                f"{row['ต้นทุนตามปี']:,.0f}",
                f"{row['PW_factor']:.4f}",
                f"{row['มูลค่าปัจจุบัน']:,.0f}",
            ])
        _add_tbl_w(doc,
            ["ปีที่","กิจกรรม","ประเภท","ต้นทุน/หน่วย","ต้นทุนตามปี","PW Factor","มูลค่าปัจจุบัน"],
            cf_rows, col_widths=[1.5, 4, 2.5, 2.5, 3, 2.5, 3])
        doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # สรุปและข้อเสนอแนะ (ท้ายรายงาน)
    # ════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    hdg("สรุปและข้อเสนอแนะ", size=18, uline=True, sb=14)

    if summary_df is not None and len(summary_df) > 0:
        best  = summary_df.iloc[0]
        worst = summary_df.iloc[-1]

        # ย่อหน้า 1 — สรุปราคาก่อสร้าง
        if summary_cost:
            sorted_c  = sorted(summary_cost, key=lambda x: x['cost_sqm'])
            cheapest  = sorted_c[0]
            body(
                f"การวิเคราะห์ราคาก่อสร้างโครงสร้างชั้นทางสำหรับ{proj_nm} "
                f"ความกว้างรวม {tw:.2f} เมตร ระยะทาง {length:.2f} กิโลเมตร "
                f"พบว่าโครงสร้างประเภท {cheapest['name']} มีราคาก่อสร้างต่ำที่สุด "
                f"ที่ {cheapest['cost_sqm']:,.2f} บาท/ตารางเมตร "
                f"หรือ {cheapest['cost_per_km']:.3f} ล้านบาท/กิโลเมตร"
            )

        # ย่อหน้า 2 — สรุป Routine Cost
        if ka_avg or kc_val:
            body(
                f"ในด้านต้นทุนบำรุงรักษาประจำปี ผิวทางแอสฟัลต์คอนกรีตมีค่าสัมประสิทธิ์ "
                f"Ka = {ka_avg:.4f} คิดเป็น {r_ac_sqm:.4f} บาท/ตร.ม./ปี "
                f"ส่วนผิวทางคอนกรีตซีเมนต์มี Kc = {kc_val:.4f} "
                f"คิดเป็น {r_cc_sqm:.4f} บาท/ตร.ม./ปี"
            )

        # ย่อหน้า 3 — สรุป LCCA หลัก
        body(
            f"เมื่อวิเคราะห์ต้นทุนตลอดอายุการใช้งาน {n} ปี ที่อัตราคิดลด {dr*100:.1f}% ต่อปี "
            f"พบว่าทางเลือกที่มีต้นทุนรวมตลอดอายุการใช้งานต่ำที่สุดคือ "
            f"{best['ทางเลือก']} ({best['ประเภทผิวทาง']}) "
            f"มีมูลค่าปัจจุบันสุทธิ NPV = {best['NPV (ล้านบาท/กม.)']:,.4f} ล้านบาท/กิโลเมตร "
            f"และต้นทุนเฉลี่ยรายปี EAC = {best['EAC (ล้านบาท/กม./ปี)']:,.4f} ล้านบาท/กิโลเมตร/ปี "
            f"({best['EAC (บาท/ตร.ม./ปี)']:,.2f} บาท/ตารางเมตร/ปี)"
        )

        # ย่อหน้า 4 — เปรียบเทียบกับอันดับสุดท้าย
        if len(summary_df) > 1:
            npv_diff = worst['NPV (ล้านบาท/กม.)'] - best['NPV (ล้านบาท/กม.)']
            body(
                f"เมื่อเปรียบเทียบกับทางเลือกที่มีต้นทุนสูงที่สุด "
                f"คือ {worst['ทางเลือก']} (NPV = {worst['NPV (ล้านบาท/กม.)']:,.4f} ล้านบาท/กม.) "
                f"ทางเลือกที่ดีที่สุดมีต้นทุนน้อยกว่า "
                f"{npv_diff:,.4f} ล้านบาท/กิโลเมตร ตลอดอายุโครงการ"
            )

        # ย่อหน้า 5 — ข้อเสนอแนะ
        body(
            f"ดังนั้น จึงมีข้อเสนอแนะให้พิจารณาเลือกใช้ {best['ทางเลือก']} "
            f"เป็นโครงสร้างชั้นทางสำหรับ{proj_nm} "
            f"เนื่องจากมีต้นทุนรวมตลอดอายุการใช้งานต่ำที่สุดในบรรดาทางเลือกทั้งหมดที่วิเคราะห์ "
            f"ทั้งนี้ผู้ออกแบบควรพิจารณาปัจจัยอื่นประกอบด้วย ได้แก่ "
            f"ความสามารถในการก่อสร้าง ความพร้อมของวัสดุในพื้นที่ และความต้องการของลูกค้า"
        )

    # อ้างอิง
    hdg("เอกสารอ้างอิง", size=16, uline=True, sb=12)
    for ref in [
        "FHWA-SA-98-079: Life-Cycle Cost Analysis in Pavement Design, USDOT (1998)",
        "AASHTO Guide for Design of Pavement Structures, American Association of State Highway and Transportation Officials (1993)",
        "NCHRP Report 703: Guide for Pavement-Type Selection (2011)",
        "คู่มือการคิดค่าปริมาณงานและงานบำรุงปกติ กองบำรุง กรมทางหลวง (มกราคม 2538)",
        "มาตรฐานงานทาง กรมทางหลวง กระทรวงคมนาคม",
    ]:
        p_r = doc.add_paragraph(); p_r.paragraph_format.left_indent = Cm(1)
        _set_rf(p_r.add_run(f"- {ref}"), size=14)

    doc.add_paragraph()
    p_f = doc.add_paragraph(); p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_rf(p_f.add_run(
        f"จัดทำโดย รศ.ดร.อิทธิพล มีผล — ภาควิชาครุศาสตร์โยธา มจพ. | "
        f"{datetime.now().strftime('%d/%m/%Y %H:%M')}"
    ), size=12)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf

    # คำนวณหมายเลขหัวข้อถัดไปจาก base_sec
# ══════════════════════════════════════════════════════════════════════════════
# ─── SESSION STATE INIT ───────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

def init_state():
    defaults = {
        # ── Shared project info ──
        'sb_project_name':   'โครงการก่อสร้างทางหลวง',
        'sb_road_length':    1.0,
        'sb_design_life':    20,
        'sb_lane_width':     3.50,
        'sb_lanes_per_dir':  2,
        'sb_shoulder_l':     2.50,
        'sb_shoulder_r':     1.50,
        'json_version':      0,
        # ── Cost Structure results ──
        'cs_all_results':    {},
        # ── Cost sync state ──
        'lc_synced_costs':   {},   # ราคาที่ sync มาแล้ว {AC:..., JPCP:..., ...}
        'lc_sync_ts':        '',   # timestamp ที่ sync ล่าสุด
        'lc_sync_source':    {},   # snapshot ของ cs_all_results ตอนที่ sync
        'lc_manual_override':False,# True เมื่อ user กด "แก้ไขเอง"
        # ── LCCA state ──
        'lc_tab_routine_done': False,
        'lc_tab_alts_done':    False,
        'lc_tab_result_done':  False,
        'lc_dirty_routine':    False,
        # Routine cost inputs
        'lc_cbr':          3.0,
        'lc_y_row':       40.0,
        'lc_y_shoulder':   1.75,
        'lc_y_terrain':    TERRAIN_KEYS[0],
        'lc_y_bridge':     0.0,
        'lc_ac_x1_key':    list(X1_MAP.keys())[0],
        'lc_ac_x3_key':    list(X3_OPTIONS.keys())[0],
        'lc_ac_x4_age':    0,
        'lc_ac_x5_width':  7.0,
        'lc_ac_x6_terrain':TERRAIN_KEYS[0],
        'lc_ac_Na':        35000.0,
        'lc_ac_Km':        1.0,
        'lc_cc_z1_idx':    1,
        'lc_cc_z3_key':    list(Z3_OPTIONS.keys())[0],
        'lc_cc_z4_width':  7.0,
        'lc_cc_Nc':        35000.0,
        'lc_cc_Km':        1.0,
        # Routine results
        'lc_ka_avg':        None,
        'lc_kc_val':        None,
        'lc_routine_ac_sqm':None,
        'lc_routine_cc_sqm':None,
        'lc_routine_ac_km': None,
        'lc_routine_cc_km': None,
        # LCCA params
        'lc_n':            20,
        'lc_dr':           0.06,
        'lc_salvage':      True,
        'lc_alternatives': None,
        'lc_base_sec':     '4.7',   # หัวข้อเริ่มต้นรายงานรวม
    }
    for k,v in defaults.items():
        if k not in st.session_state:
            st.session_state[k]=v

init_state()
ss=st.session_state

# ══════════════════════════════════════════════════════════════════════════════
# ─── SIDEBAR ─────────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

with st.sidebar:
    st.markdown("## 🛣️ PaveCost + LCCA")
    st.caption("รศ.ดร.อิทธิพล มีผล | KMUTNB")
    st.divider()

    # ── Upload Price Library Excel ──
    with st.expander("📤 Upload Price Library (Excel)",expanded=False):
        if OPENPYXL_OK:
            upxl=st.file_uploader("Price Library .xlsx",type=['xlsx'],key="sb_price_xl")
            if upxl is not None:
                fh=hashlib.md5(upxl.read()).hexdigest(); upxl.seek(0)
                if ss.get('sb_price_xl_hash')!=fh:
                    try:
                        lib=load_excel_price_library(upxl)
                        st.session_state['price_library']=lib
                        st.session_state['uploaded_price_library']=lib
                        ss['sb_price_xl_hash']=fh
                        for k in list(ss.keys()):
                            if any(p in k for p in ['tab2_bp_','tab2_conc_','tab2_ton_']): del ss[k]
                        st.success("✅ โหลด Price Library สำเร็จ")
                    except Exception as e:
                        st.error(f"❌ {e}")
                else:
                    st.success("✅ โหลดแล้ว")
        else:
            st.info("ติดตั้ง openpyxl เพื่อใช้งาน")

    # ── Load / Save JSON ──
    with st.expander("📂 โหลดโครงการ (JSON)",expanded=False):
        upj=st.file_uploader("Upload JSON",type=['json'],key="sb_upload_json")
        if upj is not None:
            try:
                fb=upj.read(); fh=hashlib.md5(fb).hexdigest()
                loaded=json.loads(fb.decode('utf-8'))
                if 'project_info' in loaded:
                    st.info(f"📌 {loaded['project_info'].get('name','-')}")
                    st.caption(f"บันทึกเมื่อ: {loaded.get('saved_at','-')}")
                if st.button("📥 นำเข้าข้อมูล",key="sb_import_json"):
                    if ss.get('sb_json_hash')!=fh:
                        ss['sb_loaded_project']=loaded; ss['sb_json_hash']=fh
                        new_v=ss.get('json_version',0)+1; ss['json_version']=new_v
                        keys_clear=[k for k in ss if any(p in k for p in ['_base_rows_','_joint_init_','_surf_'])]
                        for k in keys_clear: del ss[k]
                    st.rerun()
            except Exception as e:
                st.error(f"❌ {e}")

    st.divider()

    # ── Project Info (shared) ──
    lp=ss.get('sb_loaded_project',{}); li=lp.get('project_info',{})
    v_sb=ss.get('json_version',0)

    project_name=st.text_input("ชื่อโครงการ",value=li.get('name',ss['sb_project_name']),key=f"sb_pname_v{v_sb}")
    road_length =st.number_input("ความยาวถนน (กม.)",value=float(li.get('length',ss['sb_road_length'])),
        min_value=0.1,step=0.1,key=f"sb_length_v{v_sb}")
    design_life =st.number_input("อายุออกแบบ (ปี)",value=int(li.get('design_life',ss['sb_design_life'])),
        min_value=1,max_value=50,step=1,key=f"sb_dlife_v{v_sb}")
    st.divider()
    st.markdown("**📐 ขนาดถนน**")
    lane_width   =st.number_input("กว้างช่องจราจร (ม.)",value=float(li.get('lane_width',ss['sb_lane_width'])),
        min_value=2.5,max_value=4.5,step=0.25,key=f"sb_lw_v{v_sb}")
    lpo=[2,3,4]; lp_def=li.get('num_lanes',ss['sb_lanes_per_dir']*2)//2
    lp_idx=lpo.index(lp_def) if lp_def in lpo else 0
    lanes_per_dir=st.selectbox("ช่องจราจร/ทิศทาง",options=lpo,index=lp_idx,key=f"sb_lpd_v{v_sb}")
    num_lanes    =lanes_per_dir*2
    shoulder_l   =st.number_input("ไหล่ทางซ้าย (ม.)",value=float(li.get('shoulder_left',ss['sb_shoulder_l'])),
        min_value=0.0,max_value=4.0,step=0.25,key=f"sb_sl_v{v_sb}")
    shoulder_r   =st.number_input("ไหล่ทางขวา (ม.)",value=float(li.get('shoulder_right',ss['sb_shoulder_r'])),
        min_value=0.0,max_value=4.0,step=0.25,key=f"sb_sr_v{v_sb}")

    road_surface_w=lane_width*num_lanes
    total_shoulders=(shoulder_l+shoulder_r)*2
    total_width   =road_surface_w+total_shoulders
    area_per_km   =total_width*1000

    st.info(f"📏 ช่องรวม: **{num_lanes}** ช่อง\n"
            f"📏 ผิวจราจร: **{road_surface_w:.2f}** ม.\n"
            f"📏 ไหล่ทาง: **{total_shoulders:.2f}** ม.\n"
            f"📏 กว้างรวม: **{total_width:.2f}** ม.")

    # update shared session
    ss['sb_project_name']=project_name; ss['sb_road_length']=road_length
    ss['sb_design_life']=design_life; ss['sb_lane_width']=lane_width
    ss['sb_lanes_per_dir']=lanes_per_dir; ss['sb_shoulder_l']=shoulder_l; ss['sb_shoulder_r']=shoulder_r

project_info={
    'name':project_name,'length':road_length,'design_life':design_life,
    'lane_width':lane_width,'num_lanes':num_lanes,
    'shoulder_left':shoulder_l,'shoulder_right':shoulder_r,'total_width':total_width,
}
v=ss.get('json_version',0)

# ══════════════════════════════════════════════════════════════════════════════
# ─── MAIN HEADER ─────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

st.markdown(f"""
<div class="main-header">
  <h1>🛣️ ระบบวิเคราะห์โครงสร้างชั้นทาง + LCCA</h1>
  <p><b>{project_name}</b> &nbsp;|&nbsp; ระยะทาง {road_length:.2f} กม. &nbsp;|&nbsp; กว้าง {total_width:.2f} ม. &nbsp;|&nbsp; อายุออกแบบ {design_life} ปี</p>
</div>
""",unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# ─── MAIN TABS ────────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

# ── compute status badges for main tabs ──────────────────────────────────────
_cs_ar       = ss.get('cs_all_results', {})
_cs_has_cost = any(_cs_ar.get(pt, {}).get('cost_sqm', 0) > 0 for pt in ['AC','JPCP','JRCP','CRCP'])
_cs_badge    = " ✅" if _cs_has_cost else ""

_lc_done     = ss.get('lc_tab_result_done', False)
_lc_routine  = ss.get('lc_tab_routine_done', False)
_lc_badge    = " ✅" if _lc_done else (" 🔧" if _lc_routine else "")

tab_cs, tab_lc = st.tabs([
    f"🏗️ โครงสร้างชั้นทาง{_cs_badge}",
    f"📊 LCCA{_lc_badge}",
])

# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║  TAB A — COST STRUCTURE                                                     ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

with tab_cs:
    sub_layer, sub_price, sub_summary, sub_cross = st.tabs([
        "🏗️ กำหนดหน้าตัด",
        "💰 ราคาวัสดุ",
        "📊 สรุปต้นทุน & รายงาน",
        "🖼️ Cross-Section",
    ])

    all_results: dict = {}

    # ── Sub: Layer Editor ────────────────────────────────────────────────────
    with sub_layer:
        st.markdown(f"**โครงการ:** {project_name} &nbsp;|&nbsp; **{road_length:.2f} กม.** &nbsp;|&nbsp; **{total_width:.2f} ม.**")
        sub_ac,sub_jpcp,sub_jrcp,sub_crcp=st.tabs(["🛣️ AC","🏗️ JPCP","🏗️ JRCP","🏗️ CRCP"])
        for sub_tab,ptype in [(sub_ac,'AC'),(sub_jpcp,'JPCP'),(sub_jrcp,'JRCP'),(sub_crcp,'CRCP')]:
            with sub_tab:
                kp=ptype.lower()
                layers=render_layer_editor(ptype,kp,total_width,road_length,v=v)
                joints=[]; include_joints=True
                if ptype in ('JPCP','JRCP','CRCP'):
                    joints,include_joints=render_joint_editor(ptype,kp,area_per_km,road_length,v=v)
                layer_cost,layer_details=calculate_layer_cost(layers,road_length)
                joint_cost,joint_details=calculate_joint_cost(joints,road_length,include_joints)
                total_cost=layer_cost+joint_cost
                total_area_p=area_per_km*road_length
                cost_sqm=total_cost/total_area_p if total_area_p>0 else 0
                cost_per_km_m=total_cost/road_length/1e6 if road_length>0 else 0
                st.markdown("---")
                mc1,mc2,mc3=st.columns(3)
                with mc1: st.markdown(f'<div class="mc"><div class="lbl">💰 ราคา/ตร.ม.</div><div class="val">{cost_sqm:,.2f}</div><div class="sub">บาท</div></div>',unsafe_allow_html=True)
                with mc2: st.markdown(f'<div class="mc"><div class="lbl">📏 ราคา/กม.</div><div class="val">{cost_per_km_m:,.3f}</div><div class="sub">ล้านบาท</div></div>',unsafe_allow_html=True)
                with mc3: st.markdown(f'<div class="mc"><div class="lbl">⏱️ อายุออกแบบ</div><div class="val">{design_life}</div><div class="sub">ปี</div></div>',unsafe_allow_html=True)
                all_results[ptype]={
                    'name':ptype,'name_detail':f"{ptype} — {project_name}",
                    'layers':layers,'joints':joints,'details':layer_details+joint_details,
                    'cost_total':total_cost,'cost_sqm':cost_sqm,'cost_per_km':cost_per_km_m,
                    'include_joints':include_joints,
                }
        ss['cs_all_results']=all_results

    # ── Sub: Price Library ───────────────────────────────────────────────────
    with sub_price:
        st.header("💰 ตารางราคาวัสดุ")
        st.info("💡 แก้ราคาได้โดยตรง หรือ Upload Excel ใน Sidebar")
        lib=get_price_library()
        col_dl,_=st.columns([1,2])
        with col_dl:
            if OPENPYXL_OK:
                tpl=generate_excel_template()
                st.download_button("⬇️ Download Template Excel",data=tpl,
                    file_name="price_library_template.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        # AC: บาท/ตัน
        with st.expander("🧮 คำนวณราคา AC จากราคาต่อตัน",expanded=True):
            st.caption("ราคา (บาท/ตร.ม.) = ราคา (บาท/ตัน) × density × ความหนา (m)")
            _dc1,_dc2=st.columns([1,3])
            with _dc1:
                _density_key="tab2_density"
                if _density_key not in ss: ss[_density_key]=DEFAULT_AC_DENSITY
                st.number_input("Density (ตัน/ลบ.ม.)",min_value=2.0,max_value=2.6,step=0.05,format="%.2f",key=_density_key)
            _density=float(ss[_density_key])
            _ac_order=['PMA Wearing Course','AC Wearing Course','AC Binder Course','AC Base Course']
            _saved_ton=ss.get('ac_ton_prices',DEFAULT_AC_TON_PRICES)
            _pv1,_pv2=st.columns([3,1])
            with _pv2:
                _prev_thick_key="tab2_preview_thick"
                if _prev_thick_key not in ss: ss[_prev_thick_key]=5.0
                st.number_input("ดูตัวอย่างที่หนา (cm)",min_value=2.5,max_value=10.0,step=0.5,format="%.1f",key=_prev_thick_key)
            _prev_thick=float(ss[_prev_thick_key])
            _gh=st.columns([3,1.5,1.5])
            _gh[0].markdown("<span style='color:#6b7a8d;font-size:.82rem;font-weight:600'>วัสดุ</span>",unsafe_allow_html=True)
            _gh[1].markdown("<span style='color:#6b7a8d;font-size:.82rem;font-weight:600'>บาท/ตัน</span>",unsafe_allow_html=True)
            _gh[2].markdown(f"<span style='color:#6b7a8d;font-size:.82rem;font-weight:600'>บาท/ตร.ม. ({_prev_thick:.1f}cm)</span>",unsafe_allow_html=True)
            _ton_prices={}
            for mat in _ac_order:
                _tkey=f"tab2_ton_{mat.replace(' ','_')}"
                if _tkey not in ss: ss[_tkey]=float(_saved_ton.get(mat,DEFAULT_AC_TON_PRICES.get(mat,0)))
                _gr=st.columns([3,1.5,1.5])
                with _gr[0]: st.markdown(f"<div style='padding:8px 0;font-weight:500'>{mat}</div>",unsafe_allow_html=True)
                with _gr[1]: st.number_input(f"ton_{mat}",min_value=0.0,step=50.0,format="%.0f",key=_tkey,label_visibility="collapsed")
                ton_p=float(ss[_tkey]); _ton_prices[mat]=ton_p
                preview=round(ton_p*_density*_prev_thick/100,2) if ton_p>0 else 0.0
                with _gr[2]:
                    color='#0f2942' if preview>0 else '#94a3b8'
                    st.markdown(f"<div style='padding:8px 0;font-weight:600;color:{color}'>{'—' if preview==0 else f'{preview:,.2f}'}</div>",unsafe_allow_html=True)
            if st.button("🔄 คำนวณและอัพเดท Price Table AC",type="primary",key="tab2_calc_ton"):
                _new_ac=_calc_ac_prices_from_ton(_ton_prices,_density)
                _cur_lib=get_price_library(); _cur_lib['ac_prices']=_new_ac
                ss['price_library']=_cur_lib; ss['ac_ton_prices']=dict(_ton_prices)
                st.success(f"✅ อัพเดทราคา AC สำเร็จ (density={_density:.2f})"); st.rerun()

        # Concrete: บาท/ลบ.ม.
        st.subheader("🏗️ ราคาคอนกรีต — บาท/ลบ.ม.")
        _conc_thk=[20,25,28,30,32,35]; _conc_order=['JPCP','JRCP','CRCP']
        if 'concrete_cum_prices' not in ss: ss['concrete_cum_prices']=dict(DEFAULT_CONCRETE_CUM_PRICES)
        _conc_cum=ss['concrete_cum_prices']
        _cph=st.columns([1.2,1.2]+[1]*len(_conc_thk))
        _cph[0].markdown("<span style='color:#6b7a8d;font-size:.82rem;font-weight:600'>ประเภท</span>",unsafe_allow_html=True)
        _cph[1].markdown("<span style='color:#6b7a8d;font-size:.82rem;font-weight:600'>บาท/ลบ.ม.</span>",unsafe_allow_html=True)
        for j,t in enumerate(_conc_thk): _cph[2+j].markdown(f"<span style='color:#6b7a8d;font-size:.82rem;font-weight:600'>{t}cm</span>",unsafe_allow_html=True)
        cp_edited_rows=[]
        for ct in _conc_order:
            _cum_key=f"tab2_conc_cum_{ct}"
            if _cum_key not in ss: ss[_cum_key]=float(_conc_cum.get(ct,DEFAULT_CONCRETE_CUM_PRICES.get(ct,0)))
            _cpr=st.columns([1.2,1.2]+[1]*len(_conc_thk))
            with _cpr[0]: st.markdown(f"<div style='padding:8px 0;font-weight:600'>{ct}</div>",unsafe_allow_html=True)
            with _cpr[1]: st.number_input(f"cum_{ct}",min_value=0.0,step=50.0,format="%.0f",key=_cum_key,label_visibility="collapsed")
            cum_val=float(ss[_cum_key])
            for j,t in enumerate(_conc_thk):
                sqm_val=round(cum_val*t/100,0)
                _cpr[2+j].markdown(f"<div style='padding:8px 4px;font-size:.9rem;color:#0f2942'>{sqm_val:,.0f}</div>",unsafe_allow_html=True)
            cp_edited_rows.append({'ประเภท':ct,'บาท/ลบ.ม.':cum_val})

        # Base materials
        _bp=lib['base_prices']
        st.subheader("🪨 ราคาวัสดุพื้นทาง / รองพื้นทาง (บาท/ลบ.ม.)")
        _hcum=st.columns([3,1.5])
        _hcum[0].markdown("<span style='color:#6b7a8d;font-size:.82rem;font-weight:600'>วัสดุ</span>",unsafe_allow_html=True)
        _hcum[1].markdown("<div style='color:#6b7a8d;font-size:.82rem;font-weight:600;text-align:right'>ราคา (บาท/ลบ.ม.)</div>",unsafe_allow_html=True)
        _cum_keys=[k for k in _bp if k not in SQMKEYS]
        for _mat in _cum_keys:
            _wk=f"tab2_bp_cum_{_mat.replace(' ','_').replace('(','').replace(')','')}"
            if _wk not in ss: ss[_wk]=float(_bp.get(_mat,0))
            _rc=st.columns([3,1.5])
            with _rc[0]: st.markdown(f"<div style='padding:6px 0'>{_mat}</div>",unsafe_allow_html=True)
            with _rc[1]: st.number_input(f"ราคา {_mat}",min_value=0.0,step=10.0,format="%.2f",key=_wk,label_visibility="collapsed")

        st.subheader("🧴 ราคาวัสดุผิว / อุปกรณ์ (บาท/ตร.ม.)")
        _sqm_keys=[k for k in _bp if k in SQMKEYS]
        for _mat in _sqm_keys:
            _wk=f"tab2_bp_sqm_{_mat.replace(' ','_')}"
            if _wk not in ss: ss[_wk]=float(_bp.get(_mat,0))
            _rs=st.columns([3,1.5])
            with _rs[0]: st.markdown(f"<div style='padding:6px 0'>{_mat}</div>",unsafe_allow_html=True)
            with _rs[1]: st.number_input(f"ราคา {_mat}",min_value=0.0,step=1.0,format="%.2f",key=_wk,label_visibility="collapsed")

        if st.button("💾 บันทึกราคาที่แก้ไขลง Library",type="primary",key="cs_save_price"):
            _den=float(ss.get('tab2_density',DEFAULT_AC_DENSITY))
            new_ac=_calc_ac_prices_from_ton({mat:float(ss.get(f"tab2_ton_{mat.replace(' ','_')}",DEFAULT_AC_TON_PRICES.get(mat,0))) for mat in ['PMA Wearing Course','AC Wearing Course','AC Binder Course','AC Base Course']},_den)
            _new_conc_cum={ct:float(ss.get(f"tab2_conc_cum_{ct}",DEFAULT_CONCRETE_CUM_PRICES.get(ct,0))) for ct in _conc_order}
            ss['concrete_cum_prices']=_new_conc_cum; new_cp=_calc_concrete_prices(_new_conc_cum)
            new_bp={**{m:float(ss.get(f"tab2_bp_cum_{m.replace(' ','_').replace('(','').replace(')','')}", _bp.get(m,0))) for m in _cum_keys},
                    **{m:float(ss.get(f"tab2_bp_sqm_{m.replace(' ','_')}",_bp.get(m,0))) for m in _sqm_keys}}
            ss['price_library']={'ac_prices':new_ac,'concrete_prices':new_cp,'base_prices':new_bp}
            st.success("✅ อัพเดท Price Library สำเร็จ"); st.rerun()

    # ── Sub: Cost Summary & Report ───────────────────────────────────────────
    with sub_summary:
        st.header("📊 สรุปต้นทุนและรายงาน")
        ar=ss.get('cs_all_results',all_results)
        if not ar:
            st.info("กรุณากำหนดหน้าตัดใน Sub-tab 🏗️ กำหนดหน้าตัด ก่อน")
        else:
            ptypes_list=list(ar.keys())
            cols_m=st.columns(len(ptypes_list))
            for i,pt in enumerate(ptypes_list):
                r=ar[pt]
                with cols_m[i]:
                    st.markdown(f'<div class="mc"><div class="lbl">{pt}</div>'
                        f'<div class="val">{r["cost_sqm"]:,.0f}<span style="font-size:.85rem;color:#8a95a3"> บาท/ตร.ม.</span></div>'
                        f'<div class="sub">{r["cost_per_km"]:.3f} ล้านบาท/กม.</div></div>',unsafe_allow_html=True)

            # Auto-fill indicator สำหรับ LCCA
            st.markdown("---")
            autofill_vals={pt:ar[pt]['cost_sqm'] for pt in ptypes_list}
            has_cost=any(v>0 for v in autofill_vals.values())
            if has_cost:
                st.markdown("**🔗 ส่งต่อไป LCCA อัตโนมัติ**")
                af_cols=st.columns(len(ptypes_list))
                for i,pt in enumerate(ptypes_list):
                    with af_cols[i]:
                        v_af=autofill_vals[pt]
                        color="green" if v_af>0 else "#999"
                        st.markdown(f"<span class='autofill' style='color:{color}'>{pt}: {v_af:,.2f} บ./ตร.ม.</span>",unsafe_allow_html=True)

            st.markdown("---")
            if PLOTLY_OK:
                fig=go.Figure()
                fig.add_trace(go.Bar(
                    x=[ar[pt]['name'] for pt in ptypes_list],
                    y=[ar[pt]['cost_sqm'] for pt in ptypes_list],
                    marker_color=['#1a4a7a','#0d7377','#14a085','#52b788'],
                    text=[f"{ar[pt]['cost_sqm']:,.0f}" for pt in ptypes_list],
                    textposition='outside',
                ))
                fig.update_layout(title='เปรียบเทียบราคา/ตร.ม. ตามประเภทโครงสร้าง',
                    yaxis_title='บาท/ตร.ม.',plot_bgcolor='white',paper_bgcolor='white',height=370)
                st.plotly_chart(fig,use_container_width=True)

            st.subheader("📋 รายละเอียดต้นทุนแต่ละประเภท")
            for pt in ptypes_list:
                r=ar[pt]
                with st.expander(f"🔍 {pt} — {r['cost_sqm']:,.2f} บาท/ตร.ม.",expanded=False):
                    if r['details']:
                        st.dataframe(
                            pd.DataFrame(r['details'])[['รายการ','ปริมาณ','หน่วย','ราคา/หน่วย (แสดง)','หน่วยราคา','มูลค่า (บาท)']],
                            use_container_width=True,hide_index=True)

            st.markdown("---"); st.subheader("💾 บันทึกและส่งออก")
            col_s1, col_s2 = st.columns(2)
            with col_s1:
                construction_out={pt:{'layers':r.get('layers',[]),'joints':r.get('joints',[]),'include_joints':r.get('include_joints',True)} for pt,r in ar.items()}
                save_data={'saved_at':datetime.now().strftime('%Y-%m-%d %H:%M:%S'),'project_info':project_info,'construction':construction_out}
                json_bytes=json.dumps(save_data,ensure_ascii=False,indent=2,default=str).encode('utf-8')
                fname_json=f"{project_name.replace(' ','_')}_{datetime.now().strftime('%Y%m%d')}.json"
                st.download_button("📥 บันทึก JSON",data=json_bytes,file_name=fname_json,
                    mime="application/json",use_container_width=True)
            with col_s2:
                st.markdown(
                    '<div class="info-band" style="margin-top:4px">📄 สร้าง <b>Word Report รวม</b> (Cost Structure + Routine Cost + LCCA) '
                    'ได้ที่ Tab <b>📊 LCCA → 📄 Word Report</b> — กดปุ่มเดียวได้ไฟล์ครบ</div>',
                    unsafe_allow_html=True)

# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║  TAB A — CROSS-SECTION DRAWING                                              ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

with sub_cross:
    render_cross_section_tab(ss=ss, project_name=project_name)

# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║  TAB B — LCCA                                                               ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

with tab_lc:
    # Progress bar
    t1="prog-done" if ss["lc_tab_routine_done"] else "prog-idle"
    t2="prog-done" if ss["lc_tab_alts_done"] else ("prog-warn" if ss.get("lc_dirty_routine") else "prog-idle")
    t3="prog-done" if ss["lc_tab_result_done"] else "prog-idle"
    t4="prog-done" if ss["lc_tab_result_done"] else "prog-idle"
    i1="✅" if ss["lc_tab_routine_done"] else "⭕"
    i2="✅" if ss["lc_tab_alts_done"] else ("⚠️" if ss.get("lc_dirty_routine") else "⭕")
    i3="✅" if ss["lc_tab_result_done"] else "⭕"
    i4="📄" if ss["lc_tab_result_done"] else "🔒"
    st.markdown(f"""
    <div class="prog-wrap">
      <div class="prog-step {t1}">{i1} Routine Cost<br><small>Ka / Kc</small></div>
      <div class="prog-arr">›</div>
      <div class="prog-step {t2}">{i2} ทางเลือก<br><small>กำหนด Alt.</small></div>
      <div class="prog-arr">›</div>
      <div class="prog-step {t3}">{i3} วิเคราะห์<br><small>NPV / EAC</small></div>
      <div class="prog-arr">›</div>
      <div class="prog-step {t4}">{i4} Word Report<br><small>รายงาน</small></div>
    </div>""",unsafe_allow_html=True)
    st.divider()

    sub_routine,sub_alts,sub_result,sub_word=st.tabs([
        "🔧 Routine Cost","🏗️ กำหนดทางเลือก","📊 ผลการวิเคราะห์","📄 Word Report"])

    # ── Sub: Routine Cost (Ka / Kc) ──────────────────────────────────────────
    with sub_routine:
        st.markdown('<div class="sh">🔧 Routine Cost Calculator — Ka (AC) & Kc (Concrete)</div>',unsafe_allow_html=True)

        # ตัวแปรร่วม Y1–Y6
        st.markdown('<div class="sh-org">📌 ตัวแปรร่วม (CBR + Y factors)</div>',unsafe_allow_html=True)
        yc1,yc2,yc3,yc4,yc5=st.columns(5)
        with yc1: ss["lc_cbr"]=st.number_input("CBR ดินเดิม (%)",min_value=0.0,max_value=20.0,value=float(ss["lc_cbr"]),step=0.5,key="lc_cbr_in")
        with yc2: ss["lc_y_row"]=st.number_input("Y1: กว้างเขตทาง (ม.)",min_value=0.0,value=float(ss["lc_y_row"]),step=1.0,key="lc_yr_in")
        with yc3: ss["lc_y_shoulder"]=st.number_input("Y2: ไหล่ทางกว้างสุด 1 ข้าง (ม.)",min_value=0.0,value=float(ss["lc_y_shoulder"]),step=0.25,key="lc_ys_in")
        with yc4: ss["lc_y_terrain"]=st.selectbox("Y3/Y4/Y6: ภูมิประเทศ",TERRAIN_KEYS,index=TERRAIN_KEYS.index(ss["lc_y_terrain"]),key="lc_yt_in")
        with yc5: ss["lc_y_bridge"]=st.number_input("Y5: สะพาน (ม./กม.)",min_value=0.0,value=float(ss["lc_y_bridge"]),step=1.0,key="lc_yb_in")
        terrain_code=TERRAIN_MAP[ss["lc_y_terrain"]]

        with st.expander("ดูค่า Y factors"):
            Y1=lookup_range(ss["lc_y_row"],Y1_BREAKS); Y2=lookup_range(ss["lc_y_shoulder"],Y2_BREAKS)
            Y3=Y3_MAP[terrain_code]; Y4=Y4_MAP[terrain_code]
            Y5=lookup_range(ss["lc_y_bridge"],Y5_BREAKS); Y6=Y6_MAP[terrain_code]
            st.dataframe(pd.DataFrame({"Factor":["Y1","Y2","Y3","Y4","Y5","Y6"],
                "คำอธิบาย":["เขตทาง","ไหล่ทาง","จราจรสงเคราะห์","ท่อระบายน้ำ","สะพาน","ระบายน้ำ"],
                "ค่า":[Y1,Y2,Y3,Y4,Y5,Y6]}),hide_index=True)

        st.divider()
        col_a,col_b=st.columns(2)

        with col_a:
            st.markdown('<div class="sh">🔵 Section A: ผิวแอสฟัลท์ (Ka)</div>',unsafe_allow_html=True)
            ss["lc_ac_x1_key"]=st.selectbox("X1: ลักษณะผิวทาง",list(X1_MAP.keys()),
                index=list(X1_MAP.keys()).index(ss["lc_ac_x1_key"]),key="lc_ax1")
            ss["lc_ac_x3_key"]=st.selectbox("X3: AADT",list(X3_OPTIONS.keys()),
                index=list(X3_OPTIONS.keys()).index(ss["lc_ac_x3_key"]),key="lc_ax3")
            ss["lc_ac_x4_age"]=st.number_input("X4: อายุปัจจุบัน (ปี)",min_value=0,value=int(ss["lc_ac_x4_age"]),step=1,key="lc_ax4")
            ss["lc_ac_x5_width"]=st.number_input("X5: กว้างผิวทาง (ม.)",min_value=4.0,value=float(ss["lc_ac_x5_width"]),step=0.5,key="lc_ax5")
            ss["lc_ac_x6_terrain"]=st.selectbox("X6: ภูมิประเทศ (AC)",TERRAIN_KEYS,
                index=TERRAIN_KEYS.index(ss["lc_ac_x6_terrain"]),key="lc_ax6")
            st.divider()
            ac1,ac2=st.columns(2)
            with ac1: ss["lc_ac_Na"]=st.number_input("N มาตรฐาน AC (บาท/กม./ปี)",min_value=0.0,value=float(ss["lc_ac_Na"]),step=500.0,key="lc_aNa")
            with ac2: ss["lc_ac_Km"]=st.number_input("Km วัสดุ AC",min_value=0.0,value=float(ss["lc_ac_Km"]),step=0.01,format="%.3f",key="lc_aKm")

        with col_b:
            st.markdown('<div class="sh-org">🟠 Section B: ผิวคอนกรีต (Kc)</div>',unsafe_allow_html=True)
            ss["lc_cc_z1_idx"]=st.selectbox("Z1: ดัชนีสภาพผิวทาง",list(Z1_MAP.keys()),
                index=list(Z1_MAP.keys()).index(ss["lc_cc_z1_idx"]) if ss["lc_cc_z1_idx"] in Z1_MAP else 0,key="lc_cz1")
            ss["lc_cc_z3_key"]=st.selectbox("Z3: AADT (Concrete)",list(Z3_OPTIONS.keys()),
                index=list(Z3_OPTIONS.keys()).index(ss["lc_cc_z3_key"]) if ss["lc_cc_z3_key"] in Z3_OPTIONS else 0,key="lc_cz3")
            ss["lc_cc_z4_width"]=st.number_input("Z4: กว้างผิวทาง (ม.) (Concrete)",min_value=4.0,value=float(ss["lc_cc_z4_width"]),step=0.5,key="lc_cz4")
            st.divider()
            cc1,cc2=st.columns(2)
            with cc1: ss["lc_cc_Nc"]=st.number_input("N มาตรฐาน Concrete (บาท/กม./ปี)",min_value=0.0,value=float(ss["lc_cc_Nc"]),step=500.0,key="lc_cNc")
            with cc2: ss["lc_cc_Km"]=st.number_input("Km วัสดุ Concrete",min_value=0.0,value=float(ss["lc_cc_Km"]),step=0.01,format="%.3f",key="lc_cKm")

        st.divider()
        if st.button("🧮 คำนวณ Ka และ Kc",type="primary",key="lc_calc_k",use_container_width=True):
            n_yrs=ss.get('sb_design_life',20)
            x1=X1_MAP[ss["lc_ac_x1_key"]]
            x3=X3_OPTIONS[ss["lc_ac_x3_key"]]
            x6_code=TERRAIN_MAP[ss["lc_ac_x6_terrain"]]
            ka_avg,ka_df,ka_fixed=calc_Ka_average(x1,ss["lc_cbr"],x3,
                ss["lc_ac_x4_age"],ss["lc_ac_x5_width"],x6_code,
                ss["lc_y_row"],ss["lc_y_shoulder"],terrain_code,ss["lc_y_bridge"],n_yrs)
            z3=Z3_OPTIONS[ss["lc_cc_z3_key"]]
            kc_val,kc_fac=calc_Kc(ss["lc_cc_z1_idx"],ss["lc_cbr"],z3,
                ss["lc_cc_z4_width"],ss["lc_y_row"],ss["lc_y_shoulder"],
                terrain_code,ss["lc_y_bridge"])

            km=road_length; area_lc=total_width*1000
            r_ac_km=ka_avg*ss["lc_ac_Km"]*ss["lc_ac_Na"]*km
            r_cc_km=kc_val*ss["lc_cc_Km"]*ss["lc_cc_Nc"]*km
            r_ac_sqm=r_ac_km/area_lc if area_lc>0 else 0
            r_cc_sqm=r_cc_km/area_lc if area_lc>0 else 0

            ss.update({"lc_ka_avg":ka_avg,"lc_kc_val":kc_val,
                "lc_routine_ac_sqm":r_ac_sqm,"lc_routine_cc_sqm":r_cc_sqm,
                "lc_routine_ac_km":r_ac_km,"lc_routine_cc_km":r_cc_km,
                "_lc_ka_df":ka_df,"_lc_ka_fixed":ka_fixed,"_lc_kc_fac":kc_fac,
                "lc_tab_routine_done":True,"lc_dirty_routine":False})
            st.success(f"✅ Ka เฉลี่ย = {ka_avg:.4f} | Kc = {kc_val:.4f}")
            st.rerun()

        if ss.get("lc_ka_avg"):
            st.markdown("---")
            mc1,mc2,mc3,mc4,mc5,mc6=st.columns(6)
            for col,c_cls,lbl,val,sub in [
                (mc1,"c-blue","Ka เฉลี่ย",f"{ss['lc_ka_avg']:.4f}","ค่าสัมประสิทธิ์ AC"),
                (mc2,"c-orange","Kc",f"{ss['lc_kc_val']:.4f}","ค่าสัมประสิทธิ์ Concrete"),
                (mc3,"c-teal","บำรุง AC",f"{ss['lc_routine_ac_sqm']:.4f}","บาท/ตร.ม./ปี"),
                (mc4,"c-teal","บำรุง AC",f"{ss['lc_routine_ac_km']:,.2f}","บาท/กม./ปี"),
                (mc5,"c-purple","บำรุง Concrete",f"{ss['lc_routine_cc_sqm']:.4f}","บาท/ตร.ม./ปี"),
                (mc6,"c-purple","บำรุง Concrete",f"{ss['lc_routine_cc_km']:,.2f}","บาท/กม./ปี"),
            ]:
                with col:
                    st.markdown(f'<div class="mcc {c_cls}"><div class="lbl">{lbl}</div><div class="val">{val}</div><div class="sub">{sub}</div></div>',unsafe_allow_html=True)
            cd1,cd2=st.columns(2)
            with cd1:
                st.markdown("**Ka รายปี**"); st.dataframe(ss["_lc_ka_df"],hide_index=True,use_container_width=True,height=280)
            with cd2:
                st.markdown("**Kc Factors**"); kf=ss.get("_lc_kc_fac",{})
                st.dataframe(pd.DataFrame({"Factor":list(kf.keys()),"ค่า":[round(v,4) for v in kf.values()]}),hide_index=True,use_container_width=True)

    # ── Sub: กำหนดทางเลือก ───────────────────────────────────────────────────
    with sub_alts:
        st.markdown('<div class="sh-org">🏗️ กำหนดทางเลือกผิวทาง</div>',unsafe_allow_html=True)

        # ══════════════════════════════════════════════════════════════
        # SYNC PANEL — Auto-fill + Lock + Warning
        # ══════════════════════════════════════════════════════════════
        cs_ar     = ss.get('cs_all_results', {})
        _pts      = ['AC','JPCP','JRCP','CRCP']
        _cur_costs = {pt: cs_ar.get(pt, {}).get('cost_sqm', 0.0) for pt in _pts}
        _has_cs   = any(v > 0 for v in _cur_costs.values())

        # ตรวจว่าราคา Tab A เปลี่ยนจาก snapshot ที่ sync ไว้หรือไม่
        _synced   = ss.get('lc_synced_costs', {})
        _changed  = _has_cs and any(
            abs(_cur_costs.get(pt, 0) - _synced.get(pt, -1)) > 0.01
            for pt in _pts if _cur_costs.get(pt, 0) > 0
        )
        _is_manual = ss.get('lc_manual_override', False)

        # ── Auto-fill ครั้งแรก (ถ้ายังไม่เคย sync และไม่ได้ manual) ──────────
        if _has_cs and not _synced and not _is_manual:
            ss['lc_synced_costs'] = dict(_cur_costs)
            ss['lc_sync_ts']      = datetime.now().strftime('%H:%M:%S')
            for pt in _pts:
                ss[f'lc_cost_{pt}'] = float(_cur_costs.get(pt, 0))

        # ── Warning: ราคา Tab A เปลี่ยนหลังจาก sync ────────────────────────
        if _changed and not _is_manual:
            _diff_lines = []
            for pt in _pts:
                old = _synced.get(pt, 0); new = _cur_costs.get(pt, 0)
                if abs(new - old) > 0.01 and new > 0:
                    _diff_lines.append(f"<b>{pt}:</b> {old:,.2f} → <b>{new:,.2f}</b> บ./ตร.ม.")
            st.markdown(f"""
            <div class="sync-warn-banner">
              ⚠️ ราคาโครงสร้างชั้นทางเปลี่ยนแปลงแล้ว — กด <b>🔄 Sync ราคาใหม่</b> เพื่ออัพเดท หรือ ไม่ sync เพื่อคง Lock ไว้<br>
              {"<br>".join(_diff_lines)}
            </div>""", unsafe_allow_html=True)

        # ── Sync Panel UI ─────────────────────────────────────────────────────
        if _has_cs:
            _pills = "".join(
                f'<span class="{"sync-pill" if v > 0 else "sync-pill-zero"}">'
                f'{pt}: {"—" if v == 0 else f"{v:,.2f}"} บ./ตร.ม.</span>'
                for pt, v in _cur_costs.items()
            )
            _lock_label = '<span class="locked-badge">🔒 ล็อคราคา</span>' if not _is_manual else '<span class="manual-badge">✏️ แก้ไขเอง</span>'
            _ts = ss.get('lc_sync_ts', '')
            st.markdown(f"""
            <div class="sync-panel">
              <div class="sync-title">🔗 ราคาจาก Tab โครงสร้างชั้นทาง &nbsp; {_lock_label}</div>
              {_pills}
              <div class="sync-ts">{"Sync ล่าสุด: " + _ts if _ts else "ยังไม่ได้ Sync"}</div>
            </div>""", unsafe_allow_html=True)

            _sc1, _sc2, _sc3 = st.columns([2, 2, 3])
            with _sc1:
                if st.button("🔄 Sync ราคาใหม่", key="lc_sync_btn", use_container_width=True):
                    ss['lc_synced_costs'] = dict(_cur_costs)
                    ss['lc_sync_ts']      = datetime.now().strftime('%H:%M:%S')
                    ss['lc_manual_override'] = False
                    for pt in _pts:
                        ss[f'lc_cost_{pt}'] = float(_cur_costs.get(pt, 0))
                    # รีเซ็ต cost_ver เพื่อบังคับ widget อ่านค่าใหม่
                    ss['_lc_cost_ver'] = ss.get('_lc_cost_ver', 0) + 1
                    st.success("✅ Sync สำเร็จ"); st.rerun()
            with _sc2:
                if not _is_manual:
                    if st.button("✏️ แก้ไขเอง (ปลด Lock)", key="lc_unlock_btn", use_container_width=True):
                        ss['lc_manual_override'] = True
                        ss['_lc_cost_ver'] = ss.get('_lc_cost_ver', 0) + 1
                        st.rerun()
                else:
                    if st.button("🔒 กลับสู่ Lock", key="lc_relock_btn", use_container_width=True):
                        ss['lc_manual_override'] = False
                        # restore จาก synced
                        for pt in _pts:
                            ss[f'lc_cost_{pt}'] = float(ss['lc_synced_costs'].get(pt, 0))
                        ss['_lc_cost_ver'] = ss.get('_lc_cost_ver', 0) + 1
                        st.rerun()
        else:
            st.markdown('<div class="warn-band">⚠️ ยังไม่มีราคาจาก Tab โครงสร้างชั้นทาง — กรอกราคาก่อสร้างด้านล่างโดยตรง</div>', unsafe_allow_html=True)

        # ── ราคาก่อสร้าง input ──────────────────────────────────────────────
        st.markdown("**💰 ราคาก่อสร้าง (บาท/ตร.ม.)**")
        _cv  = ss.get("_lc_cost_ver", 0)
        _disabled_input = _has_cs and not _is_manual

        if _disabled_input:
            # แสดงเป็น read-only metric เมื่อ Locked
            _fc = st.columns(4)
            _cc_colors = ['c-blue','c-orange','c-purple','c-teal']
            for i, pt in enumerate(_pts):
                _v = float(ss.get(f'lc_cost_{pt}', 0))
                with _fc[i]:
                    st.markdown(
                        f'<div class="mcc {_cc_colors[i]}" style="padding:10px 14px">'
                        f'<div class="lbl">💰 {pt}</div>'
                        f'<div class="val" style="font-size:20px">{"—" if _v == 0 else f"{_v:,.2f}"}</div>'
                        f'<div class="sub">บาท/ตร.ม. 🔒</div></div>',
                        unsafe_allow_html=True)
        else:
            # กรอกได้ (manual mode หรือไม่มีข้อมูล Tab A)
            ca, cj, cjr, cc = st.columns(4)
            for col_, pt_, lbl_ in [(ca,'AC','AC'),(cj,'JPCP','JPCP'),(cjr,'JRCP','JRCP'),(cc,'CRCP','CRCP')]:
                with col_:
                    ss[f"lc_cost_{pt_}"] = st.number_input(
                        f"{lbl_}",
                        min_value=0.0, value=float(ss.get(f"lc_cost_{pt_}", 0)),
                        step=10.0, format="%.2f", key=f"lc_c_{pt_}_{_cv}")
            if _is_manual:
                st.caption("✏️ โหมดแก้ไขเอง — กด 🔒 กลับสู่ Lock เมื่อพร้อม")

        costs={"AC":ss.get("lc_cost_AC",0),"JPCP":ss.get("lc_cost_JPCP",0),
               "JRCP":ss.get("lc_cost_JRCP",0),"CRCP":ss.get("lc_cost_CRCP",0)}
        r_ac=ss.get("lc_routine_ac_sqm") or 0.0
        r_cc=ss.get("lc_routine_cc_sqm") or 0.0
        area_lc=total_width*1000.0
        # รวบราคาที่ lock/sync ไว้ (ใช้ทั้ง gen_alts และ info-band)
        costs = {pt: float(ss.get(f'lc_cost_{pt}', 0)) for pt in _pts}

        st.markdown(f'<div class="info-band">📐 พื้นที่คำนวณ = <b>{area_lc:,.0f} ตร.ม./กม.</b> (กว้าง {total_width:.2f} ม. × 1,000 ม.) — ผลลัพธ์ NPV/EAC เป็น <b>บาท/กม.</b></div>',unsafe_allow_html=True)

        # ── พารามิเตอร์ LCCA ──
        st.markdown("---")
        st.markdown('<div class="sh">⚙️ พารามิเตอร์การวิเคราะห์</div>',unsafe_allow_html=True)
        pc1,pc2,pc3=st.columns(3)
        with pc1: ss["lc_n"]=st.number_input("ระยะเวลาวิเคราะห์ (ปี)",min_value=5,max_value=50,value=int(ss["lc_n"]),step=1,key="lc_n_in")
        with pc2: ss["lc_dr"]=st.number_input("อัตราคิดลด (%/ปี)",min_value=1.0,max_value=20.0,value=float(ss["lc_dr"])*100,step=0.5,key="lc_dr_in")/100.0
        with pc3: ss["lc_salvage"]=st.checkbox("รวมมูลค่าซาก",value=ss["lc_salvage"],key="lc_sv_in")
        n_lc=ss["lc_n"]; dr_lc=ss["lc_dr"]

        # ── สร้าง / Reset Alternatives ──
        st.markdown("---")
        if st.button("🔄 สร้าง/รีเซ็ต Alternatives จากราคาข้างบน",key="lc_gen_alts"):
            if not ss.get("lc_tab_routine_done"):
                st.error("กรุณาคำนวณ Routine Cost ก่อน (Sub-tab 🔧)")
            else:
                alts_new=[]
                for nm,pt,cost,mc in [
                    ("ผิวทางยืดหยุ่น (AC)","Flexible",costs["AC"],r_ac),
                    ("ผิวทางคอนกรีต JPCP","JPCP",costs["JPCP"],r_cc),
                    ("ผิวทางคอนกรีต JRCP","JRCP",costs["JRCP"],r_cc),
                    ("ผิวทางคอนกรีต CRCP","CRCP",costs["CRCP"],r_cc),
                ]:
                    if cost<=0: continue
                    rehab_yr=max(10,n_lc//2) if "Flex" in pt or "AC" in pt else max(15,int(n_lc*0.75))
                    if "Flex" in pt or "AC" in pt:
                        mlist=[MaintAct("บำรุงรักษาประจำปี",mc,1,1),MaintAct("ฉาบผิว (Seal Coating)",mc*0.8,3,3)]
                        rlist=[RehabAct("Overlay",390,rehab_yr)]; sv=20.0
                    else:
                        mlist=[MaintAct("บำรุงรักษาประจำปี",mc,1,1),MaintAct("Joint Maintenance",mc*0.5,5,5)]
                        rlist=[]; sv=30.0
                    alts_new.append(PavAlt(nm,pt,cost,area_lc,mlist,rlist,sv))
                ss["lc_alternatives"]=alts_new; ss["lc_tab_alts_done"]=True
                st.success(f"✅ สร้าง {len(alts_new)} ทางเลือก"); st.rerun()

        alts=ss.get("lc_alternatives") or []
        if alts:
            st.markdown("**✏️ แก้ไขรายละเอียดทางเลือก:**")
            for ai,alt in enumerate(alts):
                with st.expander(f"✏️ {alt.name} | {alt.construction_cost:,.0f} บาท/ตร.ม.",expanded=False):
                    e1,e2=st.columns(2)
                    with e1:
                        alts[ai].construction_cost=st.number_input("ต้นทุนก่อสร้าง (บาท/ตร.ม.)",min_value=0.0,value=float(alt.construction_cost),step=10.0,key=f"lc_ec_{ai}")
                        alts[ai].salvage_pct=st.number_input("มูลค่าซาก (%)",min_value=0.0,max_value=100.0,value=float(alt.salvage_pct),step=1.0,key=f"lc_es_{ai}")
                        alts[ai].enabled=st.checkbox("แสดงผลในรายงาน",value=alt.enabled,key=f"lc_ee_{ai}")
                    with e2:
                        st.markdown("**บำรุงรักษา:**")
                        for mi,m in enumerate(alt.maintenance):
                            mc1,mc2=st.columns([2,1])
                            with mc1: alts[ai].maintenance[mi].unit_cost=st.number_input(f"{m.name} (บาท/ตร.ม./ปี)",min_value=0.0,value=float(m.unit_cost),step=1.0,key=f"lc_mc_{ai}_{mi}")
                            with mc2: alts[ai].maintenance[mi].frequency=st.number_input("ความถี่ (ปี)",min_value=0,value=int(m.frequency),step=1,key=f"lc_mf_{ai}_{mi}")
                        st.markdown("**ฟื้นฟูสภาพ:**")
                        for ri2,r in enumerate(alt.rehab):
                            rc1,rc2=st.columns([2,1])
                            with rc1: alts[ai].rehab[ri2].unit_cost=st.number_input(f"{r.name} (บาท/ตร.ม.)",min_value=0.0,value=float(r.unit_cost),step=10.0,key=f"lc_rc_{ai}_{ri2}")
                            with rc2: alts[ai].rehab[ri2].year=st.number_input("ดำเนินการปีที่",min_value=1,max_value=n_lc,value=int(r.year),step=1,key=f"lc_ry_{ai}_{ri2}")
            ss["lc_alternatives"]=alts
            st.divider()
            if st.button("🚀 คำนวณ LCCA",type="primary",key="lc_run",use_container_width=True):
                warns=[]
                if not ss.get("lc_tab_routine_done"): warns.append("ยังไม่ได้คำนวณ Routine Cost")
                if warns:
                    for w in warns: st.error(w)
                else:
                    with st.spinner("กำลังคำนวณ LCCA..."):
                        sdf,cfd=analyze_lcca(alts,n_lc,dr_lc,ss["lc_salvage"])
                        ss["_lc_sum"]=sdf; ss["_lc_cf"]=cfd; ss["lc_tab_result_done"]=True
                    st.success("✅ คำนวณ LCCA สำเร็จ — ดูผลที่ Sub-tab 📊 ผลการวิเคราะห์")

    # ── Sub: ผลการวิเคราะห์ ──────────────────────────────────────────────────
    with sub_result:
        sdf=ss.get("_lc_sum"); cfd=ss.get("_lc_cf",{})
        n_lc=ss["lc_n"]; dr_lc=ss["lc_dr"]
        if sdf is None or len(sdf)==0:
            st.info("กรุณากำหนดทางเลือกและคำนวณ LCCA ใน Sub-tab 🏗️ ก่อน")
        else:
            st.markdown('<div class="sh-grn">🏆 สรุปผล LCCA</div>',unsafe_allow_html=True)
            best=sdf.iloc[0]
            st.markdown(f"""<div class="best-row">
              🥇 <b>ทางเลือกที่ดีที่สุด: {best['ทางเลือก']}</b> ({best['ประเภทผิวทาง']})<br>
              NPV = <b>{best['NPV (ล้านบาท/กม.)']:,.4f} ล้านบาท/กม.</b> &nbsp;|&nbsp;
              EAC = <b>{best['EAC (ล้านบาท/กม./ปี)']:,.4f} ล้านบาท/กม./ปี</b> &nbsp;|&nbsp;
              EAC = <b>{best['EAC (บาท/ตร.ม./ปี)']:,.2f} บาท/ตร.ม./ปี</b>
            </div>""",unsafe_allow_html=True)

            card_c=["c-green","c-blue","c-orange","c-purple"]
            cols_r=st.columns(len(sdf))
            for i,(_,row) in enumerate(sdf.iterrows()):
                with cols_r[i]:
                    badge="🥇" if i==0 else ("🥈" if i==1 else "🥉")
                    _npv_val=f"{row['NPV (ล้านบาท/กม.)']:,.4f}"
                    st.markdown(f'<div class="mcc {card_c[i%4]}"><div class="lbl">{badge} อันดับ {int(row["อันดับ"])} — {row["ทางเลือก"]}</div><div class="val">{_npv_val}</div><div class="sub">ล้านบ./กม. (NPV)</div></div>',unsafe_allow_html=True)

            show_cols=["อันดับ","ทางเลือก","ประเภทผิวทาง",
                       "ต้นทุนก่อสร้าง (บาท/ตร.ม.)","ต้นทุนก่อสร้าง (ล้านบาท/กม.)",
                       "NPV (ล้านบาท/กม.)","EAC (ล้านบาท/กม./ปี)","EAC (บาท/ตร.ม./ปี)"]
            fmt_cols={"ต้นทุนก่อสร้าง (บาท/ตร.ม.)":"{:,.2f}","ต้นทุนก่อสร้าง (ล้านบาท/กม.)":"{:,.4f}",
                      "NPV (ล้านบาท/กม.)":"{:,.4f}","EAC (ล้านบาท/กม./ปี)":"{:,.4f}","EAC (บาท/ตร.ม./ปี)":"{:,.2f}"}
            st.dataframe(sdf[show_cols].style.format(fmt_cols),hide_index=True,use_container_width=True)

            if PLOTLY_OK:
                st.markdown('<div class="sh">📊 NPV แยกประเภทต้นทุน</div>',unsafe_allow_html=True)
                fig_bar=go.Figure()
                color_map={"ก่อสร้าง":"#1565C0","บำรุงรักษา":"#F57C00","ฟื้นฟูสภาพ":"#C62828","มูลค่าซาก":"#2E7D32"}
                for pt,col_k in [("ก่อสร้าง","PW_ก่อสร้าง"),("บำรุงรักษา","PW_บำรุงรักษา"),("ฟื้นฟูสภาพ","PW_ฟื้นฟูสภาพ"),("มูลค่าซาก","PW_มูลค่าซาก")]:
                    fig_bar.add_trace(go.Bar(name=pt,x=sdf["ทางเลือก"],y=sdf[col_k],marker_color=color_map[pt]))
                fig_bar.update_layout(barmode="relative",height=420,title="NPV แยกตามประเภทต้นทุน",
                    yaxis_title="บาท",paper_bgcolor="rgba(0,0,0,0)",plot_bgcolor="rgba(0,0,0,0)")
                st.plotly_chart(fig_bar,use_container_width=True)

                st.markdown('<div class="sh">📈 Cumulative Cost Timeline</div>',unsafe_allow_html=True)
                cum_df=build_cumulative(cfd,n_lc,dr_lc)
                fig_cum=px.line(cum_df,x="ปี",y="Cumulative NPV (บาท)",color="ทางเลือก",
                    markers=True,height=430,title=f"ต้นทุนสะสม ตลอด {n_lc} ปี")
                fig_cum.update_layout(paper_bgcolor="rgba(0,0,0,0)",plot_bgcolor="rgba(248,249,250,1)")
                st.plotly_chart(fig_cum,use_container_width=True)

                st.markdown('<div class="sh">⚖️ Breakeven Year Analysis</div>',unsafe_allow_html=True)
                be_df=calc_breakeven(cfd,n_lc,dr_lc)
                if len(be_df)>0:
                    st.dataframe(be_df,hide_index=True,use_container_width=True)
                    for _,r in be_df.iterrows():
                        band="info-band" if isinstance(r["Breakeven Year"],int) else "warn-band"
                        st.markdown(f'<div class="{band}">📌 {r["หมายเหตุ"]}</div>',unsafe_allow_html=True)

                st.markdown('<div class="sh">📉 Sensitivity Analysis — อัตราคิดลด</div>',unsafe_allow_html=True)
                alts_lc=ss.get("lc_alternatives") or []
                sens_rows=[]
                for dr_i in np.linspace(max(dr_lc-0.03,0.01),dr_lc+0.03,7):
                    for alt in [a for a in alts_lc if a.enabled]:
                        cf_i=build_cashflow(alt,n_lc,dr_i,ss["lc_salvage"])
                        sens_rows.append({"อัตราคิดลด (%)":round(dr_i*100,1),"ทางเลือก":alt.name,"NPV (บาท)":cf_i["มูลค่าปัจจุบัน"].sum()})
                fig_sens=px.line(pd.DataFrame(sens_rows),x="อัตราคิดลด (%)",y="NPV (บาท)",color="ทางเลือก",
                    markers=True,height=420,title="Sensitivity — ผลของอัตราคิดลดต่อ NPV")
                fig_sens.update_layout(paper_bgcolor="rgba(0,0,0,0)",plot_bgcolor="rgba(248,249,250,1)")
                st.plotly_chart(fig_sens,use_container_width=True)

            st.markdown('<div class="sh">💰 กระแสเงินสดรายทางเลือก</div>',unsafe_allow_html=True)
            sel=st.selectbox("เลือกทางเลือก",list(cfd.keys()),key="lc_cfsel")
            if sel in cfd:
                cf_s=cfd[sel].copy()
                cf_s["ต้นทุน/หน่วย"]=cf_s["ต้นทุน/หน่วย"].map(lambda x:f"{x:,.2f}")
                cf_s["ต้นทุนตามปี"]=cf_s["ต้นทุนตามปี"].map(lambda x:f"{x:,.0f}")
                cf_s["PW_factor"]=cf_s["PW_factor"].map(lambda x:f"{x:.4f}")
                cf_s["มูลค่าปัจจุบัน"]=cf_s["มูลค่าปัจจุบัน"].map(lambda x:f"{x:,.0f}")
                st.dataframe(cf_s,hide_index=True,use_container_width=True,height=400)

    # ── Sub: Word Report รวม (ปุ่มเดียว) ────────────────────────────────────
    with sub_word:
        st.markdown('<div class="sh">📄 Word Report รวม — กดปุ่มเดียว ได้ไฟล์ครบ</div>', unsafe_allow_html=True)

        if not DOCX_OK:
            st.error("❌ ติดตั้ง python-docx ก่อน: pip install python-docx")
        else:
            # ── ตรวจสอบความพร้อม ──────────────────────────────────────────
            cs_ar_w  = ss.get('cs_all_results', {})
            sdf_w    = ss.get("_lc_sum")
            cfd_w    = ss.get("_lc_cf", {})
            has_cs_w = any(cs_ar_w.get(pt, {}).get('cost_sqm', 0) > 0 for pt in ['AC','JPCP','JRCP','CRCP'])
            has_lc_w = sdf_w is not None and len(sdf_w) > 0

            # Progress checklist
            checks = [
                ("🏗️ ราคาโครงสร้างชั้นทาง", has_cs_w, "ไปที่ Tab โครงสร้างชั้นทาง → กำหนดหน้าตัด"),
                ("🔧 Routine Cost (Ka/Kc)", ss.get("lc_tab_routine_done", False), "ไปที่ LCCA → Routine Cost → กด คำนวณ"),
                ("📊 ผล LCCA", has_lc_w, "ไปที่ LCCA → กำหนดทางเลือก → กด คำนวณ LCCA"),
            ]
            all_ready = all(ok for _, ok, _ in checks)

            for label, ok, hint in checks:
                icon = "✅" if ok else "⭕"
                color = "#1B5E20" if ok else "#9a3412"
                st.markdown(
                    f'<div style="padding:5px 0;font-size:14px;color:{color}">'
                    f'{icon} <b>{label}</b>'
                    + (f'<span style="color:#9ca3af;font-size:12px"> — {hint}</span>' if not ok else '')
                    + '</div>', unsafe_allow_html=True)

            st.divider()

            # ── ตั้งค่าหัวข้อ ──────────────────────────────────────────────
            st.markdown("**⚙️ ตั้งค่ารายงาน**")
            _w1, _w2 = st.columns([2, 3])
            with _w1:
                ss["lc_base_sec"] = st.text_input(
                    "หมายเลขหัวข้อเริ่มต้น (เช่น 4.7)",
                    value=ss.get("lc_base_sec", "4.7"),
                    key="wr_base_sec",
                    help="โปรแกรมจะนับต่อเป็น +1, +2 อัตโนมัติ")
            _base = ss["lc_base_sec"].strip()

            def _next(s, offset):
                parts = s.split(".")
                try: return ".".join(parts[:-1] + [str(int(parts[-1]) + offset)])
                except: return s

            with _w2:
                st.markdown(
                    f'<div class="info-band" style="margin-top:24px">'
                    f'<b>{_base}</b> วัสดุ+ราคาโครงสร้างชั้นทาง &nbsp;→&nbsp; '
                    f'<b>{_next(_base,1)}</b> Routine Cost &nbsp;→&nbsp; '
                    f'<b>{_next(_base,2)}</b> LCCA'
                    f'</div>', unsafe_allow_html=True)

            # ── บทเกริ่นนำ (optional) ──────────────────────────────────────
            with st.expander("📝 บทเกริ่นนำ (optional)", expanded=False):
                _nl = project_info.get('num_lanes', 4)
                _tw = project_info.get('total_width', 0)
                _ln = project_info.get('length', 1.0)
                _auto_intro = (
                    f"รายงานฉบับนี้จัดทำขึ้นเพื่อวิเคราะห์ต้นทุนโครงสร้างชั้นทางและ"
                    f"ต้นทุนตลอดอายุการใช้งาน (LCCA) สำหรับถนน {_nl} ช่องจราจร "
                    f"ความกว้างรวม {_tw:.2f} เมตร ระยะทาง {_ln:.2f} กิโลเมตร "
                    f"ครอบคลุมผิวทางแอสฟัลต์คอนกรีต (AC) และผิวทางคอนกรีตซีเมนต์ (JPCP, JRCP, CRCP)"
                )
                _ik = f"wr_intro_{_nl}_{_tw:.1f}_{_ln:.2f}"
                if _ik not in ss: ss[_ik] = _auto_intro
                intro_txt = st.text_area("บทเกริ่นนำ", height=100, key=_ik)

            st.divider()

            # ── ปุ่มเดียว ──────────────────────────────────────────────────
            if not all_ready:
                st.markdown('<div class="warn-band">⚠️ กรุณาทำขั้นตอนที่ยังไม่ครบ (⭕) ก่อนสร้างรายงาน</div>', unsafe_allow_html=True)
                btn_disabled = True
            else:
                st.markdown('<div class="best-row">✅ ข้อมูลครบทุกส่วน — พร้อมสร้างรายงาน</div>', unsafe_allow_html=True)
                btn_disabled = False

            if st.button(
                "📋 สร้าง Word Report รวม (ไฟล์เดียว)",
                type="primary",
                use_container_width=True,
                disabled=btn_disabled,
                key="wr_gen_btn",
            ):
                with st.spinner("กำลังสร้างรายงาน — โปรดรอสักครู่..."):
                    try:
                        alts_w = ss.get("lc_alternatives") or []
                        buf = generate_word_combined(
                            project_info = project_info,
                            all_details  = cs_ar_w,
                            summary_df   = sdf_w,
                            cf_dict      = cfd_w,
                            n            = ss["lc_n"],
                            dr           = ss["lc_dr"],
                            alts         = alts_w,
                            ss           = ss,
                            base_sec     = _base,
                            intro_text   = ss.get(_ik, ""),
                        )
                        proj = project_name.replace(" ", "_")
                        fname = f"{proj}_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                        st.download_button(
                            "⬇️ ดาวน์โหลด Word Report",
                            data=buf,
                            file_name=fname,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="wr_dl_btn",
                        )
                        st.success(f"✅ สร้างรายงานสำเร็จ — {fname}")
                    except Exception as e:
                        st.error(f"สร้างรายงานไม่สำเร็จ: {e}")
                        st.exception(e)

# ── Footer ────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="footer">
  <b>รศ.ดร.อิทธิพล มีผล</b><br>
  ภาควิชาครุศาสตร์โยธา คณะครุศาสตร์อุตสาหกรรม มจพ.<br>
  <span style="color:#b0bec5">PaveCost + LCCA Integrated v1.0</span>
</div>
""",unsafe_allow_html=True)

if __name__=="__main__":
    pass
